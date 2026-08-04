#!/usr/bin/env python3
"""pdf_cli.py — PDF 处理统一 CLI (2026-05-28 · 2026-07-31 家族三合一)

9 + 1 子命令:
  read                    pdfplumber 提文本 / --list 显示 outline + pageinfo
  extract image           pdfimages -all
  extract text            pdfplumber 每页 txt (或 --single 合一)
  extract table           pdfplumber 每页 CSV
  split  by-bookmark      pypdf outline 切分
  split  by-page-range    pypdf 范围切分 (--ranges "1-10,11-20")
  merge                   pypdf PdfWriter 拼接
  decrypt                 qpdf 直解 (--password) / 无密走 cc-home pdf-decrypt skill
  convert to-docx         PDF → 可编辑 Word (结构提取: 段落重组 + 真表格)
  pipeline run            <glob> --steps <names> [--parallel ...]

底座依赖:
  pdfplumber 0.11.9 / pypdf 6.12.2 / python-docx 1.2.0 (仅 convert to-docx 用,惰性加载)
  /opt/homebrew/bin/{pdfimages,pdftotext,pdfinfo,qpdf,ocrmypdf,tesseract}
  依赖装在 /opt/homebrew/bin/python3,不在 ~/Dev/.venv —— 调用方必须用绝对路径解释器。

设计同源 `docx_cli.py` 双层 dispatch,但所有子命令本仓 native (不转发旧脚本)。
2026-07-31 家族折叠:原 `pdf_pipeline_lib.py`(pipeline 引擎)与 `pdf_to_docx.py`
(to-docx 引擎)整段并入本文件,旧件退役 ~/.Trash/consolidation-20260731/pdf/。
"""

from __future__ import annotations

import argparse
import csv
import glob as _glob
import json
import logging
import os
import re
import shutil
import statistics
import subprocess
import sys
import tempfile
import time
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Optional

# ─── parallel_contract for --workers/--max-workers cohesion ────────────
_LIB = Path.home() / "Dev" / "tools" / "dev" / "lib"
if _LIB.is_dir() and str(_LIB) not in sys.path:   # 别的机器上没有 ~/Dev，不塞死路径
    sys.path.insert(0, str(_LIB))
# 同 docx_cli.py：三条来源 = 总部工作树 / `<根>/lib` 镜像 / 装出来的 hq-devlib 包。
# append 理由同 docx_cli.py（不改本机解析优先级）。
_BUNDLED_LIB = Path(__file__).resolve().parents[2] / "lib"
if str(_BUNDLED_LIB) not in sys.path:
    sys.path.append(str(_BUNDLED_LIB))
try:
    from parallel_contract import add_parallel_args  # type: ignore
except ImportError as e:  # pragma: no cover
    # 缺它是 **fail-closed exit 2 不是降级**（2026-08-04 改，对齐 docx_cli.py）——
    # 旧 except 分支静默只留 --workers、丢 --batch/--phases/--defer/--fanout-evidence，
    # 同一版 CLI 在两台机器上 flag 面不一样是最难查的漂移。
    _seen = lambda p, tag: f"{p}（{tag}{'' if p.is_dir() else '，不存在'}）"  # noqa: E731
    print(f"[pdf_cli.py] FATAL: cannot import parallel_contract: {e}\n"
          f"  找过：{_seen(_LIB, '总部工作树')} · {_seen(_BUNDLED_LIB, '包内镜像')} · 已装的包\n"
          f"  修法：pip install hq-devlib   （= 总部 ~/Dev/tools/dev/lib 的 6 个平铺模块，\n"
          f"        本仓 pyproject 已声明为依赖；正常 `pip install doctools` 会自动带上它）",
          file=sys.stderr)
    sys.exit(2)

# Lazy imports of pdfplumber / pypdf — heavy deps; only some subcommands need
try:
    import pdfplumber  # type: ignore
except ImportError:
    pdfplumber = None  # type: ignore

try:
    import pypdf  # type: ignore
    from pypdf import PdfReader, PdfWriter  # type: ignore
except ImportError:
    pypdf = None  # type: ignore
    PdfReader = PdfWriter = None  # type: ignore


# Canonical binary paths
_PDFINFO = "/opt/homebrew/bin/pdfinfo"
_PDFIMAGES = "/opt/homebrew/bin/pdfimages"
_PDFTOTEXT = "/opt/homebrew/bin/pdftotext"
_QPDF = "/opt/homebrew/bin/qpdf"
_DECRYPT_SKILL = (
    Path.home() / "Dev" / "tools" / "cc-home"
    / "skills" / "pdf-decrypt" / "scripts" / "decrypt.py"
)
_PYTHON3 = "/opt/homebrew/bin/python3"


def _run_pdfimages(pdf_path: Path, dest_prefix: Path, page: int | None = None,
                   binary: str | None = None) -> subprocess.CompletedProcess:
    """pdfimages -all 子进程统一包装(原三处同构 argv 归一;page=N 时只抽第 N 页)。"""
    cmd = [binary or _PDFIMAGES, "-all"]
    if page is not None:
        cmd += ["-f", str(page), "-l", str(page)]
    cmd += [str(pdf_path), str(dest_prefix)]
    return subprocess.run(cmd, capture_output=True, text=True)


# ═══════════════════════════════════════════════════════════════════════
# pipeline 引擎段 (原 pdf_pipeline_lib.py 整段并入, 2026-07-31)
#
# Same design as docx pipeline (`sub/pipeline_lib.py`),swap underlay to
# pdfplumber + pypdf + Poppler CLI (pdfimages / pdftotext / pdfinfo).
#
# Two step kinds
#   1. **pdf-based step** — `fn(pdf, args, out_dir) -> dict`
#      `pdf` 是调用方已开好的 pdfplumber 对象;同一 PDF 多 pass 复用,不重复 parse
#   2. **path-based step** — `fn(pdf_path: Path, args, out_dir) -> dict`
#      自管文件 IO (典型: subprocess 包 pdfimages),无需 pdfplumber 对象
#
# Step registry: `_BUILTIN_STEPS: dict[str, tuple[str, Callable]]`,
# kind ∈ {"pdf", "path"}。加 step = 写 `_<verb>(...)` + 一行注册。
#
# Concurrency: 单 PDF × N steps → run_pipeline_single (一次 open 复用);
# N PDFs → run_pipeline_parallel (ProcessPoolExecutor 跨 PDF)。
# ⚠ _worker/_BUILTIN_STEPS 必须保持模块级顶层 (spawn 按 qualified name pickle)。
#
# Per-step out_dir: `--<step-name>-out-dir` 覆盖,默认 `<pdf-parent>/<verb>/`。
# ═══════════════════════════════════════════════════════════════════════

# ─── page-range parsing ────────────────────────────────────────────────

def parse_page_spec(spec: Optional[str], total_pages: int) -> list[int]:
    """Parse "1-5,8,10-12" → sorted 0-indexed page list.

    spec=None → all pages.
    Validates against total_pages; clamps to [1, total_pages]; raises on invalid syntax.
    """
    if not spec:
        return list(range(total_pages))
    out: set[int] = set()
    for tok in spec.split(","):
        tok = tok.strip()
        if not tok:
            continue
        m = re.match(r"^(\d+)(?:-(\d+))?$", tok)
        if not m:
            raise ValueError(f"invalid page spec token: {tok!r}")
        a = int(m.group(1))
        b = int(m.group(2)) if m.group(2) else a
        if a < 1 or b < 1 or a > b:
            raise ValueError(f"invalid page range: {tok!r}")
        for p in range(a, b + 1):
            if 1 <= p <= total_pages:
                out.add(p - 1)
    return sorted(out)


def parse_ranges_spec(spec: str) -> list[tuple[int, int]]:
    """Parse "1-10,11-20,21-30" → [(1,10),(11,20),(21,30)] (1-indexed inclusive)."""
    out: list[tuple[int, int]] = []
    for tok in spec.split(","):
        tok = tok.strip()
        if not tok:
            continue
        m = re.match(r"^(\d+)-(\d+)$", tok)
        if not m:
            raise ValueError(f"invalid range token: {tok!r} (need A-B)")
        a, b = int(m.group(1)), int(m.group(2))
        if a < 1 or b < a:
            raise ValueError(f"invalid range: {tok!r}")
        out.append((a, b))
    if not out:
        raise ValueError("no ranges parsed")
    return out


# 非法字符→"_" + 空白→"_" 的公共变换(safe_filename 与 caption stem 共底,
# 差异只在字符类范围与截断策略,由各自入口保持)。
_FN_LEGACY_BAD_RE = re.compile(r"[/\\\x00-\x1f]+")
_FN_LEGACY_WS_RE = re.compile(r"\s+")


def _sanitize_core(s: str, bad_re: re.Pattern, ws_re: re.Pattern) -> str:
    s = bad_re.sub("_", s).strip()
    return ws_re.sub("_", s)


def safe_filename(s: str, maxlen: int = 120) -> str:
    """Sanitize bookmark titles for filesystem use."""
    s = _sanitize_core(s, _FN_LEGACY_BAD_RE, _FN_LEGACY_WS_RE)
    return (s[:maxlen] or "untitled")


# ─── caption detection helpers ────────────────────────────────────────
#
# Detect lines like "图 3-7 台州市干旱指数年内分配过程图" / "表 1-1 ...".
# Anchored at start of line to avoid in-body refs like "如图 1-1 所示".
# Separator allowed: "-", "．", "."; gap between code and title allowed
# to be one or more whitespace (incl. full-width).
_IMG_CAPTION_RE = re.compile(
    r"^\s*图\s*(\d+)[\-．.](\d+)\s+(.{2,60}?)\s*$"
)
_TBL_CAPTION_RE = re.compile(
    r"^\s*表\s*(\d+)[\-．.](\d+)\s+(.{2,60}?)\s*$"
)

# Filesystem-illegal chars in caption titles → replace with "_".
_FN_BAD_CHARS_RE = re.compile(r'[/\\:*?"<>|\x00-\x1f]+')
# Whitespace (including full-width space U+3000) → "_".
_FN_WS_RE = re.compile(r"[\s　]+")


def _sanitize_caption_filename(s: str, maxlen: int = 80) -> str:
    """Sanitize a caption string into a filesystem-safe stem."""
    s = _sanitize_core(s, _FN_BAD_CHARS_RE, _FN_WS_RE)
    s = s.strip("._")
    if len(s) > maxlen:
        s = s[:maxlen].rstrip("._")
    return s or "untitled"


def _build_caption_stem(m: re.Match) -> str:
    """From a caption regex match → stem like '图3-7_台州市干旱指数年内分配过程图'."""
    kind_char = m.string.lstrip()[0]  # "图" or "表"
    major, minor, title = m.group(1), m.group(2), m.group(3)
    title_clean = _sanitize_caption_filename(title)
    return f"{kind_char}{major}-{minor}_{title_clean}"


def _find_captions_on_page(page, regex: re.Pattern) -> list[tuple[str, float]]:
    """Return [(stem, y_top), ...] for caption-matching lines on the page,
    sorted by y_top (reading order).

    Uses pdfplumber.Page.extract_text_lines() when available (gives both
    line text and top y-coord). Falls back to extract_text() split-by-line
    with synthetic y = line index when not.
    """
    out: list[tuple[str, float]] = []
    lines_data: list[tuple[str, float]] = []
    try:
        for ln in page.extract_text_lines() or []:
            text = ln.get("text") or ""
            top = float(ln.get("top", 0.0))
            lines_data.append((text, top))
    except Exception:
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        for idx, line in enumerate(txt.split("\n")):
            lines_data.append((line, float(idx)))

    for text, top in lines_data:
        m = regex.match(text)
        if not m:
            continue
        stem = _build_caption_stem(m)
        out.append((stem, top))
    out.sort(key=lambda x: x[1])
    return out


def _unique_path(out_dir: Path, stem: str, ext: str) -> Path:
    """Return out_dir/<stem><ext> avoiding collision via -2/-3 suffix.

    `ext` includes the leading dot (".jpg", ".csv", etc).
    """
    base = out_dir / f"{stem}{ext}"
    if not base.exists():
        return base
    i = 2
    while True:
        cand = out_dir / f"{stem}-{i}{ext}"
        if not cand.exists():
            return cand
        i += 1


# Threshold (bytes) below which a pdfimages output is treated as noise
# (soft-mask, single-color filler). Empirically 2-2.5KB files are noise,
# real figures start at ~25KB. 3KB cleanly separates the two.
MIN_IMAGE_BYTES = 3072


# ─── built-in step implementations ────────────────────────────────────

def _text_extract(pdf, args, out_dir: Path) -> dict:
    """pdf-based: pdfplumber extract_text() per page to .txt files.

    args.text_extract_pages — page spec (e.g. "1-5"), None=all
    args.text_extract_single — bool, concat all into full.txt
    """
    if pdf is None:
        return {"step": "text-extract",
                "error": "pdfplumber not available or pdf is None"}
    t0 = time.perf_counter()
    out_dir.mkdir(parents=True, exist_ok=True)
    spec = (getattr(args, "text_extract_pages", None)
            or getattr(args, "pages", None))
    single = bool(getattr(args, "text_extract_single", False)
                  or getattr(args, "single", False))
    total = len(pdf.pages)
    page_idx = parse_page_spec(spec, total)
    pad = max(3, len(str(total)))
    pages_written: list[str] = []
    chunks: list[str] = []
    for i in page_idx:
        try:
            txt = pdf.pages[i].extract_text() or ""
        except Exception as e:
            txt = f"[ERROR extract_text page {i+1}: {type(e).__name__}: {e}]"
        if single:
            chunks.append(f"\n\n===== page {i+1} =====\n\n{txt}")
        else:
            fn = out_dir / f"page-{str(i+1).zfill(pad)}.txt"
            fn.write_text(txt, encoding="utf-8")
            pages_written.append(fn.name)
    if single:
        full = out_dir / "full.txt"
        full.write_text("".join(chunks), encoding="utf-8")
        pages_written = ["full.txt"]
    return {
        "step": "text-extract",
        "pages_requested": len(page_idx),
        "files_written": len(pages_written),
        "out_dir": str(out_dir),
        "elapsed_s": round(time.perf_counter() - t0, 3),
        "mode": "single" if single else "per-page",
    }


def _table_extract(pdf, args, out_dir: Path) -> dict:
    """pdf-based: pdfplumber.extract_tables() per page → CSV files.

    Naming: prefer matched table-caption stems on the page
    (e.g. "表1-1_台州市主要河流特征表.csv"); fall back to
    "page-{NNN}-table-{M}.csv" when caption count < table count.

    Pairing strategy: per page, sort captions by y_top (reading order),
    and pair with tables in the order returned by pdfplumber
    (pdfplumber.extract_tables() already returns tables in page-y order).

    args.table_extract_pages — page spec; None=all
    """
    if pdf is None:
        return {"step": "table-extract",
                "error": "pdfplumber not available or pdf is None"}
    t0 = time.perf_counter()
    out_dir.mkdir(parents=True, exist_ok=True)
    spec = (getattr(args, "table_extract_pages", None)
            or getattr(args, "pages", None))
    total = len(pdf.pages)
    page_idx = parse_page_spec(spec, total)
    pad = max(3, len(str(total)))
    per_page: dict[str, int] = {}
    total_tables = 0
    caption_named = 0
    fallback_named = 0
    files: list[str] = []
    for i in page_idx:
        page = pdf.pages[i]
        try:
            tables = page.extract_tables() or []
        except Exception as e:
            per_page[str(i + 1)] = -1
            files.append(f"[ERROR p{i+1}: {type(e).__name__}: {e}]")
            continue
        per_page[str(i + 1)] = len(tables)
        if not tables:
            continue
        try:
            captions = _find_captions_on_page(page, _TBL_CAPTION_RE)
        except Exception:
            captions = []
        for j, tbl in enumerate(tables):
            if j < len(captions):
                stem = captions[j][0]
                fn = _unique_path(out_dir, stem, ".csv")
                caption_named += 1
            else:
                stem = f"page-{str(i+1).zfill(pad)}-table-{j+1}"
                fn = _unique_path(out_dir, stem, ".csv")
                fallback_named += 1
            with fn.open("w", newline="", encoding="utf-8") as f:
                w = csv.writer(f)
                for row in tbl:
                    w.writerow(["" if c is None else c for c in row])
            files.append(fn.name)
            total_tables += 1
    return {
        "step": "table-extract",
        "pages_scanned": len(page_idx),
        "total_tables": total_tables,
        "caption_named": caption_named,
        "fallback_named": fallback_named,
        "per_page_table_counts": per_page,
        "out_dir": str(out_dir),
        "elapsed_s": round(time.perf_counter() - t0, 3),
    }


def _image_extract(pdf_path: Path, args, out_dir: Path) -> dict:
    """path-based: per-page pdfimages -all + caption-based renaming.

    Algorithm (per page):
      a. open page in pdfplumber, regex-match "图 X-Y title" captions →
         list of (stem, y_top), sorted by y_top.
      b. subprocess pdfimages -f N -l N to a tempdir.
      c. drop files < MIN_IMAGE_BYTES (soft-mask / pure-color noise).
      d. zip captions ↔ remaining files in extraction order;
         caption stem → final name. Extras → fallback "page-NNN-img-MM.<ext>".

    args.image_extract_pages — page spec; None=all
    """
    t0 = time.perf_counter()
    out_dir.mkdir(parents=True, exist_ok=True)

    if pdfplumber is None:
        return {
            "step": "image-extract",
            "error": "pdfplumber not installed; required for caption-named image extraction",
        }

    # Determine page list. Need total pages → open pdfplumber once.
    spec = (getattr(args, "image_extract_pages", None)
            or getattr(args, "pages", None))

    pdfimages_bin = _PDFIMAGES
    rep_files: list[str] = []
    caption_named = 0
    fallback_named = 0
    noise_dropped = 0
    rc_nonzero_pages: list[int] = []
    stderr_first: str = ""

    try:
        pdf_obj = pdfplumber.open(str(pdf_path))
    except Exception as e:
        return {
            "step": "image-extract",
            "error": f"pdfplumber.open failed: {type(e).__name__}: {e}",
        }

    try:
        total = len(pdf_obj.pages)
        page_idx = parse_page_spec(spec, total)
        pad = max(3, len(str(total)))

        for i in page_idx:
            page_num = i + 1  # pdfimages is 1-based
            try:
                page = pdf_obj.pages[i]
                captions = _find_captions_on_page(page, _IMG_CAPTION_RE)
            except Exception:
                captions = []

            with tempfile.TemporaryDirectory(prefix="pdfimg_") as td:
                tdp = Path(td)
                cp = _run_pdfimages(pdf_path, tdp / "img", page=page_num,
                                    binary=pdfimages_bin)
                if cp.returncode != 0:
                    rc_nonzero_pages.append(page_num)
                    if not stderr_first:
                        stderr_first = (cp.stderr or "").strip()[:300]

                # Collect produced files in extraction order
                # (pdfimages numbers img-NNN sequentially in page-y order;
                # main images, masks and smasks interleave by object).
                produced = sorted(
                    [p for p in tdp.iterdir() if p.is_file()],
                    key=lambda p: p.name,
                )

                # Filter noise by size threshold.
                kept: list[Path] = []
                for p in produced:
                    try:
                        if p.stat().st_size < MIN_IMAGE_BYTES:
                            noise_dropped += 1
                            continue
                    except OSError:
                        continue
                    kept.append(p)

                # Pair captions (already sorted by y) with kept files
                # (sorted by pdfimages order = page-y order). Extras → fallback.
                for k, srcp in enumerate(kept):
                    ext = srcp.suffix.lower() or ".bin"
                    if k < len(captions):
                        stem = captions[k][0]
                        dstp = _unique_path(out_dir, stem, ext)
                        caption_named += 1
                    else:
                        stem = f"page-{str(page_num).zfill(pad)}-img-{k:02d}"
                        dstp = _unique_path(out_dir, stem, ext)
                        fallback_named += 1
                    try:
                        srcp.replace(dstp)
                    except OSError:
                        # cross-fs fallback
                        dstp.write_bytes(srcp.read_bytes())
                        try:
                            srcp.unlink()
                        except OSError:
                            pass
                    rep_files.append(dstp.name)
    finally:
        try:
            pdf_obj.close()
        except Exception:
            pass

    rep = {
        "step": "image-extract",
        "pages_scanned": len(page_idx),
        "images_written": len(rep_files),
        "caption_named": caption_named,
        "fallback_named": fallback_named,
        "noise_dropped": noise_dropped,
        "min_image_bytes": MIN_IMAGE_BYTES,
        "out_dir": str(out_dir),
        "elapsed_s": round(time.perf_counter() - t0, 3),
    }
    if rc_nonzero_pages:
        rep["pdfimages_rc_nonzero_pages"] = rc_nonzero_pages
        if stderr_first:
            rep["pdfimages_stderr_first"] = stderr_first
    return rep


# Registry: name → (kind, callable)
# kind="pdf"  → fn(pdf: pdfplumber.PDF, args, out_dir: Path) -> dict
# kind="path" → fn(pdf_path: Path,      args, out_dir: Path) -> dict
_BUILTIN_STEPS: dict[str, tuple[str, Callable[..., dict]]] = {
    "image-extract": ("path", _image_extract),
    "text-extract":  ("pdf",  _text_extract),
    "table-extract": ("pdf",  _table_extract),
}


def list_builtin_steps() -> list[str]:
    return sorted(_BUILTIN_STEPS.keys())


def is_builtin_step(name: str) -> bool:
    return name in _BUILTIN_STEPS


@dataclass
class LoadedStep:
    name: str
    kind: str  # "pdf" or "path"
    fn: Callable[..., dict]


def load_step(name: str) -> LoadedStep:
    if name not in _BUILTIN_STEPS:
        raise KeyError(
            f"unknown step: {name!r}; known: {', '.join(list_builtin_steps())}"
        )
    kind, fn = _BUILTIN_STEPS[name]
    return LoadedStep(name=name, kind=kind, fn=fn)


# ─── per-step out_dir resolution ──────────────────────────────────────

def _resolve_step_out_dir(step_name: str, args, pdf_path: Path) -> Path:
    """Resolve out_dir for a step in priority order:
    1) explicit --<step-key>-out-dir (e.g. --text-extract-out-dir DIR)
    2) global --out-dir DIR → DIR/<step-verb>/
    3) <pdf-parent>/<step-verb>/
    """
    attr = step_name.replace("-", "_") + "_out_dir"
    explicit = getattr(args, attr, None)
    if explicit:
        return Path(str(explicit))
    global_root = getattr(args, "out_dir", None)
    if global_root:
        return Path(str(global_root)) / step_name
    return pdf_path.parent / step_name


# ─── single-pdf pipeline (multi-step, one pdfplumber.open) ────────────

def run_pipeline_single(
    pdf_path: Path | str,
    step_names: list[str],
    args: argparse.Namespace | None = None,
) -> dict:
    """Run N steps against 1 PDF; open with pdfplumber at most once."""
    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.is_file():
        raise FileNotFoundError(pdf_path)

    if args is None:
        args = argparse.Namespace()
    if not hasattr(args, "pdf"):
        args.pdf = pdf_path

    t0 = time.perf_counter()
    report: dict[str, Any] = {
        "pdf": str(pdf_path),
        "steps": {},
        "timing": {},
    }

    loaded = [load_step(n) for n in step_names]
    needs_pdf = any(s.kind == "pdf" for s in loaded)

    pdf_obj = None
    if needs_pdf:
        if pdfplumber is None:
            raise RuntimeError(
                "pdfplumber not installed; required by steps: "
                + ",".join(s.name for s in loaded if s.kind == "pdf")
            )
        t_open = time.perf_counter()
        try:
            pdf_obj = pdfplumber.open(str(pdf_path))
        except Exception as e:
            report["error"] = f"pdfplumber.open failed: {type(e).__name__}: {e}"
            report["timing"]["total"] = round(time.perf_counter() - t0, 3)
            return report
        report["timing"]["open"] = round(time.perf_counter() - t_open, 3)

    try:
        for s in loaded:
            out_dir = _resolve_step_out_dir(s.name, args, pdf_path)
            t_s = time.perf_counter()
            try:
                if s.kind == "pdf":
                    rep = s.fn(pdf_obj, args, out_dir)
                else:
                    rep = s.fn(pdf_path, args, out_dir)
            except Exception as exc:
                rep = {"step": s.name,
                       "error": f"{type(exc).__name__}: {exc}"}
            report["steps"][s.name] = rep
            report["timing"][f"step:{s.name}"] = round(
                time.perf_counter() - t_s, 3
            )
    finally:
        if pdf_obj is not None:
            try:
                pdf_obj.close()
            except Exception:
                pass

    report["timing"]["total"] = round(time.perf_counter() - t0, 3)
    return report


# ─── parallel across PDFs ─────────────────────────────────────────────

def _worker_payload_to_args(payload: dict) -> argparse.Namespace:
    """Reconstruct args namespace inside child process."""
    ns = argparse.Namespace()
    for k, v in payload.get("args_dict", {}).items():
        setattr(ns, k, v)
    ns.pdf = Path(payload["pdf"])
    return ns


def _worker(payload: dict) -> dict:
    args = _worker_payload_to_args(payload)
    return run_pipeline_single(
        pdf_path=payload["pdf"],
        step_names=payload["steps"],
        args=args,
    )


def run_pipeline_parallel(
    pdf_list: list[Path | str],
    step_names: list[str],
    max_workers: int | None = None,
    args_dict: dict | None = None,
) -> dict[str, dict]:
    """Process-pool across PDFs; each worker re-runs run_pipeline_single."""
    if not pdf_list:
        return {}
    if max_workers is None:
        max_workers = min(len(pdf_list), os.cpu_count() or 4)
    args_dict = args_dict or {}
    payloads = [
        {"pdf": str(p), "steps": list(step_names), "args_dict": args_dict}
        for p in pdf_list
    ]
    results: dict[str, dict] = {}
    if max_workers <= 1 or len(payloads) == 1:
        for p in payloads:
            try:
                results[p["pdf"]] = _worker(p)
            except Exception as exc:
                results[p["pdf"]] = {
                    "pdf": p["pdf"],
                    "error": f"{type(exc).__name__}: {exc}",
                }
        return results
    with ProcessPoolExecutor(max_workers=max_workers) as ex:
        fut_map = {ex.submit(_worker, p): p["pdf"] for p in payloads}
        for fut in as_completed(fut_map):
            pdf = fut_map[fut]
            try:
                results[pdf] = fut.result()
            except Exception as exc:
                results[pdf] = {
                    "pdf": pdf,
                    "error": f"{type(exc).__name__}: {exc}",
                }
    return results


# ─── timing pretty-printer ─────────────────────────────────────────────

def format_timing_table(results: dict[str, dict], wall: float) -> str:
    lines = []
    lines.append("=" * 78)
    lines.append("PDF PIPELINE TIMING REPORT")
    lines.append("=" * 78)
    lines.append(f"{'pdf':46s}  {'open':>6s}  {'total':>8s}")
    lines.append("-" * 78)
    step_set: list[str] = []
    seen: set[str] = set()
    sum_total = 0.0
    for pdf, r in results.items():
        name = Path(pdf).name[:46]
        if not isinstance(r, dict) or "timing" not in r:
            err = r.get("error", str(r)) if isinstance(r, dict) else str(r)
            lines.append(f"{name:46s}  ERROR: {err[:60]}")
            continue
        t = r["timing"]
        open_t = t.get("open", 0.0)
        tot = t.get("total", 0.0)
        sum_total += tot
        lines.append(f"{name:46s}  {open_t:6.3f}  {tot:8.3f}")
        for k in r.get("steps", {}):
            if k not in seen:
                step_set.append(k)
                seen.add(k)
    lines.append("-" * 78)
    if step_set:
        lines.append("Per-step (sum across all PDFs):")
        for step in step_set:
            tot = 0.0
            cnt = 0
            for r in results.values():
                if isinstance(r, dict) and "timing" in r:
                    v = r["timing"].get(f"step:{step}")
                    if v is not None:
                        tot += v
                        cnt += 1
            lines.append(f"  {step:46s}  {tot:7.3f}s  ({cnt} pdfs)")
    lines.append("-" * 78)
    ratio = (sum_total / wall) if wall > 0 else 0
    lines.append(
        f"Wall clock: {wall:.3f}s  |  serial-sum: {sum_total:.3f}s  "
        f"|  ratio = {ratio:.2f}x  |  N={len(results)} PDFs"
    )
    lines.append("=" * 78)
    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════════
# to-docx 引擎段 (原 pdf_to_docx.py 整段并入, 2026-07-31)
#
# PDF → 可编辑 Word(结构提取路线,2026-07-20 立)。
# 设计取向(用户 2026-07-20 拍板):**要「干净可编辑」不要「像素级复刻」**。
#
#   pdfplumber ──抽 text lines + tables──┐
#                                         ├─→ python-docx ─→ .docx
#   (可选) pdfimages ──抽嵌入图───────────┘
#
# 保留:段落顺序 / 标题层级(按字号推断) / 列表项 / **PDF 表格 → 真 Word 表格** /
# 可选嵌图。丢弃:原版式(字体族·分栏·配色·图文环绕)—— 要那个只有
# pdf2docx(PyMuPDF),而本族已立红线「不引入 pymupdf」(AGPL 传染,
# 见 cc-home/commands/pdf.md:40)。
#
# 为什么不走 markitdown→md→md2word:md2word 会强制套「院公文模板」样式,
# 产出是公文风而非原文档结构;且复杂表格经 md 管道表中转会塌。
# ═══════════════════════════════════════════════════════════════════════

# ───────────────────────────────────────────── 版面常量
#
# 这些阈值决定「这行是标题还是正文」「这两行要不要接成一段」。都按相对量
# （相对正文字号中位数 / 页宽百分比）算，绝对 pt 值在不同 PDF 上没有可比性。

H1_RATIO = 1.45      # 字号 ≥ 正文中位数 × 此值 → 一级标题
H2_RATIO = 1.22      # 同上 → 二级标题
H3_RATIO = 1.08      # 同上 + 加粗/短行 → 三级标题
SHORT_LINE = 0.72    # 行宽 < 页面文字区宽度 × 此值 = "短行"（标题候选/段落末行）
GAP_NEW_PARA = 0.62  # 行间距 > 行高 × 此值 → 强制断段
INDENT_TOL = 3.0     # x0 相差在此 pt 内视为「同一左边界」

# 行末出现这些 → 这一行是段落终点，下一行必定另起段
_END_PUNCT = "。！？；!?;:：」』”》】…"
# 行首出现这些 → 这一行必定另起段（列表/编号/条款）
_LIST_HEAD = re.compile(
    r"^\s*(?:[•·▪◦‣▶►■□●○–—\-\*]\s+"           # 项目符号
    r"|\(?\d{1,3}[\.\)、]\s*"                     # 1. / 1) / (1) / 1、
    r"|\(?[a-zA-Z][\.\)]\s+"                      # a. / (b)
    r"|[（(][一二三四五六七八九十百]+[）)]"       # （一）
    r"|第[一二三四五六七八九十百千0-9]+[章节条款项]"  # 第三章 / 第5条
    r")"
)
_BULLET_HEAD = re.compile(r"^\s*[•·▪◦‣▶►■□●○–—]\s+")
# 西文之间接续要补空格；中日韩之间不补。
# 注意 U+FF00–FFEF 这个「全角块」里混着两类东西，**不能整块算 CJK**：
#   全角标点（，。；：（）＂）→ 算 CJK，两侧不补空格
#   全角字母数字（Ａ-Ｚ ａ-ｚ ０-９）→ 是西文，必须补空格
# 整块算 CJK 会拼出「WaterResourcesandPower」（实测踩过）。
_CJK = re.compile(
    r"[　-〿"      # CJK 标点
    r"㐀-鿿"       # 汉字
    r"豈-﫿"       # 兼容汉字
    r"！-／"       # 全角标点 ！＂＃…／
    r"：-＠"       # ：；＜＝＞？＠
    r"［-｀"       # ［＼］＾＿｀
    r"｛-･"       # ｛｜｝～、。
    r"]"
)


def _is_cjk(ch: str) -> bool:
    return bool(_CJK.match(ch))


# ───────────────────────────────────────────── 抽取层


def _word_size(words: list[dict]) -> float:
    """一组词的代表字号 = 各词字号中位数（避免上下标/脚注号拉偏）。"""
    sizes = [w["size"] for w in words if w.get("size")]
    return statistics.median(sizes) if sizes else 0.0


def _word_bold(words: list[dict]) -> bool:
    """半数以上的词字体名带 Bold/Black/Heavy → 视为加粗行。"""
    names = [str(w.get("fontname", "")) for w in words]
    if not names:
        return False
    hits = sum(1 for n in names if re.search(r"bold|black|heavy|semib", n, re.I))
    return hits * 2 > len(names)


def _join_words(words: list[dict]) -> str:
    """词序列 → 行文本。CJK 相邻不加空格，其余加一个空格。

    **必须在词级做，不能在 char 级做**：char 级要靠「间隙 > k×字宽」猜词边界，
    而学术期刊/两端对齐会把西文字距整体拉大，任何阈值都会被骗
    （实测踩过：「Vol.44 No.2」→「V o l . 4 4 N o . 2」、
    「Mann-Kendall」→「M a n n - K e n d a l l」）。
    pdfplumber 的 extract_words 已用 PDF 内部的字距/字体信息切好词，直接信它。"""
    parts: list[str] = []
    for i, w in enumerate(words):
        t = w.get("text", "")
        if i and parts:
            prev = parts[-1]
            if prev and t and not (_is_cjk(prev[-1]) or _is_cjk(t[0])):
                parts.append(" ")
        parts.append(t)
    return "".join(parts)


# 全角字母/数字 → 半角。**只转这些**，中文标点（，。；：（）「」）一律保留 ——
# 用 unicodedata.normalize("NFKC") 是错的，它会把「，」也变成「,」，毁掉中文排版。
_FW_ALNUM = {c: chr(c - 0xFEE0) for c in
             [*range(0xFF21, 0xFF3B), *range(0xFF41, 0xFF5B), *range(0xFF10, 0xFF1A)]}
# 全角句点/连字符：仅当夹在 ASCII 之间才转（Ｖｏｌ．４４ → Vol.44；中文的「．」少见）
_FW_MID = re.compile(r"(?<=[0-9A-Za-z])[．－](?=[0-9A-Za-z])")
_IDEO_SPACE_CJK = re.compile(r"(?<=[一-鿿])　+(?=[一-鿿])")


def _normalize_width(s: str) -> str:
    """全角西文规范化。学术期刊 PDF 大量用全角 ASCII 排页眉/西文
    （实测：'Ｖｏｌ．４４Ｎｏ．２'、'Ｗａｔｅｒ Ｒｅｓｏｕｒｃｅｓ'），
    原样进 Word 会搜不到、难编辑 —— 但这是**内容规范化**不是提取错误，
    故给 --keep-fullwidth 逃生。"""
    s = _IDEO_SPACE_CJK.sub("", s)          # 汉字间的表意空格 = 疏排填充，删掉
    s = s.replace("　", " ")            # 其余表意空格 → 普通空格
    s = s.translate(_FW_ALNUM)
    s = _FW_MID.sub(lambda m: chr(ord(m.group()) - 0xFEE0), s)
    return re.sub(r" {2,}", " ", s)


def _group_lines(words: list[dict]) -> list[list[dict]]:
    """词 → 行。按垂直重叠聚类（同一行的词 top 基本相同，容差取字高的 60%）。"""
    if not words:
        return []
    ws = sorted(words, key=lambda w: (w["top"], w["x0"]))
    lines: list[list[dict]] = [[ws[0]]]
    for w in ws[1:]:
        ref = lines[-1][0]
        tol = max((ref["bottom"] - ref["top"]) * 0.6, 2.0)
        if abs(w["top"] - ref["top"]) <= tol:
            lines[-1].append(w)
        else:
            lines.append([w])
    for ln in lines:
        ln.sort(key=lambda w: w["x0"])
    return lines


def _split_columns(words: list[dict]) -> list[list[dict]]:
    """把一行的词序列按「行内大空隙」切成多栏。

    为什么必须做：按 y 聚行会把左右分栏（简历头部「姓名 | 联系方式」、双栏正文）
    拍平成同一行 → 拼出「曾田力 zengtianli1@126.com」这种鬼话。

    判据不能只看绝对间隙——两端对齐会把整行词距拉大。故再加相对判据：
    真正的分栏留白必须**显著大于同一行内其它所有词距**（一枝独秀 vs 普遍偏大）。"""
    if len(words) < 2:
        return [words]
    sizes = [w["size"] for w in words if w.get("size")] or [10.0]
    em = statistics.median(sizes)

    gaps = [cur["x0"] - prev["x1"] for prev, cur in zip(words, words[1:])]
    abs_thresh = max(em * 2.2, 18.0)
    if not any(g > abs_thresh for g in gaps):
        return [words]
    others = [g for g in gaps if g <= abs_thresh]
    cut = max(abs_thresh, (max(others) if others else 0.0) * 1.8)

    groups: list[list[dict]] = [[words[0]]]
    for prev, cur in zip(words, words[1:]):
        if cur["x0"] - prev["x1"] > cut:
            groups.append([cur])
        else:
            groups[-1].append(cur)
    return groups


PAGE_FRAME_RATIO = 0.85   # 表格 bbox 占页面面积超过此比例 → 是版面边框不是表格
MIN_TABLE_FILL = 0.30     # 单元格非空率低于此 → 是排版网格不是数据表


def _real_tables(page) -> list:
    """页内**真正的数据表**。

    pdfplumber 的 find_tables 会把「整页边框」「排版分隔线围出的区域」也当表格 ——
    实测 IIQE 讲义（PPT 导出，带页框）返回一个 594×841 覆盖全页的 bbox，
    结果 1219 行正文全被当成表格内容吞掉，产出空 docx。故必须过滤：
      ① bbox 不能几乎等于整页（那是页框）
      ② 至少 2 行 × 2 列
      ③ 单元格非空率够高（排版网格大多是空的）"""
    try:
        cands = page.find_tables()
    except Exception:
        return []
    page_area = float(page.width) * float(page.height)
    out = []
    for t in cands:
        x0, y0, x1, y1 = t.bbox
        if page_area and (x1 - x0) * (y1 - y0) > page_area * PAGE_FRAME_RATIO:
            continue                                   # 页框
        try:
            data = t.extract()
        except Exception:
            continue
        if not data or len(data) < 2:
            continue
        ncols = max((len(r) for r in data), default=0)
        if ncols < 2:
            continue
        cells = sum(len(r) for r in data)
        filled = sum(1 for r in data for c in r if c and str(c).strip())
        if not cells or filled / cells < MIN_TABLE_FILL:
            continue                                   # 排版网格
        out.append((t, data))
    return out


def _in_any_bbox(w: dict, boxes: list[tuple[float, float, float, float]]) -> bool:
    """词的中心落在某表格 bbox 内 → 属于该表格。用中心点而非全包含，
    容忍 find_tables 的边界比实际文字略紧。"""
    cy = (w["top"] + w["bottom"]) / 2
    cx = (w["x0"] + w["x1"]) / 2
    return any(x0 <= cx <= x1 and y0 <= cy <= y1 for x0, y0, x1, y1 in boxes)


def _collect_blocks(pdf, drop_tables: bool, norm_width: bool = True) -> list[dict]:
    """把整个 PDF 拍平成有序 block 流：{kind: text|table|pagebreak, ...}。

    表格按它在页内的垂直位置插进正文流（不是一股脑堆到页尾），
    这样 Word 里的图文顺序才和原 PDF 一致。"""
    blocks: list[dict] = []
    for pno, page in enumerate(pdf.pages):
        if pno:
            blocks.append({"kind": "pagebreak"})

        found = [] if drop_tables else _real_tables(page)
        boxes = [t.bbox for t, _ in found]
        tables = [{"kind": "table", "top": t.bbox[1], "data": d} for t, d in found]
        if norm_width:
            for tb in tables:
                tb["data"] = [[(_normalize_width(c) if isinstance(c, str) else c)
                               for c in row] for row in tb["data"]]

        try:
            words = page.extract_words(extra_attrs=["size", "fontname"]) or []
        except Exception as e:
            print(f"  ⚠ 第 {pno + 1} 页取词失败({type(e).__name__})，跳过", file=sys.stderr)
            words = []
        if boxes:
            words = [w for w in words if not _in_any_bbox(w, boxes)]  # 表格内容不重复进正文

        page_items: list[dict] = list(tables)
        for ln_words in _group_lines(words):
            for seg in _split_columns(ln_words):
                txt = _join_words(seg).strip()
                if not txt:
                    continue
                if norm_width:
                    txt = _normalize_width(txt)
                    if not txt:
                        continue
                page_items.append({
                    "kind": "text",
                    "top": min(w["top"] for w in seg),
                    "bottom": max(w["bottom"] for w in seg),
                    "x0": min(w["x0"] for w in seg),
                    "x1": max(w["x1"] for w in seg),
                    "text": txt,
                    "size": _word_size(seg),
                    "bold": _word_bold(seg),
                })

        # 同一 y 带内先左后右（分栏切出来的右段排在左段之后）
        page_items.sort(key=lambda b: (round(b["top"], 1), b.get("x0", 0)))
        # 页面文字区宽度 —— 判「短行」的分母，用实际内容跨度而非纸张宽度
        xs = [b["x0"] for b in page_items if b["kind"] == "text"]
        xe = [b["x1"] for b in page_items if b["kind"] == "text"]
        span = (max(xe) - min(xs)) if xs and xe else float(page.width)
        for b in page_items:
            if b["kind"] == "text":
                b["page_span"] = span
                b["page_left"] = min(xs) if xs else 0.0
        blocks.extend(page_items)
    return blocks


# ───────────────────────────────────────────── 段落重组
#
# PDF 里「一行」不等于「一段」。直接一行一段，Word 里会碎成几百个短段落，
# 换行全是硬断，改一个字整段错位 —— 这是 PDF 转 Word 最劝退的地方，必须重组。


def _body_size(blocks: list[dict]) -> float:
    """正文字号 = 全文 text 行字号的中位数（标题是少数派，拉不动中位数）。"""
    sizes = [b["size"] for b in blocks if b["kind"] == "text" and b["size"]]
    return statistics.median(sizes) if sizes else 10.0


def _heading_level(b: dict, body: float) -> int | None:
    """返回 1/2/3 = 标题级别，None = 正文。

    判据 = 字号为主、加粗+短行为辅。纯加粗但字号=正文的不算标题（那多半是行内强调）。"""
    if not body:
        return None
    r = b["size"] / body
    short = (b["x1"] - b["x0"]) < b.get("page_span", 1e9) * SHORT_LINE
    if r >= H1_RATIO:
        return 1
    if r >= H2_RATIO:
        return 2
    if r >= H3_RATIO and b["bold"] and short:
        return 3
    return None


def _should_join(prev: dict, cur: dict, body: float, para: dict) -> bool:
    """prev 和 cur 是不是同一个自然段的连续两行？para = 正在累积的段落（提供首行上下文）。"""
    if _LIST_HEAD.match(cur["text"]):
        return False                                    # 新列表项/编号 → 断
    if prev["text"] and prev["text"][-1] in _END_PUNCT:
        return False                                    # 上行已收句 → 断
    gap = cur["top"] - prev["bottom"]
    line_h = max(prev["bottom"] - prev["top"], 1.0)
    if gap > line_h * GAP_NEW_PARA:
        return False                                    # 行距明显变大 → 断
    if abs(cur["size"] - prev["size"]) > 0.6:
        return False                                    # 字号跳变 → 不是同段

    # 列表项的续行是「悬挂缩进」：比 bullet 首行更靠右。这条必须**最先**判，
    # 排在「左边界跳变」「短行=段末」之前——悬挂缩进本身就是一次左边界跳变，
    # 后者会把一个 bullet 腰斩成两段（实测：「…31 个 hook（PreToolUse…」x0=46.3
    # ／「强制守卫）、3 个专用 subagent」x0=59.6，差 13.3pt 正好踩线）。
    if para.get("bullet"):
        return cur["x0"] >= para["first_x0"] - INDENT_TOL

    if abs(cur["x0"] - prev["x0"]) > INDENT_TOL * 4:
        return False                                    # 左边界大幅跳变 → 多半是另一栏
    if (prev["x1"] - prev["x0"]) < prev.get("page_span", 1e9) * SHORT_LINE:
        return False                                    # 上行是短行 = 段落末行 → 断
    if cur["x0"] > prev["x0"] + INDENT_TOL * 2:
        return False                                    # 本行明显右缩进 → 新段首行
    return True


def _join_text(a: str, b: str) -> str:
    """接续两行。中文之间直接拼；西文之间补空格；连字符断词去掉连字符。"""
    if not a:
        return b
    if a.endswith("-") and b[:1].isalpha():
        return a[:-1] + b
    if _is_cjk(a[-1]) or _is_cjk(b[:1]):
        return a + b
    return a + " " + b


def _build_paragraphs(blocks: list[dict], body: float) -> list[dict]:
    """block 流 → 段落流。相邻 text 行按 _should_join 归并；表格/分页原样穿过。"""
    out: list[dict] = []
    cur: dict | None = None

    def flush():
        nonlocal cur
        if cur:
            out.append(cur)
            cur = None

    for b in blocks:
        if b["kind"] != "text":
            flush()
            out.append(b)
            continue

        lvl = _heading_level(b, body)
        if lvl:
            flush()
            out.append({"kind": "heading", "level": lvl, "text": b["text"]})
            continue

        if cur and _should_join(cur["_last"], b, body, cur):
            cur["text"] = _join_text(cur["text"], b["text"])
            cur["_last"] = b
        else:
            flush()
            cur = {
                "kind": "para",
                "text": b["text"],
                "bullet": bool(_LIST_HEAD.match(b["text"])),
                "first_x0": b["x0"],
                "_last": b,
            }
    flush()
    return out


# ───────────────────────────────────────────── 图片（可选）


def _extract_images(pdf_path: Path, work: Path) -> list[Path]:
    """pdfimages -all 抽嵌入图。按铁律 #6 过滤噪音：< 3KB 的多半是 soft-mask/纯色分隔条。"""
    exe = shutil.which("pdfimages") or _PDFIMAGES
    if not Path(exe).exists():
        return []
    cp = _run_pdfimages(pdf_path, work / "img", binary=exe)
    if cp.returncode != 0:
        return []
    keep = []
    for p in sorted(work.glob("img-*")):
        if p.suffix.lower() in (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".ppm", ".pbm"):
            if p.stat().st_size >= MIN_IMAGE_BYTES:
                keep.append(p)
    return keep


# ── 惰性加载 python-docx + surgical 收口(只有 convert to-docx 路径需要)────────
#
# 挪进函数级的原因:read/merge 等子命令不该平白背上 python-docx 的加载;
# `import docx_safe_save` 必须先于任何 Document() 构造(lib/docx_safe_save.py
# surgical 收口,炸开面 60→1)。append 不是 insert(0) —— lib/ 和
# scripts/document/sub/ 都有 styles.py,插 0 位会顶掉脚本自己那份。

Document = None  # type: ignore[assignment]  # _ensure_docx_deps() 填充
WD_ALIGN_PARAGRAPH = Inches = Pt = None  # type: ignore[assignment]


def _ensure_docx_deps() -> None:
    global Document, WD_ALIGN_PARAGRAPH, Inches, Pt
    if Document is not None:
        return
    # pdfminer 对字体描述符不规范的 PDF 会每字符刷一条
    # "Could not get FontBBox from font descriptor" —— 纯噪音,且会淹没我们自己的
    # 警告、污染 GUI 后端的 stderr。它不影响提取结果,直接闭嘴。
    logging.getLogger("pdfminer").setLevel(logging.ERROR)
    _lib = str(Path(__file__).resolve().parents[2] / "lib")
    if _lib not in sys.path:
        sys.path.append(_lib)
    import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py
    from docx import Document as _Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
    from docx.shared import Inches as _Inches, Pt as _Pt
    Document = _Document
    WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
    Inches, Pt = _Inches, _Pt

# ───────────────────────────────────────────── 渲染层


def _set_base_style(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = "PingFang SC"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.15


def _add_table(doc: Document, data: list[list]) -> None:
    rows = [r for r in data if r and any((c or "").strip() for c in r)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            val = (row[j] if j < len(row) else "") or ""
            cell = t.cell(i, j)
            # PDF 单元格里的软换行是排版产物，不是语义换行 → 压成空格
            cell.text = re.sub(r"\s*\n\s*", " ", str(val)).strip()
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True
    doc.add_paragraph()


def convert(
    pdf_path: Path,
    out_path: Path | None = None,
    with_images: bool = False,
    with_tables: bool = True,
    norm_width: bool = True,
) -> dict:
    """PDF → docx。返回 {ok, output, pages, paragraphs, tables, images, error?}"""
    _ensure_docx_deps()
    pdf_path = Path(pdf_path).expanduser().resolve()
    if not pdf_path.exists():
        return {"ok": False, "error": f"文件不存在: {pdf_path}"}
    out = Path(out_path).expanduser() if out_path else pdf_path.with_suffix(".docx")

    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            npages = len(pdf.pages)
            blocks = _collect_blocks(pdf, drop_tables=not with_tables, norm_width=norm_width)
            # 兜底：表格识别把正文全吞了（_real_tables 的过滤没覆盖到的诡异版面）
            # → 关掉表格重来一遍。宁可丢表格结构，不可产出空文档。
            if with_tables and not any(b["kind"] == "text" for b in blocks):
                if any(b["kind"] == "table" for b in blocks):
                    print("  ⚠ 表格识别吞掉了全部正文，已关闭表格识别重试", file=sys.stderr)
                    blocks = _collect_blocks(pdf, drop_tables=True, norm_width=norm_width)
    except Exception as e:  # 加密/损坏/非 PDF
        return {"ok": False, "error": f"打开 PDF 失败: {type(e).__name__}: {e}"}

    if not any(b["kind"] == "text" for b in blocks):
        return {"ok": False,
                "error": "PDF 里没有可提取的文字层（多半是扫描件）——"
                         "先跑 ocrmypdf 加文本层，或用 vision_ocr.py 识别"}

    body = _body_size(blocks)
    items = _build_paragraphs(blocks, body)

    doc = Document()
    _set_base_style(doc)
    n_tab = 0
    for it in items:
        k = it["kind"]
        if k == "heading":
            doc.add_heading(it["text"], level=min(it["level"], 4))
        elif k == "para":
            style = "List Bullet" if it["bullet"] and _BULLET_HEAD.match(it["text"]) else None
            txt = _BULLET_HEAD.sub("", it["text"]) if style else it["text"]
            doc.add_paragraph(txt, style=style)
        elif k == "table":
            _add_table(doc, it["data"])
            n_tab += 1
        elif k == "pagebreak":
            doc.add_paragraph()

    n_img = 0
    if with_images:
        with tempfile.TemporaryDirectory() as td:
            imgs = _extract_images(pdf_path, Path(td))
            if imgs:
                doc.add_page_break()
                h = doc.add_paragraph("附：文档嵌入图")
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for p in imgs:
                    try:
                        doc.add_picture(str(p), width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        n_img += 1
                    except Exception:
                        continue  # 有些 ppm/掩码 python-docx 吃不进，跳过不中断

    try:
        doc.save(str(out))
    except Exception as e:
        return {"ok": False, "error": f"写 docx 失败: {type(e).__name__}: {e}"}

    # macOS quarantine：不清的话 Word 开出来是「受保护视图」，审阅按钮全灰
    # （memory feedback-docx-macos-quarantine）
    subprocess.run(["xattr", "-d", "com.apple.quarantine", str(out)],
                   capture_output=True)

    return {
        "ok": True,
        "output": str(out),
        "pages": npages,
        "paragraphs": sum(1 for i in items if i["kind"] in ("para", "heading")),
        "tables": n_tab,
        "images": n_img,
    }

# ═══════════════════════════════════════════════════════════════════════
# 1. read
# ═══════════════════════════════════════════════════════════════════════

def _cmd_read(args: argparse.Namespace) -> int:
    pdf_path = Path(args.pdf).resolve()
    if not pdf_path.is_file():
        print(f"ERROR: file not found: {pdf_path}", file=sys.stderr)
        return 2

    if args.list:
        # pdfinfo for page count / metadata
        cp = subprocess.run(
            [_PDFINFO, str(pdf_path)], capture_output=True, text=True
        )
        if cp.returncode == 0:
            print(cp.stdout.rstrip())
        else:
            print(f"[pdfinfo rc={cp.returncode}] {cp.stderr.strip()}",
                  file=sys.stderr)

        # pypdf outline tree
        if PdfReader is None:
            print("\n[outline] pypdf not installed", file=sys.stderr)
            return 0 if cp.returncode == 0 else cp.returncode
        try:
            reader = PdfReader(str(pdf_path))
            outline = reader.outline
        except Exception as e:
            print(f"\n[outline] error: {type(e).__name__}: {e}",
                  file=sys.stderr)
            return 0 if cp.returncode == 0 else cp.returncode

        print("\n── outline ──")
        if not outline:
            print("(no outline / bookmarks)")
            return 0 if cp.returncode == 0 else cp.returncode

        def _walk(node, depth=0):
            if isinstance(node, list):
                for item in node:
                    _walk(item, depth)
                return
            title = getattr(node, "title", str(node))
            try:
                page_num = reader.get_destination_page_number(node) + 1
            except Exception:
                page_num = "?"
            print(f"{'  ' * depth}- [{page_num}] {title}")

        _walk(outline)
        return 0

    # Default: extract text from page(s) to stdout
    if pdfplumber is None:
        print("ERROR: pdfplumber not installed", file=sys.stderr)
        return 2
    try:
        pdf = pdfplumber.open(str(pdf_path))
    except Exception as e:
        print(f"ERROR: pdfplumber.open failed: {type(e).__name__}: {e}",
              file=sys.stderr)
        return 1
    try:
        total = len(pdf.pages)
        if args.page is not None:
            spec = str(args.page)
        elif args.pages:
            spec = args.pages
        else:
            spec = None
        try:
            idx = parse_page_spec(spec, total)
        except ValueError as e:
            print(f"ERROR: {e}", file=sys.stderr)
            return 2
        for i in idx:
            try:
                txt = pdf.pages[i].extract_text() or ""
            except Exception as e:
                txt = f"[ERROR page {i+1}: {type(e).__name__}: {e}]"
            print(f"===== page {i+1} =====")
            print(txt)
        return 0
    finally:
        try:
            pdf.close()
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════════════
# 2. extract image / 3. extract text / 4. extract table
# ═══════════════════════════════════════════════════════════════════════

def _cmd_extract_image(args: argparse.Namespace) -> int:
    pdf_path = Path(args.pdf).resolve()
    if not pdf_path.is_file():
        print(f"ERROR: file not found: {pdf_path}", file=sys.stderr)
        return 2
    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    t0 = time.perf_counter()
    cp = _run_pdfimages(pdf_path, out_dir / "img")
    elapsed = time.perf_counter() - t0
    n_files = sum(1 for p in out_dir.iterdir() if p.is_file())
    print(f"[extract image] rc={cp.returncode} files={n_files} "
          f"elapsed={elapsed:.3f}s out={out_dir}")
    if cp.returncode != 0:
        print(cp.stderr, file=sys.stderr)
    return cp.returncode


def _cmd_extract_text(args: argparse.Namespace) -> int:
    pdf_path = Path(args.pdf).resolve()
    if not pdf_path.is_file():
        print(f"ERROR: file not found: {pdf_path}", file=sys.stderr)
        return 2
    if pdfplumber is None:
        print("ERROR: pdfplumber not installed", file=sys.stderr)
        return 2
    out_dir = Path(args.out_dir).resolve()
    # Wrap into pipeline-style namespace for code reuse
    ns = argparse.Namespace(
        pdf=pdf_path,
        text_extract_pages=args.pages,
        text_extract_single=args.single,
        text_extract_out_dir=out_dir,
    )
    rep = run_pipeline_single(pdf_path, ["text-extract"], args=ns)
    sub = rep.get("steps", {}).get("text-extract", {})
    if "error" in sub:
        print(f"[extract text] ERROR: {sub['error']}", file=sys.stderr)
        return 1
    print(f"[extract text] mode={sub.get('mode')} "
          f"pages={sub.get('pages_requested')} "
          f"files={sub.get('files_written')} "
          f"elapsed={sub.get('elapsed_s')}s out={sub.get('out_dir')}")
    return 0


def _cmd_extract_table(args: argparse.Namespace) -> int:
    pdf_path = Path(args.pdf).resolve()
    if not pdf_path.is_file():
        print(f"ERROR: file not found: {pdf_path}", file=sys.stderr)
        return 2
    if pdfplumber is None:
        print("ERROR: pdfplumber not installed", file=sys.stderr)
        return 2
    out_dir = Path(args.out_dir).resolve()
    ns = argparse.Namespace(
        pdf=pdf_path,
        table_extract_pages=args.pages,
        table_extract_out_dir=out_dir,
    )
    rep = run_pipeline_single(pdf_path, ["table-extract"], args=ns)
    sub = rep.get("steps", {}).get("table-extract", {})
    if "error" in sub:
        print(f"[extract table] ERROR: {sub['error']}", file=sys.stderr)
        return 1
    print(f"[extract table] tables={sub.get('total_tables')} "
          f"pages_scanned={sub.get('pages_scanned')} "
          f"elapsed={sub.get('elapsed_s')}s out={sub.get('out_dir')}")
    counts = sub.get("per_page_table_counts", {})
    nonzero = {k: v for k, v in counts.items() if v}
    if nonzero:
        print(f"  per-page (nonzero): {nonzero}")
    return 0


# ═══════════════════════════════════════════════════════════════════════
# 5. split by-bookmark / 6. split by-page-range
# ═══════════════════════════════════════════════════════════════════════

def _flatten_top_outline(reader) -> list[tuple[str, int]]:
    """Return top-level [(title, page_num_0idx), ...] from pypdf outline."""
    outline = reader.outline
    out: list[tuple[str, int]] = []
    if not outline:
        return out
    for item in outline:
        if isinstance(item, list):
            # nested children of previous top item → skip (top-only)
            continue
        title = getattr(item, "title", None) or "untitled"
        try:
            pno = reader.get_destination_page_number(item)
        except Exception:
            continue
        out.append((title, pno))
    return out


def _cmd_split_by_bookmark(args: argparse.Namespace) -> int:
    pdf_path = Path(args.pdf).resolve()
    if not pdf_path.is_file():
        print(f"ERROR: file not found: {pdf_path}", file=sys.stderr)
        return 2
    if PdfReader is None or PdfWriter is None:
        print("ERROR: pypdf not installed", file=sys.stderr)
        return 2
    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    reader = PdfReader(str(pdf_path))
    items = _flatten_top_outline(reader)
    if not items:
        print("no outline found", file=sys.stderr)
        return 3

    total_pages = len(reader.pages)
    # Build (title, start, end) — end = next start - 1, last = total - 1
    segments: list[tuple[str, int, int]] = []
    for i, (title, start) in enumerate(items):
        end = (items[i + 1][1] - 1) if i + 1 < len(items) else (total_pages - 1)
        if end < start:
            end = start
        segments.append((title, start, end))

    n_written = 0
    pad = max(2, len(str(len(segments))))
    for idx, (title, start, end) in enumerate(segments, 1):
        writer = PdfWriter()
        for p in range(start, end + 1):
            writer.add_page(reader.pages[p])
        safe = safe_filename(title)
        out_file = out_dir / f"{str(idx).zfill(pad)}-{safe}.pdf"
        with out_file.open("wb") as f:
            writer.write(f)
        n_written += 1
        print(f"  [{idx}] pages {start+1}-{end+1} → {out_file.name}")
    print(f"[split by-bookmark] {n_written} parts → {out_dir}")
    return 0


def _cmd_split_by_page_range(args: argparse.Namespace) -> int:
    pdf_path = Path(args.pdf).resolve()
    if not pdf_path.is_file():
        print(f"ERROR: file not found: {pdf_path}", file=sys.stderr)
        return 2
    if PdfReader is None or PdfWriter is None:
        print("ERROR: pypdf not installed", file=sys.stderr)
        return 2
    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        ranges = parse_ranges_spec(args.ranges)
    except ValueError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 2

    reader = PdfReader(str(pdf_path))
    total = len(reader.pages)
    pad = max(2, len(str(len(ranges))))
    for idx, (a, b) in enumerate(ranges, 1):
        if a > total:
            print(f"  [{idx}] skip {a}-{b} (file has {total} pages)",
                  file=sys.stderr)
            continue
        b_eff = min(b, total)
        writer = PdfWriter()
        for p in range(a - 1, b_eff):
            writer.add_page(reader.pages[p])
        out_file = out_dir / f"{str(idx).zfill(pad)}-pages-{a}-{b_eff}.pdf"
        with out_file.open("wb") as f:
            writer.write(f)
        print(f"  [{idx}] pages {a}-{b_eff} → {out_file.name}")
    print(f"[split by-page-range] {len(ranges)} ranges → {out_dir}")
    return 0


# ═══════════════════════════════════════════════════════════════════════
# 7. merge
# ═══════════════════════════════════════════════════════════════════════

def _cmd_merge(args: argparse.Namespace) -> int:
    if PdfWriter is None:
        print("ERROR: pypdf not installed", file=sys.stderr)
        return 2
    inputs = [Path(p).resolve() for p in args.pdfs]
    missing = [p for p in inputs if not p.is_file()]
    if missing:
        print(f"ERROR: files not found: {missing}", file=sys.stderr)
        return 2
    out = Path(args.out).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    writer = PdfWriter()
    total_pages = 0
    for p in inputs:
        try:
            reader = PdfReader(str(p))
            for page in reader.pages:
                writer.add_page(page)
                total_pages += 1
            print(f"  + {p.name} ({len(reader.pages)} pages)")
        except Exception as e:
            print(f"ERROR: {p}: {type(e).__name__}: {e}", file=sys.stderr)
            return 1
    with out.open("wb") as f:
        writer.write(f)
    print(f"[merge] {len(inputs)} files, {total_pages} pages → {out}")
    return 0


# ═══════════════════════════════════════════════════════════════════════
# 8. decrypt
# ═══════════════════════════════════════════════════════════════════════

def _cmd_convert_to_docx(args: argparse.Namespace) -> int:
    """PDF → 可编辑 Word(结构提取引擎已内嵌本文件,见「to-docx 引擎段」)。"""
    r = convert(
        Path(args.pdf),
        Path(args.out) if args.out else None,
        with_images=args.images,
        with_tables=not args.no_tables,
        norm_width=not args.keep_fullwidth,
    )
    if not r["ok"]:
        print(f"✖ {r['error']}", file=sys.stderr)
        return 1
    bits = [f"{r['pages']} 页 → {r['paragraphs']} 段"]
    if r["tables"]:
        bits.append(f"{r['tables']} 表")
    if r["images"]:
        bits.append(f"{r['images']} 图")
    print(f"✓ {Path(r['output']).name}（{' / '.join(bits)}）")
    return 0


def _cmd_decrypt(args: argparse.Namespace) -> int:
    pdf_path = Path(args.pdf).resolve()
    if not pdf_path.is_file():
        print(f"ERROR: file not found: {pdf_path}", file=sys.stderr)
        return 2
    if args.password:
        out = Path(args.out).resolve() if args.out else (
            pdf_path.with_name(pdf_path.stem + ".decrypted.pdf")
        )
        out.parent.mkdir(parents=True, exist_ok=True)
        cmd = [_QPDF, f"--password={args.password}", "--decrypt",
               str(pdf_path), str(out)]
        cp = subprocess.run(cmd, capture_output=True, text=True)
        if cp.returncode == 0:
            print(f"[decrypt] OK → {out}")
        else:
            print(f"[decrypt] qpdf rc={cp.returncode}: {cp.stderr.strip()}",
                  file=sys.stderr)
        return cp.returncode
    # No password: invoke the cc-home pdf-decrypt skill auto-guesser
    if not _DECRYPT_SKILL.is_file():
        print(f"ERROR: decrypt skill not found at {_DECRYPT_SKILL}",
              file=sys.stderr)
        return 2
    cmd = [_PYTHON3, str(_DECRYPT_SKILL), str(pdf_path)]
    if args.out:
        cmd.extend(["--out", str(Path(args.out).resolve())])
    cp = subprocess.run(cmd)
    return cp.returncode


# ═══════════════════════════════════════════════════════════════════════
# 9. pipeline run
# ═══════════════════════════════════════════════════════════════════════

def _parse_steps(s: str) -> list[str]:
    out = [t.strip() for t in s.split(",") if t.strip()]
    if not out:
        raise argparse.ArgumentTypeError("--steps cannot be empty")
    return out


def _resolve_glob(token: str) -> list[Path]:
    """Expand a glob token to existing PDF paths.

    Quoted glob from CLI arrives literal; we expand here.
    Also handles literal file paths.
    """
    p = Path(token)
    if p.is_file():
        return [p.resolve()]
    matches = _glob.glob(token, recursive=True)
    return [Path(m).resolve() for m in matches if Path(m).is_file()]


def _cmd_pipeline_run(args: argparse.Namespace) -> int:
    # Expand globs across all positional tokens
    pdfs: list[Path] = []
    for tok in args.pdfs:
        found = _resolve_glob(tok)
        if not found:
            print(f"[pipeline] WARN: no matches for {tok!r}", file=sys.stderr)
        pdfs.extend(found)
    # De-dup preserving order
    seen: set[str] = set()
    pdfs_dedup: list[Path] = []
    for p in pdfs:
        sp = str(p)
        if sp not in seen:
            seen.add(sp)
            pdfs_dedup.append(p)
    pdfs = pdfs_dedup
    if not pdfs:
        print("[pipeline] no PDFs matched after glob expansion",
              file=sys.stderr)
        return 2

    steps = args.steps
    unknown = [s for s in steps if not is_builtin_step(s)]
    if unknown:
        print(f"[pipeline] unknown steps: {unknown}; "
              f"known: {list_builtin_steps()}", file=sys.stderr)
        return 2

    print(f"[pipeline] {len(pdfs)} PDFs × {len(steps)} steps "
          f"({'parallel' if args.parallel else 'serial'} mode)")
    print(f"[pipeline] steps: {' → '.join(steps)}")
    for p in pdfs:
        print(f"  - {p.name}")

    # Build args_dict for child workers (only JSON-serializable scalars)
    args_dict = {
        "out_dir": str(args.out_dir) if args.out_dir else None,
        "image_extract_out_dir":
            str(args.image_extract_out_dir) if args.image_extract_out_dir
            else None,
        "text_extract_out_dir":
            str(args.text_extract_out_dir) if args.text_extract_out_dir
            else None,
        "text_extract_pages": args.text_extract_pages,
        "text_extract_single": args.text_extract_single,
        "table_extract_out_dir":
            str(args.table_extract_out_dir) if args.table_extract_out_dir
            else None,
        "table_extract_pages": args.table_extract_pages,
        "pages": args.pages,
        "single": args.single,
    }

    t0 = time.perf_counter()
    if not args.parallel or len(pdfs) == 1:
        results: dict[str, dict] = {}
        ns = argparse.Namespace(**args_dict)
        for p in pdfs:
            ns.pdf = p
            try:
                results[str(p)] = run_pipeline_single(p, steps, args=ns)
            except Exception as exc:
                results[str(p)] = {
                    "pdf": str(p),
                    "error": f"{type(exc).__name__}: {exc}",
                }
    else:
        results = run_pipeline_parallel(
            pdfs, steps,
            max_workers=args.max_workers,
            args_dict=args_dict,
        )
    wall = time.perf_counter() - t0

    print()
    print(format_timing_table(results, wall))

    # Per-PDF report JSON
    if args.report_dir:
        rd = Path(args.report_dir).resolve()
        rd.mkdir(parents=True, exist_ok=True)
        for pdf, rep in results.items():
            stem = Path(pdf).stem
            out = rd / f"pipeline-{stem}.json"
            out.write_text(
                json.dumps(rep, ensure_ascii=False, indent=2, default=str),
                encoding="utf-8",
            )
        print(f"[pipeline] reports → {rd}")

    any_err = any(
        isinstance(r, dict) and ("error" in r) for r in results.values()
    )
    # Also check step-level errors
    if not any_err:
        for r in results.values():
            if isinstance(r, dict):
                for sub in r.get("steps", {}).values():
                    if isinstance(sub, dict) and "error" in sub:
                        any_err = True
                        break
    return 1 if any_err else 0


# ═══════════════════════════════════════════════════════════════════════
# Argparse wiring
# ═══════════════════════════════════════════════════════════════════════

def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="pdf_cli",
        description=(
            "PDF unified CLI (2026-05-28)\n"
            "Subcommands: read / extract image|text|table / "
            "split by-bookmark|by-page-range / merge / decrypt / pipeline run"
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    add_parallel_args(p)
    sub = p.add_subparsers(dest="command", metavar="<subcommand>")

    # read
    rp = sub.add_parser("read", help="extract page text / list outline")
    rp.add_argument("pdf")
    g = rp.add_mutually_exclusive_group()
    g.add_argument("--page", type=int, help="single page number (1-indexed)")
    g.add_argument("--pages", help="page spec, e.g. '1-5,8'")
    g.add_argument("--list", action="store_true",
                   help="pdfinfo + outline tree")
    rp.set_defaults(func=_cmd_read)

    # extract <kind>
    ep = sub.add_parser("extract", help="extract image / text / table")
    ep_sub = ep.add_subparsers(dest="extract_kind", metavar="<kind>")
    ep.set_defaults(func=lambda a: (ep.print_help() or 0))

    # extract image
    eip = ep_sub.add_parser("image", help="pdfimages -all → out-dir/img-*")
    eip.add_argument("pdf")
    eip.add_argument("--out-dir", required=True)
    eip.set_defaults(func=_cmd_extract_image)

    # extract text
    etp = ep_sub.add_parser("text", help="pdfplumber per-page txt")
    etp.add_argument("pdf")
    etp.add_argument("--out-dir", required=True)
    etp.add_argument("--single", action="store_true",
                     help="concat all into full.txt")
    etp.add_argument("--pages", help="page spec, e.g. '1-5'")
    etp.set_defaults(func=_cmd_extract_text)

    # extract table
    ettp = ep_sub.add_parser("table", help="pdfplumber per-page CSV tables")
    ettp.add_argument("pdf")
    ettp.add_argument("--out-dir", required=True)
    ettp.add_argument("--pages", help="page spec, e.g. '1-5'")
    ettp.set_defaults(func=_cmd_extract_table)

    # split <kind>
    sp = sub.add_parser("split", help="split PDF by bookmark / page-range")
    sp_sub = sp.add_subparsers(dest="split_kind", metavar="<kind>")
    sp.set_defaults(func=lambda a: (sp.print_help() or 0))

    # split by-bookmark
    sbb = sp_sub.add_parser("by-bookmark",
                            help="split by top-level outline entries")
    sbb.add_argument("pdf")
    sbb.add_argument("--out-dir", required=True)
    sbb.set_defaults(func=_cmd_split_by_bookmark)

    # split by-page-range
    sbpr = sp_sub.add_parser("by-page-range",
                             help='split by --ranges "1-10,11-20"')
    sbpr.add_argument("pdf")
    sbpr.add_argument("--ranges", required=True,
                      help='comma-separated A-B ranges (1-indexed)')
    sbpr.add_argument("--out-dir", required=True)
    sbpr.set_defaults(func=_cmd_split_by_page_range)

    # merge
    mp = sub.add_parser("merge", help="concatenate multiple PDFs")
    mp.add_argument("pdfs", nargs="+", help="input PDFs in order")
    mp.add_argument("--out", required=True, help="output combined PDF")
    mp.set_defaults(func=_cmd_merge)

    # decrypt
    dp = sub.add_parser("decrypt", help="decrypt PDF via qpdf / auto-guess")
    dp.add_argument("pdf")
    dp.add_argument("--password", default=None,
                    help="explicit password (qpdf direct)")
    dp.add_argument("--out", default=None, help="output path")
    dp.set_defaults(func=_cmd_decrypt)

    # convert group
    cp = sub.add_parser("convert", help="convert PDF to other formats")
    cp_sub = cp.add_subparsers(dest="convert_cmd", metavar="<cmd>")
    cp.set_defaults(func=lambda a: (cp.print_help() or 0))

    ctd = cp_sub.add_parser("to-docx", help="PDF → editable Word (structure extraction)")
    ctd.add_argument("pdf")
    ctd.add_argument("-o", "--out", default=None, help="output .docx (default: alongside)")
    ctd.add_argument("--images", action="store_true", help="append embedded images")
    ctd.add_argument("--no-tables", action="store_true", help="skip table detection")
    ctd.add_argument("--keep-fullwidth", action="store_true",
                     help="keep full-width latin/digits (default: convert to half-width)")
    ctd.set_defaults(func=_cmd_convert_to_docx)

    # pipeline group
    pp = sub.add_parser("pipeline", help="multi-step + multi-PDF pipeline")
    pp_sub = pp.add_subparsers(dest="pipeline_cmd", metavar="<cmd>")
    pp.set_defaults(func=lambda a: (pp.print_help() or 0))

    pprun = pp_sub.add_parser(
        "run",
        help="run --steps across PDFs",
        description=(
            "Examples:\n"
            "  pdf_cli.py pipeline run report.pdf "
            "--steps text-extract,table-extract --report-dir reports/\n"
            "  pdf_cli.py pipeline run 'data/*.pdf' "
            "--steps image-extract --parallel --max-workers 4"
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    pprun.add_argument("pdfs", nargs="+",
                       help="PDF paths or globs (quote globs to defer "
                            "shell expansion)")
    pprun.add_argument("--steps", type=_parse_steps, required=True,
                       help="comma-separated step names: "
                            + ",".join(list_builtin_steps()))
    pprun.add_argument("--parallel", action="store_true",
                       help="cross-PDF process pool")
    pprun.add_argument("--max-workers", type=int, default=None,
                       help="parallel worker count "
                            "(default: min(N_pdf, cpu_count))")
    pprun.add_argument("--report-dir", default=None,
                       help="dir to write pipeline-<stem>.json per PDF")
    pprun.add_argument("--out-dir", default=None,
                       help="global out-dir root; "
                            "step uses <root>/<step-name>/")
    # per-step out_dir overrides
    pprun.add_argument("--image-extract-out-dir", default=None)
    pprun.add_argument("--text-extract-out-dir", default=None)
    pprun.add_argument("--table-extract-out-dir", default=None)
    # step-shared page selection
    pprun.add_argument("--pages", default=None,
                       help="page spec applied to all pdf-based steps "
                            "(text/table)")
    pprun.add_argument("--single", action="store_true",
                       help="text-extract: single full.txt")
    pprun.add_argument("--text-extract-pages", default=None)
    pprun.add_argument("--text-extract-single", action="store_true")
    pprun.add_argument("--table-extract-pages", default=None)
    pprun.set_defaults(func=_cmd_pipeline_run)

    return p


def main(argv: Optional[list[str]] = None) -> int:
    raw = list(argv) if argv is not None else sys.argv[1:]
    if not raw or raw[0] in ("-h", "--help"):
        _build_parser().print_help()
        return 0
    parser = _build_parser()
    try:
        args = parser.parse_args(raw)
    except SystemExit as se:
        return int(se.code) if isinstance(se.code, int) else 2

    func = getattr(args, "func", None)
    if func is None:
        parser.print_help()
        return 0
    try:
        rc = func(args)
        return int(rc) if isinstance(rc, int) else (0 if rc is None else 1)
    except SystemExit as se:
        return int(se.code) if isinstance(se.code, int) else 0
    except Exception as e:
        print(f"[pdf_cli.py] error: {type(e).__name__}: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
