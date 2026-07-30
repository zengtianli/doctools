#!/usr/bin/env python3
"""
文本格式自动修复工具 - DOCX版本
支持处理Word文档中的格式问题

功能列表：
1. 双引号格式修复
   将所有双引号统一为中文标准引号："和"
   奇数个引号 → " (中文左双引号 U+201C)
   偶数个引号 → " (中文右双引号 U+201D)

2. 英文标点符号转中文
   , → ，  (逗号)
   : → ：  (冒号)
   ; → ；  (分号)
   ! → ！  (感叹号)
   ? → ？  (问号)
   ( → （  (左括号)
   ) → ）  (右括号)

3. 中文单位转标准符号
   面积：平方米 → m²、平方公里 → km²
   体积：立方米 → m³、立方厘米 → cm³
   长度：公里 → km、厘米 → cm、毫米 → mm
   质量：公斤 → kg、毫克 → mg
   容量：毫升 → mL、微升 → μL
   时间：小时 → h、分钟 → min、秒钟 → s
   温度：摄氏度 → ℃、华氏度 → ℉

4. 单位上标格式化
   m2 → m²、m3 → m³、km2 → km²、km3 → km³

覆盖范围（2026-07-26 扩齐）：
   正文 + 表格（含嵌套表）+ 文本框 + 超链接 + **审阅修订（w:ins 插入态 / w:del 删除态）**
   + **批注 comments.xml** + **脚注 / 尾注** + 页眉页脚（默认纳入；--strip-headers 时改为删除）。
   原先走 python-docx 的 `Document.paragraphs` / `Paragraph.runs`，那两个 API 只认
   直接子节点，带修订的稿子里 w:ins/w:del 段一个字都改不到（实测某审阅稿 24 个 run /
   3101 字全漏），批注与脚注更是整个 part 没碰。现统一走 docx_xml 的元素级遍历。
   删除态文本改的是 w:delText（不退化成 w:t，修订结构原样保留）。

使用方法：
    python3 text_formatter_docx.py 文件名.docx
"""

import copy
import sys
from dataclasses import dataclass, replace
from pathlib import Path

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[2] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent.parent.parent / "lib"))
sys.path.insert(0, str(Path.home() / "Dev" / "tools" / "dev" / "lib"))  # canonical 5 modules
from docx_xml import (
    ALL_SCOPES,
    group_text,
    has_other_content,
    in_deleted,
    in_inserted,
    iter_paragraphs,
    iter_text_roots,
    para_own_runs,
    run_text,
    scope_of,
    set_group_text,
    set_run_text,
    text_groups,
    text_tag_for,
)
from file_ops import clear_quarantine
from finder import get_input_files
from progress import ProgressTracker
from text_fixes import fix_punctuation, fix_quotes, fix_units

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_write_gate import WriteGate  # 原地写回并发门（同目录 SSOT）

# ===== 配置对象 =====
# 旧实现把 5 个旋钮做成模块级可变全局,且只在 `if __name__ == "__main__"` 里赋值 ——
# 被 import 时永远拿默认值,同进程连跑两组配置会串味(第二次沿用上次残留)。
# 现在一律走冻结的 FormatConfig 显式穿参;模块级全局只作为「CLI 默认值」保留。


@dataclass(frozen=True)
class FormatConfig:
    """一次规范化的完整意图 = 规则轴 × 域轴 × 写回方式。"""

    # 规则轴
    quotes: bool = True            # 引号统一为中文弯引号
    punct: bool = True             # 英文标点 → 中文标点
    units: bool = True             # 中文单位 → 标准符号
    quote_font: bool = True        # 引号拆独立 run 并设宋体(排版动作,可单独关)
    # 域轴(空 = 全选);取值见 docx_xml.ALL_SCOPES
    scopes: frozenset = frozenset(ALL_SCOPES)
    # 破坏性 / 写回
    strip_headers: bool = False    # 删除页眉页脚引用(默认关,属排版动作)
    in_place: bool = False         # 原地写回 + .bak 备份(默认另存 _fixed)

    def wants(self, scope: str) -> bool:
        return scope in self.scopes

    @classmethod
    def from_opts(cls, opts: dict | None):
        """从 {"rule.quotes": "0", "scope.comments": "1", ...} 造配置(GUI 通路)。

        未知 key / 非法值抛 ValueError,由调用方翻成信封里的 ok:false。
        """
        cfg = cls()
        if not opts:
            return cfg
        rules, scopes = {}, set(cfg.scopes)
        for k, v in opts.items():
            on = _truthy(k, v)
            if k.startswith("rule."):
                name = k[5:]
                if name not in RULE_KEYS:
                    raise ValueError(f"未知规则:{name}(可用 {'/'.join(RULE_KEYS)})")
                rules[name] = on
            elif k.startswith("scope."):
                name = k[6:]
                if name not in ALL_SCOPES:
                    raise ValueError(f"未知范围:{name}(可用 {'/'.join(ALL_SCOPES)})")
                scopes.add(name) if on else scopes.discard(name)
            elif k in ("strip_headers", "in_place"):
                rules[k] = on
            else:
                raise ValueError(f"未知选项:{k}")
        return replace(cfg, scopes=frozenset(scopes), **rules)


RULE_KEYS = ("quotes", "punct", "units", "quote_font")


def _truthy(k: str, v) -> bool:
    if isinstance(v, bool):
        return v
    s = str(v).strip().lower()
    if s in ("1", "true", "yes", "on", "是"):
        return True
    if s in ("0", "false", "no", "off", "否"):
        return False
    raise ValueError(f"选项 {k} 的值 {v!r} 不是布尔(用 1/0)")


# ===== CLI 默认值 =====
# 删页眉页脚 = 排版动作,不属于「文本规范化」。2026-07-26 默认翻转为**保留**:
# 旧默认 True 是 initial commit 带进来的,从没人显式传过,却让每次「规范化」都把
# docx 的 headerReference/footerReference 全删光(实测 2+2 → 0),且 typeset 套完
# 院模板后第 2 步复用本引擎,把模板自带的 14/13 个页眉页脚引用一并删掉。
# 要删请显式 --strip-headers。
STRIP_HEADERS = False

# 选择性修复开关（默认全开=向后兼容；命令行 --quotes/--punct/--units 任一指定则只做指定项）
# 用途：中文期刊投稿(GB/T 7714 参考文献区标点须半角)只想修引号→ --quotes，不碰标点/单位
DO_QUOTES = True
DO_PUNCT = True
DO_UNITS = True
DO_QUOTE_FONT = True   # 引号拆独立 run 并设宋体（排版动作；--no-quote-font 可单关）
SCOPES = set(ALL_SCOPES)   # 默认全域：正文/表格/审阅修订/批注/脚注尾注/页眉页脚
IN_PLACE = False   # True 时原地写回源文件 + 备份 .bak-时间戳（Work §1.5 协议），不另存 _fixed


def _cli_config() -> "FormatConfig":
    """把模块级 CLI 默认值收成一个 FormatConfig（旧调用方零改动仍可用）。"""
    return FormatConfig(
        quotes=DO_QUOTES, punct=DO_PUNCT, units=DO_UNITS, quote_font=DO_QUOTE_FONT,
        scopes=frozenset(SCOPES), strip_headers=STRIP_HEADERS, in_place=IN_PLACE)


def _fmt_rules(cfg) -> str:
    on = [n for n, v in (("引号", cfg.quotes), ("标点", cfg.punct),
                         ("单位", cfg.units), ("引号宋体", cfg.quote_font)) if v]
    return "+".join(on) or "(无)"


def _fmt_scopes(cfg) -> str:
    return "+".join(ALL_SCOPES[k] for k in ALL_SCOPES if cfg.wants(k)) or "(无)"


QUOTE_CHARS = {"\u201c", "\u201d"}
QUOTE_FONT = "\u5b8b\u4f53"  # 宋体


def _set_run_font_songti(run_element):
    """为 run 的 rPr 设置宋体字体（ascii + hAnsi + eastAsia）"""
    rPr = run_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), QUOTE_FONT)
    rFonts.set(qn("w:hAnsi"), QUOTE_FONT)
    rFonts.set(qn("w:eastAsia"), QUOTE_FONT)
    rFonts.set(qn("w:hint"), "eastAsia")


def _split_text_at_quotes(text):
    """
    将文本在引号位置切片，返回 [(text, is_quote), ...] 片段列表。
    如果没有引号，返回 None。
    """
    if not text or not any(c in QUOTE_CHARS for c in text):
        return None

    segments = []
    buf = []
    for c in text:
        if c in QUOTE_CHARS:
            if buf:
                segments.append(("".join(buf), False))
                buf = []
            segments.append((c, True))
        else:
            buf.append(c)
    if buf:
        segments.append(("".join(buf), False))
    return segments


def _apply_quote_split(r, segments):
    """
    根据 segments 拆分 run 元素 r：第一段复用原 run，后续段插入新 run。
    引号段设置宋体。删除态 run（w:del 内）新建的文本元素仍是 w:delText。
    """
    parent = r.getparent()
    text_tag = text_tag_for(r)          # w:del / w:moveFrom 内必须继续用 w:delText

    # 第一段复用原 run
    first_text, first_is_quote = segments[0]
    set_run_text(r, first_text)
    if first_is_quote:
        _set_run_font_songti(r)

    # 后续段：复制原 run 的格式，插入到原 run 之后
    insert_after = r
    for seg_text, is_quote in segments[1:]:
        new_r = copy.deepcopy(r)
        # 只留 rPr：deepcopy 带来的旧文本、以及 drawing/footnoteReference 等
        # 一次性载荷都必须清掉，否则拆一次 run 就把图片/脚注引用复制一份出来
        for child in list(new_r):
            if child.tag != qn("w:rPr"):
                new_r.remove(child)
        t_elem = OxmlElement(text_tag)
        t_elem.text = seg_text
        # 保留空格
        t_elem.set(qn("xml:space"), "preserve")
        new_r.append(t_elem)

        if is_quote:
            _set_run_font_songti(new_r)
        else:
            # 非引号段恢复原 run 的字体（去掉宋体覆盖）
            rPr = new_r.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                orig_rPr = r.find(qn("w:rPr"))
                orig_rFonts = orig_rPr.find(qn("w:rFonts")) if orig_rPr is not None else None
                if rFonts is not None and orig_rFonts is not None:
                    # 用原始字体信息覆盖
                    rPr.replace(rFonts, copy.deepcopy(orig_rFonts))
                elif rFonts is not None and orig_rFonts is None:
                    rPr.remove(rFonts)

        parent.insert(list(parent).index(insert_after) + 1, new_r)
        insert_after = new_r


def _tally(stats, scope, r, n):
    """按域记账：修订态单独计，好让用户看见「审阅段确实改到了」。"""
    if not n:
        return
    key = scope
    if scope == "正文":
        if in_deleted(r):
            key = "修订·删除态"
        elif in_inserted(r):
            key = "修订·插入态"
    stats["scopes"][key] = stats["scopes"].get(key, 0) + n


def process_paragraph_element(p, stats, scope="正文", cfg=None):
    """
    处理一个 w:p 元素内的文本，逐 run 处理以保持每个 run 的字体格式。

    走 para_own_runs 而非 python-docx 的 `Paragraph.runs`：后者只认 `./w:r`，
    审阅修订（w:ins/w:del）、超链接、smartTag 里的 run 会被静默跳过。

    ⚠ 域过滤与引号规则**不正交**：引号左右方向靠段落级奇偶计数器，被跳过的 run
    如果不推进计数器，它后面的引号会整体反相（实测 `乙"结束` 从 `乙“` 翻成 `乙”`）。
    所以未选中的 run 仍然跑一遍 fix_quotes 推进 counter，只是**不写回**。
    """
    cfg = cfg or FormatConfig()
    original_runs = para_own_runs(p)      # 先快照:遍历中会插入新 run
    if not original_runs:
        return

    quote_counter = 0                     # 引号计数器跨 run 维护,保证配对正确

    for r in original_runs:
        # 正文 part 内按 run 归属的子域(正文/表格/审阅修订)过滤;其余 part 整体由
        # iter_text_roots 的 parts 决定,进到这里就是选中的
        selected = cfg.wants(scope_of(r)) if scope == "正文" else True

        # 逐「文本组」处理：w:tab / w:br 隔开的文本不合并，否则会被挤到末尾（对齐错位）
        groups = text_groups(r)
        for g in groups:
            original_text = group_text(g)
            if not original_text:
                continue

            fixed_text = original_text
            quote_count = punct_count = unit_count = 0
            if cfg.quotes:
                fixed_text, quote_count, quote_counter = fix_quotes(fixed_text, quote_counter)
            if not selected:
                continue                  # counter 已推进,但不改这段文字、不记账
            if cfg.punct:
                fixed_text, punct_count = fix_punctuation(fixed_text)
            if cfg.units:
                fixed_text, unit_count = fix_units(fixed_text)

            stats["quotes"] += quote_count
            stats["punctuation"] += punct_count
            stats["units"] += unit_count
            _tally(stats, scope, r, quote_count + punct_count + unit_count)

            if fixed_text != original_text:
                set_group_text(g, fixed_text)

        # 引号拆独立 run + 宋体（排版动作，可单独关）。
        # run 里夹着 tab/br/图/脚注引用时不拆 —— 复制这些一次性载荷会出复件
        if selected and cfg.quotes and cfg.quote_font and len(groups) == 1 and not has_other_content(r):
            segments = _split_text_at_quotes(run_text(r))
            if segments:
                _apply_quote_split(r, segments)


def process_root(root, stats, scope="正文", cfg=None):
    """处理一个 XML 根（正文 body / 批注 / 脚注 / 页眉…）下的全部段落。

    `iter_paragraphs` 递归到嵌套表格、文本框、修订块里的 w:p；`para_own_runs`
    在嵌套 w:p 处剪枝，所以每个 run 恰好被处理一次。
    """
    for p in list(iter_paragraphs(root)):
        process_paragraph_element(p, stats, scope, cfg)


def _strip_header_footer_refs(doc) -> int:
    """删掉所有 section 的 headerReference/footerReference，返回删掉的引用数。

    删「引用」而不是 header/footer part 本身：Word 见不到引用就按无页眉页脚渲染，
    留在包里的孤儿 part 无害；去删 part 会连带 rels 一起崩。
    """
    removed = 0
    for section in doc.sections:
        sectPr = section._sectPr
        for ref in (sectPr.findall(qn("w:headerReference"))
                    + sectPr.findall(qn("w:footerReference"))):
            sectPr.remove(ref)
            removed += 1
    return removed


def _normalize_doc(doc, cfg: FormatConfig) -> dict:
    """对**已打开的** doc 做完整规范化，返回统计。不开文件、不存盘、不打印。

    process_docx（CLI/GUI 通路，手上已经是 FormatConfig）与 apply（pipeline 通路，
    手上是 argparse.Namespace）共用这一份实现 —— 两边各抄一遍这个循环，正是
    「脚注 flush 漏一处」「新加的 part 只接进了一条通路」这类 bug 的温床。
    """
    stats = {"quotes": 0, "punctuation": 0, "units": 0, "scopes": {}}

    # 处理所有承载文本的 part：正文（含表格/嵌套表/文本框/超链接/审阅修订）
    # + 批注 comments.xml + 脚注 + 尾注 + 页眉页脚（默认保留并规范化）
    flushes = []
    parts = {k for k in ("comments", "notes", "headers") if cfg.wants(k)}
    for label, root, flush in iter_text_roots(
            doc, include_headers=not cfg.strip_headers, parts=parts):
        process_root(root, stats, label, cfg)
        if flush is not None:
            flushes.append(flush)
    for flush in flushes:          # 通用 Part（脚注/尾注）改完必须回写 blob
        flush()

    # 显式要求时才删页眉页脚（默认保留）
    stats["headers_stripped"] = _strip_header_footer_refs(doc) if cfg.strip_headers else 0
    stats["changed"] = (stats["quotes"] + stats["punctuation"] + stats["units"]
                        + stats["headers_stripped"])
    return stats


def _config_from_args(args=None) -> FormatConfig:
    """argparse.Namespace → FormatConfig（pipeline 通路）。

    逐项 getattr 带默认：pipeline 里同一个 Namespace 要喂给一串 step，缺哪项就该
    沿用本脚本的默认值，而不是 AttributeError 掀桌。dict 通路是 GUI 专用的
    FormatConfig.from_opts（键名带 rule./scope. 前缀），两套别混。
    """
    base = _cli_config()
    if args is None:
        return base
    ready = getattr(args, "format_config", None)
    if isinstance(ready, FormatConfig):     # 调用方已算好完整意图，直接用
        return ready
    scopes = getattr(args, "scopes", None)
    return replace(
        base,
        quotes=bool(getattr(args, "quotes", base.quotes)),
        punct=bool(getattr(args, "punct", base.punct)),
        units=bool(getattr(args, "units", base.units)),
        quote_font=bool(getattr(args, "quote_font", base.quote_font)),
        scopes=frozenset(scopes) if scopes else base.scopes,
        strip_headers=bool(getattr(args, "strip_headers", base.strip_headers)),
    )


# ---------------- pipeline adapter ----------------
def apply(doc, args=None) -> dict:
    """在**已打开的** doc 上做文本规范化（引号/标点/单位 + 引号宋体），返回统计。

    覆盖面照旧走 docx_xml 的 iter_text_roots：批注/脚注/尾注/页眉页脚是独立 part，
    审阅修订（w:ins/w:del）又不是 w:p 的直接子 run —— 只遍历 doc.paragraphs 会
    整块漏掉（实测某审阅稿 24 个 run / 3101 字一个都改不到）。
    """
    return _normalize_doc(doc, _config_from_args(args))


def process_docx(input_file, cfg=None):
    """处理 DOCX 文件。cfg=None 时用模块级 CLI 默认值(向后兼容既有调用方)。"""
    cfg = cfg or _cli_config()
    input_path = Path(input_file)

    if not input_path.exists():
        print(f"❌ 错误：文件不存在 - {input_file}")
        return False

    if input_path.suffix.lower() != ".docx":
        print("❌ 错误：文件必须是.docx格式")
        return False

    # 生成输出文件名（--in-place 时原地写回 + 备份；否则另存 _fixed）
    write_gate = WriteGate(input_path) if cfg.in_place else None  # 读入前 capture 基线
    if cfg.in_place:
        import shutil as _shutil
        from datetime import datetime as _dt
        bak = input_path.with_name(
            input_path.name + ".bak-" + _dt.now().strftime("%Y%m%d-%H%M%S"))
        _shutil.copy2(input_path, bak)
        print(f"🗄  已备份: {bak.name}")
        output_path = input_path
    else:
        output_path = input_path.parent / f"{input_path.stem}_fixed{input_path.suffix}"

    try:
        # 读取文档
        print(f"📖 正在读取文件: {input_path.name}")
        doc = Document(input_path)

        # 改 doc 的那件事整体在 _normalize_doc 里（统计信息由它返回，这里只负责喊）
        print(f"🔄 规则: {_fmt_rules(cfg)} | 范围: {_fmt_scopes(cfg)}")
        if cfg.strip_headers:
            print("🗑️  正在删除页眉页脚（--strip-headers）...")
        stats = _normalize_doc(doc, cfg)

        # 保存文件
        print("💾 正在保存文件...")
        if write_gate is not None:
            write_gate.assert_unchanged()  # 源文件被 WPS/其他会话改过 → 拒写(逃生 DOCX_GATE_OK=1)
        try:
            doc.save(output_path)
            clear_quarantine(output_path)
        except (PermissionError, OSError) as e:
            fallback = Path.home() / "Downloads" / Path(output_path).name
            print(f"⚠️ 源目录不可写: {e}")
            print(f"   降级到: {fallback}")
            doc.save(fallback)
            clear_quarantine(fallback)
            output_path = fallback

        print("✅ 处理完成！")
        print(f"   - 共替换了 {stats['quotes']} 个引号")
        print(f"   - 共替换了 {stats['punctuation']} 个标点符号")
        print(f"   - 共转换了 {stats['units']} 个单位")
        scopes = stats["scopes"] or {"正文": 0}
        print("   - 分域命中: " + "  ".join(f"{k} {v} 处" for k, v in scopes.items()))
        print(f"   - 输出文件: {output_path.name}")

        return True

    except Exception as e:
        print(f"❌ 处理失败：{e}")
        import traceback

        traceback.print_exc()
        return False


if __name__ == "__main__":
    # 摘选择性 flag（在 get_input_files 前），剩余为文件参数
    _argv = sys.argv[1:]
    _flags = {"--quotes", "--punct", "--units", "--no-quote-font",
              "--in-place", "--keep-headers", "--strip-headers", "--strip-only"}
    _sel = {a for a in _argv if a in _flags}

    # --scope a,b,c（域白名单；不给 = 全域）
    _scope_arg = None
    _rest = []
    _it = iter(_argv)
    for a in _it:
        if a == "--scope":
            _scope_arg = next(_it, "")
        elif a.startswith("--scope="):
            _scope_arg = a.split("=", 1)[1]
        else:
            _rest.append(a)
    _argv = _rest

    _unknown = [a for a in _argv if a.startswith("--") and a not in _flags]
    if _unknown:   # 未知 flag 不得被当成文件名静默吞掉→曾致 --help 直接跑全量改写
        print(f"❌ 未知参数：{' '.join(_unknown)}")
        print(f"   可用：{' '.join(sorted(_flags))} --scope <{'|'.join(ALL_SCOPES)}>")
        sys.exit(2)
    _argv = [a for a in _argv if a not in _flags]

    if _scope_arg is not None:
        SCOPES = {s.strip() for s in _scope_arg.split(",") if s.strip()}
        _bad = SCOPES - set(ALL_SCOPES)
        if _bad:
            print(f"❌ 未知范围：{' '.join(sorted(_bad))}")
            print(f"   可用：{' '.join(f'{k}({v})' for k, v in ALL_SCOPES.items())}")
            sys.exit(2)

    if _sel & {"--quotes", "--punct", "--units"}:
        DO_QUOTES = "--quotes" in _sel
        DO_PUNCT = "--punct" in _sel
        DO_UNITS = "--units" in _sel
    DO_QUOTE_FONT = "--no-quote-font" not in _sel
    IN_PLACE = "--in-place" in _sel
    STRIP_HEADERS = "--strip-headers" in _sel   # --keep-headers 已是默认行为,保留为兼容 no-op
    if "--strip-only" in _sel:
        # 「清页眉页脚」独立动词:只删 chrome,一个字都不改
        STRIP_HEADERS = True
        DO_QUOTES = DO_PUNCT = DO_UNITS = DO_QUOTE_FONT = False

    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(_argv, expected_ext="docx")

    if not files:
        print("❌ 错误：缺少文件名参数")
        print("\n使用方法：")
        print("  1. 在 Finder 中选中 .docx 文件，然后运行此脚本")
        print("  2. 或在命令行中提供文件路径")
        print("\n示例：")
        print("    python3 docx_text_formatter.py 文件名.docx")
        print("    python3 docx_text_formatter.py file1.docx file2.docx")
        sys.exit(1)

    print("=" * 50)
    print("文本格式自动修复工具 - DOCX版本")
    print("=" * 50)

    tracker = ProgressTracker()

    for file_path in files:
        print(f"\n处理文件: {Path(file_path).name}")
        success = process_docx(str(file_path))
        if success:
            tracker.add_success()
        else:
            tracker.add_error()

    print("\n" + "=" * 50)
    tracker.show_summary("文件处理")
