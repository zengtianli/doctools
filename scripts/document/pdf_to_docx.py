#!/usr/bin/env python3
"""pdf_to_docx.py — PDF → 可编辑 Word（结构提取路线，2026-07-20 立）

设计取向（用户 2026-07-20 拍板）：**要「干净可编辑」不要「像素级复刻」**。

  pdfplumber ──抽 text lines + tables──┐
                                        ├─→ python-docx ─→ .docx
  (可选) pdfimages ──抽嵌入图───────────┘

保留：段落顺序 / 标题层级（按字号推断）/ 列表项 / **PDF 表格 → 真 Word 表格** / 可选嵌图。
丢弃：原版式（字体族·分栏·配色·图文环绕）—— 要那个只有 pdf2docx(PyMuPDF)，
      而本族已立红线「不引入 pymupdf」（AGPL 传染，见 cc-home/commands/pdf.md:40）。

为什么不走 markitdown→md→md2word：md2word 会强制套「院公文模板」样式，
产出是公文风而非原文档结构；且复杂表格经 md 管道表中转会塌。

依赖：pdfplumber 0.11.9 / python-docx 1.2.0（**装在 /opt/homebrew/bin/python3**，
不在 ~/Dev/.venv —— 调用方必须用绝对路径解释器，同 pdf_cli.py）。

独立用法：
  /opt/homebrew/bin/python3 pdf_to_docx.py <in.pdf> [-o out.docx] [--images] [--no-tables]
经 pdf_cli：
  pdf_cli.py convert to-docx <in.pdf> [-o out.docx] [--images]
"""
from __future__ import annotations

import argparse
import re
import shutil
import statistics
import subprocess
import sys
import tempfile
from pathlib import Path

import logging

# pdfminer 对字体描述符不规范的 PDF 会每字符刷一条
# "Could not get FontBBox from font descriptor" —— 纯噪音，且会淹没我们自己的
# 警告、污染 GUI 后端的 stderr。它不影响提取结果，直接闭嘴。
logging.getLogger("pdfminer").setLevel(logging.ERROR)

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

# ───────────────────────────────────────────── 版面常量
#
# 这些阈值决定「这行是标题还是正文」「这两行要不要接成一段」。都按相对量
# （相对正文字号中位数 / 页宽百分比）算，绝对 pt 值在不同 PDF 上没有可比性。

H1_RATIO = 1.45      # 字号 ≥ 正文中位数 × 此值 → 一级标题
H2_RATIO = 1.22      # 同上 → 二级标题
H3_RATIO = 1.08      # 同上 + 加粗/短行 → 三级标题
SHORT_LINE = 0.72    # 行宽 < 页面文字区宽度 × 此值 = "短行"（标题候选/段落末行）
GAP_NEW_PARA = 0.62  # 行间距 > 行高 × 此值 → 强制断段
INDENT_TOL = 3.0     # x0 相差在此 pt 内视为「同一左边界」

# 行末出现这些 → 这一行是段落终点，下一行必定另起段
_END_PUNCT = "。！？；!?;:：」』”》】…"
# 行首出现这些 → 这一行必定另起段（列表/编号/条款）
_LIST_HEAD = re.compile(
    r"^\s*(?:[•·▪◦‣▶►■□●○–—\-\*]\s+"           # 项目符号
    r"|\(?\d{1,3}[\.\)、]\s*"                     # 1. / 1) / (1) / 1、
    r"|\(?[a-zA-Z][\.\)]\s+"                      # a. / (b)
    r"|[（(][一二三四五六七八九十百]+[）)]"       # （一）
    r"|第[一二三四五六七八九十百千0-9]+[章节条款项]"  # 第三章 / 第5条
    r")"
)
_BULLET_HEAD = re.compile(r"^\s*[•·▪◦‣▶►■□●○–—]\s+")
# 西文之间接续要补空格；中日韩之间不补。
# 注意 U+FF00–FFEF 这个「全角块」里混着两类东西，**不能整块算 CJK**：
#   全角标点（，。；：（）＂）→ 算 CJK，两侧不补空格
#   全角字母数字（Ａ-Ｚ ａ-ｚ ０-９）→ 是西文，必须补空格
# 整块算 CJK 会拼出「WaterResourcesandPower」（实测踩过）。
_CJK = re.compile(
    r"[　-〿"      # CJK 标点
    r"㐀-鿿"       # 汉字
    r"豈-﫿"       # 兼容汉字
    r"！-／"       # 全角标点 ！＂＃…／
    r"：-＠"       # ：；＜＝＞？＠
    r"［-｀"       # ［＼］＾＿｀
    r"｛-･"       # ｛｜｝～、。
    r"]"
)


def _is_cjk(ch: str) -> bool:
    return bool(_CJK.match(ch))


# ───────────────────────────────────────────── 抽取层


def _word_size(words: list[dict]) -> float:
    """一组词的代表字号 = 各词字号中位数（避免上下标/脚注号拉偏）。"""
    sizes = [w["size"] for w in words if w.get("size")]
    return statistics.median(sizes) if sizes else 0.0


def _word_bold(words: list[dict]) -> bool:
    """半数以上的词字体名带 Bold/Black/Heavy → 视为加粗行。"""
    names = [str(w.get("fontname", "")) for w in words]
    if not names:
        return False
    hits = sum(1 for n in names if re.search(r"bold|black|heavy|semib", n, re.I))
    return hits * 2 > len(names)


def _join_words(words: list[dict]) -> str:
    """词序列 → 行文本。CJK 相邻不加空格，其余加一个空格。

    **必须在词级做，不能在 char 级做**：char 级要靠「间隙 > k×字宽」猜词边界，
    而学术期刊/两端对齐会把西文字距整体拉大，任何阈值都会被骗
    （实测踩过：「Vol.44 No.2」→「V o l . 4 4 N o . 2」、
    「Mann-Kendall」→「M a n n - K e n d a l l」）。
    pdfplumber 的 extract_words 已用 PDF 内部的字距/字体信息切好词，直接信它。"""
    parts: list[str] = []
    for i, w in enumerate(words):
        t = w.get("text", "")
        if i and parts:
            prev = parts[-1]
            if prev and t and not (_is_cjk(prev[-1]) or _is_cjk(t[0])):
                parts.append(" ")
        parts.append(t)
    return "".join(parts)


# 全角字母/数字 → 半角。**只转这些**，中文标点（，。；：（）「」）一律保留 ——
# 用 unicodedata.normalize("NFKC") 是错的，它会把「，」也变成「,」，毁掉中文排版。
_FW_ALNUM = {c: chr(c - 0xFEE0) for c in
             [*range(0xFF21, 0xFF3B), *range(0xFF41, 0xFF5B), *range(0xFF10, 0xFF1A)]}
# 全角句点/连字符：仅当夹在 ASCII 之间才转（Ｖｏｌ．４４ → Vol.44；中文的「．」少见）
_FW_MID = re.compile(r"(?<=[0-9A-Za-z])[．－](?=[0-9A-Za-z])")
_IDEO_SPACE_CJK = re.compile(r"(?<=[一-鿿])　+(?=[一-鿿])")


def _normalize_width(s: str) -> str:
    """全角西文规范化。学术期刊 PDF 大量用全角 ASCII 排页眉/西文
    （实测：'Ｖｏｌ．４４Ｎｏ．２'、'Ｗａｔｅｒ Ｒｅｓｏｕｒｃｅｓ'），
    原样进 Word 会搜不到、难编辑 —— 但这是**内容规范化**不是提取错误，
    故给 --keep-fullwidth 逃生。"""
    s = _IDEO_SPACE_CJK.sub("", s)          # 汉字间的表意空格 = 疏排填充，删掉
    s = s.replace("　", " ")            # 其余表意空格 → 普通空格
    s = s.translate(_FW_ALNUM)
    s = _FW_MID.sub(lambda m: chr(ord(m.group()) - 0xFEE0), s)
    return re.sub(r" {2,}", " ", s)


def _group_lines(words: list[dict]) -> list[list[dict]]:
    """词 → 行。按垂直重叠聚类（同一行的词 top 基本相同，容差取字高的 60%）。"""
    if not words:
        return []
    ws = sorted(words, key=lambda w: (w["top"], w["x0"]))
    lines: list[list[dict]] = [[ws[0]]]
    for w in ws[1:]:
        ref = lines[-1][0]
        tol = max((ref["bottom"] - ref["top"]) * 0.6, 2.0)
        if abs(w["top"] - ref["top"]) <= tol:
            lines[-1].append(w)
        else:
            lines.append([w])
    for ln in lines:
        ln.sort(key=lambda w: w["x0"])
    return lines


def _split_columns(words: list[dict]) -> list[list[dict]]:
    """把一行的词序列按「行内大空隙」切成多栏。

    为什么必须做：按 y 聚行会把左右分栏（简历头部「姓名 | 联系方式」、双栏正文）
    拍平成同一行 → 拼出「曾田力 zengtianli1@126.com」这种鬼话。

    判据不能只看绝对间隙——两端对齐会把整行词距拉大。故再加相对判据：
    真正的分栏留白必须**显著大于同一行内其它所有词距**（一枝独秀 vs 普遍偏大）。"""
    if len(words) < 2:
        return [words]
    sizes = [w["size"] for w in words if w.get("size")] or [10.0]
    em = statistics.median(sizes)

    gaps = [cur["x0"] - prev["x1"] for prev, cur in zip(words, words[1:])]
    abs_thresh = max(em * 2.2, 18.0)
    if not any(g > abs_thresh for g in gaps):
        return [words]
    others = [g for g in gaps if g <= abs_thresh]
    cut = max(abs_thresh, (max(others) if others else 0.0) * 1.8)

    groups: list[list[dict]] = [[words[0]]]
    for prev, cur in zip(words, words[1:]):
        if cur["x0"] - prev["x1"] > cut:
            groups.append([cur])
        else:
            groups[-1].append(cur)
    return groups


PAGE_FRAME_RATIO = 0.85   # 表格 bbox 占页面面积超过此比例 → 是版面边框不是表格
MIN_TABLE_FILL = 0.30     # 单元格非空率低于此 → 是排版网格不是数据表


def _real_tables(page) -> list:
    """页内**真正的数据表**。

    pdfplumber 的 find_tables 会把「整页边框」「排版分隔线围出的区域」也当表格 ——
    实测 IIQE 讲义（PPT 导出，带页框）返回一个 594×841 覆盖全页的 bbox，
    结果 1219 行正文全被当成表格内容吞掉，产出空 docx。故必须过滤：
      ① bbox 不能几乎等于整页（那是页框）
      ② 至少 2 行 × 2 列
      ③ 单元格非空率够高（排版网格大多是空的）"""
    try:
        cands = page.find_tables()
    except Exception:
        return []
    page_area = float(page.width) * float(page.height)
    out = []
    for t in cands:
        x0, y0, x1, y1 = t.bbox
        if page_area and (x1 - x0) * (y1 - y0) > page_area * PAGE_FRAME_RATIO:
            continue                                   # 页框
        try:
            data = t.extract()
        except Exception:
            continue
        if not data or len(data) < 2:
            continue
        ncols = max((len(r) for r in data), default=0)
        if ncols < 2:
            continue
        cells = sum(len(r) for r in data)
        filled = sum(1 for r in data for c in r if c and str(c).strip())
        if not cells or filled / cells < MIN_TABLE_FILL:
            continue                                   # 排版网格
        out.append((t, data))
    return out


def _in_any_bbox(w: dict, boxes: list[tuple[float, float, float, float]]) -> bool:
    """词的中心落在某表格 bbox 内 → 属于该表格。用中心点而非全包含，
    容忍 find_tables 的边界比实际文字略紧。"""
    cy = (w["top"] + w["bottom"]) / 2
    cx = (w["x0"] + w["x1"]) / 2
    return any(x0 <= cx <= x1 and y0 <= cy <= y1 for x0, y0, x1, y1 in boxes)


def _collect_blocks(pdf, drop_tables: bool, norm_width: bool = True) -> list[dict]:
    """把整个 PDF 拍平成有序 block 流：{kind: text|table|pagebreak, ...}。

    表格按它在页内的垂直位置插进正文流（不是一股脑堆到页尾），
    这样 Word 里的图文顺序才和原 PDF 一致。"""
    blocks: list[dict] = []
    for pno, page in enumerate(pdf.pages):
        if pno:
            blocks.append({"kind": "pagebreak"})

        found = [] if drop_tables else _real_tables(page)
        boxes = [t.bbox for t, _ in found]
        tables = [{"kind": "table", "top": t.bbox[1], "data": d} for t, d in found]
        if norm_width:
            for tb in tables:
                tb["data"] = [[(_normalize_width(c) if isinstance(c, str) else c)
                               for c in row] for row in tb["data"]]

        try:
            words = page.extract_words(extra_attrs=["size", "fontname"]) or []
        except Exception as e:
            print(f"  ⚠ 第 {pno + 1} 页取词失败({type(e).__name__})，跳过", file=sys.stderr)
            words = []
        if boxes:
            words = [w for w in words if not _in_any_bbox(w, boxes)]  # 表格内容不重复进正文

        page_items: list[dict] = list(tables)
        for ln_words in _group_lines(words):
            for seg in _split_columns(ln_words):
                txt = _join_words(seg).strip()
                if not txt:
                    continue
                if norm_width:
                    txt = _normalize_width(txt)
                    if not txt:
                        continue
                page_items.append({
                    "kind": "text",
                    "top": min(w["top"] for w in seg),
                    "bottom": max(w["bottom"] for w in seg),
                    "x0": min(w["x0"] for w in seg),
                    "x1": max(w["x1"] for w in seg),
                    "text": txt,
                    "size": _word_size(seg),
                    "bold": _word_bold(seg),
                })

        # 同一 y 带内先左后右（分栏切出来的右段排在左段之后）
        page_items.sort(key=lambda b: (round(b["top"], 1), b.get("x0", 0)))
        # 页面文字区宽度 —— 判「短行」的分母，用实际内容跨度而非纸张宽度
        xs = [b["x0"] for b in page_items if b["kind"] == "text"]
        xe = [b["x1"] for b in page_items if b["kind"] == "text"]
        span = (max(xe) - min(xs)) if xs and xe else float(page.width)
        for b in page_items:
            if b["kind"] == "text":
                b["page_span"] = span
                b["page_left"] = min(xs) if xs else 0.0
        blocks.extend(page_items)
    return blocks


# ───────────────────────────────────────────── 段落重组
#
# PDF 里「一行」不等于「一段」。直接一行一段，Word 里会碎成几百个短段落，
# 换行全是硬断，改一个字整段错位 —— 这是 PDF 转 Word 最劝退的地方，必须重组。


def _body_size(blocks: list[dict]) -> float:
    """正文字号 = 全文 text 行字号的中位数（标题是少数派，拉不动中位数）。"""
    sizes = [b["size"] for b in blocks if b["kind"] == "text" and b["size"]]
    return statistics.median(sizes) if sizes else 10.0


def _heading_level(b: dict, body: float) -> int | None:
    """返回 1/2/3 = 标题级别，None = 正文。

    判据 = 字号为主、加粗+短行为辅。纯加粗但字号=正文的不算标题（那多半是行内强调）。"""
    if not body:
        return None
    r = b["size"] / body
    short = (b["x1"] - b["x0"]) < b.get("page_span", 1e9) * SHORT_LINE
    if r >= H1_RATIO:
        return 1
    if r >= H2_RATIO:
        return 2
    if r >= H3_RATIO and b["bold"] and short:
        return 3
    return None


def _should_join(prev: dict, cur: dict, body: float, para: dict) -> bool:
    """prev 和 cur 是不是同一个自然段的连续两行？para = 正在累积的段落（提供首行上下文）。"""
    if _LIST_HEAD.match(cur["text"]):
        return False                                    # 新列表项/编号 → 断
    if prev["text"] and prev["text"][-1] in _END_PUNCT:
        return False                                    # 上行已收句 → 断
    gap = cur["top"] - prev["bottom"]
    line_h = max(prev["bottom"] - prev["top"], 1.0)
    if gap > line_h * GAP_NEW_PARA:
        return False                                    # 行距明显变大 → 断
    if abs(cur["size"] - prev["size"]) > 0.6:
        return False                                    # 字号跳变 → 不是同段

    # 列表项的续行是「悬挂缩进」：比 bullet 首行更靠右。这条必须**最先**判，
    # 排在「左边界跳变」「短行=段末」之前——悬挂缩进本身就是一次左边界跳变，
    # 后者会把一个 bullet 腰斩成两段（实测：「…31 个 hook（PreToolUse…」x0=46.3
    # ／「强制守卫）、3 个专用 subagent」x0=59.6，差 13.3pt 正好踩线）。
    if para.get("bullet"):
        return cur["x0"] >= para["first_x0"] - INDENT_TOL

    if abs(cur["x0"] - prev["x0"]) > INDENT_TOL * 4:
        return False                                    # 左边界大幅跳变 → 多半是另一栏
    if (prev["x1"] - prev["x0"]) < prev.get("page_span", 1e9) * SHORT_LINE:
        return False                                    # 上行是短行 = 段落末行 → 断
    if cur["x0"] > prev["x0"] + INDENT_TOL * 2:
        return False                                    # 本行明显右缩进 → 新段首行
    return True


def _join_text(a: str, b: str) -> str:
    """接续两行。中文之间直接拼；西文之间补空格；连字符断词去掉连字符。"""
    if not a:
        return b
    if a.endswith("-") and b[:1].isalpha():
        return a[:-1] + b
    if _is_cjk(a[-1]) or _is_cjk(b[:1]):
        return a + b
    return a + " " + b


def _build_paragraphs(blocks: list[dict], body: float) -> list[dict]:
    """block 流 → 段落流。相邻 text 行按 _should_join 归并；表格/分页原样穿过。"""
    out: list[dict] = []
    cur: dict | None = None

    def flush():
        nonlocal cur
        if cur:
            out.append(cur)
            cur = None

    for b in blocks:
        if b["kind"] != "text":
            flush()
            out.append(b)
            continue

        lvl = _heading_level(b, body)
        if lvl:
            flush()
            out.append({"kind": "heading", "level": lvl, "text": b["text"]})
            continue

        if cur and _should_join(cur["_last"], b, body, cur):
            cur["text"] = _join_text(cur["text"], b["text"])
            cur["_last"] = b
        else:
            flush()
            cur = {
                "kind": "para",
                "text": b["text"],
                "bullet": bool(_LIST_HEAD.match(b["text"])),
                "first_x0": b["x0"],
                "_last": b,
            }
    flush()
    return out


# ───────────────────────────────────────────── 图片（可选）


def _extract_images(pdf_path: Path, work: Path) -> list[Path]:
    """pdfimages -all 抽嵌入图。按铁律 #6 过滤噪音：< 3KB 的多半是 soft-mask/纯色分隔条。"""
    exe = shutil.which("pdfimages") or "/opt/homebrew/bin/pdfimages"
    if not Path(exe).exists():
        return []
    try:
        subprocess.run([exe, "-all", str(pdf_path), str(work / "img")],
                       check=True, capture_output=True)
    except subprocess.CalledProcessError:
        return []
    keep = []
    for p in sorted(work.glob("img-*")):
        if p.suffix.lower() in (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".ppm", ".pbm"):
            if p.stat().st_size >= 3072:
                keep.append(p)
    return keep


# ───────────────────────────────────────────── 渲染层


def _set_base_style(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = "PingFang SC"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.15


def _add_table(doc: Document, data: list[list]) -> None:
    rows = [r for r in data if r and any((c or "").strip() for c in r)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            val = (row[j] if j < len(row) else "") or ""
            cell = t.cell(i, j)
            # PDF 单元格里的软换行是排版产物，不是语义换行 → 压成空格
            cell.text = re.sub(r"\s*\n\s*", " ", str(val)).strip()
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True
    doc.add_paragraph()


def convert(
    pdf_path: Path,
    out_path: Path | None = None,
    with_images: bool = False,
    with_tables: bool = True,
    norm_width: bool = True,
) -> dict:
    """PDF → docx。返回 {ok, output, pages, paragraphs, tables, images, error?}"""
    pdf_path = Path(pdf_path).expanduser().resolve()
    if not pdf_path.exists():
        return {"ok": False, "error": f"文件不存在: {pdf_path}"}
    out = Path(out_path).expanduser() if out_path else pdf_path.with_suffix(".docx")

    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            npages = len(pdf.pages)
            blocks = _collect_blocks(pdf, drop_tables=not with_tables, norm_width=norm_width)
            # 兜底：表格识别把正文全吞了（_real_tables 的过滤没覆盖到的诡异版面）
            # → 关掉表格重来一遍。宁可丢表格结构，不可产出空文档。
            if with_tables and not any(b["kind"] == "text" for b in blocks):
                if any(b["kind"] == "table" for b in blocks):
                    print("  ⚠ 表格识别吞掉了全部正文，已关闭表格识别重试", file=sys.stderr)
                    blocks = _collect_blocks(pdf, drop_tables=True, norm_width=norm_width)
    except Exception as e:  # 加密/损坏/非 PDF
        return {"ok": False, "error": f"打开 PDF 失败: {type(e).__name__}: {e}"}

    if not any(b["kind"] == "text" for b in blocks):
        return {"ok": False,
                "error": "PDF 里没有可提取的文字层（多半是扫描件）——"
                         "先跑 ocrmypdf 加文本层，或用 vision_ocr.py 识别"}

    body = _body_size(blocks)
    items = _build_paragraphs(blocks, body)

    doc = Document()
    _set_base_style(doc)
    n_tab = 0
    for it in items:
        k = it["kind"]
        if k == "heading":
            doc.add_heading(it["text"], level=min(it["level"], 4))
        elif k == "para":
            style = "List Bullet" if it["bullet"] and _BULLET_HEAD.match(it["text"]) else None
            txt = _BULLET_HEAD.sub("", it["text"]) if style else it["text"]
            doc.add_paragraph(txt, style=style)
        elif k == "table":
            _add_table(doc, it["data"])
            n_tab += 1
        elif k == "pagebreak":
            doc.add_paragraph()

    n_img = 0
    if with_images:
        with tempfile.TemporaryDirectory() as td:
            imgs = _extract_images(pdf_path, Path(td))
            if imgs:
                doc.add_page_break()
                h = doc.add_paragraph("附：文档嵌入图")
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for p in imgs:
                    try:
                        doc.add_picture(str(p), width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        n_img += 1
                    except Exception:
                        continue  # 有些 ppm/掩码 python-docx 吃不进，跳过不中断

    try:
        doc.save(str(out))
    except Exception as e:
        return {"ok": False, "error": f"写 docx 失败: {type(e).__name__}: {e}"}

    # macOS quarantine：不清的话 Word 开出来是「受保护视图」，审阅按钮全灰
    # （memory feedback-docx-macos-quarantine）
    subprocess.run(["xattr", "-d", "com.apple.quarantine", str(out)],
                   capture_output=True)

    return {
        "ok": True,
        "output": str(out),
        "pages": npages,
        "paragraphs": sum(1 for i in items if i["kind"] in ("para", "heading")),
        "tables": n_tab,
        "images": n_img,
    }


# ───────────────────────────────────────────── CLI


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="PDF → 可编辑 Word（结构提取：段落重组 + 真 Word 表格）")
    ap.add_argument("pdf", help="源 PDF")
    ap.add_argument("-o", "--out", default=None, help="产出 docx（默认同名同目录）")
    ap.add_argument("--images", action="store_true", help="附带抽取嵌入图（追加到文末）")
    ap.add_argument("--no-tables", action="store_true", help="不识别表格，全部当正文处理")
    ap.add_argument("--keep-fullwidth", action="store_true",
                    help="保留全角西文/数字（默认转半角，中文标点始终保留）")
    a = ap.parse_args(argv)

    r = convert(Path(a.pdf), Path(a.out) if a.out else None,
                with_images=a.images, with_tables=not a.no_tables,
                norm_width=not a.keep_fullwidth)
    if not r["ok"]:
        print(f"✖ {r['error']}", file=sys.stderr)
        return 1
    bits = [f"{r['pages']} 页 → {r['paragraphs']} 段"]
    if r["tables"]:
        bits.append(f"{r['tables']} 表")
    if r["images"]:
        bits.append(f"{r['images']} 图")
    print(f"✓ {Path(r['output']).name}（{' / '.join(bits)}）")
    return 0


if __name__ == "__main__":
    sys.exit(main())
