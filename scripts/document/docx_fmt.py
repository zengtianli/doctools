#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""docx_fmt.py — docx 版式/字体/文本规范化族（原 4 个入口脚本 2026-07-31 合并为子命令）。

子命令（CLI 表面与原脚本零改动，banner/usage/报错文案/退出码逐字保留）:
  template [apply|cleanup] ...   原 docx_apply_template.py（套模板样式 + 样式清理；
                                 无二级子命令时兼容旧调用 = 默认 apply）
  clone    {extract|apply} ...   原 docx_format_clone.py（提取/复刻版式，段落外壳克隆）
  fonts    <docx...> --check|--apply   原 docx_font_normalize.py（去等线，中文字体归一）
  text     [flags] <docx...>     原 docx_text_formatter.py（引号/标点/单位规范化，
                                 覆盖批注/脚注/尾注/审阅修订/页眉页脚）

exit 契约不变: fonts 1=有等线风险 · text 2=未知 flag/未知范围、1=缺文件 · 其余同原脚本。
text 的旧实现把 5 个旋钮做成模块级可变全局且只在 __main__ 赋值 —— 合并后 text_main
每次调用从头构造冻结 FormatConfig 显式穿参，同进程连跑两组配置不串味。
"""
from __future__ import annotations

import argparse
import copy
import datetime as _dt
import json
import os
import re
import shutil
import sys
import tempfile
import zipfile
from copy import deepcopy
from dataclasses import dataclass, replace
from io import BytesIO
from pathlib import Path

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[2] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py

from docx import Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

sys.path.insert(0, str(Path(__file__).parent.parent.parent / "lib"))
sys.path.insert(0, str(Path.home() / "Dev" / "tools" / "dev" / "lib"))  # canonical 5 modules

try:
    from lxml import etree
except ImportError:
    print("❌ 需要安装 lxml: pip install lxml")
    sys.exit(1)

from docx_parts import DEFAULT_ALLOW_CHANGED, assert_parts_intact  # noqa: E402
# verbatim repack SSOT —— fonts 子命令的就地截断重写 2026-08-01 迁入。
# clone/template/save_docx 三处**不迁**：它们是异地输出（ref→out / docx→output_path），
# lib 的 surgical_rewrite_parts 只能就地，没有 out 参数。
from docx_surgical import surgical_rewrite_parts  # noqa: E402
from docx_xml import (  # noqa: E402
    ALL_SCOPES,
    NSMAP,
    group_text,
    has_other_content,
    in_deleted,
    in_inserted,
    iter_paragraphs,
    iter_text_roots,
    para_own_runs,
    run_text,
    scope_of,
    set_group_text,
    set_run_text,
    text_groups,
    text_tag_for,
)
from file_ops import clear_quarantine  # noqa: E402
from finder import get_input_files  # noqa: E402
from progress import ProgressTracker  # noqa: E402
from text_fixes import fix_punctuation, fix_quotes, fix_units  # noqa: E402

# 放在 document/ 同级，导入需要处理路径（md_tools / docx_write_gate 同目录 SSOT）
sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_write_gate import WriteGate  # noqa: E402  原地写回并发门（同目录 SSOT）

# 复用 md_tools 的样式提取 + Finder 选中（原 md_docx_template.py，2026-07-31 并入 md_tools md2docx）
from md_tools import DEFAULT_TEMPLATE, extract_styles_xml, get_finder_selection  # noqa: E402

W = NSMAP["w"]
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"  # xml:space（内建命名空间）


# ════════════════════════════════════════════════════════════════════════════
# template — 原 docx_apply_template.py：Word 文档样式工具（套模板 + 清理）
#
# 二级子命令：
#   apply   - 给普通 docx 套上模板样式（默认；无二级子命令时兼容旧调用）
#   cleanup - 样式清理：重命名、合并、删除未使用样式
# ════════════════════════════════════════════════════════════════════════════


def apply_styles_to_docx(docx_path, styles_xml_path, output_path):
    """把模板样式注入到已有 docx 中"""

    # 读取提取的样式
    with open(styles_xml_path, "rb") as f:
        new_styles = etree.parse(f).getroot()

    with tempfile.TemporaryDirectory() as tmpdir:
        # 解压目标 docx
        with zipfile.ZipFile(docx_path, "r") as zf:
            zf.extractall(tmpdir)

        # 读取原 styles.xml
        orig_styles_path = os.path.join(tmpdir, "word", "styles.xml")
        with open(orig_styles_path, "rb") as f:
            orig_root = etree.parse(f).getroot()

        # 收集新样式的 ID
        new_style_ids = set()
        for style in new_styles.findall(".//w:style", NSMAP):
            style_id = style.get(f"{{{NSMAP['w']}}}styleId")
            new_style_ids.add(style_id)

        # 删除原文档中同名样式
        for style in orig_root.findall(".//w:style", NSMAP):
            style_id = style.get(f"{{{NSMAP['w']}}}styleId")
            if style_id in new_style_ids:
                orig_root.remove(style)

        # 添加新样式
        for style in new_styles.findall(".//w:style", NSMAP):
            orig_root.append(deepcopy(style))

        # 更新 docDefaults
        new_defaults = new_styles.find(".//w:docDefaults", NSMAP)
        if new_defaults is not None:
            old_defaults = orig_root.find(".//w:docDefaults", NSMAP)
            if old_defaults is not None:
                orig_root.remove(old_defaults)
            orig_root.insert(0, deepcopy(new_defaults))

        # 写回 styles.xml
        tree = etree.ElementTree(orig_root)
        tree.write(orig_styles_path, encoding="UTF-8", xml_declaration=True)

        # 重新打包
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for root_dir, _dirs, files in os.walk(tmpdir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, tmpdir)
                    zf.write(file_path, arcname)

    clear_quarantine(output_path)
    # 部件完整性断言（fail-closed）：基线 = 未动的源件。本函数只重写 styles.xml
    # （其余部件从 tmpdir 原字节回填），丢部件/白名单外改动即抛，不静默。
    # verbose=False 保持既有 stdout 契约。
    assert_parts_intact(docx_path, output_path,
                        allow_changed=set(DEFAULT_ALLOW_CHANGED) | {"word/styles.xml"},
                        verbose=False)
    print(f"✅ 输出: {output_path}")
    return output_path


def tmpl_cmd_apply(args):
    """template apply 子命令入口"""
    # 无参数时从 Finder 获取选中的 .docx 文件
    if not args.input:
        finder_file = get_finder_selection()
        if not finder_file:
            print("❌ Finder 中未选中任何文件")
            print("   💡 在 Finder 选中一个 .docx 文件后重新运行")
            sys.exit(1)
        if not finder_file.endswith(".docx"):
            ext = os.path.splitext(finder_file)[1] or "(无扩展名)"
            print(f"❌ 需要 .docx 文件，但选中的是: {ext}")
            print(f"   选中: {os.path.basename(finder_file)}")
            sys.exit(1)
        args.input = finder_file
        print(f"📄 从 Finder 获取: {os.path.basename(finder_file)}")

    if not os.path.exists(args.input):
        print(f"❌ 文件不存在: {args.input}")
        sys.exit(1)

    # 输出路径
    if args.output:
        output_path = args.output
    else:
        base, ext = os.path.splitext(args.input)
        output_path = f"{base}_styled{ext}"

    # 模板：优先用指定模板，缺失则降级到已提取样式（与 md_tools.py md2docx 行为一致）
    template = args.template or DEFAULT_TEMPLATE
    if os.path.exists(template):
        with tempfile.TemporaryDirectory() as tmpdir:
            print(f"📋 模板: {os.path.basename(template)}")
            styles_path, _config = extract_styles_xml(template, tmpdir)
            apply_styles_to_docx(args.input, styles_path, output_path)
    else:
        # 降级：使用同目录缓存的 heading_styles.xml
        fallback_styles = getattr(args, "styles", None) or os.path.join(
            os.path.dirname(os.path.abspath(__file__)), "heading_styles.xml"
        )
        if not os.path.exists(fallback_styles):
            print(f"❌ 模板不存在: {template}")
            print(f"❌ 缓存样式也不存在: {fallback_styles}")
            print("   请指定模板: -t 模板.docx")
            sys.exit(1)
        print(f"📋 模板缺失，使用缓存样式: {os.path.basename(fallback_styles)}")
        apply_styles_to_docx(args.input, fallback_styles, output_path)


def load_docx(path: str):
    """加载 docx，返回 (zip_bytes_dict, styles_tree, doc_tree)"""
    files = {}
    with zipfile.ZipFile(path, "r") as zf:
        for info in zf.infolist():
            files[info.filename] = zf.read(info.filename)

    styles_tree = etree.fromstring(files["word/styles.xml"])
    doc_tree = etree.fromstring(files["word/document.xml"])
    return files, styles_tree, doc_tree


def save_docx(files: dict, styles_tree, doc_tree, output_path: str):
    """保存修改后的 docx"""
    files["word/styles.xml"] = etree.tostring(styles_tree, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/document.xml"] = etree.tostring(doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in files.items():
            zf.writestr(name, data)
    with open(output_path, "wb") as f:
        f.write(buf.getvalue())
    clear_quarantine(output_path)


def get_style_map(styles_tree) -> dict:
    """从 styles.xml 提取 styleId → {name, type, basedOn}"""
    result = {}
    for s in styles_tree.findall(f".//{{{W}}}style"):
        sid = s.get(f"{{{W}}}styleId", "")
        stype = s.get(f"{{{W}}}type", "")
        ne = s.find(f"{{{W}}}name")
        name = ne.get(f"{{{W}}}val", sid) if ne is not None else sid
        based = s.find(f"{{{W}}}basedOn")
        base_id = based.get(f"{{{W}}}val", "") if based is not None else ""
        result[sid] = {"name": name, "type": stype, "basedOn": base_id, "elem": s}
    return result


def get_used_style_ids(doc_tree) -> set:
    """扫描 document.xml，找出所有被使用的样式 ID"""
    used = set()
    # 段落样式
    for ps in doc_tree.iter(f"{{{W}}}pStyle"):
        used.add(ps.get(f"{{{W}}}val", ""))
    # 字符样式
    for rs in doc_tree.iter(f"{{{W}}}rStyle"):
        used.add(rs.get(f"{{{W}}}val", ""))
    # 表格样式
    for ts in doc_tree.iter(f"{{{W}}}tblStyle"):
        used.add(ts.get(f"{{{W}}}val", ""))
    used.discard("")
    return used


def get_basedOn_ids(style_map: dict, used_ids: set) -> set:
    """递归找出所有被使用样式依赖的 basedOn 样式"""
    needed = set(used_ids)
    queue = list(used_ids)
    while queue:
        sid = queue.pop()
        info = style_map.get(sid)
        if info and info["basedOn"] and info["basedOn"] not in needed:
            needed.add(info["basedOn"])
            queue.append(info["basedOn"])
    return needed


# Word 内置样式 ID，不能删
BUILTIN_KEEP = {
    "a",  # Normal
    "a0",  # Default Paragraph Font
}


def cleanup(
    files: dict,
    styles_tree,
    doc_tree,
    renames: dict = None,
    merges: dict = None,
    delete_unused: bool = True,
    preview: bool = False,
) -> list[str]:
    """
    执行样式清理。

    Args:
        renames: {old_display_name: new_display_name} 重命名映射
        merges: {from_style_id: to_style_id} 合并映射
        delete_unused: 是否删除未使用的样式
        preview: 只生成报告不修改

    Returns:
        操作日志列表
    """
    log = []
    style_map = get_style_map(styles_tree)

    # ── 1. 合并样式 ──
    if merges:
        for from_id, to_id in merges.items():
            if from_id not in style_map:
                log.append(f"⚠️  合并跳过：源样式 ID '{from_id}' 不存在")
                continue
            if to_id not in style_map:
                log.append(f"⚠️  合并跳过：目标样式 ID '{to_id}' 不存在")
                continue

            count = 0
            if not preview:
                for ps in doc_tree.iter(f"{{{W}}}pStyle"):
                    if ps.get(f"{{{W}}}val") == from_id:
                        ps.set(f"{{{W}}}val", to_id)
                        count += 1
            else:
                for ps in doc_tree.iter(f"{{{W}}}pStyle"):
                    if ps.get(f"{{{W}}}val") == from_id:
                        count += 1

            from_name = style_map[from_id]["name"]
            to_name = style_map[to_id]["name"]
            log.append(f"🔀 合并：{from_name} ({from_id}) → {to_name} ({to_id})，{count} 个段落")

    # ── 2. 重命名样式 ──
    if renames:
        # 建立 name → styleId 映射
        name_to_id = {}
        for sid, info in style_map.items():
            name_to_id[info["name"]] = sid

        for old_name, new_name in renames.items():
            sid = name_to_id.get(old_name)
            if not sid:
                log.append(f"⚠️  重命名跳过：找不到样式 '{old_name}'")
                continue
            if not preview:
                ne = style_map[sid]["elem"].find(f"{{{W}}}name")
                if ne is not None:
                    ne.set(f"{{{W}}}val", new_name)
            log.append(f"✏️  重命名：{old_name} → {new_name}")

    # ── 3. 删除未使用样式 ──
    if delete_unused:
        # 重新扫描（合并后使用情况可能变了）
        used_ids = get_used_style_ids(doc_tree)
        needed_ids = get_basedOn_ids(style_map, used_ids)
        needed_ids |= BUILTIN_KEEP

        # 内置 heading/toc 样式保留（即使没用到，删了可能有问题）
        for sid in style_map:
            if sid.startswith("TOC") or sid in ("a", "a0"):
                needed_ids.add(sid)

        to_delete = []
        for sid, info in style_map.items():
            if sid not in needed_ids:
                to_delete.append((sid, info["name"]))

        if to_delete:
            log.append(f"\n🗑️  删除 {len(to_delete)} 个未使用样式：")
            for sid, name in sorted(to_delete, key=lambda x: x[1]):
                log.append(f"   - {name} ({sid})")
                if not preview:
                    for s in styles_tree.findall(f".//{{{W}}}style"):
                        if s.get(f"{{{W}}}styleId") == sid:
                            s.getparent().remove(s)
                            break

    return log


def tmpl_cmd_cleanup(args):
    """template cleanup 子命令入口"""
    if not args.output and not args.preview:
        print("请指定 -o 输出文件或使用 --preview 预览模式")
        sys.exit(1)

    # 加载配置
    if args.config:
        with open(args.config, encoding="utf-8") as f:
            config = json.load(f)
    else:
        config = {}

    files, styles_tree, doc_tree = load_docx(args.input)

    log = cleanup(
        files,
        styles_tree,
        doc_tree,
        renames=config.get("renames"),
        merges=config.get("merges"),
        delete_unused=config.get("delete_unused", True),
        preview=args.preview,
    )

    for line in log:
        print(line)

    if not args.preview and args.output:
        save_docx(files, styles_tree, doc_tree, args.output)
        # 部件完整性断言：save_docx 只覆写 styles.xml + document.xml
        # （document.xml 已在 DEFAULT_ALLOW_CHANGED），其余部件必须逐字节原样。
        assert_parts_intact(args.input, args.output,
                            allow_changed=set(DEFAULT_ALLOW_CHANGED) | {"word/styles.xml"},
                            verbose=False)
        print(f"\n📄 已保存到 {args.output}")
    elif args.preview:
        print("\n（预览模式，未修改文件）")


def template_main(argv):
    """template 子命令主入口（原 docx_apply_template.py main()，argv 显式穿参）。"""
    # 预扫描 argv：第一个非 flag 参数若是已知子命令则走 subparser，否则手工分流到 apply
    # （argparse subparsers 在第一个 positional 不是已知子命令时会 exit 2，故先 sniff）
    known_subcommands = {"apply", "cleanup"}
    raw_args = list(argv)
    first_positional = next((a for a in raw_args if not a.startswith("-")), None)

    if first_positional in known_subcommands:
        # 显式子命令：走 subparsers
        parser = argparse.ArgumentParser(
            prog="docx_apply_template.py",
            description="Word 文档样式工具（套模板 + 清理）",
            formatter_class=argparse.RawDescriptionHelpFormatter,
        )
        subparsers = parser.add_subparsers(dest="command", required=True)

        apply_parser = subparsers.add_parser("apply", help="给普通 docx 套上模板样式")
        _tmpl_add_apply_args(apply_parser)

        cleanup_parser = subparsers.add_parser("cleanup", help="样式清理：重命名、合并、删除未使用")
        cleanup_parser.add_argument("input", help="输入 .docx 文件")
        cleanup_parser.add_argument("-o", "--output", help="输出 .docx 文件")
        cleanup_parser.add_argument("--config", help="清理规则 JSON 文件")
        cleanup_parser.add_argument("--preview", action="store_true", help="只预览，不实际修改")

        args = parser.parse_args(raw_args)
        if args.command == "cleanup":
            tmpl_cmd_cleanup(args)
        else:
            tmpl_cmd_apply(args)
    else:
        # 无子命令（兼容旧调用：直接传路径 / Raycast 无参）→ 默认 apply
        # 顶层支持 --help / -h 列出子命令信息
        if any(a in ("-h", "--help") for a in raw_args):
            top = argparse.ArgumentParser(
                prog="docx_apply_template.py",
                description="Word 文档样式工具（套模板 + 清理）",
                formatter_class=argparse.RawDescriptionHelpFormatter,
                epilog=(
                    "子命令：\n"
                    "  apply   给普通 docx 套上模板样式（默认）\n"
                    "  cleanup 样式清理：重命名、合并、删除未使用样式\n"
                ),
            )
            _tmpl_add_apply_args(top, input_help="目标 docx 文件（默认 apply 模式）")
            top.parse_args(raw_args)  # exits with help
            return

        apply_compat = argparse.ArgumentParser(prog="docx_apply_template.py",
                                               description="给普通 docx 套上模板样式")
        _tmpl_add_apply_args(apply_compat)
        args = apply_compat.parse_args(raw_args)
        args.command = "apply"
        tmpl_cmd_apply(args)


def _tmpl_add_apply_args(parser, input_help="目标 docx 文件"):
    """input/-t/-o 三参数声明（原文件里写了 3 遍，收敛为一处；--help 文本逐字同款——
    顶层兼容分支的 input 帮助文案带「（默认 apply 模式）」后缀，用参数保差异）。"""
    parser.add_argument("input", nargs="?", help=input_help)
    parser.add_argument("-t", "--template", help="模板 docx 文件")
    parser.add_argument("-o", "--output", help="输出文件")


# ════════════════════════════════════════════════════════════════════════════
# clone — 原 docx_format_clone.py：提取 / 复刻 docx 版式（格式 SSOT）
#
# 为什么单独一族：
#   - pandoc --reference-doc 只按「样式名」映射，抓不到「直排版」（run 直接 rPr，
#     如公文标题 方正小标宋简体 二号居中 不走 heading 样式）→ 复刻不忠实。
#   - template apply -t 只换 styles.xml，同样丢直排版。
#   按「角色（标题/正文/落款）」抽出**有效格式**，用「段落外壳克隆」法复刻。
#
# 二级子命令：
#   extract <ref.docx> [-o profile.json]           解析范式件 → format-profile.json
#   apply <content.md|.docx> --ref <ref.docx> [-o out.docx]   外壳克隆复刻
# ════════════════════════════════════════════════════════════════════════════


# ─── XML 读取辅助 ────────────────────────────────────────────────────
def _read_document_xml(docx_path: str) -> etree._Element:
    with zipfile.ZipFile(docx_path) as z:
        return etree.fromstring(z.read("word/document.xml"))


def _body(root: etree._Element) -> etree._Element:
    return root.find(qn("w:body"))


def _para_text(p: etree._Element) -> str:
    return "".join(t.text or "" for t in p.iter(qn("w:t")))


def _first_rpr(p: etree._Element):
    """段落里第一个 run 的 rPr（run 直排版）。"""
    for r in p.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            return rpr
    return None


def _eff_font(p: etree._Element) -> str | None:
    rpr = _first_rpr(p)
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is not None:
        for attr in ("eastAsia", "ascii", "hAnsi"):
            v = rfonts.get(qn(f"w:{attr}"))
            if v and v != "zh-CN":
                return v
    return None


def _eff_size_halfpt(p: etree._Element) -> int | None:
    rpr = _first_rpr(p)
    if rpr is None:
        return None
    sz = rpr.find(qn("w:sz"))
    if sz is not None:
        v = sz.get(qn("w:val"))
        return int(v) if v and v.isdigit() else None
    return None


def _ppr_attr(p: etree._Element):
    """(jc, firstLineChars or firstLine) 段落直排版。"""
    ppr = p.find(qn("w:pPr"))
    jc_val = None
    ind_chars = None
    if ppr is not None:
        jc = ppr.find(qn("w:jc"))
        if jc is not None:
            jc_val = jc.get(qn("w:val"))
        ind = ppr.find(qn("w:ind"))
        if ind is not None:
            flc = ind.get(qn("w:firstLineChars"))
            fl = ind.get(qn("w:firstLine"))
            if flc:
                ind_chars = int(flc) / 100.0
            elif fl:
                ind_chars = int(fl) / 100.0  # 近似（twip 时另算，公文件多用 Chars）
    return jc_val, ind_chars


_SENT_END = "。！？；…"


# ─── 角色分类（范式件） ───────────────────────────────────────────────
def classify_shells(root: etree._Element) -> dict:
    """返回 {'title': <w:p>, 'body': <w:p>, 'signature': <w:p>}（代表性外壳）。

    规则：
      正文主导字号 = 非空段里出现最多的字号（公文件正文段最多）
      title  = 起始 居中 或 字号大于主导字号 的段（标题可能拆成多行，取首行壳）
      signature = 尾部 无句末标点 的非空段（落款）
      body   = 主导字号、非标题、非落款、且优先「无 pStyle（纯 Normal）」的代表段
    """
    from collections import Counter
    paras = [p for p in _body(root).findall(qn("w:p"))]
    nonempty = [p for p in paras if _para_text(p).strip()]
    if not nonempty:
        raise ValueError("范式件无文字段落")

    sizes = [(_eff_size_halfpt(p) or 0) for p in nonempty]
    cnt = Counter(s for s in sizes if s)
    dominant = cnt.most_common(1)[0][0] if cnt else 0

    def is_titleish(p) -> bool:
        jc, _ = _ppr_attr(p)
        sz = _eff_size_halfpt(p) or 0
        return jc == "center" or (dominant > 0 and sz > dominant)

    # 标题：首个 titleish 段
    title_p = next((p for p in nonempty if is_titleish(p)), nonempty[0])

    # 落款：尾段无句末标点 → 落款；否则退而取尾段
    last_p = nonempty[-1]
    lt = _para_text(last_p).strip()
    sig_p = last_p if (lt and lt[-1] not in _SENT_END) else last_p

    # 正文：主导字号 + 非标题 + 非落款，优先无 pStyle 的纯 Normal 段
    cands = [p for p in nonempty
             if not is_titleish(p) and p is not sig_p
             and (_eff_size_halfpt(p) or 0) == dominant]
    body_p = None
    for p in cands:
        ppr = p.find(qn("w:pPr"))
        if ppr is None or ppr.find(qn("w:pStyle")) is None:
            body_p = p
            break
    if body_p is None:
        body_p = cands[0] if cands else next(
            (p for p in nonempty if p is not title_p and p is not sig_p), title_p)

    return {"title": title_p, "body": body_p, "signature": sig_p}


def shell_profile(p: etree._Element, role: str) -> dict:
    jc, ind = _ppr_attr(p)
    sz = _eff_size_halfpt(p)
    return {
        "role": role,
        "font_cn": _eff_font(p),
        "size_pt": (sz / 2.0) if sz else None,
        "size_hao": _halfpt_to_hao(sz) if sz else None,
        "align": jc,
        "first_line_indent_chars": ind,
    }


_HAO = {84: "初号", 72: "小初", 52: "一号", 44: "二号", 36: "小二",
        32: "三号", 30: "小三", 28: "四号", 24: "小四", 21: "五号"}


def _halfpt_to_hao(halfpt: int) -> str | None:
    return _HAO.get(halfpt)


# ─── extract ────────────────────────────────────────────────────────
def clone_cmd_extract(args) -> int:
    ref = args.ref or args.input
    if not ref or not os.path.exists(ref):
        print(f"❌ 范式件不存在: {ref}", file=sys.stderr)
        return 1
    root = _read_document_xml(ref)
    shells = classify_shells(root)
    profile = {
        "source": os.path.basename(ref),
        "title": shell_profile(shells["title"], "title"),
        "body": shell_profile(shells["body"], "body"),
        "signature": shell_profile(shells["signature"], "signature"),
    }
    out = args.output or (os.path.splitext(ref)[0] + ".format-profile.json")
    with open(out, "w", encoding="utf-8") as f:
        json.dump(profile, f, ensure_ascii=False, indent=2)
    print(f"✅ 格式 profile → {out}")
    for role in ("title", "body", "signature"):
        r = profile[role]
        print(f"  [{role:9}] {r['font_cn']} · {r['size_hao'] or r['size_pt']} · "
              f"对齐={r['align']} · 首行缩进={r['first_line_indent_chars']}字符")
    return 0


# ─── content 解析 ────────────────────────────────────────────────────
def parse_content(path: str) -> dict:
    """→ {'title': str, 'body': [str], 'signature': [str]}"""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".md":
        blocks = _md_blocks(path)
    elif ext == ".docx":
        blocks = _docx_blocks(path)
    else:
        raise ValueError(f"content 须为 .md 或 .docx: {path}")
    if not blocks:
        raise ValueError("content 无文字")

    # 标题：首块（去 # 标记）
    title = re.sub(r"^#+\s*", "", blocks[0]).strip()
    rest = blocks[1:]

    # 落款：尾部连续 无句末标点 的块
    sig: list[str] = []
    while rest and rest[-1].strip() and rest[-1].strip()[-1] not in _SENT_END:
        sig.insert(0, rest.pop().strip())
    body = [b.strip() for b in rest if b.strip()]
    return {"title": title, "body": body, "signature": sig}


def _md_blocks(path: str) -> list[str]:
    raw = Path(path).read_text(encoding="utf-8")
    # 空行分块；块内多行折成一段（去硬换行）
    blocks = []
    for chunk in re.split(r"\n\s*\n", raw):
        line = " ".join(s.strip() for s in chunk.splitlines() if s.strip())
        line = re.sub(r"\s+", "", line) if _is_cjk_heavy(line) else line
        if line.strip():
            blocks.append(line)
    return blocks


def _is_cjk_heavy(s: str) -> bool:
    cjk = sum(1 for c in s if "一" <= c <= "鿿")
    return cjk > len(s) * 0.3


def _docx_blocks(path: str) -> list[str]:
    root = _read_document_xml(path)
    out = []
    for p in _body(root).findall(qn("w:p")):
        t = _para_text(p).strip()
        if t:
            out.append(t)
    return out


# ─── apply（外壳克隆） ────────────────────────────────────────────────
def _clone_with_text(shell: etree._Element, text: str) -> etree._Element:
    """克隆 shell 段落，保留 pPr + 首个 run 的 rPr，正文换为 text（单 run）。"""
    newp = copy.deepcopy(shell)
    # 取首 run 的 rPr 作模板
    rpr_tmpl = None
    src_rpr = _first_rpr(shell)
    if src_rpr is not None:
        rpr_tmpl = copy.deepcopy(src_rpr)
    # 删所有 run（保 pPr 等）
    for r in newp.findall(qn("w:r")):
        newp.remove(r)
    # 也删 hyperlink 内 run（简化：不处理超链接）
    r = etree.SubElement(newp, qn("w:r"))
    if rpr_tmpl is not None:
        r.append(rpr_tmpl)
    t = etree.SubElement(r, qn("w:t"))
    t.set(XML_SPACE, "preserve")
    t.text = text
    return newp


def _blank_like(shell: etree._Element) -> etree._Element:
    """造一个与 shell 同 pPr 的空段（间距用）。"""
    newp = copy.deepcopy(shell)
    for r in newp.findall(qn("w:r")):
        newp.remove(r)
    return newp


def clone_cmd_apply(args) -> int:
    content_path = args.input
    ref = args.ref
    if not content_path or not os.path.exists(content_path):
        print(f"❌ content 不存在: {content_path}", file=sys.stderr)
        return 1
    if not ref or not os.path.exists(ref):
        print(f"❌ 范式件不存在: {ref}", file=sys.stderr)
        return 1

    content = parse_content(content_path)
    root = _read_document_xml(ref)
    shells = classify_shells(root)
    body_el = _body(root)
    sectPr = body_el.find(qn("w:sectPr"))

    # 重建 body：清空旧段，按 标题 / 正文 / 空行 / 落款 重填
    for child in list(body_el):
        if child.tag == qn("w:sectPr"):
            continue
        body_el.remove(child)

    new_children: list[etree._Element] = []
    if content["title"]:
        new_children.append(_clone_with_text(shells["title"], content["title"]))
    for para in content["body"]:
        new_children.append(_clone_with_text(shells["body"], para))
    emit_sig = content["signature"] and not getattr(args, "no_signature", False)
    if emit_sig:
        new_children.append(_blank_like(shells["body"]))  # 落款前空一行
        for line in content["signature"]:
            new_children.append(_clone_with_text(shells["signature"], line))

    # 插到 sectPr 之前
    if sectPr is not None:
        for el in new_children:
            sectPr.addprevious(el)
    else:
        for el in new_children:
            body_el.append(el)

    # 输出：复制 ref（保 styles/fonts/numbering/media），仅换 document.xml
    out = args.output or (os.path.splitext(content_path)[0] + "_styled_fixed.docx")
    new_doc_xml = etree.tostring(root, xml_declaration=True,
                                 encoding="UTF-8", standalone=True)
    with zipfile.ZipFile(ref) as zin:
        names = zin.namelist()
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for name in names:
                data = zin.read(name)
                if name == "word/document.xml":
                    data = new_doc_xml
                zout.writestr(name, data)
    clear_quarantine(out)
    # 部件完整性断言（fail-closed）：基线 = 范式件 ref —— 产物的部件全部继承自它
    # （content 只是文字源），本函数只换 document.xml（在默认白名单），
    # 其余部件必须逐字节 verbatim。verbose=False 保持既有 stdout 契约。
    assert_parts_intact(ref, out, verbose=False)
    print(f"✅ 复刻 → {out}")
    print(f"   标题: {content['title'][:40]}")
    print(f"   正文: {len(content['body'])} 段 · 落款: {len(content['signature'])} 行")
    return 0


def clone_main(argv=None) -> int:
    p = argparse.ArgumentParser(prog="docx_format_clone.py",
                                description="提取/复刻 docx 版式")
    sub = p.add_subparsers(dest="cmd")

    pe = sub.add_parser("extract", help="提取格式 → profile.json")
    pe.add_argument("input", help="范式 docx")
    pe.add_argument("--ref", help="（同 input）")
    pe.add_argument("-o", "--output", help="输出 profile.json")
    pe.set_defaults(func=clone_cmd_extract)

    pa = sub.add_parser("apply", help="复刻格式（外壳克隆）")
    pa.add_argument("input", help="content（.md 或 .docx）")
    pa.add_argument("--ref", required=True, help="范式 docx（格式来源）")
    pa.add_argument("-o", "--output", help="输出 docx")
    pa.add_argument("--no-signature", action="store_true",
                    help="不发射落款段（交付件留空让对方自填署名）")
    pa.set_defaults(func=clone_cmd_apply)

    args = p.parse_args(argv if argv is not None else sys.argv[1:])
    if not getattr(args, "func", None):
        p.print_help()
        return 0
    return args.func(args)


# ════════════════════════════════════════════════════════════════════════════
# fonts — 原 docx_font_normalize.py：去「等线」，中文字体归一到 宋体/仿宋 体系
#
# 用户钦定（2026-07-20 /govern）：**所有 docx 文档都不用等线字体，要么宋体 要么仿宋**。
# 等线的三条来路（缺一不可全堵）：
#   ① docDefaults 走 minorEastAsia 主题而 theme <a:ea typeface=""> 为空 → Word 回退等线
#   ② run/style 显式写 等线/DengXian/Deng Xian
#   ③ 脚本新建 run 不设 rFonts → 继承 ①
# 修复动作(--apply)：docDefaults 显式化 + theme <a:ea> 补字体 + 全 part 显式等线替换。
# ════════════════════════════════════════════════════════════════════════════

DENGXIAN = re.compile(r"等线|DengXian|Deng Xian")
THEME_EA_RE = re.compile(r'(<a:ea\s+typeface=")([^"]*)(")')
DOCDEFAULT_RFONTS = re.compile(r"(<w:docDefaults>.*?<w:rPrDefault>.*?<w:rPr>.*?)(<w:rFonts[^/>]*/>)", re.S)


def _parts(zf: zipfile.ZipFile) -> list[str]:
    return [n for n in zf.namelist() if n.endswith(".xml")]


def scan(path: Path) -> list[str]:
    """返回问题列表, 空 = 干净。"""
    issues: list[str] = []
    with zipfile.ZipFile(path) as z:
        names = _parts(z)
        for n in names:
            text = z.read(n).decode("utf8", errors="ignore")
            if DENGXIAN.search(text):
                issues.append(f"{n}: 显式等线/DengXian")
        try:
            styles = z.read("word/styles.xml").decode("utf8", errors="ignore")
        except KeyError:
            styles = ""
        theme_ea = []
        for n in names:
            if n.startswith("word/theme/"):
                theme_ea += THEME_EA_RE.findall(z.read(n).decode("utf8", errors="ignore"))
        if "eastAsiaTheme" in styles and any(t[1].strip() == "" for t in theme_ea):
            issues.append("word/styles.xml: docDefaults 走 minorEastAsia 主题，而 theme 的 <a:ea> 为空 → Word 回退等线")
    return issues


def fix(path: Path, font: str, ascii_font: str) -> list[str]:
    changed: list[str] = []
    touched: set[str] = set()   # 本次真被改写的部件名（动态白名单，静态清单会误报/漏报）
    ts = _dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    shutil.copy(path, f"{path}.bak-{ts}")
    with zipfile.ZipFile(path) as zin:
        items = zin.infolist()
        data = {i.filename: zin.read(i.filename) for i in items}

    for name in list(data):
        if not name.endswith(".xml"):
            continue
        text = data[name].decode("utf8", errors="ignore")
        orig = text

        if DENGXIAN.search(text):
            text = DENGXIAN.sub(font, text)
            changed.append(f"{name}: 显式等线 → {font}")

        if name.startswith("word/theme/"):
            def _ea(m: re.Match[str]) -> str:
                return m.group(1) + (m.group(2) or font if m.group(2).strip() else font) + m.group(3)
            new = THEME_EA_RE.sub(_ea, text)
            if new != text:
                text, _ = new, changed.append(f"{name}: theme <a:ea> 空 → {font}")

        if name == "word/styles.xml":
            def _rf(m: re.Match[str]) -> str:
                return (
                    m.group(1)
                    + f'<w:rFonts w:ascii="{ascii_font}" w:eastAsia="{font}" '
                      f'w:hAnsi="{ascii_font}" w:cs="{ascii_font}"/>'
                )
            new = DOCDEFAULT_RFONTS.sub(_rf, text, count=1)
            if new != text:
                text = new
                changed.append(f"{name}: docDefaults 主题字体 → 显式 {font} / {ascii_font}")

        if text != orig:
            data[name] = text.encode("utf8")
            touched.add(name)

    if changed and touched:
        # 2026-08-01：原来是 ZipFile(path, "w") **就地截断**源件再逐部件写回 —— 全仓唯一
        # 一处写到一半崩就毁源件的写法。改走 lib/docx_surgical：写 .tmp → 部件完整性断言
        # （基线从 .bak-ts 副本换成 replace 前的源件本身，字节等价；白名单仍是动态的
        # touched，静态清单会漏掉 theme/*.xml）→ 原子 replace → 写后 CRC/parse 自检。
        surgical_rewrite_parts(path, {n: data[n] for n in touched}, backup=False)
    elif changed:
        # changed 非空而 touched 为空：规则命中但替换后字节没变（`--font 等线` 就是这样 ——
        # DENGXIAN.sub("等线") 得到同一串，changed 已 append 而 text == orig）。原来这里会
        # 把全部部件原样重写一遍（等价于不写），现在直接不写；lib 拒绝空 parts，
        # 少了这个分支会撞 ValueError 崩掉。
        pass
    else:
        Path(f"{path}.bak-{ts}").unlink(missing_ok=True)
    return changed


def fonts_main(argv) -> int:
    ap = argparse.ArgumentParser(prog="docx_font_normalize.py",
                                 description="docx 去等线（中文字体归一到 宋体/仿宋）")
    ap.add_argument("docx", nargs="+")
    ap.add_argument("--check", action="store_true", help="只报告不改（默认）")
    ap.add_argument("--apply", action="store_true", help="原地修复 + .bak-时间戳")
    ap.add_argument("--font", default="仿宋_GB2312", help="中文字体（默认 仿宋_GB2312）")
    ap.add_argument("--ascii", dest="ascii_font", default="Times New Roman")
    args = ap.parse_args(argv)

    rc = 0
    for p in args.docx:
        path = Path(p)
        if path.name.startswith("~$"):
            continue
        if not path.exists():
            print(f"[跳过] 不存在: {path}", file=sys.stderr)
            continue
        issues = scan(path)
        if not issues:
            print(f"✅ {path.name}: 无等线风险")
            continue
        rc = 1
        print(f"⚠️  {path.name}")
        for i in issues:
            print(f"    · {i}")
        if args.apply:
            for c in fix(path, args.font, args.ascii_font):
                print(f"    ✔ {c}")
            left = scan(path)
            print("    ✅ 复检干净" if not left else f"    ❌ 仍有: {left}")
            rc = 0 if not left else 1
    return rc


# ════════════════════════════════════════════════════════════════════════════
# text — 原 docx_text_formatter.py：文本格式自动修复（DOCX 版）
#
# 规则轴：① 双引号统一中文弯引号（奇=" 偶="）② 英文标点→中文 ③ 中文单位→标准符号
#         ④ 单位上标（m2→m² 等）；引号可拆独立 run 设宋体（排版动作,可单关）。
# 覆盖范围（2026-07-26 扩齐）：正文 + 表格（含嵌套表）+ 文本框 + 超链接 +
#   审阅修订（w:ins/w:del）+ 批注 comments.xml + 脚注/尾注 + 页眉页脚。
#   python-docx 的 Document.paragraphs / Paragraph.runs 只认直接子节点，带修订的
#   稿子里 w:ins/w:del 段一个字都改不到 —— 统一走 docx_xml 的元素级遍历。
#   删除态文本改 w:delText（不退化成 w:t，修订结构原样保留）。
# ════════════════════════════════════════════════════════════════════════════

# ===== 配置对象 =====
# 旧实现把 5 个旋钮做成模块级可变全局,且只在 `if __name__ == "__main__"` 里赋值 ——
# 被 import 时永远拿默认值,同进程连跑两组配置会串味(第二次沿用上次残留)。
# 现在一律走冻结的 FormatConfig 显式穿参;模块级全局只作为「CLI 默认值」保留。


@dataclass(frozen=True)
class FormatConfig:
    """一次规范化的完整意图 = 规则轴 × 域轴 × 写回方式。"""

    # 规则轴
    quotes: bool = True            # 引号统一为中文弯引号
    punct: bool = True             # 英文标点 → 中文标点
    units: bool = True             # 中文单位 → 标准符号
    quote_font: bool = True        # 引号拆独立 run 并设宋体(排版动作,可单独关)
    # 域轴(空 = 全选);取值见 docx_xml.ALL_SCOPES
    scopes: frozenset = frozenset(ALL_SCOPES)
    # 破坏性 / 写回
    strip_headers: bool = False    # 删除页眉页脚引用(默认关,属排版动作)
    in_place: bool = False         # 原地写回 + .bak 备份(默认另存 _fixed)

    def wants(self, scope: str) -> bool:
        return scope in self.scopes

    @classmethod
    def from_opts(cls, opts: dict | None):
        """从 {"rule.quotes": "0", "scope.comments": "1", ...} 造配置(GUI 通路)。

        未知 key / 非法值抛 ValueError,由调用方翻成信封里的 ok:false。
        """
        cfg = cls()
        if not opts:
            return cfg
        rules, scopes = {}, set(cfg.scopes)
        for k, v in opts.items():
            on = _truthy(k, v)
            if k.startswith("rule."):
                name = k[5:]
                if name not in RULE_KEYS:
                    raise ValueError(f"未知规则:{name}(可用 {'/'.join(RULE_KEYS)})")
                rules[name] = on
            elif k.startswith("scope."):
                name = k[6:]
                if name not in ALL_SCOPES:
                    raise ValueError(f"未知范围:{name}(可用 {'/'.join(ALL_SCOPES)})")
                scopes.add(name) if on else scopes.discard(name)
            elif k in ("strip_headers", "in_place"):
                rules[k] = on
            else:
                raise ValueError(f"未知选项:{k}")
        return replace(cfg, scopes=frozenset(scopes), **rules)


RULE_KEYS = ("quotes", "punct", "units", "quote_font")


def _truthy(k: str, v) -> bool:
    if isinstance(v, bool):
        return v
    s = str(v).strip().lower()
    if s in ("1", "true", "yes", "on", "是"):
        return True
    if s in ("0", "false", "no", "off", "否"):
        return False
    raise ValueError(f"选项 {k} 的值 {v!r} 不是布尔(用 1/0)")


# ===== CLI 默认值 =====
# 删页眉页脚 = 排版动作,不属于「文本规范化」。2026-07-26 默认翻转为**保留**:
# 旧默认 True 是 initial commit 带进来的,从没人显式传过,却让每次「规范化」都把
# docx 的 headerReference/footerReference 全删光(实测 2+2 → 0),且 typeset 套完
# 院模板后第 2 步复用本引擎,把模板自带的 14/13 个页眉页脚引用一并删掉。
# 要删请显式 --strip-headers。
STRIP_HEADERS = False

# 选择性修复开关（默认全开=向后兼容；命令行 --quotes/--punct/--units 任一指定则只做指定项）
# 用途：中文期刊投稿(GB/T 7714 参考文献区标点须半角)只想修引号→ --quotes，不碰标点/单位
DO_QUOTES = True
DO_PUNCT = True
DO_UNITS = True
DO_QUOTE_FONT = True   # 引号拆独立 run 并设宋体（排版动作；--no-quote-font 可单关）
SCOPES = set(ALL_SCOPES)   # 默认全域：正文/表格/审阅修订/批注/脚注尾注/页眉页脚
IN_PLACE = False   # True 时原地写回源文件 + 备份 .bak-时间戳（Work §1.5 协议），不另存 _fixed


def _cli_config() -> "FormatConfig":
    """把模块级 CLI 默认值收成一个 FormatConfig（旧调用方零改动仍可用）。"""
    return FormatConfig(
        quotes=DO_QUOTES, punct=DO_PUNCT, units=DO_UNITS, quote_font=DO_QUOTE_FONT,
        scopes=frozenset(SCOPES), strip_headers=STRIP_HEADERS, in_place=IN_PLACE)


def _fmt_rules(cfg) -> str:
    on = [n for n, v in (("引号", cfg.quotes), ("标点", cfg.punct),
                         ("单位", cfg.units), ("引号宋体", cfg.quote_font)) if v]
    return "+".join(on) or "(无)"


def _fmt_scopes(cfg) -> str:
    return "+".join(ALL_SCOPES[k] for k in ALL_SCOPES if cfg.wants(k)) or "(无)"


QUOTE_CHARS = {"“", "”"}
QUOTE_FONT = "宋体"  # 宋体


def _set_run_font_songti(run_element):
    """为 run 的 rPr 设置宋体字体（ascii + hAnsi + eastAsia）"""
    rPr = run_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), QUOTE_FONT)
    rFonts.set(qn("w:hAnsi"), QUOTE_FONT)
    rFonts.set(qn("w:eastAsia"), QUOTE_FONT)
    rFonts.set(qn("w:hint"), "eastAsia")


def _split_text_at_quotes(text):
    """
    将文本在引号位置切片，返回 [(text, is_quote), ...] 片段列表。
    如果没有引号，返回 None。
    """
    if not text or not any(c in QUOTE_CHARS for c in text):
        return None

    segments = []
    buf = []
    for c in text:
        if c in QUOTE_CHARS:
            if buf:
                segments.append(("".join(buf), False))
                buf = []
            segments.append((c, True))
        else:
            buf.append(c)
    if buf:
        segments.append(("".join(buf), False))
    return segments


def _apply_quote_split(r, segments):
    """
    根据 segments 拆分 run 元素 r：第一段复用原 run，后续段插入新 run。
    引号段设置宋体。删除态 run（w:del 内）新建的文本元素仍是 w:delText。
    """
    parent = r.getparent()
    text_tag = text_tag_for(r)          # w:del / w:moveFrom 内必须继续用 w:delText

    # 第一段复用原 run
    first_text, first_is_quote = segments[0]
    set_run_text(r, first_text)
    if first_is_quote:
        _set_run_font_songti(r)

    # 后续段：复制原 run 的格式，插入到原 run 之后
    insert_after = r
    for seg_text, is_quote in segments[1:]:
        new_r = copy.deepcopy(r)
        # 只留 rPr：deepcopy 带来的旧文本、以及 drawing/footnoteReference 等
        # 一次性载荷都必须清掉，否则拆一次 run 就把图片/脚注引用复制一份出来
        for child in list(new_r):
            if child.tag != qn("w:rPr"):
                new_r.remove(child)
        t_elem = OxmlElement(text_tag)
        t_elem.text = seg_text
        # 保留空格
        t_elem.set(qn("xml:space"), "preserve")
        new_r.append(t_elem)

        if is_quote:
            _set_run_font_songti(new_r)
        else:
            # 非引号段恢复原 run 的字体（去掉宋体覆盖）
            rPr = new_r.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                orig_rPr = r.find(qn("w:rPr"))
                orig_rFonts = orig_rPr.find(qn("w:rFonts")) if orig_rPr is not None else None
                if rFonts is not None and orig_rFonts is not None:
                    # 用原始字体信息覆盖
                    rPr.replace(rFonts, copy.deepcopy(orig_rFonts))
                elif rFonts is not None and orig_rFonts is None:
                    rPr.remove(rFonts)

        parent.insert(list(parent).index(insert_after) + 1, new_r)
        insert_after = new_r


def _tally(stats, scope, r, n):
    """按域记账：修订态单独计，好让用户看见「审阅段确实改到了」。"""
    if not n:
        return
    key = scope
    if scope == "正文":
        if in_deleted(r):
            key = "修订·删除态"
        elif in_inserted(r):
            key = "修订·插入态"
    stats["scopes"][key] = stats["scopes"].get(key, 0) + n


def process_paragraph_element(p, stats, scope="正文", cfg=None):
    """
    处理一个 w:p 元素内的文本，逐 run 处理以保持每个 run 的字体格式。

    走 para_own_runs 而非 python-docx 的 `Paragraph.runs`：后者只认 `./w:r`，
    审阅修订（w:ins/w:del）、超链接、smartTag 里的 run 会被静默跳过。

    ⚠ 域过滤与引号规则**不正交**：引号左右方向靠段落级奇偶计数器，被跳过的 run
    如果不推进计数器，它后面的引号会整体反相（实测 `乙"结束` 从 `乙“` 翻成 `乙”`）。
    所以未选中的 run 仍然跑一遍 fix_quotes 推进 counter，只是**不写回**。
    """
    cfg = cfg or FormatConfig()
    original_runs = para_own_runs(p)      # 先快照:遍历中会插入新 run
    if not original_runs:
        return

    quote_counter = 0                     # 引号计数器跨 run 维护,保证配对正确

    for r in original_runs:
        # 正文 part 内按 run 归属的子域(正文/表格/审阅修订)过滤;其余 part 整体由
        # iter_text_roots 的 parts 决定,进到这里就是选中的
        selected = cfg.wants(scope_of(r)) if scope == "正文" else True

        # 逐「文本组」处理：w:tab / w:br 隔开的文本不合并，否则会被挤到末尾（对齐错位）
        groups = text_groups(r)
        for g in groups:
            original_text = group_text(g)
            if not original_text:
                continue

            fixed_text = original_text
            quote_count = punct_count = unit_count = 0
            if cfg.quotes:
                fixed_text, quote_count, quote_counter = fix_quotes(fixed_text, quote_counter)
            if not selected:
                continue                  # counter 已推进,但不改这段文字、不记账
            if cfg.punct:
                fixed_text, punct_count = fix_punctuation(fixed_text)
            if cfg.units:
                fixed_text, unit_count = fix_units(fixed_text)

            stats["quotes"] += quote_count
            stats["punctuation"] += punct_count
            stats["units"] += unit_count
            _tally(stats, scope, r, quote_count + punct_count + unit_count)

            if fixed_text != original_text:
                set_group_text(g, fixed_text)

        # 引号拆独立 run + 宋体（排版动作，可单独关）。
        # run 里夹着 tab/br/图/脚注引用时不拆 —— 复制这些一次性载荷会出复件
        if selected and cfg.quotes and cfg.quote_font and len(groups) == 1 and not has_other_content(r):
            segments = _split_text_at_quotes(run_text(r))
            if segments:
                _apply_quote_split(r, segments)


def process_root(root, stats, scope="正文", cfg=None):
    """处理一个 XML 根（正文 body / 批注 / 脚注 / 页眉…）下的全部段落。

    `iter_paragraphs` 递归到嵌套表格、文本框、修订块里的 w:p；`para_own_runs`
    在嵌套 w:p 处剪枝，所以每个 run 恰好被处理一次。
    """
    for p in list(iter_paragraphs(root)):
        process_paragraph_element(p, stats, scope, cfg)


def _strip_header_footer_refs(doc) -> int:
    """删掉所有 section 的 headerReference/footerReference，返回删掉的引用数。

    删「引用」而不是 header/footer part 本身：Word 见不到引用就按无页眉页脚渲染，
    留在包里的孤儿 part 无害；去删 part 会连带 rels 一起崩。
    """
    removed = 0
    for section in doc.sections:
        sectPr = section._sectPr
        for ref in (sectPr.findall(qn("w:headerReference"))
                    + sectPr.findall(qn("w:footerReference"))):
            sectPr.remove(ref)
            removed += 1
    return removed


def _normalize_doc(doc, cfg: FormatConfig) -> dict:
    """对**已打开的** doc 做完整规范化，返回统计。不开文件、不存盘、不打印。

    process_docx（CLI/GUI 通路，手上已经是 FormatConfig）与 apply（pipeline 通路，
    手上是 argparse.Namespace）共用这一份实现 —— 两边各抄一遍这个循环，正是
    「脚注 flush 漏一处」「新加的 part 只接进了一条通路」这类 bug 的温床。
    """
    stats = {"quotes": 0, "punctuation": 0, "units": 0, "scopes": {}}

    # 处理所有承载文本的 part：正文（含表格/嵌套表/文本框/超链接/审阅修订）
    # + 批注 comments.xml + 脚注 + 尾注 + 页眉页脚（默认保留并规范化）
    flushes = []
    parts = {k for k in ("comments", "notes", "headers") if cfg.wants(k)}
    for label, root, flush in iter_text_roots(
            doc, include_headers=not cfg.strip_headers, parts=parts):
        process_root(root, stats, label, cfg)
        if flush is not None:
            flushes.append(flush)
    for flush in flushes:          # 通用 Part（脚注/尾注）改完必须回写 blob
        flush()

    # 显式要求时才删页眉页脚（默认保留）
    stats["headers_stripped"] = _strip_header_footer_refs(doc) if cfg.strip_headers else 0
    stats["changed"] = (stats["quotes"] + stats["punctuation"] + stats["units"]
                        + stats["headers_stripped"])
    return stats


def _config_from_args(args=None) -> FormatConfig:
    """argparse.Namespace → FormatConfig（pipeline 通路）。

    逐项 getattr 带默认：pipeline 里同一个 Namespace 要喂给一串 step，缺哪项就该
    沿用本脚本的默认值，而不是 AttributeError 掀桌。dict 通路是 GUI 专用的
    FormatConfig.from_opts（键名带 rule./scope. 前缀），两套别混。
    """
    base = _cli_config()
    if args is None:
        return base
    ready = getattr(args, "format_config", None)
    if isinstance(ready, FormatConfig):     # 调用方已算好完整意图，直接用
        return ready
    scopes = getattr(args, "scopes", None)
    return replace(
        base,
        quotes=bool(getattr(args, "quotes", base.quotes)),
        punct=bool(getattr(args, "punct", base.punct)),
        units=bool(getattr(args, "units", base.units)),
        quote_font=bool(getattr(args, "quote_font", base.quote_font)),
        scopes=frozenset(scopes) if scopes else base.scopes,
        strip_headers=bool(getattr(args, "strip_headers", base.strip_headers)),
    )


# ---------------- pipeline adapter ----------------
def apply(doc, args=None) -> dict:
    """在**已打开的** doc 上做文本规范化（引号/标点/单位 + 引号宋体），返回统计。

    覆盖面照旧走 docx_xml 的 iter_text_roots：批注/脚注/尾注/页眉页脚是独立 part，
    审阅修订（w:ins/w:del）又不是 w:p 的直接子 run —— 只遍历 doc.paragraphs 会
    整块漏掉（实测某审阅稿 24 个 run / 3101 字一个都改不到）。
    """
    return _normalize_doc(doc, _config_from_args(args))


def process_docx(input_file, cfg=None):
    """处理 DOCX 文件。cfg=None 时用模块级 CLI 默认值(向后兼容既有调用方)。"""
    cfg = cfg or _cli_config()
    input_path = Path(input_file)

    if not input_path.exists():
        print(f"❌ 错误：文件不存在 - {input_file}")
        return False

    if input_path.suffix.lower() != ".docx":
        print("❌ 错误：文件必须是.docx格式")
        return False

    # 生成输出文件名（--in-place 时原地写回 + 备份；否则另存 _fixed）
    write_gate = WriteGate(input_path) if cfg.in_place else None  # 读入前 capture 基线
    if cfg.in_place:
        import shutil as _shutil
        from datetime import datetime as _dtt
        bak = input_path.with_name(
            input_path.name + ".bak-" + _dtt.now().strftime("%Y%m%d-%H%M%S"))
        _shutil.copy2(input_path, bak)
        print(f"🗄  已备份: {bak.name}")
        output_path = input_path
    else:
        output_path = input_path.parent / f"{input_path.stem}_fixed{input_path.suffix}"

    try:
        # 读取文档
        print(f"📖 正在读取文件: {input_path.name}")
        doc = Document(input_path)

        # 改 doc 的那件事整体在 _normalize_doc 里（统计信息由它返回，这里只负责喊）
        print(f"🔄 规则: {_fmt_rules(cfg)} | 范围: {_fmt_scopes(cfg)}")
        if cfg.strip_headers:
            print("🗑️  正在删除页眉页脚（--strip-headers）...")
        stats = _normalize_doc(doc, cfg)

        # 保存文件
        print("💾 正在保存文件...")
        if write_gate is not None:
            write_gate.assert_unchanged()  # 源文件被 WPS/其他会话改过 → 拒写(逃生 DOCX_GATE_OK=1)
        try:
            doc.save(output_path)
            clear_quarantine(output_path)
        except (PermissionError, OSError) as e:
            fallback = Path.home() / "Downloads" / Path(output_path).name
            print(f"⚠️ 源目录不可写: {e}")
            print(f"   降级到: {fallback}")
            doc.save(fallback)
            clear_quarantine(fallback)
            output_path = fallback

        print("✅ 处理完成！")
        print(f"   - 共替换了 {stats['quotes']} 个引号")
        print(f"   - 共替换了 {stats['punctuation']} 个标点符号")
        print(f"   - 共转换了 {stats['units']} 个单位")
        scopes = stats["scopes"] or {"正文": 0}
        print("   - 分域命中: " + "  ".join(f"{k} {v} 处" for k, v in scopes.items()))
        print(f"   - 输出文件: {output_path.name}")

        return True

    except Exception as e:
        print(f"❌ 处理失败：{e}")
        import traceback

        traceback.print_exc()
        return False


def text_main(argv):
    """text 子命令主入口（原 docx_text_formatter.py __main__ 块，argv 显式穿参）。

    与旧实现的唯一结构差异：不再改写模块级 DO_* 全局，每次调用从模块默认值
    从头构造冻结 FormatConfig —— 同进程连跑两组配置不串味（docx_cli 连发场景）。
    stdout / 报错文案 / 退出码逐字保留。
    """
    # 摘选择性 flag（在 get_input_files 前），剩余为文件参数
    _argv = list(argv)
    _flags = {"--quotes", "--punct", "--units", "--no-quote-font",
              "--in-place", "--keep-headers", "--strip-headers", "--strip-only"}
    _sel = {a for a in _argv if a in _flags}

    # --scope a,b,c（域白名单；不给 = 全域）
    _scope_arg = None
    _rest = []
    _it = iter(_argv)
    for a in _it:
        if a == "--scope":
            _scope_arg = next(_it, "")
        elif a.startswith("--scope="):
            _scope_arg = a.split("=", 1)[1]
        else:
            _rest.append(a)
    _argv = _rest

    _unknown = [a for a in _argv if a.startswith("--") and a not in _flags]
    if _unknown:   # 未知 flag 不得被当成文件名静默吞掉→曾致 --help 直接跑全量改写
        print(f"❌ 未知参数：{' '.join(_unknown)}")
        print(f"   可用：{' '.join(sorted(_flags))} --scope <{'|'.join(ALL_SCOPES)}>")
        sys.exit(2)
    _argv = [a for a in _argv if a not in _flags]

    scopes = set(ALL_SCOPES)
    if _scope_arg is not None:
        scopes = {s.strip() for s in _scope_arg.split(",") if s.strip()}
        _bad = scopes - set(ALL_SCOPES)
        if _bad:
            print(f"❌ 未知范围：{' '.join(sorted(_bad))}")
            print(f"   可用：{' '.join(f'{k}({v})' for k, v in ALL_SCOPES.items())}")
            sys.exit(2)

    do_quotes, do_punct, do_units = DO_QUOTES, DO_PUNCT, DO_UNITS
    if _sel & {"--quotes", "--punct", "--units"}:
        do_quotes = "--quotes" in _sel
        do_punct = "--punct" in _sel
        do_units = "--units" in _sel
    do_quote_font = "--no-quote-font" not in _sel
    in_place = "--in-place" in _sel
    strip_headers = "--strip-headers" in _sel   # --keep-headers 已是默认行为,保留为兼容 no-op
    if "--strip-only" in _sel:
        # 「清页眉页脚」独立动词:只删 chrome,一个字都不改
        strip_headers = True
        do_quotes = do_punct = do_units = do_quote_font = False

    cfg = FormatConfig(
        quotes=do_quotes, punct=do_punct, units=do_units, quote_font=do_quote_font,
        scopes=frozenset(scopes), strip_headers=strip_headers, in_place=in_place)

    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(_argv, expected_ext="docx")

    if not files:
        print("❌ 错误：缺少文件名参数")
        print("\n使用方法：")
        print("  1. 在 Finder 中选中 .docx 文件，然后运行此脚本")
        print("  2. 或在命令行中提供文件路径")
        print("\n示例：")
        print("    python3 docx_text_formatter.py 文件名.docx")
        print("    python3 docx_text_formatter.py file1.docx file2.docx")
        sys.exit(1)

    print("=" * 50)
    print("文本格式自动修复工具 - DOCX版本")
    print("=" * 50)

    tracker = ProgressTracker()

    for file_path in files:
        print(f"\n处理文件: {Path(file_path).name}")
        success = process_docx(str(file_path), cfg)
        if success:
            tracker.add_success()
        else:
            tracker.add_error()

    print("\n" + "=" * 50)
    tracker.show_summary("文件处理")


# ════════════════════════════════════════════════════════════════════════════
# 首 token 分发
# ════════════════════════════════════════════════════════════════════════════

USAGE = """docx_fmt.py — docx 版式/字体/文本规范化族

子命令:
  template [apply|cleanup] <docx> [-t 模板.docx] [-o out.docx] [--config rules.json] [--preview]
                                        套模板样式 + 样式清理（原 docx_apply_template.py）
  clone    {extract|apply} ...          提取/复刻版式（原 docx_format_clone.py）
  fonts    <docx...> [--check|--apply] [--font F] [--ascii A]
                                        去等线字体归一（原 docx_font_normalize.py）
  text     [--quotes|--punct|--units|--no-quote-font|--in-place|--strip-headers|--strip-only]
           [--scope a,b,c] <docx...>    文本规范化（原 docx_text_formatter.py）
"""


def main(argv=None) -> int:
    argv = list(sys.argv[1:] if argv is None else argv)
    if not argv:
        print(USAGE)
        return 1
    if argv[0] in ("-h", "--help"):
        print(USAGE)
        return 0
    cmd, rest = argv[0], argv[1:]
    if cmd == "template":
        rc = template_main(rest)
        return 0 if rc is None else rc
    if cmd == "clone":
        return clone_main(rest)
    if cmd == "fonts":
        return fonts_main(rest)
    if cmd == "text":
        rc = text_main(rest)
        return 0 if rc is None else rc
    print(USAGE)
    print(f"未知子命令: {cmd}", file=sys.stderr)
    return 2


if __name__ == "__main__":
    sys.exit(main())
