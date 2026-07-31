#!/usr/bin/env python3
"""typeset_apply — spec(yaml) 驱动的排版引擎：按固定顺序把各动作施加到同一份 docx 上。

    python3 typeset_apply.py <docx> --spec <spec.yaml> [--dry-run] [--no-backup] [--report x.json]
    python3 typeset_apply.py --list            # 动作表 + 每条为什么排在这个位置
    python3 typeset_apply.py --dump-schema     # 由 ACTIONS 派生出 spec 的全量字段说明

## 为什么有这个（和 typeset_pipeline.py 的分工）

`typeset_pipeline.py` 是 **subprocess driver**：每步 fork 一个脚本，各自开文件、各自存盘。
10 步 = 开合 10 次 = **10 次重打包机会**。python-docx 每存一次就重写约 60 个部件
（实测 301 部件的真报告，语义只变 1 个），十次下来是 600 次「Word 可能渲染得不一样」。

本引擎把这条路收成：**parse → 在同一棵内存树上顺序施加各动作 → save**（详见下面「三段」）。
炸开面从 N×60 收到两次存盘的量，再由 `lib/docx_safe_save` 的 surgical 收口压到接近 1。

## 为什么不直接用 pipeline_lib.run_pipeline（(a) 薄层 vs (b) 自带执行的取舍）

本引擎**复用** pipeline_lib 的 `load_step` / `lsof_check` / `make_backup_path` ——
动态加载、占用检查、备份命名这三件事它已经做对了，重写就是造轮子。
但「跑」这一步没有走 `run_pipeline`，三条具体原因（不是「重写更省事」）：

1. **Namespace 隔离**：`run_pipeline` 对全部 step 共用**同一个扁平 Namespace**，于是
   `set_table_borders` 的 `color` 会一并递给 `normalize_fonts`，两个动作撞同名选项时
   谁也不会报错、只会互相串味。本引擎给每个动作单独构造 Namespace（见 `_ns`），
   换来「选项跨动作不泄漏」+「未知选项能报错」—— 那正是 spec 这一层的全部意义。
2. **fail-closed 语义相反**：`run_pipeline` 把 step 异常吞成 `{"error": ...}` 并**继续跑后面的
   step、照常存盘**。排版链上一步炸了继续跑，产出的是一份「一半新一半旧」的交付件，
   比报错难查得多。本引擎是一炸就停，doc 阶段炸了**不存盘**（原件保持进来的样子）。
3. **两段不够用**：`run_pipeline` 只有 doc→path 两段，且 doc 段恒在 path 段之前。
   `normalize_fonts` 必须跑在 `freeze_*`（path）**之后**，在两段模型里表达不出来
   —— 见下面「三段」。

所以结论是 (b)，但只自带**执行循环**这一层，加载/备份/lsof 仍然消费总部的 pipeline_lib。

## 三段：doc-pre → save → path → 重开 → doc-post

2026-07-30 从两段改成三段。起因：`normalize_fonts` / `set_table_borders` /
`set_table_align` 这三个脚本被补上了 `apply(doc, args)`（原来只有 `apply_path`），
`load_step` 优先取 `apply` → 加载到的 kind 变成 doc，与本表声明的 path 对不上，
引擎 fail-closed 直接 exit 2，整条 spec 跑不了。修法有两条：

  下策：让 ACTIONS 的 kind 具备强制力（声明 path 就绑 `apply_path`）。
        跑得通，但 `normalize_fonts`/`set_table_borders`/`set_table_align` 各自
        开一次文件、存一次盘 = **3 次多余的重打包**，正是本引擎要消灭的东西。
  采用：加一个 doc-post 段。path 段跑完**重开一次** doc，把这三个当 doc 动作跑完
        再存一次。文件开合从 3 次降到 1 次（重开这次本来 `apply_path` 内部也要开），
        同时 `normalize_fonts` 仍然稳稳排在 `freeze_*` 之后。

实际落位：只有 `normalize_fonts` 真需要 doc-post（它必须看见 freeze 生成的编号 run）；
`set_table_borders` / `set_table_align` 与文本轴无依赖，放 doc-pre 是**白捡的**
——不启用 `normalize_fonts` 时连重开都省了。

## 顺序为什么写死在引擎里

排版动作之间有真依赖（先冻结域再改字体、先配对题注再编号……），顺序排错不会报错，
只会**安静地产出错的版面**。所以 spec 只控制「跑不跑」和参数，**顺序由 ACTIONS 表定死**，
每条的理由写在 `why_here` 里。用户想改顺序 = 改引擎并留下理由，不是改配置。

## fail-closed 的几处

- spec 里出现引擎不认识的键（顶层 / 动作名 / 动作内选项）→ 报错退出。
  静默忽略配置项是最难查的一类 bug：用户以为配了，其实一直没生效。
- spec 不存在 / yaml 解析失败 / 一个动作都没启用 → 非 0 退出，拒绝在空 spec 上报成功。
- 动作声明的 kind 与实际加载到的接口对不上（脚本被改过）→ 报错，不猜。
- 任一动作抛异常 → 立刻停，后面的不跑。doc-pre 出错时**不存盘**，原件保持不变。
- ACTIONS 表自身也自检（phase/kind 是否自洽、每个选项是否都写了说明）→ import 期就炸。

退出码：0 成功 / 1 spec 有问题（不认识的键、空 spec、找不到文件）/ 2 运行期故障
        （docx 被占用、动作报错、动作加载失败）
"""
from __future__ import annotations

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[2] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py

import argparse  # noqa: E402
import json  # noqa: E402
import shutil  # noqa: E402
import sys  # noqa: E402
import time  # noqa: E402
from dataclasses import dataclass, field  # noqa: E402
from pathlib import Path  # noqa: E402
from typing import Any  # noqa: E402

HERE = Path(__file__).resolve().parent
SUB = HERE / "sub"
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))

import yaml  # noqa: E402
from docx import Document  # noqa: E402

# 只复用 pipeline_lib 的加载/占用/备份三件事；不走 run_pipeline 的三条理由见模块 docstring。
from sub.pipeline_lib import load_step, lsof_check, make_backup_path  # noqa: E402


# ──────────────────────────────── 动作表 ────────────────────────────────

# 动作脚本住在哪个目录。多数在 sub/，少数（早于 sub/ 分层就存在的）在 scripts/document/ 根。
HOMES = {"sub": SUB, "root": HERE}

# 执行阶段。doc-pre 与 doc-post 的接口都是 apply(doc,args)，差别只在跑在存盘/path 段的哪一侧。
# path-pre 为什么存在（2026-07-30 加）：restyle / sync_toc 这类「照 golden 对齐」的动作
# 是**按段落文本匹配**去找对应段的。排版链里 convert_chapter_format / renumber_headings /
# number_captions / fix_superscript_refs 全都会改文本，跑完之后目标件的文本已经和 golden
# 对不上，匹配率塌到接近零 —— 而且它不报错，只是安静地「没什么可移植」。
# 所以这类动作必须排在**整条链之前**，它们又是 zip 级的（path kind），doc-pre 之前没有
# 可放的位置，故新开一段。别为了省一个 phase 把它们塞进 path：那等于让它们必然失效。
PHASES = ("path-pre", "doc-pre", "path", "doc-post")


@dataclass(frozen=True)
class Action:
    name: str                     # spec 里写的动作名 = 脚本名（无 .py）
    kind: str                     # 接口：doc = apply(doc,args) / path = apply_path(path,args)
    phase: str                    # 执行阶段，见 PHASES
    what: str                     # 干什么（人读）
    why_here: str                 # 为什么排在这个位置（顺序的理由，不是复述干了啥）
    opts: dict[str, Any] = field(default_factory=dict)      # 认的选项 → 默认值
    opt_docs: dict[str, str] = field(default_factory=dict)  # 选项 → 人读说明（--dump-schema 用）
    home: str = "sub"             # 脚本住哪，见 HOMES
    unavailable: str = ""         # 非空 = 本仓当前跑不了，值是原因
    readonly: bool = False        # True = 只读自检，不改树 → 全是只读时跳过存盘（见 run()）
    count_keys: tuple[str, ...] = ()   # 该脚本用什么键报「改了多少」。空 = 用通用的 "changed"。
                                  # 为什么要声明而不是猜：全仓 49 处用 changed，但 strip_* 一族
                                  # 用的是领域名（bookmarks_removed / outlinelvl_removed）。
                                  # 靠启发式去猜哪个键是「主计数」，猜错时的表现是**报个错数**，
                                  # 比不报还糟。声明在表里 = 猜不着的时候必然报「键对不上」。
    tail_path: str = ""           # 非空 = 本动作在 path 段还有「另一半」(同模块的 apply_path)，
                                  # 值 = 为什么那一半必须跟着跑。两半绑定，spec 里只有一个开关。


# 顺序 = 列表顺序（同一 phase 内）。三段的分界见模块 docstring「三段」。
ACTIONS: list[Action] = [
    # ── 〇 · path-pre：照 golden 对齐。按段落文本匹配，故必须跑在改文本的动作之前 ──
    Action(
        "restyle", "path", "path-pre",
        "从同源 golden 整段克隆段落/字符格式（含媒体段兜底补 pStyle）",
        "整条链的第一步。它拿目标件的段落文本去 golden 里找对应段，找到才移植格式；"
        "而后面 convert_chapter_format / renumber_headings / number_captions / "
        "fix_superscript_refs 每一个都会改文本。排在它们后面 = 匹配率塌到接近零，"
        "而且**不报错**，只是安静地说「无需修改」。不给 ref 时 noop 并如实报 skipped。",
        opts={"restyle_ref": None},
        opt_docs={"restyle_ref": "同源 golden docx 路径（格式源）；null = 本动作跳过"},
    ),
    Action(
        "sync_toc", "path", "path-pre",
        "从同源 golden 移植目录块 + 对账目录样式（TOC 1/2/3 等）",
        "同 restyle，也是照 golden 对账，必须在改文本之前。排在 restyle 之后："
        "先把段落格式对齐，目录样式对账才是在同一套样式表上做的。返回值里带 "
        "anchors_resolved/anchors_total —— 目录锚点解析不到时 Word 里看着完全正常、"
        "只是点不动，这个数不进报告就没人会发现。",
        opts={"sync_toc_ref": None},
        opt_docs={"sync_toc_ref": "同源 golden docx 路径（目录源）；null = 本动作跳过"},
    ),

    # ── 一 · doc-pre：parse 一次，下面这些顺序施加在同一棵内存树上 ────────────
    Action(
        "strip_revisions", "doc", "doc-pre",
        "接受/清除正文里的审阅修订标记（w:ins/w:del/w:moveFrom/w:moveTo）",
        "必须最先。带修订标记时 doc.paragraphs 读到的是「显示态」而非最终文本，"
        "后面所有按文本匹配的动作（章号解析、题注识别）拿到的都是半成品。",
        opts={"revision_mode": "accept-all", "keep_comments": False},
        count_keys=("ins_accepted", "del_accepted", "comment_reference_removed"),
        opt_docs={"revision_mode": "修订处理模式；脚本 choices 当前只有 accept-all"
                                  "（ins 展开 + del 删除）",
                  "keep_comments": "true = 保留 word/comments.xml 里的批注定义。"
                                   "默认 false：连同正文锚点一起清干净"},
        # 【Q4 原先的答案是「不做」，2026-07-30 对抗核验之后改了 —— 那个取舍是错的。】
        #
        # 原答案：本引擎只走 apply(doc)（清正文锚点），清空 comments.xml 那半边在
        # apply_path 里、doc 阶段够不着；为一个 bool 改 Action 结构不值。
        #
        # 实测打脸：跑完 report-generic.yaml，成品里 **正文锚点 0 个、comments.xml 里
        # 还躺着 2 条完整批注**（作者 dh9304、时间戳 2025-12-08、正文是内部核算口径）。
        # Word 里看不见（没锚点），`bid_residue_scan` 前后都报 FAIL 75、一个字都没变——
        # **没有任何闸门抓得到它**，随包发给业主。
        #
        # 所以这不是「少配一个 bool」，是「清修订」这个动作被拆掉了一半还宣称做完了。
        # 用户在 spec 里写 strip_revisions，意思绝不可能是「清正文但把批注定义留着」。
        # 两半必须绑定 → tail_path。多出来的结构成本（一个字段 + path 段一个循环）
        # 远小于「交付件里藏着真名真时间戳」的代价。
        tail_path="它清的是 comments.xml + settings.xml 的 trackChanges，Document 对象够不着，"
                  "只能在 path 段改 zip。与正文锚点是同一件事的两半：只跑前一半 = 留下"
                  "Word 里看不见、闸门也抓不到的孤儿批注（真作者账号 + 真时间戳）。",
    ),
    Action(
        "strip_bookmarks", "doc", "doc-pre",
        "清 _Toc/_Ref/_Hlk/_GoBack 这类残留书签锚点",
        "要在任何搬段/删段动作之前。书签是 bookmarkStart/End 成对的空元素，"
        "段被搬走或删掉之后它们就变成对不上的孤儿，再清就得处理孤儿分支。",
        opts={"bookmark_prefixes": None},
        count_keys=("bookmarks_removed",),
        opt_docs={"bookmark_prefixes": "要清的书签名前缀；逗号串或 list，"
                                       "\"all\"=清所有，_GoBack 恒清；null=脚本默认 _Toc,_Ref,_Hlk"},
    ),
    Action(
        "strip_empty_captions", "doc", "doc-pre",
        "删掉内容为空的题注段（模板残留的空「表名/图名」壳子）",
        "必须在 pair_table_captions / number_captions 之前：空题注段会被配对算法当成"
        "候选表名去抢配对，也会被编号动作编上一个指向空气的「表 X-Y」。先删干净，"
        "后面两步看到的才是真题注集合。",
        opts={},
    ),
    Action(
        "delete_empty_h1", "doc", "doc-pre",
        "删掉空的一级标题段（并把段上挂的 sectPr 迁到前一段，保住分节）",
        "所有「按段落 index 生成计划」的动作都必须排在它后面 —— 它一删段，"
        "先前记下的 index 全部漂移。",
        opts={},
    ),
    Action(
        "relocate_orphan_blocks", "doc", "doc-pre",
        "按外部 plan JSON 把游离块搬回它该在的章下",
        "结构定型的最后一步，必须在标题编号之前：搬完之后章的先后次序才是最终的，"
        "编号才有意义。没给 relocate_plan 时它自己 noop。",
        opts={"relocate_plan": None},
        opt_docs={"relocate_plan": "游离块搬迁 plan JSON 的路径；null 时脚本自己 noop"},
    ),
    Action(
        "convert_chapter_format", "doc", "doc-pre",
        "H1 章号格式统一（「第X章 Y」/「X、Y」→「N Y」）",
        "必须在 renumber_headings 之前：重编号解析的是转换之后的形态，"
        "顺序反了就等于拿旧格式去套新编号规则。",
        opts={},
    ),
    Action(
        "renumber_headings", "doc", "doc-pre",
        "按层级重排标题编号",
        "必须在 number_captions 之前：题注号「图 X-Y」的 X 取自所在章的章号，"
        "章号还没定就编题注号 = 编出来的号下一步就作废。",
        opts={"h1_base": 1},
        opt_docs={"h1_base": "首个 H1 的目标编号；整数，或 \"auto\"=从文档现状探测"
                             "（标书从第 10 章起编那种）"},
        # 【调查者 Q2：h1_base 在 CLI 有、在 apply() 通路上丢了，要不要接上？】
        # 已接。这不是新增业务逻辑，是补一处**抽取时漏搬的**参数：main() 里一直是
        # `plan_renumber(doc, h1_base)`，apply() 抽出来时写成了 `plan_renumber(doc)`。
        # 不接的话本引擎只能提供一个「引擎悄悄按 1 起编」的动作 —— 而「配了没生效」
        # 正是这一层最该防的失败模式，宁可补上也不留一个够不着的旋钮。
    ),
    Action(
        "strip_outlinelvl_from_captions", "doc", "doc-pre",
        "摘掉题注段上被误设的 outlineLvl（段级直接格式）",
        "要在题注配对/编号之前。带 outlineLvl 的题注段会被当成大纲标题，"
        "配对与编号两步都按「这是不是标题」分流，先清干净它们看到的才是真题注集合。",
        opts={},
        count_keys=("outlinelvl_removed",),
    ),
    Action(
        "pair_table_captions", "doc", "doc-pre",
        "按外部 decision JSON 把表与表名配对（调整相对位置）",
        "必须在 number_captions 之前 —— 没配好就编号 = 把号写到错的段上，"
        "之后再纠正就要二次改写正文文本。没给 pair_decision 时它自己 noop。",
        opts={"pair_decision": None},
        opt_docs={"pair_decision": "表↔表名配对 decision JSON 的路径。这是「你先给判断结果」"
                                   "而不是配置值（脚本自己猜不准）；null=noop"},
    ),
    Action(
        "number_captions", "doc", "doc-pre",
        "给表名/图名段补「表 X-Y」/「图 X-Y」章节序号",
        "排在配对之后、字体统一之前：它会**新增文本 run**，那些新 run 得让后面的"
        "字体统一盖到，否则新补的号是继承字体、和正文不一致。",
        opts={},
    ),
    Action(
        "fix_superscript_refs", "doc", "doc-pre",
        "把 [1] / [3-4] 这类文献引用拆成独立 run 并设上标",
        "排在编号类动作之后、字体统一之前：它同样是**拆 run + 新增 run**，"
        "得让 doc-post 的字体统一盖到。放在 number_captions 之后是因为题注号里的"
        "方括号数字不该被误当成文献引用 —— 编完号再扫，看到的是最终文本。",
        opts={},
        home="root",
    ),
    Action(
        "docx_apply_image_caption", "doc", "doc-pre",
        "给图片段与它下一行的图名段套图名样式，并在图名后补空行",
        "必须在 normalize_fonts 之前：它改的是**段落样式**，字体统一要按最终样式"
        "去判「这段算正文还是标题」，样式换在后面就等于按旧身份统一了字体。"
        "放在 number_captions 之后：图名文本先定下来再套样式，避免样式套到"
        "随后又被改写的段上。样式名走模糊匹配，匹配不到时它 noop 并如实报"
        "style_resolved=null，不静默假装干过。",
        opts={"style_name": "ZDWP图名"},
        opt_docs={"style_name": "套给图片段与图名段的段落样式名；模糊匹配，"
                                "\"ZDWP 图名\" 这类空格变体也命中。文档里没有这个样式 → 整步 noop"},
        home="root",
        # 【调查者 Q3：图片这一节是空的，要不要这轮收编？】收了，位置按调查者建议
        # （number_captions 与 add_header_footer 之间）。center_images 仍不收：它走 lxml
        # 直接读 zip，apply / apply_path 两个接口都没有，收编要先给它补接口。
    ),
    Action(
        "add_header_footer", "doc", "doc-pre",
        "设置页眉页脚（含页码域）",
        "doc-pre 的最后一步：它按 doc.sections 遍历，而前面的删段/搬块都可能改变节结构"
        "（delete_empty_h1 里就有 sectPr 迁移），节稳定下来再写页眉页脚。",
        opts={
            "header": "",
            "footer_prefix": "",
            "page_number": True,
            "font_size": 10.0,
            "gap_spaces": 13,
            "header_align": "right",
            "footer_align": "center",
        },
        opt_docs={
            "header": "页眉文字；空串 = 不写字（CLI 里它是 required，spec 里默认空）",
            "footer_prefix": "页脚前缀，通常是院名；空串时 apply() 内部兜底"
                             "「浙江省水利水电勘测设计院」",
            "page_number": "页脚是否插 PAGE 域",
            "font_size": "页眉页脚字号 pt",
            "gap_spaces": "院名与页码之间的空格数",
            "header_align": "页眉对齐 left/center/right",
            "footer_align": "页脚对齐 left/center/right",
        },
    ),
    Action(
        "set_table_borders", "doc", "doc-pre",
        "统一表格框线（线型 / 磅值 / 颜色 / 间距）",
        "表格框线与文本轴互不依赖，所以它待在 doc-pre 而不是 doc-post —— 白捡一次"
        "不用重开文件。（它同时有 apply 和 apply_path，本引擎按声明的 doc 走内存版。）",
        opts={"val": "single", "sz": 4, "color": "auto", "space": 0,
              "keep_cell_borders": False},
        opt_docs={
            "val": "框线线型 single/double/dashed/none…",
            "sz": "框线磅值（1/8 pt 为单位，4 = 0.5pt）",
            "color": "框线颜色 hex 或 auto(=黑)",
            "space": "框线边距",
            "keep_cell_borders": "true = 保留单元格自带 tcBorders，只把非实线改成实线；"
                                 "false = 直接删 tcBorders 让表级框线生效",
        },
    ),
    Action(
        "set_table_align", "doc", "doc-pre",
        "统一表格对齐（表级居中，可选单元格内文字也居中）",
        "同 set_table_borders，与文本轴无依赖，留在 doc-pre。",
        opts={"cell_center": False},
        opt_docs={"cell_center": "true = 单元格内文字水平+垂直也居中；"
                                 "表级 jc=center 是脚本写死的，没有 left/right 可选"},
    ),

    # ── doc-pre 收尾自检（只读，不改树）─────────────────────────────
    # 它们报的是 **doc 阶段**的终态，不含 path 阶段（冻域 / 摘样式 outlineLvl）的效果。
    # 没把它们挪进 doc-post 是因为：只读动作单独触发一次「重开一整份 docx」不划算，
    # 而它们盯的三样（题注段级 outlineLvl / 表配对 / 书签）全部在 doc 阶段定型。
    # 唯一的近似：audit_caption_outline 看不到 strip_style_outlinelvl 摘掉的样式级 outlineLvl。
    Action(
        "audit_caption_outline", "doc", "doc-pre",
        "只读自检：题注段的 outlineLvl / 样式污染还剩多少",
        "放在所有 doc 改动之后 —— 自检要报的是终态，不是中间态。",
        opts={}, readonly=True,
    ),
    Action(
        "audit_table_pairing", "doc", "doc-pre",
        "只读自检：表与表名的配对情况（孤儿表名 / 孤儿表 / 重名）",
        "同上，终态自检。",
        opts={}, readonly=True,
    ),
    Action(
        "audit_bookmarks", "doc", "doc-pre",
        "只读自检：书签总数与孤儿 start/end",
        "同上，终态自检。",
        opts={}, readonly=True,
    ),

    # ── 二 · path：存盘之后逐个改 zip 里 document.xml 以外的部件 ──────────────
    Action(
        "strip_doc_protection", "path", "path",
        "清 settings.xml 里的文档保护位（只读保护 / 编辑限制）",
        "path 段第一个：后面几步都要改 styles.xml / document.xml，先把保护摘掉"
        "语义上更顺（技术上 python-docx 不受保护位约束，摘不摘都能写）。",
        opts={},
    ),
    Action(
        "freeze_heading_numbers", "path", "path",
        "把标题的 Word 自动编号固化成真文本，并摘掉 styles.xml 里的 numPr",
        "必须在 normalize_fonts 之前。域没冻结时编号是 Word 渲染期算出来的，"
        "document.xml 里根本没有承载它的 run，字体统一无从下手；冻结之后它们才是"
        "普通 run，字体才盖得住。",
        opts={"freeze_levels": None, "unlink_style": True, "num_id": "4"},
        opt_docs={
            "freeze_levels": "冻结哪几级标题的自动编号；逗号串或 list，元素只能是 1/2/3/4；"
                             "null = 全部四级",
            "unlink_style": "是否同时摘掉 styles.xml 里标题样式的 numPr",
            "num_id": "从 numbering.xml 的哪个 numId 取 lvlText 模板",
        },
    ),
    Action(
        "freeze_all_fields", "path", "path",
        "冻结其余各类域（TOC / PAGE / SEQ / REF …）成静态文本",
        "同上，也必须在字体统一之前。放在 freeze_heading_numbers 之后："
        "标题编号是专门处理（还要动 styles.xml 的 numPr），先做专项再做通扫，"
        "反过来会让通扫把标题域也冻成普通文本、专项那步就找不到目标了。",
        opts={"freeze_types": "all"},
        opt_docs={"freeze_types": "冻结哪些域类型；\"all\" 或 \"TOC,PAGEREF,SEQ\" 这样的逗号串"},
    ),
    Action(
        "strip_style_outlinelvl", "path", "path",
        "在 styles.xml 里摘掉题注**样式**上的 outlineLvl",
        "和 doc 阶段那步 strip_outlinelvl_from_captions 是**两处**：那步摘段级直接格式，"
        "这步摘样式定义。只做一处的话，另一处会继续把题注推进大纲。样式表要在 doc 存盘"
        "之后改，否则 python-docx 存盘会把它盖回去。",
        opts={"strip_styles": None},
        opt_docs={"strip_styles": "要从 styles.xml 摘 outlineLvl 的样式名集合，逗号串；"
                                  "null = 脚本内置的默认题注样式集"},
    ),
    Action(
        "center_images", "path", "path",
        "含图片的段显式居中 + 缩进清零",
        "排在 path 段末尾。它写的是**段级直接格式**（jc/ind），直接格式在 Word 里压过"
        "样式，所以任何改样式的动作（docx_apply_image_caption 套图名样式、restyle 克隆"
        "格式）都得先跑完，否则它写的居中会被后面的样式变更盖住语义 —— 表现是「有的图"
        "居中有的没有」，且不报错。",
        opts={},
    ),
    Action(
        "line_spacing", "path", "path",
        "对照 golden 给正文段补固定行距（只补没有显式行距的段）",
        "排在 center_images 之后、doc-post 之前。它按段落文本去 golden 取行距值，"
        "所以理论上也怕文本被改 —— 但它只认**正文段**且逐段独立，题注/标题编号的改动"
        "波及不到，实测放这里比放 path-pre 好：放前面的话，此后新增的段（题注号补出来的"
        "那些）永远拿不到行距。不给 ref 时 noop 并如实报 skipped。",
        opts={"line_spacing_ref": None},
        opt_docs={"line_spacing_ref": "同源 golden docx 路径（行距值取其众数）；null = 本动作跳过"},
    ),

    # ── 三 · doc-post：path 段改完重开一次 doc，跑「必须看见冻结结果」的动作 ────
    Action(
        "normalize_fonts", "doc", "doc-post",
        "全文字体统一（正文中文 / 标题中文 / 西文 / 标题色）",
        "整条链最后一步，也是**唯一**要跑在 path 段之后的 doc 动作：它得盖到所有"
        "新增文本 —— number_captions 补的题注号、fix_superscript_refs 拆出的引用 run、"
        "以及 freeze_* 冻出来的编号文本。冻结产物只在 path 段落盘之后才存在于 "
        "document.xml 里，所以必须重开一次 doc 才看得见。顺序反了，成品里会看到"
        "零星几处字体不一样，而且不报任何错。",
        opts={
            "body_cjk": "宋体",
            "heading_cjk": "黑体",
            "latin": "Times New Roman",
            "heading_color": "000000",
        },
        opt_docs={
            "body_cjk": "正文中文字体（等线系在 spec 层就被 BANNED_FONTS 拦下）",
            "heading_cjk": "标题中文字体",
            "latin": "中西文混排里的西文字体",
            "heading_color": "标题字体颜色 hex",
        },
    ),

    # ── 本仓当前跑不了的：注明原因，被启用时明确报错，不静默跳过 ──────────
    Action(
        "fix_heading_disorder", "doc", "doc-pre",
        "修标题层级乱序（假晋级 / 级差跳跃 / 相邻重复）",
        "本该排在 convert_chapter_format 之前（层级先正确，章号格式才好统一）。",
        opts={},
        unavailable="import 了 apply_body_styles —— 那是 qual-supply 的脚本，"
                    "doctools 仓里没有这个模块，加载即 ModuleNotFoundError。"
                    "本引擎不替它补依赖（那是改业务逻辑）；要用先把该模块搬进来。",
    ),
    Action(
        "reorder_heading_blocks", "doc", "doc-pre",
        "按标题块重排/去重整块内容",
        "本该紧跟 fix_heading_disorder。",
        opts={},
        unavailable="它 from fix_heading_disorder import …，连坐上面那条同样的缺失依赖。",
    ),
]

BY_NAME: dict[str, Action] = {a.name: a for a in ACTIONS}

# ACTIONS 表自检：表是本引擎唯一的事实源，它自己错了下面全错，所以 import 期就炸。
for _a in ACTIONS:
    if _a.phase not in PHASES:
        raise SystemExit(f"[ACTIONS 表错] {_a.name}: phase={_a.phase!r} 不在 {PHASES}")
    if (_a.phase in ("path", "path-pre")) != (_a.kind == "path"):
        raise SystemExit(
            f"[ACTIONS 表错] {_a.name}: phase={_a.phase} 与 kind={_a.kind} 不自洽 —— "
            f"path 段只能放 apply_path 动作，doc-pre/doc-post 只能放 apply 动作"
        )
    if _a.home not in HOMES:
        raise SystemExit(f"[ACTIONS 表错] {_a.name}: home={_a.home!r} 不在 {sorted(HOMES)}")
    _undoc = sorted(set(_a.opts) - set(_a.opt_docs))
    if _undoc:
        # 强制每个选项都有人读说明：--dump-schema 派生出的 schema 文件全靠它，
        # 漏一个就等于给用户一个「不知道是干嘛的」旋钮。
        raise SystemExit(f"[ACTIONS 表错] {_a.name}: 选项 {_undoc} 没写 opt_docs")
    _stray = sorted(set(_a.opt_docs) - set(_a.opts))
    if _stray:
        raise SystemExit(f"[ACTIONS 表错] {_a.name}: opt_docs 里有不存在的选项 {_stray}")

# 顶层允许的键。多一个都报错 —— 顶层拼错（actions 写成 action）最容易表现为
# 「配了但一个动作都没跑」，报错比静默跑空强。
#
# 【调查者 Q5：七件事要不要做成 yaml 层级（页眉页脚: / 表格: / …）？】不做，actions 保持扁平。
#   分节在语义上不承载任何执行含义（顺序由 ACTIONS 定死），多一层就多一处能拼错的地方，
#   还要给 TOP_KEYS 和 load_spec 加一层展平。分节改用注释表达 —— 见 docx_spec.schema.yaml。
# 【调查者 Q6：要不要支持顶层 profile: zdwp 让所有动作继承？】不做。
#   它会打破「选项不跨动作泄漏」这条刻意设计（见 _ns 的注释）——一旦开了这个口子，
#   下一个「反正大家都用」的选项就会照着往顶层放，最后又回到扁平共用 Namespace。
#   而且当前收编的动作里**一个都不吃 profile**（吃它的 outline / fix_styleset 都没收编，
#   见下面「显式不收编」），等真出现第二个吃 profile 的动作再说。
TOP_KEYS = {"version", "actions"}

# 全局 CLAUDE.md #12：docx 中文字体只用宋体系/仿宋系，标题可黑体/方正小标宋，禁等线。
# 在 spec 层就拦掉，别等成品发出去才发现整本是等线。
BANNED_FONTS = ("等线", "DengXian", "Dengxian", "dengxian")
FONT_OPT_KEYS = {"body_cjk", "heading_cjk"}

# 显式不收编的动作，写在这里是为了让「为什么没有它」有答案，不用下次再查一遍：
#   delete_chapter / delete_table_rows —— 不可逆的删章删行。CLAUDE.md 立过
#   「破坏性动作必须自己占一个动词」：把它们混进排版 spec，等于让一份「调版面」的
#   配置顺手删内容，正是那条规矩要防的反模式。要删章就单独敲 sub/delete_chapter.py
#   （delete_table_rows 的值：--table-index / --rows FROM:TO / --expected-first-col
#    / --expected-residue / --expected-last，后三个是删前后的安全校验）。
#   center_images / line_spacing —— 走 lxml 直接读 zip，既没有 apply 也没有 apply_path，
#   本引擎不给它们开 subprocess 兜底（一开就退回「多开合一次」，正好抵消本引擎的收益）。
#   需要它们的场景走 typeset_pipeline.py，那本来就是 subprocess driver。
#   outline —— 一个模块三个互斥子动作（promote-h1 / demote-h2 / normalize-arabic），
#   缺 outline_cmd 直接 ValueError；且它与已收编的 renumber_headings / convert_chapter_format
#   在「谁来定标题层级与编号」上直接重叠，两个都启用等于让 spec 自己跟自己打架。
#   它还吃 --profile，见上面 Q6。真要用单独敲 sub/outline.py。
#   fix_styleset —— 实测在本引擎的加载路径下 `ImportError: attempted relative import
#   with no known parent package`（它 import 了同包兄弟模块），load_step 根本加载不到。
#   收编前得先把它的包内相对 import 改成可独立加载 —— 那是改它的业务代码。
#   docx_text_formatter —— 有 apply(doc,args) 了，但它的能力是「4 个规则开关 + 一条域白名单」，
#   规则本身（引号/标点/单位替换表）硬编码在 lib/text_fixes.py。收编它得先决定
#   scope 白名单怎么在 spec 里表达（它是唯一一个选项不是标量的动作），本轮不动。


# ──────────────────────────────── spec 解析 ────────────────────────────────

class SpecError(Exception):
    """spec 本身有问题（不认识的键 / 空 / 解析失败）→ 退出码 1。"""


def load_spec(path: Path) -> dict[str, dict]:
    """读 spec，返回 {动作名: 已合并默认值的 opts}，顺序无关（顺序由 ACTIONS 定）。

    任何不认识的键都抛 SpecError —— 这是本函数存在的主要理由，不是顺手做的校验。
    """
    if not path.is_file():
        raise SpecError(f"spec 不存在: {path}")
    try:
        raw = yaml.safe_load(path.read_text(encoding="utf-8"))
    except yaml.YAMLError as exc:
        raise SpecError(f"spec 不是合法 yaml: {path}\n{exc}") from exc
    if raw is None:
        raise SpecError(f"spec 是空文件: {path}")
    if not isinstance(raw, dict):
        raise SpecError(f"spec 顶层必须是映射，实际是 {type(raw).__name__}: {path}")

    unknown_top = sorted(set(raw) - TOP_KEYS)
    if unknown_top:
        raise SpecError(
            f"spec 顶层有不认识的键: {unknown_top}；只认 {sorted(TOP_KEYS)}"
        )

    actions_raw = raw.get("actions")
    if not isinstance(actions_raw, dict) or not actions_raw:
        raise SpecError("spec 缺 actions:（或它是空的）—— 拒绝在空 spec 上报成功")

    enabled: dict[str, dict] = {}
    for name, cfg in actions_raw.items():
        act = BY_NAME.get(name)
        if act is None:
            raise SpecError(
                f"不认识的动作 '{name}'。可用动作：\n  " +
                "\n  ".join(f"{a.name:32s} {a.what}" for a in ACTIONS)
            )
        if cfg is False or cfg is None:
            continue                      # 显式关掉 / 留个占位，都是不跑
        if cfg is True:
            cfg = {}
        if not isinstance(cfg, dict):
            raise SpecError(
                f"动作 '{name}' 的值必须是 true/false 或选项映射，实际是 "
                f"{type(cfg).__name__}"
            )
        cfg = dict(cfg)
        if cfg.pop("enabled", True) is False:
            continue
        unknown_opt = sorted(set(cfg) - set(act.opts))
        if unknown_opt:
            raise SpecError(
                f"动作 '{name}' 有不认识的选项 {unknown_opt}；"
                f"它只认 {sorted(act.opts) or '（无选项）'}"
            )
        for k in FONT_OPT_KEYS & set(cfg):
            if any(b in str(cfg[k]) for b in BANNED_FONTS):
                raise SpecError(
                    f"动作 '{name}' 的 {k}={cfg[k]!r} 用了等线系字体 —— "
                    f"docx 中文只用宋体系/仿宋系，标题可黑体/方正小标宋"
                )
        merged = dict(act.opts)
        merged.update(cfg)
        enabled[name] = merged

    if not enabled:
        raise SpecError("spec 里一个动作都没启用 —— 拒绝在空 spec 上报成功")

    dead = [n for n in enabled if BY_NAME[n].unavailable]
    if dead:
        msg = "\n".join(f"  {n}: {BY_NAME[n].unavailable}" for n in dead)
        raise SpecError(f"下面这些动作本仓当前跑不了，不能启用：\n{msg}")

    return enabled


# ──────────────────────────────── 执行 ────────────────────────────────

def _ns(opts: dict, *, docx: Path, dry_run: bool) -> argparse.Namespace:
    """把一个动作的 opts dict 变成它 apply() 认的 Namespace。

    各脚本的 apply/apply_path 一律 `getattr(args, "xxx", 默认)` 取参，所以这里
    必须给 Namespace 而不是 dict（dict 没有属性访问，全部 getattr 会静默落到默认值 ——
    那就是「配了没生效」的经典死法）。
    每个动作一个独立 Namespace：选项不跨动作泄漏。
    docx/dry_run/no_backup 是引擎注入的公共字段：
      no_backup 恒 True —— 引擎已经统一备份过一次，各动作不许再自己备份，
      否则跑一趟排版会在交付目录里堆一排 .bak-N-日期.docx。
    """
    return argparse.Namespace(docx=docx, dry_run=dry_run, no_backup=True,
                              report=None, **opts)


def run(docx: Path, spec: dict[str, dict], *, dry_run: bool,
        no_backup: bool) -> tuple[dict, int]:
    """按 ACTIONS 的固定顺序执行 spec 启用的动作。返回 (报告, 退出码)。"""
    t0 = time.perf_counter()
    plan = [a for a in ACTIONS if a.name in spec]
    report: dict[str, Any] = {
        "docx": str(docx),
        "dry_run": dry_run,
        "顺序": [a.name for a in plan],
        "分段": {p: [a.name for a in plan if a.phase == p] for p in PHASES},
        "改了什么": [],
        "计数": {},
        "actions": {},
        "timing": {},
    }

    # 先把全部动作加载起来再动文件（fail fast）：加载失败发生在改文件之前，
    # 原件一定还是干净的。
    loaded = {}
    for a in plan:
        try:
            st = load_step(a.name, step_dir=HOMES[a.home])
        except Exception as exc:
            report["error"] = f"动作 '{a.name}' 加载失败: {type(exc).__name__}: {exc}"
            return report, 2
        if st.kind != a.kind:
            # 脚本被改过（apply ←→ apply_path 换了），顺序假设可能已经不成立，不猜。
            report["error"] = (
                f"动作 '{a.name}' 声明 kind={a.kind}，实际加载到 kind={st.kind} —— "
                f"脚本接口变了，请同步 typeset_apply.py 的 ACTIONS 表"
            )
            return report, 2
        loaded[a.name] = st

    # 占用检查只在真要写盘时做。纯只读的 spec（全是 audit_*）不碰文件，却会被
    # Spotlight 的 mdworker_ 之类的读句柄判成「被占用」而 rc=2 —— 实测就是这么挂的。
    # 「Word 开着时想跑个只读体检」是完全正当的用法，不该被拦。
    writes = any(not a.readonly for a in plan)
    if not dry_run and writes:
        occ = lsof_check(docx)
        if occ:
            report["error"] = f"docx 被 Word/WPS 占用，先关掉:\n{occ}"
            return report, 2

    # 纯只读的 spec 不备份：备份是为了「改坏了能回退」，一个字节都不改时它只是在
    # 交付目录里多丢一份 59MB 垃圾。judge 用上面那个 writes（基于整个 plan，
    # 因为 path 段和 tail 也会真写盘）。
    if not dry_run and not no_backup and writes:
        bak = make_backup_path(docx)
        shutil.copy2(docx, bak)
        report["backup"] = str(bak)

    pathpre_plan = [a for a in plan if a.phase == "path-pre"]
    pre_plan = [a for a in plan if a.phase == "doc-pre"]
    path_plan = [a for a in plan if a.phase == "path"]
    post_plan = [a for a in plan if a.phase == "doc-post"]

    def _record(a: Action, rep: dict) -> None:
        report["actions"][a.name] = rep
        if a.count_keys:
            present = [k for k in a.count_keys if isinstance(rep.get(k), (int, float))]
            changed = sum(rep[k] for k in present) if present else None
        else:
            changed = rep.get("changed") if isinstance(rep.get("changed"), (int, float)) else None

        if changed is not None:
            report["计数"][a.name] = changed
            why = f"（{rep['skipped']}）" if rep.get("skipped") else ""
            report["改了什么"].append(f"{a.name}: {a.what} → changed={changed}{why}")
        elif a.readonly:
            report["改了什么"].append(f"{a.name}: {a.what} → 只读自检，无计数")
        else:
            # 会写盘的动作却报不出计数 —— 从前这里和「只读」走同一条分支，于是
            # strip_bookmarks 删掉 77 个书签、摘要上写「只读/无 changed 计数」。
            # 报告说不出自己改了多少，就等于没有报告。这里判红，不静默降级。
            keys = sorted(k for k, v in rep.items() if isinstance(v, (int, float)))
            report.setdefault("计数键缺失", []).append(
                {"action": a.name, "declared": list(a.count_keys) or ["changed"],
                 "rep 里的数值键": keys})
            report["改了什么"].append(
                f"⚠ {a.name}: {a.what} → **计数键对不上**："
                f"表里声明 {list(a.count_keys) or ['changed']}，脚本实际返回 {keys}")

    def _run_doc_phase(doc, actions: list[Action], *, phase_label: str) -> str | None:
        """在同一棵内存树上顺序施加 actions。返回 None = 全过，否则返回错误说明。"""
        for a in actions:
            t = time.perf_counter()
            try:
                rep = loaded[a.name].call(doc=doc, args=_ns(spec[a.name], docx=docx,
                                                            dry_run=dry_run))
            except Exception as exc:
                report["actions"][a.name] = {"error": f"{type(exc).__name__}: {exc}"}
                return f"动作 '{a.name}'（{phase_label}）抛异常"
            report["timing"][f"step:{a.name}"] = round(time.perf_counter() - t, 3)
            _record(a, rep if isinstance(rep, dict) else {"raw": repr(rep)})
        return None

    def _run_path_phase(actions: list[Action], tails: list[Action]) -> str | None:
        """逐个跑 path-kind 动作（直接改 zip）。path-pre 与 path 两段共用这一份实现 ——
        抄第二份的那天，两段的错误处理就会开始分叉。返回 None = 全过。"""
        for a in actions:
            is_tail = a in tails
            key = f"{a.name}::path" if is_tail else a.name
            fn = (getattr(loaded[a.name].module, "apply_path", None) if is_tail
                  else loaded[a.name].call)
            if is_tail and fn is None:
                report["error"] = (
                    f"动作 '{a.name}' 声明了 tail_path，但它的模块里没有 apply_path —— "
                    f"脚本接口变了，请同步 ACTIONS 表"
                )
                return report["error"]
            t = time.perf_counter()
            try:
                args = _ns(spec[a.name], docx=docx, dry_run=dry_run)
                rep = fn(docx, args) if is_tail else fn(docx_path=docx, args=args)
            except Exception as exc:
                report["actions"][key] = {"error": f"{type(exc).__name__}: {exc}"}
                report["error"] = (
                    f"动作 '{key}' 抛异常，已中止。注意 path 阶段在存盘之后，"
                    f"此前的动作已经落盘 —— 要回退用备份 {report.get('backup', '(没备份)')}"
                )
                return report["error"]
            report["timing"][f"step:{key}"] = round(time.perf_counter() - t, 3)
            report["actions"][key] = rep if isinstance(rep, dict) else {"raw": repr(rep)}
            if not is_tail:
                _record(a, rep if isinstance(rep, dict) else {"raw": repr(rep)})
            else:
                report["改了什么"].append(f"{key}: {a.what} 的另一半（comments/settings 等部件）")
        return None

    # ⓪ path-pre：照 golden 对齐格式。必须在任何改文本的动作之前跑，理由见 PHASES 注释。
    #    dry-run 时它们的 apply_path 自己认 dry_run（本轮接入时逐个加的），所以照跑。
    if pathpre_plan:
        err = _run_path_phase(pathpre_plan, [])
        if err:
            return report, 2

    # ① doc-pre：parse 一次，全部动作跑在同一棵树上，再 save 一次。
    doc = None
    if pre_plan:
        t = time.perf_counter()
        doc = Document(str(docx))
        report["timing"]["parse"] = round(time.perf_counter() - t, 3)
        err = _run_doc_phase(doc, pre_plan, phase_label="doc-pre")
        if err:
            # 中途炸了就不存盘：原件保持进来时的样子，比存一个半成品好。
            report["error"] = f"{err}，已中止且**未存盘**（原件未变）"
            return report, 2
        # 全是只读自检时不存盘（2026-07-30 对抗核验发现）：原来只要 pre_plan 非空就无条件
        # save，于是「跑一份纯 audit 的 spec 体检一下」会把 59MB 交付件整包重写一遍、
        # 顺手在交付目录里落一个同样 59MB 的 .bak，mtime 也变。语义上没坏（收口救了），
        # 但「只读体检」这件事本来就不该碰文件 —— 何况 lsof 门还会让「Word 开着时想体检」
        # 直接 rc=2。
        if not dry_run and not all(a.readonly for a in pre_plan):
            t = time.perf_counter()
            doc.save(str(docx))
            report["timing"]["save"] = round(time.perf_counter() - t, 3)
        elif all(a.readonly for a in pre_plan):
            report["改了什么"].append("doc-pre 全是只读自检 → 跳过存盘（文件一个字节没动）")

    # ② path：这些动作要改 styles.xml / settings.xml / numbering.xml 等 document.xml
    #    之外的部件，只能在 doc-pre 存盘之后逐个改 zip —— 放在存盘之前，python-docx
    #    的存盘会把它们刚写进去的内容整个盖掉。
    #    dry-run 时它们读到的是**未经 doc-pre 修改**的文件，所以报的数是对原件的估计，
    #    不等于真跑一遍的结果。这一条是 dry-run 的固有近似，不是 bug。
    #    tail：某些 doc 动作还有「另一半」在同模块的 apply_path 里（见 Action.tail_path）。
    #    两半是同一件事，spec 里只有一个开关，所以这里跟着跑，不给用户漏掉的机会。
    #    dry-run 不跑 tail：`strip_revisions.apply_path` **没有 dry_run 守卫**，无条件
    #    `shutil.move(tmp, docx)`（2026-07-30 实测：加上 tail 之后 --dry-run 的 md5 变了）。
    #    在引擎里挡住，而不是去改那个脚本 —— 它的 main() 自己在调用前就判了 dry-run，
    #    改它等于改一个跑得好好的脚本的业务逻辑。tail 是本引擎新加的调用路径，
    #    这条路径上的安全由本引擎负责。
    tails = [a for a in (pre_plan + post_plan) if a.tail_path]
    if dry_run and tails:
        report["改了什么"].append(
            "dry-run：跳过 " + "、".join(f"{a.name}::path" for a in tails)
            + "（它们的 apply_path 不认 dry_run，跑了就会真写盘）")
        tails = []
    err = _run_path_phase(tails + path_plan, tails)
    if err:
        return report, 2

    # ③ doc-post：重开一次 doc，跑「必须看见 path 阶段产物」的动作（当前只有 normalize_fonts）。
    if post_plan:
        if dry_run:
            # dry-run 里 doc-pre 没存盘、path 段也没写盘，重开等于把 doc-pre 的改动全丢掉，
            # 报出来的数会比真跑少（新补的题注号、拆出的引用 run 都不在）。
            # 所以 dry-run 沿用同一棵内存树：宁可只近似在「看不到 path 段产物」这一处，
            # 也不要两处都近似。没有 doc-pre 时才退回读盘。
            doc2 = doc if doc is not None else Document(str(docx))
        else:
            t = time.perf_counter()
            doc2 = Document(str(docx))
            report["timing"]["parse2"] = round(time.perf_counter() - t, 3)
        err = _run_doc_phase(doc2, post_plan, phase_label="doc-post")
        if err:
            report["error"] = (
                f"{err}，已中止。注意 doc-post 在 path 阶段之后，此前的改动都已落盘 —— "
                f"要回退用备份 {report.get('backup', '(没备份)')}"
            )
            return report, 2
        if not dry_run and not all(a.readonly for a in post_plan):
            t = time.perf_counter()
            doc2.save(str(docx))
            report["timing"]["save2"] = round(time.perf_counter() - t, 3)

    report["timing"]["total"] = round(time.perf_counter() - t0, 3)
    if report.get("计数键缺失"):
        # 文件已经改好了（rc 不代表交付件坏了），但报告说不出某个动作改了多少 ——
        # 那正是「守卫哑掉」的形状：跑完全绿、摘要上一行「无计数」，没人会去查。
        # 全表 29 条实测 0 缺失，所以这里判红不会误伤存量；将来加动作忘了声明
        # count_keys，第一次跑就会被拦下。
        report["error"] = ("有动作报不出计数（见「计数键缺失」）—— 文件已改，但报告不完整。"
                           "去 ACTIONS 表给它声明 count_keys。")
        return report, 1
    return report, 0


# ─────────────────────────── schema 派生（SSOT = ACTIONS）───────────────────────

PHASE_TITLE = {
    "path-pre": "〇 · path-pre —— 在任何动作之前，按 golden 参照件对齐格式（zip 级）",
    "doc-pre":  "一 · doc-pre —— parse 一次，下面这些按顺序施加在同一棵内存树上，然后存盘",
    "path":     "二 · path —— 存盘之后逐个改 zip 里 document.xml 以外的部件",
    "doc-post": "三 · doc-post —— 重开一次 doc，跑「必须看见冻结结果」的动作，再存一次",
}


def _yaml_scalar(v: Any) -> str:
    """把默认值渲染成 yaml 字面量。中文字符串不加引号（yaml 不需要），
    空串/纯数字串/hex 串必须加，否则读回来类型就变了。"""
    if v is None:
        return "null"
    if isinstance(v, bool):
        return "true" if v else "false"
    if isinstance(v, (int, float)):
        return str(v)
    s = str(v)
    if s == "" or s.strip() != s or s.replace(".", "", 1).isdigit() or s.startswith("0"):
        return f'"{s}"'
    return s


def dump_schema() -> str:
    """由 ACTIONS 派生 spec 的全量字段说明（= 一份把所有动作都写出来的示例 spec）。

    之所以派生而不是手写：动作表改了、手写的 schema 不会跟着改，两处必然漂开，
    而「文档说有这个选项、引擎其实不认」正是 spec 层最要命的一类谎。
    """
    L: list[str] = []
    L.append("# docx_spec.schema.yaml —— spec(yaml) 的全量字段说明")
    L.append("#")
    L.append("# ⚠️ 本文件由引擎派生，**别手改**。改动作/选项请改 ACTIONS 表再重新生成：")
    L.append("#     cd ~/Dev/tools/doctools && python3 scripts/document/typeset_apply.py \\")
    L.append("#         --dump-schema > config/docx_spec.schema.yaml")
    L.append("#")
    L.append("# 用法：")
    L.append("#     python3 scripts/document/typeset_apply.py <docx> --spec <spec.yaml> [--dry-run]")
    L.append("#")
    L.append("# 顶层只认 " + " / ".join(sorted(TOP_KEYS)) + " —— 多一个键就报错退出")
    L.append("#（顶层拼错最容易表现为「配了但一个动作都没跑」，报错比静默跑空强）。")
    L.append("# actions 下每个键 = 一个动作；值 = false/null（不跑）| true（跑，全默认）| 选项映射。")
    L.append("# 选项映射里还认一个 enabled: false，等价于整条关掉（便于保留参数不删）。")
    L.append("#")
    L.append("# **顺序不在这里控**：动作之间有真依赖（先冻域再改字体、先配对题注再编号），")
    L.append("# 顺序写死在引擎 ACTIONS 表里，每条带「为什么在这个位置」。spec 只管")
    L.append("# 「跑不跑」和「参数是什么」。分节是注释不是 yaml 层级 —— actions 是扁平映射，")
    L.append("# 套一层 `页眉页脚:` 会被判「不认识的动作」。")
    L.append("#")
    L.append("# 本文件把**所有可用动作全部列出**，直接拿去用等于「全都跑」。真实交付建议照抄")
    L.append("# config/spec-examples/report-generic.yaml 再按需增删。")
    L.append("")
    L.append("version: 1")
    L.append("")
    L.append("actions:")
    for phase in PHASES:
        acts = [a for a in ACTIONS if a.phase == phase and not a.unavailable]
        if not acts:
            continue
        L.append("")
        L.append(f"  # ══ {PHASE_TITLE[phase]}")
        for a in acts:
            L.append("")
            L.append(f"  # {a.name} —— {a.what}")
            wrapped = _wrap("为什么在这个位置：" + a.why_here, 76)
            L.append(f"  #   {wrapped[0]}")
            for line in wrapped[1:]:
                L.append(f"  #     {line}")
            if a.home != "sub":
                L.append(f"  #   脚本位置：scripts/document/{a.name}.py（不在 sub/ 下）")
            if not a.opts:
                L.append(f"  {a.name}: true".ljust(40) + "# 无选项，只有跑/不跑")
                continue
            L.append(f"  {a.name}:")
            for k, v in a.opts.items():
                for line in _wrap(a.opt_docs[k], 70):
                    L.append(f"    # {line}")
                L.append(f"    {k}: {_yaml_scalar(v)}")
    dead = [a for a in ACTIONS if a.unavailable]
    if dead:
        L.append("")
        L.append("# ── 本仓当前跑不了的动作：写进 actions: 会被判红并退出 ──────────────────")
        for a in dead:
            L.append(f"#   {a.name} —— {a.what}")
            for line in _wrap(a.unavailable, 76):
                L.append(f"#       {line}")
    L.append("")
    return "\n".join(L)


# 收尾标点不另起一行：中文按列宽硬折时，「。」「）」落到下一行开头很扎眼。
_NO_LEAD = set("。，、；：）」』】》!?！？.,;:)]}")


def _wrap(text: str, width: int) -> list[str]:
    """按显示宽度折行（中文按 2 列算）。不用 textwrap 是因为它按字符数折，
    中文行会短一半、看起来参差不齐。"""
    out, cur, w = [], "", 0
    for ch in text:
        cw = 2 if ord(ch) > 0x2E7F else 1
        if w + cw > width and cur and ch not in _NO_LEAD:
            out.append(cur)
            cur, w = "", 0
        cur += ch
        w += cw
    if cur:
        out.append(cur)
    return out or [""]


# ──────────────────────────────── CLI ────────────────────────────────

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="spec 驱动的排版引擎：按固定顺序把各动作施加到同一份 docx 上",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="动作清单与顺序理由：python3 typeset_apply.py --list",
    )
    ap.add_argument("docx", nargs="?", type=Path, help="目标 docx（原地改，默认先备份）")
    ap.add_argument("--spec", type=Path, help="spec.yaml 路径")
    ap.add_argument("--dry-run", action="store_true", help="不存盘，只报会改什么")
    ap.add_argument("--no-backup", action="store_true", help="跳过 .bak-N-日期 备份")
    ap.add_argument("--report", type=Path, default=None, help="把完整报告写成 JSON")
    ap.add_argument("--list", action="store_true", help="列出全部动作、执行顺序与理由")
    ap.add_argument("--dump-schema", action="store_true",
                    help="把 spec 的全量字段说明打到 stdout（由 ACTIONS 派生）")
    a = ap.parse_args(argv)

    if a.dump_schema:
        print(dump_schema())
        return 0

    if a.list:
        for phase in PHASES:
            acts = [x for x in ACTIONS if x.phase == phase]
            if not acts:
                continue
            print(f"\n══ {PHASE_TITLE[phase]}")
            for i, act in enumerate(acts, 1):
                flag = f"\n      ⛔ 不可用：{act.unavailable}" if act.unavailable else ""
                print(f"{i:2d}. [{act.kind:4s}] {act.name}\n"
                      f"      做什么 : {act.what}\n"
                      f"      为什么在这个位置 : {act.why_here}\n"
                      f"      选项   : {act.opts or '（无）'}" + flag)
        return 0

    if a.docx is None or a.spec is None:
        ap.error("需要 <docx> 与 --spec（或只用 --list / --dump-schema）")
    if not a.docx.is_file():
        print(f"⛔ docx 不存在: {a.docx}", file=sys.stderr)
        return 1

    try:
        spec = load_spec(a.spec)
    except SpecError as exc:
        print(f"⛔ spec 有问题：{exc}", file=sys.stderr)
        return 1

    docx = a.docx.resolve()
    report, rc = run(docx, spec, dry_run=a.dry_run, no_backup=a.no_backup)
    report["spec"] = str(a.spec.resolve())

    print(f"=== typeset_apply · {docx.name} ===")
    print(f"spec      : {a.spec}")
    for phase in PHASES:
        names = report["分段"].get(phase) or []
        if names:
            print(f"{phase:9s}: {' → '.join(names)}")
    if report.get("backup"):
        print(f"备份      : {report['backup']}")
    for line in report["改了什么"]:
        print(f"  · {line}")
    if report.get("error"):
        print(f"⛔ {report['error']}", file=sys.stderr)
    tm = report.get("timing", {})
    if tm:
        print(f"耗时      : parse={tm.get('parse', 0)}s save={tm.get('save', 0)}s "
              f"parse2={tm.get('parse2', 0)}s save2={tm.get('save2', 0)}s "
              f"total={tm.get('total', 0)}s")

    if a.report:
        a.report.parent.mkdir(parents=True, exist_ok=True)
        a.report.write_text(json.dumps(report, ensure_ascii=False, indent=2),
                            encoding="utf-8")
        print(f"[report] {a.report}")
    return rc


if __name__ == "__main__":
    sys.exit(main())
