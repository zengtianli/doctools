#!/usr/bin/env python3
"""
md_tools.py - Markdown 工具集

将 6 个独立的 Markdown 处理脚本合并为一个统一入口。

子命令:
    format      文本格式自动修复（引号/标点/单位）
    merge       合并多个 Markdown 文件为一个
    split       按一级标题拆分 Markdown 文件
    strip       删除 Markdown 文件中所有 blockquote
    to-html     Markdown 渲染为 HTML 并在浏览器中打开
    frontmatter 批量生成 YAML frontmatter（LLM）
    md2docx     Markdown 转 Docx（样式复刻版，原 md_docx_template.py）

用法:
    python3 md_tools.py <subcommand> [args...]
    python3 md_tools.py format file.md
    python3 md_tools.py merge a.md b.md
    python3 md_tools.py split input.md
    python3 md_tools.py strip md_final/ --fix
    python3 md_tools.py to-html file.md
    python3 md_tools.py frontmatter docs/ --dry-run
    python3 md_tools.py md2docx input.md -t 模板.docx -o output.docx
"""

import argparse
import json
import os
import re
import subprocess
import sys
import tempfile
import time
import webbrowser
import zipfile
from concurrent.futures import ThreadPoolExecutor
from copy import deepcopy
from pathlib import Path

# ── lib path setup ──────────────────────────────────────────────────
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "lib"))
sys.path.insert(0, str(Path.home() / "Dev" / "tools" / "dev" / "lib"))  # canonical 5 modules
from display import show_error, show_info, show_processing, show_success, show_warning
from file_ops import (
    check_file_extension,
    clear_quarantine,
    fatal_error,
    validate_input_file,
)
from finder import get_input_files
from progress import ProgressTracker
from text_fixes import fix_punctuation, fix_quotes, fix_units

# ── md2docx（样式复刻版）依赖：原 md_docx_template.py 迁入 ──────────────
try:
    from lxml import etree
except ImportError:
    print("❌ 需要安装 lxml: pip install lxml")
    sys.exit(1)

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT  # noqa: F401
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import nsmap, qn  # noqa: F401
    from docx.shared import Cm, Pt, Twips  # noqa: F401
except ImportError:
    print("❌ 需要安装 python-docx: pip install python-docx")
    sys.exit(1)

import contextlib

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[2] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py

from docx_parts import DEFAULT_ALLOW_CHANGED, assert_parts_intact  # noqa: E402
from docx_xml import NSMAP  # noqa: E402

# ── version info ────────────────────────────────────────────────────
SCRIPT_VERSION = "3.0.0"
SCRIPT_AUTHOR = "tianli"
SCRIPT_UPDATED = "2026-03-25"


def format_process_file(input_file, rules=("quotes", "punct", "units")):
    """处理单个文件的格式修复。rules = 要应用的规则子集(与 docx 引擎口径一致)。"""
    input_path = Path(input_file)

    if not input_path.exists():
        show_error(f"文件不存在 - {input_file}")
        return False

    output_path = input_path.parent / f"{input_path.stem}_fixed{input_path.suffix}"

    try:
        show_processing(f"正在读取文件: {input_path.name}")
        with open(input_path, encoding="utf-8") as f:
            content = f.read()

        fixed_content = content
        quote_count = punct_count = unit_count = 0
        if "quotes" in rules:
            show_processing("正在处理引号...")
            fixed_content, quote_count, _ = fix_quotes(fixed_content)
        if "punct" in rules:
            show_processing("正在处理标点符号...")
            fixed_content, punct_count = fix_punctuation(fixed_content)
        if "units" in rules:
            show_processing("正在转换单位...")
            fixed_content, unit_count = fix_units(fixed_content)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(fixed_content)

        show_success("处理完成!")
        print(f"   - 共替换了 {quote_count} 个引号")
        print(f"   - 共替换了 {punct_count} 个标点符号")
        print(f"   - 共转换了 {unit_count} 个单位")
        print(f"   - 输出文件: {output_path.name}")

        return True

    except Exception as e:
        show_error(f"处理失败: {e}")
        return False


def cmd_format(args):
    """format 子命令入口"""
    sel = tuple(r for r in ("quotes", "punct", "units") if getattr(args, r, False))
    rules = sel or ("quotes", "punct", "units")   # 一个都没给 = 全做(向后兼容)
    files = args.files
    if not files:
        files = get_input_files([], expected_ext="md")

    if not files:
        fatal_error(
            "缺少文件名参数\n\n"
            "使用方法:\n"
            "    python3 md_tools.py format 文件名.md\n"
            "    python3 md_tools.py format file1.md file2.md\n"
            "    或在 Finder 中选择 .md 文件后运行"
        )

    print("=" * 50)
    print("文本格式自动修复工具")
    print("=" * 50)

    tracker = ProgressTracker()

    # 纯本地 I/O，文件间独立 → 文件级 ThreadPoolExecutor 并发（max_workers=8）。
    # map 保序，逐文件独立读写各自 *_fixed.md，无共享状态。
    def _one(file_path) -> bool:
        print(f"\n处理文件: {Path(file_path).name}")
        try:
            return format_process_file(str(file_path), rules)
        except Exception as e:  # 失败隔离：单文件崩不拖垮整批
            show_error(f"处理失败: {e}")
            return False

    if len(files) == 1:
        results = [_one(files[0])]
    else:
        workers = max(1, min(8, len(files)))
        with ThreadPoolExecutor(max_workers=workers) as ex:
            results = list(ex.map(_one, files))

    for success in results:
        if success:
            tracker.add_success()
        else:
            tracker.add_error()

    print("\n" + "=" * 50)
    tracker.show_summary("文件处理")


# ════════════════════════════════════════════════════════════════════
#  merge - 合并多个 Markdown 文件
# ════════════════════════════════════════════════════════════════════

def merge_md_files(md_files: list, output_file: Path):
    """合并多个 Markdown 文件"""
    tracker = ProgressTracker()

    md_files = sorted(md_files)

    show_info(f"准备合并 {len(md_files)} 个 Markdown 文件")

    try:
        with open(output_file, "w", encoding="utf-8") as f_out:
            for i, md_file in enumerate(md_files, 1):
                md_path = Path(md_file)

                if not validate_input_file(md_path):
                    tracker.add_skip()
                    continue

                if not check_file_extension(md_path, "md"):
                    show_warning(f"跳过非 Markdown 文件: {md_path.name}")
                    tracker.add_skip()
                    continue

                show_info(f"处理 ({i}/{len(md_files)}): {md_path.name}")

                with open(md_path, encoding="utf-8") as f_in:
                    content = f_in.read()
                    f_out.write(content)
                    if i < len(md_files):
                        f_out.write("\n\n")

                tracker.add_success()

        show_success(f"合并完成，已保存为: {output_file.name}")
        tracker.show_summary("文件合并")

    except Exception as e:
        fatal_error(f"合并失败: {e}")


def cmd_merge(args):
    """merge 子命令入口"""
    files = args.files
    if not files:
        files = get_input_files([], expected_ext="md")

    if not files:
        fatal_error("请提供至少一个 Markdown 文件，或在 Finder 中选择文件后运行")

    # 判断最后一个参数是否为输出文件
    if args.output:
        output_file = Path(args.output)
        md_files = files
    elif files[-1].endswith(".md") and len(files) > 1 and not Path(files[-1]).exists():
        output_file = Path(files[-1])
        md_files = files[:-1]
    else:
        first_file = Path(files[0]).resolve()
        output_file = first_file.parent / "merged.md"
        md_files = files

    merge_md_files(md_files, output_file)


# ════════════════════════════════════════════════════════════════════
#  split - 按一级标题拆分 Markdown 文件
# ════════════════════════════════════════════════════════════════════

def split_slugify(title: str) -> str:
    """将标题转换为文件名友好格式"""
    title = re.sub(r"^#*\s*\d*\.?\s*", "", title)
    slug = title.lower().strip()
    slug = re.sub(r"[^\w\s-]", "", slug)
    slug = re.sub(r"[\s-]+", "_", slug)
    return slug[:50]


def split_markdown(input_path: Path) -> list[tuple[str, str]]:
    """
    拆分 Markdown 文件
    返回: [(filename, content), ...]
    """
    content = input_path.read_text(encoding="utf-8")

    pattern = r"^(# .+)$"
    parts = re.split(pattern, content, flags=re.MULTILINE)

    results = []
    idx = 0

    # 第一部分: # 之前的内容（标题、摘要等）
    if parts[0].strip():
        results.append((f"{idx:02d}_title.md", parts[0].strip()))
        idx += 1

    # 后续部分: 成对出现（标题, 内容）
    for i in range(1, len(parts), 2):
        if i + 1 < len(parts):
            heading = parts[i]
            body = parts[i + 1] if i + 1 < len(parts) else ""

            slug = split_slugify(heading)
            filename = f"{idx:02d}_{slug}.md"

            full_content = f"{heading}\n{body}".strip()
            results.append((filename, full_content))
            idx += 1

    return results


def cmd_split(args):
    """split 子命令入口"""
    input_file = args.input
    if not input_file:
        files = get_input_files([], expected_ext="md")
        if files:
            input_file = files[0]

    if not input_file:
        fatal_error("请提供一个 Markdown 文件，或在 Finder 中选择文件后运行")

    input_path = Path(input_file).resolve()

    if not validate_input_file(input_path):
        sys.exit(1)
    if not check_file_extension(input_path, "md"):
        fatal_error(f"不是 Markdown 文件: {input_path.name}")

    # 输出目录
    output_dir = input_path.parent / f"{input_path.stem}_split"
    output_dir.mkdir(exist_ok=True)

    parts = split_markdown(input_path)

    show_info(f"输出目录: {output_dir}")
    show_info(f"拆分为 {len(parts)} 个文件:")

    for filename, content in parts:
        output_path = output_dir / filename
        output_path.write_text(content, encoding="utf-8")
        lines = content.count("\n") + 1
        show_success(f"{filename} ({lines} 行)")

    show_success("拆分完成!")


# ════════════════════════════════════════════════════════════════════
#  strip - 删除 Markdown 文件中所有 blockquote
# ════════════════════════════════════════════════════════════════════

def strip_blockquotes(text: str) -> str:
    """删除所有 blockquote 行，并清理残留的连续空行"""
    lines = text.split("\n")
    result = []
    in_code = False

    for line in lines:
        stripped = line.strip()

        # 跟踪代码块
        if stripped.startswith("```"):
            in_code = not in_code
            result.append(line)
            continue

        # 代码块内不处理
        if in_code:
            result.append(line)
            continue

        # 跳过 blockquote 行
        if stripped.startswith(">"):
            continue

        result.append(line)

    # 清理连续空行（最多保留 1 个）
    cleaned = []
    prev_empty = False
    for line in result:
        is_empty = line.strip() == ""
        if is_empty and prev_empty:
            continue
        cleaned.append(line)
        prev_empty = is_empty

    return "\n".join(cleaned)


def strip_count_blockquotes(text: str) -> int:
    """统计 blockquote 行数（排除代码块内）"""
    lines = text.split("\n")
    count = 0
    in_code = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        if stripped.startswith(">"):
            count += 1
    return count


def strip_process_file(filepath: Path, do_fix: bool, output_dir: Path | None) -> dict:
    """处理单个文件"""
    text = filepath.read_text(encoding="utf-8")
    bq_count = strip_count_blockquotes(text)

    if bq_count > 0:
        print(f"  {filepath.name}: {bq_count} 行 blockquote")
    else:
        print(f"  {filepath.name}: 无 blockquote")

    if do_fix and bq_count > 0:
        fixed_text = strip_blockquotes(text)

        if output_dir:
            output_dir.mkdir(parents=True, exist_ok=True)
            out_path = output_dir / filepath.name
        else:
            out_path = filepath

        out_path.write_text(fixed_text, encoding="utf-8")
        show_success(f"已删除 {bq_count} 行 -> {out_path}")

    return {"file": filepath.name, "blockquotes": bq_count}


def cmd_strip(args):
    """strip 子命令入口"""
    input_path = Path(args.input)
    if not input_path.exists():
        show_error(f"路径不存在: {input_path}")
        sys.exit(1)

    if input_path.is_dir():
        md_files = sorted(input_path.glob("*.md"))
        if not md_files:
            show_error(f"目录中没有 .md 文件: {input_path}")
            sys.exit(1)
        show_info(f"发现 {len(md_files)} 个 MD 文件")
    else:
        md_files = [input_path]

    output_dir = Path(args.output_dir) if args.output_dir else None

    all_stats = []
    for f in md_files:
        stats = strip_process_file(f, args.fix, output_dir)
        all_stats.append(stats)

    total_bq = sum(s["blockquotes"] for s in all_stats)
    files_with_bq = sum(1 for s in all_stats if s["blockquotes"] > 0)

    if len(all_stats) > 1:
        print(f"\n合计: {total_bq} 行 blockquote（{files_with_bq}/{len(all_stats)} 个文件）")

    if total_bq > 0 and not args.fix:
        print("\n使用 --fix 执行删除")
        sys.exit(1)


# ════════════════════════════════════════════════════════════════════
#  md2docx 共用件（to-docx pandoc 版 2026-08-04 用户拍板退役：守卫旁路 + 零消费者）
# ════════════════════════════════════════════════════════════════════

# 默认模板路径（doctools SoT）
DEFAULT_TEMPLATE = str(Path(__file__).parent.parent.parent / "templates" / "template.docx")


def get_finder_selection():
    """获取 Finder 选中的文件"""
    script = """
    tell application "Finder"
        set sel to selection
        if (count of sel) > 0 then
            return POSIX path of (item 1 of sel as alias)
        end if
    end tell
    """
    result = subprocess.run(["osascript", "-e", script], capture_output=True, text=True)
    return result.stdout.strip()


# ════════════════════════════════════════════════════════════════════
#  md2docx - Markdown 转 Docx（样式复刻版，原 md_docx_template.py 迁入）
# ════════════════════════════════════════════════════════════════════

# 需要提取的样式 ID
TARGET_STYLES = {
    "a": "Normal",  # 基础样式
    "10": "Heading 1",  # 标题1
    "2": "Heading 2",  # 标题2
    "3": "Heading 3",  # 标题3
    "4": "Heading 4",  # 标题4
    "ZDWP": "ZDWP正文",  # 水利正文
    "ZDWP1": "ZDWP表名",  # 表格标题
    "ZDWP3": "ZDWP表格内容",  # 表格单元格文字
    "ZDWP4": "ZDWP图名",  # 图片标题
}

# 样式依赖关系
STYLE_DEPS = {
    "10": ["2"],  # Heading 1 依赖 Heading 2
    "2": ["a"],  # Heading 2 依赖 Normal
    "3": ["4"],  # Heading 3 依赖 Heading 4
    "4": ["a"],  # Heading 4 依赖 Normal
    "ZDWP": ["a"],  # ZDWP正文 依赖 Normal
    "ZDWP1": ["a"],  # ZDWP表名 依赖 Normal
    "ZDWP3": ["a"],  # ZDWP表格内容 依赖 Normal
    "ZDWP4": ["a"],  # ZDWP图名 依赖 Normal
}

# 默认配置文件路径
DEFAULT_STYLES_DIR = os.path.dirname(os.path.abspath(__file__))


def extract_styles_xml(docx_path, output_dir=None):
    """从 docx 提取样式 XML"""

    if output_dir is None:
        output_dir = os.path.dirname(docx_path) or "."

    print(f"📄 提取样式: {docx_path}")

    # 解压 docx
    with zipfile.ZipFile(docx_path, "r") as zf:
        styles_xml = zf.read("word/styles.xml")

    # 解析 XML
    root = etree.fromstring(styles_xml)

    # 收集需要的样式
    collected_styles = {}
    found_body_style = None

    for style in root.findall(".//w:style", NSMAP):
        style_id = style.get(f"{{{NSMAP['w']}}}styleId")
        style_type = style.get(f"{{{NSMAP['w']}}}type")

        # 只要段落样式
        if style_type != "paragraph":
            continue

        # 获取样式名
        name_elem = style.find("w:name", NSMAP)
        style_name = name_elem.get(f"{{{NSMAP['w']}}}val") if name_elem is not None else ""

        # 检查是否是目标样式
        if style_id in TARGET_STYLES:
            collected_styles[style_id] = {"element": deepcopy(style), "name": style_name, "id": style_id}
            print(f"  ✓ 找到: {style_id} ({style_name})")

            # 记录正文样式
            if style_id == "ZDWP":
                found_body_style = "ZDWP"

        # 如果没有 ZDWP，用 Normal 作为正文
        if style_id == "a" and found_body_style is None:
            found_body_style = "a"

    # 检查必要样式
    required = ["a", "10", "2", "3", "4"]
    missing = [s for s in required if s not in collected_styles]
    if missing:
        print(f"⚠️  缺少样式: {missing}")

    # 提取 docDefaults（默认字体设置）
    doc_defaults = root.find(".//w:docDefaults", NSMAP)

    # 生成精简的 styles.xml
    new_root = etree.Element(f"{{{NSMAP['w']}}}styles", nsmap=NSMAP)

    # 添加 docDefaults
    if doc_defaults is not None:
        new_root.append(deepcopy(doc_defaults))

    # 按依赖顺序添加样式
    added = set()

    def add_style(style_id):
        if style_id in added or style_id not in collected_styles:
            return
        # 先添加依赖
        for dep in STYLE_DEPS.get(style_id, []):
            add_style(dep)
        new_root.append(collected_styles[style_id]["element"])
        added.add(style_id)

    for style_id in collected_styles:
        add_style(style_id)

    # eastAsia 只准中文字体（2026-09-04 修根）
    # 模板 template.docx 的 styles.xml 里混进了 eastAsia="Times New Roman"
    # /"TimesNewRomanPS-BoldMT"/"DengXian" —— eastAsia 槽位放西文字体本身就是错的，
    # 而 DengXian＝等线，全局规则明令禁止出现在中文 docx 里。以前靠出件方每次手工
    # 打补丁改回宋体（ip-legal docx-outfile 规程第 3 项），漏一次就把等线印进递交件。
    # 这里在样式提取时一次性归一，所有下游消费者自动对齐。
    # ⚠ 只动 w:rFonts/@w:eastAsia；w:lang 也有 eastAsia 属性（值是 zh-CN/en-US 这类
    # 语言代码），碰了会把语言标记写成字体名。
    _cjk_ok = ("宋", "仿", "黑", "楷", "圆", "魏", "隶", "方正", "华文", "思源",
               "SimSun", "SimHei", "FangSong", "KaiTi", "STSong", "STFangsong",
               "STKaiti", "STHeiti", "Source Han", "Noto Sans CJK", "Noto Serif CJK",
               "PingFang", "Hiragino Sans GB", "Microsoft YaHei", "微软雅黑")
    _fixed = 0
    for _rf in new_root.iter(f"{{{NSMAP['w']}}}rFonts"):
        _ea = _rf.get(f"{{{NSMAP['w']}}}eastAsia")
        if not _ea:
            continue
        if any(k in _ea for k in _cjk_ok):
            continue
        _rf.set(f"{{{NSMAP['w']}}}eastAsia", "宋体")
        _fixed += 1
    if _fixed:
        print(f"🔤 eastAsia 归一：{_fixed} 处非中文字体（含等线/DengXian）改为宋体")

    # 保存 XML
    styles_xml_path = os.path.join(output_dir, "heading_styles.xml")
    tree = etree.ElementTree(new_root)
    tree.write(styles_xml_path, encoding="UTF-8", xml_declaration=True, pretty_print=True)
    print(f"💾 样式 XML: {styles_xml_path}")

    # 生成人类可读的说明
    info_path = os.path.join(output_dir, "styles_info.txt")
    with open(info_path, "w", encoding="utf-8") as f:
        f.write("=" * 60 + "\n")
        f.write(f"样式提取自: {os.path.basename(docx_path)}\n")
        f.write("=" * 60 + "\n\n")

        f.write("提取的样式:\n")
        for _, info in collected_styles.items():
            f.write(f"  - {info['id']}: {info['name']}\n")

        f.write(f"\n正文样式: {found_body_style}\n")

        f.write("\nMarkdown 映射:\n")
        f.write("  # 标题    → Heading 1\n")
        f.write("  ## 标题   → Heading 2\n")
        f.write("  ### 标题  → Heading 3\n")
        f.write("  #### 标题 → Heading 4\n")
        f.write(f"  普通段落  → {collected_styles.get(found_body_style, {}).get('name', 'Normal')}\n")
        f.write(f"  表x ...   → {collected_styles.get('ZDWP1', {}).get('name', 'Normal')}\n")
        f.write(f"  表格内容  → {collected_styles.get('ZDWP3', {}).get('name', 'Normal')}\n")
        f.write(f"  图x ...   → {collected_styles.get('ZDWP4', {}).get('name', 'Normal')}\n")

    print(f"📝 样式说明: {info_path}")

    # 保存配置
    config = {
        "source": os.path.basename(docx_path),
        "body_style": found_body_style,
        "body_style_name": collected_styles.get(found_body_style, {}).get("name", "Normal"),
        "styles": {k: v["name"] for k, v in collected_styles.items()},
    }
    config_path = os.path.join(output_dir, "styles_config.json")
    with open(config_path, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

    return styles_xml_path, config


def create_docx_with_styles(styles_xml_path, output_path):
    """创建带样式的空白 docx"""

    # 先创建空白文档
    doc = Document()

    # 保存到临时文件
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    doc.save(tmp_path)

    # 读取提取的样式
    with open(styles_xml_path, "rb") as f:
        new_styles = etree.parse(f).getroot()

    # 解压 docx，修改 styles.xml，重新打包
    with tempfile.TemporaryDirectory() as tmpdir:
        # 解压
        with zipfile.ZipFile(tmp_path, "r") as zf:
            zf.extractall(tmpdir)

        # 读取原 styles.xml
        orig_styles_path = os.path.join(tmpdir, "word", "styles.xml")
        with open(orig_styles_path, "rb") as f:
            orig_root = etree.parse(f).getroot()

        # 合并样式：把新样式添加到原样式中
        # 先删除同名样式
        existing_ids = set()
        for style in new_styles.findall(".//w:style", NSMAP):
            style_id = style.get(f"{{{NSMAP['w']}}}styleId")
            existing_ids.add(style_id)

        for style in orig_root.findall(".//w:style", NSMAP):
            style_id = style.get(f"{{{NSMAP['w']}}}styleId")
            if style_id in existing_ids:
                orig_root.remove(style)

        # 添加新样式
        for style in new_styles.findall(".//w:style", NSMAP):
            orig_root.append(deepcopy(style))

        # 更新 docDefaults
        new_defaults = new_styles.find(".//w:docDefaults", NSMAP)
        if new_defaults is not None:
            old_defaults = orig_root.find(".//w:docDefaults", NSMAP)
            if old_defaults is not None:
                orig_root.remove(old_defaults)
            orig_root.insert(0, deepcopy(new_defaults))

        # 写回 styles.xml
        tree = etree.ElementTree(orig_root)
        tree.write(orig_styles_path, encoding="UTF-8", xml_declaration=True)

        # 重新打包
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for root_dir, _dirs, files in os.walk(tmpdir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, tmpdir)
                    zf.write(file_path, arcname)

    clear_quarantine(output_path)

    # 部件完整性断言（fail-closed）：虽是造新文件，但基线 = python-docx 空白模板
    # tmp_path（本函数不增删部件，只换 styles.xml）——对它 diff 恰能抓
    # 「os.walk 重打包截断/漏拷/tmpdir 混入杂文件」。必须在 unlink(tmp_path) 之前比。
    assert_parts_intact(tmp_path, output_path,
                        allow_changed=set(DEFAULT_ALLOW_CHANGED) | {"word/styles.xml"},
                        verbose=False)

    # 清理临时文件
    os.unlink(tmp_path)

    return output_path


def set_table_border(table, border_color="000000", border_size=4):
    """给表格设置边框

    Args:
        table: python-docx Table 对象
        border_color: 边框颜色（十六进制，默认黑色）
        border_size: 边框粗细（单位：1/8磅，4=0.5磅）
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # 创建边框元素
    tblBorders = OxmlElement("w:tblBorders")

    # 六种边框：top, left, bottom, right, insideH, insideV
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(border_size))
        border.set(qn("w:color"), border_color)
        border.set(qn("w:space"), "0")
        tblBorders.append(border)

    # 移除旧边框设置
    old_borders = tblPr.find(qn("w:tblBorders"))
    if old_borders is not None:
        tblPr.remove(old_borders)

    tblPr.append(tblBorders)


def parse_table_row(line):
    """解析表格行，返回单元格列表"""
    # 去掉首尾的 |
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    # 分割单元格
    cells = [cell.strip() for cell in line.split("|")]
    return cells


def is_separator_row(line):
    """检查是否是表格分隔行 |---|---|"""
    line = line.strip()
    if not line.startswith("|"):
        return False
    # 去掉 | 后检查是否只有 - : 空格
    content = line.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
    return len(content) == 0


def clean_markdown_text(text):
    """清理行内噪音(code/math)。**bold**/*italic* 标记保留,由 add_md_runs 渲染成真 Word run。"""
    text = re.sub(r"`(.+?)`", r"\1", text)  # `code`
    text = re.sub(r"\$(.+?)\$", r"\1", text)  # $math$ (简单处理)
    return text


def strip_md_markers(text):
    """整段剥掉 markdown 行内标记(标题/表名/图名本身有样式,不需要 run 级加粗)。"""
    text = clean_markdown_text(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # **bold**
    text = re.sub(r"\*(.+?)\*", r"\1", text)  # *italic*
    return text


# **bold** 或 *italic*(首字符非空格/星号,避免 3*4 这类误伤)
_MD_INLINE = re.compile(r"(\*\*[^*]+?\*\*|\*[^*\s][^*]*?\*)")


def add_md_runs(para, text):
    """把 **bold**/*italic* 渲染成真 Word run —— 星号字面量绝不落进 docx。"""
    for part in _MD_INLINE.split(clean_markdown_text(text)):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            para.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            para.add_run(part[1:-1]).italic = True
        else:
            para.add_run(part)


def parse_list_item(line):
    """解析列表项，返回 (缩进级别, 内容)"""
    # 计算缩进（空格或 tab）
    stripped = line.lstrip()
    indent = len(line) - len(stripped)
    indent_level = indent // 2  # 每 2 空格一级

    # 去掉列表标记
    if (
        stripped.startswith("- ")
        or stripped.startswith("* ")
        and not stripped.startswith("**")
        or stripped.startswith("> ")
    ):
        content = stripped[2:]
    elif re.match(r"^\d+\.\s", stripped):
        # 有序列表：编号是作者信息（书状「请求 1-5」「附件 1-5」），保留字面编号，只规范空格。
        # 2026-08-20 前此处剥掉编号 → 递交件里五条请求全无序号（ip-legal 实测），故改为保留。
        content = re.sub(r"^(\d+)\.\s+", r"\1. ", stripped)
    else:
        content = stripped

    return indent_level, content.strip()


def merge_list_items(items):
    """合并列表项为段落文本

    智能处理：如果前一项以冒号结尾，则不加分号分隔
    """
    if not items:
        return ""
    if len(items) == 1:
        return items[0]

    result = items[0]
    for item in items[1:]:
        # 如果前面以冒号结尾，直接连接
        if result.endswith("：") or result.endswith(":"):
            result += item
        else:
            result += "；" + item

    return result


def parse_markdown(md_content):
    """解析 Markdown，返回元素列表"""
    elements = []
    lines = md_content.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行跳过
        if not stripped:
            i += 1
            continue

        # 分隔线 ---
        if stripped == "---":
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            elements.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # 表名：以 "表" + 数字 开头（允许"表 6.5-1"式空格间隔——中文题注主流写法）
        if re.match(r"^表\s?\d", stripped):
            elements.append({"type": "table_title", "text": stripped})
            i += 1
            continue

        # 图名：以 "图" + 数字 开头（允许"图 6.5-1"式空格间隔）
        if re.match(r"^图\s?\d", stripped):
            elements.append({"type": "figure_title", "text": stripped})
            i += 1
            continue

        # 表格：以 | 开头
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # 解析表格
            if len(table_lines) >= 2:
                headers = []
                rows = []

                # 第一行是表头
                headers = parse_table_row(table_lines[0])

                # 第二行是分隔符，跳过
                # 后面是数据行
                for tl in table_lines[2:]:
                    if not is_separator_row(tl):
                        rows.append(parse_table_row(tl))

                elements.append({"type": "table", "headers": headers, "rows": rows})
            continue

        # 列表项：以 - 、* 、> 或数字. 开头
        # 注意：区分 "* 列表项"（单星号+空格）和 "**加粗**"（双星号，不是列表）
        is_list_start = (
            stripped.startswith("- ")
            or (stripped.startswith("* ") and not stripped.startswith("**"))
            or stripped.startswith("> ")
            or re.match(r"^\d+\.\s", stripped)
        )

        if is_list_start:
            # 收集连续的列表项
            list_items = []
            while i < len(lines):
                current = lines[i]
                current_stripped = current.strip()

                # 空行结束列表
                if not current_stripped:
                    break

                # 检查是否是列表项（包括缩进的子项）
                is_list_item = (
                    current_stripped.startswith("- ")
                    or (current_stripped.startswith("* ") and not current_stripped.startswith("**"))
                    or current_stripped.startswith("> ")
                    or re.match(r"^\d+\.\s", current_stripped)
                    or (current.startswith("  ") and list_items)  # 缩进的续行
                )

                if not is_list_item:
                    break

                indent_level, content = parse_list_item(current)
                if content:  # 只添加非空内容
                    list_items.append((indent_level, content))
                i += 1

            # 把列表项转换为段落
            # 策略：每个列表项作为一个段落，缩进项合并到上一项
            if list_items:
                current_para = []
                for indent_level, content in list_items:
                    clean_content = clean_markdown_text(content)
                    if indent_level == 0:
                        # 新的一级列表项
                        if current_para:
                            # 保存之前的段落
                            elements.append({"type": "paragraph", "text": merge_list_items(current_para)})
                        current_para = [clean_content]
                    else:
                        # 缩进的子项，合并到当前段落
                        current_para.append(clean_content)

                # 保存最后一个段落
                if current_para:
                    elements.append({"type": "paragraph", "text": merge_list_items(current_para)})
            continue

        # 普通段落
        para_lines = [line]
        i += 1
        # 继续收集段落内容，直到遇到特殊行
        while i < len(lines):
            next_line = lines[i]
            next_stripped = next_line.strip()

            # 空行结束段落
            if not next_stripped:
                break

            # 特殊行结束段落
            if (
                re.match(r"^#{1,4}\s", next_line)  # 标题
                or next_stripped.startswith("|")  # 表格
                or next_stripped.startswith("- ")  # 列表
                or (next_stripped.startswith("* ") and not next_stripped.startswith("**"))
                or next_stripped.startswith("> ")
                or re.match(r"^\d+\.\s", next_stripped)
                or next_stripped == "---"
            ):
                break

            para_lines.append(next_line)
            i += 1

        text = " ".join(para_lines)
        text = clean_markdown_text(text)

        if text.strip():  # 只添加非空段落
            elements.append({"type": "paragraph", "text": text})

    return elements


def convert_md_to_docx(md_path, styles_xml_path, output_path, config=None):
    """转换 Markdown 到 Docx"""

    print(f"📖 读取: {md_path}")

    # 读取 markdown
    with open(md_path, encoding="utf-8") as f:
        md_content = f.read()

    # 解析
    elements = parse_markdown(md_content)
    print(f"📊 解析: {len(elements)} 个元素")

    # 创建带样式的 docx
    create_docx_with_styles(styles_xml_path, output_path)

    # 打开并写入内容
    doc = Document(output_path)

    # 确定正文样式名
    body_style = "ZDWP正文"
    if config and config.get("body_style_name"):
        body_style = config["body_style_name"]

    # 表名/表格内容/图名:优先用模板提取到的真实样式名(模板里可能带空格,如「ZDWP 表名」),
    # 硬编码名只作 fallback —— 否则名字对不上就整体退化成正文样式。
    styles_map = (config or {}).get("styles", {})
    table_cell_style = styles_map.get("ZDWP3", "ZDWP表格内容")
    table_title_style = styles_map.get("ZDWP1", "ZDWP表名")
    figure_title_style = styles_map.get("ZDWP4", "ZDWP图名")

    # 检查样式是否存在
    available_styles = {s.name for s in doc.styles}
    if body_style not in available_styles:
        body_style = "Normal"
        print("⚠️  使用 Normal 作为正文样式")

    if table_cell_style not in available_styles:
        table_cell_style = body_style
        print(f"⚠️  使用 {body_style} 作为表格内容样式")

    if table_title_style not in available_styles:
        table_title_style = body_style
        print(f"⚠️  使用 {body_style} 作为表名样式")

    if figure_title_style not in available_styles:
        figure_title_style = body_style
        print(f"⚠️  使用 {body_style} 作为图名样式")

    # 样式映射
    heading_styles = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
    }

    def _styled_paragraph(text, style_name, inline=True):
        """加一段:样式尽力套,行内 markdown 一律消化(inline=真 run 加粗;否则整段剥标记)。"""
        try:
            para = doc.add_paragraph(style=style_name)
        except KeyError:
            para = doc.add_paragraph()
        if inline:
            add_md_runs(para, text)
        else:
            para.add_run(strip_md_markers(text))
        return para

    # 写入内容
    for elem in elements:
        if elem["type"] == "heading":
            level = elem["level"]
            style_name = heading_styles.get(level, "Heading 4")
            # 标题样式自带加粗,行内标记整段剥掉
            text = strip_md_markers(elem["text"])
            try:
                doc.add_paragraph(text, style=style_name)
            except KeyError:
                doc.add_heading(text, level=level)

        elif elem["type"] == "paragraph":
            _styled_paragraph(elem["text"], body_style)

        elif elem["type"] == "table_title":
            _styled_paragraph(elem["text"], table_title_style, inline=False)

        elif elem["type"] == "figure_title":
            _styled_paragraph(elem["text"], figure_title_style, inline=False)

        elif elem["type"] == "table":
            headers = elem["headers"]
            rows = elem["rows"]

            # 计算行列数
            num_cols = len(headers)
            num_rows = 1 + len(rows)  # 表头 + 数据行

            # 创建表格
            table = doc.add_table(rows=num_rows, cols=num_cols)

            # 设置表格边框
            set_table_border(table)

            # 填充表头
            header_row = table.rows[0]
            for j, header in enumerate(headers):
                cell = header_row.cells[j]
                # 垂直居中
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # 清空默认段落，设置样式
                cell.text = ""
                para = cell.paragraphs[0]
                add_md_runs(para, header)
                # 水平居中
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                with contextlib.suppress(KeyError):
                    para.style = table_cell_style

            # 填充数据行
            for i, row_data in enumerate(rows):
                row = table.rows[i + 1]
                for j, cell_text in enumerate(row_data):
                    if j < num_cols:
                        cell = row.cells[j]
                        # 垂直居中
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.text = ""
                        para = cell.paragraphs[0]
                        add_md_runs(para, cell_text)
                        # 水平居中
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        with contextlib.suppress(KeyError):
                            para.style = table_cell_style

    # 保存
    doc.save(output_path)
    clear_quarantine(output_path)
    print(f"✅ 输出: {output_path}")

    return output_path


# md2docx 子命令 parser（build_parser 里赋值；fall-through 时打印它的 help）
_MD2DOCX_PARSER = None


def cmd_md2docx(args):
    """md2docx 子命令入口（原 md_docx_template.py main() 三分支 + Finder fallback）"""
    # 无参数时从 Finder 获取选中的 .md 文件
    if not args.input:
        finder_file = get_finder_selection()
        if not finder_file:
            print("❌ Finder 中未选中任何文件")
            print("   💡 在 Finder 选中一个 .md 文件后重新运行")
            sys.exit(1)
        if not finder_file.endswith(".md"):
            ext = os.path.splitext(finder_file)[1] or "(无扩展名)"
            print(f"❌ 需要 .md 文件，但选中的是: {ext}")
            print(f"   选中: {os.path.basename(finder_file)}")
            sys.exit(1)
        args.input = finder_file
        print(f"📄 从 Finder 获取: {os.path.basename(finder_file)}")

    # 命令: extract
    if args.input == "extract":
        if not args.template_or_md:
            print("❌ 用法: python md_docx_heading_template.py extract 模板.docx")
            sys.exit(1)
        extract_styles_xml(args.template_or_md)
        return

    # 命令: convert
    if args.input == "convert":
        if not args.template_or_md:
            print("❌ 用法: python md_docx_heading_template.py convert input.md -o output.docx")
            sys.exit(1)

        md_path = args.template_or_md
        styles_path = args.styles or os.path.join(DEFAULT_STYLES_DIR, "heading_styles.xml")
        output_path = args.output or os.path.splitext(md_path)[0] + ".docx"

        if not os.path.exists(styles_path):
            print(f"❌ 样式文件不存在: {styles_path}")
            print("   请先运行: python md_docx_heading_template.py extract 模板.docx")
            sys.exit(1)

        # 读取配置
        config_path = os.path.join(os.path.dirname(styles_path), "styles_config.json")
        config = None
        if os.path.exists(config_path):
            with open(config_path, encoding="utf-8") as f:
                config = json.load(f)

        convert_md_to_docx(md_path, styles_path, output_path, config)
        return

    # 一步完成: input.md -t 模板.docx -o output.docx
    if args.input.endswith(".md"):
        md_path = args.input

        # 使用指定模板或默认模板
        template = args.template or DEFAULT_TEMPLATE

        if os.path.exists(template):
            # 先提取样式
            with tempfile.TemporaryDirectory() as tmpdir:
                print(f"📋 使用模板: {os.path.basename(template)}")
                styles_path, config = extract_styles_xml(template, tmpdir)
                output_path = args.output or os.path.splitext(md_path)[0] + ".docx"
                convert_md_to_docx(md_path, styles_path, output_path, config)
        else:
            # 使用已提取的样式
            styles_path = args.styles or os.path.join(DEFAULT_STYLES_DIR, "heading_styles.xml")
            if not os.path.exists(styles_path):
                print(f"❌ 样式文件不存在: {styles_path}")
                print(f"❌ 默认模板不存在: {DEFAULT_TEMPLATE}")
                print("   请指定模板: -t 模板.docx")
                sys.exit(1)

            config_path = os.path.join(os.path.dirname(styles_path), "styles_config.json")
            config = None
            if os.path.exists(config_path):
                with open(config_path, encoding="utf-8") as f:
                    config = json.load(f)

            output_path = args.output or os.path.splitext(md_path)[0] + ".docx"
            convert_md_to_docx(md_path, styles_path, output_path, config)
        return

    (_MD2DOCX_PARSER or build_parser()).print_help()


# ════════════════════════════════════════════════════════════════════
#  to-html - Markdown 渲染为 HTML 并在浏览器中打开
# ════════════════════════════════════════════════════════════════════

# 内联 CSS，不依赖外部文件
HTML_CSS = """
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Helvetica, Arial, sans-serif;
    line-height: 1.6; color: #1f2937; background: #f9fafb;
    max-width: 900px; margin: 0 auto; padding: 2rem;
  }
  h1 { font-size: 1.8rem; border-bottom: 2px solid #e5e7eb; padding-bottom: 0.5rem; margin: 1.5rem 0 1rem; color: #111827; }
  h2 { font-size: 1.4rem; border-bottom: 1px solid #e5e7eb; padding-bottom: 0.3rem; margin: 1.3rem 0 0.8rem; color: #1f2937; }
  h3 { font-size: 1.15rem; margin: 1rem 0 0.5rem; color: #374151; }
  p { margin: 0.5rem 0; }
  a { color: #2563eb; text-decoration: none; }
  a:hover { text-decoration: underline; }
  code { background: #f3f4f6; padding: 0.15rem 0.4rem; border-radius: 4px; font-size: 0.9em; }
  pre { background: #1f2937; color: #f9fafb; padding: 1rem; border-radius: 8px; overflow-x: auto; margin: 0.8rem 0; }
  pre code { background: none; padding: 0; color: inherit; }
  blockquote { border-left: 4px solid #3b82f6; padding: 0.5rem 1rem; margin: 0.8rem 0; background: #eff6ff; color: #1e40af; }
  table { border-collapse: collapse; width: 100%; margin: 0.8rem 0; }
  th, td { border: 1px solid #d1d5db; padding: 0.5rem 0.75rem; text-align: left; }
  th { background: #f3f4f6; font-weight: 600; }
  tr:nth-child(even) { background: #f9fafb; }
  ul, ol { padding-left: 1.5rem; margin: 0.5rem 0; }
  li { margin: 0.2rem 0; }
  hr { border: none; border-top: 1px solid #e5e7eb; margin: 1.5rem 0; }
  .nav { background: #1f2937; color: white; padding: 0.8rem 1.5rem; margin: -2rem -2rem 2rem; }
  .nav a { color: #93c5fd; margin-right: 1rem; }
  .file-list { list-style: none; padding: 0; }
  .file-list li { padding: 0.5rem 0; border-bottom: 1px solid #e5e7eb; }
  .file-list .date { color: #6b7280; font-size: 0.85em; margin-left: 0.5rem; }
  .emoji { font-size: 1.1em; }
  input[type="checkbox"] { margin-right: 0.3rem; }
</style>
"""


def html_md_to_html_simple(md_text):
    """简易 MD -> HTML 转换，不依赖第三方库"""
    html = md_text

    # 先处理代码块
    code_blocks = []

    def save_code(m):
        code_blocks.append(m.group(1))
        return f"__CODE_BLOCK_{len(code_blocks) - 1}__"

    html = re.sub(r"```[\w]*\n(.*?)```", save_code, html, flags=re.DOTALL)

    inline_codes = []

    def save_inline(m):
        inline_codes.append(m.group(1))
        return f"__INLINE_CODE_{len(inline_codes) - 1}__"

    html = re.sub(r"`([^`]+)`", save_inline, html)

    # 表格
    def convert_table(m):
        lines = m.group(0).strip().split("\n")
        rows = []
        for i, line in enumerate(lines):
            line = line.strip()
            if not line.startswith("|"):
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if i == 1 and all(re.match(r"^[-:]+$", c) for c in cells):
                continue  # separator row
            tag = "th" if i == 0 else "td"
            row = "".join(f"<{tag}>{c}</{tag}>" for c in cells)
            rows.append(f"<tr>{row}</tr>")
        return f"<table>{''.join(rows)}</table>"

    html = re.sub(r"(\|.+\|[\n\r]+)+", convert_table, html)

    # Headers
    html = re.sub(r"^### (.+)$", r"<h3>\1</h3>", html, flags=re.MULTILINE)
    html = re.sub(r"^## (.+)$", r"<h2>\1</h2>", html, flags=re.MULTILINE)
    html = re.sub(r"^# (.+)$", r"<h1>\1</h1>", html, flags=re.MULTILINE)

    # Bold, italic
    html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html)
    html = re.sub(r"\*(.+?)\*", r"<em>\1</em>", html)

    # Checkbox
    html = re.sub(r"- \[x\]", r'<li><input type="checkbox" checked disabled>', html)
    html = re.sub(r"- \[ \]", r'<li><input type="checkbox" disabled>', html)

    # Lists
    html = re.sub(r"^- (.+)$", r"<li>\1</li>", html, flags=re.MULTILINE)

    # Blockquote
    html = re.sub(r"^> (.+)$", r"<blockquote>\1</blockquote>", html, flags=re.MULTILINE)

    # HR
    html = re.sub(r"^---+$", r"<hr>", html, flags=re.MULTILINE)

    # Links
    html = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', html)

    # Paragraphs (lines that aren't already HTML)
    lines = html.split("\n")
    result = []
    for line in lines:
        stripped = line.strip()
        if (
            stripped
            and not stripped.startswith("<")
            and not stripped.startswith("__CODE")
            and not stripped.startswith("__INLINE")
        ):
            result.append(f"<p>{line}</p>")
        else:
            result.append(line)
    html = "\n".join(result)

    # Restore code blocks
    for i, block in enumerate(code_blocks):
        html = html.replace(f"__CODE_BLOCK_{i}__", f"<pre><code>{block}</code></pre>")
    for i, code in enumerate(inline_codes):
        html = html.replace(f"__INLINE_CODE_{i}__", f"<code>{code}</code>")

    return html


def html_render_file(md_path, output_dir=None):
    """渲染单个 MD 文件为 HTML"""
    md_path = Path(md_path)
    md_text = md_path.read_text(encoding="utf-8")

    title = md_path.stem
    title_match = re.search(r"^# (.+)$", md_text, re.MULTILINE)
    if title_match:
        title = title_match.group(1)

    body = html_md_to_html_simple(md_text)

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    {HTML_CSS}
</head>
<body>
    <div class="nav">
        <strong>{md_path.name}</strong>
        <span style="float:right;color:#9ca3af;font-size:0.85em">{md_path.parent}</span>
    </div>
    {body}
</body>
</html>"""

    if output_dir:
        out = Path(output_dir) / f"{md_path.stem}.html"
    else:
        out = Path(tempfile.mkdtemp()) / f"{md_path.stem}.html"

    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(html, encoding="utf-8")
    return out


def html_render_directory(dir_path, output_dir=None):
    """渲染整个目录为带索引的 HTML"""
    dir_path = Path(dir_path)
    md_files = sorted(dir_path.glob("**/*.md"), key=lambda p: p.stat().st_mtime, reverse=True)

    if not md_files:
        show_error(f"目录 {dir_path} 中没有找到 .md 文件")
        sys.exit(1)

    if output_dir is None:
        output_dir = Path(tempfile.mkdtemp())
    else:
        output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    file_links = []
    for md in md_files:
        out = html_render_file(md, output_dir)
        rel = md.relative_to(dir_path)
        mtime = os.path.getmtime(md)
        from datetime import datetime

        date_str = datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M")
        file_links.append((rel, out.name, date_str))

    items = "\n".join(
        f'<li><a href="{name}">{rel}</a><span class="date">{date}</span></li>'
        for rel, name, date in file_links
    )

    index_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{dir_path.name} — 文件浏览</title>
    {HTML_CSS}
</head>
<body>
    <div class="nav"><strong>{dir_path}</strong> <span style="color:#9ca3af">({len(md_files)} 个文件)</span></div>
    <h1>{dir_path.name}</h1>
    <ul class="file-list">{items}</ul>
</body>
</html>"""

    index_path = output_dir / "index.html"
    index_path.write_text(index_html, encoding="utf-8")
    return index_path


def html_render_directory_from_files(file_list, output_dir):
    """从文件列表生成带索引的 HTML 目录"""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    file_links = []
    for f in file_list:
        md_path = Path(f)
        out = html_render_file(md_path, output_dir)
        mtime = os.path.getmtime(md_path)
        from datetime import datetime

        date_str = datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M")
        file_links.append((md_path.name, out.name, date_str))

    items = "\n".join(
        f'<li><a href="{name}">{rel}</a><span class="date">{date}</span></li>'
        for rel, name, date in file_links
    )

    index_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>文件浏览</title>
    {HTML_CSS}
</head>
<body>
    <div class="nav"><strong>多文件预览</strong> <span style="color:#9ca3af">({len(file_list)} 个文件)</span></div>
    <h1>文件列表</h1>
    <ul class="file-list">{items}</ul>
</body>
</html>"""

    index_path = output_dir / "index.html"
    index_path.write_text(index_html, encoding="utf-8")
    return index_path


def cmd_to_html(args):
    """to-html 子命令入口"""
    targets = args.targets

    if not targets:
        show_error("用法: md_tools.py to-html <file.md|directory/> [file2.md ...]")
        sys.exit(1)

    target = targets[0]

    if os.path.isdir(target):
        out = html_render_directory(target, args.output_dir)
        show_success(f"索引页: {out}")
        if not args.no_open:
            webbrowser.open(f"file://{out}")
    elif len(targets) > 1:
        tmpdir = Path(tempfile.mkdtemp())
        for f in targets:
            html_render_file(f, tmpdir)
        out = html_render_directory_from_files(targets, tmpdir)
        if not args.no_open:
            webbrowser.open(f"file://{out}")
    else:
        out = html_render_file(target, args.output_dir)
        show_success(f"{out}")
        if not args.no_open:
            webbrowser.open(f"file://{out}")


# ════════════════════════════════════════════════════════════════════
#  frontmatter - 批量生成 MD frontmatter（LLM）
# ════════════════════════════════════════════════════════════════════

FRONTMATTER_IGNORE_DIRS = {".git", "_site", "__pycache__", "node_modules", ".DS_Store"}

FRONTMATTER_SYSTEM_PROMPT = (
    "你是文档分析助手。根据文档内容生成 YAML frontmatter。"
    "只返回 frontmatter 块（含 --- 分隔符），不要其他文字。"
    "description 用中文，≤50字。tags 3-5 个，用中英文均可。"
    '格式：\n---\ndescription: "描述"\ntags: [tag1, tag2, tag3]\n---'
)


def frontmatter_has_frontmatter(text: str) -> bool:
    """检查文本是否已有 YAML frontmatter（以 --- 开头且有第二个 ---）"""
    if not text.startswith("---\n") and not text.startswith("---\r\n"):
        return False
    second = text.find("---", 4)
    return second != -1


def frontmatter_call_llm(content: str) -> str:
    """调用 LLM 生成 frontmatter。"""
    from llm_client import chat

    return chat(
        system=FRONTMATTER_SYSTEM_PROMPT,
        message=f"为以下文档生成 frontmatter：\n\n{content[:2000]}",
    )


def frontmatter_parse_response(text: str) -> str | None:
    """从 API 响应中提取 --- 之间的 frontmatter 块。

    容错处理模型返回多余文字的情况。
    返回完整的 frontmatter 块（含 --- 分隔符和尾部换行），失败返回 None。
    """
    match = re.search(r"---\s*\n(.*?)\n---", text, re.DOTALL)
    if not match:
        return None
    inner = match.group(1).strip()
    if not inner:
        return None
    return f"---\n{inner}\n---\n"


def frontmatter_scan_files(src_dir: str) -> list[tuple[str, str]]:
    """递归扫描目录下的 .md 文件，跳过 FRONTMATTER_IGNORE_DIRS。

    返回 [(relative_path, absolute_path), ...] 按相对路径排序。
    """
    src = Path(src_dir).resolve()
    results = []
    for root, dirs, files in os.walk(src):
        dirs[:] = [d for d in dirs if d not in FRONTMATTER_IGNORE_DIRS]
        for f in files:
            if f.endswith(".md"):
                abs_path = os.path.join(root, f)
                rel_path = os.path.relpath(abs_path, src)
                results.append((rel_path, abs_path))
    results.sort(key=lambda x: x[0])
    return results


def frontmatter_process_file(filepath: str, content: str) -> str | None:
    """对单个文件调用 API 生成 frontmatter。

    返回 frontmatter 字符串，失败返回 None。
    """
    try:
        raw = frontmatter_call_llm(content)
    except Exception as e:
        print(f"  Error: {e}")
        return None

    fm = frontmatter_parse_response(raw)
    if fm is None:
        print(f"  Failed to parse response: {raw[:100]}")
        return None
    return fm


def cmd_frontmatter(args):
    """frontmatter 子命令入口"""
    src_dir = args.src_dir
    if not os.path.isdir(src_dir):
        show_error(f"'{src_dir}' is not a directory")
        sys.exit(1)

    # 收集文件
    if args.file:
        abs_path = os.path.join(os.path.abspath(src_dir), args.file)
        if not os.path.isfile(abs_path):
            show_error(f"file not found: {abs_path}")
            sys.exit(1)
        files = [(args.file, abs_path)]
    else:
        files = frontmatter_scan_files(src_dir)

    if not files:
        print("No .md files found.")
        return

    # 过滤已有 frontmatter 的文件
    to_process = []
    skipped = 0
    for rel, abs_path in files:
        try:
            content = Path(abs_path).read_text(encoding="utf-8")
        except Exception as e:
            print(f"Warning: cannot read {rel}: {e}")
            skipped += 1
            continue
        if frontmatter_has_frontmatter(content):
            skipped += 1
            continue
        to_process.append((rel, abs_path, content))

    total = len(to_process)

    # dry-run 模式
    if args.dry_run:
        print(f"Files to process ({total}):")
        for rel, _, _ in to_process:
            print(f"  {rel}")
        print(f"\nTotal: {total} files to process, {skipped} skipped (already have frontmatter)")
        return

    if total == 0:
        print(f"All files already have frontmatter. Skipped {skipped} files.")
        return

    print(f"Processing {total} files (skipping {skipped} with existing frontmatter)...\n")

    processed = 0
    failed = 0

    for i, (rel, abs_path, content) in enumerate(to_process, 1):
        print(f"[{i}/{total}] Processing: {rel}")

        fm = frontmatter_process_file(abs_path, content)
        if fm is None:
            failed += 1
            print("  FAILED")
            continue

        # 写入文件头部
        try:
            new_content = fm + "\n" + content
            Path(abs_path).write_text(new_content, encoding="utf-8")
            processed += 1
            desc_match = re.search(r'description:\s*"(.+?)"', fm)
            desc = desc_match.group(1) if desc_match else "(parsed)"
            print(f"  OK: {desc}")
        except Exception as e:
            failed += 1
            print(f"  Write error: {e}")

        # Rate limiting: 0.5s 间隔
        if i < total:
            time.sleep(0.5)

    print(f"\nDone! Processed {processed} files, skipped {skipped}, failed {failed}")


# ════════════════════════════════════════════════════════════════════
#  CLI - argparse 主入口
# ════════════════════════════════════════════════════════════════════

def build_parser():
    """构建 argparse 解析器"""
    parser = argparse.ArgumentParser(
        prog="md_tools.py",
        description="Markdown 工具集 - 格式修复/合并/拆分/去引用/转Docx/转HTML/Frontmatter",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""子命令示例:
  %(prog)s format file.md                        格式修复（引号/标点/单位）
  %(prog)s merge a.md b.md                       合并多个文件
  %(prog)s merge a.md b.md -o combined.md        合并并指定输出
  %(prog)s split input.md                        按一级标题拆分
  %(prog)s strip md_final/                       检查 blockquote（只统计）
  %(prog)s strip md_final/ --fix                 删除 blockquote
  %(prog)s to-html file.md                       渲染为 HTML
  %(prog)s to-html dir/                          目录批量渲染
  %(prog)s frontmatter docs/                     批量生成 frontmatter
  %(prog)s frontmatter docs/ --dry-run           只列出待处理文件
  %(prog)s frontmatter docs/ --file README.md    只处理单个文件
  %(prog)s md2docx input.md -t template.docx     样式复刻转 docx（原 md_docx_template.py）
""",
    )
    parser.add_argument(
        "--version",
        action="version",
        version=f"%(prog)s {SCRIPT_VERSION} (by {SCRIPT_AUTHOR}, {SCRIPT_UPDATED})",
    )

    subparsers = parser.add_subparsers(dest="command", help="可用子命令")

    # ── format ──
    p_format = subparsers.add_parser(
        "format",
        help="文本格式自动修复（引号/标点/单位）",
        description="修复 Markdown 文件中的引号、英文标点和中文单位格式",
    )
    p_format.add_argument("files", nargs="*", help="要处理的 Markdown 文件（支持多个）")
    p_format.add_argument("--quotes", action="store_true", help="只修引号")
    p_format.add_argument("--punct", action="store_true", help="只修英文标点")
    p_format.add_argument("--units", action="store_true", help="只转中文单位")

    # ── merge ──
    p_merge = subparsers.add_parser(
        "merge",
        help="合并多个 Markdown 文件为一个",
        description="将多个 Markdown 文件按文件名排序后合并为一个文件",
    )
    p_merge.add_argument("files", nargs="*", help="要合并的 Markdown 文件")
    p_merge.add_argument("-o", "--output", help="输出文件名（默认 merged.md）")

    # ── split ──
    p_split = subparsers.add_parser(
        "split",
        help="按一级标题拆分 Markdown 文件",
        description="将一个 Markdown 文件按 # 一级标题拆分为多个文件",
    )
    p_split.add_argument("input", nargs="?", help="要拆分的 Markdown 文件")

    # ── strip ──
    p_strip = subparsers.add_parser(
        "strip",
        help="删除 Markdown 文件中所有 blockquote",
        description="删除 Markdown 文件中所有 blockquote（> 开头的行）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""示例:
  md_tools.py strip md_final/                            检查模式（只统计）
  md_tools.py strip md_final/ --fix                      删除（覆盖原文件）
  md_tools.py strip md_final/ --fix --output-dir md_out/ 删除（输出到新目录）
""",
    )
    p_strip.add_argument("input", help="MD 文件或目录路径")
    p_strip.add_argument("--fix", action="store_true", help="执行删除")
    p_strip.add_argument("--output-dir", help="输出目录（默认覆盖原文件）")

    # ── to-html ──
    p_html = subparsers.add_parser(
        "to-html",
        help="Markdown 渲染为 HTML 并在浏览器中打开",
        description="将 Markdown 文件或整个目录渲染为 HTML 页面",
    )
    p_html.add_argument("targets", nargs="*", help="MD 文件或目录（支持多个文件）")
    p_html.add_argument("-o", "--output-dir", help="输出目录（默认用临时目录）")
    p_html.add_argument("--no-open", action="store_true", help="不自动打开浏览器")

    # ── frontmatter ──
    p_fm = subparsers.add_parser(
        "frontmatter",
        help="批量生成 YAML frontmatter（LLM）",
        description="扫描目录下的 Markdown 文件，调用 LLM 生成 description + tags 的 YAML frontmatter",
    )
    p_fm.add_argument("src_dir", help="要扫描的目录路径")
    p_fm.add_argument("--dry-run", action="store_true", help="只列出会处理的文件，不调用 API")
    p_fm.add_argument("--file", dest="file", help="只处理单个文件（相对于 src_dir 的路径）")

    # ── md2docx ──（原 md_docx_template.py，argv 契约逐字保留）
    global _MD2DOCX_PARSER
    p_m2d = subparsers.add_parser(
        "md2docx",
        help="Markdown 转 Docx（样式复刻版，原 md_docx_template.py）",
        description="Markdown 转 Docx（样式复刻版）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 提取样式
  python md_tools.py md2docx extract 模板.docx

  # 转换（Raycast 调用：选中 .md 文件即可）
  python md_tools.py md2docx input.md

  # 指定模板
  python md_tools.py md2docx input.md -t 模板.docx -o output.docx
        """,
    )
    p_m2d.add_argument("input", nargs="?", help="Markdown 文件或命令 (extract/convert)")
    p_m2d.add_argument("template_or_md", nargs="?", help="模板 docx (extract) 或 md 文件 (convert)")
    p_m2d.add_argument("-t", "--template", help="模板 docx 文件")
    p_m2d.add_argument("-s", "--styles", help="样式 XML 文件")
    p_m2d.add_argument("-o", "--output", help="输出文件")
    _MD2DOCX_PARSER = p_m2d

    return parser


def main():
    parser = build_parser()
    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    # 路由到对应子命令
    commands = {
        "format": cmd_format,
        "merge": cmd_merge,
        "split": cmd_split,
        "strip": cmd_strip,
        "to-html": cmd_to_html,
        "frontmatter": cmd_frontmatter,
        "md2docx": cmd_md2docx,
    }

    handler = commands.get(args.command)
    if handler:
        handler(args)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
