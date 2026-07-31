#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""chapter.py — 章操作家族三合一（2026-07-31 家族折叠）

子命令 ↔ 原脚本（函数体逐字搬移；模块级 apply/main 改名 _<sub> 后缀，
.bak-N-日期 备份路径生成收敛到 _cli_common（机制同构，文案不变））：

    convert-arabic   ← convert_chapter_format.py
    delete           ← delete_chapter.py
    delete-empty-h1  ← delete_empty_h1.py

各子命令 CLI 与原独立脚本逐字一致：python3 sub/chapter.py <sub> <docx> …。
退役原件在 ~/.Trash/consolidation-20260731/chapter/（含 MANIFEST.md）。
"""
from __future__ import annotations

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[3] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py

# sub/ 自身进 sys.path —— docx_cli 的 _dispatch 用 spec_from_file_location 加载,
# 不带脚本目录, 裸 import _cli_common 会 ImportError (append 不是 insert(0))
_sys.path.append(str(_Path(__file__).resolve().parent))
import _cli_common as _cc  # noqa: E402  备份路径机制 SSOT

import argparse  # noqa: E402
import json  # noqa: E402
import re  # noqa: E402
import shutil  # noqa: E402
import sys  # noqa: E402
from datetime import date  # noqa: E402,F401
from pathlib import Path  # noqa: E402
from typing import Optional  # noqa: E402

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


# ══════════ convert-arabic ← convert_chapter_format.py ══════════

_CN_DIGIT = {
    "零": 0, "一": 1, "二": 2, "三": 3, "四": 4,
    "五": 5, "六": 6, "七": 7, "八": 8, "九": 9,
}
_CN_UNIT = {"十": 10, "百": 100, "千": 1000}


def chinese_to_arabic(s: str) -> int:
    """中文数字串转 int。支持「一」「十」「十一」「二十」「三十五」「一百零三」等。

    若 s 已是阿拉伯数字串,直接 int(s)。无法解析抛 ValueError。
    """
    s = s.strip()
    if not s:
        raise ValueError("empty numeral")
    if s.isdigit():
        return int(s)

    total = 0
    current = 0  # 当前累积位段
    last_unit = 0

    for ch in s:
        if ch in _CN_DIGIT:
            current = _CN_DIGIT[ch]
        elif ch in _CN_UNIT:
            unit = _CN_UNIT[ch]
            if current == 0:
                # 「十X」开头省略一: 十 = 10
                current = 1
            total += current * unit
            current = 0
            last_unit = unit
        else:
            raise ValueError(f"unrecognized char in numeral: {ch!r}")

    total += current
    return total


# ---------------------------------------------------------------------------
# 编号匹配
# ---------------------------------------------------------------------------
# 模式 A: 第X章 Y   (X 可中文数字或阿拉伯)
_RE_CHAPTER = re.compile(r"^第([一二三四五六七八九十百零\d]+)章\s+(.+)$")
# 模式 B: X、Y      (X 必须中文数字; 阿拉伯+顿号不动)
_RE_DUNHAO = re.compile(r"^([一二三四五六七八九十百零]+)、\s*(.+)$")
# 模式 C: 已是阿拉伯「N Y」 → 跳过
_RE_ARABIC = re.compile(r"^\d+\s+\S")


def parse_h1(text: str) -> tuple[int, str] | None:
    """返回 (num, rest) 或 None(不匹配/已 normalized/无法转换)。"""
    if _RE_ARABIC.match(text):
        return None
    m = _RE_CHAPTER.match(text)
    if m:
        try:
            return chinese_to_arabic(m.group(1)), m.group(2).strip()
        except ValueError:
            return None
    m = _RE_DUNHAO.match(text)
    if m:
        try:
            return chinese_to_arabic(m.group(1)), m.group(2).strip()
        except ValueError:
            return None
    return None


# ---------------------------------------------------------------------------
# Run 级替换 (保留首 run 格式)
# ---------------------------------------------------------------------------
def rewrite_paragraph_runs(paragraph, new_text: str) -> None:
    """把段的可见文字改成 new_text,保留 runs[0] 的格式属性。

    做法:把 new_text 写进 runs[0].text;清空其余 run text(保 run 节点以防
    被段落级 pPr 引用;实测安全)。这样 H1 段编号 + 标题文字全部走 run[0] 的
    rPr(bold/size/宋体),格式 100% 保留。
    """
    runs = paragraph.runs
    if not runs:
        # 罕见: paragraph 无 run 直接放裸文本,fallback 用 text=
        paragraph.text = new_text
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def process(docx_path: Path, dry_run: bool, do_backup: bool) -> dict:
    doc = Document(str(docx_path))
    plan: list[dict] = []

    for idx, p in enumerate(doc.paragraphs):
        if p.style.name != "Heading 1":
            continue
        old = p.text
        parsed = parse_h1(old)
        if parsed is None:
            plan.append({
                "para_idx": idx,
                "action": "skip",
                "before": old,
                "after": old,
            })
            continue
        num, rest = parsed
        new = f"{num} {rest}"
        plan.append({
            "para_idx": idx,
            "action": "rewrite",
            "before": old,
            "after": new,
        })
        if not dry_run:
            rewrite_paragraph_runs(p, new)

    report = {
        "docx": str(docx_path),
        "dry_run": dry_run,
        "total_h1": len(plan),
        "rewrite_count": sum(1 for x in plan if x["action"] == "rewrite"),
        "skip_count": sum(1 for x in plan if x["action"] == "skip"),
        "items": plan,
        "backup": None,
    }

    if not dry_run:
        if do_backup:
            backup = _next_backup_path(docx_path)
            shutil.copy2(docx_path, backup)
            report["backup"] = str(backup)
        doc.save(str(docx_path))

    return report


def _next_backup_path(docx_path: Path) -> Path:
    """生成 `<stem>.bak-N-YYYY-MM-DD.docx`,N 递增直至不冲突。"""
    return _cc.find_next_backup(docx_path)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main_convert_arabic() -> int:
    ap = argparse.ArgumentParser(
        description="H1 段「第X章 Y」/「X、Y」→「N Y」阿拉伯数字格式转换。",
    )
    ap.add_argument("docx", type=Path, help="目标 docx 路径")
    ap.add_argument("--dry-run", action="store_true",
                    help="只打印 plan 不改文件")
    ap.add_argument("--no-backup", action="store_true",
                    help="真改时跳过 .bak 备份")
    ap.add_argument("--report", type=Path,
                    help="把 JSON 报告写到此路径")
    args = ap.parse_args()

    if not args.docx.exists():
        print(f"ERROR: docx not found: {args.docx}", file=sys.stderr)
        return 2

    report = process(
        args.docx,
        dry_run=args.dry_run,
        do_backup=not args.no_backup,
    )

    # 终端摘要
    mode = "DRY-RUN" if args.dry_run else "APPLIED"
    print(f"[{mode}] {args.docx}")
    print(f"  H1 总数: {report['total_h1']}")
    print(f"  改写: {report['rewrite_count']}  跳过: {report['skip_count']}")
    if report.get("backup"):
        print(f"  备份: {report['backup']}")
    for item in report["items"]:
        tag = "REWRITE" if item["action"] == "rewrite" else "skip   "
        print(f"  [{tag}] para#{item['para_idx']}: {item['before']!r}"
              f" -> {item['after']!r}")

    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(
            json.dumps(report, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        print(f"  JSON 报告: {args.report}")

    return 0


# ---------------- pipeline adapter ----------------
def apply_convert_arabic(doc, args=None) -> dict:
    """pipeline: H1 「第X章 Y」/「X、Y」→「N Y」"""
    dry = bool(getattr(args, "dry_run", False)) if args else False
    plan = []
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name != "Heading 1":
            continue
        old = p.text
        parsed = parse_h1(old)
        if parsed is None:
            plan.append({"para_idx": idx, "action": "skip", "before": old, "after": old})
            continue
        num, rest = parsed
        new = f"{num} {rest}"
        plan.append({"para_idx": idx, "action": "rewrite", "before": old, "after": new})
        if not dry:
            rewrite_paragraph_runs(p, new)
    return {
        "changed": sum(1 for x in plan if x["action"] == "rewrite"),
        "total_h1": len(plan),
        "skip_count": sum(1 for x in plan if x["action"] == "skip"),
        "items": plan,
    }


# ══════════ delete ← delete_chapter.py ══════════

H1_STYLE = "Heading 1"


def find_h1_paragraphs(doc):
    """返回 [(idx_in_paragraphs, paragraph)] 仅 Heading 1."""
    return [(i, p) for i, p in enumerate(doc.paragraphs) if p.style and p.style.name == H1_STYLE]


def match_h1_by_number(h1_list, numbers):
    """numbers: list[str] like ['6','7','8']. 匹配 text 开头 '^N\\s+'."""
    matched = []
    for n in numbers:
        pat = re.compile(rf"^{re.escape(n)}\s+")
        hits = [(i, p) for i, p in h1_list if pat.match(p.text)]
        if not hits:
            raise SystemExit(f"[ERR] 未找到 H1 编号 {n!r} 的段")
        if len(hits) > 1:
            raise SystemExit(f"[ERR] H1 编号 {n!r} 匹配到 {len(hits)} 段 (期望 1)")
        matched.append(hits[0])
    return matched


def match_h1_by_text_prefix(h1_list, prefixes):
    matched = []
    for pref in prefixes:
        hits = [(i, p) for i, p in h1_list if p.text.startswith(pref)]
        if not hits:
            raise SystemExit(f"[ERR] 未找到 H1 text 前缀 {pref!r}")
        if len(hits) > 1:
            raise SystemExit(f"[ERR] H1 text 前缀 {pref!r} 匹配 {len(hits)} 段")
        matched.append(hits[0])
    return matched


def collect_block_elements(body):
    """body 的直接子节点中,顺序保留 <w:p> 与 <w:tbl>(<w:sectPr> 排除)."""
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    return [el for el in body.iterchildren() if el.tag in (p_tag, tbl_tag)]


def build_delete_plan(doc, target_h1_elems):
    """对每个 target H1 <w:p>,找其在 body block-list 中位置,延伸到下一 H1 前(或末尾).
    返回 list of dict: {h1_text, start_block_idx, end_block_idx_exclusive, n_paragraphs, n_tables, elements}
    """
    body = doc.element.body
    blocks = collect_block_elements(body)
    block_index = {id(el): i for i, el in enumerate(blocks)}

    # 找出所有 H1 段 <w:p> 在 blocks 中的位置(按 docx style)
    h1_positions = []  # list of (block_idx, element, text)
    for p in doc.paragraphs:
        if p.style and p.style.name == H1_STYLE:
            idx = block_index.get(id(p._element))
            if idx is not None:
                h1_positions.append((idx, p._element, p.text))

    h1_positions.sort(key=lambda x: x[0])
    # 给每个 H1 找 "下一个 H1 起始 idx"
    next_h1_idx = {}
    for i, (bi, _el, _txt) in enumerate(h1_positions):
        next_h1_idx[bi] = h1_positions[i + 1][0] if i + 1 < len(h1_positions) else len(blocks)

    plans = []
    target_ids = {id(el) for el in target_h1_elems}
    for bi, el, txt in h1_positions:
        if id(el) not in target_ids:
            continue
        end = next_h1_idx[bi]
        chunk = blocks[bi:end]
        n_p = sum(1 for e in chunk if e.tag == qn("w:p"))
        n_t = sum(1 for e in chunk if e.tag == qn("w:tbl"))
        plans.append({
            "h1_text": txt,
            "start_block_idx": bi,
            "end_block_idx_exclusive": end,
            "n_blocks": end - bi,
            "n_paragraphs": n_p,
            "n_tables": n_t,
            "_elements": chunk,
        })
    return plans


def apply_plan(plans):
    total_p = 0
    total_t = 0
    for plan in plans:
        for el in plan["_elements"]:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        total_p += plan["n_paragraphs"]
        total_t += plan["n_tables"]
    return total_p, total_t


def next_backup_path(docx_path: Path) -> Path:
    return _cc.find_next_backup(docx_path)


def main_delete_chapter():
    ap = argparse.ArgumentParser(description="按 H1 编号或 text 前缀删整章 (H1 + 其下所有段表 → 下一 H1 前/文末).")
    ap.add_argument("docx", type=Path, help="目标 docx 路径")
    grp = ap.add_mutually_exclusive_group(required=True)
    grp.add_argument("--h1", help="H1 编号列表, 逗号分隔. 例: 6,7,8")
    grp.add_argument("--h1-text", help="H1 text 前缀列表, 逗号分隔. 例: '6 保障措施,7 嘉兴...'")
    ap.add_argument("--dry-run", action="store_true", help="只列 plan, 不写盘")
    ap.add_argument("--no-backup", action="store_true", help="不生成 .bak-N 备份")
    ap.add_argument("--report", type=Path, help="JSON 报告输出路径")
    args = ap.parse_args()

    if not args.docx.exists():
        raise SystemExit(f"[ERR] docx 不存在: {args.docx}")

    doc = Document(str(args.docx))
    h1_list = find_h1_paragraphs(doc)
    if not h1_list:
        raise SystemExit("[ERR] 文档无 Heading 1")

    if args.h1:
        nums = [s.strip() for s in args.h1.split(",") if s.strip()]
        matched = match_h1_by_number(h1_list, nums)
    else:
        prefs = [s.strip() for s in args.h1_text.split(",") if s.strip()]
        matched = match_h1_by_text_prefix(h1_list, prefs)

    target_elems = [p._element for _, p in matched]
    plans = build_delete_plan(doc, target_elems)

    # 打印 plan
    print(f"[plan] docx: {args.docx}")
    print(f"[plan] 全文 H1 段总数: {len(h1_list)}")
    print(f"[plan] 拟删章数: {len(plans)}")
    for pl in plans:
        print(f"  - H1='{pl['h1_text']}' blocks[{pl['start_block_idx']}:{pl['end_block_idx_exclusive']}] "
              f"=> {pl['n_paragraphs']} 段 + {pl['n_tables']} 表 (共 {pl['n_blocks']} 块)")

    report = {
        "docx": str(args.docx),
        "dry_run": args.dry_run,
        "h1_total_before": len(h1_list),
        "plans": [{k: v for k, v in pl.items() if not k.startswith("_")} for pl in plans],
    }

    if args.dry_run:
        print("[dry-run] 未写盘")
    else:
        if not args.no_backup:
            bak = next_backup_path(args.docx)
            shutil.copy2(args.docx, bak)
            print(f"[backup] {bak}")
            report["backup"] = str(bak)
        total_p, total_t = apply_plan(plans)
        doc.save(str(args.docx))
        print(f"[done] 已删 {total_p} 段 + {total_t} 表, 保存 → {args.docx}")
        report["deleted_paragraphs"] = total_p
        report["deleted_tables"] = total_t

        # verify 重读 + H1 列表
        doc2 = Document(str(args.docx))
        h1_after = [p.text for p in doc2.paragraphs if p.style and p.style.name == H1_STYLE]
        print(f"[verify] 删后 H1 段数: {len(h1_after)}")
        for t in h1_after:
            print(f"  · {t}")
        report["h1_total_after"] = len(h1_after)
        report["h1_texts_after"] = h1_after

    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"[report] {args.report}")


# ---------------- pipeline adapter ----------------
def apply_delete_chapter(doc, args=None) -> dict:
    h1 = getattr(args, "delete_h1", None) if args else None
    h1_text = getattr(args, "delete_h1_text", None) if args else None
    dry = bool(getattr(args, "dry_run", False)) if args else False
    if not h1 and not h1_text:
        return {"changed": 0, "skipped": "no delete_h1 / delete_h1_text in args"}
    h1_list = find_h1_paragraphs(doc)
    if h1:
        nums = [s.strip() for s in h1.split(",") if s.strip()]
        matched = match_h1_by_number(h1_list, nums)
    else:
        prefs = [s.strip() for s in h1_text.split(",") if s.strip()]
        matched = match_h1_by_text_prefix(h1_list, prefs)
    target_elems = [p._element for _, p in matched]
    plans = build_delete_plan(doc, target_elems)
    if not dry:
        total_p, total_t = apply_plan(plans)
    else:
        total_p = sum(pl["n_paragraphs"] for pl in plans)
        total_t = sum(pl["n_tables"] for pl in plans)
    return {
        "changed": total_p + total_t,
        "deleted_paragraphs": total_p,
        "deleted_tables": total_t,
        "plans_count": len(plans),
    }


# ══════════ delete-empty-h1 ← delete_empty_h1.py ══════════

H1_STYLES = {
    "Heading 1",
    "标题 1",
    "heading 1",
    "10",
    "1.1.1.1 N级标题",
}


def get_style_name(p) -> str:
    try:
        return (p.style.name or "") if p.style is not None else ""
    except Exception:
        return ""


def is_empty_h1(p) -> bool:
    sname = get_style_name(p)
    if sname not in H1_STYLES:
        return False
    text = "".join(r.text or "" for r in p.runs)
    return text.strip() == ""


def extract_sectPr(p_elem):
    """从段 <w:p> 的 <w:pPr> 里取 <w:sectPr> 子元素 (若有), 返回元素或 None。"""
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return None
    return pPr.find(qn("w:sectPr"))


def ensure_pPr(p_elem):
    """确保段有 <w:pPr>, 没有就建一个插在 <w:p> 头部, 返回 pPr 元素。"""
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        return pPr
    from lxml import etree
    pPr = etree.SubElement(p_elem, qn("w:pPr"))
    # SubElement 默认追加在末尾; 但 <w:pPr> 在 <w:p> 必须排首位 — 重排
    p_elem.remove(pPr)
    p_elem.insert(0, pPr)
    return pPr


def migrate_sectPr(empty_p_elem, prev_p_elem) -> bool:
    """把 empty_p 的 sectPr 迁到 prev_p 的 pPr 末尾。返回是否迁移过。"""
    sectPr = extract_sectPr(empty_p_elem)
    if sectPr is None:
        return False
    # 从 empty 的 pPr 删
    empty_pPr = empty_p_elem.find(qn("w:pPr"))
    if empty_pPr is not None:
        empty_pPr.remove(sectPr)
    # 装到 prev 的 pPr 末尾
    prev_pPr = ensure_pPr(prev_p_elem)
    # 若 prev_pPr 已含 sectPr (罕见 — prev 也是节末), 替掉
    existing = prev_pPr.find(qn("w:sectPr"))
    if existing is not None:
        prev_pPr.remove(existing)
    prev_pPr.append(sectPr)
    return True


def apply_delete_empty_h1(doc, args=None) -> dict:
    """pipeline-compatible adapter.

    扫 doc.paragraphs 找空 H1 段; 倒序删并把 sectPr 迁到前段。
    不读写文件 (driver 已 parse / 将 save)。

    args (Namespace, 可为 None):
        dry_run: bool — True 时只 plan 不改 doc

    返回 dict: {"changed": int, "issues": [...], "deleted": [...],
                "sectPr_migrated": int, "h1_total_before": int, "h1_total_after": int}
    """
    dry_run = bool(getattr(args, "dry_run", False)) if args is not None else False
    issues: list[str] = []

    paragraphs = doc.paragraphs
    targets = [idx for idx, p in enumerate(paragraphs) if is_empty_h1(p)]

    h1_before = sum(1 for p in paragraphs if get_style_name(p) in H1_STYLES)

    report: dict = {
        "changed": 0,
        "issues": issues,
        "deleted": [],
        "sectPr_migrated": 0,
        "h1_total_before": h1_before,
        "empty_h1_idx": targets,
    }

    if not targets:
        report["h1_total_after"] = h1_before
        return report

    if dry_run:
        report["h1_total_after"] = h1_before  # plan only
        return report

    # 倒序删 (避免索引漂移)
    for idx in sorted(targets, reverse=True):
        empty_p = paragraphs[idx]
        empty_elem = empty_p._element
        parent = empty_elem.getparent()
        prev_elem = empty_elem.getprevious()
        # 跳过非 <w:p> 的 sibling (如 <w:tbl>)
        while prev_elem is not None and prev_elem.tag != qn("w:p"):
            prev_elem = prev_elem.getprevious()

        migrated = False
        if prev_elem is not None:
            migrated = migrate_sectPr(empty_elem, prev_elem)
            if migrated:
                report["sectPr_migrated"] += 1
        else:
            issues.append(f"idx={idx} 是首段, sectPr 将随删丢失")

        parent.remove(empty_elem)
        report["deleted"].append({"idx": idx, "sectPr_migrated": migrated})
        report["changed"] += 1

    # 重扫 H1 数 (paragraphs 是 live view, 删后立刻新)
    report["h1_total_after"] = sum(1 for p in doc.paragraphs if get_style_name(p) in H1_STYLES)
    return report


def main_delete_empty_h1(argv=None) -> int:
    ap = argparse.ArgumentParser(description="删空壳 H1 段 + 迁 sectPr 到前段")
    ap.add_argument("docx_path")
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--no-backup", action="store_true")
    ap.add_argument("--report", default=None)
    args = ap.parse_args(argv)

    docx_path = Path(args.docx_path)
    if not docx_path.exists():
        print(f"ERROR: {docx_path} 不存在", file=sys.stderr)
        return 2

    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs

    # 找空 H1 的物理索引 + 前段索引
    targets = []
    for idx, p in enumerate(paragraphs):
        if is_empty_h1(p):
            targets.append(idx)

    # 收集 H1 序列 (before)
    h1_before = []
    for idx, p in enumerate(paragraphs):
        if get_style_name(p) in H1_STYLES:
            text = "".join(r.text or "" for r in p.runs)
            h1_before.append({"idx": idx, "text": text[:60], "empty": text.strip() == ""})

    print(f"[scan] 段总数={len(paragraphs)} H1总数={len(h1_before)} 空H1={len(targets)}")
    for t in targets:
        print(f"  - idx={t} (空 H1, 待删)")

    report = {
        "total_paragraphs_before": len(paragraphs),
        "h1_total_before": len(h1_before),
        "h1_sequence_before": h1_before,
        "empty_h1_idx": targets,
        "deleted": [],
        "sectPr_migrated_to_idx": [],
    }

    if args.dry_run:
        if args.report:
            Path(args.report).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"[report] {args.report}")
        print("[dry-run] 不动 docx")
        return 0

    if not targets:
        print("[noop] 无空 H1, 无事可做")
        if args.report:
            Path(args.report).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return 0

    # backup
    if not args.no_backup:
        bak = _cc.make_backup(docx_path)
        print(f"[backup] {bak.name}")

    # 按倒序删 (避免索引漂移)
    for idx in sorted(targets, reverse=True):
        empty_p = paragraphs[idx]
        empty_elem = empty_p._element
        # 找前段元素 (同父下)
        parent = empty_elem.getparent()
        prev_elem = empty_elem.getprevious()
        # 跳过非 <w:p> 的 sibling (如 <w:tbl>)
        while prev_elem is not None and prev_elem.tag != qn("w:p"):
            prev_elem = prev_elem.getprevious()

        migrated = False
        if prev_elem is not None:
            migrated = migrate_sectPr(empty_elem, prev_elem)
            if migrated:
                # prev_elem 在 doc.paragraphs 里的 idx 不易直接拿,记元素 hash 替代
                report["sectPr_migrated_to_idx"].append({"empty_idx": idx, "prev_elem_tag_ok": True})
        else:
            # 无前段 (文档第一段就是空 H1 — 罕见), sectPr 留 empty 段一起删 → 节信息丢
            print(f"WARN: idx={idx} 是首段, sectPr 将随删丢失", file=sys.stderr)

        # 删该段
        parent.remove(empty_elem)
        report["deleted"].append({"idx": idx, "sectPr_migrated": migrated})
        print(f"[deleted] idx={idx} sectPr_migrated={migrated}")

    doc.save(str(docx_path))
    print(f"[saved] {docx_path}")

    # verify 重读
    try:
        doc2 = Document(str(docx_path))
    except Exception as e:
        print(f"ERROR: 重读失败 (OOXML 可能损坏): {e}", file=sys.stderr)
        return 3

    h1_after = []
    for idx, p in enumerate(doc2.paragraphs):
        if get_style_name(p) in H1_STYLES:
            text = "".join(r.text or "" for r in p.runs)
            h1_after.append({"idx": idx, "text": text[:60], "empty": text.strip() == ""})
    report["total_paragraphs_after"] = len(doc2.paragraphs)
    report["h1_total_after"] = len(h1_after)
    report["h1_sequence_after"] = h1_after

    print(f"[verify] 段总数={len(doc2.paragraphs)} H1总数={len(h1_after)} (空H1={sum(1 for x in h1_after if x['empty'])})")
    for x in h1_after:
        print(f"  - idx={x['idx']} {x['text']!r}")

    if args.report:
        Path(args.report).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"[report] {args.report}")

    return 0


# ──────────────────────────── 家族入口（子命令分发）────────────────────────────

SUBCOMMANDS = {
    "convert-arabic": main_convert_arabic,
    "delete": main_delete_chapter,
    "delete-empty-h1": main_delete_empty_h1,
}


def main(argv: list[str] | None = None) -> int:
    args = list(sys.argv[1:] if argv is None else argv)
    if not args or args[0] in ("-h", "--help"):
        print("usage: chapter.py {" + ",".join(SUBCOMMANDS) + "} <args…>\n"
              "每个子命令的参数与原独立脚本逐字一致：chapter.py <sub> --help 查看。")
        return 0 if args else 2
    sub, rest = args[0], args[1:]
    fn = SUBCOMMANDS.get(sub)
    if fn is None:
        print(f"[chapter] unknown subcommand: {sub!r}; choices={list(SUBCOMMANDS)}",
              file=sys.stderr)
        return 2
    saved = sys.argv[:]
    sys.argv = [sys.argv[0]] + rest
    try:
        rc = fn()
        return int(rc) if isinstance(rc, int) else 0
    finally:
        sys.argv = saved


if __name__ == "__main__":
    sys.exit(main())
