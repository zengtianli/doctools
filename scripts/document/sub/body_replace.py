#!/usr/bin/env python3
# distilled from eco-flow/taizhou-天台 body-replace need (2026-05-26 W1)
r"""body_replace.py — 保留 shell (styles/numbering/sectPr/cover/H1 标题) + body 替换为 content 正文。

场景:
    - shell docx = 健康 styleset 的模板(磐安切片), 含封面/H1 标题/页眉页脚/样式表
    - content docx = 真实内容 docx(天台章节), styleset 可能不一致
    - 目标: shell 的 styles + 第一个 H1 标题段 + content 正文(默认丢 content 首个 H1, 避免重复)

CLI:
    python3 scripts/document/sub/body_replace.py \
        --shell <path> --content <path> --out <path> \
        [--keep-shell-h1 | --no-keep-shell-h1] \
        [--dry-run]

策略:
    1. shutil.copy(shell, out) — 全套 styles/numbering/media/sectPr/cover/headers/footers
    2. open out → body children = list(body):
         - keep-shell-h1 (默认): 找到第一个 H1 段, 保留 [0, first_h1+1), 删它之后的内容
         - no-keep-shell-h1: 删 body 全部内容(保留 sectPr)
    3. open content → body children:
         - keep-shell-h1: 跳过 content 第一个 H1, 从其后开始取
         - no-keep-shell-h1: 取全部
    4. 取出元素 deepcopy 进 out body 末尾(在 sectPr 之前)
    5. style fallback: content 段引用 shell 不存在的 style → 改 Normal + stderr warning
    6. 保存 out

Known limitations (W1 first pass):
    - content 内嵌图片(inline shape) 不复制 image part 到 out 的 media/, rId 会指空。
      时间紧, 先做文字+段落+表格, 图片留 follow-up。stderr 打 warning 列被丢的图。
    - content 跨章节引用(bookmarks/refs) out of scope。
    - footnote/comment part 不复制(罕见)。
"""
from __future__ import annotations

import argparse
import shutil
import sys
from copy import deepcopy
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx 未安装 (pip install python-docx)", file=sys.stderr)
    sys.exit(2)


H1_STYLES = {
    "Heading 1",
    "标题 1",
    "heading 1",
    "1",
    "10",
    "Heading1",
    "1.1.1.1 N级标题",
}


def _is_h1_elem(elem) -> bool:
    """Body-level <w:p> with H1-class style."""
    if elem.tag != qn("w:p"):
        return False
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return False
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return False
    style_val = pStyle.get(qn("w:val")) or ""
    if style_val in H1_STYLES:
        return True
    if style_val.lower() in {s.lower() for s in H1_STYLES}:
        return True
    return False


def _paragraph_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def _find_first_h1_index(children: list, sect_idx: int) -> int:
    """Return index of first H1 element in children[0:sect_idx), or -1 if none."""
    for i in range(sect_idx):
        if _is_h1_elem(children[i]):
            return i
    return -1


def _find_sect_idx(body) -> int:
    """Find index of trailing <w:sectPr> in body. Returns len(children) if none."""
    children = list(body)
    for i in range(len(children) - 1, -1, -1):
        if children[i].tag == qn("w:sectPr"):
            return i
    return len(children)


def _collect_used_styles(doc) -> set:
    """Collect all pStyle val referenced by paragraphs in body. Used for fallback check."""
    styles = set()
    for p in doc.element.body.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        v = pStyle.get(qn("w:val"))
        if v:
            styles.add(v)
    return styles


def _available_style_ids(doc) -> set:
    """Return set of styleId defined in doc's styles.xml (canonical resolvable keys).

    Note: w:val on w:pStyle MUST be a styleId, not a name. Names are only useful
    as a lookup target — we surface the name→id map separately via
    _available_name_to_id.
    """
    ids = set()
    try:
        for s in doc.styles:
            try:
                if s.style_id:
                    ids.add(s.style_id)
            except Exception:
                pass
    except Exception:
        pass
    return ids


def _available_name_to_id(doc) -> dict:
    """Map style.name -> style.style_id (for content val that uses names not ids)."""
    name_to_id: dict = {}
    try:
        for s in doc.styles:
            try:
                if s.name and s.style_id:
                    name_to_id[s.name] = s.style_id
            except Exception:
                pass
    except Exception:
        pass
    return name_to_id


def _has_image_refs(elem) -> int:
    """Count <a:blip> or <w:drawing> descendants — proxy for inline images."""
    n = 0
    for _ in elem.iter(qn("w:drawing")):
        n += 1
    return n


# Common heading style aliases (content styleId → shell preferred styleId/name).
# Tried in order; first one present in `available` wins.
_HEADING_ALIASES = {
    "Heading1":  ["Heading 1", "1", "标题 1"],
    "Heading2":  ["Heading 2", "2", "标题 2"],
    "Heading3":  ["Heading 3", "3", "标题 3"],
    "Heading4":  ["Heading 4", "4", "标题 4"],
    "Heading5":  ["Heading 5", "5", "标题 5"],
    "Heading6":  ["Heading 6", "6", "标题 6"],
    "Heading 1": ["Heading 1", "1", "标题 1"],
    "Heading 2": ["Heading 2", "2", "标题 2"],
    "Heading 3": ["Heading 3", "3", "标题 3"],
    "Heading 4": ["Heading 4", "4", "标题 4"],
    "Heading 5": ["Heading 5", "5", "标题 5"],
    "Heading 6": ["Heading 6", "6", "标题 6"],
    "标题 1":    ["Heading 1", "1", "标题 1"],
    "标题 2":    ["Heading 2", "2", "标题 2"],
    "标题 3":    ["Heading 3", "3", "标题 3"],
    "标题 4":    ["Heading 4", "4", "标题 4"],
    # Common pandoc-produced body styles → shell's body
    "FirstParagraph": ["Normal", "正文"],
    "BodyText":       ["Normal", "正文"],
    "Compact":        ["Normal", "正文"],
}


def _resolve_alias(v: str, available_ids: set, name_to_id: dict):
    """Return target styleId via alias map; tries direct id, then via name_to_id.

    Returns None if no alias hit resolvable styleId.
    """
    aliases = _HEADING_ALIASES.get(v, [])
    for a in aliases:
        # direct styleId hit
        if a in available_ids:
            return a
        # alias is a name → resolve to id
        if a in name_to_id:
            return name_to_id[a]
    return None


def _remap_styles_in_elem(elem, available_ids: set, name_to_id: dict,
                           fallback: str = "Normal") -> list:
    """Walk <w:p> descendants, rewrite pStyle val to a valid styleId.

    Order: (1) val is already a valid styleId → keep
           (2) val is a known name in shell → swap to its styleId
           (3) val matches _HEADING_ALIASES → use alias's resolved styleId
           (4) fall back to `fallback` styleId + record warning
    """
    warnings: list = []
    for p in elem.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        v = pStyle.get(qn("w:val"))
        if not v:
            continue
        # (1) already a valid styleId in shell
        if v in available_ids:
            continue
        # (2) val is a style.name in shell → swap to id
        if v in name_to_id:
            pStyle.set(qn("w:val"), name_to_id[v])
            continue
        # (3) alias map
        alias_target = _resolve_alias(v, available_ids, name_to_id)
        if alias_target is not None:
            pStyle.set(qn("w:val"), alias_target)
            continue
        # (4) fallback
        txt = _paragraph_text(p)[:30]
        warnings.append((v, txt))
        pStyle.set(qn("w:val"), fallback)
    return warnings


def body_replace(
    shell_path: Path,
    content_path: Path,
    out_path: Path,
    keep_shell_h1: bool = True,
) -> dict:
    """Core operation. Returns dict of stats for reporting."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(str(shell_path), str(out_path))

    out_doc = Document(str(out_path))
    out_body = out_doc.element.body
    out_children = list(out_body)
    out_sect_idx = _find_sect_idx(out_body)

    # Step 1: prune out body
    if keep_shell_h1:
        first_h1 = _find_first_h1_index(out_children, out_sect_idx)
        if first_h1 < 0:
            print(
                "WARN: --keep-shell-h1 set but shell has 0 H1 — will keep nothing from shell body",
                file=sys.stderr,
            )
            keep_until = 0
        else:
            keep_until = first_h1 + 1  # keep [0, first_h1] inclusive
    else:
        keep_until = 0

    sect_elem = None
    if out_sect_idx < len(out_children):
        sect_elem = out_children[out_sect_idx]

    # Remove everything from index keep_until up to (but not including) sectPr
    for i, elem in enumerate(out_children):
        if elem is sect_elem:
            continue
        if i >= keep_until and i < out_sect_idx:
            out_body.remove(elem)

    # Step 2: collect from content
    content_doc = Document(str(content_path))
    content_body = content_doc.element.body
    content_children = list(content_body)
    content_sect_idx = _find_sect_idx(content_body)

    if keep_shell_h1:
        c_first_h1 = _find_first_h1_index(content_children, content_sect_idx)
        if c_first_h1 < 0:
            content_start = 0  # no H1 in content → take everything
        else:
            content_start = c_first_h1 + 1  # skip content's first H1
    else:
        content_start = 0

    # Step 3: available styles in out (= shell's styles)
    available_ids = _available_style_ids(out_doc)
    name_to_id = _available_name_to_id(out_doc)

    # Step 4: deepcopy content children into out body, before sectPr
    appended = 0
    image_drops = 0
    style_warnings: list = []

    for i in range(content_start, content_sect_idx):
        src_elem = content_children[i]
        # Skip sectPr if it appears mid-body (rare; nested section breaks)
        if src_elem.tag == qn("w:sectPr"):
            continue
        new_elem = deepcopy(src_elem)
        # style fallback
        warns = _remap_styles_in_elem(new_elem, available_ids, name_to_id, fallback="Normal")
        style_warnings.extend(warns)
        # count images (will be broken — rId points to content's media)
        n_img = _has_image_refs(new_elem)
        image_drops += n_img
        # insert before sectPr
        if sect_elem is not None:
            sect_elem.addprevious(new_elem)
        else:
            out_body.append(new_elem)
        appended += 1

    out_doc.save(str(out_path))

    # Build stats
    out_doc2 = Document(str(out_path))
    final_paras = len(out_doc2.paragraphs)
    final_h1 = [p for p in out_doc2.paragraphs if p.style.name in H1_STYLES]
    first_h1_text = final_h1[0].text if final_h1 else ""
    used_styles = sorted({p.style.name for p in out_doc2.paragraphs if p.style is not None})

    # Print warnings to stderr
    if style_warnings:
        unique_styles = {}
        for v, txt in style_warnings:
            unique_styles.setdefault(v, txt)
        print(
            f"WARN: {len(style_warnings)} paragraphs referenced styles missing in shell — fell back to Normal",
            file=sys.stderr,
        )
        for v, txt in unique_styles.items():
            print(f"  · style {v!r} not in shell (sample: {txt!r})", file=sys.stderr)

    if image_drops:
        print(
            f"WARN: {image_drops} <w:drawing> elements copied without remapping image parts — "
            f"images will likely render broken. (known limitation, follow-up)",
            file=sys.stderr,
        )

    return {
        "shell": str(shell_path),
        "content": str(content_path),
        "out": str(out_path),
        "keep_shell_h1": keep_shell_h1,
        "appended_elements": appended,
        "image_drops": image_drops,
        "style_warnings": len(style_warnings),
        "final_paragraphs": final_paras,
        "first_h1_text": first_h1_text,
        "used_styles": used_styles,
    }


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Replace docx body keeping shell's styles + (optional) first H1.",
    )
    ap.add_argument("--shell", required=True, help="shell docx (styles/cover/H1 source)")
    ap.add_argument("--content", required=True, help="content docx (body source)")
    ap.add_argument("--out", required=True, help="output docx path")

    # Default: keep-shell-h1 ON (mutually exclusive)
    grp = ap.add_mutually_exclusive_group()
    grp.add_argument(
        "--keep-shell-h1", dest="keep_shell_h1", action="store_true",
        default=True,
        help="(default) preserve shell's first H1 paragraph; drop content's first H1",
    )
    grp.add_argument(
        "--no-keep-shell-h1", dest="keep_shell_h1", action="store_false",
        help="drop all shell body; take content from its first element",
    )
    ap.add_argument(
        "--dry-run", action="store_true",
        help="print plan only, don't write output",
    )
    args = ap.parse_args()

    shell = Path(args.shell).expanduser().resolve()
    content = Path(args.content).expanduser().resolve()
    out = Path(args.out).expanduser().resolve()

    if not shell.exists():
        print(f"ERROR: shell docx not found: {shell}", file=sys.stderr)
        return 2
    if not content.exists():
        print(f"ERROR: content docx not found: {content}", file=sys.stderr)
        return 2

    if args.dry_run:
        # Inspect both, print plan
        s_doc = Document(str(shell))
        s_children = list(s_doc.element.body)
        s_sect = _find_sect_idx(s_doc.element.body)
        s_h1 = _find_first_h1_index(s_children, s_sect)
        c_doc = Document(str(content))
        c_children = list(c_doc.element.body)
        c_sect = _find_sect_idx(c_doc.element.body)
        c_h1 = _find_first_h1_index(c_children, c_sect)
        print(f"# DRY RUN — body-replace plan")
        print(f"  shell:   {shell}")
        print(f"    body children: {s_sect}, first H1 idx: {s_h1}")
        if s_h1 >= 0:
            print(f"    shell H1 text: {_paragraph_text(s_children[s_h1])!r}")
        print(f"  content: {content}")
        print(f"    body children: {c_sect}, first H1 idx: {c_h1}")
        if c_h1 >= 0:
            print(f"    content H1 text: {_paragraph_text(c_children[c_h1])!r}")
        print(f"  out:     {out}")
        print(f"  keep_shell_h1: {args.keep_shell_h1}")
        if args.keep_shell_h1:
            keep_until = (s_h1 + 1) if s_h1 >= 0 else 0
            c_start = (c_h1 + 1) if c_h1 >= 0 else 0
        else:
            keep_until = 0
            c_start = 0
        print(f"  → keep shell children [0:{keep_until}) + content children [{c_start}:{c_sect})")
        return 0

    stats = body_replace(shell, content, out, keep_shell_h1=args.keep_shell_h1)
    print(f"OK: wrote {out}")
    print(f"  appended {stats['appended_elements']} elements from content")
    print(f"  final paragraphs: {stats['final_paragraphs']}")
    print(f"  first H1: {stats['first_h1_text']!r}")
    print(f"  styles used: {stats['used_styles']}")
    if stats["image_drops"]:
        print(f"  WARN: {stats['image_drops']} images may render broken (see stderr)")
    if stats["style_warnings"]:
        print(f"  WARN: {stats['style_warnings']} style fallbacks (see stderr)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
