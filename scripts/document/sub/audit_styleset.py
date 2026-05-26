"""audit_styleset.py — group module: style-set health audits (5 subcommands)

  audit-styleset style-coherence            ← 段实际套样式集 ⊆ profile.roles ∪ tolerated
  audit-styleset role-coverage              ← 7 角色全覆盖且无空角色
  audit-styleset body-style-concentration   ← 主 body 样式占比 ≥ profile.concentration_threshold
  audit-styleset style-pool-cleanliness     ← 定义但未用样式 / 总定义 ≤ pool_unused_ratio_max (warn)
  audit-styleset style-pane-filter          ← stylePaneFormatFilter 已设 (warn)

距 distill 自 ad-hoc docx-health 实验 (2026-05-26 W-audit-styleset).
audit-only (零修改) · 各 subcommand 返回 JSON 严重度: fail | warn | pass.

profile yaml: profiles/eco_flow_health.yaml (default), 可 --profile 覆盖.

严重度策略 (按 GOAL):
- fail: style-coherence / role-coverage / body-style-concentration
- warn: style-pool-cleanliness / style-pane-filter
"""
from __future__ import annotations

import argparse
import json
import os
import sys
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

import yaml
from docx import Document

_HERE = Path(__file__).resolve().parent
DEFAULT_PROFILE_PATH = _HERE.parent / "profiles" / "eco_flow_health.yaml"


# ---------- profile loader ----------

def load_profile(profile_path: Path) -> dict:
    """Load YAML profile. Expand ~ in anchor_docx."""
    p = Path(os.path.expanduser(str(profile_path))).resolve()
    if not p.exists():
        raise FileNotFoundError(f"profile not found: {p}")
    with open(p, "r", encoding="utf-8") as f:
        prof = yaml.safe_load(f) or {}
    prof["_profile_path"] = str(p)
    return prof


def _all_role_styles(prof: dict) -> set[str]:
    """Flatten roles dict → set of style names."""
    out: set[str] = set()
    for role, styles in (prof.get("roles") or {}).items():
        for s in styles or []:
            out.add(s)
    return out


def _tolerated(prof: dict) -> set[str]:
    return set(prof.get("tolerated_styles") or [])


# ---------- docx introspection helpers ----------

def _paragraphs_style_count(doc) -> tuple[Counter, Counter, int, int]:
    """Return (style_count_all, style_count_nonempty, total, nonempty_total)."""
    c_all: Counter = Counter()
    c_ne: Counter = Counter()
    nonempty = 0
    for para in doc.paragraphs:
        try:
            sn = para.style.name if para.style else ""
        except Exception:
            sn = ""
        c_all[sn] += 1
        if (para.text or "").strip():
            c_ne[sn] += 1
            nonempty += 1
    return c_all, c_ne, len(doc.paragraphs), nonempty


def _defined_styles(doc) -> set[str]:
    out: set[str] = set()
    for s in doc.styles:
        try:
            out.add(s.name)
        except Exception:
            continue
    return out


def _settings_has_stylepanefilter(docx_path: Path) -> bool:
    try:
        with zipfile.ZipFile(str(docx_path)) as z:
            try:
                xml = z.read("word/settings.xml").decode("utf-8", errors="ignore")
            except KeyError:
                return False
        return "stylePaneFormatFilter" in xml
    except Exception:
        return False


# ---------- 5 audit functions ----------

def audit_style_coherence(docx_path: Path, prof: dict) -> dict:
    """段实际套样式集 ⊆ profile.roles ∪ tolerated。
    fail: 非空段使用 profile-外 + 非 tolerated 样式比例 > coherence_unknown_ratio_max
       OR 任一 tolerated 样式 dominance > tolerated_dominance_max
    """
    doc = Document(str(docx_path))
    _, c_ne, _, ne_total = _paragraphs_style_count(doc)
    role_set = _all_role_styles(prof)
    tol = _tolerated(prof)

    unknown_ne = 0
    unknown_styles: dict[str, int] = {}
    tolerated_overuse: dict[str, float] = {}
    role_seg = 0
    for sn, cnt in c_ne.items():
        if sn in role_set:
            role_seg += cnt
        elif sn in tol:
            # tolerated 但不算 role,检查是否 dominance 过高
            ratio = cnt / ne_total if ne_total else 0.0
            tdmax = float(prof.get("tolerated_dominance_max", 0.30))
            if ratio > tdmax:
                tolerated_overuse[sn] = round(ratio, 4)
        else:
            unknown_ne += cnt
            unknown_styles[sn] = cnt

    unknown_ratio = unknown_ne / ne_total if ne_total else 0.0
    umax = float(prof.get("coherence_unknown_ratio_max", 0.10))

    severity = "pass"
    failure_reasons = []
    if unknown_ratio > umax:
        severity = "fail"
        failure_reasons.append(
            f"unknown_style_ratio {unknown_ratio:.3f} > {umax:.3f}")
    if tolerated_overuse:
        severity = "fail"
        for sn, r in tolerated_overuse.items():
            failure_reasons.append(
                f"tolerated style '{sn}' dominance {r:.3f} > "
                f"{prof.get('tolerated_dominance_max', 0.30):.3f}")

    return {
        "check_id": "style-coherence",
        "severity": severity,
        "found": severity != "pass",
        "details": {
            "nonempty_total": ne_total,
            "role_segment_count": role_seg,
            "unknown_segment_count": unknown_ne,
            "unknown_ratio": round(unknown_ratio, 4),
            "unknown_styles_top10": dict(Counter(unknown_styles).most_common(10)),
            "tolerated_overuse": tolerated_overuse,
            "thresholds": {
                "coherence_unknown_ratio_max": umax,
                "tolerated_dominance_max": float(prof.get("tolerated_dominance_max", 0.30)),
            },
            "failure_reasons": failure_reasons,
        },
        "summary": (
            f"unknown {unknown_ne}/{ne_total} ({unknown_ratio*100:.1f}%); "
            f"overuse {len(tolerated_overuse)} → {severity}"
        ),
    }


def audit_role_coverage(docx_path: Path, prof: dict) -> dict:
    """7 角色全覆盖 (每个角色至少 1 个 style 在 nonempty 段出现)。
    fail: 任一 mandatory role 在 docx nonempty 段一个都没出现。
    """
    doc = Document(str(docx_path))
    c_all, c_ne, _, _ = _paragraphs_style_count(doc)
    roles = prof.get("roles") or {}
    # mandatory roles (按 task GOAL): body / title / heading / caption-figure / caption-table
    #   layout / toc / hyperlink 是 optional warn 不 fail
    mandatory = ["body", "title", "heading", "caption-figure", "caption-table"]
    optional = ["layout", "toc", "hyperlink"]

    coverage: dict[str, dict] = {}
    missing_mandatory: list[str] = []
    missing_optional: list[str] = []

    for role in mandatory + optional:
        candidates = roles.get(role) or []
        # 算覆盖时用 c_all (含空段),允许 Title/Heading 段实际套上但文本空 (常见封面/分页占位)
        found_styles = {s: c_all.get(s, 0) for s in candidates if c_all.get(s, 0) > 0}
        coverage[role] = {
            "candidates": candidates,
            "found_with_count": found_styles,
            "covered": bool(found_styles),
        }
        if not found_styles:
            if role in mandatory:
                missing_mandatory.append(role)
            else:
                missing_optional.append(role)

    severity = "pass"
    if missing_mandatory:
        severity = "fail"
    elif missing_optional:
        severity = "warn"

    return {
        "check_id": "role-coverage",
        "severity": severity,
        "found": severity != "pass",
        "details": {
            "coverage": coverage,
            "missing_mandatory": missing_mandatory,
            "missing_optional": missing_optional,
        },
        "summary": (
            f"mandatory missing {missing_mandatory}; "
            f"optional missing {missing_optional} → {severity}"
        ),
    }


def audit_body_style_concentration(docx_path: Path, prof: dict) -> dict:
    """主 body 样式占比 ≥ profile.concentration_threshold。
    算法: sum(body-role 样式段在 nonempty 段) / nonempty 段总数
    fail: 比例 < concentration_threshold。
    """
    doc = Document(str(docx_path))
    _, c_ne, _, ne_total = _paragraphs_style_count(doc)
    body_styles = (prof.get("roles") or {}).get("body") or []
    body_count = sum(c_ne.get(s, 0) for s in body_styles)
    ratio = body_count / ne_total if ne_total else 0.0
    thr = float(prof.get("concentration_threshold", 0.60))

    severity = "fail" if ratio < thr else "pass"

    return {
        "check_id": "body-style-concentration",
        "severity": severity,
        "found": severity != "pass",
        "details": {
            "body_styles": body_styles,
            "body_segment_count": body_count,
            "nonempty_total": ne_total,
            "concentration_ratio": round(ratio, 4),
            "threshold": thr,
            "per_body_style": {s: c_ne.get(s, 0) for s in body_styles},
        },
        "summary": (
            f"body {body_count}/{ne_total} ({ratio*100:.1f}%) "
            f"vs thr {thr*100:.0f}% → {severity}"
        ),
    }


def audit_style_pool_cleanliness(docx_path: Path, prof: dict) -> dict:
    """未用样式 / 总定义 ≤ pool_unused_ratio_max (warn 级)。"""
    doc = Document(str(docx_path))
    c_all, _, _, _ = _paragraphs_style_count(doc)
    defined = _defined_styles(doc)
    used = {sn for sn, n in c_all.items() if sn and n > 0}
    unused = defined - used
    total = len(defined)
    ratio = len(unused) / total if total else 0.0
    thr = float(prof.get("pool_unused_ratio_max", 0.85))

    severity = "warn" if ratio > thr else "pass"

    return {
        "check_id": "style-pool-cleanliness",
        "severity": severity,
        "found": severity != "pass",
        "details": {
            "styles_defined": total,
            "styles_used": len(used),
            "styles_unused": len(unused),
            "unused_ratio": round(ratio, 4),
            "threshold": thr,
            "unused_styles_sample": sorted(unused)[:20],
        },
        "summary": (
            f"unused {len(unused)}/{total} ({ratio*100:.1f}%) "
            f"vs thr {thr*100:.0f}% → {severity}"
        ),
    }


def audit_style_pane_filter(docx_path: Path, prof: dict) -> dict:
    """stylePaneFormatFilter 已设 (warn 级)。"""
    require = bool(prof.get("require_stylepanefilter", True))
    has = _settings_has_stylepanefilter(docx_path)
    severity = "pass"
    if require and not has:
        severity = "warn"
    return {
        "check_id": "style-pane-filter",
        "severity": severity,
        "found": severity != "pass",
        "details": {
            "require_stylepanefilter": require,
            "has_stylepanefilter": has,
        },
        "summary": f"stylePaneFormatFilter={'set' if has else 'absent'} → {severity}",
    }


# ---------- dispatcher ----------

_AUDITS = {
    "style-coherence":           audit_style_coherence,
    "role-coverage":             audit_role_coverage,
    "body-style-concentration":  audit_body_style_concentration,
    "style-pool-cleanliness":    audit_style_pool_cleanliness,
    "style-pane-filter":         audit_style_pane_filter,
}


def run_one(target: str, docx_path: Path, profile_path: Path) -> dict:
    prof = load_profile(profile_path)
    fn = _AUDITS.get(target)
    if fn is None:
        return {
            "check_id": target,
            "severity": "fail",
            "found": True,
            "details": {"error": f"unknown audit target: {target}"},
            "summary": f"unknown target {target}",
        }
    report = fn(docx_path, prof)
    report["docx_path"] = str(docx_path)
    report["profile_path"] = prof.get("_profile_path")
    return report


def _print_human(report: dict) -> None:
    print(f"=== audit-styleset · {report['check_id']} ===")
    print(f"  docx     : {Path(report['docx_path']).name}")
    print(f"  profile  : {Path(report.get('profile_path') or '').name}")
    print(f"  severity : {report['severity']}")
    print(f"  summary  : {report['summary']}")
    det = report.get("details") or {}
    for k, v in det.items():
        if isinstance(v, (dict, list)) and len(str(v)) > 200:
            s = json.dumps(v, ensure_ascii=False)[:200] + "..."
            print(f"    {k}: {s}")
        else:
            print(f"    {k}: {v}")


def _dispatch(target: str, args) -> int:
    docx_path = Path(os.path.expanduser(str(args.docx_path)))
    if not docx_path.exists():
        print(f"ERROR: docx not found: {docx_path}", file=sys.stderr)
        return 2
    profile_path = Path(os.path.expanduser(str(getattr(args, "profile", None) or DEFAULT_PROFILE_PATH)))
    report = run_one(target, docx_path, profile_path)
    if getattr(args, "json", False):
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        _print_human(report)
    if getattr(args, "report", None):
        Path(args.report).write_text(
            json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"[report] {args.report}")
    # exit code: fail=1, warn=0, pass=0
    return 1 if report.get("severity") == "fail" else 0


# ---------- registration ----------

def register(subparsers) -> None:
    """Register `audit-styleset <target>` (aliased `styleset`)."""
    p = subparsers.add_parser(
        "audit-styleset",
        help="style-set health audits (style-coherence / role-coverage / "
             "body-style-concentration / style-pool-cleanliness / style-pane-filter)",
        aliases=["styleset"],
    )
    sp = p.add_subparsers(dest="styleset_target", metavar="<target>", required=True)
    for tgt in _AUDITS.keys():
        spp = sp.add_parser(tgt, help=f"audit-styleset {tgt} (read-only)")
        spp.add_argument("docx_path", help="target docx path")
        spp.add_argument(
            "--profile",
            default=str(DEFAULT_PROFILE_PATH),
            help=f"yaml profile path (default: {DEFAULT_PROFILE_PATH.name})",
        )
        spp.add_argument("--json", action="store_true", help="print JSON to stdout")
        spp.add_argument("--report", help="also write JSON report to this path")
        spp.set_defaults(func=lambda a, _t=tgt: _dispatch(_t, a))


# ---------- standalone CLI ----------

def main():
    ap = argparse.ArgumentParser(description="audit-styleset · 5 style-set health checks")
    ap.add_argument("target", choices=list(_AUDITS.keys()),
                    help="audit target")
    ap.add_argument("docx_path", help="target docx path")
    ap.add_argument("--profile", default=str(DEFAULT_PROFILE_PATH),
                    help=f"yaml profile (default: {DEFAULT_PROFILE_PATH.name})")
    ap.add_argument("--json", action="store_true")
    ap.add_argument("--report", help="write JSON report path")
    args = ap.parse_args()

    class _Ns:
        pass
    ns = _Ns()
    ns.docx_path = args.docx_path
    ns.profile = args.profile
    ns.json = args.json
    ns.report = args.report
    rc = _dispatch(args.target, ns)
    sys.exit(rc)


if __name__ == "__main__":
    main()
