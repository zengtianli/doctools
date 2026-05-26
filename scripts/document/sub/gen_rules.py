#!/usr/bin/env python3
"""解析 改动草稿 MD，生成合并 rules JSON 用于 docx_tools.py track-changes。

含 3 道守卫：
  1. 直引号 ↔ 弯引号自动 swap
  2. title-only rule 跳过（短标题改动通常被用户反馈"标题别动"）
  3. 说明性括号"（面向...）""（含...）"等自动剥除

Usage:
  python3 gen_rules.py --drafts-dir 成果/md --docx 目标.docx --out _revise_rules.json
"""
import argparse
import json
import re
from pathlib import Path

from docx import Document

PAIR_BLOCK_RE = re.compile(
    r"\*\*原文(?:\s*#?([\w\d]+))?[^*]*\*\*[^\n]*\n((?:>\s?.*\n)+)", re.MULTILINE
)
REVISE_BLOCK_RE = re.compile(
    r"\*\*改为(?:\s*#?([\w\d]+))?[^*]*\*\*[^\n]*\n((?:>\s?.*\n)+)", re.MULTILINE
)
SECTION_RE = re.compile(r"^## 改动\s*\d+", re.MULTILINE)


def extract_quote(block):
    lines = []
    for ln in block.split("\n"):
        if not ln.startswith(">"):
            continue
        content = re.sub(r"^>\s?", "", ln).rstrip()
        if content.startswith("**") and "原文" in content[:20]:
            continue
        lines.append(content)
    return "\n".join(lines).strip()


def parse_md(md_path):
    text = md_path.read_text(encoding="utf-8")
    rules = []
    starts = [m.start() for m in SECTION_RE.finditer(text)]
    if not starts:
        return rules
    sections = []
    for i, start in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(text)
        sections.append(text[start:end])

    for sec_idx, sec in enumerate(sections, 1):
        first_line = sec.split("\n", 1)[0]
        title = first_line.replace("## 改动", "").strip()
        originals = [(m.group(1), extract_quote(m.group(2))) for m in PAIR_BLOCK_RE.finditer(sec)]
        revises = [(m.group(1), extract_quote(m.group(2))) for m in REVISE_BLOCK_RE.finditer(sec)]
        if not originals or not revises:
            continue
        orig_ids = {oid: txt for oid, txt in originals if oid}
        rev_ids = {rid: txt for rid, txt in revises if rid}
        common = orig_ids.keys() & rev_ids.keys()
        if common:
            for cid in sorted(common, key=lambda x: int(re.sub(r"\D", "", x) or 0)):
                rules.append({
                    "find": orig_ids[cid],
                    "replace": rev_ids[cid],
                    "comment": f"[{md_path.stem} 改动{sec_idx} #{cid}] {title[:40]}"
                })
            continue
        if len(originals) == len(revises):
            for (oid, otxt), (rid, rtxt) in zip(originals, revises):
                rules.append({
                    "find": otxt,
                    "replace": rtxt,
                    "comment": f"[{md_path.stem} 改动{sec_idx}] {title[:40]}"
                })
        else:
            print(f"  ⚠️ {md_path.name} 改动{sec_idx}: 原文{len(originals)}/改为{len(revises)} 不匹配，跳过")
    return rules


def try_quote_swap(s):
    """直引号 → 弯引号候选。"""
    out = []
    toggle = False
    for ch in s:
        if ch == '"':
            out.append("\u201d" if toggle else "\u201c")
            toggle = not toggle
        elif ch == "'":
            out.append("\u2019" if toggle else "\u2018")
        else:
            out.append(ch)
    return "".join(out)


def is_title(text):
    t = text.strip()
    if len(t) > 80:
        return False
    if re.match(r"^[（(][一二三四五六七八九十]+[）)]", t):
        return True
    if re.match(r"^第[一二三四五六七八九十]+[章节部分]", t):
        return True
    if any(kw in t for kw in ["子方案", "总体技术", "工作部署", "技术路线", "技术思路", "技术方案"]) and len(t) < 60:
        return True
    return False


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--drafts-dir", required=True, help="改动草稿 MD 目录")
    ap.add_argument("--docx", required=True, help="目标 docx（用于 find 存在性校验 + 引号自适配）")
    ap.add_argument("--out", required=True, help="输出 rules JSON 路径")
    ap.add_argument("--glob", default="*-改动草稿.md", help="MD 文件 glob（默认 *-改动草稿.md）")
    ap.add_argument("--no-title-skip", action="store_true", help="不跳过 title-only rule（默认跳过）")
    ap.add_argument("--no-paren-strip", action="store_true", help="不剥除说明性括号（默认剥除）")
    args = ap.parse_args()

    md_files = sorted(Path(args.drafts_dir).glob(args.glob))
    md_files = [f for f in md_files if "vs" not in f.name and "清理-清单" not in f.name]
    print(f"待处理 MD: {len(md_files)} 份")
    all_rules = []
    for f in md_files:
        rules = parse_md(f)
        print(f"  {f.name}: {len(rules)} 条")
        all_rules.extend(rules)

    # 加载 docx 全文用于 find 验证
    doc = Document(args.docx)
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                docx_text += "\n" + cell.text

    # Pass 1: 去空 + 去 find==replace + 引号 swap 守卫
    valid = []
    quote_fixed = 0
    for r in all_rules:
        if not r["find"].strip():
            continue
        if r["find"] == r["replace"]:
            continue
        if r["find"] not in docx_text:
            alt = try_quote_swap(r["find"])
            if alt != r["find"] and alt in docx_text:
                r["find"] = alt
                r["replace"] = try_quote_swap(r["replace"])
                quote_fixed += 1
        for m in re.finditer(r"〔\s*〕", r["find"] + r["replace"]):
            print(f"  ⚠️ 空〔〕：{r['comment']}")
        valid.append(r)
    if quote_fixed:
        print(f"  ✓ 引号 swap 修复 {quote_fixed} 条")

    # Pass 2: title-only 跳过守卫
    if not args.no_title_skip:
        skipped = 0
        kept = []
        for r in valid:
            if is_title(r["find"]) and is_title(r["replace"]):
                skipped += 1
                continue
            kept.append(r)
        if skipped:
            print(f"  ✓ 跳过 title-only rule: {skipped} 条")
        valid = kept

    # Pass 3: 说明性括号剥除守卫
    if not args.no_paren_strip:
        PAREN = re.compile(r"[（(](?:面向|针对|含|本节|破除|重写|说明)[^）)]{0,40}[）)]")
        stripped = 0
        for r in valid:
            new_r = PAREN.sub("", r["replace"])
            if new_r != r["replace"]:
                r["replace"] = new_r.strip()
                stripped += 1
        if stripped:
            print(f"  ✓ 剥括号说明: {stripped} 条")

    Path(args.out).parent.mkdir(parents=True, exist_ok=True)
    with open(args.out, "w", encoding="utf-8") as fp:
        json.dump(valid, fp, ensure_ascii=False, indent=2)
    print(f"\nOK -> {args.out}")
    print(f"最终 {len(valid)} 条规则（原 {len(all_rules)}，过滤 {len(all_rules)-len(valid)}）")


# ---------------- pipeline adapter ----------------
def apply_path(docx_path, args=None) -> dict:
    """pipeline-compatible adapter (跨文件 analyzer).

    docx_path = 目标 docx (用于 find 存在性校验 + 引号自适配)。
    args 透传:
      - drafts_dir (必需): 改动草稿 MD 目录
      - glob: MD glob 模式 (默认 *-改动草稿.md)
      - no_title_skip / no_paren_strip: 守卫开关
      - out / out_dir: 输出 JSON 路径
    """
    from pathlib import Path as _P
    drafts_dir = getattr(args, "drafts_dir", None) if args else None
    if not drafts_dir:
        return {"skipped": "no --drafts-dir; gen_rules needs MD drafts dir"}
    glob_pattern = getattr(args, "glob", None) or "*-改动草稿.md"
    out_path = getattr(args, "out", None)
    out_dir = getattr(args, "out_dir", None)
    if out_path:
        out = _P(out_path)
    elif out_dir:
        out = _P(out_dir) / f"_revise_rules-{_P(docx_path).stem}.json"
    else:
        out = _P(docx_path).parent / "reports" / f"_revise_rules-{_P(docx_path).stem}.json"
    out.parent.mkdir(parents=True, exist_ok=True)

    no_title_skip = bool(getattr(args, "no_title_skip", False))
    no_paren_strip = bool(getattr(args, "no_paren_strip", False))

    md_files = sorted(_P(drafts_dir).glob(glob_pattern))
    md_files = [f for f in md_files if "vs" not in f.name and "清理-清单" not in f.name]
    all_rules = []
    for f in md_files:
        all_rules.extend(parse_md(f))

    doc = Document(str(docx_path))
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                docx_text += "\n" + cell.text

    valid = []
    quote_fixed = 0
    for r in all_rules:
        if not r["find"].strip():
            continue
        if r["find"] == r["replace"]:
            continue
        if r["find"] not in docx_text:
            alt = try_quote_swap(r["find"])
            if alt != r["find"] and alt in docx_text:
                r["find"] = alt
                r["replace"] = try_quote_swap(r["replace"])
                quote_fixed += 1
        valid.append(r)
    skipped = 0
    if not no_title_skip:
        kept = []
        for r in valid:
            if is_title(r["find"]) and is_title(r["replace"]):
                skipped += 1
                continue
            kept.append(r)
        valid = kept
    stripped = 0
    if not no_paren_strip:
        PAREN = re.compile(r"[(（](?:面向|针对|含|本节|破除|重写|说明)[^)）]{0,40}[)）]")
        for r in valid:
            new_r = PAREN.sub("", r["replace"])
            if new_r != r["replace"]:
                r["replace"] = new_r.strip()
                stripped += 1

    with open(out, "w", encoding="utf-8") as fp:
        json.dump(valid, fp, ensure_ascii=False, indent=2)

    return {
        "md_files": len(md_files),
        "raw_rules": len(all_rules),
        "valid_rules": len(valid),
        "quote_fixed": quote_fixed,
        "title_skipped": skipped,
        "paren_stripped": stripped,
        "out": str(out),
    }


if __name__ == "__main__":
    main()
