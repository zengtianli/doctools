"""pipeline_lib.py — docx batch pipeline 基础设施 (distilled from qual-supply · 2026-05-26)

原 qual-supply `scripts/_pipeline_lib.py`，上提到总部 doctools sub/ 后去掉前缀下划线。
qual-supply 的 `_pipeline_lib.py` 已改为 thin shim (`from sub.pipeline_lib import *`)。

设计目标
--------
- **不破坏「一脚本一功能」铁律**: 每脚本仍单功能、独立 CLI 仍能跑、main() 不动
- **复用 parse + write + backup + lsof**: 23 脚本对同一 docx 串行调用时, Document 只 parse 1 次, 写 1 次, lsof 1 次, backup 1 次
- **跨 docx 并行**: 用 ProcessPoolExecutor 满核扇出, 避开 GIL
- **同 docx 内串行**: 写冲突避免, step 顺序由用户控制

两类 step 接口
--------------
1. **doc-based step** — 脚本暴露 `apply(doc, args) -> dict`
   - args 是 argparse.Namespace (或 dict-like with attribute access)
   - apply 改 doc 内存对象, **不读不写文件**
   - 返回 report dict (含 changed / issues / 任何 stats)
   - 适用: 17 个改 doc.paragraphs / doc.tables / pStyle 的脚本

2. **path-based step** — 脚本暴露 `apply_path(docx_path, args) -> dict`
   - 必须直接操作 zip (改 word/styles.xml / numbering.xml / media/ 等)
   - apply_path 自己负责读写 zip; pipeline 在 doc-based steps 写入磁盘后才调
   - 返回 report dict
   - 适用: 6 个 zip-write 脚本 (freeze_heading_numbers / freeze_all_fields /
     strip_style_outlinelvl / audit_heading_numbers / audit_word_fields /
     relink_images_from_source)

pipeline 执行顺序
-----------------
1. lsof_check 1 次
2. backup 1 次 (除非 backup_once=False)
3. 把 steps 按声明顺序分两段: pre = 所有 doc-based, post = 所有 path-based
   (若混合声明, 用户负责把 path-based 放后面; driver 不重排)
4. 依次:
   - Document(path) parse 1 次
   - for step in doc-based: step.apply(doc, args)
   - doc.save(path) 1 次
   - for step in path-based: step.apply_path(path, args)
5. 返回 {step_name: report, "_meta": {timing, backup, ...}}

并行模式 (跨 docx)
------------------
ProcessPoolExecutor(max_workers=min(N_docx, cpu_count))
每个 docx 一个 process 跑完整 pipeline

step_dir 参数
-------------
load_step() 的 step_dir 参数指定从哪个目录动态加载脚本模块。
默认 None = 从 cwd/scripts/ 找（qual-supply 兼容）。
可传 --step-dir 指向任意项目的 scripts/ 目录。
"""

from __future__ import annotations

import argparse
import importlib
import importlib.util
import json
import os
import shutil
import subprocess
import sys
import time
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any, Callable

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


# ----------------- helpers -----------------

def lsof_check(docx_path: Path) -> str | None:
    """返回非 None 即被占用; None 表示空闲"""
    try:
        out = subprocess.run(
            ["lsof", str(docx_path)],
            capture_output=True, text=True, timeout=5,
        )
        if out.returncode == 0 and out.stdout.strip():
            return out.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def make_backup_path(src: Path) -> Path:
    """统一备份命名: <stem>.bak-N-YYYY-MM-DD<suffix>, N 自增"""
    today = date.today().isoformat()
    parent, stem, suffix = src.parent, src.stem, src.suffix
    n = 1
    while True:
        cand = parent / f"{stem}.bak-{n}-{today}{suffix}"
        if not cand.exists():
            return cand
        n += 1


# ----------------- step loading -----------------

@dataclass
class LoadedStep:
    name: str
    kind: str  # "doc" or "path"
    fn: Callable[..., dict]
    module: Any

    def call(self, *, doc=None, docx_path: Path | None = None, args=None) -> dict:
        if self.kind == "doc":
            return self.fn(doc, args)
        return self.fn(docx_path, args)


def load_step(name: str, step_dir: Path | str | None = None) -> LoadedStep:
    """Load a step module by name.

    name = 脚本名(无.py).
    step_dir = 从哪个目录加载脚本。默认 None → cwd/scripts/（qual-supply 兼容）。
    Dynamic import; 优先 apply(doc, args), 否则 apply_path(path, args).

    Built-in steps (audit-styleset-all / split-by-h1) take precedence over
    on-disk scripts; they reuse the pipeline's already-parsed doc to avoid
    re-parsing the same 53MB docx 5-6 times.
    """
    if name in _BUILTIN_STEPS:
        kind, fn = _BUILTIN_STEPS[name]
        return LoadedStep(name=name, kind=kind, fn=fn, module=None)

    # Determine search path
    if step_dir is not None:
        search_dirs = [Path(step_dir).resolve()]
    else:
        # Default: cwd/scripts/ (qual-supply compat) + cwd itself
        cwd = Path.cwd()
        search_dirs = [cwd / "scripts", cwd]

    # Try loading from search_dirs via spec_from_file_location
    for d in search_dirs:
        candidate = d / f"{name}.py"
        if candidate.is_file():
            alias = f"_pipeline_step__{name}"
            spec = importlib.util.spec_from_file_location(alias, str(candidate))
            if spec and spec.loader:
                mod = importlib.util.module_from_spec(spec)
                # Insert directory for relative sibling imports within step scripts
                dir_str = str(d)
                inserted = dir_str not in sys.path
                if inserted:
                    sys.path.insert(0, dir_str)
                try:
                    spec.loader.exec_module(mod)
                finally:
                    if inserted:
                        try:
                            sys.path.remove(dir_str)
                        except ValueError:
                            pass
                if hasattr(mod, "apply"):
                    return LoadedStep(name=name, kind="doc", fn=mod.apply, module=mod)
                if hasattr(mod, "apply_path"):
                    return LoadedStep(name=name, kind="path", fn=mod.apply_path, module=mod)
                raise AttributeError(
                    f"script '{name}' ({candidate}) has neither apply(doc,args) nor apply_path(path,args)"
                )

    # Fallback: importlib.import_module (original qual-supply behaviour)
    try:
        module = importlib.import_module(name)
    except ImportError:
        module = importlib.import_module(f"scripts.{name}")
    if hasattr(module, "apply"):
        return LoadedStep(name=name, kind="doc", fn=module.apply, module=module)
    if hasattr(module, "apply_path"):
        return LoadedStep(name=name, kind="path", fn=module.apply_path, module=module)
    raise AttributeError(
        f"script '{name}' has neither apply(doc,args) nor apply_path(path,args)"
    )


# ----------------- built-in steps (single-parse reuse) -----------------
# Registered names short-circuit on-disk script loading; each reuses the
# pipeline's already-parsed `doc` to amortize 53MB-docx parse cost across
# multiple analyses.

def _builtin_audit_styleset_all(doc, args) -> dict:
    """Run all 5 audit-styleset checks against the already-parsed doc.

    Resolves profile path from args.styleset_profile or default profile.
    Reads zip-level word/settings.xml once (for stylePaneFormatFilter check).
    """
    from . import audit_styleset
    docx_path = Path(str(getattr(args, "docx", "")))
    profile_path = getattr(args, "styleset_profile", None) or str(
        audit_styleset.DEFAULT_PROFILE_PATH
    )
    return audit_styleset.run_all_on_doc(doc, docx_path, Path(profile_path))


def _builtin_health_diagnose(doc, args) -> dict:
    """Run docx health 8-check diagnose against source path.

    NOTE: most health checks parse docx independently via zipfile and lxml
    (not python-docx), so they don't benefit from the pipeline's already-
    parsed `doc`. We accept `doc` for API uniformity but pass docx_path
    through to the HealthChecker.
    """
    from . import health
    docx_path = Path(str(getattr(args, "docx", "")))
    workers = int(getattr(args, "health_workers", 8) or 8)
    checker = health.HealthChecker(docx_path, workers=workers)
    results = checker.run_all()
    # severity rollup
    sev_levels = []
    for cid, r in results.items():
        if r.get("found"):
            sev_levels.append(health.SEVERITY.get(cid, "?"))
    if "High" in sev_levels:
        overall = "fail"
    elif sev_levels:
        overall = "warn"
    else:
        overall = "pass"
    return {
        "overall_severity": overall,
        "checks": results,
        "found_count": sum(1 for r in results.values() if r.get("found")),
    }


def _builtin_image_extract(doc, args) -> dict:
    """Extract embedded images from docx by neighboring caption.

    Read-only — does not mutate parsed doc. Reuses parsed doc only for API
    uniformity (image_extract parses zip directly via lxml + zipfile).

    Args:
      args.docx — source docx path (injected by pipeline driver)
      args.image_extract_out_dir — explicit out dir; if None → <docx-parent>/images/
      args.image_extract_quiet — suppress per-file log lines (default True in pipeline)
    """
    from . import image_extract
    docx_path = Path(str(getattr(args, "docx", "")))
    out_dir = getattr(args, "image_extract_out_dir", None)
    if not out_dir:
        out_dir = docx_path.parent / "images"
    else:
        out_dir = Path(str(out_dir))
    quiet = bool(getattr(args, "image_extract_quiet", True))
    exit_code = image_extract.extract_images(docx_path, out_dir, quiet=quiet)
    # Count what landed
    n_files = 0
    if out_dir.is_dir():
        n_files = sum(1 for p in out_dir.iterdir() if p.is_file())
    return {
        "exit_code": exit_code,
        "out_dir": str(out_dir),
        "files_in_out_dir": n_files,
    }


def _builtin_table_extract(doc, args) -> dict:
    """Extract each table in docx as an independent minimal docx (caption-named).

    Read-only against source — does not mutate parsed doc. Reuses parsed doc
    for plan_extracts() to avoid re-parsing.

    Args:
      args.docx — source docx path (injected by pipeline driver)
      args.table_extract_out_dir — explicit out dir; if None → <docx-parent>/tables/
      args.table_extract_name_pattern — filename pattern (default '{stem}.docx')
      args.table_extract_dry_run — print plan only, don't write files
    """
    from . import extract_tables
    docx_path = Path(str(getattr(args, "docx", "")))
    out_dir = getattr(args, "table_extract_out_dir", None)
    if not out_dir:
        out_dir = docx_path.parent / "tables"
    else:
        out_dir = Path(str(out_dir))
    name_pattern = getattr(args, "table_extract_name_pattern", None) or "{stem}.docx"
    dry_run = bool(getattr(args, "table_extract_dry_run", False))
    return extract_tables.run_extract(
        src_docx=docx_path,
        out_dir=out_dir,
        dry_run=dry_run,
        name_pattern=name_pattern,
        doc=doc,
    )


def _builtin_split_by_h1(doc, args) -> dict:
    """Split docx by H1 reusing the already-parsed doc for slice planning.

    Required args: split_out_dir.
    Optional args: include_frontmatter (bool), allow_no_h1 (bool),
                   split_name_pattern (str), split_dry_run (bool).
    Raises RuntimeError on h1_count=0 + not allow_no_h1 (per fail-fast contract).
    """
    from . import split_by_h1
    docx_path = Path(str(getattr(args, "docx", "")))
    out_dir = getattr(args, "split_out_dir", None)
    if not out_dir:
        raise RuntimeError("split-by-h1 step requires --split-out-dir")
    rep = split_by_h1.run_split(
        src_docx=docx_path,
        out_dir=Path(str(out_dir)),
        include_frontmatter=bool(getattr(args, "include_frontmatter", False)),
        allow_no_h1=bool(getattr(args, "allow_no_h1", False)),
        dry_run=bool(getattr(args, "split_dry_run", False)),
        name_pattern=getattr(args, "split_name_pattern", None)
            or "{idx:02d}-{title}.docx",
        doc=doc,
    )
    # propagate fail-fast exit code 3 (0 H1 detected) as exception so pipeline
    # surfaces it instead of silent partial success
    if rep.get("exit_code") == 3:
        raise RuntimeError(rep.get("error", "0 Heading-1 detected"))
    return rep


# (name → (kind, callable)); kind = "doc" → fn(doc, args)
_BUILTIN_STEPS: dict[str, tuple[str, Callable[..., dict]]] = {
    "audit-styleset-all": ("doc", _builtin_audit_styleset_all),
    "split-by-h1":        ("doc", _builtin_split_by_h1),
    "health-diagnose":    ("doc", _builtin_health_diagnose),
    "image-extract":      ("doc", _builtin_image_extract),
    "table-extract":      ("doc", _builtin_table_extract),
}

# Read-only built-ins (don't mutate parsed doc → no need to save source docx).
# When ALL doc-based steps are non-mutating, run_pipeline skips doc.save().
_NON_MUTATING_STEPS: set[str] = {
    "audit-styleset-all",
    "split-by-h1",
    "health-diagnose",
    "image-extract",
    "table-extract",
}


def is_builtin_step(name: str) -> bool:
    return name in _BUILTIN_STEPS


def is_non_mutating_step(name: str) -> bool:
    return name in _NON_MUTATING_STEPS


# ----------------- single docx pipeline -----------------

def run_pipeline(
    docx_path: Path | str,
    step_names: list[str],
    args: argparse.Namespace | None = None,
    backup_once: bool = True,
    lsof_once: bool = True,
    dry_run: bool = False,
    no_backup: bool = False,
    step_dir: Path | str | None = None,
) -> dict:
    """对单个 docx 顺序执行 steps; parse/save/lsof/backup 各 1 次.

    step_dir: 传给 load_step，指定 step 脚本目录。None = cwd/scripts/（qual-supply 兼容）。
    """
    docx_path = Path(docx_path).resolve()
    if not docx_path.is_file():
        raise FileNotFoundError(docx_path)

    t0 = time.perf_counter()
    report: dict[str, Any] = {"docx": str(docx_path), "steps": {}, "timing": {}}

    # lsof
    if lsof_once and not dry_run:
        t_ls = time.perf_counter()
        occ = lsof_check(docx_path)
        report["timing"]["lsof"] = time.perf_counter() - t_ls
        if occ:
            raise RuntimeError(
                f"docx 被占用 (Word/WPS): {docx_path}\n{occ}"
            )

    # load all steps first (fail fast)
    t_load = time.perf_counter()
    loaded = [load_step(n, step_dir=step_dir) for n in step_names]
    report["timing"]["load_steps"] = time.perf_counter() - t_load

    # backup
    backup_path = None
    if backup_once and not dry_run and not no_backup:
        backup_path = make_backup_path(docx_path)
        shutil.copy2(docx_path, backup_path)
        report["backup"] = str(backup_path)

    # default args
    if args is None:
        args = argparse.Namespace()
    # inject standard fields if missing
    if not hasattr(args, "docx"):
        args.docx = docx_path
    if not hasattr(args, "dry_run"):
        args.dry_run = dry_run
    if not hasattr(args, "no_backup"):
        args.no_backup = True  # pipeline 已经备份过, step 内别再备
    if not hasattr(args, "report"):
        args.report = None

    # split steps: doc-based first, path-based last
    doc_steps = [s for s in loaded if s.kind == "doc"]
    path_steps = [s for s in loaded if s.kind == "path"]
    # warn if order interleaved
    declared_kinds = [s.kind for s in loaded]
    if declared_kinds != [s.kind for s in doc_steps + path_steps]:
        report["warnings"] = [
            "step 声明顺序混合 doc/path; pipeline 已自动重排为 doc-first, path-last"
        ]

    # doc-based pass
    if doc_steps:
        t_parse = time.perf_counter()
        doc = Document(str(docx_path))
        report["timing"]["parse"] = time.perf_counter() - t_parse
        for s in doc_steps:
            t_s = time.perf_counter()
            try:
                rep = s.call(doc=doc, args=args)
            except Exception as exc:
                rep = {"error": repr(exc)}
            report["steps"][s.name] = rep
            report["timing"][f"step:{s.name}"] = time.perf_counter() - t_s
        # Skip save when all doc-based steps are non-mutating built-ins (audit
        # / split = read-only against source). Avoids touching mtime + bytes
        # of the source 53MB docx (and lets baseline split byte-for-byte match).
        all_non_mutating = all(is_non_mutating_step(s.name) for s in doc_steps)
        if not dry_run and not all_non_mutating:
            t_w = time.perf_counter()
            doc.save(str(docx_path))
            report["timing"]["save"] = time.perf_counter() - t_w

    # path-based pass (after save)
    for s in path_steps:
        t_s = time.perf_counter()
        try:
            rep = s.call(docx_path=docx_path, args=args)
        except Exception as exc:
            rep = {"error": repr(exc)}
        report["steps"][s.name] = rep
        report["timing"][f"step:{s.name}"] = time.perf_counter() - t_s

    report["timing"]["total"] = time.perf_counter() - t0
    return report


# ----------------- parallel across docs -----------------

def _worker(payload: dict) -> dict:
    """ProcessPoolExecutor worker entry"""
    # Reconstruct args namespace from payload for built-in step options
    ns_args = argparse.Namespace(
        docx=Path(payload["docx"]),
        dry_run=payload.get("dry_run", False),
        no_backup=True,
        report=None,
        styleset_profile=payload.get("styleset_profile"),
        split_out_dir=payload.get("split_out_dir"),
        split_name_pattern=payload.get("split_name_pattern"),
        include_frontmatter=payload.get("include_frontmatter", False),
        allow_no_h1=payload.get("allow_no_h1", False),
        split_dry_run=payload.get("split_dry_run", False),
        image_extract_out_dir=payload.get("image_extract_out_dir"),
        image_extract_quiet=payload.get("image_extract_quiet", True),
        table_extract_out_dir=payload.get("table_extract_out_dir"),
        table_extract_name_pattern=payload.get("table_extract_name_pattern"),
        table_extract_dry_run=payload.get("table_extract_dry_run", False),
    )
    return run_pipeline(
        docx_path=payload["docx"],
        step_names=payload["steps"],
        args=ns_args,
        backup_once=payload.get("backup_once", True),
        lsof_once=payload.get("lsof_once", True),
        dry_run=payload.get("dry_run", False),
        no_backup=payload.get("no_backup", False),
        step_dir=payload.get("step_dir"),
    )


def run_pipeline_parallel(
    docx_list: list[Path | str],
    step_names: list[str],
    max_workers: int | None = None,
    dry_run: bool = False,
    no_backup: bool = False,
    step_dir: Path | str | None = None,
    builtin_opts: dict | None = None,
) -> dict[str, dict]:
    """对 N docx 并行执行 pipeline (跨 docx process-level 并行).

    builtin_opts: forwarded to workers for built-in step options
                  (styleset_profile / split_out_dir / include_frontmatter / ...).
    """
    if not docx_list:
        return {}
    if max_workers is None:
        max_workers = min(len(docx_list), os.cpu_count() or 4)
    builtin_opts = builtin_opts or {}
    payloads = [
        {
            "docx": str(d),
            "steps": step_names,
            "dry_run": dry_run,
            "no_backup": no_backup,
            "step_dir": str(step_dir) if step_dir else None,
            **builtin_opts,
        }
        for d in docx_list
    ]
    results: dict[str, dict] = {}
    if max_workers == 1 or len(docx_list) == 1:
        for p in payloads:
            try:
                results[p["docx"]] = _worker(p)
            except Exception as exc:
                results[p["docx"]] = {"error": repr(exc)}
        return results
    with ProcessPoolExecutor(max_workers=max_workers) as ex:
        fut_map = {ex.submit(_worker, p): p["docx"] for p in payloads}
        for fut in as_completed(fut_map):
            d = fut_map[fut]
            try:
                results[d] = fut.result()
            except Exception as exc:
                results[d] = {"error": repr(exc)}
    return results


# ----------------- pretty timing reporter -----------------

def format_timing_table(results: dict[str, dict]) -> str:
    """渲染墙钟统计表"""
    lines = []
    lines.append("=" * 78)
    lines.append("PIPELINE TIMING REPORT")
    lines.append("=" * 78)

    # collect all step names
    step_set: list[str] = []
    seen = set()
    for r in results.values():
        if not isinstance(r, dict) or "steps" not in r:
            continue
        for k in r.get("steps", {}):
            if k not in seen:
                step_set.append(k)
                seen.add(k)

    # header
    lines.append(f"{'docx':40s}  {'parse':>7s}  {'save':>7s}  {'total':>8s}")
    lines.append("-" * 78)
    total_wall = 0.0
    for docx, r in results.items():
        if not isinstance(r, dict) or "timing" not in r:
            lines.append(f"{Path(docx).name[:40]:40s}  ERROR: {r}")
            continue
        t = r["timing"]
        name = Path(docx).name[:40]
        parse_t = t.get("parse", 0.0)
        save_t = t.get("save", 0.0)
        tot = t.get("total", 0.0)
        total_wall = max(total_wall, tot)  # parallel: max
        lines.append(
            f"{name:40s}  {parse_t:7.3f}  {save_t:7.3f}  {tot:8.3f}"
        )
    lines.append("-" * 78)
    # per-step timing (sum across docs)
    if step_set:
        lines.append("Per-step (sum across all docs):")
        for step in step_set:
            tot = 0.0
            cnt = 0
            for r in results.values():
                if isinstance(r, dict) and "timing" in r:
                    v = r["timing"].get(f"step:{step}")
                    if v is not None:
                        tot += v
                        cnt += 1
            lines.append(f"  {step:50s}  {tot:7.3f}s  ({cnt} docs)")
    lines.append("-" * 78)
    lines.append(
        f"Max wall (parallel mode dominator): {total_wall:.3f}s "
        f"over {len(results)} docs"
    )
    lines.append("=" * 78)
    return "\n".join(lines)


__all__ = [
    "lsof_check",
    "make_backup_path",
    "load_step",
    "run_pipeline",
    "run_pipeline_parallel",
    "format_timing_table",
]
