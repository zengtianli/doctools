#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""strip.py — strip_* 家族七合一（2026-07-31 家族折叠）

子命令 ↔ 原脚本（函数体逐字搬移；模块级 apply/apply_path/main 改名
apply_<sub>/apply_path_<sub>/main_<sub>，其余公有名一律保留）：

    bookmarks         ← strip_bookmarks.py
    revisions         ← strip_revisions.py
    empty-captions    ← strip_empty_captions.py
    orphan-media      ← strip_orphan_media.py（scan_orphans / rewrite_skip 公有名保留：
                                               split.py 与 health.py 靠它）
    outlinelvl        ← strip_outlinelvl_from_captions.py
    style-outlinelvl  ← strip_style_outlinelvl.py
    doc-protection    ← strip_doc_protection.py

各子命令 CLI 与原独立脚本逐字一致：python3 sub/strip.py <sub> <docx> [flags…]。
原脚本完整文档保留在各段 _DOC_* 常量（--help 仍引用）。
退役原件在 ~/.Trash/consolidation-20260731/strip/（含 MANIFEST.md）。
"""
from __future__ import annotations

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[3] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py
from docx_parts import (  # noqa: E402  部件完整性断言
    DEFAULT_ALLOW_CHANGED, PartIntegrityError, assert_parts_intact, diff_parts,
)

# sub/ 自身进 sys.path —— docx_cli 的 _dispatch 用 spec_from_file_location 加载,
# 不带脚本目录, 裸 import _cli_common 会 ImportError (append 不是 insert(0))
_sys.path.append(str(_Path(__file__).resolve().parent))
import _cli_common as _cc  # noqa: E402  家族 main() 样板 SSOT

import argparse  # noqa: E402
import re  # noqa: E402
import shutil  # noqa: E402
import sys  # noqa: E402
import zipfile  # noqa: E402
from collections import Counter  # noqa: E402
from pathlib import Path  # noqa: E402

from lxml import etree  # noqa: E402

try:
    from docx import Document
except ImportError:
    print("[ERR] 缺 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(2)
from docx.oxml.ns import qn  # noqa: E402

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{NS}}}"


# ══════════ bookmarks ← strip_bookmarks.py ══════════

DEFAULT_PREFIXES = ["_Toc", "_Ref", "_Hlk"]


def parse_prefixes(s: str) -> list[str] | None:
    s = (s or "").strip()
    if not s:
        return list(DEFAULT_PREFIXES)
    if s.lower() == "all":
        return None  # None 表示删全部
    return [p.strip() for p in s.split(",") if p.strip()]


def _should_remove(name: str, prefixes: list[str] | None) -> bool:
    """name 是 bookmark name; prefixes=None 表示删全部."""
    if name == "_GoBack":
        # Word 默认自动 bookmark, 任何场景都安全清掉
        return True
    if prefixes is None:
        return True
    for p in prefixes:
        if name.startswith(p):
            return True
    return False


def strip(doc, prefixes: list[str] | None) -> dict:
    """从 doc 移除匹配的 bookmark, 返回 report dict."""
    body = doc.element.body

    # 1. 扫所有 bookmarkStart, 按 id 收集要删的 (按 name prefix 判)
    ids_to_delete: set[str] = set()
    removed_names: list[str] = []
    kept_names: list[str] = []

    for bs in body.iter(f"{W}bookmarkStart"):
        bid = bs.get(f"{W}id")
        name = bs.get(f"{W}name") or ""
        if bid is None:
            continue
        if _should_remove(name, prefixes):
            ids_to_delete.add(bid)
            removed_names.append(name)
        else:
            kept_names.append(name)

    # 2. 删 bookmarkStart
    removed_start = 0
    for bs in list(body.iter(f"{W}bookmarkStart")):
        bid = bs.get(f"{W}id")
        if bid in ids_to_delete:
            parent = bs.getparent()
            if parent is not None:
                parent.remove(bs)
                removed_start += 1

    # 3. 删配对 bookmarkEnd
    removed_end = 0
    for be in list(body.iter(f"{W}bookmarkEnd")):
        bid = be.get(f"{W}id")
        if bid in ids_to_delete:
            parent = be.getparent()
            if parent is not None:
                parent.remove(be)
                removed_end += 1

    by_prefix_removed: Counter = Counter()
    for n in removed_names:
        if n == "_GoBack":
            by_prefix_removed["_GoBack"] += 1
            continue
        matched = False
        for p in ("_Toc", "_Ref", "_Hlk", "_Hyperlink", "OLE_LINK"):
            if n.startswith(p):
                by_prefix_removed[p] += 1
                matched = True
                break
        if not matched:
            by_prefix_removed["user_defined"] += 1

    return {
        "bookmarks_removed": len(ids_to_delete),
        "starts": removed_start,
        "ends": removed_end,
        "removed_by_prefix": dict(by_prefix_removed),
        "kept_count": len(kept_names),
        "kept_sample": kept_names[:20],
    }


# ---------------- pipeline adapter ----------------
def apply_bookmarks(doc, args=None) -> dict:
    """pipeline 调用 — 改 doc 内存对象, 不写盘."""
    prefixes_attr = getattr(args, "bookmark_prefixes", None) if args else None
    if prefixes_attr is None:
        prefixes = list(DEFAULT_PREFIXES)
    elif isinstance(prefixes_attr, str):
        prefixes = parse_prefixes(prefixes_attr)
    else:
        prefixes = list(prefixes_attr) if prefixes_attr else list(DEFAULT_PREFIXES)
    return strip(doc, prefixes)


def main_bookmarks(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Strip bookmarks (default: _Toc/_Ref/_Hlk/_GoBack) from docx.")
    ap.add_argument("docx", type=Path)
    ap.add_argument("--prefixes", default=",".join(DEFAULT_PREFIXES),
                    help="Comma-separated bookmark name prefixes to strip (e.g. _Toc,_Ref,_Hlk). "
                         "'all' = strip every bookmark. _GoBack is always stripped. "
                         f"Default: {','.join(DEFAULT_PREFIXES)}")
    _cc.add_write_flags(
        ap,
        dry_run_help="Report only, do not modify file",
        no_backup_help="Skip writing .bak-N-<date>.docx",
        report_help="Write JSON report to this path",
    )
    args = ap.parse_args(argv)

    if not args.docx.exists():
        print(f"[ERR] 文件不存在: {args.docx}", file=sys.stderr)
        return 1

    prefixes = parse_prefixes(args.prefixes)
    p = args.docx.resolve()

    if not args.dry_run:
        occ = _cc.lsof_check(p)
        if occ:
            print(f"[ERR] 文件被占用 (Word/WPS 没关?):\n{occ}", file=sys.stderr)
            sys.exit(2)

    doc = Document(str(p))
    report = strip(doc, prefixes)
    report["docx"] = str(p)
    report["dry_run"] = args.dry_run
    report["prefixes"] = "all" if prefixes is None else prefixes

    _cc.print_json(report)

    _cc.write_report(report, args.report)

    if args.dry_run:
        print("[dry-run] no file written")
        return 0

    if not args.no_backup:
        bp = _cc.make_backup(p)
        print(f"[backup] {bp.name}")

    doc.save(str(p))
    print(f"[done] wrote {p}")
    return 0


# ══════════ revisions ← strip_revisions.py ══════════

_DOC_REVISIONS = """
strip_revisions.py
==================

单功能描述
----------
从 docx 移除所有「修订相关」残留, 让合稿前文档干净:
  1. 接受所有插入 ``<w:ins>`` (展开内部 run, 删 ins 包裹)
  2. 接受所有删除 ``<w:del>`` (整体移除 del 元素 + 内部 run, 即删除"被标记删除"的内容)
  3. 删 ``<w:commentRangeStart>`` / ``<w:commentRangeEnd>`` / ``<w:commentReference>``
     (commentReference 自身移除, 保留外层 run 防破坏段结构)
  4. 清空 ``word/comments.xml`` (留个 valid 空壳防 Word 抱怨; --keep-comments 关闭)
  5. 关 trackChanges (``word/settings.xml`` 移除 ``<w:trackChanges/>``)

不做
----
- 不删 rsid / rsidR / rsidRDefault 等修订追踪 ID 属性 (数量巨大, 不影响合稿)
- 不动 sectPr / pPr 中其他修订属性 (pPrChange / rPrChange 等暂忽略, 接受范围仅 ins/del/comments/trackChanges)
- 不裸解 zip 改 XML 字符串 — settings.xml / comments.xml 走 lxml.etree 解析改写

触发场景
--------
- 合稿前 docx 还残留 Word 修订模式产物 (ins/del/批注/trackChanges 开关)
- 需要 "扁平化" 文档让后续 normalize / renumber 等脚本不被 ins/del 包裹打乱段结构

CLI
---
    python3 scripts/strip_revisions.py <docx_path> \\
        [--mode accept-all] [--keep-comments] [--dry-run] [--no-backup] [--report <json>]

  --mode accept-all   接受所有修订 (ins 展开 + del 删除); 默认且当前唯一支持模式
  --keep-comments     保留 word/comments.xml 原内容 (默认清空)
  --dry-run           列每类计数不写
  --no-backup         不备份 .bak-N-YYYY-MM-DD.docx
  --report <json>     写 JSON 报告

Pipeline ready
--------------
- ``apply(doc, args) -> dict`` : 改 doc 内存对象 (body 内 ins/del/comment*)
- ``apply_path(docx_path, args) -> dict`` : 改 zip 内 word/comments.xml + word/settings.xml

约束
----
- python-docx + lxml, 禁裸 zip XML 字符串改 (settings/comments 走 etree.fromstring)
- 默认 --backup, 写前 lsof 自检 Word/WPS 占用 (占用立即退出)
- 一脚本一功能: 仅处理修订/批注/trackChanges, 不动样式/编号/段顺序
"""

# ---------------- 内存改 (body 内 ins/del/comment*) ----------------

def _strip_in_element(root) -> dict:
    """对一个 lxml element (body 或 header/footer root) 处理 ins/del/comment*"""
    ins_count = 0
    del_count = 0
    crs_count = 0
    cre_count = 0
    cref_count = 0

    # 1. <w:ins>: 展开内部子节点到父
    for ins in list(root.iter(f"{W}ins")):
        parent = ins.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ins)
        for child in list(ins):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)
        ins_count += 1

    # 2. <w:del>: 整体删除 (内部 run 一起走)
    for del_el in list(root.iter(f"{W}del")):
        parent = del_el.getparent()
        if parent is None:
            continue
        parent.remove(del_el)
        del_count += 1

    # 3. comment 引用 三种全删
    for tag, counter_name in [
        (f"{W}commentRangeStart", "crs"),
        (f"{W}commentRangeEnd", "cre"),
        (f"{W}commentReference", "cref"),
    ]:
        for el in list(root.iter(tag)):
            parent = el.getparent()
            if parent is None:
                continue
            parent.remove(el)
            if counter_name == "crs":
                crs_count += 1
            elif counter_name == "cre":
                cre_count += 1
            else:
                cref_count += 1

    return {
        "ins_accepted": ins_count,
        "del_accepted": del_count,
        "comment_range_start_removed": crs_count,
        "comment_range_end_removed": cre_count,
        "comment_reference_removed": cref_count,
    }


def apply_revisions(doc, args=None) -> dict:
    """改 doc 内存对象的 body — pipeline 接口"""
    mode = getattr(args, "revision_mode", "accept-all") if args else "accept-all"
    body = doc.element.body
    result = _strip_in_element(body)
    result["mode"] = mode
    return result


# ---------------- zip 改 (comments.xml + settings.xml) ----------------

def apply_path_revisions(docx_path, args=None) -> dict:
    """改 word/comments.xml (清空, 除非 keep_comments) + word/settings.xml (删 trackChanges)"""
    keep_comments = bool(getattr(args, "keep_comments", False)) if args else False
    docx_path = Path(docx_path)
    tmp = docx_path.with_suffix(docx_path.suffix + ".tmp")

    track_changes_removed = 0
    comments_emptied = False
    has_comments = False
    has_settings = False

    parser = etree.XMLParser(remove_blank_text=False)

    with zipfile.ZipFile(str(docx_path), 'r') as zin:
        with zipfile.ZipFile(str(tmp), 'w', zipfile.ZIP_DEFLATED) as zout:
            for it in zin.infolist():
                data = zin.read(it.filename)
                if it.filename == "word/settings.xml":
                    has_settings = True
                    try:
                        root = etree.fromstring(data, parser=parser)
                        for tc in root.findall(f".//{W}trackChanges"):
                            tc.getparent().remove(tc)
                            track_changes_removed += 1
                        data = etree.tostring(
                            root, xml_declaration=True,
                            encoding='UTF-8', standalone=True,
                        )
                    except etree.XMLSyntaxError:
                        pass
                elif it.filename == "word/comments.xml" and not keep_comments:
                    has_comments = True
                    empty_xml = (
                        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
                        b'<w:comments xmlns:w="' + NS.encode() + b'"/>'
                    )
                    data = empty_xml
                    comments_emptied = True
                zout.writestr(it, data)

    # 部件完整性断言(move 前:docx_path 仍是未动源件=天然基线;炸则 tmp 被清、源件无损)。
    # settings.xml 不在 DEFAULT 白名单,显式报备;comments.xml/document.xml 已在。
    # CLI main 与 pipeline 都汇到本函数,一处覆盖双入口。
    try:
        assert_parts_intact(docx_path, tmp,
                            allow_changed=set(DEFAULT_ALLOW_CHANGED) | {"word/settings.xml"},
                            verbose=False)
    except Exception:
        tmp.unlink(missing_ok=True)
        raise
    shutil.move(str(tmp), str(docx_path))
    return {
        "trackchanges_removed": track_changes_removed,
        "comments_emptied": comments_emptied,
        "has_comments_xml": has_comments,
        "has_settings_xml": has_settings,
    }


# ---------------- pre-scan (用 zip + lxml 不改, 算 before counts) ----------------

def scan_counts(docx_path: Path) -> dict:
    """只读扫 before 数量, 供报告用"""
    counts = {
        "ins": 0,
        "del": 0,
        "comment_range_start": 0,
        "comment_range_end": 0,
        "comment_reference": 0,
        "trackChanges": 0,
        "comments_xml_present": False,
    }
    try:
        with zipfile.ZipFile(str(docx_path)) as z:
            names = set(z.namelist())
            if "word/document.xml" in names:
                doc = etree.fromstring(z.read("word/document.xml"))
                counts["ins"] = len(doc.findall(f".//{W}ins"))
                counts["del"] = len(doc.findall(f".//{W}del"))
                counts["comment_range_start"] = len(doc.findall(f".//{W}commentRangeStart"))
                counts["comment_range_end"] = len(doc.findall(f".//{W}commentRangeEnd"))
                counts["comment_reference"] = len(doc.findall(f".//{W}commentReference"))
            if "word/settings.xml" in names:
                s = etree.fromstring(z.read("word/settings.xml"))
                counts["trackChanges"] = len(s.findall(f".//{W}trackChanges"))
            counts["comments_xml_present"] = "word/comments.xml" in names
    except (zipfile.BadZipFile, etree.XMLSyntaxError, KeyError):
        pass
    return counts


# ---------------- CLI ----------------

def main_revisions() -> int:
    parser = argparse.ArgumentParser(
        description=_DOC_REVISIONS,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("docx", type=Path)
    parser.add_argument(
        "--mode", choices=["accept-all"], default="accept-all",
        help="目前仅支持 accept-all (接受所有修订: ins 展开 + del 删除)",
    )
    parser.add_argument("--keep-comments", action="store_true",
                        help="保留 word/comments.xml (默认清空)")
    _cc.add_write_flags(parser)
    args = parser.parse_args()

    # 兼容 pipeline adapter 名字
    args.revision_mode = args.mode

    if not args.docx.exists():
        print(f"[ERR] 文件不存在: {args.docx}", file=sys.stderr)
        return 2

    if not args.dry_run:
        occ = _cc.lsof_check(args.docx)
        if occ:
            print(f"[ERR] 文件被占用 (Word/WPS 在开?), 立即停止:\n{occ}", file=sys.stderr)
            return 3

    before = scan_counts(args.docx)
    print(f"[INFO] 扫描 {args.docx.name}")
    print(f"  [before] ins={before['ins']} del={before['del']} "
          f"crs={before['comment_range_start']} cre={before['comment_range_end']} "
          f"cref={before['comment_reference']} trackChanges={before['trackChanges']} "
          f"comments_xml={'yes' if before['comments_xml_present'] else 'no'}")

    # ---- 1) 内存改 body
    doc = Document(str(args.docx))
    body_result = apply_revisions(doc, args)

    report = {
        "docx": str(args.docx.resolve()),
        "dry_run": args.dry_run,
        "mode": args.mode,
        "keep_comments": args.keep_comments,
        "backup": None,
        "wrote": False,
        "before": before,
        "body_changes": body_result,
        "zip_changes": None,
    }

    print(f"  [body] ins_accepted={body_result['ins_accepted']} "
          f"del_accepted={body_result['del_accepted']} "
          f"crs={body_result['comment_range_start_removed']} "
          f"cre={body_result['comment_range_end_removed']} "
          f"cref={body_result['comment_reference_removed']}")

    if args.dry_run:
        print("[DRY-RUN] 不写文件")
        _cc.write_report(report, args.report, announce="[INFO] report -> {path}")
        return 0

    # 是否要写? 只要 body 有改 或 settings/comments 需要清就写
    need_write_body = sum([
        body_result['ins_accepted'],
        body_result['del_accepted'],
        body_result['comment_range_start_removed'],
        body_result['comment_range_end_removed'],
        body_result['comment_reference_removed'],
    ]) > 0
    need_zip = (before['trackChanges'] > 0) or (
        before['comments_xml_present'] and not args.keep_comments
    )

    if not need_write_body and not need_zip:
        print("[INFO] 无需变更, 不写文件")
        _cc.write_report(report, args.report, announce="[INFO] report -> {path}")
        return 0

    if not args.no_backup:
        bak = _cc.make_backup(args.docx)
        report["backup"] = str(bak)
        print(f"[INFO] 备份 -> {bak.name}")

    # ---- 2) 写 body (即使 body 没变也 save 一次保证一致, 但只在 need_write_body 时 save)
    if need_write_body:
        doc.save(str(args.docx))

    # ---- 3) zip 改 comments.xml + settings.xml
    if need_zip:
        zip_result = apply_path_revisions(args.docx, args)
        report["zip_changes"] = zip_result
        print(f"  [zip] trackchanges_removed={zip_result['trackchanges_removed']} "
              f"comments_emptied={zip_result['comments_emptied']}")

    report["wrote"] = True
    print(f"[OK] 已清理修订残留, 写回 {args.docx.name}")

    # ---- 4) verify (after counts)
    after = scan_counts(args.docx)
    report["after"] = after
    print(f"  [after]  ins={after['ins']} del={after['del']} "
          f"crs={after['comment_range_start']} cre={after['comment_range_end']} "
          f"cref={after['comment_reference']} trackChanges={after['trackChanges']}")

    _cc.write_report(report, args.report, announce="[INFO] report -> {path}")

    return 0


# ══════════ empty-captions ← strip_empty_captions.py ══════════

_DOC_EMPTY_CAPTIONS = """
strip_empty_captions.py
=======================

单功能描述
----------
从 docx 移除**样式为 caption (图名称 / 表格标题 / 图名 / 表名 / Caption 等)
且段落文本完全为空**的段。模板里常残留 "图前 placeholder 空段" 这种结构性
噪音, 影响 audit / renumber 等下游脚本判定。

算法
----
1. ``python-docx`` 打开 docx
2. 遍历 ``doc.paragraphs``, 识别 caption 样式 (style.name 大小写不敏感匹配:
   "图名称" / "图名" / "表格标题" / "表标题" / "表名" / 含 "caption" )
3. 段文本 ``.text.strip() == ""`` **且** 无 ``<w:drawing>`` / ``<w:pict>`` /
   ``<w:object>`` (inline 图/对象) **且** 无 ``<w:r>`` 含非空 ``<w:t>`` →
   strict 空段, 标记删除
4. 经 lxml 真删 (``parent.remove(p_elem)``), 不是设隐藏
5. ``doc.save``

谨慎边界
--------
- **不删 inline drawing**: 段含 ``<w:drawing>`` (浮动图也含) 一律保留
- **不删非 strict 空**: 段含 inline image / 含非空 run / 含表格引用都保留
- **不动正文段**: 仅处理 caption 样式段
- 同一 caption 段紧贴下一个非空 caption 段时 (典型 "图占位空段 + 图名" 配对),
  本脚本会删 placeholder 段; 这是预期 (用户已确认 strict 空段为模板残留)

CLI
---
    python3 sub/strip_empty_captions.py <docx_path> \\
        [-o OUT | --inplace] [--dry-run] [--no-backup] [--report <json>]

默认行为
--------
- 默认 ``--inplace`` (留 ``.bak-N-YYYY-MM-DD.docx``)
- ``-o OUT`` 写到新路径, 不动原文件 + 不留 bak
- ``--dry-run`` 列将删的段 (样式 + 段索引 + 上下文 1 行), 不写

不做
----
- 不删非空 caption 段
- 不删含 inline drawing 的 caption 段 (图本身不动)
- 不改 caption 编号 (caption renumber 的活)
- 不动正文 / heading 段
"""

# 触发删除的样式名关键词 (大小写不敏感, substring 匹配)
CAPTION_STYLE_KEYWORDS = (
    "图名称", "图名", "表格标题", "表标题", "表名", "caption",
)


def is_caption_style(style_name: str | None) -> bool:
    if not style_name:
        return False
    low = style_name.lower()
    for kw in CAPTION_STYLE_KEYWORDS:
        if kw.lower() in low:
            return True
    return False


def is_strict_empty_paragraph(p_elem) -> bool:
    """段是 strict 空: 无 <w:t> 含文字 且 无 <w:drawing>/<w:pict>/<w:object>."""
    # 任何 inline 图/对象都视为非空 (内容保留)
    for tag in (f"{W}drawing", f"{W}pict", f"{W}object"):
        if p_elem.find(f".//{tag}") is not None:
            return False
    # 任何 <w:t> 含非空文本视为非空
    for t in p_elem.findall(f".//{W}t"):
        if t.text and t.text.strip():
            return False
    # 任何 <w:tab> / <w:br> / <w:noBreakHyphen> 也不算文本 (这些段就是结构噪音);
    # 但 <w:sym> / <w:fldChar> 含字段 (如 SEQ 编号) 不算文本但可能是有意义的占位 —
    # 保守起见, 段含 <w:fldChar> (字段) 视为非空, 保留
    if p_elem.find(f".//{W}fldChar") is not None:
        return False
    if p_elem.find(f".//{W}sym") is not None:
        return False
    return True


def scan_empty_captions(doc) -> list[dict]:
    """扫: 返回每个待删段的 {idx, style, prev_text, next_text}."""
    hits: list[dict] = []
    paragraphs = list(doc.paragraphs)
    for i, p in enumerate(paragraphs):
        try:
            style_name = p.style.name if p.style is not None else None
        except (AttributeError, KeyError):
            style_name = None
        if not is_caption_style(style_name):
            continue
        if not is_strict_empty_paragraph(p._element):
            continue
        prev_text = paragraphs[i - 1].text.strip()[:40] if i > 0 else ""
        next_text = paragraphs[i + 1].text.strip()[:40] if i + 1 < len(paragraphs) else ""
        hits.append({
            "idx": i,
            "style": style_name,
            "prev_text": prev_text,
            "next_text": next_text,
        })
    return hits


def delete_marked(doc, hits: list[dict]) -> int:
    """按段索引删除 (从后往前以保索引有效)."""
    paragraphs = list(doc.paragraphs)
    idx_set = {h["idx"] for h in hits}
    deleted = 0
    for i in sorted(idx_set, reverse=True):
        if i >= len(paragraphs):
            continue
        p_elem = paragraphs[i]._element
        parent = p_elem.getparent()
        if parent is None:
            continue
        parent.remove(p_elem)
        deleted += 1
    return deleted


# ---------------- pipeline adapter ----------------
def apply_empty_captions(doc, args=None) -> dict:
    """在内存 doc 上删掉 strict 空 caption 段, 返回改动报告.

    识别规则是启发式(样式关键词 + strict 空判定), 刻意不做成参数 ——
    它靠的是模板残留的客观形态, 一旦开放配置就会被调成"顺手多删几段"。
    所以本函数不读 args, 但仍保留该形参以符合 pipeline 的 step 契约。
    """
    hits = scan_empty_captions(doc)
    deleted = delete_marked(doc, hits)
    return {"changed": deleted, "deleted": deleted, "hits": hits}


def main_empty_captions() -> int:
    parser = argparse.ArgumentParser(
        description=_DOC_EMPTY_CAPTIONS,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("docx", type=Path)
    mx = parser.add_mutually_exclusive_group()
    mx.add_argument("-o", "--output", type=Path, default=None,
                    help="写到新路径(不动原文件,不留 bak)")
    mx.add_argument("--inplace", action="store_true", default=True,
                    help="原地改写(默认),自动留 .bak-N-YYYY-MM-DD")
    _cc.add_write_flags(parser)
    args = parser.parse_args()

    if not args.docx.exists():
        print(f"[ERR] 文件不存在: {args.docx}", file=sys.stderr)
        return 2

    inplace = args.output is None
    if not args.dry_run and inplace:
        occ = _cc.lsof_check(args.docx)
        if occ:
            print(f"[ERR] 文件被占用 (Word/WPS 在开?), 立即停止:\n{occ}", file=sys.stderr)
            return 3

    print(f"[INFO] 扫描 {args.docx.name}")
    doc = Document(str(args.docx))
    hits = scan_empty_captions(doc)
    print(f"  [scan] empty_caption_paragraphs={len(hits)}")

    report = {
        "docx": str(args.docx.resolve()),
        "dry_run": args.dry_run,
        "inplace": inplace,
        "output": str(args.output.resolve()) if args.output else None,
        "backup": None,
        "wrote": False,
        "hits": hits,
        "deleted": 0,
    }

    if not hits:
        print("[INFO] 无空 caption 段, 不写")
        _cc.write_report(report, args.report, announce="[INFO] report -> {path}")
        return 0

    if args.dry_run:
        print("[DRY-RUN] 将删除以下空 caption 段:")
        for h in hits[:50]:
            print(f"  - idx={h['idx']:>4}  style={h['style']!r:<20}  "
                  f"prev='{h['prev_text']}'  next='{h['next_text']}'")
        if len(hits) > 50:
            print(f"  ... 共 {len(hits)} 段 (仅显示前 50)")
        _cc.write_report(report, args.report, announce="[INFO] report -> {path}")
        return 0

    if inplace and not args.no_backup:
        bak = _cc.make_backup(args.docx)
        report["backup"] = str(bak)
        print(f"[INFO] 备份 -> {bak.name}")

    # dry-run / 无命中 两条分支已在上面用 scan 结果提前返回, 走到这里必然是真删
    deleted = apply_empty_captions(doc, args)["deleted"]
    report["deleted"] = deleted

    out_path = args.output if args.output else args.docx
    doc.save(str(out_path))
    report["wrote"] = True
    print(f"[OK] 删除 {deleted} 个空 caption 段, 写回 {out_path.name}")

    _cc.write_report(report, args.report, announce="[INFO] report -> {path}")

    return 0


# ══════════ orphan-media ← strip_orphan_media.py ══════════

_DOC_ORPHAN_MEDIA = """
strip_orphan_media.py
=====================

单功能描述
----------
从 docx (zip) 移除 ``word/media/*`` 中**未被任何 rId 引用**的孤儿媒体文件
(常见为模板 placeholder 残留的 ``.wmf`` / ``.emf`` / ``.png`` 等)。

算法
----
1. ``zipfile`` 打开 docx (只读)
2. 收集所有关系文件 (``word/_rels/document.xml.rels`` + 所有
   ``word/_rels/header*.xml.rels`` / ``footer*.xml.rels`` /
   ``footnotes.xml.rels`` / ``endnotes.xml.rels`` / ``comments.xml.rels``)
3. 解析每个 rels, 收集 ``Target`` 指向 ``media/*`` 的项 (相对路径转 ``word/media/<name>``)
4. 列 ``word/media/`` 下所有 entry, 差集 = orphan
5. 重打包: 写新 zip 跳过 orphan 文件

触发场景
--------
- 模板继承时图替换不彻底, 留下旧 ``.wmf`` placeholder 占体积
- ``docx_cli.py audit images`` 报 orphan media >0 时清场
- 减小 docx 体积 (有时 50%+)

CLI
---
    python3 sub/strip_orphan_media.py <docx_path> \\
        [-o OUT | --inplace] [--dry-run] [--no-backup] [--report <json>]

默认行为
--------
- 默认 ``--inplace`` (留 ``.bak-N-YYYY-MM-DD.docx``)
- ``-o OUT`` 写到新路径, **不**改原文件 + **不**留 bak
- ``--dry-run`` 列将删的文件名 + 释放字节数, 不写
- 写前 ``lsof`` 自检 Word/WPS 占用

不做
----
- 不动 ``word/media/`` 之外的资源 (embeddings/charts 等)
- 不试图"修复"挂的 rId (那是 image relink 的活)
- 仅删未被引用的 media; 被引用的一律保留
"""

NS_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
REL = f"{{{NS_REL}}}"


# ---------------- 核心扫描 ----------------

# rels 文件可能在 word/_rels/ 下任意 *.xml.rels
_RELS_RE = re.compile(r"^word/_rels/.+\.xml\.rels$")
_MEDIA_RE = re.compile(r"^word/media/.+$")

# Office namespaces used to detect rId usage in body / header / footer xml
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_R_ATTRS = (f"{{{NS_R}}}embed", f"{{{NS_R}}}link", f"{{{NS_R}}}id")


def _collect_referenced_media(z: zipfile.ZipFile) -> set[str]:
    """收集所有 rels 文件里 Target 指向 media/* 的项,返回 zip 内规范化路径集合 (word/media/<name>)."""
    referenced: set[str] = set()
    parser = etree.XMLParser(remove_blank_text=False, recover=True)
    for name in z.namelist():
        if not _RELS_RE.match(name):
            continue
        try:
            data = z.read(name)
            root = etree.fromstring(data, parser=parser)
        except (etree.XMLSyntaxError, KeyError):
            continue
        if root is None:
            continue
        for rel in root.findall(f"{REL}Relationship"):
            target = rel.get("Target") or ""
            # Target 一般是相对 word/ 的: "media/image1.wmf" 或 "../media/x" 罕见
            # 也可能是绝对 "/word/media/x" (少见)
            t = target.replace("\\", "/").lstrip("/")
            if t.startswith("media/"):
                referenced.add("word/" + t)
            elif "/media/" in t:
                # 罕见情况, 提取 media/ 后缀
                idx = t.find("media/")
                referenced.add("word/" + t[idx:])
    return referenced


def _scan_used_rids_in_xml(z: zipfile.ZipFile, xml_name: str) -> set[str]:
    """Scan an OOXML part for actually-used rIds (r:embed / r:link / r:id attrs)."""
    used: set[str] = set()
    try:
        data = z.read(xml_name)
    except KeyError:
        return used
    parser = etree.XMLParser(remove_blank_text=False, recover=True)
    try:
        root = etree.fromstring(data, parser=parser)
    except etree.XMLSyntaxError:
        return used
    if root is None:
        return used
    for elem in root.iter():
        for attr in NS_R_ATTRS:
            rid = elem.get(attr)
            if rid:
                used.add(rid)
    return used


def _collect_used_media_via_body(z: zipfile.ZipFile) -> set[str]:
    """Deep scan: only count media actually referenced by an rId used in
    document.xml / header*.xml / footer*.xml / footnotes.xml / endnotes.xml / comments.xml.

    Returns set of zip-paths like "word/media/imageN.png" that are TRULY used.
    """
    # 1. Map each rels file to its used-rIds (from the corresponding XML part)
    body_xmls = [
        "word/document.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
        "word/comments.xml",
    ]
    for n in z.namelist():
        if re.match(r"^word/(header|footer)\d*\.xml$", n):
            body_xmls.append(n)

    used_per_part: dict[str, set[str]] = {}
    for xn in body_xmls:
        used_per_part[xn] = _scan_used_rids_in_xml(z, xn)

    # 2. For each rels file, intersect its Relationship rIds with the used-rIds
    #    of the corresponding XML part, then collect Target → zip-path.
    parser = etree.XMLParser(remove_blank_text=False, recover=True)
    truly_referenced: set[str] = set()

    def _rels_for(part_xml: str) -> str:
        # word/document.xml -> word/_rels/document.xml.rels
        # word/header1.xml -> word/_rels/header1.xml.rels
        d, fn = part_xml.rsplit("/", 1)
        return f"{d}/_rels/{fn}.rels"

    for part_xml, used_rids in used_per_part.items():
        if not used_rids:
            continue
        rels_name = _rels_for(part_xml)
        try:
            data = z.read(rels_name)
            root = etree.fromstring(data, parser=parser)
        except (KeyError, etree.XMLSyntaxError):
            continue
        if root is None:
            continue
        for rel in root.findall(f"{REL}Relationship"):
            rid = rel.get("Id")
            if rid not in used_rids:
                continue
            target = (rel.get("Target") or "").replace("\\", "/").lstrip("/")
            if target.startswith("media/"):
                truly_referenced.add("word/" + target)
            elif "/media/" in target:
                idx = target.find("media/")
                truly_referenced.add("word/" + target[idx:])
    return truly_referenced


def scan_orphans(docx_path: Path, deep: bool = False) -> dict:
    """只读扫: 返回 {referenced, media_in_zip, orphans, orphan_bytes}.

    deep=False (默认): 只看 rels 是否仍 list 此 media (相容旧行为).
    deep=True: 还看 document.xml/header*/footer*/footnotes/endnotes/comments 里
               有没有真的用到此 rId. 用于 split / table-extract 这类
               "body 已裁但 rels 未裁" 的场景, 此时多数 media 在 rels 里
               还在但已无 body 引用 -> 标 orphan.
    """
    with zipfile.ZipFile(str(docx_path), "r") as z:
        if deep:
            referenced = _collect_used_media_via_body(z)
        else:
            referenced = _collect_referenced_media(z)
        media_in_zip: dict[str, int] = {}  # name -> compressed size (近似释放空间)
        for info in z.infolist():
            if _MEDIA_RE.match(info.filename):
                media_in_zip[info.filename] = info.compress_size
        orphans = sorted(set(media_in_zip) - referenced)
        orphan_bytes = sum(media_in_zip[n] for n in orphans)
    return {
        "referenced_count": len(referenced),
        "media_in_zip_count": len(media_in_zip),
        "orphans": orphans,
        "orphan_count": len(orphans),
        "orphan_compressed_bytes": orphan_bytes,
        "deep": deep,
    }


def _rewrite_rels_drop_orphans(rels_xml: bytes, orphan_targets: set[str]) -> bytes:
    """Drop <Relationship> entries whose Target points to a removed media file.

    orphan_targets: zip-paths like "word/media/imageN.png" that have been deleted.
    Returns rewritten rels XML bytes; if nothing changed, returns input unchanged.
    """
    parser = etree.XMLParser(remove_blank_text=False, recover=True)
    try:
        root = etree.fromstring(rels_xml, parser=parser)
    except etree.XMLSyntaxError:
        return rels_xml
    if root is None:
        return rels_xml
    changed = False
    for rel in list(root.findall(f"{REL}Relationship")):
        target = (rel.get("Target") or "").replace("\\", "/").lstrip("/")
        if target.startswith("media/"):
            zp = "word/" + target
        elif "/media/" in target:
            idx = target.find("media/")
            zp = "word/" + target[idx:]
        else:
            continue
        if zp in orphan_targets:
            root.remove(rel)
            changed = True
    if not changed:
        return rels_xml
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def rewrite_skip(docx_path: Path, out_path: Path, orphans: set[str]) -> int:
    """重打包: 跳过 orphans 集合中的文件 + 同步清理 rels 里指向它们的 Relationship.
    返回实际跳过数 (媒体文件)."""
    skipped = 0
    tmp = out_path.with_suffix(out_path.suffix + ".tmp")
    with zipfile.ZipFile(str(docx_path), "r") as zin:
        with zipfile.ZipFile(str(tmp), "w", zipfile.ZIP_DEFLATED) as zout:
            for it in zin.infolist():
                if it.filename in orphans:
                    skipped += 1
                    continue
                data = zin.read(it.filename)
                # If it's a rels file, drop any Relationship pointing to a deleted media
                if _RELS_RE.match(it.filename):
                    data = _rewrite_rels_drop_orphans(data, orphans)
                zout.writestr(it, data)
    # B 类完整性对账（move 之前，docx_path 未动 = 天然基线，断言炸则源件/产物均无损）：
    # ① 丢的 == 判定的 orphan 集（一个不多不少）② 无未报备新增 ③ 字节变化只落在 rels 上
    d = diff_parts(docx_path, tmp, allow_changed=frozenset())
    problems = []
    if set(d.lost) != set(orphans):
        problems.append(
            f"lost != 声明删除集: 多删 {sorted(set(d.lost) - set(orphans))} "
            f"/ 漏删 {sorted(set(orphans) - set(d.lost))}")
    if d.added:
        problems.append(f"未报备新增部件: {d.added}")
    bad_changed = [n for n in d.changed if not _RELS_RE.match(n)]
    if bad_changed:
        problems.append(f"rels 之外的部件被改写: {bad_changed}")
    if problems:
        tmp.unlink(missing_ok=True)
        raise PartIntegrityError(
            "strip_orphan_media 完整性校验未通过: " + "; ".join(problems))
    shutil.move(str(tmp), str(out_path))
    return skipped


# ---------------- CLI ----------------

def main_orphan_media() -> int:
    parser = argparse.ArgumentParser(
        description=_DOC_ORPHAN_MEDIA,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("docx", type=Path)
    mx = parser.add_mutually_exclusive_group()
    mx.add_argument("-o", "--output", type=Path, default=None,
                    help="写到新路径(不动原文件,不留 bak)")
    mx.add_argument("--inplace", action="store_true", default=True,
                    help="原地改写(默认),自动留 .bak-N-YYYY-MM-DD")
    _cc.add_write_flags(parser, no_backup_help="inplace 模式下不留 .bak")
    parser.add_argument(
        "--deep", action="store_true",
        help="深扫: 只保留 document.xml/header*/footer*/footnotes/endnotes/"
             "comments 真正引用的 rId 对应 media; 其余即使 rels 还在 list 也算 orphan "
             "(split / table-extract 等 body 裁但 rels 未裁的场景必须开)",
    )
    args = parser.parse_args()

    if not args.docx.exists():
        print(f"[ERR] 文件不存在: {args.docx}", file=sys.stderr)
        return 2

    inplace = args.output is None
    if not args.dry_run and inplace:
        occ = _cc.lsof_check(args.docx)
        if occ:
            print(f"[ERR] 文件被占用 (Word/WPS 在开?), 立即停止:\n{occ}", file=sys.stderr)
            return 3

    print(f"[INFO] 扫描 {args.docx.name}{' (deep)' if args.deep else ''}")
    scan = scan_orphans(args.docx, deep=args.deep)
    print(f"  [scan] referenced={scan['referenced_count']} "
          f"media_in_zip={scan['media_in_zip_count']} "
          f"orphans={scan['orphan_count']} "
          f"orphan_bytes(compressed)={scan['orphan_compressed_bytes']}")

    report = {
        "docx": str(args.docx.resolve()),
        "dry_run": args.dry_run,
        "inplace": inplace,
        "output": str(args.output.resolve()) if args.output else None,
        "backup": None,
        "wrote": False,
        "scan": scan,
        "skipped": 0,
        "size_before": args.docx.stat().st_size,
        "size_after": None,
    }

    if scan["orphan_count"] == 0:
        print("[INFO] 无 orphan media, 不写")
        _cc.write_report(report, args.report, announce="[INFO] report -> {path}")
        return 0

    # dry-run: 列名
    if args.dry_run:
        print("[DRY-RUN] 将删除以下 orphan media:")
        for n in scan["orphans"][:50]:
            print(f"  - {n}")
        if scan["orphan_count"] > 50:
            print(f"  ... 共 {scan['orphan_count']} 个 (仅显示前 50)")
        print(f"[DRY-RUN] 预计释放 (compressed) {scan['orphan_compressed_bytes']} bytes")
        _cc.write_report(report, args.report, announce="[INFO] report -> {path}")
        return 0

    # 真写
    out_path = args.output if args.output else args.docx
    if inplace and not args.no_backup:
        bak = _cc.make_backup(args.docx)
        report["backup"] = str(bak)
        print(f"[INFO] 备份 -> {bak.name}")

    skipped = rewrite_skip(args.docx, out_path, set(scan["orphans"]))
    report["skipped"] = skipped
    report["size_after"] = out_path.stat().st_size
    report["wrote"] = True

    delta = report["size_before"] - report["size_after"]
    print(f"[OK] 删除 {skipped} 个 orphan media, "
          f"体积 {report['size_before']} -> {report['size_after']} "
          f"(减 {delta} bytes / {delta/1024:.1f} KB)")

    _cc.write_report(report, args.report, announce="[INFO] report -> {path}")

    return 0


# ══════════ outlinelvl ← strip_outlinelvl_from_captions.py ══════════

_DOC_OUTLINELVL = """
strip_outlinelvl_from_captions.py
==================================

单功能描述
----------
从 docx 中匹配 `^(表|图)\\s*\\d+\\.\\d+-\\d+` 的段移除 ``<w:outlineLvl>`` 子元素,
让这些 caption 段不再出现在 Word "视图 > 导航窗格"的标题大纲里。

触发场景
--------
- docx 整合后表名/图名段被错误地写入了 outlineLvl=6,导致 Word 导航大纲污染。
- 仅处理"表 X.Y-Z" / "图 X.Y-Z" 三段式编号的 caption 段。
- 不动 H1-H4 章节标题段(它们的 outlineLvl 由 Heading 1-4 样式合法控制)。

CLI
---
    python3 scripts/strip_outlinelvl_from_captions.py <docx_path> \\
        [--dry-run] [--no-backup] [--report <json>]

启发规则
--------
- 命中: ``re.match(r'^(表|图)\\s*\\d+\\.\\d+-\\d+', paragraph.text.strip())``
- 移除: ``paragraph._p / w:pPr / w:outlineLvl`` 子元素(用 lxml ``pPr.remove(ol)``)
- 段已无 outlineLvl → skip (no_outlinelvl_skip++)

不许做
------
- 不改段文本(编号 ``表X.Y-Z`` 用户已锁定保留)
- 不改段 style.name (交给 apply_caption_styles.py)
- 不裸解 zip 改 XML (走 python-docx + lxml)
- 不动 H1-H4 / Normal 正文段 / 19 张图

约束
----
- 单文件实现, 仅依赖 python-docx
- 自动备份 .bak-N-YYYY-MM-DD.docx (除非 --no-backup)
- 写前 lsof 自检 Word/WPS 占用, 占用立即退出
"""

CAPTION_PATTERN = re.compile(r'^(表|图)\s*\d+\.\d+-\d+')


def scan_and_strip(doc: Document, apply: bool) -> dict:
    """扫所有段, 命中 caption pattern 的从 pPr 移除 outlineLvl"""
    captions_processed = 0
    outlinelvl_removed = 0
    no_outlinelvl_skip = 0
    details = []

    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not CAPTION_PATTERN.match(text):
            continue
        captions_processed += 1
        pPr = p._p.find(qn('w:pPr'))
        ol = None if pPr is None else pPr.find(qn('w:outlineLvl'))
        if ol is None:
            no_outlinelvl_skip += 1
            details.append({
                "idx": idx,
                "text": text[:60],
                "outlineLvl_before": None,
                "action": "skip",
            })
            continue
        ol_val = ol.get(qn('w:val'))
        details.append({
            "idx": idx,
            "text": text[:60],
            "outlineLvl_before": ol_val,
            "action": "remove",
        })
        if apply:
            pPr.remove(ol)
        outlinelvl_removed += 1

    return {
        "captions_processed": captions_processed,
        "outlinelvl_removed": outlinelvl_removed,
        "no_outlinelvl_skip": no_outlinelvl_skip,
        "details": details,
    }


def main_outlinelvl() -> int:
    parser = argparse.ArgumentParser(description=_DOC_OUTLINELVL, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("docx", type=Path)
    _cc.add_write_flags(parser)
    args = parser.parse_args()

    if not args.docx.exists():
        print(f"[ERR] 文件不存在: {args.docx}", file=sys.stderr)
        return 2

    if not args.dry_run:
        occ = _cc.lsof_check(args.docx)
        if occ:
            print(f"[ERR] 文件被占用 (Word/WPS 在开?), 立即停止:\n{occ}", file=sys.stderr)
            return 3

    doc = Document(str(args.docx))
    result = scan_and_strip(doc, apply=not args.dry_run)

    report = {
        "docx": str(args.docx.resolve()),
        "dry_run": args.dry_run,
        "backup": None,
        "wrote": False,
        **result,
    }

    print(f"[INFO] 扫描 {args.docx.name}")
    print(f"  captions_processed   = {result['captions_processed']}")
    print(f"  outlinelvl_removed   = {result['outlinelvl_removed']}")
    print(f"  no_outlinelvl_skip   = {result['no_outlinelvl_skip']}")

    # 列前几条
    for d in result["details"][:5]:
        print(f"  idx={d['idx']:4d} | ol={d['outlineLvl_before']!s:4s} | {d['action']:6s} | {d['text']}")

    if args.dry_run:
        print("[DRY-RUN] 不写文件")
    elif result["outlinelvl_removed"] == 0:
        print("[INFO] 无需移除, 不写文件")
    else:
        if not args.no_backup:
            bak = _cc.make_backup(args.docx)
            report["backup"] = str(bak)
            print(f"[INFO] 备份 -> {bak.name}")
        doc.save(str(args.docx))
        report["wrote"] = True
        print(f"[OK] 已移除 {result['outlinelvl_removed']} 个 outlineLvl, 写回 {args.docx.name}")

    _cc.write_report(report, args.report, announce="[INFO] report -> {path}")

    return 0


# ---------------- pipeline adapter ----------------
def apply_outlinelvl_from_captions(doc, args=None) -> dict:
    dry = bool(getattr(args, "dry_run", False)) if args else False
    return scan_and_strip(doc, apply=not dry)


# ══════════ style-outlinelvl ← strip_style_outlinelvl.py ══════════

_DOC_STYLE_OUTLINELVL = """strip_style_outlinelvl.py — 从 word/styles.xml 移除 caption 族样式自身的 <w:outlineLvl>

## 单功能
docx 的 `word/styles.xml` 里,样式定义自带 <w:pPr><w:outlineLvl/> 时,
所有引用该样式的段落都会被 Word 视为进入大纲层级(显示在导航窗格),
即使段级 <w:outlineLvl> 已被清理也无效——段级会继承样式级。
本脚本仅删除指定 caption 族样式(默认 ZDWP 表名 / ZDWP 图名 / Caption / 0 表名 /
0 图名称 / 0图)上的 <w:outlineLvl> 元素,使图/表 caption 段彻底退出 Word 导航大纲。

## 触发场景
- 跑 strip_outlinelvl_from_captions.py(段级清理)后 Word 导航仍显示 caption
- audit_caption_outline.py 报 caption 段被视为大纲层级
- styles.xml 自查发现 caption 族样式带 outlineLvl

## 严禁动什么
- Heading 1-9 / Title / TOC Heading / 任何 "标题N" / "1.1.1.x N级标题" 等真章节样式
- docx 段级 outlineLvl(那是 strip_outlinelvl_from_captions.py 的职责)
- 样式的 type / rPr / 其他 pPr 子元素
"""

DEFAULT_CAPTION_STYLES = {
    "ZDWP 表名",
    "ZDWP 图名",
    "Caption",
    "0 表名",
    "0 图名称",
    "0图",
}

# 永远不许动的"真标题"样式(白名单兜底,即使用户误把它们传进 --styles 也拒绝)
PROTECTED_STYLE_PREFIXES = (
    "heading ",   # heading 1..9
    "Heading ",
)
PROTECTED_STYLE_NAMES = {
    "Title",
    "TOC Heading",
    "标题1", "标题2", "标题3", "标题4", "标题5", "标题6",
    "1.1.1.1四级标题",
    "--1.1.1三级标题",
}


def is_protected(name: str) -> bool:
    if name in PROTECTED_STYLE_NAMES:
        return True
    for p in PROTECTED_STYLE_PREFIXES:
        if name.startswith(p):
            return True
    # 形如 "标题N" / "X级标题" 兜底
    if "标题" in name and any(c.isdigit() for c in name):
        return True
    return False


def find_outlinelvl_styles(root) -> list[tuple[str, str, int]]:
    """返回所有带 outlineLvl 的样式列表 [(name, styleId, lvl), ...]"""
    out = []
    for s in root.findall(f".//{W}style"):
        n = s.find(f"{W}name")
        nm = n.get(W + "val") if n is not None else "?"
        sid = s.get(W + "styleId")
        ol = s.find(f"{W}pPr/{W}outlineLvl")
        if ol is not None:
            out.append((nm, sid, int(ol.get(W + "val"))))
    return out


def patch_styles_xml(docx_path: Path, target_names: set[str], dry_run: bool):
    with zipfile.ZipFile(docx_path, "r") as z:
        styles_bytes = z.read("word/styles.xml")

    root = etree.fromstring(styles_bytes)
    removed = []
    skipped_protected = []

    for s in root.findall(f".//{W}style"):
        n = s.find(f"{W}name")
        nm = n.get(W + "val") if n is not None else None
        if nm is None or nm not in target_names:
            continue
        if is_protected(nm):
            skipped_protected.append(nm)
            continue
        pPr = s.find(f"{W}pPr")
        if pPr is None:
            continue
        ol = pPr.find(f"{W}outlineLvl")
        if ol is None:
            continue
        lvl = ol.get(W + "val")
        sid = s.get(W + "styleId")
        if not dry_run:
            pPr.remove(ol)
        removed.append({"name": nm, "styleId": sid, "outlineLvl": lvl})

    if dry_run:
        return removed, skipped_protected, None

    new_bytes = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    tmp = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for it in zin.infolist():
            if it.filename == "word/styles.xml":
                zout.writestr(it, new_bytes)
            else:
                zout.writestr(it, zin.read(it.filename))
    # 部件完整性断言(move 前:docx_path 仍是未动源件=天然基线;炸则 tmp 被清、源件无损)。
    # styles.xml 不在 DEFAULT 白名单,显式报备。main 与 pipeline apply_path 都汇到本函数。
    try:
        assert_parts_intact(docx_path, tmp,
                            allow_changed=set(DEFAULT_ALLOW_CHANGED) | {"word/styles.xml"},
                            verbose=False)
    except Exception:
        tmp.unlink(missing_ok=True)
        raise
    shutil.move(str(tmp), str(docx_path))
    return removed, skipped_protected, new_bytes


def main_style_outlinelvl():
    ap = argparse.ArgumentParser(description=_DOC_STYLE_OUTLINELVL.split("\n\n")[0])
    ap.add_argument("docx", help="目标 docx")
    ap.add_argument(
        "--styles",
        help=f"逗号分隔的样式名集合,默认: {','.join(sorted(DEFAULT_CAPTION_STYLES))}",
    )
    _cc.add_write_flags(
        ap,
        dry_run_help="只统计不写",
        no_backup_help="跳过 .bak 备份",
        report_help="写 JSON 报告路径",
    )
    args = ap.parse_args()

    docx_path = Path(args.docx).resolve()
    if not docx_path.exists():
        print(f"[abort] 不存在: {docx_path}", file=sys.stderr)
        sys.exit(2)

    if not args.dry_run:
        occ = _cc.lsof_check(docx_path)
        if occ:
            print(f"[abort] {docx_path} 被打开(Word/WPS),请关闭后重试:\n{occ}",
                  file=sys.stderr)
            sys.exit(2)

    if args.styles:
        target = {s.strip() for s in args.styles.split(",") if s.strip()}
    else:
        target = set(DEFAULT_CAPTION_STYLES)

    # 诊断:全列 outlineLvl 样式
    with zipfile.ZipFile(docx_path, "r") as z:
        sb = z.read("word/styles.xml")
    root = etree.fromstring(sb)
    pre_all = find_outlinelvl_styles(root)
    print(f"[diag] styles.xml 共 {len(pre_all)} 个样式带 outlineLvl(before)")
    pre_in_target = [t for t in pre_all if t[0] in target and not is_protected(t[0])]
    print(f"[diag] 命中 target 且非保护样式: {len(pre_in_target)}")
    for nm, sid, lvl in pre_in_target:
        print(f"  - {nm} (id={sid}, lvl={lvl})")

    # 备份
    backup_path = None
    if not args.dry_run and not args.no_backup:
        backup_path = _cc.make_backup(docx_path)
        print(f"[backup] {backup_path.name}")

    removed, skipped_protected, _ = patch_styles_xml(
        docx_path, target, dry_run=args.dry_run
    )

    print(f"[result] removed={len(removed)}  dry_run={args.dry_run}")
    for r in removed:
        print(f"  - {r['name']} (id={r['styleId']}, was lvl={r['outlineLvl']})")
    if skipped_protected:
        print(f"[guard] 拒绝处理保护样式: {skipped_protected}")

    if args.report:
        rep = {
            "docx": str(docx_path),
            "dry_run": args.dry_run,
            "backup": str(backup_path) if backup_path else None,
            "target_styles": sorted(target),
            "pre_outlinelvl_styles": [
                {"name": n, "styleId": s, "outlineLvl": l} for n, s, l in pre_all
            ],
            "removed": removed,
            "skipped_protected": skipped_protected,
        }
        _cc.write_report(rep, args.report, mkdir=False, announce="[report] {path}")


# ---------------- pipeline adapter ----------------
def apply_path_style_outlinelvl(docx_path, args=None) -> dict:
    styles = getattr(args, "strip_styles", None) if args else None
    if styles:
        target = {s.strip() for s in styles.split(",") if s.strip()}
    else:
        target = set(DEFAULT_CAPTION_STYLES)
    dry = bool(getattr(args, "dry_run", False)) if args else False
    removed, skipped_protected, _ = patch_styles_xml(
        Path(docx_path), target, dry_run=dry
    )
    return {
        "changed": len(removed),
        "removed": removed,
        "skipped_protected": skipped_protected,
    }


# ══════════ doc-protection ← strip_doc_protection.py ══════════

# 默认清除的保护类元素 local-name
PROTECTION_ELEMENTS = [
    "documentProtection",
    "writeProtection",
    "readModeInkLockDown",
    "formsDesign",
    "enforcement",
    "trackChanges",  # 兜底, 与 strip_revisions.py 重叠
]

SETTINGS_PATH = "word/settings.xml"


def _rewrite_settings(p: Path, new_settings_bytes: bytes) -> None:
    """重打包 zip，仅替换 settings.xml —— CLI main 与 pipeline apply_path 共汇的唯一写盘点。

    replace 之前挂部件完整性断言（此刻 p 未动 = 天然基线，断言炸则源件无损）；
    settings.xml 不在 DEFAULT_ALLOW_CHANGED，必须显式报备。
    """
    tmp = p.with_suffix(p.suffix + ".tmp")
    with zipfile.ZipFile(p, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == SETTINGS_PATH:
                data = new_settings_bytes
            zout.writestr(item, data)
    try:
        assert_parts_intact(p, tmp,
                            allow_changed=set(DEFAULT_ALLOW_CHANGED) | {SETTINGS_PATH},
                            verbose=False)
    except Exception:
        tmp.unlink(missing_ok=True)
        raise
    tmp.replace(p)


def _strip(root, dry_run: bool):
    """Remove protection elements from settings root. Return Counter of removed by tag."""
    removed = Counter()
    for tag in PROTECTION_ELEMENTS:
        # findall on settings root: protection elements are direct children of <w:settings>
        # but tolerate descendant for safety
        for el in list(root.iter(f"{W}{tag}")):
            parent = el.getparent()
            if parent is None:
                continue
            removed[tag] += 1
            if not dry_run:
                parent.remove(el)
    return removed


def _audit(root) -> Counter:
    """Pre/post audit: count protection elements in settings root."""
    c = Counter()
    for tag in PROTECTION_ELEMENTS:
        c[tag] = len(list(root.iter(f"{W}{tag}")))
    return c


def main_doc_protection():
    ap = argparse.ArgumentParser(description="Strip document protection / edit limits / "
                                              "viewProtection from docx settings.xml.")
    ap.add_argument("docx", help="Path to .docx")
    _cc.add_write_flags(
        ap,
        dry_run_help="Report only, do not modify file",
        no_backup_help="Skip writing .bak-N-<date>.docx",
        report_help="Write JSON report to this path",
    )
    args = ap.parse_args()

    p = Path(args.docx).resolve()
    if not p.exists():
        print(f"[ERR] 文件不存在: {p}", file=sys.stderr)
        sys.exit(1)

    if not args.dry_run:
        occ = _cc.lsof_check(p)
        if occ:
            print(f"[ERR] 文件被占用 (Word/WPS 没关?):\n{occ}", file=sys.stderr)
            sys.exit(2)

    # Load settings.xml
    with zipfile.ZipFile(p, "r") as z:
        names = z.namelist()
        if SETTINGS_PATH not in names:
            report = {
                "docx": str(p),
                "dry_run": args.dry_run,
                "note": f"no {SETTINGS_PATH}",
                "removed": {},
                "total": 0,
            }
            _cc.print_json(report)
            _cc.write_report(report, args.report, mkdir=False)
            return
        settings_bytes = z.read(SETTINGS_PATH)

    parser = etree.XMLParser(remove_blank_text=False)
    root = etree.fromstring(settings_bytes, parser)

    before = _audit(root)
    removed = _strip(root, dry_run=args.dry_run)
    after = _audit(root) if not args.dry_run else {k: before[k] for k in before}

    report = {
        "docx": str(p),
        "dry_run": args.dry_run,
        "before": {k: v for k, v in before.items() if v > 0},
        "removed": {k: v for k, v in removed.items() if v > 0},
        "total": sum(removed.values()),
        "after": {k: v for k, v in after.items() if v > 0},
    }

    _cc.print_json(report)

    _cc.write_report(report, args.report, mkdir=False)

    if args.dry_run:
        print("[dry-run] no file written")
        return

    if sum(removed.values()) == 0:
        print("[noop] no protection elements found, file unchanged")
        return

    if not args.no_backup:
        bp = _cc.make_backup(p)
        print(f"[backup] {bp.name}")

    new_settings_bytes = etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                                        standalone=True)

    _rewrite_settings(p, new_settings_bytes)
    print(f"[done] wrote {p}")


# ---------------- pipeline adapter ----------------
def apply_path_doc_protection(docx_path, args=None) -> dict:
    """pipeline 入口: 移除全部保护类元素, dry_run from args."""
    dry = bool(getattr(args, "dry_run", False)) if args else False
    p = Path(docx_path).resolve()

    with zipfile.ZipFile(p, "r") as z:
        names = z.namelist()
        if SETTINGS_PATH not in names:
            return {"changed": 0, "removed": {}, "total": 0, "note": f"no {SETTINGS_PATH}"}
        settings_bytes = z.read(SETTINGS_PATH)

    parser = etree.XMLParser(remove_blank_text=False)
    root = etree.fromstring(settings_bytes, parser)

    before = _audit(root)
    removed = _strip(root, dry_run=dry)
    n_changed = sum(removed.values())

    if not dry and n_changed > 0:
        new_settings_bytes = etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                                            standalone=True)
        _rewrite_settings(p, new_settings_bytes)

    after = _audit(root) if not dry else {k: before[k] for k in before}

    return {
        "changed": n_changed,
        "before": {k: v for k, v in before.items() if v > 0},
        "removed": {k: v for k, v in removed.items() if v > 0},
        "after": {k: v for k, v in after.items() if v > 0},
        "total": n_changed,
    }


# ──────────────────────────── 家族入口（子命令分发）────────────────────────────

SUBCOMMANDS = {
    "bookmarks": main_bookmarks,
    "revisions": main_revisions,
    "empty-captions": main_empty_captions,
    "orphan-media": main_orphan_media,
    "outlinelvl": main_outlinelvl,
    "style-outlinelvl": main_style_outlinelvl,
    "doc-protection": main_doc_protection,
}


def main(argv: list[str] | None = None) -> int:
    return _cc.family_main(SUBCOMMANDS, argv, file=__file__,
                           usage_args="<docx> [flags…]")


if __name__ == "__main__":
    sys.exit(main())
