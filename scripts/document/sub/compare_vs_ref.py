#!/usr/bin/env python3
"""对比多份改动 MD 里的"改为"段 vs 参考 docx 全文（主标 or 其他基准），找出雷同风险。

Usage:
  python3 compare_vs_ref.py --drafts-dir 成果/md --ref 主标.docx --out vs-主标-雷同检查.md
"""
import argparse
import re
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document

THRESHOLD_HIGH = 0.85
THRESHOLD_MID = 0.6


def normalize(t):
    t = re.sub(r"\s+", "", t)
    pairs = [("，", ","), ("。", "."), ("；", ";"), ("（", "("), ("）", ")"),
             ("\u201c", '"'), ("\u201d", '"')]
    for a, b in pairs:
        t = t.replace(a, b)
    return t


def extract_ref_paragraphs(path):
    doc = Document(str(path))
    out = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t and len(t) >= 8:
            out.append((f"P{i}", t, normalize(t)))
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                t = cell.text.strip()
                if t and len(t) >= 8:
                    out.append((f"T{ti}.{ri}.{ci}", t, normalize(t)))
    return out


def extract_revisions(drafts_dir: Path, glob_pattern: str):
    out = []
    for f in sorted(drafts_dir.glob(glob_pattern)):
        if "vs" in f.name or "清理" in f.name:
            continue
        text = f.read_text(encoding="utf-8")
        for m in re.finditer(r"\*\*改为[^*]*\*\*：?\s*\n((?:>\s.*\n)+)", text):
            block = m.group(1)
            content_lines = [re.sub(r"^>\s?", "", ln).strip() for ln in block.split("\n") if ln.strip()]
            content = " ".join(content_lines).strip()
            if len(content) >= 8:
                out.append((f.name, m.start(), content, normalize(content)))
    return out


def best_match(target_norm, ref_index, ref_norms_list):
    if target_norm in ref_index:
        return ("EXACT", 1.0, ref_index[target_norm])
    tl = len(target_norm)
    best_r, best_id = 0.0, None
    sm = SequenceMatcher(autojunk=False)
    sm.set_seq2(target_norm)
    for rn, ri in ref_norms_list:
        if abs(len(rn) - tl) > tl * 0.6:
            continue
        sm.set_seq1(rn)
        r = sm.ratio()
        if r > best_r:
            best_r, best_id = r, ri
            if r >= 0.99:
                break
    if best_r >= THRESHOLD_HIGH:
        return ("HIGH", best_r, best_id)
    if best_r >= THRESHOLD_MID:
        return ("MID", best_r, best_id)
    return ("OK", best_r, best_id)


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--drafts-dir", required=True, help="改动草稿 MD 所在目录")
    ap.add_argument("--ref", required=True, help="参考 docx（主标或基准）")
    ap.add_argument("--out", required=True, help="输出 MD 路径")
    ap.add_argument("--glob", default="*改动草稿.md", help="MD 文件 glob 模式（默认 *改动草稿.md）")
    args = ap.parse_args()

    ref_paras = extract_ref_paragraphs(Path(args.ref))
    print(f"参考 {len(ref_paras)} 段")
    ref_index = {n: i for i, _, n in ref_paras}
    ref_text = {i: t for i, t, _ in ref_paras}
    ref_norms_list = [(n, i) for i, _, n in ref_paras]

    revs = extract_revisions(Path(args.drafts_dir), args.glob)
    print(f"改为段 {len(revs)} 条")

    risks = []
    for fname, _, content, norm in revs:
        status, ratio, ref_id = best_match(norm, ref_index, ref_norms_list)
        if status in ("EXACT", "HIGH", "MID"):
            risks.append((fname, content, status, ratio, ref_id, ref_text.get(ref_id, "")))

    L = [f"# 改动后内容 vs 参考 雷同检查\n",
         f"> 参考：`{Path(args.ref).name}`（{len(ref_paras)} 段）",
         f"> 改为段：{len(revs)} 条",
         f"> 阈值：HIGH ratio≥{THRESHOLD_HIGH} → 必须重写；MID {THRESHOLD_MID}≤ratio<{THRESHOLD_HIGH} → 建议调整\n"]
    if not risks:
        L.append("✅ 无雷同风险\n")
    else:
        n_h = sum(1 for r in risks if r[2] in ("EXACT", "HIGH"))
        n_m = sum(1 for r in risks if r[2] == "MID")
        L.append(f"## 总览：{len(risks)} 条风险（高 {n_h} / 中 {n_m}）\n")
        by_file = {}
        for r in risks:
            by_file.setdefault(r[0], []).append(r)
        for fname in sorted(by_file):
            L.append(f"\n## {fname}\n")
            for _, content, status, ratio, ref_id, ref_t in by_file[fname]:
                L.append(f"### [{status} ratio={ratio:.2f}] vs 参考 {ref_id}")
                L.append(f"**改为**：\n> {content[:300]}")
                L.append(f"**参考对应段**：\n> {ref_t[:300]}\n")

    Path(args.out).write_text("\n".join(L), encoding="utf-8")
    print(f"OK -> {args.out}")
    print(f"风险 {len(risks)} 条")


# ---------------- pipeline adapter ----------------
def apply_path(docx_path, args=None) -> dict:
    """pipeline-compatible adapter (跨文件 analyzer).

    docx_path 充当 *主目标* (此 step 它就是 ``--ref`` 默认值);
    可被 args.ref 覆盖。args 透传:
      - drafts_dir (必需): 改动草稿 MD 目录
      - ref: 参考 docx (覆盖 docx_path)
      - glob: MD glob 模式 (默认 *改动草稿.md)
      - out / out_dir: 输出路径 (out 优先; 否则 out_dir/vs-<ref-stem>-雷同检查.md)
    """
    from pathlib import Path as _P
    drafts_dir = getattr(args, "drafts_dir", None) if args else None
    if not drafts_dir:
        return {"skipped": "no --drafts-dir; compare_vs_ref needs MD drafts dir"}
    ref_path = _P(getattr(args, "ref", None) or docx_path)
    glob_pattern = getattr(args, "glob", None) or "*改动草稿.md"
    out_path = getattr(args, "out", None)
    out_dir = getattr(args, "out_dir", None)
    if out_path:
        out = _P(out_path)
    elif out_dir:
        out = _P(out_dir) / f"vs-{ref_path.stem}-雷同检查.md"
    else:
        out = _P(docx_path).parent / "reports" / f"vs-{ref_path.stem}-雷同检查.md"
    out.parent.mkdir(parents=True, exist_ok=True)

    ref_paras = extract_ref_paragraphs(ref_path)
    ref_index = {n: i for i, _, n in ref_paras}
    ref_text = {i: t for i, t, _ in ref_paras}
    ref_norms_list = [(n, i) for i, _, n in ref_paras]
    revs = extract_revisions(_P(drafts_dir), glob_pattern)
    risks = []
    for fname, _, content, norm in revs:
        status, ratio, ref_id = best_match(norm, ref_index, ref_norms_list)
        if status in ("EXACT", "HIGH", "MID"):
            risks.append((fname, content, status, ratio, ref_id, ref_text.get(ref_id, "")))
    L = [f"# 改动后内容 vs 参考 雷同检查\n",
         f"> 参考：`{ref_path.name}`({len(ref_paras)} 段)",
         f"> 改为段：{len(revs)} 条\n"]
    if not risks:
        L.append("✅ 无雷同风险\n")
    else:
        L.append(f"## 总览：{len(risks)} 条风险\n")
        by_file = {}
        for r in risks:
            by_file.setdefault(r[0], []).append(r)
        for fname in sorted(by_file):
            L.append(f"\n## {fname}\n")
            for _, content, status, ratio, ref_id, ref_t in by_file[fname]:
                L.append(f"### [{status} ratio={ratio:.2f}] vs 参考 {ref_id}")
                L.append(f"**改为**：\n> {content[:300]}")
                L.append(f"**参考对应段**：\n> {ref_t[:300]}\n")
    out.write_text("\n".join(L), encoding="utf-8")
    return {
        "ref": str(ref_path),
        "ref_paras": len(ref_paras),
        "revisions": len(revs),
        "risks": len(risks),
        "high": sum(1 for r in risks if r[2] in ("EXACT", "HIGH")),
        "mid": sum(1 for r in risks if r[2] == "MID"),
        "out": str(out),
    }


if __name__ == "__main__":
    main()
