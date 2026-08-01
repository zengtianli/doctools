#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""table.py — 表格家族四合一（2026-07-31 家族折叠）

子命令 ↔ 原脚本（函数体逐字搬移；模块级 apply/apply_path/main/process/_fmt 改名
带 _<sub> 后缀，其余公有名一律保留）：

    delete-rows  ← delete_table_rows.py
    extract      ← extract_tables.py（run_extract / plan_extracts 公有名保留：
                                      pipeline_lib builtin step 靠它）
    center       ← set_table_align.py
    borders      ← set_table_borders.py（--center 链式触发 center，由 docx_cli 层编排）

各子命令 CLI 与原独立脚本逐字一致：python3 sub/table.py <sub> …。
退役原件在 ~/.Trash/consolidation-20260731/table/（含 MANIFEST.md）。
"""
from __future__ import annotations

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[3] / "lib"))
import caption_re  # noqa: E402  题注判据 SSOT
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py
from docx_parts import PartIntegrityError, diff_parts  # noqa: E402  部件完整性(B类:精确部件集)

import argparse  # noqa: E402
import re  # noqa: E402
import shutil  # noqa: E402
import sys  # noqa: E402
import zipfile  # noqa: E402
from datetime import datetime  # noqa: E402
from pathlib import Path  # noqa: E402
from typing import Optional  # noqa: E402

from lxml import etree  # noqa: E402

from docx import Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


# ══════════ delete-rows ← delete_table_rows.py ══════════

def apply_delete_rows(doc, args=None) -> dict:
    """在**已打开**的 doc 上删掉指定表的指定行区间，删前跑安全校验。

    校验失败**不 sys.exit**，而是回一个带 error_kind 的 report，由调用方决定后果
    （CLI 打中文提示后退 1；pipeline 记进报告接着跑别的 step）。这样校验只有一份：
    之前 main() 和 apply_path() 各写了一遍，两边判据已经漂开（关键字校验一处看
    「起始行是否够两列」、另一处不看），改一处必漏一处。

    删行只动 document.xml 的 <w:tr>，不碰别的部件，所以是 doc-based step。
    """
    if args is None:
        return {"skipped": "no args; delete_table_rows needs --table-index/--rows"}
    table_index = getattr(args, "table_index", None)
    rows = getattr(args, "rows", None)
    if table_index is None or not rows:
        return {"skipped": "delete_table_rows needs --table-index and --rows"}

    if table_index >= len(doc.tables):
        return {"error_kind": "table_index_out_of_range",
                "error": f"table-index {table_index} 超出范围 (共 {len(doc.tables)} 表)",
                "tables": len(doc.tables)}

    tbl = doc.tables[table_index]
    from_idx, to_idx = map(int, rows.split(":"))
    before = len(tbl.rows)
    dry_run = bool(getattr(args, "dry_run", False))
    expected_first_col = getattr(args, "expected_first_col", "") or ""
    expected_residue = getattr(args, "expected_residue", "") or ""
    expected_last = getattr(args, "expected_last", "") or ""

    base = {"table_index": table_index, "rows_range": rows,
            "before": before, "dry_run": dry_run}

    # 安全校验 1：保留行的第一列必须逐字对上
    if expected_first_col:
        expected = [x.strip() for x in expected_first_col.split(",")]
        actual = [tbl.rows[i].cells[0].text.strip()
                  for i in range(min(from_idx, before))]
        if actual != expected:
            return {**base, "error_kind": "first_col_mismatch",
                    "error": f"first-col 校验失败: 期望 {expected} 实际 {actual}",
                    "expected": expected, "actual": actual}

    # 安全校验 2：起始行不足两列时整条跳过 —— 合并单元格/小标题行本来就取不到
    # 第二列文本，硬查会把正常情况误判成「删错行」而挡住删除（沿用 main() 判据）。
    if expected_residue and len(tbl.rows[from_idx].cells) >= 2:
        head = tbl.rows[from_idx].cells[0].text + tbl.rows[from_idx].cells[1].text
        if expected_residue not in head:
            found = any(
                expected_residue in " ".join(c.text for c in tbl.rows[i].cells)
                for i in range(from_idx, min(to_idx + 1, before))
            )
            if not found:
                return {**base, "error_kind": "residue_not_found",
                        "error": f"被删行未含关键字「{expected_residue}」"}

    if dry_run:
        planned = to_idx - from_idx + 1
        return {**base, "after": before - planned, "deleted": planned, "changed": 0}

    # 倒序删除：正序删会让后续下标整体前移
    for i in range(to_idx, from_idx - 1, -1):
        if i < len(tbl.rows):
            row = tbl.rows[i]
            row._element.getparent().remove(row._element)

    after = len(tbl.rows)
    result = {**base, "after": after, "deleted": before - after,
              "changed": before - after}
    if expected_last and after >= 1:
        last = tbl.rows[-1]
        last_text = last.cells[1].text.strip() if len(last.cells) >= 2 else ""
        result["last_cell"] = last_text
        result["last_check"] = expected_last in last_text
    return result


def main_delete_rows():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--docx", required=True, help="目标 docx（原地修改）")
    ap.add_argument("--table-index", type=int, required=True, help="表格索引（0-based）")
    ap.add_argument("--rows", required=True, help="要删的行范围 FROM:TO（闭区间，0-based）")
    ap.add_argument("--expected-first-col", default="",
                    help="删除前保留行第一列期望值（逗号分隔），用于安全校验；留空则跳过校验")
    ap.add_argument("--expected-residue", default="",
                    help="要删的第一行第二列应含此关键字，确保删对了")
    ap.add_argument("--expected-last", default="",
                    help="删后末行第二列期望值，用于校验删对了")
    args = ap.parse_args()

    docx_path = Path(args.docx)
    doc = Document(str(docx_path))
    rep = apply_delete_rows(doc, args)
    err = rep.get("error_kind")

    if err == "table_index_out_of_range":
        print(f"❌ table-index 超出范围（文档只有 {rep['tables']} 个表）", file=sys.stderr)
        sys.exit(1)

    from_idx, to_idx = map(int, args.rows.split(":"))
    print(f"删除前 T{args.table_index} 行数: {rep['before']}")
    print(f"计划删除: R{from_idx}-R{to_idx}（共 {to_idx-from_idx+1} 行）")

    # 安全校验 1：保留行的第一列
    if err == "first_col_mismatch":
        print(f"❌ 前 {from_idx} 行第一列与期望不一致："
              f"\n  期望 {rep['expected']}\n  实际 {rep['actual']}", file=sys.stderr)
        sys.exit(1)
    if args.expected_first_col:
        print("✓ 保留行结构校验通过")

    # 安全校验 2：要删的起始行应含关键字
    if err == "residue_not_found":
        print(f"❌ 被删行中未找到关键字「{args.expected_residue}」，可能删错行", file=sys.stderr)
        sys.exit(1)
    if args.expected_residue:
        print(f"✓ 被删行包含期望关键字「{args.expected_residue}」")

    doc.save(str(docx_path))

    # 验证
    doc2 = Document(str(docx_path))
    tbl2 = doc2.tables[args.table_index]
    print(f"删除后行数: {len(tbl2.rows)}")
    if args.expected_last and len(tbl2.rows) >= 1:
        last = tbl2.rows[-1]
        if len(last.cells) >= 2:
            actual_last = last.cells[1].text.strip()
            if args.expected_last and args.expected_last not in actual_last:
                print(f"⚠️ 末行第二列「{actual_last}」不含期望「{args.expected_last}」", file=sys.stderr)
            else:
                print(f"✓ 末行校验通过: {actual_last[:40]}")
    print(f"OK -> {docx_path}")


# ---------------- pipeline adapter ----------------
def apply_path_delete_rows(docx_path, args=None) -> dict:
    """原地 mutator：开文件 → apply() → 存盘。校验/删行逻辑全在 apply() 里。

    args 透传:
      - table_index (必需): 表索引 0-based
      - rows (必需): FROM:TO 闭区间 0-based
      - expected_first_col / expected_residue / expected_last: 安全校验
      - dry_run: 不写盘

    留着是为了不掐断已经按路径调它的老调用方；pipeline_lib.load_step 优先取
    apply()，新链路一律走纯内存版本，这里不会被 pipeline 选中。
    """
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))
    rep = apply_delete_rows(doc, args)
    # 校验没过 / 参数不全 / dry-run —— 一律不落盘
    if rep.get("error_kind") or rep.get("skipped") or rep.get("dry_run"):
        return rep
    doc.save(str(docx_path))
    return rep


# ══════════ extract ← extract_tables.py ══════════

_ILLEGAL_FILENAME_RE = re.compile(r'[/\\:*?"<>|\r\n\t]')
_MULTI_WS_RE = re.compile(r"\s+")

# Caption heuristic: starts with 表 / Table; length < 80 chars
# 2026-08-01 判据下沉 lib/caption_re.KIND_PREFIX_TABLE。**有意保持只看首字、不看编号**
# —— 切表要给无编号题注也能命名，收紧成要求编号会大面积掉回 `table-{idx:02d}` fallback
# （违反 CLAUDE.md §抽取类工具默认契约第 ① 条）。
_CAPTION_PREFIX_RE = caption_re.pattern(caption_re.KIND_PREFIX_TABLE)
_CAPTION_MAX_LEN = 80

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ---------- static skeleton templates ----------

_CONTENT_TYPES_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
  <Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>
</Types>
"""

_CONTENT_TYPES_XML_NO_NUMBERING = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>
</Types>
"""

_CONTENT_TYPES_XML_NO_THEME = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
</Types>
"""

_CONTENT_TYPES_XML_MIN = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>
"""

_ROOT_RELS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""


def _build_doc_rels(has_numbering: bool, has_theme: bool) -> str:
    parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">',
        '  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>',
    ]
    next_id = 2
    if has_numbering:
        parts.append(
            f'  <Relationship Id="rId{next_id}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" '
            'Target="numbering.xml"/>'
        )
        next_id += 1
    if has_theme:
        parts.append(
            f'  <Relationship Id="rId{next_id}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" '
            'Target="theme/theme1.xml"/>'
        )
    parts.append('</Relationships>')
    return "\n".join(parts) + "\n"


def _pick_content_types(has_numbering: bool, has_theme: bool) -> str:
    if has_numbering and has_theme:
        return _CONTENT_TYPES_XML
    if has_theme and not has_numbering:
        return _CONTENT_TYPES_XML_NO_NUMBERING
    if has_numbering and not has_theme:
        return _CONTENT_TYPES_XML_NO_THEME
    return _CONTENT_TYPES_XML_MIN


# ---------- filename helpers ----------

def sanitize_filename(name: str, max_len: int = 100) -> str:
    """Replace illegal filename chars with _, compress whitespace, strip, truncate."""
    if not name:
        return "untitled"
    n = _ILLEGAL_FILENAME_RE.sub("_", name)
    n = _MULTI_WS_RE.sub(" ", n).strip()
    if not n:
        return "untitled"
    if len(n) > max_len:
        n = n[:max_len].rstrip()
    return n


def _paragraph_text(p_elem) -> str:
    """Extract concatenated text from <w:t> nodes under a <w:p>."""
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def _is_caption_like(text: str) -> bool:
    """True if text looks like a table caption (starts with 表/Table, short)."""
    if not text:
        return False
    t = text.strip()
    if not t or len(t) >= _CAPTION_MAX_LEN:
        return False
    return bool(_CAPTION_PREFIX_RE.match(t))


def _find_caption(children: list, tbl_idx: int) -> Optional[str]:
    """Scan back up to 2 paragraphs, else forward 1, for a caption-like <w:p>."""
    seen_paras = 0
    for i in range(tbl_idx - 1, -1, -1):
        elem = children[i]
        if elem.tag != qn("w:p"):
            continue
        text = _paragraph_text(elem).strip()
        if not text:
            continue
        seen_paras += 1
        if _is_caption_like(text):
            return text
        if seen_paras >= 2:
            break

    for i in range(tbl_idx + 1, len(children)):
        elem = children[i]
        if elem.tag == qn("w:tbl"):
            break
        if elem.tag != qn("w:p"):
            continue
        text = _paragraph_text(elem).strip()
        if not text:
            continue
        if _is_caption_like(text):
            return text
        break

    return None


# ---------- planning ----------

def plan_extracts(docx_path: Path, doc=None) -> tuple[list[dict], int]:
    """Inspect docx → return (extracts, tbl_count).

    extracts: list[{idx, caption, stem, tbl_idx_in_body, caption_idx_in_body}]
    """
    if doc is None:
        doc = Document(str(docx_path))
    body = doc.element.body
    children = list(body)

    sect_idx = len(children)
    for i in range(len(children) - 1, -1, -1):
        if children[i].tag == qn("w:sectPr"):
            sect_idx = i
            break

    extracts: list[dict] = []
    idx_counter = 0
    used_stems: dict[str, int] = {}

    for i in range(sect_idx):
        elem = children[i]
        if elem.tag != qn("w:tbl"):
            continue
        caption = _find_caption(children, i)
        if caption:
            stem_raw = caption
        else:
            stem_raw = f"table-{idx_counter:02d}"
        stem = sanitize_filename(stem_raw)
        base_stem = stem
        n = used_stems.get(base_stem, 0)
        if n > 0:
            stem = f"{base_stem}-{n + 1}"
        used_stems[base_stem] = n + 1

        caption_idx = -1
        if caption:
            seen = 0
            for j in range(i - 1, -1, -1):
                e2 = children[j]
                if e2.tag != qn("w:p"):
                    continue
                t2 = _paragraph_text(e2).strip()
                if not t2:
                    continue
                seen += 1
                if t2 == caption.strip():
                    caption_idx = j
                    break
                if seen >= 2:
                    break
            if caption_idx == -1:
                for j in range(i + 1, len(children)):
                    e2 = children[j]
                    if e2.tag == qn("w:tbl"):
                        break
                    if e2.tag != qn("w:p"):
                        continue
                    t2 = _paragraph_text(e2).strip()
                    if not t2:
                        continue
                    if t2 == caption.strip():
                        caption_idx = j
                    break

        extracts.append({
            "idx": idx_counter,
            "caption": caption,
            "stem": stem,
            "tbl_idx_in_body": i,
            "caption_idx_in_body": caption_idx,
        })
        idx_counter += 1

    return extracts, idx_counter


# ---------- minimal-skeleton writer ----------

def _read_source_parts(src_docx: Path) -> dict:
    """Read needed parts from source docx zip (single open).

    Returns dict with keys: document_xml_root (lxml Element of <w:document>),
        sect_pr_xml (bytes or None), styles_xml (bytes), numbering_xml (bytes or None),
        theme_xml (bytes or None).
    """
    with zipfile.ZipFile(str(src_docx)) as z:
        names = set(z.namelist())
        doc_bytes = z.read("word/document.xml")
        styles_bytes = z.read("word/styles.xml") if "word/styles.xml" in names else None
        numbering_bytes = z.read("word/numbering.xml") if "word/numbering.xml" in names else None
        theme_bytes = z.read("word/theme/theme1.xml") if "word/theme/theme1.xml" in names else None

    # Parse document.xml; keep root element for namespace declarations + sectPr extraction.
    root = etree.fromstring(doc_bytes)
    # Locate body/sectPr (last child of body)
    body = root.find(qn("w:body"))
    sect_pr = None
    if body is not None:
        for child in reversed(list(body)):
            if child.tag == qn("w:sectPr"):
                sect_pr = child
                break

    return {
        "src_path": Path(src_docx),   # 部件完整性断言要用源件当基线
        "doc_root": root,
        "doc_body": body,
        "sect_pr": sect_pr,
        "styles": styles_bytes,
        "numbering": numbering_bytes,
        "theme": theme_bytes,
    }


def _build_minimal_document_xml(
    doc_root,
    tbl_elem,
    caption_elem,
    sect_pr,
) -> bytes:
    """Construct minimal word/document.xml with given table + optional caption + sectPr.

    Re-uses root namespace declarations of source document.xml so any
    namespaced attrs on the copied <w:tbl> still resolve.
    """
    # Build new root with same nsmap as source.
    nsmap = dict(doc_root.nsmap)
    new_root = etree.Element(qn("w:document"), nsmap=nsmap)
    # Preserve mc:Ignorable etc. attributes from source root.
    for k, v in doc_root.attrib.items():
        new_root.set(k, v)
    new_body = etree.SubElement(new_root, qn("w:body"))

    # Deep-copy caption paragraph (if any) and table, sectPr.
    from copy import deepcopy
    if caption_elem is not None:
        new_body.append(deepcopy(caption_elem))
    new_body.append(deepcopy(tbl_elem))
    if sect_pr is not None:
        new_body.append(deepcopy(sect_pr))
    else:
        # Minimal sectPr fallback
        etree.SubElement(new_body, qn("w:sectPr"))

    return b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + etree.tostring(
        new_root, xml_declaration=False, encoding="utf-8"
    )


def _has_inline_image(tbl_elem) -> bool:
    """Detect inline images (drawing or pict) inside a table."""
    drawing_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    pict_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
    for el in tbl_elem.iter():
        if el.tag == drawing_tag or el.tag == pict_tag:
            return True
    return False


def write_extract_minimal(
    src_parts: dict,
    dst_docx: Path,
    tbl_idx: int,
    caption_idx: int,
    body_children_cache: list,
) -> tuple[int, int, bool]:
    """Write a minimal docx containing the target table + caption + sectPr only.

    Returns (tbl_count_in_output, file_bytes, has_inline_image).
    """
    dst_docx.parent.mkdir(parents=True, exist_ok=True)

    tbl_elem = body_children_cache[tbl_idx]
    caption_elem = body_children_cache[caption_idx] if caption_idx >= 0 else None
    inline_img = _has_inline_image(tbl_elem)

    doc_xml = _build_minimal_document_xml(
        src_parts["doc_root"], tbl_elem, caption_elem, src_parts["sect_pr"]
    )

    has_numbering = src_parts["numbering"] is not None
    has_theme = src_parts["theme"] is not None

    content_types = _pick_content_types(has_numbering, has_theme)
    doc_rels = _build_doc_rels(has_numbering, has_theme)

    # Write zip.
    with zipfile.ZipFile(str(dst_docx), "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", _ROOT_RELS_XML)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
        z.writestr("word/document.xml", doc_xml)
        if src_parts["styles"] is not None:
            z.writestr("word/styles.xml", src_parts["styles"])
        else:
            # Should never hit (Word docs always have styles.xml) but degrade safely.
            z.writestr(
                "word/styles.xml",
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
            )
        if has_numbering:
            z.writestr("word/numbering.xml", src_parts["numbering"])
        if has_theme:
            z.writestr("word/theme/theme1.xml", src_parts["theme"])

    # ── 部件完整性断言（B 类：最小骨架**本意就要减部件**，assert_parts_intact 会把
    #    「其余全丢」误判成事故 → 改用精确部件集断言，白名单与上面的写入代码同源）
    expected = {
        "[Content_Types].xml",
        "_rels/.rels",
        "word/_rels/document.xml.rels",
        "word/document.xml",
        "word/styles.xml",
    }
    if has_numbering:
        expected.add("word/numbering.xml")
    if has_theme:
        expected.add("word/theme/theme1.xml")
    with zipfile.ZipFile(str(dst_docx)) as zchk:
        got = set(zchk.namelist())
    if got != expected:
        raise PartIntegrityError(
            f"{dst_docx.name} 部件集漂移: "
            f"多出 {sorted(got - expected)} / 缺失 {sorted(expected - got)}")
    # 搬运的部件（styles/numbering/theme）必须逐字节 verbatim；lost 是骨架策略的
    # 本意（media/header 等有意不带，CLI 已 NOTE 声明），不算错。
    d = diff_parts(src_parts["src_path"], dst_docx, allow_changed={
        "word/document.xml", "[Content_Types].xml",
        "word/_rels/document.xml.rels", "_rels/.rels",
    })
    fallback_added = {"word/styles.xml"} if src_parts["styles"] is None else set()
    bad_added = [n for n in d.added if n not in fallback_added]
    if d.changed or bad_added:
        raise PartIntegrityError(
            f"{dst_docx.name} 搬运部件被改写或未报备新增: "
            f"changed={d.changed} added={bad_added}")

    # Verify with python-docx.
    try:
        doc2 = Document(str(dst_docx))
        tbl_count = len(doc2.tables)
    except Exception:
        tbl_count = 0

    file_bytes = dst_docx.stat().st_size
    return tbl_count, file_bytes, inline_img


# ---------- runner / CLI ----------

def run_extract(
    src_docx: Path,
    out_dir: Path,
    dry_run: bool = False,
    name_pattern: str = "{stem}.docx",
    doc=None,
) -> dict:
    """Execute table-extract; reuses provided `doc` for planning if given."""
    src = Path(src_docx).expanduser().resolve()
    if not src.exists():
        return {"error": f"input docx not found: {src}", "exit_code": 2}
    out_dir = Path(out_dir).expanduser().resolve()

    extracts, tbl_count = plan_extracts(src, doc=doc)

    if tbl_count == 0:
        return {"tbl_count": 0, "extracts_emitted": 0, "note": "no tables in docx"}

    if dry_run:
        plan = []
        for e in extracts:
            fname = name_pattern.format(stem=e["stem"], idx=e["idx"])
            plan.append({
                "idx": e["idx"], "fname": fname,
                "caption": e["caption"], "tbl_idx": e["tbl_idx_in_body"],
            })
        return {"tbl_count": tbl_count, "extracts_planned": plan, "dry_run": True}

    out_dir.mkdir(parents=True, exist_ok=True)

    # Read source parts ONCE (avoid re-opening zip per table).
    src_parts = _read_source_parts(src)
    body_children = list(src_parts["doc_body"]) if src_parts["doc_body"] is not None else []

    emitted = []
    failed = []
    inline_img_count = 0
    for e in extracts:
        fname = name_pattern.format(stem=e["stem"], idx=e["idx"])
        dst = out_dir / fname
        try:
            tc, nbytes, has_img = write_extract_minimal(
                src_parts, dst, e["tbl_idx_in_body"], e["caption_idx_in_body"],
                body_children,
            )
            if has_img:
                inline_img_count += 1
            emitted.append({
                "idx": e["idx"], "fname": fname,
                "tables": tc, "bytes": nbytes,
                "caption": e["caption"],
                "inline_image": has_img,
            })
        except Exception as exc:
            failed.append({
                "idx": e["idx"], "fname": fname,
                "error": f"{type(exc).__name__}: {exc}",
            })
    return {
        "tbl_count": tbl_count,
        "extracts_emitted": len(emitted),
        "extracts_failed": len(failed),
        "inline_image_count": inline_img_count,
        "out_dir": str(out_dir),
        "emitted": emitted,
        "failed": failed,
    }


def main_table_extract() -> int:
    ap = argparse.ArgumentParser(
        description="Extract each table in a DOCX as an independent minimal DOCX "
                    "(filename = neighboring caption text, or table-XX fallback).",
    )
    ap.add_argument("--docx", required=True, help="input docx path")
    ap.add_argument("--out-dir", required=True, help="output directory (mkdir -p)")
    ap.add_argument(
        "--name-pattern",
        default="{stem}.docx",
        help="output filename pattern, default '{stem}.docx' "
             "(available: {stem}, {idx})",
    )
    ap.add_argument(
        "--dry-run", action="store_true",
        help="print plan only, don't write files",
    )
    args = ap.parse_args()

    src = Path(args.docx).expanduser().resolve()
    if not src.exists():
        print(f"ERROR: input docx not found: {src}", file=sys.stderr)
        return 2
    out_dir = Path(args.out_dir).expanduser().resolve()

    extracts, tbl_count = plan_extracts(src)

    if tbl_count == 0:
        print(f"[extract-tables] 0 tables in docx — nothing to extract")
        return 0

    if args.dry_run:
        print(f"# DRY RUN — would write {len(extracts)} files to {out_dir}/")
        for e in extracts:
            fname = args.name_pattern.format(stem=e["stem"], idx=e["idx"])
            cap = (e["caption"] or "[no caption]")[:60]
            print(f"  [{e['idx']:>2}] tbl@body[{e['tbl_idx_in_body']}]  → {fname}  (caption={cap!r})")
        return 0

    out_dir.mkdir(parents=True, exist_ok=True)
    src_parts = _read_source_parts(src)
    body_children = list(src_parts["doc_body"]) if src_parts["doc_body"] is not None else []

    n_ok = 0
    inline_img_count = 0
    for e in extracts:
        fname = args.name_pattern.format(stem=e["stem"], idx=e["idx"])
        dst = out_dir / fname
        try:
            tc, nbytes, has_img = write_extract_minimal(
                src_parts, dst, e["tbl_idx_in_body"], e["caption_idx_in_body"],
                body_children,
            )
            if has_img:
                inline_img_count += 1
            img_flag = " [img]" if has_img else ""
            print(f"  [{e['idx']:>2}] {fname}  · {tc} tables · {nbytes:,} bytes{img_flag}")
            n_ok += 1
        except Exception as exc:
            print(f"  [{e['idx']:>2}] FAILED {fname}: {type(exc).__name__}: {exc}",
                  file=sys.stderr)
    print(f"OK: {n_ok} files written to {out_dir}")
    if inline_img_count:
        print(f"NOTE: {inline_img_count} table(s) reference inline images; "
              "the minimal docx omits media files (Word will show broken refs).")
    return 0 if n_ok == len(extracts) else 1


# ══════════ center ← set_table_align.py ══════════

_TBLPR_AFTER_CENTER = ("tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook",
                "tblCaption", "tblDescription")
# pPr 子元素顺序：jc 须排在这些之前（足够覆盖常见情况）
_PPR_AFTER = ("rPr",)


def _insert_in_order(parent, child, after_tags):
    """把 child 插到 parent 中第一个属于 after_tags 的元素之前，否则追加到末尾。"""
    after_qn = {qn(f"w:{t}") for t in after_tags}
    for existing in parent:
        if existing.tag in after_qn:
            existing.addprevious(child)
            return
    parent.append(child)


def _set_table_jc(tbl_el, val: str = "center") -> None:
    """设/重置表级 tblPr/jc。tblPr 是 tbl 的必有首子元素。"""
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        _insert_in_order(tblPr, jc, _TBLPR_AFTER_CENTER)
    jc.set(qn("w:val"), val)


def _center_cells(tbl_el) -> int:
    """单元格内段落水平居中 + 单元格垂直居中。返回处理的单元格数。"""
    n = 0
    for tc in tbl_el.iter(qn("w:tc")):
        # 垂直居中：tcPr/vAlign=center
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        v = tcPr.find(qn("w:vAlign"))
        if v is None:
            v = OxmlElement("w:vAlign")
            tcPr.append(v)
        v.set(qn("w:val"), "center")
        # 水平居中：每个直属段落 pPr/jc=center
        for p in tc.findall(qn("w:p")):
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p.insert(0, pPr)
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                _insert_in_order(pPr, jc, _PPR_AFTER)
            jc.set(qn("w:val"), "center")
        n += 1
    return n


def _all_tbl_elements(doc):
    """所有 <w:tbl>，含单元格内嵌套表。"""
    return list(doc.element.body.iter(qn("w:tbl")))


def apply_center(doc, args=None) -> dict:
    """在**已打开**的 doc 上把所有表格整体居中（可选连单元格内文字一起居中）。

    改动全在 document.xml 的 <w:tbl>/<w:tc>/<w:p> 上，不碰 zip 里别的部件，
    所以是 doc-based step —— 存盘、备份、复检都归调用方。
    """
    cell_center = bool(getattr(args, "cell_center", False)) if args else False

    tbls = _all_tbl_elements(doc)
    cells = 0
    for tbl_el in tbls:
        _set_table_jc(tbl_el, "center")
        if cell_center:
            cells += _center_cells(tbl_el)

    return {
        "changed": len(tbls),
        "tables": len(tbls),
        "cell_center": cell_center,
        "cells": cells,
    }


def process_center(docx_path: Path, cell_center: bool, dry_run: bool, backup: bool) -> dict:
    doc = Document(str(docx_path))
    result = {
        "file": str(docx_path),
        **apply_center(doc, argparse.Namespace(cell_center=cell_center)),
        "dry_run": dry_run,
    }
    n_tbl = result["tables"]
    if dry_run or n_tbl == 0:
        result["written"] = False
        return result

    if backup:
        bak = docx_path.with_name(
            docx_path.name + ".bak-" + datetime.now().strftime("%Y%m%d-%H%M%S"))
        shutil.copy2(docx_path, bak)
        result["backup"] = str(bak)

    doc.save(str(docx_path))
    result["written"] = True

    # 验证：重读，统计表级 jc=center 的表数
    doc2 = Document(str(docx_path))
    centered = 0
    for tbl_el in _all_tbl_elements(doc2):
        tblPr = tbl_el.find(qn("w:tblPr"))
        jc = tblPr.find(qn("w:jc")) if tblPr is not None else None
        if jc is not None and jc.get(qn("w:val")) == "center":
            centered += 1
    result["verify_centered"] = centered
    return result


def _fmt_center(r: dict) -> str:
    head = f"[table center] {Path(r['file']).name}"
    if r["tables"] == 0:
        return f"{head}: 无表格，跳过"
    tail = ""
    if r.get("written"):
        tail = f" → centered {r.get('verify_centered')}/{r['tables']}"
        if r["cell_center"]:
            tail += f" · cells={r['cells']}"
        if r.get("backup"):
            tail += f" · bak={Path(r['backup']).name}"
    elif r["dry_run"]:
        tail = " (dry-run)"
    return f"{head}: {r['tables']} 表 · 整体居中{tail}"


def main_table_center() -> int:
    ap = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("docx_pos", nargs="?", help="(positional) docx 路径，等价 --docx")
    ap.add_argument("--docx", dest="docx_kw", help="目标 docx（原地修改）")
    ap.add_argument("--cell-center", action="store_true",
                    help="同时把单元格内文字水平+垂直居中（默认只表格整体居中）")
    ap.add_argument("--no-backup", action="store_true", help="不创建 .bak-时间戳 备份")
    ap.add_argument("--dry-run", action="store_true", help="只报告不写盘")
    args = ap.parse_args()

    docx = args.docx_kw or args.docx_pos
    if not docx:
        print("[table center] missing docx (positional or --docx)", file=sys.stderr)
        return 2
    docx_path = Path(docx)
    if not docx_path.exists():
        print(f"[table center] not found: {docx_path}", file=sys.stderr)
        return 2

    r = process_center(docx_path, cell_center=args.cell_center,
                dry_run=args.dry_run, backup=not args.no_backup)
    print(_fmt_center(r))
    return 0


# ---------------- pipeline adapter ----------------
def apply_path_center(docx_path, args=None) -> dict:
    """原地 mutator（自己开文件、备份、存盘、复检）。

    留着是为了不掐断已经按路径调它的老调用方；pipeline_lib.load_step 优先取
    apply()，新链路一律走纯内存版本，这里不会被 pipeline 选中。
    """
    cell = bool(getattr(args, "cell_center", False)) if args else False
    dry = bool(getattr(args, "dry_run", False)) if args else False
    backup = not bool(getattr(args, "no_backup", False)) if args else True
    return process_center(Path(docx_path), cell, dry, backup)


# ══════════ borders ← set_table_borders.py ══════════

EDGES = ("top", "left", "bottom", "right", "insideH", "insideV")
# tblPr 子元素中必须排在 tblBorders 之后者（用于定位插入点）
_TBLPR_AFTER_BORDERS = ("shd", "tblLayout", "tblCellMar", "tblLook",
                "tblCaption", "tblDescription")
# tcPr 子元素中必须排在 tcBorders 之后者
_TCPR_AFTER = ("shd", "noWrap", "tcMar", "textDirection", "tcFitText",
               "vAlign", "hideMark")


def _mk_border(tag: str, val: str, sz: int, color: str, space: int):
    el = OxmlElement(f"w:{tag}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    return el


def _insert_in_order(parent, child, after_tags):
    """把 child 插到 parent 中第一个属于 after_tags 的元素之前，否则追加到末尾。"""
    after_qn = {qn(f"w:{t}") for t in after_tags}
    for existing in parent:
        if existing.tag in after_qn:
            existing.addprevious(child)
            return
    parent.append(child)


def _set_table_borders(tbl_el, val: str, sz: int, color: str, space: int) -> None:
    """设/重置表级 tblBorders 为指定 6 边。tblPr 是 tbl 的必有首子元素。"""
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:  # 理论上不会发生（schema 必有）
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in EDGES:
        borders.append(_mk_border(edge, val, sz, color, space))
    _insert_in_order(tblPr, borders, _TBLPR_AFTER_BORDERS)


def _strip_cell_borders(tbl_el) -> int:
    """删除每个单元格的 tcBorders（让表级统一生效）。返回删除个数。"""
    n = 0
    for tcPr in tbl_el.iter(qn("w:tcPr")):
        tcB = tcPr.find(qn("w:tcBorders"))
        if tcB is not None:
            tcPr.remove(tcB)
            n += 1
    return n


def _solidify_cell_borders(tbl_el, val: str, sz: int, color: str, space: int) -> int:
    """保留 tcBorders，但把任何非 single 的边（nil/none/dashed…）改写为实线。
    返回改写的边数。"""
    n = 0
    for tcB in tbl_el.iter(qn("w:tcBorders")):
        for edge in EDGES:
            el = tcB.find(qn(f"w:{edge}"))
            if el is None:
                continue
            if el.get(qn("w:val")) != val:
                el.set(qn("w:val"), val)
                el.set(qn("w:sz"), str(sz))
                el.set(qn("w:space"), str(space))
                el.set(qn("w:color"), color)
                n += 1
    return n


def _all_tbl_elements_iter(doc):
    """所有 <w:tbl>，含单元格内嵌套表。"""
    return doc.element.body.iter(qn("w:tbl"))


def apply_borders(doc, args=None) -> dict:
    """在**已打开**的 doc 上把所有表格改成满格实线（表级 + 单元格级两手抓）。

    之所以能纯内存做完：全部改动只落在 document.xml 的 <w:tbl> 子树上，
    styles.xml / numbering.xml 一概不碰 —— 所以它是 doc-based step，能和别的
    apply() 串在同一棵树上跑，全程只 parse 一次、只 save 一次（存盘是调用方的事）。
    """
    val = getattr(args, "val", "single") if args else "single"
    sz = int(getattr(args, "sz", 4)) if args else 4
    color = getattr(args, "color", "auto") if args else "auto"
    space = int(getattr(args, "space", 0)) if args else 0
    keep_cell = bool(getattr(args, "keep_cell_borders", False)) if args else False

    tbls = list(_all_tbl_elements_iter(doc))
    cell_changed = 0
    for tbl_el in tbls:
        _set_table_borders(tbl_el, val, sz, color, space)
        if keep_cell:
            cell_changed += _solidify_cell_borders(tbl_el, val, sz, color, space)
        else:
            cell_changed += _strip_cell_borders(tbl_el)

    return {
        "changed": len(tbls),
        "tables": len(tbls),
        "mode": "keep-cell" if keep_cell else "strip-cell",
        "cell_changed": cell_changed,
        "border": f"{val}/sz{sz}/{color}",
    }


def process_borders(docx_path: Path, val: str, sz: int, color: str, space: int,
            keep_cell: bool, dry_run: bool, backup: bool) -> dict:
    doc = Document(str(docx_path))
    result = {
        "file": str(docx_path),
        **apply_borders(doc, argparse.Namespace(
            val=val, sz=sz, color=color, space=space, keep_cell_borders=keep_cell)),
        "dry_run": dry_run,
    }
    n_tbl = result["tables"]
    if dry_run or n_tbl == 0:
        result["written"] = False
        return result

    if backup:
        bak = docx_path.with_name(
            docx_path.name + ".bak-" + datetime.now().strftime("%Y%m%d-%H%M%S"))
        shutil.copy2(docx_path, bak)
        result["backup"] = str(bak)

    doc.save(str(docx_path))
    result["written"] = True

    # 验证：重读，统计表级 6 边齐全 + 残留 nil 边
    doc2 = Document(str(docx_path))
    full = 0
    residual_nil = 0
    for tbl_el in _all_tbl_elements_iter(doc2):
        tblB = tbl_el.find(qn("w:tblPr"))
        tblB = tblB.find(qn("w:tblBorders")) if tblB is not None else None
        if tblB is not None and all(
            (e := tblB.find(qn(f"w:{edge}"))) is not None
            and e.get(qn("w:val")) == val for edge in EDGES
        ):
            full += 1
        for tcB in tbl_el.iter(qn("w:tcBorders")):
            for edge in EDGES:
                e = tcB.find(qn(f"w:{edge}"))
                if e is not None and e.get(qn("w:val")) in ("nil", "none"):
                    residual_nil += 1
    result["verify_full_grid"] = full
    result["verify_residual_nil"] = residual_nil
    return result


def _fmt_borders(r: dict) -> str:
    head = f"[table borders] {Path(r['file']).name}"
    if r["tables"] == 0:
        return f"{head}: 无表格，跳过"
    tail = ""
    if r.get("written"):
        tail = (f" → full-grid {r.get('verify_full_grid')}/{r['tables']}"
                f" · residual-nil {r.get('verify_residual_nil')}")
        if r.get("backup"):
            tail += f" · bak={Path(r['backup']).name}"
    elif r["dry_run"]:
        tail = " (dry-run)"
    return (f"{head}: {r['tables']} 表 · {r['mode']} · cell_changed={r['cell_changed']}"
            f" · {r['border']}{tail}")


def main_table_borders():
    ap = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("docx_pos", nargs="?", help="(positional) docx 路径，等价 --docx")
    ap.add_argument("--docx", dest="docx_kw", help="目标 docx（原地修改）")
    ap.add_argument("--val", default="single",
                    help="边框线型 (single/double/dashed/...)，默认 single 实线")
    ap.add_argument("--sz", type=int, default=4,
                    help="线宽，单位 1/8 pt（4=0.5pt），默认 4")
    ap.add_argument("--color", default="auto", help="边框颜色，默认 auto（=黑）")
    ap.add_argument("--space", type=int, default=0, help="边距，默认 0")
    ap.add_argument("--keep-cell-borders", action="store_true",
                    help="保留单元格 tcBorders，仅把非实线边改写为实线（默认=直接删 tcBorders）")
    ap.add_argument("--no-backup", action="store_true", help="不创建 .bak-时间戳 备份")
    ap.add_argument("--dry-run", action="store_true", help="只报告不写盘")
    args = ap.parse_args()

    docx = args.docx_kw or args.docx_pos
    if not docx:
        print("[table borders] missing docx (positional or --docx)", file=sys.stderr)
        return 2
    docx_path = Path(docx)
    if not docx_path.exists():
        print(f"[table borders] not found: {docx_path}", file=sys.stderr)
        return 2

    r = process_borders(docx_path, args.val, args.sz, args.color, args.space,
                keep_cell=args.keep_cell_borders, dry_run=args.dry_run,
                backup=not args.no_backup)
    print(_fmt_borders(r))
    return 0


# ---------------- pipeline adapter ----------------
def apply_path_borders(docx_path, args=None) -> dict:
    """原地 mutator（自己开文件、备份、存盘、复检）。

    留着是为了不掐断已经按路径调它的老调用方；pipeline_lib.load_step 优先取
    apply()，所以新链路一律走上面那个纯内存版本，这里不会被 pipeline 选中。
    """
    val = getattr(args, "val", "single") if args else "single"
    sz = int(getattr(args, "sz", 4)) if args else 4
    color = getattr(args, "color", "auto") if args else "auto"
    space = int(getattr(args, "space", 0)) if args else 0
    keep = bool(getattr(args, "keep_cell_borders", False)) if args else False
    dry = bool(getattr(args, "dry_run", False)) if args else False
    backup = not bool(getattr(args, "no_backup", False)) if args else True
    return process_borders(Path(docx_path), val, sz, color, space, keep, dry, backup)


# ──────────────────────────── 家族入口（子命令分发）────────────────────────────

SUBCOMMANDS = {
    "delete-rows": main_delete_rows,
    "extract": main_table_extract,
    "center": main_table_center,
    "borders": main_table_borders,
}


def main(argv: list[str] | None = None) -> int:
    args = list(sys.argv[1:] if argv is None else argv)
    if not args or args[0] in ("-h", "--help"):
        print("usage: table.py {" + ",".join(SUBCOMMANDS) + "} <args…>\n"
              "每个子命令的参数与原独立脚本逐字一致：table.py <sub> --help 查看。")
        return 0 if args else 2
    sub, rest = args[0], args[1:]
    fn = SUBCOMMANDS.get(sub)
    if fn is None:
        print(f"[table] unknown subcommand: {sub!r}; choices={list(SUBCOMMANDS)}",
              file=sys.stderr)
        return 2
    saved = sys.argv[:]
    sys.argv = [sys.argv[0]] + rest
    try:
        rc = fn()
        return int(rc) if isinstance(rc, int) else 0
    finally:
        sys.argv = saved


if __name__ == "__main__":
    sys.exit(main())
