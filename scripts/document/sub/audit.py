#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""audit.py — audit_* 家族六合一（2026-07-31 家族折叠，全部 read-only）

子命令 ↔ 原脚本（函数体逐字搬移；模块级 audit/apply/apply_path/main 改名
audit_<sub>/apply_<sub>/apply_path_<sub>/main_<sub>，其余公有名一律保留——
health.py 靠 audit_table_pairing / PREFIX_RE / _normalize_style）：

    bookmarks      ← audit_bookmarks.py
    captions       ← audit_caption_outline.py
    headings       ← audit_heading_numbers.py
    images         ← audit_images.py
    table-pairing  ← audit_table_pairing.py
    fields         ← audit_word_fields.py

各子命令 CLI 与原独立脚本逐字一致：python3 sub/audit.py <sub> <docx> [flags…]。
退役原件在 ~/.Trash/consolidation-20260731/audit/（含 MANIFEST.md）。
"""
from __future__ import annotations

import argparse
import os
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

from lxml import etree

# sub/ 自身进 sys.path —— docx_cli 的 _dispatch 用 spec_from_file_location 加载,
# 不带脚本目录, 裸 import _cli_common 会 ImportError (append 不是 insert(0))
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parent))
import _cli_common as _cc  # noqa: E402  家族 main() 样板 SSOT

try:
    from docx import Document
except ImportError:
    print("[ERR] 缺 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(2)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
NSMAP = {"w": W_NS}


# ══════════ bookmarks ← audit_bookmarks.py ══════════

# Word 内部自动 bookmark 前缀 (启发集合)
KNOWN_INTERNAL_PREFIXES = ("_Toc", "_Ref", "_Hlk", "_GoBack", "_Hyperlink", "OLE_LINK")


def _classify_name(name: str) -> str:
    """返回 prefix 类别 (用于 bookmark_by_prefix 计数)."""
    if not name:
        return "EMPTY"
    if name == "_GoBack":
        return "_GoBack"
    for p in ("_Toc", "_Ref", "_Hlk", "_Hyperlink", "OLE_LINK"):
        if name.startswith(p):
            return p
    return "user_defined"


def _para_idx_map(body) -> dict:
    """每个 <w:p> 元素 → 物理顺序 idx."""
    return {p: i for i, p in enumerate(body.iter(f"{W}p"))}


def audit_bookmarks(doc) -> dict:
    """扫 doc body 全部 bookmarkStart/bookmarkEnd, 返回 report dict."""
    body = doc.element.body
    para_idx_map = _para_idx_map(body)

    starts: dict[str, dict] = {}  # id → {name, para_idx}
    ends: dict[str, int] = {}     # id → end para_idx

    for bs in body.iter(f"{W}bookmarkStart"):
        bid = bs.get(f"{W}id")
        name = bs.get(f"{W}name") or ""
        if bid is None:
            continue
        starts[bid] = {
            "name": name,
            "start_para_idx": _ancestor_para_idx(bs, para_idx_map),
        }
    for be in body.iter(f"{W}bookmarkEnd"):
        bid = be.get(f"{W}id")
        if bid is None:
            continue
        ends[bid] = _ancestor_para_idx(be, para_idx_map)

    bookmarks: list[dict] = []
    orphan_starts: list[dict] = []
    orphan_ends: list[str] = []
    by_prefix: Counter = Counter()
    para_idx_for_each: dict[str, dict] = {}

    for bid, info in starts.items():
        name = info["name"]
        rec = {
            "id": bid,
            "name": name,
            "start_para_idx": info["start_para_idx"],
            "end_para_idx": ends.get(bid, -1),
            "prefix": _classify_name(name),
        }
        bookmarks.append(rec)
        by_prefix[rec["prefix"]] += 1
        para_idx_for_each[bid] = {
            "name": name,
            "start": info["start_para_idx"],
            "end": ends.get(bid, -1),
        }
        if bid not in ends:
            orphan_starts.append({"id": bid, "name": name,
                                  "start_para_idx": info["start_para_idx"]})
    for bid in ends:
        if bid not in starts:
            orphan_ends.append(bid)

    return {
        "bookmark_count": len(bookmarks),
        "bookmark_by_prefix": dict(by_prefix),
        "bookmarks": bookmarks,
        "orphan_starts": orphan_starts,
        "orphan_ends": orphan_ends,
        "para_idx_for_each": para_idx_for_each,
    }


# ---------------- pipeline adapter ----------------
def apply_bookmarks(doc, args=None) -> dict:
    """pipeline 调用 — 只读 audit, 不改 doc."""
    return audit_bookmarks(doc)


def main_bookmarks(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Audit bookmarks in a docx (read-only).")
    ap.add_argument("docx", type=Path)
    _cc.add_report_flag(ap, help="Write full JSON report to this path")
    args = ap.parse_args(argv)

    if not args.docx.exists():
        print(f"[ERR] 文件不存在: {args.docx}", file=sys.stderr)
        return 1

    doc = Document(str(args.docx))
    report = audit_bookmarks(doc)
    report["docx"] = str(args.docx.resolve())

    summary = {
        "docx": report["docx"],
        "bookmark_count": report["bookmark_count"],
        "bookmark_by_prefix": report["bookmark_by_prefix"],
        "orphan_starts_count": len(report["orphan_starts"]),
        "orphan_ends_count": len(report["orphan_ends"]),
    }
    _cc.print_json(summary)

    _cc.write_report(report, args.report)
    return 0


# ══════════ captions ← audit_caption_outline.py ══════════

# 兼容 `图1.2-1` 三段、`表1-1` 两段、`表 1.2-1`、`图1.1` 两段(纯点)
CAPTION_RE = re.compile(r"^\s*(表|图)\s*\d+([\.\-]\d+){1,2}")

# Caption family styles — docx 里允许的 caption 类样式名(集团 ZDWP 命名 + 通用)
CAPTION_FAMILY_KEYWORDS = ("caption", "表名", "图名", "0图", "0表", "zdwp 表名", "zdwp 图名")


def is_caption_style(style_name: str) -> bool:
    if not style_name:
        return False
    low = style_name.lower().strip()
    for kw in CAPTION_FAMILY_KEYWORDS:
        if kw in low:
            return True
    return False


def get_outline_lvl(p_elem) -> str | None:
    """返回 paragraph 的 outlineLvl 值(字符串)或 None。"""
    pPr = p_elem.find(f"{{{W_NS}}}pPr")
    if pPr is None:
        return None
    ol = pPr.find(f"{{{W_NS}}}outlineLvl")
    if ol is None:
        return None
    return ol.get(f"{{{W_NS}}}val")


def get_style_name(p, doc) -> str:
    """返回 paragraph 的 style 名称(从 styles.xml 解析 styleId → name)。"""
    try:
        return p.style.name if p.style is not None else ""
    except Exception:
        return ""


def collect_available_caption_styles(doc) -> list[str]:
    """列 docx styles.xml 里所有 caption-family 样式名。"""
    result = []
    for s in doc.styles:
        try:
            if is_caption_style(s.name):
                result.append(s.name)
        except Exception:
            continue
    return sorted(set(result))


def _audit_from_doc_caption_outline(doc, docx_path_label: str = "") -> dict:
    body_paragraphs = doc.paragraphs

    total = len(body_paragraphs)
    captions_total = 0
    captions_with_outlinelvl: list[dict] = []
    captions_by_style: Counter = Counter()
    captions_clean_examples: list[dict] = []
    h_count: Counter = Counter()
    polluted_count = 0
    wrong_style_count = 0
    empty_caption = 0
    all_caption_records: list[dict] = []

    for idx, p in enumerate(body_paragraphs):
        style_name = get_style_name(p, doc)
        text = (p.text or "").strip()

        # H styles 统计
        if style_name and style_name.startswith("Heading "):
            h_count[style_name] += 1

        # caption 识别:文本匹配 caption 正则 OR style 已套 caption-family (但 text 是空的也算被识别为 caption)
        text_is_caption = bool(CAPTION_RE.match(text))
        style_is_caption = is_caption_style(style_name)

        if text_is_caption or style_is_caption:
            captions_total += 1
            outline_lvl = get_outline_lvl(p._p)
            captions_by_style[style_name or "(no-style)"] += 1

            record = {
                "idx": idx,
                "style": style_name,
                "outlineLvl": outline_lvl,
                "text": text[:80],
            }
            all_caption_records.append(record)

            # 空 caption(style 套了但文本空)
            if not text:
                empty_caption += 1

            # 污染:有 outlineLvl 且数值 <= 4(0-4 表示 H1-H5 级被错继承)
            if outline_lvl is not None:
                try:
                    lvl_int = int(outline_lvl)
                    if lvl_int <= 6:  # 0-6 都算污染(<=6 = 进 outline 大纲)
                        polluted_count += 1
                        if len(captions_with_outlinelvl) < 30:
                            captions_with_outlinelvl.append(record)
                except ValueError:
                    pass

            # style 错配:文本是 caption 形态但 style 不是 caption-family 也不是空(像 Normal/正文)
            if text_is_caption and not style_is_caption:
                wrong_style_count += 1

            # clean 样本:outline 清 + style 正确
            if outline_lvl is None and style_is_caption and len(captions_clean_examples) < 10:
                captions_clean_examples.append(record)

    h_styles_present = sorted(h_count.keys())
    available_caption_styles = collect_available_caption_styles(doc)

    return {
        "docx_path": docx_path_label,
        "total_paragraphs": total,
        "captions_total": captions_total,
        "captions_with_outlinelvl": captions_with_outlinelvl,
        "captions_by_style": dict(captions_by_style),
        "captions_clean_examples": captions_clean_examples,
        "h_styles_present": h_styles_present,
        "h_count": dict(h_count),
        "caption_styles_available": available_caption_styles,
        "all_caption_records": all_caption_records,  # 用于 sample 5 抽样
        "issues": {
            "polluted_outline_count": polluted_count,
            "wrong_style_count": wrong_style_count,
            "empty_caption": empty_caption,
        },
    }


def audit_caption_outline(docx_path: Path) -> dict:
    doc = Document(str(docx_path))
    return _audit_from_doc_caption_outline(doc, str(docx_path))


def apply_caption_outline(doc, args=None) -> dict:
    """pipeline read-only adapter"""
    label = str(getattr(args, "docx", "")) if args else ""
    return _audit_from_doc_caption_outline(doc, label)


def main_caption_outline():
    ap = argparse.ArgumentParser(description="audit-only · caption 段 outlineLvl + style 污染审计")
    ap.add_argument("docx", help="docx 路径")
    _cc.add_report_flag(ap, help="JSON 报告输出路径(可选)")
    args = ap.parse_args()

    docx_path = Path(args.docx)
    if not docx_path.exists():
        print(f"ERROR: docx 不存在: {docx_path}", file=sys.stderr)
        sys.exit(1)

    report = audit_caption_outline(docx_path)

    # 控制台 summary
    print(f"=== audit_caption_outline · {docx_path.name} ===")
    print(f"total_paragraphs       : {report['total_paragraphs']}")
    print(f"captions_total         : {report['captions_total']}")
    print(f"polluted_outline_count : {report['issues']['polluted_outline_count']}")
    print(f"wrong_style_count      : {report['issues']['wrong_style_count']}")
    print(f"empty_caption          : {report['issues']['empty_caption']}")
    print(f"captions_by_style      : {report['captions_by_style']}")
    print(f"h_count                : {report['h_count']}")
    print(f"caption_styles_avail   : {report['caption_styles_available']}")

    _cc.write_report(report, args.report, mkdir=False, announce="[report] {path}")


# ══════════ headings ← audit_heading_numbers.py ══════════

PREFIX_RE = re.compile(r"^\d+(?:\.\d+)*\s")


def _strip_ns(tag: str) -> str:
    return tag.split("}", 1)[-1] if "}" in tag else tag


def _read_xml(z: zipfile.ZipFile, name: str) -> str:
    with z.open(name) as f:
        return f.read().decode("utf-8")


def _para_text(p_xml: str) -> str:
    """concat all w:t in paragraph xml string"""
    texts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", p_xml)
    return "".join(texts)


def _para_style(p_xml: str) -> str | None:
    m = re.search(r'<w:pStyle w:val="([^"]+)"', p_xml)
    return m.group(1) if m else None


def _para_numpr(p_xml: str) -> tuple[bool, str | None, str | None]:
    """returns (has_numPr, numId, ilvl) at paragraph level only"""
    # only look in <w:pPr> .. </w:pPr>
    ppr_m = re.search(r"<w:pPr>(.*?)</w:pPr>", p_xml, re.DOTALL)
    if not ppr_m:
        return (False, None, None)
    ppr = ppr_m.group(1)
    if "<w:numPr>" not in ppr and "<w:numPr/>" not in ppr:
        return (False, None, None)
    numId_m = re.search(r'<w:numId w:val="(\d+)"', ppr)
    ilvl_m = re.search(r'<w:ilvl w:val="(\d+)"', ppr)
    return (True, numId_m.group(1) if numId_m else None, ilvl_m.group(1) if ilvl_m else None)


def _split_paragraphs(doc_xml: str) -> list[str]:
    # naive splitter on top-level w:p; good enough for read-only flat scan
    paras = []
    i = 0
    while True:
        start = doc_xml.find("<w:p ", i)
        s2 = doc_xml.find("<w:p>", i)
        if start < 0 and s2 < 0:
            break
        if start < 0 or (s2 >= 0 and s2 < start):
            start = s2
        # find matching </w:p> (no nested w:p in OOXML)
        end = doc_xml.find("</w:p>", start)
        if end < 0:
            break
        paras.append(doc_xml[start:end + 6])
        i = end + 6
    return paras


HEADING_STYLE_MAP = {
    "1": "Heading 1",
    "2": "Heading 2",
    "3": "Heading 3",
    "4": "Heading 4",
    "Heading1": "Heading 1",
    "Heading2": "Heading 2",
    "Heading3": "Heading 3",
    "Heading4": "Heading 4",
}


def _normalize_style(style_id: str | None) -> str | None:
    if style_id is None:
        return None
    return HEADING_STYLE_MAP.get(style_id, style_id)


def _scan_styles(styles_xml: str) -> dict:
    """find heading 1-4 style blocks; report numPr status"""
    out: dict[str, dict] = {}
    for sid in ["1", "2", "3", "4", "Heading1", "Heading2", "Heading3", "Heading4"]:
        pat = f'<w:style w:type="paragraph" w:styleId="{sid}"'
        idx = styles_xml.find(pat)
        if idx < 0:
            # fallback: any w:style w:styleId="{sid}"
            idx = styles_xml.find(f'w:styleId="{sid}"')
            if idx < 0:
                continue
            idx = styles_xml.rfind("<w:style ", 0, idx)
            if idx < 0:
                continue
        end = styles_xml.find("</w:style>", idx)
        block = styles_xml[idx:end + 10]
        name_m = re.search(r'<w:name w:val="([^"]+)"', block)
        name = name_m.group(1) if name_m else f"styleId={sid}"
        if not name.lower().startswith("heading"):
            continue
        norm = _normalize_style(name) or name
        if norm in out:
            continue
        has_numpr = "<w:numPr>" in block or "<w:numPr/>" in block
        numId_m = re.search(r'<w:numPr>.*?<w:numId w:val="(\d+)"', block, re.DOTALL)
        ilvl_m = re.search(r'<w:numPr>.*?<w:ilvl w:val="(\d+)"', block, re.DOTALL)
        if has_numpr:
            status = f"exists numId={numId_m.group(1) if numId_m else None} ilvl={ilvl_m.group(1) if ilvl_m else None}"
        else:
            status = "removed (no w:numPr)"
        out[norm.lower()] = {
            "styleId": sid,
            "name": name,
            "has_numpr": has_numpr,
            "numId": numId_m.group(1) if numId_m else None,
            "ilvl": ilvl_m.group(1) if ilvl_m else None,
            "status": status,
        }
    return out


def _scan_numbering(numbering_xml: str, num_ids: list[str]) -> dict:
    """for each numId, resolve abstractNumId and dump lvlText per ilvl 0..4"""
    out: dict = {}
    for num_id in num_ids:
        if num_id is None or num_id in out:
            continue
        m = re.search(rf'<w:num w:numId="{num_id}"[^>]*>(.*?)</w:num>', numbering_xml, re.DOTALL)
        if not m:
            out[num_id] = {"error": "numId not found"}
            continue
        abs_m = re.search(r'<w:abstractNumId w:val="(\d+)"', m.group(1))
        if not abs_m:
            out[num_id] = {"error": "abstractNumId not found"}
            continue
        abs_id = abs_m.group(1)
        start = numbering_xml.find(f'<w:abstractNum w:abstractNumId="{abs_id}"')
        end = numbering_xml.find("</w:abstractNum>", start)
        block = numbering_xml[start:end + 16] if start >= 0 else ""
        lvltexts: dict[str, dict] = {}
        for ilvl in range(5):
            lvl_start = block.find(f'<w:lvl w:ilvl="{ilvl}"')
            if lvl_start < 0:
                continue
            lvl_end = block.find("</w:lvl>", lvl_start)
            lvl_block = block[lvl_start:lvl_end]
            lt = re.search(r'<w:lvlText w:val="([^"]*)"', lvl_block)
            nfmt = re.search(r'<w:numFmt w:val="([^"]+)"', lvl_block)
            lvltexts[str(ilvl)] = {
                "numFmt": nfmt.group(1) if nfmt else None,
                "lvlText": lt.group(1) if lt else None,
            }
        out[num_id] = {"abstractNumId": abs_id, "lvltexts": lvltexts}
    return out


def audit_heading_numbers(docx_path: Path) -> dict:
    with zipfile.ZipFile(docx_path) as z:
        doc_xml = _read_xml(z, "word/document.xml")
        styles_xml = _read_xml(z, "word/styles.xml")
        try:
            numbering_xml = _read_xml(z, "word/numbering.xml")
        except KeyError:
            numbering_xml = ""

    paras = _split_paragraphs(doc_xml)

    h_count_by_level: dict[str, int] = {}
    h_with_prefix = 0
    h_without_prefix = 0
    h_with_para_numpr = 0
    samples_no_prefix: list[list] = []
    samples_with_prefix: list[list] = []
    h_details: list[dict] = []

    for idx, p in enumerate(paras):
        style_raw = _para_style(p)
        style_norm = _normalize_style(style_raw)
        if style_norm not in {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}:
            continue
        text = _para_text(p)
        text50 = text[:50]
        has_npr, p_numId, p_ilvl = _para_numpr(p)
        has_prefix = bool(PREFIX_RE.match(text))
        prefix_literal = ""
        if has_prefix:
            m = PREFIX_RE.match(text)
            if m:
                prefix_literal = m.group(0).strip()
        h_count_by_level[style_norm] = h_count_by_level.get(style_norm, 0) + 1
        if has_prefix:
            h_with_prefix += 1
            if len(samples_with_prefix) < 10:
                samples_with_prefix.append([idx, style_norm, prefix_literal, text50])
        else:
            h_without_prefix += 1
            if len(samples_no_prefix) < 10:
                samples_no_prefix.append([idx, style_norm, text50])
        if has_npr:
            h_with_para_numpr += 1
        h_details.append({
            "idx": idx,
            "style": style_norm,
            "text50": text50,
            "para_numpr": has_npr,
            "para_numId": p_numId,
            "para_ilvl": p_ilvl,
            "has_prefix": has_prefix,
            "prefix": prefix_literal,
        })

    styles_status = _scan_styles(styles_xml)

    # numIds used by heading styles
    heading_num_ids = sorted({
        v.get("numId") for v in styles_status.values() if v.get("numId")
    } | {"4"})  # always inspect numId=4

    numbering_status = _scan_numbering(numbering_xml, list(heading_num_ids)) if numbering_xml else {}

    # styles_numpr_status flat
    flat_styles_status = {}
    for k, v in styles_status.items():
        flat_styles_status[v["name"]] = v["status"]

    return {
        "docx": str(docx_path),
        "total_h_paragraphs": h_with_prefix + h_without_prefix,
        "h_count_by_level": h_count_by_level,
        "h_with_prefix": h_with_prefix,
        "h_without_prefix": h_without_prefix,
        "h_with_para_numpr": h_with_para_numpr,
        "styles_numpr_status": flat_styles_status,
        "numbering_lvltext": {
            num_id: v.get("lvltexts", {}) for num_id, v in numbering_status.items()
        },
        "samples_no_prefix": samples_no_prefix,
        "samples_with_prefix": samples_with_prefix,
        "h_details": h_details,
    }


def main_heading_numbers():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", type=Path)
    _cc.add_report_flag(ap)
    args = ap.parse_args()
    if not args.docx.exists():
        print(f"ERR: {args.docx} not found", file=sys.stderr)
        sys.exit(2)
    result = audit_heading_numbers(args.docx)
    _cc.write_report(result, args.report, mkdir=False)
    # summary to stdout
    summary = {k: v for k, v in result.items() if k != "h_details"}
    _cc.print_json(summary)


# ---------------- pipeline adapter ----------------
def apply_path_heading_numbers(docx_path, args=None) -> dict:
    return audit_heading_numbers(Path(docx_path))


# ══════════ images ← audit_images.py ══════════

# rels 文件可能在 word/_rels/ 下任意 *.xml.rels (document / header* / footer* /
# footnotes / endnotes / comments / numbering ...). 主 document 引用走
# document.xml.rels, 但 header/footer/footnotes 等各自带 rels 文件且各自的
# rId 命名空间独立 (rId1 在 document.xml.rels vs header1.xml.rels 可指不同 target).
_RELS_RE = re.compile(r"^word/_rels/(.+)\.xml\.rels$")
# 主 part XML (排除 _rels/ / theme/ / settings 等纯样式), 凡可能 embed image 的:
_PART_XML_RE = re.compile(
    r"^word/(document|header\d*|footer\d*|footnotes|endnotes|comments|numbering)\.xml$"
)

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "o": "urn:schemas-microsoft-com:office:office",
    "rels": "http://schemas.openxmlformats.org/package/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def list_media(zf: zipfile.ZipFile) -> list[dict]:
    media = []
    for info in zf.infolist():
        if info.filename.startswith("word/media/") and not info.is_dir():
            name = info.filename[len("word/media/") :]
            if not name:
                continue
            ext = os.path.splitext(name)[1].lstrip(".").lower()
            media.append(
                {
                    "name": name,
                    "full_path": info.filename,
                    "size": info.file_size,
                    "ext": ext,
                    "referenced_by": [],
                }
            )
    return media


def parse_rels(zf: zipfile.ZipFile) -> dict[str, dict]:
    """[legacy] 返回 document.xml.rels 的 rId -> {target, type} (向后兼容).

    新代码应该用 parse_all_rels(), 该函数只覆盖主 document 的 rels, 不含
    header/footer/footnotes 等 part 各自的 rels. 保留只为兼容已有调用者.
    """
    try:
        data = zf.read("word/_rels/document.xml.rels")
    except KeyError:
        return {}
    tree = etree.fromstring(data)
    out = {}
    for rel in tree.findall("rels:Relationship", NS):
        rid = rel.get("Id")
        target = rel.get("Target") or ""
        rtype = rel.get("Type") or ""
        out[rid] = {"target": target, "type": rtype}
    return out


def parse_all_rels(zf: zipfile.ZipFile) -> dict[str, dict[str, dict]]:
    """扫所有 word/_rels/*.xml.rels, 返回 {part_name -> {rId -> {target, type}}}.

    part_name = rels 所属的 part 的 stem, 例: 'document', 'header1', 'footer3'.
    每个 part 的 rId 命名空间独立, 必须按 part 隔离查询.

    复用 strip_orphan_media._collect_referenced_media 的扫法 (但保留 rId 维度).
    """
    parser = etree.XMLParser(remove_blank_text=False, recover=True)
    out: dict[str, dict[str, dict]] = {}
    for name in zf.namelist():
        m = _RELS_RE.match(name)
        if not m:
            continue
        part_name = m.group(1)  # 'document' / 'header1' / 'footer3' ...
        try:
            data = zf.read(name)
            tree = etree.fromstring(data, parser=parser)
        except (etree.XMLSyntaxError, KeyError):
            continue
        if tree is None:
            continue
        part_rels: dict[str, dict] = {}
        for rel in tree.findall("rels:Relationship", NS):
            rid = rel.get("Id")
            target = rel.get("Target") or ""
            rtype = rel.get("Type") or ""
            part_rels[rid] = {"target": target, "type": rtype}
        out[part_name] = part_rels
    return out


def collect_rid_refs_in_part(zf: zipfile.ZipFile, part_path: str) -> list[dict]:
    """扫一个 part XML 的 body, 返回 [{rid, kind}] (kind=drawing/pict/chart/diagram).

    drawing/pict 都查 r:embed / r:link / r:id 各种姿势.
    用于扫 header*/footer*/footnotes/endnotes/comments 的 body, 把它们
    引用的 rId 也算 referenced, 避免 audit-images 漏报这些 part 引用的 media.
    """
    refs: list[dict] = []
    try:
        data = zf.read(part_path)
    except KeyError:
        return refs
    try:
        root = etree.fromstring(data, parser=etree.XMLParser(recover=True))
    except etree.XMLSyntaxError:
        return refs
    if root is None:
        return refs
    R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

    # DrawingML blip / chart / diagram
    for blip in root.findall(".//a:blip", NS):
        rid = blip.get(f"{R}embed") or blip.get(f"{R}link")
        if rid:
            refs.append({"rid": rid, "kind": "drawing"})
    for chart in root.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/chart}chart"
    ):
        rid = chart.get(f"{R}id")
        if rid:
            refs.append({"rid": rid, "kind": "chart"})
    for dgm in root.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/diagram}relIds"
    ):
        rid = dgm.get(f"{R}dm")
        if rid:
            refs.append({"rid": rid, "kind": "diagram"})

    # legacy VML pict / imagedata
    for imgdata in root.findall(".//v:imagedata", NS):
        rid = imgdata.get(f"{R}id") or imgdata.get(f"{R}href")
        if rid:
            refs.append({"rid": rid, "kind": "pict"})

    # OLE object (oleObject r:id 也算引用, 走 embeddings 路径不算 media 但要登记)
    return refs


def snippet(text: str, n: int = 60) -> str:
    text = (text or "").strip().replace("\n", " ")
    if len(text) <= n:
        return text
    return text[:n] + "…"


def scan_body(doc) -> tuple[list[dict], list[dict], list[dict]]:
    """扫所有段, 返回 (drawings, picts, paragraph_index_with_text)"""
    drawings = []
    picts = []
    paragraphs = []

    body = doc.element.body
    para_idx = -1
    for child in body.iterchildren():
        tag = etree.QName(child).localname
        if tag == "p":
            para_idx += 1
            text = "".join(child.itertext())
            paragraphs.append({"idx": para_idx, "text": text})

            # drawings (DrawingML — image blip / chart / diagram)
            for d in child.findall(".//w:drawing", NS):
                rid = None
                subtype = "image"
                blip = d.find(".//a:blip", NS)
                if blip is not None:
                    rid = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    ) or blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link"
                    )
                else:
                    # chart: <c:chart r:id="rIdN">
                    chart = d.find(
                        ".//{http://schemas.openxmlformats.org/drawingml/2006/chart}chart"
                    )
                    if chart is not None:
                        rid = chart.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                        )
                        subtype = "chart"
                    else:
                        # diagram: <dgm:relIds r:dm="rIdN">
                        dgm = d.find(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/diagram}relIds"
                        )
                        if dgm is not None:
                            rid = dgm.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}dm"
                            )
                            subtype = "diagram"
                drawings.append(
                    {
                        "para_idx": para_idx,
                        "type": "drawing",
                        "subtype": subtype,
                        "rid": rid,
                        "para_text_snippet": snippet(text),
                    }
                )

            # legacy VML pict
            for p in child.findall(".//w:pict", NS):
                rid = None
                imgdata = p.find(".//v:imagedata", NS)
                if imgdata is not None:
                    rid = imgdata.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                    ) or imgdata.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}href"
                    )
                picts.append(
                    {
                        "para_idx": para_idx,
                        "type": "pict",
                        "rid": rid,
                        "para_text_snippet": snippet(text),
                    }
                )

        elif tag == "tbl":
            # 表格内段不递增主 idx, 但需扫 drawing/pict
            for tp in child.findall(".//w:p", NS):
                text = "".join(tp.itertext())
                for d in tp.findall(".//w:drawing", NS):
                    rid = None
                    blip = d.find(".//a:blip", NS)
                    if blip is not None:
                        rid = blip.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                    drawings.append(
                        {
                            "para_idx": para_idx,  # 表前最后一段 idx, 标识 in-table
                            "in_table": True,
                            "type": "drawing",
                            "rid": rid,
                            "para_text_snippet": snippet(text),
                        }
                    )
                for p in tp.findall(".//w:pict", NS):
                    rid = None
                    imgdata = p.find(".//v:imagedata", NS)
                    if imgdata is not None:
                        rid = imgdata.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                        )
                    picts.append(
                        {
                            "para_idx": para_idx,
                            "in_table": True,
                            "type": "pict",
                            "rid": rid,
                            "para_text_snippet": snippet(text),
                        }
                    )
    return drawings, picts, paragraphs


def resolve_target(rel_target: str) -> str:
    """rels Target 是相对 word/ 的, 拼成 zip 路径"""
    if not rel_target:
        return ""
    t = rel_target.lstrip("/")
    if t.startswith("media/"):
        return "word/" + t
    return "word/" + t


def audit_images(docx_path: Path) -> dict:
    with zipfile.ZipFile(docx_path, "r") as zf:
        zip_names = set(zf.namelist())
        media = list_media(zf)
        rels = parse_rels(zf)  # legacy: 只 document.xml.rels
        all_rels = parse_all_rels(zf)  # 全 part rels (含 header/footer/footnotes/...)
        # 扫所有 part XML body 收集 rId 引用 (含 header/footer/footnotes/endnotes/comments/numbering)
        # 这是修正旧版漏扫的关键: 旧版只标 document.xml body 的 drawing/pict rId,
        # 导致 header/footer 里 imagedata 引用的 media 全被误判 orphan.
        part_refs: dict[str, list[dict]] = {}  # part_name -> [{rid, kind}]
        for zname in zf.namelist():
            pm = _PART_XML_RE.match(zname)
            if not pm:
                continue
            part_name = pm.group(1)
            part_refs[part_name] = collect_rid_refs_in_part(zf, zname)

    doc = Document(str(docx_path))
    drawings, picts, paragraphs = scan_body(doc)

    media_by_name = {m["name"]: m for m in media}

    def annotate(items: list[dict]) -> None:
        """标注 document.xml body 内 drawing/pict 的 status + 落 referenced_by.

        rId 解析走 document part 的 rels (`all_rels['document']`),
        与旧 `rels` 字典等价.
        """
        doc_rels = all_rels.get("document", rels)
        for it in items:
            rid = it.get("rid")
            if not rid:
                it["status"] = "no-rid"
                it["target"] = None
                continue
            rel = doc_rels.get(rid)
            if rel is None:
                it["status"] = "dangling-no-rel"
                it["target"] = None
                continue
            target = rel["target"]
            it["target"] = target
            full = resolve_target(target)
            if full not in zip_names:
                it["status"] = "dangling-target-missing"
                continue
            # mark media referenced (带来源标注: "document:rIdN")
            mname = os.path.basename(target)
            if mname in media_by_name:
                media_by_name[mname]["referenced_by"].append(f"document:{rid}")
            it["status"] = "ok"

    annotate(drawings)
    annotate(picts)

    # 全 part 扫描: 把 header/footer/footnotes/... 引用的 media 也标 referenced.
    # 这是消除磐安 .bak-1 47 个 .wmf 误报为 orphan 的核心修复.
    for part_name, refs in part_refs.items():
        if part_name == "document":
            # document.xml body 已通过 annotate() 标过, 跳过避免重复 (但下面也兜底)
            pass
        part_rels = all_rels.get(part_name, {})
        for ref in refs:
            rid = ref["rid"]
            rel = part_rels.get(rid)
            if rel is None:
                continue
            target = rel["target"]
            full = resolve_target(target)
            if full not in zip_names:
                continue
            mname = os.path.basename(target)
            if mname in media_by_name:
                tag = f"{part_name}:{rid}"
                if tag not in media_by_name[mname]["referenced_by"]:
                    media_by_name[mname]["referenced_by"].append(tag)

    # 兜底: 严谨算法对账 — 若 rels 里 Target 直接指向 word/media/<name>,
    # 也算 referenced (与 strip_orphan_media.scan_orphans 语义对齐).
    # 这覆盖某些非常规 part 引用 / 模板残留 rels 仍真指 media 的边界情况.
    for part_name, part_rels in all_rels.items():
        for rid, rel in part_rels.items():
            target = (rel.get("target") or "").replace("\\", "/").lstrip("/")
            if not target:
                continue
            if target.startswith("media/") or "/media/" in target:
                mname = os.path.basename(target)
                if mname in media_by_name:
                    tag = f"rels[{part_name}]:{rid}"
                    if not any(
                        t.startswith(f"{part_name}:") or t == tag
                        for t in media_by_name[mname]["referenced_by"]
                    ):
                        media_by_name[mname]["referenced_by"].append(tag)

    orphan_media = [m["name"] for m in media if not m["referenced_by"]]
    dangling_rids = [
        {"type": x["type"], "para_idx": x["para_idx"], "rid": x["rid"], "status": x["status"]}
        for x in drawings + picts
        if x["status"] != "ok"
    ]

    # anchor 段: 找含"图"/"示意"/"布局"/"分布" 等图标题关键词的段, 列其相邻 drawing
    anchor_keywords = ["图", "示意", "布局", "分布", "structure", "map"]
    anchors = []
    drawing_by_para: dict[int, list[int]] = {}
    for i, d in enumerate(drawings):
        drawing_by_para.setdefault(d["para_idx"], []).append(i)
    for p in paragraphs:
        text = p["text"].strip()
        if not text:
            continue
        if any(k in text for k in anchor_keywords) and len(text) < 80:
            # 找紧邻的 drawing (该段或下一段)
            next_d = None
            for delta in (0, 1, 2, -1):
                cand = drawing_by_para.get(p["idx"] + delta, [])
                if cand:
                    next_d = {
                        "delta": delta,
                        "drawing_idx_in_list": cand[0],
                        "status": drawings[cand[0]]["status"],
                    }
                    break
            anchors.append(
                {
                    "para_idx": p["idx"],
                    "text": snippet(text, 80),
                    "nearest_drawing": next_d,
                }
            )

    summary = {
        "media_files_count": len(media),
        "drawings_count": len(drawings),
        "picts_count": len(picts),
        "orphan_media_count": len(orphan_media),
        "orphan_media": orphan_media,
        "dangling_rids_count": len(dangling_rids),
        "dangling_rids": dangling_rids,
        "issues_count": len(orphan_media) + len(dangling_rids),
        # 修复 2026-05-26: 列出已扫的 rels parts (含 header/footer/footnotes/...),
        # 验证 orphan_media_count 与 strip_orphan_media 对齐.
        "rels_parts_scanned": sorted(all_rels.keys()),
        "part_xmls_scanned": sorted(part_refs.keys()),
    }

    return {
        "docx_path": str(docx_path),
        "summary": summary,
        "media_files": media,
        "drawings": drawings,
        "picts": picts,
        "anchor_paragraphs": anchors,
    }


def main_images() -> int:
    ap = argparse.ArgumentParser(description="Audit docx images (audit-only).")
    ap.add_argument("docx_path", type=Path)
    _cc.add_report_flag(ap)
    args = ap.parse_args()

    if not args.docx_path.exists():
        print(f"ERR: not found: {args.docx_path}", file=sys.stderr)
        return 2

    report_path = args.report or Path(f"/tmp/audit-images-{args.docx_path.stem}.json")
    result = audit_images(args.docx_path)
    _cc.write_report(result, report_path, mkdir=False)
    s = result["summary"]
    print(
        f"audit done -> {report_path}\n"
        f"  media={s['media_files_count']}  drawings={s['drawings_count']}  "
        f"picts={s['picts_count']}  orphan_media={s['orphan_media_count']}  "
        f"dangling={s['dangling_rids_count']}  issues={s['issues_count']}"
    )
    return 0


# ---------------- pipeline adapter ----------------
def apply_path_images(docx_path, args=None) -> dict:
    """pipeline read-only: 走 zip 路径 audit"""
    return audit_images(Path(docx_path))


# ══════════ table-pairing ← audit_table_pairing.py ══════════

_DOC_TABLE_PAIRING = """audit_table_pairing.py — 只读 audit docx 中"表名段 ↔ <w:tbl>"配对状态.

检测 5 类问题:
  1. caption-name-content-mismatch  表名内容关键词与紧邻 tbl 首行列名语义不匹配
  2. orphan-caption-no-downstream-tbl  表名段下游 8 段内无 tbl
  3. orphan-tbl-no-upstream-caption    tbl 上游 8 段内无表名
  4. duplicate-caption-name            两个或多个表名段名字完全相同 (合并冲突)
  5. two-captions-compete-same-tbl     两个表名段都紧邻同一 tbl

输出 audit JSON 让人 (主会话/用户) 拍板 decision, 再喂给 caption.py pair 改.

接口:
  python3 audit_table_pairing.py <docx> [--report <json>] [--quiet]

默认 stdout 打印 summary + issues; --report 写完整 JSON.
"""

W_P = f"{{{W_NS}}}p"
W_TBL = f"{{{W_NS}}}tbl"
W_T = f"{{{W_NS}}}t"
W_PPR = f"{{{W_NS}}}pPr"
W_PSTYLE = f"{{{W_NS}}}pStyle"
W_VAL = f"{{{W_NS}}}val"
W_TR = f"{{{W_NS}}}tr"
W_TC = f"{{{W_NS}}}tc"

# 兼容两种表号: 扁平 "表3-1" 与中文章节式 "表3.1-1" (章.节-序, /docx renumber --cn-section 产出)。
# group(1)=章节号(3 或 3.1), group(2)=序号, group(3)=表名。向后兼容: (?:\.\d+)* 可选, 扁平号仍匹配。
CAP_PATTERN = re.compile(r"^\s*表\s*(\d+(?:\.\d+)*)\s*[-–—]\s*(\d+)\s*(.*)$")
# 附录式表号 "附表N 表名" (院模板结论章/附录常用, 单号无短横)。与 CAP_PATTERN 互斥:
# 仅当 CAP_PATTERN 不匹配时回退此式, 把附表识别为合法 caption (否则真附表被误判孤儿表)。
CAP_APPENDIX_PATTERN = re.compile(r"^\s*附表\s*(\d+)\s*(.*)$")

# 关键词→同义词字典 (字面包含即视为命中)
KEYWORD_SYNONYMS: dict[str, list[str]] = {
    "成本": ["成本", "占比", "费用", "造价", "运营成本"],
    "价格": ["价格", "水价", "价差", "标准", "元/m³", "调价", "调整"],
    "目标": ["目标", "指标", "阶段", "任务"],
    "职责": ["职责", "部门", "分工", "主导"],
    "原则": ["原则", "措施", "核心要求"],
    "对比": ["对比", "占比", "嘉兴", "浙江省", "全省"],
    "格局": ["水源类型", "供水", "占比", "水质", "等级"],
    "政策": ["层级", "政策", "文件", "时间", "文号", "要求"],
    "实践": ["城市", "模式", "规模", "利用率", "置换"],
    "必要性": ["维度", "问题", "效果", "案例"],
    "前置": ["环节", "嵌入", "管控", "约束", "前置"],
    "差异化": ["用户类别", "水源类型", "价格", "价差", "适用"],
    "保障": ["保障类型", "措施", "实施", "依据", "主体"],
    "联动": ["阶段", "时间", "目标", "指标", "效益", "任务"],
    "征收": ["征收", "监督", "标准", "主体"],
}


def get_style_id(p) -> str:
    ppr = p.find(W_PPR)
    if ppr is None:
        return ""
    ps = ppr.find(W_PSTYLE)
    if ps is None:
        return ""
    return ps.get(W_VAL) or ""


def get_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter(W_T)).strip()


def get_first_row_cells(tbl) -> list[str]:
    cells = []
    for tr in tbl.iter(W_TR):
        for tc in tr.iter(W_TC):
            txt = get_text(tc)
            cells.append(txt[:40])
        break
    return cells


def count_rows_cols(tbl) -> tuple[int, int]:
    rows = list(tbl.iter(W_TR))
    if not rows:
        return 0, 0
    cols = len(list(rows[0].iter(W_TC)))
    return len(rows), cols


def heuristic_match(caption_name: str, tbl_first_row: list[str]) -> tuple[float, str]:
    """对 caption 名字关键词与 tbl 首行 cell 命中度打分.

    返回 (score, method).
      score = 命中关键词数 / 名字关键词总数 (0.0 .. 1.0)
      method = "kw-match:k1+k2..." 或 "no-match"
    """
    if not caption_name:
        return 0.0, "no-caption-name"
    first_row_text = "|".join(tbl_first_row)
    hits = []
    total_kws = 0
    for kw, syns in KEYWORD_SYNONYMS.items():
        if kw in caption_name:
            total_kws += 1
            if any(s in first_row_text for s in syns):
                hits.append(kw)
    if total_kws == 0:
        return 0.0, "no-keyword-extracted"
    score = len(hits) / total_kws
    if hits:
        return score, f"kw-match:{'+'.join(hits)}"
    return 0.0, "no-match"


def _audit_from_doc_table_pairing(doc, docx_path_label: str = "") -> dict:
    elems = list(doc.element.body.iterchildren())

    # 收集 captions
    captions = []
    cap_counter = 0
    for i, e in enumerate(elems):
        if e.tag == W_P:
            text = get_text(e)
            m = CAP_PATTERN.match(text)
            ma = CAP_APPENDIX_PATTERN.match(text) if not m else None
            if m or ma:
                cap_counter += 1
                if m:
                    ch, num, name = m.group(1), m.group(2), m.group(3).strip()
                    number = f"表{ch}-{num}"
                else:  # 附录式 "附表N"
                    ch, num, name = "附", ma.group(1), ma.group(2).strip()
                    number = f"附表{num}"
                # 找紧邻下游"注:..."段 (style ZDWP, 文本以"注"开头)
                notes_idx = []
                j = i + 1
                while j < len(elems) and elems[j].tag == W_P:
                    nt = get_text(elems[j])
                    if nt.startswith(("注:", "注：", "注", "Note", "note")) and len(nt) < 200:
                        notes_idx.append(j)
                        j += 1
                    else:
                        break
                captions.append({
                    "id": f"cap-{cap_counter}",
                    "elem_idx": i,
                    "number": number,
                    "chapter": ch,  # 字符串: 兼容 cn-section "3.1" (旧扁平 "3" 亦为字符串)；附录式 ch="附"
                    "seq": int(num),
                    "name": name,
                    "style": get_style_id(e),
                    "raw_text": text,
                    "notes_idx": notes_idx,
                })

    # 收集 tbls
    tbls = []
    tbl_counter = 0
    for i, e in enumerate(elems):
        if e.tag == W_TBL:
            tbl_counter += 1
            cells = get_first_row_cells(e)
            rc, cc = count_rows_cols(e)
            tbls.append({
                "id": f"tbl-{tbl_counter}",
                "elem_idx": i,
                "first_row_cells": cells,
                "row_count": rc,
                "col_count": cc,
            })

    # 为每个 caption 算 nearest_downstream_tbl + heuristic_match
    for cap in captions:
        nxt = None
        for t in tbls:
            if t["elem_idx"] > cap["elem_idx"]:
                nxt = t
                break
        if nxt:
            cap["nearest_downstream_tbl"] = nxt["id"]
            cap["nearest_downstream_tbl_idx"] = nxt["elem_idx"]
            cap["distance"] = nxt["elem_idx"] - cap["elem_idx"]
            # 启发匹配下游 5 个候选 tbl 中得分最高的
            candidates = [t for t in tbls
                          if cap["elem_idx"] < t["elem_idx"] <= cap["elem_idx"] + 30]
            best_score, best_method, best_tid = 0.0, "no-match", nxt["id"]
            for c in candidates:
                s, m = heuristic_match(cap["name"], c["first_row_cells"])
                if s > best_score:
                    best_score, best_method, best_tid = s, m, c["id"]
            cap["heuristic_match"] = {
                "tbl_id": best_tid, "score": round(best_score, 3), "method": best_method,
            }
        else:
            cap["nearest_downstream_tbl"] = None
            cap["distance"] = None
            cap["heuristic_match"] = {"tbl_id": None, "score": 0.0, "method": "no-downstream"}

    # 为每个 tbl 算上游 8 段内的 captions
    for tbl in tbls:
        ups = [c["id"] for c in captions
               if tbl["elem_idx"] - 8 <= c["elem_idx"] < tbl["elem_idx"]]
        tbl["upstream_captions_within_8_paras"] = ups

    # 检测 issues
    issues = []

    # 1. orphan-caption-no-downstream-tbl  (孤儿表名: 与紧邻下游 tbl 距离 > 5)
    for cap in captions:
        if cap.get("distance") is None or cap["distance"] > 5:
            issues.append({
                "type": "orphan-caption-no-downstream-tbl",
                "caption_id": cap["id"],
                "caption_number": cap["number"],
                "caption_name": cap["name"],
                "elem_idx": cap["elem_idx"],
                "details": f"下游最近 tbl 距离 = {cap.get('distance')}",
            })

    # 2. orphan-tbl-no-upstream-caption (孤儿表: 上游 8 段无表名)
    for tbl in tbls:
        if not tbl["upstream_captions_within_8_paras"]:
            issues.append({
                "type": "orphan-tbl-no-upstream-caption",
                "tbl_id": tbl["id"],
                "elem_idx": tbl["elem_idx"],
                "first_row": tbl["first_row_cells"],
                "details": "上游 8 段内无表名段",
            })

    # 3. duplicate-caption-name (重名: 名字完全相同, 名字非空)
    from collections import defaultdict
    by_name = defaultdict(list)
    for cap in captions:
        if cap["name"]:
            by_name[cap["name"]].append(cap)
    for name, lst in by_name.items():
        if len(lst) > 1:
            issues.append({
                "type": "duplicate-caption-name",
                "name": name,
                "caption_ids": [c["id"] for c in lst],
                "details": f"{len(lst)} 个表名同名: " + ", ".join(
                    f"{c['id']}({c['number']})" for c in lst
                ),
            })

    # 4. two-captions-compete-same-tbl (两表名抢同一 tbl)
    by_tbl = defaultdict(list)
    for tbl in tbls:
        for cid in tbl["upstream_captions_within_8_paras"]:
            by_tbl[tbl["id"]].append(cid)
    for tid, caps in by_tbl.items():
        if len(caps) >= 2:
            issues.append({
                "type": "two-captions-compete-same-tbl",
                "tbl_id": tid,
                "competing_captions": caps,
                "details": f"tbl {tid} 上游 8 段内同时存在 {len(caps)} 个表名: " + ", ".join(caps),
            })

    # 5. caption-name-content-mismatch (启发分 = 0 且名字含已知关键词)
    for cap in captions:
        if not cap["name"]:
            continue
        hm = cap["heuristic_match"]
        if hm["score"] == 0.0 and hm["method"] == "no-match" and cap.get("distance") and cap["distance"] <= 5:
            # 名字含已知关键词但 tbl 首行不命中
            issues.append({
                "type": "caption-name-content-mismatch",
                "caption_id": cap["id"],
                "caption_number": cap["number"],
                "caption_name": cap["name"],
                "tbl_id": hm["tbl_id"],
                "details": f"表名关键词与下游 tbl 首行 cell 字面无交集",
            })

    # 6. empty-caption-name
    for cap in captions:
        if not cap["name"]:
            issues.append({
                "type": "empty-caption-name",
                "caption_id": cap["id"],
                "caption_number": cap["number"],
                "elem_idx": cap["elem_idx"],
                "details": "表名段编号后无标题文字",
            })

    summary = {
        "captions": len(captions),
        "tbls": len(tbls),
        "orphan_captions": sum(1 for i in issues if i["type"] == "orphan-caption-no-downstream-tbl"),
        "orphan_tbls": sum(1 for i in issues if i["type"] == "orphan-tbl-no-upstream-caption"),
        "duplicate_caption_names": sum(1 for i in issues if i["type"] == "duplicate-caption-name"),
        "competing_pairs": sum(1 for i in issues if i["type"] == "two-captions-compete-same-tbl"),
        "content_mismatches": sum(1 for i in issues if i["type"] == "caption-name-content-mismatch"),
        "empty_names": sum(1 for i in issues if i["type"] == "empty-caption-name"),
    }

    return {
        "docx_path": docx_path_label,
        "summary": summary,
        "captions": captions,
        "tbls": tbls,
        "issues": issues,
    }


def audit_table_pairing(docx_path: Path) -> dict:
    doc = Document(str(docx_path))
    return _audit_from_doc_table_pairing(doc, str(docx_path))


def apply_table_pairing(doc, args=None) -> dict:
    """pipeline read-only adapter"""
    label = str(getattr(args, "docx", "")) if args else ""
    return _audit_from_doc_table_pairing(doc, label)


def print_summary(audit_data: dict) -> None:
    s = audit_data["summary"]
    print(f"\n{'='*78}")
    print(f"audit: {audit_data['docx_path']}")
    print(f"{'='*78}")
    print(f"captions = {s['captions']} | tbls = {s['tbls']} | "
          f"差 = {s['captions'] - s['tbls']} (孤儿表名数)")
    print(f"orphan_captions  = {s['orphan_captions']}")
    print(f"orphan_tbls      = {s['orphan_tbls']}")
    print(f"duplicate_names  = {s['duplicate_caption_names']}")
    print(f"competing_pairs  = {s['competing_pairs']}")
    print(f"content_mismatch = {s['content_mismatches']}")
    print(f"empty_names      = {s['empty_names']}")
    print(f"\n{'─'*78}")
    print("issues:")
    for issue in audit_data["issues"]:
        head = f"  [{issue['type']}]"
        rest = " ".join(f"{k}={v}" for k, v in issue.items() if k != "type")
        print(f"{head} {rest}")


def main_table_pairing(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=_DOC_TABLE_PAIRING.split("\n")[0])
    ap.add_argument("docx", help="输入 docx 路径")
    _cc.add_report_flag(ap, help="完整 audit JSON 输出路径")
    ap.add_argument("--quiet", action="store_true", help="不打印 summary stdout")
    args = ap.parse_args(argv)

    src = Path(args.docx)
    if not src.exists():
        print(f"[error] 找不到 {src}", file=sys.stderr)
        return 2

    audit_data = audit_table_pairing(src)
    if not args.quiet:
        print_summary(audit_data)
    _cc.write_report(audit_data, args.report, announce="\n[report] {path}")
    return 0


# ══════════ fields ← audit_word_fields.py ══════════

def _localname(tag: str) -> str:
    return tag.split("}", 1)[-1] if "}" in tag else tag


def _iter_paragraph_idx(root: etree._Element) -> dict[etree._Element, int]:
    """给文档每个 <w:p> 标注一个 idx(物理顺序)。"""
    idx_map: dict[etree._Element, int] = {}
    for i, p in enumerate(root.iter(f"{W}p")):
        idx_map[p] = i
    return idx_map


def _ancestor_para_idx(elem: etree._Element, para_idx_map: dict[etree._Element, int]) -> int:
    cur = elem
    while cur is not None:
        if cur.tag == f"{W}p":
            return para_idx_map.get(cur, -1)
        cur = cur.getparent()
    return -1


def _classify_instr(instr: str) -> str:
    """从 instrText 文本里取字段类型(如 TOC / PAGEREF / SEQ / STYLEREF / REF / DATE / TIME / HYPERLINK)。
    Special:开头是 '=' 视作 '='(formula)。"""
    s = instr.strip()
    if not s:
        return "EMPTY"
    if s.startswith("="):
        return "="
    m = re.match(r"^([A-Z][A-Z0-9_]*)", s)
    if m:
        return m.group(1)
    return "OTHER"


def _parse_complex_fields(root: etree._Element, para_idx_map: dict[etree._Element, int]) -> list[dict]:
    """按文档物理顺序遍历 fldChar,依 begin/separate/end 配对组装字段。
    支持嵌套(begin 栈),返回每个字段记录。"""
    fields: list[dict] = []
    stack: list[dict] = []  # 栈,每个元素 = 一个未闭合 field 的临时状态

    for elem in root.iter():
        tag = _localname(elem.tag)
        if tag == "fldChar":
            ftype = elem.get(f"{W}fldCharType")
            if ftype == "begin":
                stack.append({
                    "instr_parts": [],
                    "result_parts": [],
                    "phase": "instr",  # instr -> result(after separate)
                    "para_idx": _ancestor_para_idx(elem, para_idx_map),
                    "depth": len(stack),  # 0 = top-level
                })
            elif ftype == "separate":
                if stack:
                    stack[-1]["phase"] = "result"
            elif ftype == "end":
                if stack:
                    rec = stack.pop()
                    instr = "".join(rec["instr_parts"])
                    result = "".join(rec["result_parts"])
                    fields.append({
                        "kind": "complex",
                        "type": _classify_instr(instr),
                        "instr": instr.strip(),
                        "result_sample": result.strip()[:200],
                        "para_idx": rec["para_idx"],
                        "depth": rec["depth"],
                    })
        elif tag == "instrText":
            if stack:
                stack[-1]["instr_parts"].append(elem.text or "")
        elif tag == "t":
            # plain text — 在 result phase 时,把外层 field 的 result 累计
            if stack and stack[-1]["phase"] == "result":
                # 只算最内层的 result(避免外层重复累计 inner field 文本)
                stack[-1]["result_parts"].append(elem.text or "")

    # 处理未闭合(异常)
    while stack:
        rec = stack.pop()
        instr = "".join(rec["instr_parts"])
        fields.append({
            "kind": "complex-unclosed",
            "type": _classify_instr(instr),
            "instr": instr.strip(),
            "result_sample": "",
            "para_idx": rec["para_idx"],
            "depth": rec["depth"],
        })
    return fields


def _parse_simple_fields(root: etree._Element, para_idx_map: dict[etree._Element, int]) -> list[dict]:
    fields: list[dict] = []
    for elem in root.iter(f"{W}fldSimple"):
        instr = elem.get(f"{W}instr", "") or ""
        # result = 内部所有 w:t 拼接
        result = "".join((t.text or "") for t in elem.iter(f"{W}t"))
        fields.append({
            "kind": "simple",
            "type": _classify_instr(instr),
            "instr": instr.strip(),
            "result_sample": result.strip()[:200],
            "para_idx": _ancestor_para_idx(elem, para_idx_map),
            "depth": 0,
        })
    return fields


def _count_artifacts(root: etree._Element) -> dict[str, int]:
    return {
        "fldChar_count": sum(1 for _ in root.iter(f"{W}fldChar")),
        "instrText_count": sum(1 for _ in root.iter(f"{W}instrText")),
        "fldSimple_count": sum(1 for _ in root.iter(f"{W}fldSimple")),
    }


def _scan_one_xml(xml_bytes: bytes, xml_name: str) -> dict[str, Any]:
    parser = etree.XMLParser(huge_tree=True, recover=False)
    root = etree.fromstring(xml_bytes, parser=parser)
    para_idx_map = _iter_paragraph_idx(root)
    counts = _count_artifacts(root)
    complex_fields = _parse_complex_fields(root, para_idx_map)
    simple_fields = _parse_simple_fields(root, para_idx_map)
    all_fields = complex_fields + simple_fields

    type_dist: dict[str, int] = {}
    for f in all_fields:
        type_dist[f["type"]] = type_dist.get(f["type"], 0) + 1

    depths = [f.get("depth", 0) for f in complex_fields]
    return {
        "xml_name": xml_name,
        **counts,
        "type_distribution": type_dist,
        "nested_depth_max": (max(depths) if depths else 0),
        "field_total": len(all_fields),
        "fields": all_fields,
    }


def audit_word_fields(docx_path: Path, include_headers: bool) -> dict[str, Any]:
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)
    targets: list[str] = ["word/document.xml"]
    files_scanned: list[str] = []
    aggregate = {
        "fldChar_count": 0,
        "instrText_count": 0,
        "fldSimple_count": 0,
        "type_distribution": {},
        "nested_depth_max": 0,
        "field_total": 0,
        "fields": [],
        "per_file": [],
    }
    with zipfile.ZipFile(docx_path, "r") as z:
        names = set(z.namelist())
        if include_headers:
            for n in names:
                if n.startswith("word/header") and n.endswith(".xml"):
                    targets.append(n)
                if n.startswith("word/footer") and n.endswith(".xml"):
                    targets.append(n)
        for name in targets:
            if name not in names:
                continue
            data = z.read(name)
            try:
                res = _scan_one_xml(data, name)
            except etree.XMLSyntaxError as e:
                aggregate["per_file"].append({"xml_name": name, "error": str(e)})
                continue
            files_scanned.append(name)
            aggregate["fldChar_count"] += res["fldChar_count"]
            aggregate["instrText_count"] += res["instrText_count"]
            aggregate["fldSimple_count"] += res["fldSimple_count"]
            aggregate["field_total"] += res["field_total"]
            aggregate["nested_depth_max"] = max(aggregate["nested_depth_max"], res["nested_depth_max"])
            for k, v in res["type_distribution"].items():
                aggregate["type_distribution"][k] = aggregate["type_distribution"].get(k, 0) + v
            aggregate["fields"].extend(res["fields"])
            aggregate["per_file"].append({
                "xml_name": name,
                "fldChar_count": res["fldChar_count"],
                "instrText_count": res["instrText_count"],
                "fldSimple_count": res["fldSimple_count"],
                "field_total": res["field_total"],
                "type_distribution": res["type_distribution"],
            })
    aggregate["files_scanned"] = files_scanned
    aggregate["docx"] = str(docx_path)
    return aggregate


def main_word_fields(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Audit Word field artifacts in a docx (read-only).")
    ap.add_argument("docx", type=Path)
    ap.add_argument("--include-headers", action="store_true",
                    help="Also scan word/header*.xml and word/footer*.xml")
    _cc.add_report_flag(ap, help="Write full JSON report to this path")
    args = ap.parse_args(argv)

    report = audit_word_fields(args.docx, include_headers=args.include_headers)

    # stdout 摘要 (短)
    summary = {
        "docx": report["docx"],
        "files_scanned": report["files_scanned"],
        "fldChar_count": report["fldChar_count"],
        "instrText_count": report["instrText_count"],
        "fldSimple_count": report["fldSimple_count"],
        "field_total": report["field_total"],
        "nested_depth_max": report["nested_depth_max"],
        "type_distribution": report["type_distribution"],
    }
    _cc.print_json(summary)

    # 全报告(含 fields 详情)
    _cc.write_report(report, args.report)
    return 0


# ---------------- pipeline adapter ----------------
def apply_path_word_fields(docx_path, args=None) -> dict:
    include_headers = bool(getattr(args, "include_headers", False)) if args else False
    return audit_word_fields(Path(docx_path), include_headers=include_headers)


# ──────────────────────────── 家族入口（子命令分发）────────────────────────────

SUBCOMMANDS = {
    "bookmarks": main_bookmarks,
    "captions": main_caption_outline,
    "headings": main_heading_numbers,
    "images": main_images,
    "table-pairing": main_table_pairing,
    "fields": main_word_fields,
}


def main(argv: list[str] | None = None) -> int:
    args = list(sys.argv[1:] if argv is None else argv)
    if not args or args[0] in ("-h", "--help"):
        print("usage: audit.py {" + ",".join(SUBCOMMANDS) + "} <docx> [flags…]\n"
              "每个子命令的参数与原独立脚本逐字一致：audit.py <sub> --help 查看。")
        return 0 if args else 2
    sub, rest = args[0], args[1:]
    fn = SUBCOMMANDS.get(sub)
    if fn is None:
        print(f"[audit] unknown subcommand: {sub!r}; choices={list(SUBCOMMANDS)}",
              file=sys.stderr)
        return 2
    saved = sys.argv[:]
    sys.argv = [sys.argv[0]] + rest
    try:
        rc = fn()
        return int(rc) if isinstance(rc, int) else 0
    finally:
        sys.argv = saved


if __name__ == "__main__":
    sys.exit(main())
