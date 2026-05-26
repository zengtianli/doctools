#!/usr/bin/env python3
"""reorder_heading_blocks.py — 检测+修 docx「标题块错位 heading-block-misorder」

算法:
  1. 切段块: heading 段 + 下属段(直到下一个同级或更高级 heading)
  2. 检测错位: 同父块下兄弟按 number tuple 比较,排序后不同 = 错位
  3. 检测重复: 同父下 number tuple 相同的块 → 保留 styled + 段数多 + 后出现
  4. 重排策略: 用 lxml 直接重排 <w:body> 下 <w:p>; 删重复块 = remove <w:p>
  5. 跨章节边界: H1 (章) 独立, 不跨章重排

接口:
  python3 reorder_heading_blocks.py <docx> [--dry-run] [--no-backup] [--report <json>]

默认 真改 (用户已批准). dry-run 只输出 plan.
"""
from __future__ import annotations

import argparse
import json
import shutil
import sys
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document

# 复用 fix_heading_disorder 的 number/style 解析,不重写
sys.path.insert(0, str(Path(__file__).parent))
from fix_heading_disorder import (  # noqa: E402
    detect_heading_form,
    is_heading_style,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_P = f"{{{W_NS}}}p"


# ---------- 段分类 ----------

def classify_paragraph(p, idx: int) -> dict:
    """返回段信息 dict: {idx, text, style_id, style_name, h_level, h_number}.

    h_level: 1-5 = heading 层级 (styled OR form-only); None = 非 heading.
    h_number: tuple 编号 (如 (3,1,2)) 或 None.

    判定优先:
      - styled Heading X → level = X, number 用 text 形态解析(如能)
      - 否则 detect_heading_form: 若 form depth ∈ {1..5} 且像 heading → level = depth
        但要小心: 普通正文段以数字开头(如"1. xxx 是 ..." 列表)别误判 → 仅当文本 ≤ 50 字算"像 heading"
    """
    text = (p.text or "").strip()
    sid = p.style.style_id if p.style is not None else None
    nm = p.style.name if p.style is not None else None

    styled_level = is_heading_style(sid, nm)
    form = detect_heading_form(text) if text else None

    h_level = None
    h_number = None

    if styled_level is not None:
        h_level = styled_level
        if form and form[2] is not None:
            h_number = form[2]
    elif form is not None and form[1] in (1, 2, 3, 4, 5) and form[2] is not None:
        # form-only heading: 文本短且形态明确 → 认作 heading
        # 但纯 "1 xxx" 单数字 H1_form 可能是正文列表项 → 仅在很短(≤30 字)时认
        label, depth, _ = form
        if label in ("H2_form", "H3_form", "H4_form", "chapter_H1"):
            if len(text) <= 80:
                h_level = depth
                h_number = form[2]
        # H1_form / H1_dot_form / Title_cn 不当 heading 处理(避免列表项误判)

    return {
        "idx": idx,
        "text": text,
        "style_id": sid,
        "style_name": nm,
        "h_level": h_level,
        "h_number": h_number,
        "p": p,
        "styled": styled_level is not None,
    }


# ---------- 块切分 ----------

def slice_blocks(paras: list[dict]) -> list[dict]:
    """切段块. 每个块 = heading + 下属直到下一个同级或更高级 heading.

    返回 list[block], block 含:
      level, heading_idx, heading_text, number, end_idx (exclusive),
      paragraph_indices, styled, parent_block_idx (即包含本块的更高级块的 heading_idx;
      若为 H1 块或无更高级, 则 -1).
    """
    headings = [pi for pi in paras if pi["h_level"] is not None]
    blocks: list[dict] = []
    # 用栈追踪当前每层 heading 的 heading_idx, 以确定 parent
    # stack[level] = heading_idx of currently open block at that level
    open_stack: dict[int, int] = {}
    for i, h in enumerate(headings):
        # 当前 heading 出现 → 关闭所有 level >= h.level 的 open block
        for lvl in list(open_stack.keys()):
            if lvl >= h["h_level"]:
                del open_stack[lvl]
        # parent = max level < h.level still open
        parent_levels = [lvl for lvl in open_stack if lvl < h["h_level"]]
        parent_idx = open_stack[max(parent_levels)] if parent_levels else -1

        # 下一个同级或更高级别 heading = level <= h.level
        end_idx = len(paras)
        for nh in headings[i + 1:]:
            if nh["h_level"] <= h["h_level"]:
                end_idx = nh["idx"]
                break
        blocks.append({
            "level": h["h_level"],
            "heading_idx": h["idx"],
            "heading_text": h["text"],
            "number": h["h_number"],
            "end_idx": end_idx,
            "paragraph_indices": list(range(h["idx"], end_idx)),
            "styled": h["styled"],
            "style_id": h["style_id"],
            "parent_block_idx": parent_idx,
        })
        open_stack[h["h_level"]] = h["idx"]

    # Rehome pass — 让每个 H{N} 块尽量找到 number 匹配的 L{N-1} 父块.
    # 用于修两种常见反模式:
    #   (a) H3 出现在其 L2 父之前 (3.1.2 物理在 3.1 之前) — parent fall-through 到 H1
    #   (b) H3 跨章漂浮 (在错误的 L2 之下) — number prefix 与 parent number 不匹配
    by_idx = {b["heading_idx"]: b for b in blocks}
    for b in blocks:
        if b["number"] is None or len(b["number"]) <= 1:
            continue
        want_prefix = b["number"][:-1]
        want_level = len(want_prefix)
        # 候选 = number 完全匹配 want_prefix 且 level 严格 = want_level 的块
        candidates = [c for c in blocks
                      if c["number"] == want_prefix
                      and c["level"] == want_level]
        if not candidates:
            continue
        cur_parent = by_idx.get(b["parent_block_idx"])
        if cur_parent and cur_parent["heading_idx"] in [c["heading_idx"] for c in candidates]:
            continue
        # 同 H1 范围内的 candidates 优先
        # 若 cur_parent 没有 → 用所有 candidates
        cur_h1 = _find_h1_ancestor(b, by_idx)
        same_h1_cands = [c for c in candidates
                         if _find_h1_ancestor(c, by_idx) == cur_h1]
        pool = same_h1_cands if same_h1_cands else candidates
        # 在 pool 内选 heading_idx <= b.heading_idx 且最大的 (本块前最近); 否则第一个
        before = [c for c in pool if c["heading_idx"] < b["heading_idx"]]
        if before:
            target = max(before, key=lambda x: x["heading_idx"])
        else:
            target = min(pool, key=lambda x: x["heading_idx"])
        if not _same_h1(b, target, by_idx):
            continue
        b["parent_block_idx"] = target["heading_idx"]
    return blocks


_RE_NUM_PREFIX = None
def _strip_number_prefix(text: str) -> str:
    """剥离 heading 文本的编号前缀, 如 '3.1.2 办公室工作机制' → '办公室工作机制'."""
    global _RE_NUM_PREFIX
    if _RE_NUM_PREFIX is None:
        import re
        _RE_NUM_PREFIX = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s+")
    m = _RE_NUM_PREFIX.match(text or "")
    if m:
        return (text or "")[m.end():].strip()
    return (text or "").strip()


def _texts_similar(texts: list[str]) -> bool:
    """判定一组文本是否"实质相同". 严格 — 完全相同或一方是另一方前缀(>= 4 字符)."""
    if not texts:
        return True
    base = texts[0]
    for t in texts[1:]:
        if t == base:
            continue
        # 一方是另一方前缀且长度 >=4
        if len(t) >= 4 and len(base) >= 4:
            if t.startswith(base) or base.startswith(t):
                continue
        return False
    return True


def _find_h1_ancestor(b: dict, by_idx: dict) -> Optional[int]:
    """沿 parent_block_idx 链向上找 H1 块的 heading_idx; 找不到返回 None."""
    seen_path = set()
    cur = b
    while cur and cur["heading_idx"] not in seen_path:
        seen_path.add(cur["heading_idx"])
        if cur["level"] == 1:
            return cur["heading_idx"]
        pidx = cur["parent_block_idx"]
        if pidx < 0 or pidx not in by_idx:
            return None
        cur = by_idx[pidx]
    return None


def _same_h1(b1: dict, b2: dict, by_idx: dict) -> bool:
    h1a = _find_h1_ancestor(b1, by_idx)
    h1b = _find_h1_ancestor(b2, by_idx)
    if h1a is None or h1b is None:
        # 无 H1 上下文 → 视为同一(顶级)
        return h1a == h1b
    return h1a == h1b


def parent_prefix(number: Optional[tuple]) -> Optional[tuple]:
    """块的父块编号前缀. (3,1,2) -> (3,1); (3,1) -> (3,); (3,) -> ()."""
    if number is None:
        return None
    return number[:-1]


# ---------- 错位 / 重复检测 ----------

def detect_issues(blocks: list[dict]) -> tuple[list[dict], list[dict]]:
    """返回 (misordered_groups, duplicate_groups).

    misordered_group: 同父下兄弟块按文档顺序 vs number 排序顺序不一致.
    duplicate_group: 同父下 number 完全相同的多个块.
    """
    # 按 (level, parent_block_idx) 分组兄弟块 — parent_block_idx 是文档位置上
    # 的真实父块, 不是 number prefix. 这样跨章/重排过的章节里的同号 heading 不会被
    # 误归为兄弟.
    # **不处理 H1 (level=1)** — 顶级章节由多份合并文档而来, 同号 H1 (如多个 "2 xxx") 是
    # 不同章节的合法重复, 不能当 duplicate 删. H1 顺序按文档顺序保留.
    groups: dict[tuple, list[dict]] = {}
    for b in blocks:
        if b["number"] is None:
            continue
        if b["level"] == 1:
            continue
        key = (b["level"], b["parent_block_idx"])
        groups.setdefault(key, []).append(b)

    # 找每个 block 的父块, 取出父块 number prefix 用于"合法兄弟"过滤
    block_by_idx = {b["heading_idx"]: b for b in blocks}

    misordered = []
    duplicates = []
    for key, sibs in groups.items():
        level, parent_idx = key
        # 合法兄弟过滤: 只考虑 number 的 prefix 与父块 number 一致的块.
        # 比如父块=(4,) "4 水价管理机制" 下, 只算 (4,x) 的 H2 块为兄弟; (3,4)/(3,5) 跨章
        # 漂浮块不当兄弟 (它们结构上不该在这里, 但 number-sort 会越权把它们排进来).
        if parent_idx >= 0 and parent_idx in block_by_idx:
            parent_num = block_by_idx[parent_idx]["number"]
        else:
            parent_num = ()
        if parent_num is None:
            parent_num = ()

        valid_sibs = []
        invalid_sibs = []
        for b in sibs:
            if b["number"][:len(parent_num)] == parent_num:
                valid_sibs.append(b)
            else:
                invalid_sibs.append(b)

        if invalid_sibs:
            # 记录但不处理(留 issue 给后续人工/其他脚本)
            pass

        if not valid_sibs:
            continue

        # 文档顺序 (按 heading_idx)
        sibs_doc = sorted(valid_sibs, key=lambda x: x["heading_idx"])

        # 检重复 — 仅当 number 相同 AND heading 文本除编号外一致 → 真重复.
        # number 相同但文本不同 = number-collision (不同标题用同一编号), 不删.
        # 在同 number 内部按 stripped-text 聚类, 每簇若 >= 2 则记一组 duplicates.
        seen: dict[tuple, list[dict]] = {}
        for b in sibs_doc:
            seen.setdefault(b["number"], []).append(b)
        for num, group in seen.items():
            if len(group) <= 1:
                continue
            # 按 stripped text 聚类
            clusters: dict[str, list[dict]] = {}
            for b in group:
                k = _strip_number_prefix(b["heading_text"])
                # 简化: 用前 20 字作 key
                key_text = k[:20]
                clusters.setdefault(key_text, []).append(b)
            for ktext, cluster in clusters.items():
                if len(cluster) >= 2:
                    duplicates.append({
                        "key": key,
                        "number": num,
                        "blocks": cluster,
                    })

        # 检错位 — 从 sibs 中剔除"真 duplicate 的非 keeper" (那些块会被删, 不参与排序),
        # 但 number-collision 的两块都保留参与排序.
        to_delete_idx: set[int] = set()
        for num, group in seen.items():
            if len(group) <= 1:
                continue
            clusters: dict[str, list[dict]] = {}
            for b in group:
                key_text = _strip_number_prefix(b["heading_text"])[:20]
                clusters.setdefault(key_text, []).append(b)
            for ktext, cluster in clusters.items():
                if len(cluster) >= 2:
                    keeper = pick_keeper(cluster)
                    for b in cluster:
                        if b is not keeper:
                            to_delete_idx.add(b["heading_idx"])
        survivors = [b for b in valid_sibs if b["heading_idx"] not in to_delete_idx]
        # 排序键: 主键 number, 副键 heading_idx 保稳定
        kept_doc = sorted(survivors, key=lambda x: x["heading_idx"])
        kept_sorted = sorted(survivors, key=lambda x: (x["number"], x["heading_idx"]))
        if [b["heading_idx"] for b in kept_doc] != [b["heading_idx"] for b in kept_sorted]:
            misordered.append({
                "key": key,
                "before": [b["number"] for b in kept_doc],
                "after": [b["number"] for b in kept_sorted],
                "kept_doc": kept_doc,
                "kept_sorted": kept_sorted,
            })

    return misordered, duplicates


def pick_keeper(group: list[dict]) -> dict:
    """从同 number 的重复块中挑保留者.

    优先: ① styled (Heading X 样式) ② 块内段数多 ③ 后出现 (heading_idx 大)
    """
    def score(b):
        styled_score = 1 if b["styled"] else 0
        size_score = len(b["paragraph_indices"])
        idx_score = b["heading_idx"]  # 后出现优先
        return (styled_score, size_score, idx_score)

    return max(group, key=score)


# ---------- 重排 / 删除 (lxml) ----------

def apply_changes(doc, blocks: list[dict], misordered: list[dict],
                  duplicates: list[dict]) -> tuple[int, int]:
    """对 docx body 应用重排+删除.

    返回 (moves_applied, deletions_applied).

    策略:
      1. 先标记所有要删除的段 idx (重复块的非 keeper)
      2. 对每个 misordered group, 按 number 排序后, 找出兄弟块的 XML 片段集合,
         在 body 中 detach 再按新顺序 insert
      3. 兄弟块范围 = [heading_idx, end_idx); 跨章节边界检查: 所有兄弟在同 H1 内

    实现: 全用 lxml `<w:body>` 上操作 `<w:p>`.
    """
    body = doc.element.body
    # 收集 body 下所有顶层 <w:p> (注意 body 还含 <w:sectPr> 等其他元素)
    # paragraphs idx 对齐 doc.paragraphs (跳过 sectPr/tbl)
    # doc.paragraphs 只含 body 下 <w:p>, idx 与 body 中 <w:p> 顺序一致.
    all_p_elements = list(body.iter(W_P))
    # 但 doc.paragraphs 可能含 table 内的 p — 用 doc.paragraphs[i]._element 更安全
    para_elems = [p._element for p in doc.paragraphs]

    # ---- step 1: 删除重复 ----
    delete_idx_set: set[int] = set()
    delete_records = []
    for dup in duplicates:
        keeper = pick_keeper(dup["blocks"])
        for b in dup["blocks"]:
            if b is keeper:
                continue
            for idx in b["paragraph_indices"]:
                delete_idx_set.add(idx)
            delete_records.append({
                "number": list(dup["number"]),
                "kept_idx": keeper["heading_idx"],
                "deleted_idx": b["heading_idx"],
                "reason": ("kept styled" if keeper["styled"] and not b["styled"]
                           else "kept larger/later block"),
            })

    # ---- step 2: 重排兄弟块 ----
    # 关键: 兄弟块的段范围在 docx 中可能 交错(因为错位本身就是源 idx 乱). 我们要按 "新顺序" 把每个 keeper 块的 <w:p> 连续起来, 放在原最早兄弟的位置.
    # 实现:
    #   a) 收集本 group 所有 keeper 的段 idx 集合 union (排除 delete_idx_set 里的)
    #   b) 这些段在 body 中 detach
    #   c) 按 keeper 的 number 排序顺序, 依次把每个 keeper 块的段 XML 列表合并
    #   d) 在原最早段位置之前 insert 进 body
    moves_applied = 0

    # 同一个 idx 不能被两个 group 重排 → 但 group 按 (level, parent_prefix) 分,
    # 不同 group 的段不相交. 安全.
    # 先处理: 收集要重排的段 idx 集合, 这些段不能在 delete 阶段被 remove (会丢失).
    # 顺序: 先处理删除(纯 remove), 再处理重排(detach + insert).
    # 但删除涉及的段如果在某 group keeper 块内 — 不会, keeper 不会被标记为删.
    # 非 keeper 块的段 idx 在 delete_idx_set 中, 重排时这些块也不在 keeper 列表 → ok.

    # 但还有一种情况: 非 keeper 块的段, 如果它在某 misordered group 中作为 "before" 出现 → 因 pick_keeper 已被剔除, 不在 kept_sorted 中. ok.

    # 先执行 删除
    deletions_applied = 0
    for idx in sorted(delete_idx_set, reverse=True):
        elem = para_elems[idx]
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            deletions_applied += 1

    # 重新刷新 para_elems? — 用旧引用, 因为 lxml elem 即使 detach 仍可用,
    # 而我们只对剩余 keeper 的段 idx 操作, 它们的 elem 引用仍有效.

    # 跨章检查: 仅在同一 H1 (chapter) 范围内允许重排.
    # 找出每个 H1 块的 idx 范围
    h1_ranges = []
    for b in blocks:
        if b["level"] == 1:
            h1_ranges.append((b["heading_idx"], b["end_idx"]))
    if not h1_ranges:
        # 无 H1 — 整 doc 一个范围
        h1_ranges = [(0, len(para_elems))]

    def same_chapter(idx_list: list[int]) -> bool:
        if not idx_list:
            return True
        for lo, hi in h1_ranges:
            if all(lo <= i < hi for i in idx_list):
                return True
        return False

    for mis in misordered:
        kept_sorted = mis["kept_sorted"]
        kept_doc = mis["kept_doc"]
        if [b["number"] for b in kept_doc] == [b["number"] for b in kept_sorted]:
            continue  # 已经对了

        # 跨章检查
        heading_idxs = [b["heading_idx"] for b in kept_sorted]
        if not same_chapter(heading_idxs):
            print(f"[skip] cross-chapter group {mis['key']} — refused")
            continue

        # 收集每个 keeper 块的 elem list (按 keeper 的 paragraph_indices, 跳过 delete)
        keeper_elems_per_block = []
        for b in kept_sorted:
            elems = [para_elems[i] for i in b["paragraph_indices"]
                     if i not in delete_idx_set]
            keeper_elems_per_block.append(elems)

        all_elems_flat = [e for lst in keeper_elems_per_block for e in lst]
        if not all_elems_flat:
            continue
        # anchor = 文档原顺序中最早的 keeper 段 (kept_doc 第一个块的首段)
        first_doc_block = min(kept_doc, key=lambda x: x["heading_idx"])
        anchor_elem = para_elems[first_doc_block["heading_idx"]]
        anchor_parent = anchor_elem.getparent()
        if anchor_parent is None:
            continue
        # 用 anchor 前面的非 detach 元素作为 "前驱锚"
        # 找 anchor_elem 的 previous sibling 不在 all_elems_flat 集合里的
        detach_set = set(id(e) for e in all_elems_flat)
        pred_elem = anchor_elem.getprevious()
        while pred_elem is not None and id(pred_elem) in detach_set:
            pred_elem = pred_elem.getprevious()

        # detach
        for e in all_elems_flat:
            parent = e.getparent()
            if parent is not None:
                parent.remove(e)

        # 重新计算 insert_pos: pred_elem 之后第一个位置
        if pred_elem is None:
            insert_pos = 0
        else:
            insert_pos = list(anchor_parent).index(pred_elem) + 1

        for block_elems in keeper_elems_per_block:
            for e in block_elems:
                anchor_parent.insert(insert_pos, e)
                insert_pos += 1

        moves_applied += 1

    return moves_applied, deletions_applied


# ---------- main ----------

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("docx", help="输入 docx 路径")
    ap.add_argument("--dry-run", action="store_true", help="只生成 plan, 不动 docx")
    ap.add_argument("--no-backup", action="store_true", help="不自动备份")
    ap.add_argument("--report", help="JSON plan 输出路径")
    args = ap.parse_args(argv)

    src = Path(args.docx)
    if not src.exists():
        print(f"[error] 找不到 {src}", file=sys.stderr)
        return 2

    # 备份
    backup_path = None
    if not args.dry_run and not args.no_backup:
        today = date.today().isoformat()
        n = 1
        while True:
            cand = src.with_name(f"{src.stem}.bak-{n}-{today}{src.suffix}")
            if not cand.exists():
                break
            n += 1
        shutil.copy2(src, cand)
        backup_path = cand
        print(f"[backup] {cand.name}")

    doc = Document(str(src))

    # 段分类
    paras = [classify_paragraph(p, i) for i, p in enumerate(doc.paragraphs)]
    blocks = slice_blocks(paras)

    misordered, duplicates = detect_issues(blocks)

    plan = {
        "input": str(src),
        "backup": str(backup_path) if backup_path else None,
        "blocks_total": len(blocks),
        "misordered_pairs": [
            {
                "level": m["key"][0],
                "parent_idx": m["key"][1],
                "before": [list(n) for n in m["before"]],
                "after": [list(n) for n in m["after"]],
            }
            for m in misordered
        ],
        "duplicates": [
            {
                "level": d["key"][0],
                "parent_idx": d["key"][1],
                "number": list(d["number"]),
                "kept_idx": pick_keeper(d["blocks"])["heading_idx"],
                "deleted_idxs": [b["heading_idx"] for b in d["blocks"]
                                 if b is not pick_keeper(d["blocks"])],
            }
            for d in duplicates
        ],
        "dry_run": args.dry_run,
    }

    print(f"[plan] blocks={len(blocks)}  misordered_groups={len(misordered)}  "
          f"duplicate_groups={len(duplicates)}")
    for m in misordered[:20]:
        b = [".".join(str(x) for x in n) for n in m["before"]]
        a = [".".join(str(x) for x in n) for n in m["after"]]
        print(f"  misorder L{m['key'][0]} parent_idx={m['key'][1]}: {b} → {a}")
    for d in duplicates[:20]:
        num = ".".join(str(x) for x in d["number"])
        kept = pick_keeper(d["blocks"])["heading_idx"]
        delkk = [b["heading_idx"] for b in d["blocks"] if b is not pick_keeper(d["blocks"])]
        print(f"  dup L{d['key'][0]} parent_idx={d['key'][1]} num={num}: keep idx={kept} del={delkk}")

    if not args.dry_run:
        moves, dels = apply_changes(doc, blocks, misordered, duplicates)
        plan["moves_applied"] = moves
        plan["deletions_applied"] = dels
        doc.save(str(src))
        print(f"[apply] moves={moves} deletions={dels} → saved")
    else:
        plan["moves_applied"] = 0
        plan["deletions_applied"] = 0
        print("[dry-run] no changes written")

    if args.report:
        Path(args.report).parent.mkdir(parents=True, exist_ok=True)
        Path(args.report).write_text(json.dumps(plan, indent=2, ensure_ascii=False))
        print(f"[report] {args.report}")

    return 0


# ---------------- pipeline adapter ----------------
def apply(doc, args=None) -> dict:
    dry = bool(getattr(args, "dry_run", False)) if args else False
    paras = [classify_paragraph(p, i) for i, p in enumerate(doc.paragraphs)]
    blocks = slice_blocks(paras)
    misordered, duplicates = detect_issues(blocks)
    moves = dels = 0
    if not dry:
        moves, dels = apply_changes(doc, blocks, misordered, duplicates)
    return {
        "changed": moves + dels,
        "blocks_total": len(blocks),
        "misordered_groups": len(misordered),
        "duplicate_groups": len(duplicates),
        "moves_applied": moves,
        "deletions_applied": dels,
    }


if __name__ == "__main__":
    sys.exit(main())
