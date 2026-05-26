#!/usr/bin/env python3
"""section_read.py — 按标题关键字读取 DOCX 章节内容

Distilled from panan-rigid-2026/scripts/read_section.py (B级通用, 2026-05-26).
原版 DOCX_GLOB 硬编码改为必须显式传入 docx 路径.
核心逻辑(heading_level / list_headings / read_section)完全通用.

用法:
    python3 section_read.py <docx_file> <query>     # 按标题关键字读章节
    python3 section_read.py <docx_file> --list      # 列出所有标题

参数:
    docx_file : DOCX 文件路径（必须显式指定）
    query     : 标题关键字（支持模糊匹配 — 任意标题含 query 则命中）
    --list    : 改为列出所有标题（含段落索引，方便后续 md merge-into-docx 使用）

触发场景:
    - 查看 DOCX 某章节当前内容（知道关键字但不知索引）
    - 先 --list 获取段落索引，再决定 md merge-into-docx 的 start/end
    - 快速浏览大型 Word 报告的章节结构
"""
from __future__ import annotations

import re
import sys
from pathlib import Path
from docx import Document


def heading_level(style_name: str) -> int:
    """Return heading level from style name ('Heading 2' → 2); 99 for non-headings."""
    m = re.search(r"Heading\s*(\d)", style_name)
    return int(m.group(1)) if m else 99


def list_headings(doc: Document) -> None:
    """Print all headings with their paragraph indices."""
    for i, p in enumerate(doc.paragraphs):
        lvl = heading_level(p.style.name)
        if lvl < 99 and p.text.strip():
            indent = "  " * (lvl - 1)
            print(f"[{i:>4}] {indent}{p.text.strip()[:100]}")


def read_section(doc: Document, query: str, doc_path: str = "") -> None:
    """Print all paragraphs in the section whose heading matches query."""
    query_clean = re.sub(r"[（(].*$", "", query).strip()

    # Exact substring match first
    start_idx = None
    start_level = None
    for i, p in enumerate(doc.paragraphs):
        lvl = heading_level(p.style.name)
        if lvl < 99 and query_clean in p.text:
            start_idx = i
            start_level = lvl
            break

    # Fuzzy: all keywords must appear in heading
    if start_idx is None:
        keywords = [w for w in query_clean.split() if len(w) > 1]
        if keywords:
            for i, p in enumerate(doc.paragraphs):
                lvl = heading_level(p.style.name)
                if lvl < 99 and all(k in p.text for k in keywords):
                    start_idx = i
                    start_level = lvl
                    break

    if start_idx is None:
        sys.exit(f"找不到匹配「{query}」的标题")

    # Collect paragraphs until next same-or-higher heading
    end_idx = len(doc.paragraphs)
    for i in range(start_idx + 1, len(doc.paragraphs)):
        lvl = heading_level(doc.paragraphs[i].style.name)
        if lvl <= start_level:
            end_idx = i
            break

    label = doc_path or "<docx>"
    print(f"# 文件: {label}")
    print(f"# 段落 [{start_idx}] ~ [{end_idx - 1}]，共 {end_idx - start_idx} 段\n")
    for i in range(start_idx, end_idx):
        p = doc.paragraphs[i]
        lvl = heading_level(p.style.name)
        text = p.text.strip()
        if not text:
            continue
        if lvl < 99:
            marker = "#" * lvl
            print(f"[{i}] {marker} {text}")
        else:
            print(f"[{i}] {text[:200]}")


def apply_path(docx_path: str | None = None, args=None) -> dict:
    """pipeline adapter — delegates to main(); docx_path injected as sys.argv[1]."""
    try:
        saved = sys.argv[:]
        if docx_path and (len(sys.argv) < 2 or sys.argv[1] != str(docx_path)):
            sys.argv = [sys.argv[0], str(docx_path)] + sys.argv[2:]
        rc = main()
        sys.argv = saved
        return {"status": "ok", "rc": rc, "script": "section_read.py"}
    except SystemExit as e:
        return {"status": "sysexit", "code": e.code, "script": "section_read.py"}
    except Exception as e:
        return {"status": "error", "error": repr(e), "script": "section_read.py"}


def main() -> int:
    if len(sys.argv) < 2 or sys.argv[1] in ("-h", "--help"):
        print(__doc__)
        return 0

    docx_path = sys.argv[1]
    rest_args = sys.argv[2:]

    if not rest_args or (len(rest_args) == 1 and rest_args[0] == "--list"):
        # Allow: script docx --list  OR  script docx  (prints headings)
        doc = Document(docx_path)
        list_headings(doc)
        return 0

    if rest_args[0] == "--list":
        doc = Document(docx_path)
        list_headings(doc)
        return 0

    query = " ".join(rest_args)
    doc = Document(docx_path)
    read_section(doc, query, doc_path=docx_path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
