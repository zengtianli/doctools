#!/usr/bin/env python3
"""delete_chapter.py — 按 H1 段编号或 text 前缀,删整章(H1 + 下属所有段+表 直到下一个 H1 或文末).

用法:
    python3 scripts/delete_chapter.py <docx> --h1 6,7,8 [--dry-run] [--no-backup] [--report <json>]
    python3 scripts/delete_chapter.py <docx> --h1-text "6 保障措施,7 嘉兴..."

设计:
    - 用 python-docx 找 H1 段
    - 用 lxml 操作 <w:body>: 收集每章 [H1 段 <w:p>, ..., 下一 H1 前/文末) 之间所有 <w:p>+<w:tbl>
    - 保留末尾 <w:sectPr>
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

H1_STYLE = "Heading 1"


def find_h1_paragraphs(doc):
    """返回 [(idx_in_paragraphs, paragraph)] 仅 Heading 1."""
    return [(i, p) for i, p in enumerate(doc.paragraphs) if p.style and p.style.name == H1_STYLE]


def match_h1_by_number(h1_list, numbers):
    """numbers: list[str] like ['6','7','8']. 匹配 text 开头 '^N\\s+'."""
    matched = []
    for n in numbers:
        pat = re.compile(rf"^{re.escape(n)}\s+")
        hits = [(i, p) for i, p in h1_list if pat.match(p.text)]
        if not hits:
            raise SystemExit(f"[ERR] 未找到 H1 编号 {n!r} 的段")
        if len(hits) > 1:
            raise SystemExit(f"[ERR] H1 编号 {n!r} 匹配到 {len(hits)} 段 (期望 1)")
        matched.append(hits[0])
    return matched


def match_h1_by_text_prefix(h1_list, prefixes):
    matched = []
    for pref in prefixes:
        hits = [(i, p) for i, p in h1_list if p.text.startswith(pref)]
        if not hits:
            raise SystemExit(f"[ERR] 未找到 H1 text 前缀 {pref!r}")
        if len(hits) > 1:
            raise SystemExit(f"[ERR] H1 text 前缀 {pref!r} 匹配 {len(hits)} 段")
        matched.append(hits[0])
    return matched


def collect_block_elements(body):
    """body 的直接子节点中,顺序保留 <w:p> 与 <w:tbl>(<w:sectPr> 排除)."""
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    return [el for el in body.iterchildren() if el.tag in (p_tag, tbl_tag)]


def build_delete_plan(doc, target_h1_elems):
    """对每个 target H1 <w:p>,找其在 body block-list 中位置,延伸到下一 H1 前(或末尾).
    返回 list of dict: {h1_text, start_block_idx, end_block_idx_exclusive, n_paragraphs, n_tables, elements}
    """
    body = doc.element.body
    blocks = collect_block_elements(body)
    block_index = {id(el): i for i, el in enumerate(blocks)}

    # 找出所有 H1 段 <w:p> 在 blocks 中的位置(按 docx style)
    h1_positions = []  # list of (block_idx, element, text)
    for p in doc.paragraphs:
        if p.style and p.style.name == H1_STYLE:
            idx = block_index.get(id(p._element))
            if idx is not None:
                h1_positions.append((idx, p._element, p.text))

    h1_positions.sort(key=lambda x: x[0])
    # 给每个 H1 找 "下一个 H1 起始 idx"
    next_h1_idx = {}
    for i, (bi, _el, _txt) in enumerate(h1_positions):
        next_h1_idx[bi] = h1_positions[i + 1][0] if i + 1 < len(h1_positions) else len(blocks)

    plans = []
    target_ids = {id(el) for el in target_h1_elems}
    for bi, el, txt in h1_positions:
        if id(el) not in target_ids:
            continue
        end = next_h1_idx[bi]
        chunk = blocks[bi:end]
        n_p = sum(1 for e in chunk if e.tag == qn("w:p"))
        n_t = sum(1 for e in chunk if e.tag == qn("w:tbl"))
        plans.append({
            "h1_text": txt,
            "start_block_idx": bi,
            "end_block_idx_exclusive": end,
            "n_blocks": end - bi,
            "n_paragraphs": n_p,
            "n_tables": n_t,
            "_elements": chunk,
        })
    return plans


def apply_plan(plans):
    total_p = 0
    total_t = 0
    for plan in plans:
        for el in plan["_elements"]:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        total_p += plan["n_paragraphs"]
        total_t += plan["n_tables"]
    return total_p, total_t


def next_backup_path(docx_path: Path) -> Path:
    today = date.today().isoformat()
    stem = docx_path.stem
    parent = docx_path.parent
    n = 1
    while True:
        cand = parent / f"{stem}.bak-{n}-{today}.docx"
        if not cand.exists():
            return cand
        n += 1


def main():
    ap = argparse.ArgumentParser(description="按 H1 编号或 text 前缀删整章 (H1 + 其下所有段表 → 下一 H1 前/文末).")
    ap.add_argument("docx", type=Path, help="目标 docx 路径")
    grp = ap.add_mutually_exclusive_group(required=True)
    grp.add_argument("--h1", help="H1 编号列表, 逗号分隔. 例: 6,7,8")
    grp.add_argument("--h1-text", help="H1 text 前缀列表, 逗号分隔. 例: '6 保障措施,7 嘉兴...'")
    ap.add_argument("--dry-run", action="store_true", help="只列 plan, 不写盘")
    ap.add_argument("--no-backup", action="store_true", help="不生成 .bak-N 备份")
    ap.add_argument("--report", type=Path, help="JSON 报告输出路径")
    args = ap.parse_args()

    if not args.docx.exists():
        raise SystemExit(f"[ERR] docx 不存在: {args.docx}")

    doc = Document(str(args.docx))
    h1_list = find_h1_paragraphs(doc)
    if not h1_list:
        raise SystemExit("[ERR] 文档无 Heading 1")

    if args.h1:
        nums = [s.strip() for s in args.h1.split(",") if s.strip()]
        matched = match_h1_by_number(h1_list, nums)
    else:
        prefs = [s.strip() for s in args.h1_text.split(",") if s.strip()]
        matched = match_h1_by_text_prefix(h1_list, prefs)

    target_elems = [p._element for _, p in matched]
    plans = build_delete_plan(doc, target_elems)

    # 打印 plan
    print(f"[plan] docx: {args.docx}")
    print(f"[plan] 全文 H1 段总数: {len(h1_list)}")
    print(f"[plan] 拟删章数: {len(plans)}")
    for pl in plans:
        print(f"  - H1='{pl['h1_text']}' blocks[{pl['start_block_idx']}:{pl['end_block_idx_exclusive']}] "
              f"=> {pl['n_paragraphs']} 段 + {pl['n_tables']} 表 (共 {pl['n_blocks']} 块)")

    report = {
        "docx": str(args.docx),
        "dry_run": args.dry_run,
        "h1_total_before": len(h1_list),
        "plans": [{k: v for k, v in pl.items() if not k.startswith("_")} for pl in plans],
    }

    if args.dry_run:
        print("[dry-run] 未写盘")
    else:
        if not args.no_backup:
            bak = next_backup_path(args.docx)
            shutil.copy2(args.docx, bak)
            print(f"[backup] {bak}")
            report["backup"] = str(bak)
        total_p, total_t = apply_plan(plans)
        doc.save(str(args.docx))
        print(f"[done] 已删 {total_p} 段 + {total_t} 表, 保存 → {args.docx}")
        report["deleted_paragraphs"] = total_p
        report["deleted_tables"] = total_t

        # verify 重读 + H1 列表
        doc2 = Document(str(args.docx))
        h1_after = [p.text for p in doc2.paragraphs if p.style and p.style.name == H1_STYLE]
        print(f"[verify] 删后 H1 段数: {len(h1_after)}")
        for t in h1_after:
            print(f"  · {t}")
        report["h1_total_after"] = len(h1_after)
        report["h1_texts_after"] = h1_after

    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"[report] {args.report}")


# ---------------- pipeline adapter ----------------
def apply(doc, args=None) -> dict:
    h1 = getattr(args, "delete_h1", None) if args else None
    h1_text = getattr(args, "delete_h1_text", None) if args else None
    dry = bool(getattr(args, "dry_run", False)) if args else False
    if not h1 and not h1_text:
        return {"changed": 0, "skipped": "no delete_h1 / delete_h1_text in args"}
    h1_list = find_h1_paragraphs(doc)
    if h1:
        nums = [s.strip() for s in h1.split(",") if s.strip()]
        matched = match_h1_by_number(h1_list, nums)
    else:
        prefs = [s.strip() for s in h1_text.split(",") if s.strip()]
        matched = match_h1_by_text_prefix(h1_list, prefs)
    target_elems = [p._element for _, p in matched]
    plans = build_delete_plan(doc, target_elems)
    if not dry:
        total_p, total_t = apply_plan(plans)
    else:
        total_p = sum(pl["n_paragraphs"] for pl in plans)
        total_t = sum(pl["n_tables"] for pl in plans)
    return {
        "changed": total_p + total_t,
        "deleted_paragraphs": total_p,
        "deleted_tables": total_t,
        "plans_count": len(plans),
    }


if __name__ == "__main__":
    main()
