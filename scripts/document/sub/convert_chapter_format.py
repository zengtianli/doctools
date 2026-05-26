#!/usr/bin/env python3
# distilled from qual-supply/scripts/convert_chapter_format.py (2026-05-25 W1)
"""convert_chapter_format.py — H1 段「第X章 Y」/「X、Y」格式 → 「N Y」阿拉伯数字。

仅动 style.name == "Heading 1" 段;不碰 Title / H2 / H3 / 正文 / 表 cell。
run 级替换保留 bold / 字号 / 中文宋体。

接口:
    python3 scripts/convert_chapter_format.py <docx_path> [--dry-run]
                                                         [--no-backup]
                                                         [--report <json>]
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document

# ---------------------------------------------------------------------------
# 中文数字 → 阿拉伯
# ---------------------------------------------------------------------------
_CN_DIGIT = {
    "零": 0, "一": 1, "二": 2, "三": 3, "四": 4,
    "五": 5, "六": 6, "七": 7, "八": 8, "九": 9,
}
_CN_UNIT = {"十": 10, "百": 100, "千": 1000}


def chinese_to_arabic(s: str) -> int:
    """中文数字串转 int。支持「一」「十」「十一」「二十」「三十五」「一百零三」等。

    若 s 已是阿拉伯数字串,直接 int(s)。无法解析抛 ValueError。
    """
    s = s.strip()
    if not s:
        raise ValueError("empty numeral")
    if s.isdigit():
        return int(s)

    total = 0
    current = 0  # 当前累积位段
    last_unit = 0

    for ch in s:
        if ch in _CN_DIGIT:
            current = _CN_DIGIT[ch]
        elif ch in _CN_UNIT:
            unit = _CN_UNIT[ch]
            if current == 0:
                # 「十X」开头省略一: 十 = 10
                current = 1
            total += current * unit
            current = 0
            last_unit = unit
        else:
            raise ValueError(f"unrecognized char in numeral: {ch!r}")

    total += current
    return total


# ---------------------------------------------------------------------------
# 编号匹配
# ---------------------------------------------------------------------------
# 模式 A: 第X章 Y   (X 可中文数字或阿拉伯)
_RE_CHAPTER = re.compile(r"^第([一二三四五六七八九十百零\d]+)章\s+(.+)$")
# 模式 B: X、Y      (X 必须中文数字; 阿拉伯+顿号不动)
_RE_DUNHAO = re.compile(r"^([一二三四五六七八九十百零]+)、\s*(.+)$")
# 模式 C: 已是阿拉伯「N Y」 → 跳过
_RE_ARABIC = re.compile(r"^\d+\s+\S")


def parse_h1(text: str) -> tuple[int, str] | None:
    """返回 (num, rest) 或 None(不匹配/已 normalized/无法转换)。"""
    if _RE_ARABIC.match(text):
        return None
    m = _RE_CHAPTER.match(text)
    if m:
        try:
            return chinese_to_arabic(m.group(1)), m.group(2).strip()
        except ValueError:
            return None
    m = _RE_DUNHAO.match(text)
    if m:
        try:
            return chinese_to_arabic(m.group(1)), m.group(2).strip()
        except ValueError:
            return None
    return None


# ---------------------------------------------------------------------------
# Run 级替换 (保留首 run 格式)
# ---------------------------------------------------------------------------
def rewrite_paragraph_runs(paragraph, new_text: str) -> None:
    """把段的可见文字改成 new_text,保留 runs[0] 的格式属性。

    做法:把 new_text 写进 runs[0].text;清空其余 run text(保 run 节点以防
    被段落级 pPr 引用;实测安全)。这样 H1 段编号 + 标题文字全部走 run[0] 的
    rPr(bold/size/宋体),格式 100% 保留。
    """
    runs = paragraph.runs
    if not runs:
        # 罕见: paragraph 无 run 直接放裸文本,fallback 用 text=
        paragraph.text = new_text
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def process(docx_path: Path, dry_run: bool, do_backup: bool) -> dict:
    doc = Document(str(docx_path))
    plan: list[dict] = []

    for idx, p in enumerate(doc.paragraphs):
        if p.style.name != "Heading 1":
            continue
        old = p.text
        parsed = parse_h1(old)
        if parsed is None:
            plan.append({
                "para_idx": idx,
                "action": "skip",
                "before": old,
                "after": old,
            })
            continue
        num, rest = parsed
        new = f"{num} {rest}"
        plan.append({
            "para_idx": idx,
            "action": "rewrite",
            "before": old,
            "after": new,
        })
        if not dry_run:
            rewrite_paragraph_runs(p, new)

    report = {
        "docx": str(docx_path),
        "dry_run": dry_run,
        "total_h1": len(plan),
        "rewrite_count": sum(1 for x in plan if x["action"] == "rewrite"),
        "skip_count": sum(1 for x in plan if x["action"] == "skip"),
        "items": plan,
        "backup": None,
    }

    if not dry_run:
        if do_backup:
            backup = _next_backup_path(docx_path)
            shutil.copy2(docx_path, backup)
            report["backup"] = str(backup)
        doc.save(str(docx_path))

    return report


def _next_backup_path(docx_path: Path) -> Path:
    """生成 `<stem>.bak-N-YYYY-MM-DD.docx`,N 递增直至不冲突。"""
    today = date.today().isoformat()
    parent = docx_path.parent
    stem = docx_path.stem
    suffix = docx_path.suffix
    n = 1
    while True:
        candidate = parent / f"{stem}.bak-{n}-{today}{suffix}"
        if not candidate.exists():
            return candidate
        n += 1


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main() -> int:
    ap = argparse.ArgumentParser(
        description="H1 段「第X章 Y」/「X、Y」→「N Y」阿拉伯数字格式转换。",
    )
    ap.add_argument("docx", type=Path, help="目标 docx 路径")
    ap.add_argument("--dry-run", action="store_true",
                    help="只打印 plan 不改文件")
    ap.add_argument("--no-backup", action="store_true",
                    help="真改时跳过 .bak 备份")
    ap.add_argument("--report", type=Path,
                    help="把 JSON 报告写到此路径")
    args = ap.parse_args()

    if not args.docx.exists():
        print(f"ERROR: docx not found: {args.docx}", file=sys.stderr)
        return 2

    report = process(
        args.docx,
        dry_run=args.dry_run,
        do_backup=not args.no_backup,
    )

    # 终端摘要
    mode = "DRY-RUN" if args.dry_run else "APPLIED"
    print(f"[{mode}] {args.docx}")
    print(f"  H1 总数: {report['total_h1']}")
    print(f"  改写: {report['rewrite_count']}  跳过: {report['skip_count']}")
    if report.get("backup"):
        print(f"  备份: {report['backup']}")
    for item in report["items"]:
        tag = "REWRITE" if item["action"] == "rewrite" else "skip   "
        print(f"  [{tag}] para#{item['para_idx']}: {item['before']!r}"
              f" -> {item['after']!r}")

    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(
            json.dumps(report, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        print(f"  JSON 报告: {args.report}")

    return 0


# ---------------- pipeline adapter ----------------
def apply(doc, args=None) -> dict:
    """pipeline: H1 「第X章 Y」/「X、Y」→「N Y」"""
    dry = bool(getattr(args, "dry_run", False)) if args else False
    plan = []
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name != "Heading 1":
            continue
        old = p.text
        parsed = parse_h1(old)
        if parsed is None:
            plan.append({"para_idx": idx, "action": "skip", "before": old, "after": old})
            continue
        num, rest = parsed
        new = f"{num} {rest}"
        plan.append({"para_idx": idx, "action": "rewrite", "before": old, "after": new})
        if not dry:
            rewrite_paragraph_runs(p, new)
    return {
        "changed": sum(1 for x in plan if x["action"] == "rewrite"),
        "total_h1": len(plan),
        "skip_count": sum(1 for x in plan if x["action"] == "skip"),
        "items": plan,
    }


if __name__ == "__main__":
    sys.exit(main())
