#!/usr/bin/env python3
# distilled from qual-supply/scripts/ (2026-05-25 W2)
# merged 3 scripts:
#   promote_misclassified_h1.py       → outline promote-h1
#   demote_h2_with_h3_format.py       → outline demote-h2
#   normalize_outline_to_arabic.py    → outline normalize-arabic
#
# 项目硬编码样式集合 → 抽到 doctools/config/styles_registry.yaml SSOT,通过 --profile <name> 切换。
"""doctools outline group · profile-driven 大纲层级规范化.

子命令:
    outline promote-h1          把伪 H1 段 (文本 "N、…" 但挂 H2/H3/Title) 升级为真 H1 + 编号阿拉伯化
    outline demote-h2           样式=H2 但文本是 H3 形态 (^\\d+、) 短段 → 降为 H3 (不改文本)
    outline normalize-arabic    一次性识别 H1/H2/H3 文本形态 + 改样式 + 阿拉伯化编号

通用 CLI:
    <docx_path> [--dry-run] [--no-backup] [--report <json>] [--profile <name>]

profile 默认 'zdwp' (兼容 qual-supply);eco-flow / generic 等通过 --profile 切换。
profile 提供 H1_STYLES / H2_STYLES / H3_STYLES / TITLE_STYLES 等集合用于 candidate detection。
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
from datetime import date
from pathlib import Path
from typing import Optional

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[3] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py
from cn_number import chinese_to_arabic  # noqa: E402,F401  中文数字 SSOT

from docx import Document

# loader (支持包 import 或脚本独立 import)
try:
    from doctools.lib.styles import load_profile, StylesProfile  # type: ignore
except ImportError:
    import importlib.util as _ilu
    _styles_path = Path(__file__).resolve().parent.parent.parent.parent / "lib" / "styles.py"
    _spec = _ilu.spec_from_file_location("_doctools_styles", _styles_path)
    _m = _ilu.module_from_spec(_spec)
    _spec.loader.exec_module(_m)  # type: ignore
    load_profile = _m.load_profile
    StylesProfile = _m.StylesProfile


# =============================================================================
# 公共辅助
# =============================================================================

# 中文数字转换：本文件原有一份局部实现（与 chapter.py 那份逐行等价），
# 2026-08-01 下沉到 lib/cn_number.py，见文件顶部 import。


def lsof_check(path: Path) -> Optional[str]:
    try:
        r = subprocess.run(["lsof", str(path)], capture_output=True, text=True, timeout=5)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def pick_backup_path(src: Path) -> Path:
    today = date.today().isoformat()
    parent, stem, suffix = src.parent, src.stem, src.suffix
    n = 1
    while True:
        cand = parent / f"{stem}.bak-{n}-{today}{suffix}"
        if not cand.exists():
            return cand
        n += 1


def _common_setup(args) -> Path:
    src = Path(args.docx)
    if not src.exists():
        print(f"[ERR] 文件不存在: {src}", file=sys.stderr)
        sys.exit(2)
    src = src.resolve()
    if not args.dry_run:
        occ = lsof_check(src)
        if occ:
            print(f"[ABORT] 文件被占用 (Word/WPS):\n{occ}", file=sys.stderr)
            sys.exit(3)
    return src


def _save_with_backup(src: Path, doc, args, wrote_needed: bool = True) -> Optional[Path]:
    if args.dry_run or not wrote_needed:
        return None
    bak = None
    if not args.no_backup:
        bak = pick_backup_path(src)
        shutil.copy2(src, bak)
    doc.save(str(src))
    return bak


def _emit_report(report: dict, args):
    rp = getattr(args, "report", None)
    if rp:
        rp = Path(rp)
        rp.parent.mkdir(parents=True, exist_ok=True)
        rp.write_text(json.dumps(report, ensure_ascii=False, indent=2, default=str), encoding="utf-8")


def _rewrite_paragraph_runs(paragraph, new_text: str) -> None:
    """run 级替换 runs[0].text = new_text; runs[1:].text = '' (保 bold/字号/字体)."""
    runs = paragraph.runs
    if not runs:
        paragraph.text = new_text
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _is_in_table_cell(paragraph) -> bool:
    parent = paragraph._element.getparent()
    while parent is not None:
        tag = parent.tag
        if tag.endswith("}tc"):
            return True
        if tag.endswith("}body"):
            return False
        parent = parent.getparent()
    return False


def _resolve_heading_style(doc, profile: StylesProfile, level: int):
    """从 profile.H{level}_STYLES 列表里挑第一个 doc.styles 实际可用的 style obj.

    Returns: docx style object or None.
    """
    candidates = getattr(profile, f"H{level}_STYLES", []) or []
    available = {s.name: s for s in doc.styles}
    for name in candidates:
        if name in available:
            return available[name]
    return None


# =============================================================================
# 子命令 1: outline promote-h1 — promote_misclassified_h1
# =============================================================================
# 「N、 rest」: N = 中文数字或阿拉伯, 顿号后允许 0+ 空白
_RE_DUNHAO = re.compile(r"^([一二三四五六七八九十百零\d]+)、\s*(.+)$")
PROMOTE_MAX_TEXT_LEN = 80


def _detect_pseudo_h1(paragraph, candidate_styles: set[str]) -> Optional[tuple[int, str]]:
    text = paragraph.text
    if not text or not text.strip():
        return None
    if len(text) >= PROMOTE_MAX_TEXT_LEN:
        return None
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name not in candidate_styles:
        return None
    if style_name == "Heading 1":  # 防御性
        return None
    m = _RE_DUNHAO.match(text)
    if not m:
        return None
    try:
        num = chinese_to_arabic(m.group(1))
    except ValueError:
        return None
    rest = m.group(2).strip()
    if not rest:
        return None
    return num, rest


def _process_promote_h1(doc, profile: StylesProfile, dry_run: bool) -> dict:
    h1_style = _resolve_heading_style(doc, profile, 1)
    if h1_style is None:
        raise SystemExit(
            f"ERROR: no H1 style available in document; profile.H1_STYLES={profile.H1_STYLES}"
        )

    # candidate = H2 ∪ H3 ∪ Title (profile-driven, 不含 H1 / Body)
    candidate_styles: set[str] = set()
    candidate_styles |= set(profile.H2_STYLES or [])
    candidate_styles |= set(profile.H3_STYLES or [])
    candidate_styles |= set(profile.TITLE_STYLES or [])

    plan: list[dict] = []
    for idx, p in enumerate(doc.paragraphs):
        if _is_in_table_cell(p):
            continue
        hit = _detect_pseudo_h1(p, candidate_styles)
        if hit is None:
            continue
        num, rest = hit
        old_text = p.text
        old_style = p.style.name
        new_text = f"{num} {rest}"
        plan.append({
            "para_idx": idx,
            "action": "promote",
            "before_style": old_style,
            "after_style": h1_style.name,
            "before_text": old_text,
            "after_text": new_text,
        })
        if not dry_run:
            _rewrite_paragraph_runs(p, new_text)
            p.style = h1_style

    return {
        "candidate_styles": sorted(candidate_styles),
        "max_text_len": PROMOTE_MAX_TEXT_LEN,
        "promote_count": len(plan),
        "items": plan,
    }


def cmd_promote_h1(args) -> int:
    # load_profile 先于 _common_setup:profile 名字打错要在碰文件之前就炸,
    # 而不是等备份都做完了才发现。(cache 命中,apply 里再取一次是零成本)
    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))
    result = _drop_pipeline_keys(apply(doc, args, subcommand="promote-h1"))
    bak = _save_with_backup(src, doc, args, wrote_needed=result["promote_count"] > 0)

    report = {
        "subcommand": "outline promote-h1",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": args.dry_run,
        "backup": str(bak) if bak else None,
        **result,
    }
    _emit_report(report, args)
    mode = "DRY-RUN" if args.dry_run else "APPLIED"
    print(f"[outline promote-h1] {mode} promote_count={result['promote_count']}")
    for it in result["items"][:10]:
        print(f"  idx={it['para_idx']} {it['before_style']!r}→{it['after_style']!r}")
        print(f"    before: {it['before_text']!r}")
        print(f"    after : {it['after_text']!r}")
    if bak:
        print(f"backup: {bak}")
    return 0


# =============================================================================
# 子命令 2: outline demote-h2 — demote_h2_with_h3_format
# =============================================================================
_H3_TEXT_PATTERN = re.compile(r"^\d+、\s*\S")
DEMOTE_LENGTH_LIMIT = 80


def _scan_demote_candidates(doc, profile: StylesProfile) -> list[dict]:
    """扫所有 H2 段(profile-driven), 找文本是 H3 形态(阿拉伯+顿号短段)的."""
    candidates = []
    h2_set = set(profile.H2_STYLES or [])
    for idx, p in enumerate(doc.paragraphs):
        sn = p.style.name if p.style else ""
        if sn not in h2_set:
            continue
        text = p.text.strip()
        if not text or len(text) >= DEMOTE_LENGTH_LIMIT:
            continue
        if _H3_TEXT_PATTERN.match(text):
            candidates.append({
                "idx": idx,
                "text": text,
                "style_before": sn,
                "style_after": "Heading 3",   # 实际用 _resolve_heading_style 拿真 style obj
            })
    return candidates


def _apply_demotion(doc, candidates: list[dict], h3_style):
    for c in candidates:
        p = doc.paragraphs[c["idx"]]
        p.style = h3_style
        c["style_after"] = h3_style.name


def _process_demote_h2(doc, profile: StylesProfile, dry_run: bool) -> dict:
    """扫 H2 里文本形态是 H3 的段并降级。

    「文档里根本没有可用的 H3 样式」不是异常而是常态(不同院模板样式名不同),
    所以不抛,只记进 warnings 由调用方决定怎么喊 —— pipeline 里没有 stderr 可喊。
    """
    candidates = _scan_demote_candidates(doc, profile)
    h3_style = _resolve_heading_style(doc, profile, 3)

    warnings: list[str] = []
    if h3_style is None and candidates:
        warnings.append(f"no H3 style in document; profile.H3_STYLES={profile.H3_STYLES}")

    applied = bool(candidates) and not dry_run and h3_style is not None
    if applied:
        _apply_demotion(doc, candidates, h3_style)

    return {
        "candidates_count": len(candidates),
        "candidates": candidates,
        "h3_style_used": h3_style.name if h3_style else None,
        "demoted": len(candidates) if applied else 0,
        "warnings": warnings,
    }


def cmd_demote_h2(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))
    result = _drop_pipeline_keys(apply(doc, args, subcommand="demote-h2"))
    candidates = result["candidates"]
    h3_style_name = result["h3_style_used"]

    for w in result["warnings"]:
        print(f"[WARN] {w}", file=sys.stderr)

    wrote_needed = bool(candidates) and h3_style_name is not None
    bak = _save_with_backup(src, doc, args, wrote_needed=wrote_needed)

    report = {
        "subcommand": "outline demote-h2",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": args.dry_run,
        "backup": str(bak) if bak else None,
        "candidates_count": result["candidates_count"],
        "candidates": candidates,
        "h3_style_used": h3_style_name,
    }
    _emit_report(report, args)
    print(f"[outline demote-h2] candidates={len(candidates)} h3_style={h3_style_name!r}")
    for c in candidates[:10]:
        print(f"  idx={c['idx']:4d} {c['style_before']!r}→{c['style_after']!r} | {c['text']}")
    if bak:
        print(f"backup: {bak}")
    return 0


# =============================================================================
# 子命令 3: outline normalize-arabic — normalize_outline_to_arabic
# =============================================================================
_RE_H1 = re.compile(r"^([一二三四五六七八九十百零]+)、\s*(.+)$")
_RE_H2 = re.compile(r"^[（(]([一二三四五六七八九十百零]+)[）)]\s*(.+)$")
_RE_H3 = re.compile(r"^(\d+)、\s*(.+)$")

NORM_H1_MAX_LEN = 60
NORM_H2_MAX_LEN = 80
NORM_H3_MAX_LEN = 80


def _try_h1(text: str, style_name: str, h1_names: set[str]) -> Optional[tuple[int, str]]:
    if len(text) >= NORM_H1_MAX_LEN:
        # 仍可放过 — 但只在样式本来就是 H1 时
        if style_name not in h1_names:
            return None
    m = _RE_H1.match(text)
    if not m:
        return None
    try:
        num = chinese_to_arabic(m.group(1))
    except ValueError:
        return None
    rest = m.group(2).strip()
    if not rest:
        return None
    return num, rest


def _try_h2_norm(text: str) -> Optional[tuple[int, str]]:
    if len(text) >= NORM_H2_MAX_LEN:
        return None
    m = _RE_H2.match(text)
    if not m:
        return None
    try:
        num = chinese_to_arabic(m.group(1))
    except ValueError:
        return None
    rest = m.group(2).strip()
    if not rest:
        return None
    return num, rest


def _try_h3_norm(text: str) -> Optional[tuple[int, str]]:
    if len(text) >= NORM_H3_MAX_LEN:
        return None
    m = _RE_H3.match(text)
    if not m:
        return None
    try:
        num = int(m.group(1))
    except ValueError:
        return None
    rest = m.group(2).strip()
    if not rest:
        return None
    return num, rest


def _process_normalize(doc, profile: StylesProfile, dry_run: bool) -> dict:
    h1_style = _resolve_heading_style(doc, profile, 1)
    h2_style = _resolve_heading_style(doc, profile, 2)
    h3_style = _resolve_heading_style(doc, profile, 3)
    missing = []
    if h1_style is None: missing.append("H1")
    if h2_style is None: missing.append("H2")
    if h3_style is None: missing.append("H3")
    if missing:
        raise SystemExit(f"ERROR: missing styles in document: {missing}")

    h1_names = set(profile.H1_STYLES or [])
    title_names = set(profile.TITLE_STYLES or [])

    plan: list[dict] = []
    h1_count = h2_count = h3_count = 0

    for idx, p in enumerate(doc.paragraphs):
        if _is_in_table_cell(p):
            continue
        text = p.text
        if not text or not text.strip():
            continue
        text_stripped = text.strip()
        style_name = p.style.name if p.style is not None else ""
        if style_name in title_names:
            continue

        h1_hit = _try_h1(text_stripped, style_name, h1_names)
        h2_hit = _try_h2_norm(text_stripped)
        h3_hit = _try_h3_norm(text_stripped)

        if h1_hit is not None:
            num, rest = h1_hit
            h1_count += 1
            h2_count = 0
            h3_count = 0
            new_text = profile.format_heading(1, h1_count) + rest
            plan.append({
                "para_idx": idx, "level": "H1",
                "before_style": style_name, "after_style": h1_style.name,
                "detected_num": num, "assigned_num": str(h1_count),
                "before_text": text_stripped, "after_text": new_text,
            })
            if not dry_run:
                _rewrite_paragraph_runs(p, new_text)
                p.style = h1_style
            continue

        if h2_hit is not None:
            if h1_count == 0:
                continue
            num, rest = h2_hit
            h2_count += 1
            h3_count = 0
            new_text = profile.format_heading(2, h1_count, h2_count) + rest
            plan.append({
                "para_idx": idx, "level": "H2",
                "before_style": style_name, "after_style": h2_style.name,
                "detected_num": num, "assigned_num": f"{h1_count}.{h2_count}",
                "before_text": text_stripped, "after_text": new_text,
            })
            if not dry_run:
                _rewrite_paragraph_runs(p, new_text)
                p.style = h2_style
            continue

        if h3_hit is not None:
            if h1_count == 0 or h2_count == 0:
                continue
            num, rest = h3_hit
            h3_count += 1
            new_text = profile.format_heading(3, h1_count, h2_count, h3_count) + rest
            plan.append({
                "para_idx": idx, "level": "H3",
                "before_style": style_name, "after_style": h3_style.name,
                "detected_num": num, "assigned_num": f"{h1_count}.{h2_count}.{h3_count}",
                "before_text": text_stripped, "after_text": new_text,
            })
            if not dry_run:
                _rewrite_paragraph_runs(p, new_text)
                p.style = h3_style
            continue

    by_level = {"H1": 0, "H2": 0, "H3": 0}
    for it in plan:
        by_level[it["level"]] += 1
    return {
        "h1_max_len": NORM_H1_MAX_LEN,
        "h2_max_len": NORM_H2_MAX_LEN,
        "h3_max_len": NORM_H3_MAX_LEN,
        "change_count": len(plan),
        "by_level": by_level,
        "items": plan,
    }


def cmd_normalize_arabic(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))
    result = _drop_pipeline_keys(apply(doc, args, subcommand="normalize-arabic"))
    bak = _save_with_backup(src, doc, args, wrote_needed=result["change_count"] > 0)

    report = {
        "subcommand": "outline normalize-arabic",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": args.dry_run,
        "backup": str(bak) if bak else None,
        **result,
    }
    _emit_report(report, args)
    mode = "DRY-RUN" if args.dry_run else "APPLIED"
    print(f"[outline normalize-arabic] {mode} change_count={result['change_count']} "
          f"H1={result['by_level']['H1']} H2={result['by_level']['H2']} H3={result['by_level']['H3']}")
    for it in result["items"][:10]:
        print(f"  idx={it['para_idx']:>4} {it['level']} num={it['assigned_num']}")
        print(f"      before: {it['before_text']!r}")
        print(f"      after : {it['after_text']!r}")
    if bak:
        print(f"backup: {bak}")
    return 0


# =============================================================================
# pipeline adapter (sub/pipeline_lib.py 的 doc-based step 契约)
# =============================================================================
_SUBCOMMANDS = {
    "promote-h1": _process_promote_h1,
    "demote-h2": _process_demote_h2,
    "normalize-arabic": _process_normalize,
}

# apply() 为 pipeline 补的通用键。CLI 的 JSON 报告有自己的键名与键序
# (profile 排在 dry_run 前面),摘掉再拼报告,报告结构才不会因为
# 「给 pipeline 加了个键」而悄悄漂移。
_PIPELINE_KEYS = ("outline_cmd", "profile", "changed")


def _drop_pipeline_keys(result: dict) -> dict:
    for k in _PIPELINE_KEYS:
        result.pop(k, None)
    return result


def apply(doc, args=None, *, subcommand: str | None = None) -> dict:
    """在**已打开的** doc 上跑一个 outline 子动作,返回改动报告。

    本脚本一个模块带三个互斥动作,所以 apply 必须先知道跑哪个:
      - pipeline 通路:读 args.outline_cmd(与 CLI subparser 的 dest 同名,
        所以 pipeline 和 CLI 喂的是同一个 Namespace,不用两套字段);
      - CLI 通路:cmd_* 显式传 subcommand= 钉死。**不能只靠 args**,
        因为 cmd_promote_h1 收到的 Namespace 若缺 outline_cmd(被程序化调用时),
        落到默认值就会安静地跑成另一个动作。
    未知/缺省一律抛,不静默 no-op —— 那会让 pipeline 报绿却什么都没干。
    """
    cmd = subcommand or (getattr(args, "outline_cmd", None) if args else None)
    if cmd not in _SUBCOMMANDS:
        raise ValueError(
            f"outline.apply 需要子动作(args.outline_cmd 或 subcommand=),"
            f"收到 {cmd!r};可用 {'/'.join(_SUBCOMMANDS)}"
        )

    requested = getattr(args, "profile", None) if args else None
    profile = load_profile(requested)
    dry_run = bool(getattr(args, "dry_run", False)) if args else False

    result = _SUBCOMMANDS[cmd](doc, profile, dry_run)
    # 三个动作各自的计数键不同名,再给一个统一的 changed 供 pipeline 汇总
    changed = {
        "promote-h1": lambda r: r["promote_count"],
        "demote-h2": lambda r: r["demoted"],          # 无可用 H3 样式时确实一个没改
        "normalize-arabic": lambda r: r["change_count"],
    }[cmd](result)
    return {
        **result,
        "outline_cmd": cmd,
        "profile": getattr(profile, "_name", requested),
        "changed": changed,
    }


# =============================================================================
# CLI register & main
# =============================================================================
def _add_common_args(p: argparse.ArgumentParser):
    p.add_argument("docx", type=Path, help="目标 docx 路径")
    p.add_argument("--dry-run", action="store_true", help="只规划不写入")
    p.add_argument("--no-backup", action="store_true", help="跳过备份(慎用)")
    p.add_argument("--report", type=Path, default=None, help="写 JSON 报告到此路径")
    p.add_argument("--profile", type=str, default=None,
                   help="styles_registry profile (zdwp / eco-flow / generic / ...)")


def register(subparsers):
    """register 3 outline subcommands into a doctools CLI subparsers.

    嵌套命名:
        outline promote-h1
        outline demote-h2
        outline normalize-arabic
    """
    outline_p = subparsers.add_parser("outline", help="大纲层级规范化(promote-h1/demote-h2/normalize-arabic)")
    outline_sub = outline_p.add_subparsers(dest="outline_cmd", required=True)

    ph1 = outline_sub.add_parser("promote-h1", help="把伪 H1(挂 H2/H3/Title) 升为 H1 + 阿拉伯化")
    _add_common_args(ph1)
    ph1.set_defaults(func=cmd_promote_h1)

    dh2 = outline_sub.add_parser("demote-h2", help="样式=H2 但文本是 H3 形态 → 降为 H3")
    _add_common_args(dh2)
    dh2.set_defaults(func=cmd_demote_h2)

    nm = outline_sub.add_parser("normalize-arabic", help="一次识别 H1/H2/H3 文本 + 改样式 + 阿拉伯化")
    _add_common_args(nm)
    nm.set_defaults(func=cmd_normalize_arabic)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="doctools-outline",
        description="docx outline group · profile-driven 大纲层级规范化 (distilled from qual-supply)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)
    register(sub)
    args = parser.parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    sys.exit(main())
