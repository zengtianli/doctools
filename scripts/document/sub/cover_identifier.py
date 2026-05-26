"""cover_identifier.py — LLM 驱动的封面段角色识别 (W-cover 2026-05-26).

识别 docx 封面前 N 段对应的 4 个角色 (primary_title / subtitle / author / date)
在 doc.paragraphs 数组的 0-based 段 idx.

公开接口:
    identify_cover_roles(docx_path, scan_first_n=30) -> dict

  返回: {
    "primary_title_idx": int | None,
    "subtitle_idx":      int | None,
    "author_idx":        int | None,
    "date_idx":          int | None,
  }

LLM 通路: 走总部 SSOT `tools.llm_client` (`~/Dev/tools/llm_client.py`),
fork claude CLI haiku.  ⚠ 不要在 CC agent 内调本模块（会嵌套）—
本模块作 cmd_restore 在外部 standalone 跑时 LLM 路径；CC 会话内
请配合 fallback (`identify_cover_roles_heuristic`) 用。

Fallback: paragraph[3]/[4]/作者扫"院/局/所/公司" + 日期扫"年.*月" 启发式.
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document


# ─── llm_client 定位 ───────────────────────────────────────────────────
_TOOLS_DIR = Path.home() / "Dev" / "tools"
if str(_TOOLS_DIR) not in sys.path:
    sys.path.insert(0, str(_TOOLS_DIR))


def _load_llm_client():
    """import tools.llm_client (SSOT canonical at ~/Dev/tools/llm_client.py)."""
    try:
        import importlib
        return importlib.import_module("llm_client")
    except ImportError as e:
        raise ImportError(
            f"tools.llm_client 不可达 (~/Dev/tools/llm_client.py): {e}\n"
            "不要自造 httpx 调 API — 改 fallback (heuristic) 或加 PYTHONPATH"
        ) from e


# ─── docx → 前 N 个非空段 ─────────────────────────────────────────────
def _extract_front_paragraphs(docx_path: Path, n: int) -> list[tuple[int, str]]:
    """Return [(idx_in_doc.paragraphs, text), ...] for first N non-empty paras."""
    doc = Document(str(docx_path))
    out: list[tuple[int, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if not t:
            continue
        out.append((idx, t))
        if len(out) >= n:
            break
    return out


# ─── LLM prompt build + call ───────────────────────────────────────────
_PROMPT_TMPL = """以下是一份 Word 文档封面的前几段文本 (idx: 文本):

{paragraphs_block}

请识别 4 个封面角色对应的段落 idx (在原 doc.paragraphs 数组中的 0-based 位置):
- 主标题 (primary_title): 通常是地名+主题, 如"某县小型水库生态流量核定与保障"
- 副标题 (subtitle): 通常是文档类型, 如"实施方案"/"技术报告"
- 作者 (author): 编制单位, 如"XX设计研究院"/"XX局"
- 日期 (date): 时间, 如"二○二六年五月"/"2026年5月"

严格返回 JSON 格式 (无 markdown 代码块):
{{"primary_title_idx": N, "subtitle_idx": N, "author_idx": N, "date_idx": N}}

找不到某个角色时该字段写 null。"""


def _build_prompt(paras: list[tuple[int, str]]) -> str:
    block = "\n".join(f"{idx}: {t}" for idx, t in paras)
    return _PROMPT_TMPL.format(paragraphs_block=block)


_JSON_FENCE_RE = re.compile(r"^```(?:json)?\s*|\s*```$", re.MULTILINE)


def _parse_llm_json(raw: str) -> dict:
    """Strip markdown fences and parse JSON. Tolerates extra prose around the block."""
    s = (raw or "").strip()
    # remove markdown code fences
    s = _JSON_FENCE_RE.sub("", s).strip()
    # find first {...} block (greedy)
    m = re.search(r"\{[^{}]*\}", s, re.DOTALL)
    if m:
        s = m.group(0)
    try:
        return json.loads(s)
    except json.JSONDecodeError as e:
        raise ValueError(f"LLM 返回非 JSON: {raw[:200]!r}") from e


def _validate_idx(val, scan_n: int) -> Optional[int]:
    if val is None:
        return None
    try:
        i = int(val)
    except (TypeError, ValueError):
        return None
    if i < 0 or i > scan_n * 4:  # 留宽容度: idx in doc.paragraphs 不必 ≤ scan_n
        return None
    return i


# ─── public: LLM-driven identify ───────────────────────────────────────
def identify_cover_roles(docx_path: Path, scan_first_n: int = 30,
                         model: str = "haiku", timeout: int = 60) -> dict:
    """LLM-driven cover role identification (4 idx, 0-based in doc.paragraphs).

    Returns: {primary_title_idx, subtitle_idx, author_idx, date_idx} all int|None.

    Raises:
        ImportError if tools.llm_client unreachable.
        ValueError  if LLM returns un-parseable JSON.
    """
    docx_path = Path(docx_path).expanduser().resolve()
    paras = _extract_front_paragraphs(docx_path, scan_first_n)
    if not paras:
        return {"primary_title_idx": None, "subtitle_idx": None,
                "author_idx": None, "date_idx": None}

    llm = _load_llm_client()
    system = "你是一名熟悉中文政府/工程报告排版的助手。严格按要求返回 JSON。"
    message = _build_prompt(paras)
    try:
        reply = llm.chat(system, message, model=model, timeout=timeout)
    except Exception as e:
        raise RuntimeError(f"LLM call failed: {type(e).__name__}: {e}") from e

    parsed = _parse_llm_json(reply)
    scan_n = scan_first_n
    return {
        "primary_title_idx": _validate_idx(parsed.get("primary_title_idx"), scan_n),
        "subtitle_idx":      _validate_idx(parsed.get("subtitle_idx"),      scan_n),
        "author_idx":        _validate_idx(parsed.get("author_idx"),        scan_n),
        "date_idx":          _validate_idx(parsed.get("date_idx"),          scan_n),
    }


# ─── public: heuristic fallback (no LLM) ───────────────────────────────
_AUTHOR_KEY_RE = re.compile(r"(院|局|所|公司|中心|大学|学院|研究所|设计|规划|集团)$")
_DATE_RE = re.compile(r"(?:二[〇○零一二三四五六七八九]+年|\d{4}\s*年).{0,10}月")
_DOC_TYPE_RE = re.compile(
    r"(实施方案|技术报告|可行性研究报告|初步设计|工作报告|工程报告|"
    r"专项规划|规划报告|核定方案|保障方案|评估报告)$"
)


def identify_cover_roles_heuristic(docx_path: Path, scan_first_n: int = 30) -> dict:
    """Heuristic fallback when LLM unavailable.

    Rules (rough but robust on eco-flow样板):
      - primary_title: longest non-empty para in first 6 entries (>= 8 chars)
                       且其后跟一个明显是 doc-type 的短段 — 取那个长段
      - subtitle:      首个 (在 primary_title 之后) 匹配 _DOC_TYPE_RE 的段
      - author:        首个匹配 _AUTHOR_KEY_RE 的段 (≤ 30 字)
      - date:          首个匹配 _DATE_RE 的段
    """
    paras = _extract_front_paragraphs(docx_path, scan_first_n)
    if not paras:
        return {"primary_title_idx": None, "subtitle_idx": None,
                "author_idx": None, "date_idx": None}

    primary_title_idx: Optional[int] = None
    subtitle_idx: Optional[int] = None
    author_idx: Optional[int] = None
    date_idx: Optional[int] = None

    # primary_title: 前 6 非空段中最长的 (>= 8 字符) — 多数封面首屏即主标题
    head_window = paras[:6]
    if head_window:
        cand = max(head_window, key=lambda x: len(x[1]))
        if len(cand[1]) >= 8:
            primary_title_idx = cand[0]

    for idx, text in paras:
        if subtitle_idx is None and _DOC_TYPE_RE.search(text):
            if primary_title_idx is None or idx > primary_title_idx:
                subtitle_idx = idx
        if author_idx is None and len(text) <= 30 and _AUTHOR_KEY_RE.search(text):
            author_idx = idx
        if date_idx is None and _DATE_RE.search(text):
            date_idx = idx

    return {
        "primary_title_idx": primary_title_idx,
        "subtitle_idx":      subtitle_idx,
        "author_idx":        author_idx,
        "date_idx":          date_idx,
    }


# ─── CLI (standalone test) ─────────────────────────────────────────────
def _cli(argv: list[str] | None = None) -> int:
    import argparse
    ap = argparse.ArgumentParser(
        prog="cover_identifier",
        description="Identify docx cover role paragraph idx (LLM + heuristic fallback)",
    )
    ap.add_argument("docx", type=Path, help="target docx path")
    ap.add_argument("--no-llm", action="store_true", help="use heuristic only")
    ap.add_argument("--scan", type=int, default=30, help="scan first N non-empty paragraphs")
    args = ap.parse_args(argv)

    docx = args.docx.expanduser().resolve()
    if not docx.exists():
        print(f"[ERR] not found: {docx}", file=sys.stderr)
        return 2

    if args.no_llm:
        result = identify_cover_roles_heuristic(docx, scan_first_n=args.scan)
        method = "heuristic"
    else:
        try:
            result = identify_cover_roles(docx, scan_first_n=args.scan)
            method = "llm"
        except (ImportError, RuntimeError, ValueError) as e:
            print(f"[WARN] LLM unavailable, falling back to heuristic: {e}", file=sys.stderr)
            result = identify_cover_roles_heuristic(docx, scan_first_n=args.scan)
            method = "heuristic (LLM failed)"

    print(json.dumps({"method": method, "result": result},
                     ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(_cli())
