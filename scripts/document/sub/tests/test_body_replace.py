"""tests for body_replace.py"""
from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest
from docx import Document

_HERE = Path(__file__).resolve().parent
_SUB_DIR = _HERE.parent

# Load module directly without requiring package context.
spec = importlib.util.spec_from_file_location(
    "body_replace_under_test", str(_SUB_DIR / "body_replace.py")
)
br = importlib.util.module_from_spec(spec)
sys.modules["body_replace_under_test"] = br
assert spec.loader is not None
spec.loader.exec_module(br)


def _make_shell(path: Path) -> None:
    """Shell: cover line + H1 'shell title' + shell body para."""
    doc = Document()
    doc.add_paragraph("Cover line — should be kept")
    doc.add_heading("shell title", level=1)
    doc.add_paragraph("shell content (should be dropped by body_replace)")
    doc.save(str(path))


def _make_content(path: Path) -> None:
    """Content: H1 'content title' + 2 paras (1 with custom style)."""
    doc = Document()
    doc.add_heading("content title", level=1)
    doc.add_paragraph("new para 1")
    p = doc.add_paragraph("new para 2 (custom style)")
    # custom style ref → will trigger fallback warning when applied to shell
    # We do it by directly poking pStyle
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pStyle = pPr.makeelement(qn("w:pStyle"), {qn("w:val"): "NonexistentStyle"})
    pPr.append(pStyle)
    doc.save(str(path))


def test_keep_shell_h1_default(tmp_path: Path):
    shell = tmp_path / "shell.docx"
    content = tmp_path / "content.docx"
    out = tmp_path / "out.docx"
    _make_shell(shell)
    _make_content(content)

    stats = br.body_replace(shell, content, out, keep_shell_h1=True)
    assert out.exists()

    d = Document(str(out))
    paras = d.paragraphs
    # Should contain: cover, H1 'shell title', new para 1, new para 2
    texts = [p.text for p in paras]
    assert "Cover line — should be kept" in texts
    # First H1 should be 'shell title', NOT 'content title'
    h1s = [p for p in paras if p.style.name in br.H1_STYLES]
    assert h1s, "expected at least one H1 in output"
    assert h1s[0].text == "shell title", f"expected 'shell title', got {h1s[0].text!r}"
    # content title must be dropped
    assert "content title" not in texts
    # new paras must be present
    assert "new para 1" in texts
    assert "new para 2 (custom style)" in texts
    # shell body content (post-H1) must be removed
    assert "shell content (should be dropped by body_replace)" not in texts


def test_no_keep_shell_h1(tmp_path: Path):
    shell = tmp_path / "shell.docx"
    content = tmp_path / "content.docx"
    out = tmp_path / "out.docx"
    _make_shell(shell)
    _make_content(content)

    stats = br.body_replace(shell, content, out, keep_shell_h1=False)
    assert out.exists()

    d = Document(str(out))
    paras = d.paragraphs
    texts = [p.text for p in paras]
    h1s = [p for p in paras if p.style.name in br.H1_STYLES]
    # With keep_shell_h1=False, shell body is fully dropped → first H1 is content's
    assert h1s, "expected at least one H1 from content"
    assert h1s[0].text == "content title"
    # Cover line is part of shell body (before first H1) → should also be dropped
    assert "Cover line — should be kept" not in texts


def test_dry_run_does_not_create_output(tmp_path: Path, capsys):
    shell = tmp_path / "shell.docx"
    content = tmp_path / "content.docx"
    out = tmp_path / "out.docx"
    _make_shell(shell)
    _make_content(content)

    saved = sys.argv[:]
    sys.argv = [
        "body_replace.py",
        "--shell", str(shell),
        "--content", str(content),
        "--out", str(out),
        "--dry-run",
    ]
    try:
        rc = br.main()
    finally:
        sys.argv = saved
    assert rc == 0
    assert not out.exists(), "dry-run must not create output file"
    captured = capsys.readouterr()
    assert "DRY RUN" in captured.out


def test_style_fallback_warning(tmp_path: Path, capsys):
    """content references NonexistentStyle → should fallback to Normal + warn."""
    shell = tmp_path / "shell.docx"
    content = tmp_path / "content.docx"
    out = tmp_path / "out.docx"
    _make_shell(shell)
    _make_content(content)

    stats = br.body_replace(shell, content, out, keep_shell_h1=True)
    captured = capsys.readouterr()
    # WARN should land on stderr
    assert "NonexistentStyle" in captured.err or stats["style_warnings"] >= 1
    assert stats["style_warnings"] >= 1


def test_main_smoke(tmp_path: Path):
    """Drive via main() entry point with argparse."""
    shell = tmp_path / "shell.docx"
    content = tmp_path / "content.docx"
    out = tmp_path / "out.docx"
    _make_shell(shell)
    _make_content(content)

    saved = sys.argv[:]
    sys.argv = [
        "body_replace.py",
        "--shell", str(shell),
        "--content", str(content),
        "--out", str(out),
    ]
    try:
        rc = br.main()
    finally:
        sys.argv = saved
    assert rc == 0
    assert out.exists()
    d = Document(str(out))
    h1s = [p for p in d.paragraphs if p.style.name in br.H1_STYLES]
    assert h1s[0].text == "shell title"


def test_missing_input_returns_error(tmp_path: Path):
    out = tmp_path / "out.docx"
    saved = sys.argv[:]
    sys.argv = [
        "body_replace.py",
        "--shell", str(tmp_path / "nope.docx"),
        "--content", str(tmp_path / "nope2.docx"),
        "--out", str(out),
    ]
    try:
        rc = br.main()
    finally:
        sys.argv = saved
    assert rc == 2
