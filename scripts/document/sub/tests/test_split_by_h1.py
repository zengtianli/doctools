"""tests for split_by_h1.py"""
from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest
from docx import Document

_HERE = Path(__file__).resolve().parent
_SUB_DIR = _HERE.parent

# Load module directly so test runs without requiring scripts to be a package.
spec = importlib.util.spec_from_file_location(
    "split_by_h1_under_test", str(_SUB_DIR / "split_by_h1.py")
)
sbh = importlib.util.module_from_spec(spec)
sys.modules["split_by_h1_under_test"] = sbh
assert spec.loader is not None
spec.loader.exec_module(sbh)


def _make_docx(path: Path) -> None:
    """Build minimal docx: H1×2 + 3 body paragraphs + 1 frontmatter line."""
    doc = Document()
    # 1 normal paragraph as frontmatter (before first H1)
    doc.add_paragraph("Front cover line")
    # H1 #1
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("body 1.1")
    doc.add_paragraph("body 1.2")
    # H1 #2
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("body 2.1")
    doc.save(str(path))


def test_sanitize_filename():
    assert sbh.sanitize_filename("foo/bar:baz") == "foo_bar_baz"
    assert sbh.sanitize_filename("  hello   world  ") == "hello world"
    assert sbh.sanitize_filename("") == "untitled"
    assert sbh.sanitize_filename("   ") == "untitled"
    long = "a" * 200
    assert len(sbh.sanitize_filename(long)) == 100
    # illegal chars covered
    assert sbh.sanitize_filename('a*b?c"d<e>f|g\\h/i:j') == "a_b_c_d_e_f_g_h_i_j"


def test_split_dry_run(tmp_path: Path, capsys):
    src = tmp_path / "in.docx"
    _make_docx(src)
    out = tmp_path / "out"
    rc = sbh.main.__wrapped__ if hasattr(sbh.main, "__wrapped__") else None
    # Use argparse path
    saved = sys.argv[:]
    sys.argv = [
        "split_by_h1.py",
        "--docx", str(src),
        "--out-dir", str(out),
        "--include-frontmatter",
        "--dry-run",
    ]
    try:
        rc = sbh.main()
    finally:
        sys.argv = saved
    assert rc == 0
    # No files should have been written
    assert not out.exists() or not any(out.iterdir())
    captured = capsys.readouterr().out
    assert "DRY RUN" in captured
    assert "Chapter One" in captured
    assert "Chapter Two" in captured


def test_split_without_frontmatter(tmp_path: Path):
    src = tmp_path / "in.docx"
    _make_docx(src)
    out = tmp_path / "out"
    saved = sys.argv[:]
    sys.argv = [
        "split_by_h1.py",
        "--docx", str(src),
        "--out-dir", str(out),
    ]
    try:
        rc = sbh.main()
    finally:
        sys.argv = saved
    assert rc == 0
    files = sorted(out.glob("*.docx"))
    # 2 H1 → 2 files
    assert len(files) == 2
    # All openable + correct titles
    titles = []
    for f in files:
        d = Document(str(f))
        # First paragraph in each slice should be the H1
        first_heading = next((p.text for p in d.paragraphs if p.style.name in {"Heading 1", "标题 1"}), None)
        if first_heading:
            titles.append(first_heading)
    assert "Chapter One" in titles
    assert "Chapter Two" in titles


def test_split_with_frontmatter(tmp_path: Path):
    src = tmp_path / "in.docx"
    _make_docx(src)
    out = tmp_path / "out"
    saved = sys.argv[:]
    sys.argv = [
        "split_by_h1.py",
        "--docx", str(src),
        "--out-dir", str(out),
        "--include-frontmatter",
    ]
    try:
        rc = sbh.main()
    finally:
        sys.argv = saved
    assert rc == 0
    files = sorted(out.glob("*.docx"))
    # frontmatter + 2 H1 = 3 files
    assert len(files) == 3
    # frontmatter file (idx 0) should contain "Front cover line"
    frontmatter_file = files[0]
    d = Document(str(frontmatter_file))
    all_text = " ".join(p.text for p in d.paragraphs)
    assert "Front cover line" in all_text
    # ... and should NOT contain Chapter One heading
    assert "Chapter One" not in all_text


def test_illegal_filename_in_h1_title(tmp_path: Path):
    src = tmp_path / "in.docx"
    doc = Document()
    doc.add_heading("Bad/Title:With*Illegal?Chars", level=1)
    doc.add_paragraph("body")
    doc.save(str(src))
    out = tmp_path / "out"
    saved = sys.argv[:]
    sys.argv = [
        "split_by_h1.py",
        "--docx", str(src),
        "--out-dir", str(out),
    ]
    try:
        rc = sbh.main()
    finally:
        sys.argv = saved
    assert rc == 0
    files = sorted(out.glob("*.docx"))
    assert len(files) == 1
    name = files[0].name
    for bad in '/\\:*?"<>|':
        assert bad not in name, f"illegal char {bad!r} survived in filename {name!r}"
