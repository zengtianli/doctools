"""tests for fix_superscript_refs.py — apply(doc, args) 最小回归。

覆盖三件事：
  1. 真改：[1] / [3-4] 被拆成独立 run 且设了 w:vertAlign=superscript；
  2. 参考文献列表行（行首 [n]）被跳过；
  3. dry-run：计数与真改一致、details 给引用列表、且不动树。
"""
from __future__ import annotations

import argparse
import importlib.util
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

_HERE = Path(__file__).resolve().parent
_SUB_DIR = _HERE.parent

# Load module directly without requiring package context.
spec = importlib.util.spec_from_file_location(
    "fix_superscript_refs_under_test", str(_SUB_DIR / "fix_superscript_refs.py")
)
fsr = importlib.util.module_from_spec(spec)
sys.modules["fix_superscript_refs_under_test"] = fsr
assert spec.loader is not None
spec.loader.exec_module(fsr)


def _make_doc(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("正文引用见文献[1]，另见[3-4]的论述。")
    doc.add_paragraph("[1] 张三. 某研究[J]. 某学报, 2020.")  # 参考文献行，须跳过
    doc.add_paragraph("没有引用的普通段落。")
    doc.save(str(path))


def _superscript_ref_texts(doc) -> list[str]:
    out = []
    for para in doc.paragraphs:
        for r in para._element.iter(qn("w:r")):
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                continue
            vert = rPr.find(qn("w:vertAlign"))
            if vert is None or vert.get(qn("w:val")) != "superscript":
                continue
            t = r.find(qn("w:t"))
            if t is not None and t.text:
                out.append(t.text)
    return out


def test_apply_superscripts_refs_and_skips_ref_list(tmp_path):
    p = tmp_path / "in.docx"
    _make_doc(p)
    doc = Document(str(p))

    rep = fsr.apply(doc, argparse.Namespace(dry_run=False))

    assert rep["paragraphs"] == 1          # 只有正文段被改；文献行跳过
    assert rep["refs"] == 2                # [1] 与 [3-4]
    assert rep["changed"] == 2
    assert sorted(_superscript_ref_texts(doc)) == ["[1]", "[3-4]"]


def test_apply_dry_run_counts_match_and_tree_untouched(tmp_path):
    p = tmp_path / "in.docx"
    _make_doc(p)
    doc = Document(str(p))

    rep = fsr.apply(doc, argparse.Namespace(dry_run=True))

    assert rep["dry_run"] is True
    assert rep["paragraphs"] == 1
    assert rep["refs"] == 2
    # dry-run 的 details.refs 是命中的引用列表（真改时是计数）——沿用原行为
    assert rep["details"][0]["refs"] == ["[1]", "[3-4]"]
    # 不动树：没有任何 run 被设为上标
    assert _superscript_ref_texts(doc) == []
