#!/usr/bin/env python3
# distilled from qual-supply/scripts/add_header_footer.py (2026-05-25 W1)
"""add_header_footer.py — 给 docx 所有 section 加水利院标准格式页眉/页脚.

水利院标准:
  - 页眉: 报告/项目名 (右对齐, 字号 10pt = sz 20 半磅)
  - 页脚: 院名 + 空格 + PAGE 字段 (居中对齐, 字号 10pt)

实现:
  python-docx section.header / section.footer 接口 (不裸解 zip 改 XML),
  操作所有 section 统一加, 不分首页/奇偶/封面特殊 (用户后续 Word 里可自调).

接口:
  python3 add_header_footer.py <docx> \\
    --header "<页眉文字>" \\
    --footer-prefix "<院名>" \\
    [--page-number] [--font-size 10] [--gap-spaces 13] \\
    [--dry-run] [--no-backup] [--report <json>]

默认 真改 + 自动备份 .bak-N-<date>.docx.

复制实例 (改本项目主报告):
  python3 scripts/add_header_footer.py docs/分质供水机制政策研究报告-整合版-2026-05-24.docx \\
    --header "分质供水机制政策研究报告" \\
    --footer-prefix "浙江省水利水电勘测设计院有限责任公司" \\
    --page-number
"""
from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import sys
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def lsof_check(docx_path: Path) -> Optional[str]:
    try:
        out = subprocess.run(
            ["lsof", "--", str(docx_path)],
            capture_output=True, text=True, timeout=5,
        )
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return None
    if out.returncode == 0 and out.stdout.strip():
        lines = out.stdout.strip().split("\n")
        if len(lines) > 1:
            return "\n".join(lines)
    return None


def clear_paragraph(p) -> None:
    """清空段内所有 run / sdt, 保留 pPr."""
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)


def add_run_with_font(p, text: str, font_size_pt: float):
    run = p.add_run(text)
    run.font.size = Pt(font_size_pt)
    # 显式设 East-Asian font 字号 (确保中文也按 sz 显示)
    rPr = run._r.get_or_add_rPr()
    sz = rPr.find(qn("w:sz"))
    szCs = rPr.find(qn("w:szCs"))
    half_pt = int(font_size_pt * 2)
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), str(half_pt))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(half_pt))
    return run


def add_page_field(run) -> None:
    """在 run 后追加 PAGE 字段 (3 元素: fldChar begin + instrText + fldChar end)."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    instr.set(qn("xml:space"), "preserve")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_header(section, text: str, font_size_pt: float, alignment: str = "right") -> None:
    """设 section 页眉为单行文字."""
    header = section.header
    header.is_linked_to_previous = False
    # 确保有至少 1 段
    if not header.paragraphs:
        header.add_paragraph()
    para = header.paragraphs[0]
    clear_paragraph(para)
    # 对齐
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.RIGHT)
    if text:
        add_run_with_font(para, text, font_size_pt)
    # 删多余段 (header 模板有时自带额外空段)
    for extra in header.paragraphs[1:]:
        try:
            extra._p.getparent().remove(extra._p)
        except Exception:
            pass


def set_footer(section, prefix: str, with_page: bool,
               gap_spaces: int, font_size_pt: float,
               alignment: str = "center") -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    if not footer.paragraphs:
        footer.add_paragraph()
    para = footer.paragraphs[0]
    clear_paragraph(para)
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
    if prefix:
        add_run_with_font(para, prefix, font_size_pt)
    if with_page:
        # 中间空格分隔
        if prefix:
            add_run_with_font(para, " " * gap_spaces, font_size_pt)
        page_run = add_run_with_font(para, "", font_size_pt)
        add_page_field(page_run)
    for extra in footer.paragraphs[1:]:
        try:
            extra._p.getparent().remove(extra._p)
        except Exception:
            pass


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("docx", help="输入 docx 路径")
    ap.add_argument("--header", required=True, help="页眉文字")
    ap.add_argument("--footer-prefix", required=True, help="页脚前缀 (院名)")
    ap.add_argument("--page-number", action="store_true",
                    help="页脚加 PAGE 字段 (自动页码)")
    ap.add_argument("--font-size", type=float, default=10.0,
                    help="字号 pt (默认 10 = 5号; ref 水利院实测 sz=20 半磅 = 10pt)")
    ap.add_argument("--gap-spaces", type=int, default=13,
                    help="页脚院名与页码之间空格数 (默认 13)")
    ap.add_argument("--header-align", default="right",
                    choices=["left", "center", "right"],
                    help="页眉对齐 (默认 right)")
    ap.add_argument("--footer-align", default="center",
                    choices=["left", "center", "right"],
                    help="页脚对齐 (默认 center)")
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--no-backup", action="store_true")
    ap.add_argument("--report", help="结果 JSON 输出路径")
    args = ap.parse_args(argv)

    src = Path(args.docx)
    if not src.exists():
        print(f"[error] 找不到 {src}", file=sys.stderr)
        return 2

    if not args.dry_run:
        lsof = lsof_check(src)
        if lsof:
            print(f"[error] docx 被进程占用 (关 Word/WPS 后重试):\n{lsof}",
                  file=sys.stderr)
            return 3

    backup_path = None
    if not args.dry_run and not args.no_backup:
        today = date.today().isoformat()
        n = 1
        while True:
            cand = src.with_name(f"{src.stem}.bak-{n}-{today}{src.suffix}")
            if not cand.exists():
                break
            n += 1
        shutil.copy2(src, cand)
        backup_path = cand
        print(f"[backup] {cand.name}")

    doc = Document(str(src))
    sections = doc.sections
    print(f"[info] section 数: {len(sections)}")
    report_sections = []
    for i, sec in enumerate(sections):
        set_header(sec, args.header, args.font_size, args.header_align)
        set_footer(sec, args.footer_prefix, args.page_number,
                   args.gap_spaces, args.font_size, args.footer_align)
        report_sections.append({
            "index": i,
            "header": args.header,
            "header_align": args.header_align,
            "footer_prefix": args.footer_prefix,
            "footer_align": args.footer_align,
            "page_number": args.page_number,
            "font_size_pt": args.font_size,
        })
        action = "would-set" if args.dry_run else "set"
        print(f"  section {i}: {action} header={args.header[:40]!r}  "
              f"footer={args.footer_prefix[:40]!r}{'  +PAGE' if args.page_number else ''}")

    if not args.dry_run:
        doc.save(str(src))
        print(f"[saved] {src}")
    else:
        print("[dry-run] 不落盘")

    out = {
        "docx_path": str(src),
        "backup": str(backup_path) if backup_path else None,
        "dry_run": args.dry_run,
        "sections": report_sections,
    }
    if args.report:
        Path(args.report).parent.mkdir(parents=True, exist_ok=True)
        Path(args.report).write_text(json.dumps(out, indent=2, ensure_ascii=False))
        print(f"[report] {args.report}")
    return 0


# ---------------- pipeline adapter (2026-05-25 · _pipeline_lib) ----------------
def apply(doc, args=None) -> dict:
    """pipeline-compatible: 在内存 doc 上施加 header/footer.

    args 可含 header/footer_prefix/page_number/font_size/gap_spaces/header_align/footer_align.
    未传则用默认值 (header='', footer_prefix='浙江省水利水电勘测设计院', page_number=True).
    """
    header = getattr(args, "header", "") if args else ""
    footer_prefix = getattr(args, "footer_prefix", "浙江省水利水电勘测设计院") if args else "浙江省水利水电勘测设计院"
    page_number = getattr(args, "page_number", True) if args else True
    font_size = getattr(args, "font_size", 10.0) if args else 10.0
    gap_spaces = getattr(args, "gap_spaces", 13) if args else 13
    header_align = getattr(args, "header_align", "right") if args else "right"
    footer_align = getattr(args, "footer_align", "center") if args else "center"
    n_secs = 0
    for sec in doc.sections:
        set_header(sec, header, font_size, header_align)
        set_footer(sec, footer_prefix, page_number, gap_spaces, font_size, footer_align)
        n_secs += 1
    return {"changed": n_secs, "sections": n_secs}


if __name__ == "__main__":
    sys.exit(main())
