#!/usr/bin/env python3
# distilled from eco-flow/taizhou-天台 table extract need (2026-05-28)
r"""extract_tables.py — 从 docx 抽出每张表为独立 docx (最小骨架策略).

文件名 = 邻近 caption 段文字 (如 "表2.2-1-水库基本情况.docx") 或 fallback `table-XX`.

**实现策略 (2026-05-28 重写)**:
旧策略「shutil.copy 源 docx → 删 body 非目标元素」在部分 case 上 zip 损坏 / media
瘦身未生效 (54MB), 已废.

新策略「构造最小 docx 骨架」:
    target.docx
    ├── [Content_Types].xml       (静态模板)
    ├── _rels/.rels               (静态模板, 只关联 word/document.xml)
    └── word/
        ├── _rels/document.xml.rels  (静态模板, styles/numbering/theme 3 rel)
        ├── document.xml             (新构造: 源 namespace + <w:tbl> + sectPr)
        ├── styles.xml               (从源 docx 原样复制)
        ├── numbering.xml            (从源 docx 原样复制, 可缺)
        └── theme/theme1.xml         (从源 docx 原样复制, 可缺)

**不复制**: word/media/* · header* · footer* · comments* · footnotes · endnotes ·
fontTable · settings · webSettings · customXml · docProps · embeddings.

inline image (w:drawing/w:pict) → 保 XML 引用, **不**带 media, Word 显示断链占位符.

CLI:
    python3 scripts/document/sub/extract_tables.py \\
        --docx <path> --out-dir <dir> \\
        [--name-pattern '{stem}.docx'] [--dry-run]
"""
from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
except ImportError:
    print("ERROR: python-docx / lxml 未安装 (pip install python-docx lxml)", file=sys.stderr)
    sys.exit(2)


_ILLEGAL_FILENAME_RE = re.compile(r'[/\\:*?"<>|\r\n\t]')
_MULTI_WS_RE = re.compile(r"\s+")

# Caption heuristic: starts with 表 / Table; length < 80 chars
_CAPTION_PREFIX_RE = re.compile(r"^\s*(?:表|Table\b)", re.IGNORECASE)
_CAPTION_MAX_LEN = 80

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ---------- static skeleton templates ----------

_CONTENT_TYPES_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
  <Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>
</Types>
"""

_CONTENT_TYPES_XML_NO_NUMBERING = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>
</Types>
"""

_CONTENT_TYPES_XML_NO_THEME = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
</Types>
"""

_CONTENT_TYPES_XML_MIN = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>
"""

_ROOT_RELS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""


def _build_doc_rels(has_numbering: bool, has_theme: bool) -> str:
    parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">',
        '  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>',
    ]
    next_id = 2
    if has_numbering:
        parts.append(
            f'  <Relationship Id="rId{next_id}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" '
            'Target="numbering.xml"/>'
        )
        next_id += 1
    if has_theme:
        parts.append(
            f'  <Relationship Id="rId{next_id}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" '
            'Target="theme/theme1.xml"/>'
        )
    parts.append('</Relationships>')
    return "\n".join(parts) + "\n"


def _pick_content_types(has_numbering: bool, has_theme: bool) -> str:
    if has_numbering and has_theme:
        return _CONTENT_TYPES_XML
    if has_theme and not has_numbering:
        return _CONTENT_TYPES_XML_NO_NUMBERING
    if has_numbering and not has_theme:
        return _CONTENT_TYPES_XML_NO_THEME
    return _CONTENT_TYPES_XML_MIN


# ---------- filename helpers ----------

def sanitize_filename(name: str, max_len: int = 100) -> str:
    """Replace illegal filename chars with _, compress whitespace, strip, truncate."""
    if not name:
        return "untitled"
    n = _ILLEGAL_FILENAME_RE.sub("_", name)
    n = _MULTI_WS_RE.sub(" ", n).strip()
    if not n:
        return "untitled"
    if len(n) > max_len:
        n = n[:max_len].rstrip()
    return n


def _paragraph_text(p_elem) -> str:
    """Extract concatenated text from <w:t> nodes under a <w:p>."""
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def _is_caption_like(text: str) -> bool:
    """True if text looks like a table caption (starts with 表/Table, short)."""
    if not text:
        return False
    t = text.strip()
    if not t or len(t) >= _CAPTION_MAX_LEN:
        return False
    return bool(_CAPTION_PREFIX_RE.match(t))


def _find_caption(children: list, tbl_idx: int) -> Optional[str]:
    """Scan back up to 2 paragraphs, else forward 1, for a caption-like <w:p>."""
    seen_paras = 0
    for i in range(tbl_idx - 1, -1, -1):
        elem = children[i]
        if elem.tag != qn("w:p"):
            continue
        text = _paragraph_text(elem).strip()
        if not text:
            continue
        seen_paras += 1
        if _is_caption_like(text):
            return text
        if seen_paras >= 2:
            break

    for i in range(tbl_idx + 1, len(children)):
        elem = children[i]
        if elem.tag == qn("w:tbl"):
            break
        if elem.tag != qn("w:p"):
            continue
        text = _paragraph_text(elem).strip()
        if not text:
            continue
        if _is_caption_like(text):
            return text
        break

    return None


# ---------- planning ----------

def plan_extracts(docx_path: Path, doc=None) -> tuple[list[dict], int]:
    """Inspect docx → return (extracts, tbl_count).

    extracts: list[{idx, caption, stem, tbl_idx_in_body, caption_idx_in_body}]
    """
    if doc is None:
        doc = Document(str(docx_path))
    body = doc.element.body
    children = list(body)

    sect_idx = len(children)
    for i in range(len(children) - 1, -1, -1):
        if children[i].tag == qn("w:sectPr"):
            sect_idx = i
            break

    extracts: list[dict] = []
    idx_counter = 0
    used_stems: dict[str, int] = {}

    for i in range(sect_idx):
        elem = children[i]
        if elem.tag != qn("w:tbl"):
            continue
        caption = _find_caption(children, i)
        if caption:
            stem_raw = caption
        else:
            stem_raw = f"table-{idx_counter:02d}"
        stem = sanitize_filename(stem_raw)
        base_stem = stem
        n = used_stems.get(base_stem, 0)
        if n > 0:
            stem = f"{base_stem}-{n + 1}"
        used_stems[base_stem] = n + 1

        caption_idx = -1
        if caption:
            seen = 0
            for j in range(i - 1, -1, -1):
                e2 = children[j]
                if e2.tag != qn("w:p"):
                    continue
                t2 = _paragraph_text(e2).strip()
                if not t2:
                    continue
                seen += 1
                if t2 == caption.strip():
                    caption_idx = j
                    break
                if seen >= 2:
                    break
            if caption_idx == -1:
                for j in range(i + 1, len(children)):
                    e2 = children[j]
                    if e2.tag == qn("w:tbl"):
                        break
                    if e2.tag != qn("w:p"):
                        continue
                    t2 = _paragraph_text(e2).strip()
                    if not t2:
                        continue
                    if t2 == caption.strip():
                        caption_idx = j
                    break

        extracts.append({
            "idx": idx_counter,
            "caption": caption,
            "stem": stem,
            "tbl_idx_in_body": i,
            "caption_idx_in_body": caption_idx,
        })
        idx_counter += 1

    return extracts, idx_counter


# ---------- minimal-skeleton writer ----------

def _read_source_parts(src_docx: Path) -> dict:
    """Read needed parts from source docx zip (single open).

    Returns dict with keys: document_xml_root (lxml Element of <w:document>),
        sect_pr_xml (bytes or None), styles_xml (bytes), numbering_xml (bytes or None),
        theme_xml (bytes or None).
    """
    with zipfile.ZipFile(str(src_docx)) as z:
        names = set(z.namelist())
        doc_bytes = z.read("word/document.xml")
        styles_bytes = z.read("word/styles.xml") if "word/styles.xml" in names else None
        numbering_bytes = z.read("word/numbering.xml") if "word/numbering.xml" in names else None
        theme_bytes = z.read("word/theme/theme1.xml") if "word/theme/theme1.xml" in names else None

    # Parse document.xml; keep root element for namespace declarations + sectPr extraction.
    root = etree.fromstring(doc_bytes)
    # Locate body/sectPr (last child of body)
    body = root.find(qn("w:body"))
    sect_pr = None
    if body is not None:
        for child in reversed(list(body)):
            if child.tag == qn("w:sectPr"):
                sect_pr = child
                break

    return {
        "doc_root": root,
        "doc_body": body,
        "sect_pr": sect_pr,
        "styles": styles_bytes,
        "numbering": numbering_bytes,
        "theme": theme_bytes,
    }


def _build_minimal_document_xml(
    doc_root,
    tbl_elem,
    caption_elem,
    sect_pr,
) -> bytes:
    """Construct minimal word/document.xml with given table + optional caption + sectPr.

    Re-uses root namespace declarations of source document.xml so any
    namespaced attrs on the copied <w:tbl> still resolve.
    """
    # Build new root with same nsmap as source.
    nsmap = dict(doc_root.nsmap)
    new_root = etree.Element(qn("w:document"), nsmap=nsmap)
    # Preserve mc:Ignorable etc. attributes from source root.
    for k, v in doc_root.attrib.items():
        new_root.set(k, v)
    new_body = etree.SubElement(new_root, qn("w:body"))

    # Deep-copy caption paragraph (if any) and table, sectPr.
    from copy import deepcopy
    if caption_elem is not None:
        new_body.append(deepcopy(caption_elem))
    new_body.append(deepcopy(tbl_elem))
    if sect_pr is not None:
        new_body.append(deepcopy(sect_pr))
    else:
        # Minimal sectPr fallback
        etree.SubElement(new_body, qn("w:sectPr"))

    return b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + etree.tostring(
        new_root, xml_declaration=False, encoding="utf-8"
    )


def _has_inline_image(tbl_elem) -> bool:
    """Detect inline images (drawing or pict) inside a table."""
    drawing_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    pict_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
    for el in tbl_elem.iter():
        if el.tag == drawing_tag or el.tag == pict_tag:
            return True
    return False


def write_extract_minimal(
    src_parts: dict,
    dst_docx: Path,
    tbl_idx: int,
    caption_idx: int,
    body_children_cache: list,
) -> tuple[int, int, bool]:
    """Write a minimal docx containing the target table + caption + sectPr only.

    Returns (tbl_count_in_output, file_bytes, has_inline_image).
    """
    dst_docx.parent.mkdir(parents=True, exist_ok=True)

    tbl_elem = body_children_cache[tbl_idx]
    caption_elem = body_children_cache[caption_idx] if caption_idx >= 0 else None
    inline_img = _has_inline_image(tbl_elem)

    doc_xml = _build_minimal_document_xml(
        src_parts["doc_root"], tbl_elem, caption_elem, src_parts["sect_pr"]
    )

    has_numbering = src_parts["numbering"] is not None
    has_theme = src_parts["theme"] is not None

    content_types = _pick_content_types(has_numbering, has_theme)
    doc_rels = _build_doc_rels(has_numbering, has_theme)

    # Write zip.
    with zipfile.ZipFile(str(dst_docx), "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", _ROOT_RELS_XML)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
        z.writestr("word/document.xml", doc_xml)
        if src_parts["styles"] is not None:
            z.writestr("word/styles.xml", src_parts["styles"])
        else:
            # Should never hit (Word docs always have styles.xml) but degrade safely.
            z.writestr(
                "word/styles.xml",
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
            )
        if has_numbering:
            z.writestr("word/numbering.xml", src_parts["numbering"])
        if has_theme:
            z.writestr("word/theme/theme1.xml", src_parts["theme"])

    # Verify with python-docx.
    try:
        doc2 = Document(str(dst_docx))
        tbl_count = len(doc2.tables)
    except Exception:
        tbl_count = 0

    file_bytes = dst_docx.stat().st_size
    return tbl_count, file_bytes, inline_img


# ---------- runner / CLI ----------

def run_extract(
    src_docx: Path,
    out_dir: Path,
    dry_run: bool = False,
    name_pattern: str = "{stem}.docx",
    doc=None,
) -> dict:
    """Execute table-extract; reuses provided `doc` for planning if given."""
    src = Path(src_docx).expanduser().resolve()
    if not src.exists():
        return {"error": f"input docx not found: {src}", "exit_code": 2}
    out_dir = Path(out_dir).expanduser().resolve()

    extracts, tbl_count = plan_extracts(src, doc=doc)

    if tbl_count == 0:
        return {"tbl_count": 0, "extracts_emitted": 0, "note": "no tables in docx"}

    if dry_run:
        plan = []
        for e in extracts:
            fname = name_pattern.format(stem=e["stem"], idx=e["idx"])
            plan.append({
                "idx": e["idx"], "fname": fname,
                "caption": e["caption"], "tbl_idx": e["tbl_idx_in_body"],
            })
        return {"tbl_count": tbl_count, "extracts_planned": plan, "dry_run": True}

    out_dir.mkdir(parents=True, exist_ok=True)

    # Read source parts ONCE (avoid re-opening zip per table).
    src_parts = _read_source_parts(src)
    body_children = list(src_parts["doc_body"]) if src_parts["doc_body"] is not None else []

    emitted = []
    failed = []
    inline_img_count = 0
    for e in extracts:
        fname = name_pattern.format(stem=e["stem"], idx=e["idx"])
        dst = out_dir / fname
        try:
            tc, nbytes, has_img = write_extract_minimal(
                src_parts, dst, e["tbl_idx_in_body"], e["caption_idx_in_body"],
                body_children,
            )
            if has_img:
                inline_img_count += 1
            emitted.append({
                "idx": e["idx"], "fname": fname,
                "tables": tc, "bytes": nbytes,
                "caption": e["caption"],
                "inline_image": has_img,
            })
        except Exception as exc:
            failed.append({
                "idx": e["idx"], "fname": fname,
                "error": f"{type(exc).__name__}: {exc}",
            })
    return {
        "tbl_count": tbl_count,
        "extracts_emitted": len(emitted),
        "extracts_failed": len(failed),
        "inline_image_count": inline_img_count,
        "out_dir": str(out_dir),
        "emitted": emitted,
        "failed": failed,
    }


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Extract each table in a DOCX as an independent minimal DOCX "
                    "(filename = neighboring caption text, or table-XX fallback).",
    )
    ap.add_argument("--docx", required=True, help="input docx path")
    ap.add_argument("--out-dir", required=True, help="output directory (mkdir -p)")
    ap.add_argument(
        "--name-pattern",
        default="{stem}.docx",
        help="output filename pattern, default '{stem}.docx' "
             "(available: {stem}, {idx})",
    )
    ap.add_argument(
        "--dry-run", action="store_true",
        help="print plan only, don't write files",
    )
    args = ap.parse_args()

    src = Path(args.docx).expanduser().resolve()
    if not src.exists():
        print(f"ERROR: input docx not found: {src}", file=sys.stderr)
        return 2
    out_dir = Path(args.out_dir).expanduser().resolve()

    extracts, tbl_count = plan_extracts(src)

    if tbl_count == 0:
        print(f"[extract-tables] 0 tables in docx — nothing to extract")
        return 0

    if args.dry_run:
        print(f"# DRY RUN — would write {len(extracts)} files to {out_dir}/")
        for e in extracts:
            fname = args.name_pattern.format(stem=e["stem"], idx=e["idx"])
            cap = (e["caption"] or "[no caption]")[:60]
            print(f"  [{e['idx']:>2}] tbl@body[{e['tbl_idx_in_body']}]  → {fname}  (caption={cap!r})")
        return 0

    out_dir.mkdir(parents=True, exist_ok=True)
    src_parts = _read_source_parts(src)
    body_children = list(src_parts["doc_body"]) if src_parts["doc_body"] is not None else []

    n_ok = 0
    inline_img_count = 0
    for e in extracts:
        fname = args.name_pattern.format(stem=e["stem"], idx=e["idx"])
        dst = out_dir / fname
        try:
            tc, nbytes, has_img = write_extract_minimal(
                src_parts, dst, e["tbl_idx_in_body"], e["caption_idx_in_body"],
                body_children,
            )
            if has_img:
                inline_img_count += 1
            img_flag = " [img]" if has_img else ""
            print(f"  [{e['idx']:>2}] {fname}  · {tc} tables · {nbytes:,} bytes{img_flag}")
            n_ok += 1
        except Exception as exc:
            print(f"  [{e['idx']:>2}] FAILED {fname}: {type(exc).__name__}: {exc}",
                  file=sys.stderr)
    print(f"OK: {n_ok} files written to {out_dir}")
    if inline_img_count:
        print(f"NOTE: {inline_img_count} table(s) reference inline images; "
              "the minimal docx omits media files (Word will show broken refs).")
    return 0 if n_ok == len(extracts) else 1


if __name__ == "__main__":
    sys.exit(main())
