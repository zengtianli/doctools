#!/usr/bin/env python3
# distilled from eco-flow/taizhou-天台 split-by-h1 need (2026-05-26 W1)
r"""split_by_h1.py — 按 Heading 1 切分 docx → N 个独立 docx。

每段保留原 docx 全套 styles / numbering / sectPr / page setup / 图片 / 表格 /
rels (整套 zip parts 复制, 只改 document.xml 的 body)。

CLI:
    python3 scripts/document/sub/split_by_h1.py \
        --docx <path> --out-dir <dir> \
        [--name-pattern '{idx:02d}-{title}.docx'] \
        [--include-frontmatter] [--dry-run]

策略:
    1. 用 python-docx 打开源 docx → 顺序遍历 body XML 子元素(<w:p>/<w:tbl>等),
       记录每个 <w:p style="Heading 1"> 的 element index 与标题文本。
    2. 区段切片 = [h1_i.idx, h1_{i+1}.idx) 半开区间;
       第一个 H1 前的内容: 默认丢弃, --include-frontmatter 时存 idx=0。
    3. 每段输出: shutil.copy 源 docx 到目标(保留 styles/numbering/media 等),
       打开 copy → 删 body 内非目标区段的子元素(保留 sectPr 全文末节标记) → save。
    4. 文件名: 非法字符 / \ : * ? " < > | 替换 _; 多空格压一个; strip; 截 100 char。
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx 未安装 (pip install python-docx)", file=sys.stderr)
    sys.exit(2)


H1_STYLES = {
    "Heading 1",
    "标题 1",
    "heading 1",
    "1",
    "10",
    "1.1.1.1 N级标题",
    # 院方系列报告（浙江省水利水电勘测设计院可用水量/生态流量报告）样式名
    "1一级标题",
}


_ILLEGAL_FILENAME_RE = re.compile(r'[/\\:*?"<>|\r\n\t]')
_MULTI_WS_RE = re.compile(r"\s+")


def sanitize_filename(name: str, max_len: int = 100) -> str:
    """Replace illegal filename chars with _, compress whitespace, strip, truncate."""
    if not name:
        return "untitled"
    n = _ILLEGAL_FILENAME_RE.sub("_", name)
    n = _MULTI_WS_RE.sub(" ", n).strip()
    if not n:
        return "untitled"
    if len(n) > max_len:
        n = n[:max_len].rstrip()
    return n


def get_style_name(p) -> str:
    try:
        return (p.style.name or "") if p.style is not None else ""
    except Exception:
        return ""


def is_h1_paragraph(p) -> bool:
    """True if paragraph carries an H1-class style (per H1_STYLES set)."""
    return get_style_name(p) in H1_STYLES


def _body_children(doc) -> list:
    """Return ordered list of direct child elements of <w:body> (excludes sectPr)."""
    body = doc.element.body
    # body 的最后一个子通常是 sectPr (文档级节属性), 不动它
    return list(body)


def _iter_paragraphs_in_element(elem):
    """Find all <w:p> descendants of a body-level element (for tables, etc.)."""
    return elem.iter(qn("w:p"))


def _paragraph_text(p_elem) -> str:
    """Extract concatenated text from <w:t> nodes under a <w:p>."""
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def _build_style_id_to_name(doc) -> dict:
    """Map styleId → styleName for the doc (院方系列 styleId 如 '1f8',
    与 styleName '1一级标题' 不同 → 必须建映射才能命中 H1_STYLES)."""
    out = {}
    try:
        styles_xml = doc.styles.element
        for s in styles_xml.findall(qn("w:style")):
            sid = s.get(qn("w:styleId"))
            name_el = s.find(qn("w:name"))
            nm = name_el.get(qn("w:val")) if name_el is not None else None
            if sid and nm:
                out[sid] = nm
    except Exception:
        pass
    return out


def _is_h1_elem(elem, style_id_to_name: Optional[dict] = None) -> bool:
    """Check if a body-level <w:p> element carries an H1 style."""
    if elem.tag != qn("w:p"):
        return False
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return False
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return False
    style_val = pStyle.get(qn("w:val")) or ""
    if style_val in H1_STYLES:
        return True
    # docx style id may differ from style name; also try common normalizations
    if style_val.lower() in {s.lower() for s in H1_STYLES}:
        return True
    # Heuristic: some templates use "1" or "10" as Heading 1 styleId
    if style_val in {"1", "10", "Heading1"}:
        return True
    # 院方系列: styleId (如 '1f8') ≠ styleName (如 '1一级标题') → 查 map
    if style_id_to_name is not None:
        nm = style_id_to_name.get(style_val)
        if nm and (nm in H1_STYLES or nm.lower() in {s.lower() for s in H1_STYLES}):
            return True
    return False


def plan_slices(docx_path: Path, include_frontmatter: bool, doc=None):
    """Inspect docx → return (slices, sect_idx, h1_count).

    slices: list[{idx, title, start, end, is_frontmatter}]
    h1_count: number of H1 elements detected in body (independent of slice emission)

    `doc`: if provided, skip the source parse (pipeline reuse).
    """
    if doc is None:
        doc = Document(str(docx_path))
    body = doc.element.body
    children = list(body)
    style_id_to_name = _build_style_id_to_name(doc)
    # Locate sectPr (final node) — we exclude it from slicing range
    sect_idx = len(children)
    for i in range(len(children) - 1, -1, -1):
        if children[i].tag == qn("w:sectPr"):
            sect_idx = i
            break

    # Find H1 positions among children [0, sect_idx)
    h1_positions: list[tuple[int, str]] = []
    for i in range(sect_idx):
        elem = children[i]
        if _is_h1_elem(elem, style_id_to_name):
            title = _paragraph_text(elem).strip() or "untitled"
            h1_positions.append((i, title))

    slices: list[dict] = []
    idx_counter = 0
    if h1_positions:
        first_h1 = h1_positions[0][0]
        if include_frontmatter and first_h1 > 0:
            slices.append({
                "idx": idx_counter,
                "title": "frontmatter",
                "start": 0,
                "end": first_h1,
                "is_frontmatter": True,
            })
            idx_counter += 1
        for k, (pos, title) in enumerate(h1_positions):
            end = h1_positions[k + 1][0] if k + 1 < len(h1_positions) else sect_idx
            slices.append({
                "idx": idx_counter,
                "title": title,
                "start": pos,
                "end": end,
                "is_frontmatter": False,
            })
            idx_counter += 1
    else:
        # No H1 found; treat full body as one slice if frontmatter requested
        if include_frontmatter and sect_idx > 0:
            slices.append({
                "idx": 0,
                "title": "frontmatter",
                "start": 0,
                "end": sect_idx,
                "is_frontmatter": True,
            })
    return slices, sect_idx, len(h1_positions)


def write_slice(src_docx: Path, dst_docx: Path, start: int, end: int) -> tuple[int, int]:
    """Copy src→dst, then prune body keeping only children [start, end) + sectPr.

    Returns (paragraph_count_in_slice, file_bytes).
    """
    dst_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(str(src_docx), str(dst_docx))
    doc = Document(str(dst_docx))
    body = doc.element.body
    children = list(body)
    # Identify sectPr (last)
    sect_elem = None
    for i in range(len(children) - 1, -1, -1):
        if children[i].tag == qn("w:sectPr"):
            sect_elem = children[i]
            break
    # Determine elements to KEEP — by their original index in children list
    keep_set = set(range(start, end))
    # Remove all body children NOT in keep_set (and not the trailing sectPr)
    for i, elem in enumerate(children):
        if elem is sect_elem:
            continue
        if i not in keep_set:
            body.remove(elem)
    doc.save(str(dst_docx))
    # Count paragraphs in saved slice
    doc2 = Document(str(dst_docx))
    para_count = len(doc2.paragraphs)
    file_bytes = dst_docx.stat().st_size
    return para_count, file_bytes


def run_split(
    src_docx: Path,
    out_dir: Path,
    include_frontmatter: bool = False,
    allow_no_h1: bool = False,
    dry_run: bool = False,
    name_pattern: str = "{idx:02d}-{title}.docx",
    doc=None,
) -> dict:
    """Execute split-by-h1; reuses provided `doc` for planning if given.

    Returns a report dict (used by pipeline built-in step). Raises on
    fail-fast conditions (0 H1 + no allow_no_h1) for caller to handle.
    """
    src = Path(src_docx).expanduser().resolve()
    if not src.exists():
        return {"error": f"input docx not found: {src}", "exit_code": 2}
    out_dir = Path(out_dir).expanduser().resolve()

    slices, sect_idx, h1_count = plan_slices(src, include_frontmatter, doc=doc)

    if h1_count == 0 and not allow_no_h1:
        return {
            "error": "0 Heading-1 detected (docx unhealthy); run /docx health first or pass --allow-no-h1",
            "exit_code": 3,
            "h1_count": 0,
        }

    if not slices:
        return {"h1_count": h1_count, "slices_emitted": 0, "note": "no slices to emit"}

    if dry_run:
        plan = []
        for s in slices:
            safe = sanitize_filename(s["title"])
            fname = name_pattern.format(idx=s["idx"], title=safe)
            plan.append({"idx": s["idx"], "fname": fname, "title": s["title"]})
        return {"h1_count": h1_count, "slices_planned": plan, "dry_run": True}

    out_dir.mkdir(parents=True, exist_ok=True)
    emitted = []
    failed = []
    for s in slices:
        safe = sanitize_filename(s["title"])
        fname = name_pattern.format(idx=s["idx"], title=safe)
        dst = out_dir / fname
        try:
            paras, nbytes = write_slice(src, dst, s["start"], s["end"])
            emitted.append({
                "idx": s["idx"], "fname": fname, "paragraphs": paras, "bytes": nbytes,
            })
        except Exception as e:
            failed.append({"idx": s["idx"], "fname": fname,
                           "error": f"{type(e).__name__}: {e}"})
    return {
        "h1_count": h1_count,
        "slices_emitted": len(emitted),
        "slices_failed": len(failed),
        "out_dir": str(out_dir),
        "emitted": emitted,
        "failed": failed,
    }


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Split a DOCX by Heading 1 into N independent DOCX files.",
    )
    ap.add_argument("--docx", required=True, help="input docx path")
    ap.add_argument("--out-dir", required=True, help="output directory (mkdir -p)")
    ap.add_argument(
        "--name-pattern",
        default="{idx:02d}-{title}.docx",
        help="output filename pattern, default '{idx:02d}-{title}.docx'",
    )
    ap.add_argument(
        "--include-frontmatter", action="store_true",
        help="emit content before first H1 as 00-frontmatter.docx",
    )
    ap.add_argument(
        "--dry-run", action="store_true",
        help="print plan only, don't write files",
    )
    ap.add_argument(
        "--allow-no-h1", action="store_true",
        help="suppress unhealthy-docx fail-fast when 0 H1 detected (rarely needed; "
             "default behavior is to FAIL and instruct user to run /docx health first)",
    )
    args = ap.parse_args()

    src = Path(args.docx).expanduser().resolve()
    if not src.exists():
        print(f"ERROR: input docx not found: {src}", file=sys.stderr)
        return 2
    out_dir = Path(args.out_dir).expanduser().resolve()

    slices, sect_idx, h1_count = plan_slices(src, args.include_frontmatter)

    # Health gate: 0 H1 detected = docx unhealthy signal (default fail-fast).
    # Iron rule [[docx-split-fail-run-health-first]]: don't patch around bad data,
    # tell the user to run /docx health first (scaffold already exists).
    if h1_count == 0 and not args.allow_no_h1:
        print(
            "\n".join([
                "",
                "ERROR: 0 Heading-1 detected in body — docx likely UNHEALTHY.",
                f"  file: {src}",
                "",
                "  Recognized H1 styles: " + ", ".join(sorted(H1_STYLES)),
                "",
                "  Likely causes:",
                "    - chapter titles styled as Normal/body paragraphs (not Heading 1)",
                "    - heading-level-skew (real H1 demoted to H2/H3/...)",
                "    - caption-outline-pollution (figure/table captions stole H1 slot)",
                "    - custom Chinese style name not in H1_STYLES whitelist",
                "",
                "  Fix path (scaffold already in place):",
                f"    /docx health diagnose '{src}'    # see which病种 hits",
                f"    /docx health full     '{src}'    # diagnose + auto-fix safe + re-diagnose",
                "    # then re-run this split script",
                "",
                "  Escape hatch (rare): pass --allow-no-h1 if you really want the current behavior",
                "  (emits only frontmatter.docx, requires --include-frontmatter).",
                "",
            ]),
            file=sys.stderr,
        )
        return 3

    if not slices:
        # h1_count > 0 path can't reach here; this is the --allow-no-h1 + no frontmatter case
        print(f"[split-by-h1] no H1 (and no frontmatter requested) — nothing to do")
        return 0

    if args.dry_run:
        print(f"# DRY RUN — would write {len(slices)} files to {out_dir}/")
        for s in slices:
            safe = sanitize_filename(s["title"])
            fname = args.name_pattern.format(idx=s["idx"], title=safe)
            print(f"  [{s['idx']:>2}] children[{s['start']}:{s['end']}]  → {fname}  (title={s['title']!r})")
        return 0

    out_dir.mkdir(parents=True, exist_ok=True)
    n_ok = 0
    for s in slices:
        safe = sanitize_filename(s["title"])
        fname = args.name_pattern.format(idx=s["idx"], title=safe)
        dst = out_dir / fname
        try:
            paras, nbytes = write_slice(src, dst, s["start"], s["end"])
            print(f"  [{s['idx']:>2}] {fname}  · {paras} paragraphs · {nbytes:,} bytes")
            n_ok += 1
        except Exception as e:
            print(f"  [{s['idx']:>2}] FAILED {fname}: {type(e).__name__}: {e}",
                  file=sys.stderr)
    print(f"OK: {n_ok} files written to {out_dir}")
    return 0 if n_ok == len(slices) else 1


if __name__ == "__main__":
    sys.exit(main())
