#!/usr/bin/env python3
"""relocate_orphan_blocks.py — 按外部 plan JSON 把 docx 中孤儿段块挪到正确位置.

输入: plan JSON, 含 moves[] 列表; 每个 move 描述一个孤儿块的 source 范围与
target 插入位置. 脚本对每个 move 独立执行:

  1. 用 source_heading_text 二次锚定 source_heading_idx (防 idx 漂移)
  2. 用 target_context_text 二次锚定 target_insert_after_idx
  3. lxml 操作 <w:body>: detach source 段块的所有 block 子元素 (含 <w:p>/<w:tbl>),
     在 target 锚之后顺序 insert

关键设计:
  - **二次锚定 (text-based re-resolve)**: 每个 move 执行前都基于 fresh doc 重算
    idx, 不复用 plan 里的 idx 字段做最终定位 — 因 上一个 move 把段移走,
    后续 idx 漂移. plan 里的 idx 是给 worker 算 source range 用的, text 是
    最终锚.
  - **保留 <w:tbl>**: source range 内若含 table, 一并搬走 (用 body 直系子元素
    扫描, 不只 <w:p>).
  - **不算 number / 不重排**: 纯机械搬, 信任 plan 内容.

接口:
  python3 relocate_orphan_blocks.py <docx> --plan <plan_json> [--dry-run] [--no-backup]

默认: 真改 + 自动备份 .bak-N-<date>.docx
"""
from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import sys
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_P = f"{{{W_NS}}}p"
W_TBL = f"{{{W_NS}}}tbl"
W_SECTPR = f"{{{W_NS}}}sectPr"


# ---------- helpers ----------

def lsof_check(docx_path: Path) -> Optional[str]:
    """检测 docx 是否被 Word 等进程占用. 返回占用进程描述, 否则 None."""
    try:
        out = subprocess.run(
            ["lsof", "--", str(docx_path)],
            capture_output=True, text=True, timeout=5,
        )
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return None
    if out.returncode == 0 and out.stdout.strip():
        # lsof 列出占用行
        lines = out.stdout.strip().split("\n")
        if len(lines) > 1:  # 含 header + 至少一条
            return "\n".join(lines)
    return None


def get_body_block_children(body):
    """返回 body 直系子元素中的 block-level 元素列表 (按文档顺序).

    block-level = <w:p> 或 <w:tbl>; <w:sectPr> 与其他元素跳过.
    这是与 docx.paragraphs 的关键差异: paragraphs 不含 <w:tbl>.
    我们用全集合 blocks (p+tbl) 做 detach/insert, idx 体系也基于 blocks.
    """
    blocks = []
    for child in body:
        tag = child.tag
        if tag == W_P or tag == W_TBL:
            blocks.append(child)
    return blocks


def block_text(elem) -> str:
    """提取 block (p or tbl) 文本, strip 空白. tbl 取所有 <w:t> 文本拼接."""
    texts = []
    for t in elem.iter(f"{{{W_NS}}}t"):
        if t.text:
            texts.append(t.text)
    return "".join(texts).strip()


def find_block_idx_by_text(blocks: list, expected_text: str,
                          hint_idx: Optional[int] = None,
                          search_window: int = 50) -> Optional[int]:
    """在 blocks 中找首个文本完全匹配 expected_text 的 block idx.

    优先策略:
      1. 若 hint_idx 给了, 先检查 hint_idx (零漂移情形)
      2. 否则在 hint_idx ± search_window 范围内找
      3. 仍不中 → 全文档扫
      4. 找到多个 → 取最接近 hint_idx 的; 无 hint 取首个
    """
    expected = (expected_text or "").strip()
    if not expected:
        return None

    # 完整候选: 文本完全相等
    matches_exact = [i for i, b in enumerate(blocks) if block_text(b) == expected]
    if not matches_exact:
        # fallback: 前缀匹配 (>= 10 chars 或全 text 长度)
        min_len = min(len(expected), 30)
        prefix = expected[:min_len]
        matches_exact = [i for i, b in enumerate(blocks)
                        if block_text(b).startswith(prefix)
                        and len(block_text(b)) >= min_len]

    if not matches_exact:
        return None

    if hint_idx is None:
        return matches_exact[0]

    # 取最接近 hint 的
    return min(matches_exact, key=lambda i: abs(i - hint_idx))


def find_block_idx_by_context(blocks: list, context_text: str,
                              hint_idx: Optional[int] = None) -> Optional[int]:
    """target_context_text 是 "插入位置前一段的 text 前 50 字" — 按前缀匹配查."""
    ctx = (context_text or "").strip()
    if not ctx:
        return None

    # 完全匹配优先
    exact = [i for i, b in enumerate(blocks) if block_text(b) == ctx]
    if exact:
        if hint_idx is None:
            return exact[0]
        return min(exact, key=lambda i: abs(i - hint_idx))

    # 前缀匹配 (block text 以 ctx 开头, 或 ctx 以 block text 开头)
    # 用 ctx 前 30-50 字片段
    snippet = ctx[: min(50, len(ctx))]
    cand = []
    for i, b in enumerate(blocks):
        bt = block_text(b)
        if not bt:
            continue
        # 双向前缀容错
        n = min(len(snippet), len(bt), 30)
        if n < 8:
            continue
        if snippet[:n] == bt[:n]:
            cand.append(i)
    if not cand:
        return None
    if hint_idx is None:
        return cand[0]
    return min(cand, key=lambda i: abs(i - hint_idx))


# ---------- move 执行 ----------

def execute_move(body, move: dict, move_no: int, dry_run: bool = False) -> dict:
    """执行单个 move. 返回 result dict.

    流程:
      1. blocks = body 当前 block 子元素 (fresh)
      2. text 锚定 source_heading_idx → src_h_idx
      3. text 锚定 target_insert_after_idx → tgt_idx
      4. source range = [src_h_idx, src_h_idx + (orig_end - orig_head)]
         (用 plan 里的 end - head 差作长度)
      5. detach source 段 (collect refs), 在 tgt_idx 之后顺序 insert
      6. 防自杀: target 不能落在 source range 内
    """
    blocks = get_body_block_children(body)
    n = len(blocks)

    src_head_hint = move.get("source_heading_idx")
    src_end_hint = move.get("source_block_end_idx")
    tgt_hint = move.get("target_insert_after_idx")
    src_text = move.get("source_heading_text", "")
    tgt_text = move.get("target_context_text", "")

    # 二次锚定
    src_h_idx = find_block_idx_by_text(blocks, src_text, hint_idx=src_head_hint)
    if src_h_idx is None:
        return {
            "move_no": move_no, "status": "skip",
            "reason": f"source_heading_text not found: {src_text[:40]!r}",
        }

    tgt_idx = find_block_idx_by_context(blocks, tgt_text, hint_idx=tgt_hint)
    if tgt_idx is None:
        return {
            "move_no": move_no, "status": "skip",
            "reason": f"target_context_text not found: {tgt_text[:40]!r}",
        }

    # 计算 source range 长度 — 用 plan 里 end-head 差 (块长度对漂移不敏感)
    if src_head_hint is None or src_end_hint is None:
        return {
            "move_no": move_no, "status": "skip",
            "reason": "plan missing source_heading_idx / source_block_end_idx",
        }
    block_len = src_end_hint - src_head_hint + 1
    if block_len < 1:
        return {
            "move_no": move_no, "status": "skip",
            "reason": f"invalid block_len={block_len}",
        }

    src_end_idx = src_h_idx + block_len - 1
    if src_end_idx >= n:
        # 块长度超出当前 blocks 尾部 (可能 plan 与现状对不上) → 截到 n-1
        src_end_idx = n - 1

    # 自杀检测
    if src_h_idx <= tgt_idx <= src_end_idx:
        return {
            "move_no": move_no, "status": "skip",
            "reason": f"target idx {tgt_idx} falls inside source range [{src_h_idx},{src_end_idx}]",
        }

    if dry_run:
        return {
            "move_no": move_no, "status": "dry-ok",
            "resolved_source_head_idx": src_h_idx,
            "resolved_source_end_idx": src_end_idx,
            "resolved_target_idx": tgt_idx,
            "block_len": block_len,
        }

    # 收集 source 段的 elem refs
    source_elems = blocks[src_h_idx: src_end_idx + 1]

    # target elem ref (在 detach 前抓引用)
    target_elem = blocks[tgt_idx]

    # detach source
    parent = body
    for e in source_elems:
        if e.getparent() is parent:
            parent.remove(e)

    # 在 target_elem 之后顺序 insert
    # parent 是 body, 用 addnext 链或 index-based insert
    # 我们用 index-based: 先取 target 在 body 中的当前 index, 然后 insert(index+1, ...)
    # 注意 detach 后 target 索引会变, 重新 query
    body_children = list(body)
    try:
        tgt_pos = body_children.index(target_elem)
    except ValueError:
        return {
            "move_no": move_no, "status": "error",
            "reason": "target_elem lost from body after source detach",
        }

    insert_at = tgt_pos + 1
    for offset, e in enumerate(source_elems):
        body.insert(insert_at + offset, e)

    return {
        "move_no": move_no, "status": "ok",
        "resolved_source_head_idx": src_h_idx,
        "resolved_source_end_idx": src_end_idx,
        "resolved_target_idx": tgt_idx,
        "block_len": block_len,
        "moved_block_count": len(source_elems),
    }


# ---------- main ----------

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="按 plan JSON 把 docx 孤儿段块挪到正确位置 (机械操作, 不算编号)"
    )
    ap.add_argument("docx", help="输入 docx 路径")
    ap.add_argument("--plan", required=True, help="plan JSON 路径 (含 moves[])")
    ap.add_argument("--dry-run", action="store_true", help="只 verify plan, 不动 docx")
    ap.add_argument("--no-backup", action="store_true", help="不自动备份")
    ap.add_argument("--self-test", action="store_true",
                    help="生成内置 fake plan 做 self-test (开发用)")
    args = ap.parse_args(argv)

    src = Path(args.docx)
    if not src.exists():
        print(f"[error] docx 不存在: {src}", file=sys.stderr)
        return 2

    plan_path = Path(args.plan)
    if not plan_path.exists():
        print(f"[error] plan 不存在: {plan_path}", file=sys.stderr)
        return 2

    # lsof 检查
    occupied = lsof_check(src)
    if occupied:
        print(f"[error] docx 被进程占用, 关闭 Word/WPS 后重试:\n{occupied}",
              file=sys.stderr)
        return 3

    # 读 plan + schema 校验 (doctools v1)
    try:
        plan = json.loads(plan_path.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"[error] plan JSON 解析失败: {e}", file=sys.stderr)
        return 2

    try:
        from lib.schemas import validate as _validate_schema
        _err = _validate_schema(plan, "plan")
        if _err:
            print(f"[error] plan schema 校验失败 (v1): {_err}", file=sys.stderr)
            return 2
    except Exception:
        # schemas lib 不可用时降级到原 ad-hoc 检查
        pass

    moves = plan.get("moves", [])
    if not isinstance(moves, list):
        print(f"[error] plan.moves 不是 list", file=sys.stderr)
        return 2

    print(f"[plan] {plan_path.name} — moves={len(moves)}  source_docx={plan.get('source_docx','?')}")

    # 备份
    backup_path = None
    if not args.dry_run and not args.no_backup:
        today = date.today().isoformat()
        m = 1
        while True:
            cand = src.with_name(f"{src.stem}.bak-{m}-{today}{src.suffix}")
            if not cand.exists():
                break
            m += 1
        shutil.copy2(src, cand)
        backup_path = cand
        print(f"[backup] {cand.name}")

    # 加载 docx
    doc = Document(str(src))
    body = doc.element.body

    results = []
    for i, move in enumerate(moves, start=1):
        res = execute_move(body, move, i, dry_run=args.dry_run)
        results.append(res)
        # 简报
        status = res["status"]
        if status in ("ok", "dry-ok"):
            print(f"  [{status}] move#{i}  src_h={res.get('resolved_source_head_idx')} "
                  f"end={res.get('resolved_source_end_idx')} "
                  f"→ after idx={res.get('resolved_target_idx')}  "
                  f"len={res.get('block_len')}  "
                  f"src={move.get('source_heading_text','')[:30]!r}")
        else:
            print(f"  [{status}] move#{i}  {res.get('reason','?')}")

    # 写盘
    if not args.dry_run:
        ok_cnt = sum(1 for r in results if r["status"] == "ok")
        if ok_cnt > 0:
            doc.save(str(src))
            print(f"[save] {src.name} ← {ok_cnt} moves applied")
            # 重读自证 OOXML 合法
            try:
                _verify = Document(str(src))
                _para_n = len(_verify.paragraphs)
                print(f"[verify] re-read ok, paragraphs={_para_n}")
            except Exception as e:
                print(f"[error] re-read failed: {e}", file=sys.stderr)
                if backup_path:
                    print(f"  → 可用备份恢复: {backup_path.name}")
                return 4
        else:
            print(f"[save] 无 ok move, 不写盘")

    # 抽样 print 末尾段顺序 (前后各 5 段)
    if not args.dry_run:
        post_doc = Document(str(src))
        post_blocks = get_body_block_children(post_doc.element.body)
        print(f"\n[post-sample] total blocks={len(post_blocks)}, first 5 + last 5:")
        for i, b in enumerate(post_blocks[:5]):
            t = block_text(b)[:60]
            print(f"  [{i}] {t!r}")
        if len(post_blocks) > 10:
            print(f"  ...")
            for i, b in enumerate(post_blocks[-5:], start=len(post_blocks) - 5):
                t = block_text(b)[:60]
                print(f"  [{i}] {t!r}")

    # 汇总
    ok = sum(1 for r in results if r["status"] == "ok")
    dry = sum(1 for r in results if r["status"] == "dry-ok")
    skip = sum(1 for r in results if r["status"] == "skip")
    err = sum(1 for r in results if r["status"] == "error")
    print(f"\n[summary] ok={ok}  dry-ok={dry}  skip={skip}  error={err}  total={len(results)}")

    return 0 if err == 0 else 5


# ---------------- pipeline adapter ----------------
def apply(doc, args=None) -> dict:
    """pipeline: 仅当 args.relocate_plan 提供时执行;否则 noop"""
    plan_path = getattr(args, "relocate_plan", None) if args else None
    if not plan_path:
        return {"changed": 0, "skipped": "no relocate_plan in args"}
    dry = bool(getattr(args, "dry_run", False)) if args else False
    try:
        plan = json.loads(Path(plan_path).read_text(encoding="utf-8"))
    except Exception as exc:
        return {"error": f"plan read failed: {exc}"}
    moves = plan.get("moves", [])
    body = doc.element.body
    results = []
    for i, move in enumerate(moves, start=1):
        res = execute_move(body, move, i, dry_run=dry)
        results.append(res)
    ok = sum(1 for r in results if r["status"] == "ok")
    return {
        "changed": ok,
        "ok": ok,
        "skip": sum(1 for r in results if r["status"] == "skip"),
        "error": sum(1 for r in results if r["status"] == "error"),
    }


if __name__ == "__main__":
    sys.exit(main())
