#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""renumber_headings_seq.py — 标题编号按**现有编号深度**重排序（修断号/重号/错号，不动层级）。

与 renumber_headings.py 的分工（一脚本一功能）:
  renumber_headings.py   层级取自**样式**(Heading 1/2/3),全篇推倒重编 —— 适合样式可信的文档。
  本脚本                 层级取自**现有编号的段数**("12.1"=2 级,不管它套什么样式),只修
                         序列错误 —— 适合样式≠层级的文档(如标书 12.1~12.7 套 Heading 1)、
                         以及"只想改错号、别动结构"的场景。序号修正(doc_dispatch renum)默认用本脚本。

算法(顺序扫 Heading 1-4 段, 跳过无编号标题——不发明编号):
  解析段首点分编号 → depth = 段数。
  depth=1: 首个编号作 base(保留,标书从 10 起就还是 10),后续 = base+i。
  depth>1: 前缀 = 当前上级修正后编号,末位按同父内出现序 1..k。
  深度突跳(1 级直接接 3 级)记 warning,按"父链最近已知编号"补链。

接口: python3 renumber_headings_seq.py <docx> [--dry-run] [--no-backup] [--report <json>]
退出码: 0 成功且复检连贯 / 2 输入错误 / 4 复检失败
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from renumber_headings import (  # noqa: E402  复用 run 级安全改写,不重写
    get_paragraph_text,
    get_style_name,
    make_backup,
    rewrite_paragraph_with_prefix,
)

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx 未安装", file=sys.stderr)
    sys.exit(2)

HEAD_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
NUM_RE = re.compile(r"^\s*(\d+(?:[.．]\d+)*)[\s　.．、]")


def plan(doc):
    """返回 (plan_items, warnings)。plan_items=[(para_idx, old_num, new_num)]。"""
    items, warns = [], []
    path: list[int] = []  # 修正后的编号链,如 [12, 3] = 12.3
    base_locked = False
    for idx, p in enumerate(doc.paragraphs):
        if get_style_name(p) not in HEAD_STYLES:
            continue
        m = NUM_RE.match(get_paragraph_text(p))
        if not m:
            continue  # 无编号标题(前言/附录)不发明编号
        old = m.group(1).replace("．", ".")
        depth = old.count(".") + 1
        if depth == 1:
            if not base_locked:
                path = [int(old)]  # 首个一级号 = base,原样保留(标书从10起仍是10)
                base_locked = True
            else:
                path = [_next_top(items, int(old))]
        else:
            if depth > len(path) + 1:
                warns.append(f"#{idx} 深度突跳: {old} (当前链 {'.'.join(map(str, path))})")
                while depth > len(path) + 1:
                    path.append(1)
            path = path[: depth - 1] + [_next_sibling(items, path[: depth - 1])]
        new = ".".join(map(str, path))
        items.append((idx, old, new, depth))
    return items, warns


def _next_top(items, _old):
    tops = [int(n.split(".")[0]) for _, _, n, d in items if d == 1]
    return tops[-1] + 1 if tops else 1


def _next_sibling(items, parent: list[int]):
    pre = ".".join(map(str, parent)) + "."
    sibs = [int(n[len(pre):]) for _, _, n, d in items
            if d == len(parent) + 1 and n.startswith(pre) and "." not in n[len(pre):]]
    return sibs[-1] + 1 if sibs else 1


def verify(doc):
    """复检: 再按 plan 模型扫一遍,所有 old==new 即连贯。"""
    items, _ = plan(doc)
    return [(i, o, n) for i, o, n, _ in items if o != n]


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="标题编号按现有深度重排序(修断号/重号/错号,不动层级)")
    ap.add_argument("docx_path")
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--no-backup", action="store_true")
    ap.add_argument("--report", default=None)
    a = ap.parse_args(argv)
    path = Path(a.docx_path)
    if not path.exists():
        print(f"ERROR: {path} 不存在", file=sys.stderr)
        return 2

    doc = Document(str(path))
    items, warns = plan(doc)
    changes = [(i, o, n) for i, o, n, _ in items if o != n]
    print(f"[plan] 带编号标题 {len(items)} 段, 需修 {len(changes)} 段")
    for i, o, n in changes:
        print(f"  #{i} {o} → {n}")
    for w in warns:
        print(f"  ⚠ {w}")
    if a.report:
        Path(a.report).write_text(json.dumps(
            {"changes": [{"idx": i, "old": o, "new": n} for i, o, n in changes],
             "warnings": warns}, ensure_ascii=False, indent=2), encoding="utf-8")
    if a.dry_run:
        print("[dry-run] 未写文件")
        return 0
    if not changes:
        print("✓ 标题编号已连贯,无需改写")
        return 0
    if not a.no_backup:
        print(f"[backup] {make_backup(path).name}")
    paras = doc.paragraphs
    for i, _o, n in changes:
        rewrite_paragraph_with_prefix(paras[i], f"{n} ")
    doc.save(str(path))
    left = verify(Document(str(path)))
    if left:
        print(f"✗ 复检仍有 {len(left)} 处不连贯: {left[:5]}", file=sys.stderr)
        return 4
    print(f"✓ 已修 {len(changes)} 处,复检连贯")
    return 0


if __name__ == "__main__":
    sys.exit(main())
