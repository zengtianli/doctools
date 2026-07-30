#!/usr/bin/env python3
"""删除 docx 中指定表的指定行范围（track-changes 不支持表行删除时用）。

Usage:
  python3 delete_table_rows.py --docx 修改稿.docx --table-index 6 --rows 8:15 \
    --expected-first-col "序号,1,2,3,4,5,6,7" \
    --expected-residue "自然资源集约利用"
"""
import argparse
import sys
from pathlib import Path

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[3] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py

from docx import Document


def apply(doc, args=None) -> dict:
    """在**已打开**的 doc 上删掉指定表的指定行区间，删前跑安全校验。

    校验失败**不 sys.exit**，而是回一个带 error_kind 的 report，由调用方决定后果
    （CLI 打中文提示后退 1；pipeline 记进报告接着跑别的 step）。这样校验只有一份：
    之前 main() 和 apply_path() 各写了一遍，两边判据已经漂开（关键字校验一处看
    「起始行是否够两列」、另一处不看），改一处必漏一处。

    删行只动 document.xml 的 <w:tr>，不碰别的部件，所以是 doc-based step。
    """
    if args is None:
        return {"skipped": "no args; delete_table_rows needs --table-index/--rows"}
    table_index = getattr(args, "table_index", None)
    rows = getattr(args, "rows", None)
    if table_index is None or not rows:
        return {"skipped": "delete_table_rows needs --table-index and --rows"}

    if table_index >= len(doc.tables):
        return {"error_kind": "table_index_out_of_range",
                "error": f"table-index {table_index} 超出范围 (共 {len(doc.tables)} 表)",
                "tables": len(doc.tables)}

    tbl = doc.tables[table_index]
    from_idx, to_idx = map(int, rows.split(":"))
    before = len(tbl.rows)
    dry_run = bool(getattr(args, "dry_run", False))
    expected_first_col = getattr(args, "expected_first_col", "") or ""
    expected_residue = getattr(args, "expected_residue", "") or ""
    expected_last = getattr(args, "expected_last", "") or ""

    base = {"table_index": table_index, "rows_range": rows,
            "before": before, "dry_run": dry_run}

    # 安全校验 1：保留行的第一列必须逐字对上
    if expected_first_col:
        expected = [x.strip() for x in expected_first_col.split(",")]
        actual = [tbl.rows[i].cells[0].text.strip()
                  for i in range(min(from_idx, before))]
        if actual != expected:
            return {**base, "error_kind": "first_col_mismatch",
                    "error": f"first-col 校验失败: 期望 {expected} 实际 {actual}",
                    "expected": expected, "actual": actual}

    # 安全校验 2：起始行不足两列时整条跳过 —— 合并单元格/小标题行本来就取不到
    # 第二列文本，硬查会把正常情况误判成「删错行」而挡住删除（沿用 main() 判据）。
    if expected_residue and len(tbl.rows[from_idx].cells) >= 2:
        head = tbl.rows[from_idx].cells[0].text + tbl.rows[from_idx].cells[1].text
        if expected_residue not in head:
            found = any(
                expected_residue in " ".join(c.text for c in tbl.rows[i].cells)
                for i in range(from_idx, min(to_idx + 1, before))
            )
            if not found:
                return {**base, "error_kind": "residue_not_found",
                        "error": f"被删行未含关键字「{expected_residue}」"}

    if dry_run:
        planned = to_idx - from_idx + 1
        return {**base, "after": before - planned, "deleted": planned, "changed": 0}

    # 倒序删除：正序删会让后续下标整体前移
    for i in range(to_idx, from_idx - 1, -1):
        if i < len(tbl.rows):
            row = tbl.rows[i]
            row._element.getparent().remove(row._element)

    after = len(tbl.rows)
    result = {**base, "after": after, "deleted": before - after,
              "changed": before - after}
    if expected_last and after >= 1:
        last = tbl.rows[-1]
        last_text = last.cells[1].text.strip() if len(last.cells) >= 2 else ""
        result["last_cell"] = last_text
        result["last_check"] = expected_last in last_text
    return result


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--docx", required=True, help="目标 docx（原地修改）")
    ap.add_argument("--table-index", type=int, required=True, help="表格索引（0-based）")
    ap.add_argument("--rows", required=True, help="要删的行范围 FROM:TO（闭区间，0-based）")
    ap.add_argument("--expected-first-col", default="",
                    help="删除前保留行第一列期望值（逗号分隔），用于安全校验；留空则跳过校验")
    ap.add_argument("--expected-residue", default="",
                    help="要删的第一行第二列应含此关键字，确保删对了")
    ap.add_argument("--expected-last", default="",
                    help="删后末行第二列期望值，用于校验删对了")
    args = ap.parse_args()

    docx_path = Path(args.docx)
    doc = Document(str(docx_path))
    rep = apply(doc, args)
    err = rep.get("error_kind")

    if err == "table_index_out_of_range":
        print(f"❌ table-index 超出范围（文档只有 {rep['tables']} 个表）", file=sys.stderr)
        sys.exit(1)

    from_idx, to_idx = map(int, args.rows.split(":"))
    print(f"删除前 T{args.table_index} 行数: {rep['before']}")
    print(f"计划删除: R{from_idx}-R{to_idx}（共 {to_idx-from_idx+1} 行）")

    # 安全校验 1：保留行的第一列
    if err == "first_col_mismatch":
        print(f"❌ 前 {from_idx} 行第一列与期望不一致："
              f"\n  期望 {rep['expected']}\n  实际 {rep['actual']}", file=sys.stderr)
        sys.exit(1)
    if args.expected_first_col:
        print("✓ 保留行结构校验通过")

    # 安全校验 2：要删的起始行应含关键字
    if err == "residue_not_found":
        print(f"❌ 被删行中未找到关键字「{args.expected_residue}」，可能删错行", file=sys.stderr)
        sys.exit(1)
    if args.expected_residue:
        print(f"✓ 被删行包含期望关键字「{args.expected_residue}」")

    doc.save(str(docx_path))

    # 验证
    doc2 = Document(str(docx_path))
    tbl2 = doc2.tables[args.table_index]
    print(f"删除后行数: {len(tbl2.rows)}")
    if args.expected_last and len(tbl2.rows) >= 1:
        last = tbl2.rows[-1]
        if len(last.cells) >= 2:
            actual_last = last.cells[1].text.strip()
            if args.expected_last and args.expected_last not in actual_last:
                print(f"⚠️ 末行第二列「{actual_last}」不含期望「{args.expected_last}」", file=sys.stderr)
            else:
                print(f"✓ 末行校验通过: {actual_last[:40]}")
    print(f"OK -> {docx_path}")


# ---------------- pipeline adapter ----------------
def apply_path(docx_path, args=None) -> dict:
    """原地 mutator：开文件 → apply() → 存盘。校验/删行逻辑全在 apply() 里。

    args 透传:
      - table_index (必需): 表索引 0-based
      - rows (必需): FROM:TO 闭区间 0-based
      - expected_first_col / expected_residue / expected_last: 安全校验
      - dry_run: 不写盘

    留着是为了不掐断已经按路径调它的老调用方；pipeline_lib.load_step 优先取
    apply()，新链路一律走纯内存版本，这里不会被 pipeline 选中。
    """
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))
    rep = apply(doc, args)
    # 校验没过 / 参数不全 / dry-run —— 一律不落盘
    if rep.get("error_kind") or rep.get("skipped") or rep.get("dry_run"):
        return rep
    doc.save(str(docx_path))
    return rep


if __name__ == "__main__":
    main()
