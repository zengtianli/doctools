"""shape_contract.py — structural invariants check (跑前后对账, 防 fix 工具破坏结构)

shape_contract / 结构对账 / _verify_structure_invariants (GOAL I6 keyword anchors)

W13 (2026-05-26): docx-health-v2 共享 helper · 所有 fix 工具跑前后 capture_structure +
verify_no_structural_drift, 任一结构指标漂移即拒跑 (除非 --force)。

8(+1) 项结构指标 (snapshot dict 字段):
  1. paragraph_count        (int)            — len(doc.paragraphs)
  2. table_count            (int)            — len(doc.tables)
  3. section_count          (int)            — len(doc.sections)
  4. heading_counts         (dict[str,int])  — {Heading 1..9, Title, Subtitle} Counter
  5. caption_figure_count   (int)            — paragraphs whose style.name matches 图名/figure caption
  6. caption_table_count    (int)            — paragraphs whose style.name matches 表名/表格标题/table caption
  7. drawings_count         (int)            — document.xml <w:drawing> count + zipfile word/media/* file count
  8. figure_number_set      (set[str])       — caption text 里的 "图X-Y" 集合
     table_number_set       (set[str])       — caption text 里的 "表X-Y" 集合
  9. track_changes_count    (int, bonus)     — w:ins + w:del element count (strip revisions 允许减)

Usage:
    from .shape_contract import capture_structure, verify_no_structural_drift

    before = capture_structure(docx_path)
    # ... run fix work that may modify docx ...
    after = capture_structure(docx_path)
    passed, violations = verify_no_structural_drift(before, after,
                                                    allowed_deltas={"paragraph_count": -10},
                                                    raise_on_fail=True)
"""

from __future__ import annotations

import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Tuple

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


# ───────────────────────────────────────────────────────────────────────────────
# style-name regex (覆盖中英两条线)
# ───────────────────────────────────────────────────────────────────────────────
_FIG_CAPTION_NAME_RE = re.compile(
    r"(图\s*名|图名称|figure\s*caption|image\s*caption|图\s*注|^figure$)",
    re.IGNORECASE,
)
_TABLE_CAPTION_NAME_RE = re.compile(
    r"(表\s*名|表名称|表格标题|table\s*caption|表\s*注|表\s*题|^table$)",
    re.IGNORECASE,
)

_HEADING_NAMES = (
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "Heading 5", "Heading 6", "Heading 7", "Heading 8", "Heading 9",
    "Title", "Subtitle",
)

# caption 文本里的 "图X-Y" / "表X-Y" 编号 (Y 可缺省, 兼容 "图3-2"/"图3.2-1" 等)
# 第 2 个 alt = 附录式 "附图N"/"附表N" (单号无短横, 院模板结论章/附录常用)。仅在带"附"
# 前缀时放宽无短横, 故不会误吃正文内联 "见图1" (无附前缀仍需短横); 且 set 去重防重计。
_FIG_NUM_TEXT_RE = re.compile(r"图\s*[\d．.]+[-—–][\d．.]+|附图\s*[\d．.]+")
_TBL_NUM_TEXT_RE = re.compile(r"表\s*[\d．.]+[-—–][\d．.]+|附表\s*[\d．.]+")


# ───────────────────────────────────────────────────────────────────────────────
# capture
# ───────────────────────────────────────────────────────────────────────────────
def _style_name_of(p) -> str:
    try:
        s = p.style
        if s is None:
            return ""
        return getattr(s, "name", "") or ""
    except Exception:
        return ""


def _count_drawings_in_doc(doc) -> int:
    """Count <w:drawing> elements in document body (paragraphs + tables)."""
    body = doc.element.body
    return len(body.findall(".//" + qn("w:drawing")))


def _count_media_files_in_zip(docx_path: Path) -> int:
    try:
        with zipfile.ZipFile(str(docx_path)) as z:
            return sum(1 for n in z.namelist()
                       if n.startswith("word/media/") and not n.endswith("/"))
    except (zipfile.BadZipFile, FileNotFoundError, OSError):
        return 0


def _count_track_changes(doc) -> int:
    body = doc.element.body
    n = 0
    n += len(body.findall(".//" + qn("w:ins")))
    n += len(body.findall(".//" + qn("w:del")))
    return n


def capture_structure(docx_path: Path | str) -> dict:
    """打开 docx, 返回 8+1 项结构指标 snapshot dict.

    Why double source (python-docx + zipfile): caption_figure_count via style 是
    用户视角的"题注段数", drawings_count via body+zip 是物理"图片对象数",
    二者结合能同时抓「caption 段被误删」+ 「图片对象丢失」两类灾难。
    """
    p = Path(docx_path)
    doc = Document(str(p))

    heading_counter: Counter = Counter()
    cap_fig = 0
    cap_tbl = 0
    fig_nums: set[str] = set()
    tbl_nums: set[str] = set()

    for para in doc.paragraphs:
        name = _style_name_of(para)
        if name in _HEADING_NAMES:
            heading_counter[name] += 1
        # caption by style name — 空段不计入 caption 数: 带「图名/表名」样式但无文本 =
        # 排版占位空行 (审定模板常在题注前后留同样式空段), 非真题注; 计入会虚增 caption 段数
        # 致 caption-count-consistency 误报 (编号集 ≠ 段数).
        is_fig = bool(_FIG_CAPTION_NAME_RE.search(name)) if name else False
        is_tbl = bool(_TABLE_CAPTION_NAME_RE.search(name)) if name else False
        text = para.text or ""
        if is_fig and text.strip():
            cap_fig += 1
        if is_tbl and text.strip():
            cap_tbl += 1
        # caption text → "图X-Y" / "表X-Y" 集合
        if is_fig or "图" in text[:6]:
            for m in _FIG_NUM_TEXT_RE.findall(text):
                fig_nums.add(re.sub(r"\s+", "", m))
        if is_tbl or "表" in text[:6]:
            for m in _TBL_NUM_TEXT_RE.findall(text):
                tbl_nums.add(re.sub(r"\s+", "", m))

    # drawings_count: body <w:drawing> + zip word/media/* (取 max, 双源对账)
    drawings_body = _count_drawings_in_doc(doc)
    media_files = _count_media_files_in_zip(p)
    drawings_count = max(drawings_body, media_files)

    return {
        "paragraph_count":      len(doc.paragraphs),
        "table_count":          len(doc.tables),
        "section_count":        len(doc.sections),
        "heading_counts":       dict(heading_counter),
        "caption_figure_count": cap_fig,
        "caption_table_count":  cap_tbl,
        "drawings_count":       drawings_count,
        "drawings_body":        drawings_body,
        "drawings_media_files": media_files,
        "figure_number_set":    sorted(fig_nums),
        "table_number_set":     sorted(tbl_nums),
        "track_changes_count":  _count_track_changes(doc),
    }


# ───────────────────────────────────────────────────────────────────────────────
# diff
# ───────────────────────────────────────────────────────────────────────────────
_SCALAR_FIELDS = (
    "paragraph_count",
    "table_count",
    "section_count",
    "caption_figure_count",
    "caption_table_count",
    "drawings_count",
    "track_changes_count",
)


def diff_structure(before: dict, after: dict,
                   allowed_deltas: dict | None = None) -> list[str]:
    """对比 before / after, 返回 violation 列表 (空 = 全过).

    allowed_deltas 字段名 → 允许的 delta (after - before).
      正数 = 允许增 N (或更少); 负数 = 允许减 N (或更少, 即 |after-before| <= |delta|).
      未列入 = 必须 0.
    特例:
      "heading_counts" 接受 int (各 heading 合计允许 delta) 或 dict (按 name).
      "figure_number_set" / "table_number_set" 接受 int (集合 size delta).
    """
    allowed = dict(allowed_deltas or {})
    violations: list[str] = []

    # 1. scalar fields
    for f in _SCALAR_FIELDS:
        b = int(before.get(f, 0) or 0)
        a = int(after.get(f, 0) or 0)
        delta = a - b
        if delta == 0:
            continue
        max_allowed = allowed.get(f, 0)
        # negative max_allowed = 允许 delta 至少这么少 (例 -10 → delta >= -10 且 <= 0)
        # positive max_allowed = 允许 delta 至多这么多 (例 +5 → delta <= 5 且 >= 0)
        if max_allowed < 0 and (delta >= max_allowed and delta <= 0):
            continue
        if max_allowed > 0 and (delta <= max_allowed and delta >= 0):
            continue
        if max_allowed == 0 and delta == 0:
            continue
        violations.append(
            f"{f}: {b} -> {a} (delta={delta:+d}, allowed_delta={max_allowed:+d})"
        )

    # 2. heading_counts (per name)
    hb = dict(before.get("heading_counts", {}) or {})
    ha = dict(after.get("heading_counts", {}) or {})
    h_allowed = allowed.get("heading_counts", 0)
    all_h_names = set(hb) | set(ha)
    for n in sorted(all_h_names):
        d = int(ha.get(n, 0)) - int(hb.get(n, 0))
        if d == 0:
            continue
        # heading_counts 漂移默认零容忍, 除非传 dict 或非零 int
        if isinstance(h_allowed, dict):
            allow = int(h_allowed.get(n, 0))
        else:
            allow = int(h_allowed)
        if allow < 0 and (d >= allow and d <= 0):
            continue
        if allow > 0 and (d <= allow and d >= 0):
            continue
        if allow == 0:
            violations.append(
                f"heading_counts[{n!r}]: "
                f"{hb.get(n,0)} -> {ha.get(n,0)} (delta={d:+d})"
            )

    # 3. set diff (figure/table numbers)
    for f in ("figure_number_set", "table_number_set"):
        sb = set(before.get(f, []) or [])
        sa = set(after.get(f, []) or [])
        if sb == sa:
            continue
        added = sa - sb
        removed = sb - sa
        allow = allowed.get(f, 0)
        # size_delta 是 net size change
        size_delta = len(sa) - len(sb)
        if isinstance(allow, int):
            # net size 范围允许
            if (allow >= 0 and 0 <= size_delta <= allow) or \
               (allow <= 0 and allow <= size_delta <= 0):
                # size 允许, 但若是 add+remove 混合 (内容漂移) 仍记违规
                if added and removed:
                    violations.append(
                        f"{f}: content drift "
                        f"+{sorted(added)[:5]} -{sorted(removed)[:5]} "
                        f"(size_delta={size_delta:+d}, allowed={allow:+d})"
                    )
                continue
        violations.append(
            f"{f}: size {len(sb)} -> {len(sa)} "
            f"added={sorted(added)[:5]} removed={sorted(removed)[:5]} "
            f"(allowed_size_delta={allow})"
        )

    return violations


# ───────────────────────────────────────────────────────────────────────────────
# verify (entry for fix tools)
# ───────────────────────────────────────────────────────────────────────────────
def verify_no_structural_drift(before: dict, after: dict,
                               allowed_deltas: dict | None = None,
                               raise_on_fail: bool = True) -> Tuple[bool, list[str]]:
    """跑后立即调用此函数对账 — shape_contract / 结构对账 entry.

    Args:
        before / after: capture_structure() 输出
        allowed_deltas: 字段 → 允许 delta (见 diff_structure)
        raise_on_fail: True 时违规抛 RuntimeError; False 时仅返回

    Returns:
        (passed: bool, violations: list[str])
    """
    violations = diff_structure(before, after, allowed_deltas)
    passed = not violations
    if violations and raise_on_fail:
        raise RuntimeError(
            "[shape_contract] FAIL — structural drift detected:\n  "
            + "\n  ".join(violations)
        )
    return passed, violations


def format_snapshot(snap: dict, indent: str = "  ") -> str:
    """Pretty-print snapshot for logs."""
    lines = []
    for f in _SCALAR_FIELDS:
        lines.append(f"{indent}{f:25s} = {snap.get(f, 0)}")
    h = snap.get("heading_counts", {})
    if h:
        lines.append(f"{indent}heading_counts            = {dict(sorted(h.items()))}")
    fs = snap.get("figure_number_set", [])
    ts = snap.get("table_number_set", [])
    lines.append(f"{indent}figure_number_set (n={len(fs)})  = {fs[:8]}{'...' if len(fs)>8 else ''}")
    lines.append(f"{indent}table_number_set  (n={len(ts)})  = {ts[:8]}{'...' if len(ts)>8 else ''}")
    return "\n".join(lines)


# allow standalone smoke run: python3 shape_contract.py <docx>
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("usage: shape_contract.py <docx_path>")
        sys.exit(2)
    snap = capture_structure(Path(sys.argv[1]))
    print(f"[shape_contract] snapshot of {sys.argv[1]}:")
    print(format_snapshot(snap))
