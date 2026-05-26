#!/usr/bin/env python3
# distilled from qual-supply/scripts/delete_empty_h1.py (2026-05-25 W1)
r"""delete_empty_h1.py — 删空壳 H1 段 + 迁 sectPr 到前段。

单功能: 仅删除 style 命中 H1 集合且 text.strip() 为空的段;
若该段含 <w:pPr>/<w:sectPr> 节标记, 迁移到前一段的 pPr 末尾, 不丢节布局。

何时跑:
    子报告整合后 (如 freeze_heading_numbers 完毕) 发现物理上多出空 H1
    (idx=N 段 style="Heading 1" 但 text=""), 仅作为 sectPr 承载。
    用于让 renumber_headings 重派编号前先压缩 H1 序列。

CLI:
    python3 scripts/delete_empty_h1.py <docx_path> [--dry-run] [--no-backup] [--report <json>]

启发规则:
    H1 命中样式集合: {"Heading 1", "标题 1", "heading 1", "10", "1.1.1.1 N级标题"}
    text.strip() == "" 才算空 H1。

不许做:
    - 不删 text 非空的 H1 (即使带 sectPr)
    - 不删 H2/H3/H4 / Title / 普通段
    - 不动 zdwp表名 等其他样式
    - 不裸解 zip 改 XML, 走 python-docx + lxml
    - 不用 paragraph.text = ""
"""
from __future__ import annotations

import argparse
import json
import shutil
import sys
from datetime import date
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx 未安装", file=sys.stderr)
    sys.exit(2)


H1_STYLES = {
    "Heading 1",
    "标题 1",
    "heading 1",
    "10",
    "1.1.1.1 N级标题",
}


def get_style_name(p) -> str:
    try:
        return (p.style.name or "") if p.style is not None else ""
    except Exception:
        return ""


def is_empty_h1(p) -> bool:
    sname = get_style_name(p)
    if sname not in H1_STYLES:
        return False
    text = "".join(r.text or "" for r in p.runs)
    return text.strip() == ""


def extract_sectPr(p_elem):
    """从段 <w:p> 的 <w:pPr> 里取 <w:sectPr> 子元素 (若有), 返回元素或 None。"""
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return None
    return pPr.find(qn("w:sectPr"))


def ensure_pPr(p_elem):
    """确保段有 <w:pPr>, 没有就建一个插在 <w:p> 头部, 返回 pPr 元素。"""
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        return pPr
    from lxml import etree
    pPr = etree.SubElement(p_elem, qn("w:pPr"))
    # SubElement 默认追加在末尾; 但 <w:pPr> 在 <w:p> 必须排首位 — 重排
    p_elem.remove(pPr)
    p_elem.insert(0, pPr)
    return pPr


def migrate_sectPr(empty_p_elem, prev_p_elem) -> bool:
    """把 empty_p 的 sectPr 迁到 prev_p 的 pPr 末尾。返回是否迁移过。"""
    sectPr = extract_sectPr(empty_p_elem)
    if sectPr is None:
        return False
    # 从 empty 的 pPr 删
    empty_pPr = empty_p_elem.find(qn("w:pPr"))
    if empty_pPr is not None:
        empty_pPr.remove(sectPr)
    # 装到 prev 的 pPr 末尾
    prev_pPr = ensure_pPr(prev_p_elem)
    # 若 prev_pPr 已含 sectPr (罕见 — prev 也是节末), 替掉
    existing = prev_pPr.find(qn("w:sectPr"))
    if existing is not None:
        prev_pPr.remove(existing)
    prev_pPr.append(sectPr)
    return True


def apply(doc, args=None) -> dict:
    """pipeline-compatible adapter.

    扫 doc.paragraphs 找空 H1 段; 倒序删并把 sectPr 迁到前段。
    不读写文件 (driver 已 parse / 将 save)。

    args (Namespace, 可为 None):
        dry_run: bool — True 时只 plan 不改 doc

    返回 dict: {"changed": int, "issues": [...], "deleted": [...],
                "sectPr_migrated": int, "h1_total_before": int, "h1_total_after": int}
    """
    dry_run = bool(getattr(args, "dry_run", False)) if args is not None else False
    issues: list[str] = []

    paragraphs = doc.paragraphs
    targets = [idx for idx, p in enumerate(paragraphs) if is_empty_h1(p)]

    h1_before = sum(1 for p in paragraphs if get_style_name(p) in H1_STYLES)

    report: dict = {
        "changed": 0,
        "issues": issues,
        "deleted": [],
        "sectPr_migrated": 0,
        "h1_total_before": h1_before,
        "empty_h1_idx": targets,
    }

    if not targets:
        report["h1_total_after"] = h1_before
        return report

    if dry_run:
        report["h1_total_after"] = h1_before  # plan only
        return report

    # 倒序删 (避免索引漂移)
    for idx in sorted(targets, reverse=True):
        empty_p = paragraphs[idx]
        empty_elem = empty_p._element
        parent = empty_elem.getparent()
        prev_elem = empty_elem.getprevious()
        # 跳过非 <w:p> 的 sibling (如 <w:tbl>)
        while prev_elem is not None and prev_elem.tag != qn("w:p"):
            prev_elem = prev_elem.getprevious()

        migrated = False
        if prev_elem is not None:
            migrated = migrate_sectPr(empty_elem, prev_elem)
            if migrated:
                report["sectPr_migrated"] += 1
        else:
            issues.append(f"idx={idx} 是首段, sectPr 将随删丢失")

        parent.remove(empty_elem)
        report["deleted"].append({"idx": idx, "sectPr_migrated": migrated})
        report["changed"] += 1

    # 重扫 H1 数 (paragraphs 是 live view, 删后立刻新)
    report["h1_total_after"] = sum(1 for p in doc.paragraphs if get_style_name(p) in H1_STYLES)
    return report


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="删空壳 H1 段 + 迁 sectPr 到前段")
    ap.add_argument("docx_path")
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--no-backup", action="store_true")
    ap.add_argument("--report", default=None)
    args = ap.parse_args(argv)

    docx_path = Path(args.docx_path)
    if not docx_path.exists():
        print(f"ERROR: {docx_path} 不存在", file=sys.stderr)
        return 2

    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs

    # 找空 H1 的物理索引 + 前段索引
    targets = []
    for idx, p in enumerate(paragraphs):
        if is_empty_h1(p):
            targets.append(idx)

    # 收集 H1 序列 (before)
    h1_before = []
    for idx, p in enumerate(paragraphs):
        if get_style_name(p) in H1_STYLES:
            text = "".join(r.text or "" for r in p.runs)
            h1_before.append({"idx": idx, "text": text[:60], "empty": text.strip() == ""})

    print(f"[scan] 段总数={len(paragraphs)} H1总数={len(h1_before)} 空H1={len(targets)}")
    for t in targets:
        print(f"  - idx={t} (空 H1, 待删)")

    report = {
        "total_paragraphs_before": len(paragraphs),
        "h1_total_before": len(h1_before),
        "h1_sequence_before": h1_before,
        "empty_h1_idx": targets,
        "deleted": [],
        "sectPr_migrated_to_idx": [],
    }

    if args.dry_run:
        if args.report:
            Path(args.report).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"[report] {args.report}")
        print("[dry-run] 不动 docx")
        return 0

    if not targets:
        print("[noop] 无空 H1, 无事可做")
        if args.report:
            Path(args.report).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return 0

    # backup
    if not args.no_backup:
        today = date.today().isoformat()
        n = 1
        while True:
            bak = docx_path.parent / f"{docx_path.stem}.bak-{n}-{today}.docx"
            if not bak.exists():
                shutil.copy2(docx_path, bak)
                print(f"[backup] {bak.name}")
                break
            n += 1

    # 按倒序删 (避免索引漂移)
    for idx in sorted(targets, reverse=True):
        empty_p = paragraphs[idx]
        empty_elem = empty_p._element
        # 找前段元素 (同父下)
        parent = empty_elem.getparent()
        prev_elem = empty_elem.getprevious()
        # 跳过非 <w:p> 的 sibling (如 <w:tbl>)
        while prev_elem is not None and prev_elem.tag != qn("w:p"):
            prev_elem = prev_elem.getprevious()

        migrated = False
        if prev_elem is not None:
            migrated = migrate_sectPr(empty_elem, prev_elem)
            if migrated:
                # prev_elem 在 doc.paragraphs 里的 idx 不易直接拿,记元素 hash 替代
                report["sectPr_migrated_to_idx"].append({"empty_idx": idx, "prev_elem_tag_ok": True})
        else:
            # 无前段 (文档第一段就是空 H1 — 罕见), sectPr 留 empty 段一起删 → 节信息丢
            print(f"WARN: idx={idx} 是首段, sectPr 将随删丢失", file=sys.stderr)

        # 删该段
        parent.remove(empty_elem)
        report["deleted"].append({"idx": idx, "sectPr_migrated": migrated})
        print(f"[deleted] idx={idx} sectPr_migrated={migrated}")

    doc.save(str(docx_path))
    print(f"[saved] {docx_path}")

    # verify 重读
    try:
        doc2 = Document(str(docx_path))
    except Exception as e:
        print(f"ERROR: 重读失败 (OOXML 可能损坏): {e}", file=sys.stderr)
        return 3

    h1_after = []
    for idx, p in enumerate(doc2.paragraphs):
        if get_style_name(p) in H1_STYLES:
            text = "".join(r.text or "" for r in p.runs)
            h1_after.append({"idx": idx, "text": text[:60], "empty": text.strip() == ""})
    report["total_paragraphs_after"] = len(doc2.paragraphs)
    report["h1_total_after"] = len(h1_after)
    report["h1_sequence_after"] = h1_after

    print(f"[verify] 段总数={len(doc2.paragraphs)} H1总数={len(h1_after)} (空H1={sum(1 for x in h1_after if x['empty'])})")
    for x in h1_after:
        print(f"  - idx={x['idx']} {x['text']!r}")

    if args.report:
        Path(args.report).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"[report] {args.report}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
