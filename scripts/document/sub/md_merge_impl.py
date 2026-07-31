#!/usr/bin/env python3
"""md_merge_impl.py — 将 MD 内容安全替换到 DOCX 指定章节

Distilled from panan-rigid-2026/scripts/merge_md_to_docx.py (A级通用, 2026-05-26).
零业务硬编码; 纯 python-docx XML 操作.

用法:
    python3 md_merge_impl.py <md_file> <docx_file> <start_idx> <end_idx> [output_file]

参数:
    md_file     : 要并入的 MD 文件路径
    docx_file   : 目标 DOCX 文件路径
    start_idx   : 替换起始段落索引（Heading 段落，会保留并更新标题）
    end_idx     : 替换结束段落索引（不含，即下一个章节的 Heading）
    output_file : 可选，输出文件路径，默认在 docx 同目录加 -merged 后缀

安全保证:
    - 只删除 w:p 段落元素
    - 保留 w:tbl 表格和其他非段落 XML 元素
    - 表格按前导段落文本锚点回插

触发场景:
    - 把 MD 内容合入 docx 某章节（知道起止段落索引）
    - 用新写的 MD 草稿替换 Word 交付物已有章节内容
    - 配合 `section read-section --list` 先确认段落索引再合入
"""
from __future__ import annotations

import argparse
import os
import re
import sys
import shutil
from datetime import datetime
from pathlib import Path

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[3] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 复用 md_docx_template 的 md-table 解析/边框 helper, 不重写 (铁律 #5)
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from md_docx_template import (  # noqa: E402
    parse_table_row,
    is_separator_row,
    set_table_border,
    clean_markdown_text,
)


# ─────────────────────────── 修订标记 (track changes) ───────────────────────────
# 场景：目标 docx 本身处于 Word「修订模式」协作中（正文大量 w:ins 未接受）。
# 直接删旧段/插新段会让协作方看不出改了什么，也无法逐条接受/拒绝。
# 开启后：新增内容包 <w:ins>、被替换的旧内容转 <w:del>（w:t → w:delText），
# 段落标记同步打 ins/del，与 Word 原生「修订」完全一致。


def _next_rev_id(state: dict) -> int:
    state["rid"] += 1
    return state["rid"]


def _rev_el(tag: str, state: dict) -> "OxmlElement":
    el = OxmlElement(f"w:{tag}")
    el.set(qn("w:id"), str(_next_rev_id(state)))
    el.set(qn("w:author"), state["author"])
    el.set(qn("w:date"), state["date"])
    return el


def _para_mark(p_elem, tag: str, state: dict) -> None:
    """给段落标记本身打 ins/del（表示这个段落是新增的/被删的）。"""
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.insert(0, rPr)
    rPr.append(_rev_el(tag, state))


def _wrap_runs(container, tag: str, state: dict) -> None:
    """把 container 内所有 w:r 逐个包进 <w:ins>/<w:del>。

    已被同类标记包住的跳过；del 时 w:t → w:delText（OOXML 要求）。
    对已在 w:ins 内的 run 用 del 包裹是合法的（插入后又删除）。
    """
    for run in list(container.iter(qn("w:r"))):
        parent = run.getparent()
        if parent is None or parent.tag == qn(f"w:{tag}"):
            continue
        if tag == "del":
            for t in run.findall(qn("w:t")):
                t.tag = qn("w:delText")
        idx = list(parent).index(run)
        wrapper = _rev_el(tag, state)
        parent.remove(run)
        wrapper.append(run)
        parent.insert(idx, wrapper)


def mark_inserted(elem, state: dict) -> None:
    """整个元素（段落或表格）标记为新增。"""
    if elem.tag == qn("w:tbl"):
        for tr in elem.findall(qn("w:tr")):
            trPr = tr.find(qn("w:trPr"))
            if trPr is None:
                trPr = OxmlElement("w:trPr")
                tr.insert(0, trPr)
            trPr.append(_rev_el("ins", state))
        for p in elem.iter(qn("w:p")):
            _para_mark(p, "ins", state)
        _wrap_runs(elem, "ins", state)
    else:
        _para_mark(elem, "ins", state)
        _wrap_runs(elem, "ins", state)


def mark_deleted(elem, state: dict) -> None:
    """整个元素（段落或表格）标记为删除，元素保留在文档中。"""
    if elem.tag == qn("w:tbl"):
        for tr in elem.findall(qn("w:tr")):
            trPr = tr.find(qn("w:trPr"))
            if trPr is None:
                trPr = OxmlElement("w:trPr")
                tr.insert(0, trPr)
            trPr.append(_rev_el("del", state))
        for p in elem.iter(qn("w:p")):
            _para_mark(p, "del", state)
        _wrap_runs(elem, "del", state)
    else:
        _para_mark(elem, "del", state)
        _wrap_runs(elem, "del", state)


def scan_max_rev_id(doc) -> int:
    """扫描文档现有修订/批注 id 上界，新 id 从这里往后取，避免撞号。"""
    mx = 0
    for el in doc.element.body.iter():
        v = el.get(qn("w:id"))
        if v and v.lstrip("-").isdigit():
            mx = max(mx, int(v))
    return mx


def parse_md(filepath: str) -> list:
    """解析 MD 为 block 列表,支持 标题 / 段落 / markdown 表格→w:tbl。

    block 形态:
      ("heading", level:int, text)     —— ## .. ##### → level 2..5
      ("para", text)                   —— 普通段落(经 clean_markdown_text 去 **bold**/`code`/$math$ 语法噪音)
      ("table", headers:list, rows:list[list]) —— markdown 表格(管道符 + 分隔行)

    2026-05-29 (GOAL report-automation Phase 0-A): 修复原 parse_md 只认标题+段、
    markdown 表格静默写成字面管道符 Normal 段的 🔴 缺口。表格解析复用 md_docx_template
    的 parse_table_row/is_separator_row(铁律 #5 不重写)。
    """
    blocks: list = []
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped == "":
            i += 1
            continue
        # 表格: 连续 | 开头行 (≥2 行: 表头 + 分隔 + 数据)
        if stripped.startswith("|"):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            if len(tbl_lines) >= 2:
                headers = [clean_markdown_text(c) for c in parse_table_row(tbl_lines[0])]
                rows = [
                    [clean_markdown_text(c) for c in parse_table_row(tl)]
                    for tl in tbl_lines[2:]
                    if not is_separator_row(tl)
                ]
                blocks.append(("table", headers, rows))
            else:
                for tl in tbl_lines:  # 退化: 不足 2 行不成表 → 普通段落
                    blocks.append(("para", clean_markdown_text(tl.strip())))
            continue
        # 图片 ![alt](path) — 路径相对 md 文件目录解析
        img_m = re.match(r"^!\[(.*?)\]\((.+?)\)\s*$", stripped)
        if img_m:
            alt, path = img_m.group(1), img_m.group(2).strip()
            if not os.path.isabs(path):
                path = os.path.normpath(
                    os.path.join(os.path.dirname(os.path.abspath(filepath)), path)
                )
            blocks.append(("image", path, alt))
            i += 1
            continue
        # 标题 #..######（含一级：首个 heading 用于更新被合并节的标题，见 merge_md_into_docx() 内 blocks[0] 分支）
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            blocks.append(("heading", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue
        # 普通段落
        blocks.append(("para", clean_markdown_text(stripped)))
        i += 1
    return blocks


def resolve_anchor(doc, anchor: str, *, from_idx: int = 0) -> int:
    """返回 from_idx 起首个 stripped 文本 startswith/contains anchor 的段落 idx。

    未匹配 → ValueError。用于 --start-anchor/--end-anchor 省掉手查索引那步。
    """
    norm = anchor.strip()
    for i in range(from_idx, len(doc.paragraphs)):
        t = (doc.paragraphs[i].text or "").strip()
        if t.startswith(norm) or norm in t:
            return i
    raise ValueError(f"anchor 未匹配: {anchor!r}")


# 2026-07-30 从 `apply` 改名：顶层 `apply` 是 pipeline 契约的保留名 `apply(doc, args)`。
# 本函数要 (md_file, docx_file, start_idx, end_idx, ...)，load_step 选中它必 TypeError。
def merge_md_into_docx(
    md_file: str,
    docx_file: str,
    start_idx: int,
    end_idx: int,
    output_file: str | None = None,
    in_place: bool = False,
    no_backup: bool = False,
    track_changes: bool = False,
    tc_author: str = "CC",
) -> str:
    """Merge MD into DOCX section and return output path.

    Args:
        md_file: Path to source Markdown file.
        docx_file: Path to target DOCX file.
        start_idx: Index of section heading paragraph (kept, title updated from MD).
        end_idx: Index of next section heading (exclusive; content up to here replaced).
        output_file: Optional output path. Defaults to <docx>-merged.docx.

    Returns:
        Absolute path of saved output file.
    """
    if in_place:
        if not no_backup:
            bak = f"{docx_file}.bak-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            shutil.copy2(docx_file, bak)
            print(f"已备份: {bak}")
        output_file = docx_file
    elif output_file is None:
        p = Path(docx_file)
        output_file = str(p.parent / (p.stem + "-merged" + p.suffix))

    if output_file != docx_file:
        shutil.copy2(docx_file, output_file)
    doc = Document(output_file)
    body = doc.element.body

    start_elem = doc.paragraphs[start_idx]._element
    end_elem = doc.paragraphs[end_idx]._element

    print(f"替换范围: 段落[{start_idx}]~[{end_idx - 1}]")
    print(f"  起始: {doc.paragraphs[start_idx].text[:60]}")
    print(f"  结束前: {doc.paragraphs[end_idx - 1].text[:60]}")
    print(f"  下一章: {doc.paragraphs[end_idx].text[:60]}")

    # Collect children between start and end (exclusive)
    children = []
    in_range = False
    for child in list(body):
        if child is start_elem:
            in_range = True
            continue  # keep start (section heading)
        if child is end_elem:
            break
        if in_range:
            children.append(child)

    # Classify: paragraphs vs non-paragraph elements (tables etc.)
    paras = []
    tables: list[tuple] = []  # (element, anchor_text)
    last_para_text = ""
    for child in children:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            last_para_text = "".join(child.itertext()).strip()[:80]
            paras.append(child)
        else:
            tables.append((child, last_para_text))

    print(f"\n范围内: {len(paras)} 段落, {len(tables)} 个非段落元素")

    tc_state = None
    if track_changes:
        tc_state = {
            "rid": scan_max_rev_id(doc) + 1000,  # 留空档，避开既有 id 段
            "author": tc_author,
            "date": datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"),
        }

    if track_changes:
        # 旧内容原位保留并标记为删除；新内容随后追加在其后
        for p_elem in paras:
            mark_deleted(p_elem, tc_state)
        for elem, _ in tables:
            mark_deleted(elem, tc_state)
        print(f"修订模式: {len(paras)} 段落 + {len(tables)} 非段落元素标记为删除")
        tables = []  # 已原位保留，无需回插
    else:
        # Detach non-paragraph elements then delete paragraphs
        for elem, _ in tables:
            body.remove(elem)
        for p_elem in paras:
            body.remove(p_elem)

    blocks = parse_md(md_file)

    # If MD starts with a heading, update the section heading text
    if blocks and blocks[0][0] == "heading":
        title_text = blocks[0][2]
        p_start = doc.paragraphs[start_idx]
        cur_title = (p_start.text or "").strip()
        if track_changes:
            # 修订模式下不静默重写既有标题（会抹掉协作方的修订痕迹）。
            # start_idx 落在节标题本身时：标题已在文档里 → 丢弃 MD 的重复标题。
            # start_idx 落在区间前一段时（原标题在替换区间内、已被标删）：
            #   必须把 MD 的标题作为新 block 插入，否则整节标题丢失。
            if cur_title == title_text:
                blocks = blocks[1:]
                print(f"标题已在文档中，不重复插入: {title_text!r}")
            else:
                print(f"标题作为新增内容插入: {title_text!r}（start 段为 {cur_title!r}）")
        else:
            for run in p_start.runs:
                run.text = ""
            if p_start.runs:
                p_start.runs[0].text = title_text
            blocks = blocks[1:]
            print(f"标题更新为: {title_text}")

    n_tbl = sum(1 for b in blocks if b[0] == "table")
    print(f"插入 {len(blocks)} 个 block (含 {n_tbl} 个表格)")

    # Insert blocks after the section heading (para/heading/table)
    # 修订模式下旧内容仍在原位（已标删），新内容接在其后
    ref = paras[-1] if (track_changes and paras) else start_elem
    for blk in blocks:
        if blk[0] == "heading":
            new_p = doc.add_paragraph(blk[2], style=f"Heading {min(blk[1], 5)}")
            new_elem = new_p._element
        elif blk[0] == "table":
            headers, rows = blk[1], blk[2]
            ncols = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=ncols)
            set_table_border(table)
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
            for ri, rd in enumerate(rows):
                for j, ct in enumerate(rd):
                    if j < ncols:
                        table.rows[ri + 1].cells[j].text = ct
            new_elem = table._tbl  # add_table 追加在 body 末; addnext 移到 anchor 后
        elif blk[0] == "image":
            img_path, _alt = blk[1], blk[2]
            new_p = doc.add_paragraph(style="Normal")
            if os.path.exists(img_path):
                pic = new_p.add_run().add_picture(img_path)
                maxw = Inches(5.8)  # 超过正文宽则等比缩放
                if pic.width > maxw:
                    ratio = maxw / pic.width
                    pic.width = int(pic.width * ratio)
                    pic.height = int(pic.height * ratio)
            else:
                new_p.add_run(f"[图片缺失: {img_path}]")
                print(f"WARN: 图片不存在, 插占位文本: {img_path}")
            new_elem = new_p._element
        else:  # para
            new_p = doc.add_paragraph(blk[1], style="Normal")
            new_elem = new_p._element
        if track_changes:
            mark_inserted(new_elem, tc_state)
        ref.addnext(new_elem)
        ref = new_elem

    # Reinsert non-paragraph elements by anchor text matching
    for tbl_elem, anchor in tables:
        inserted = False
        if anchor:
            for child in list(body):
                if child.tag.split("}")[-1] == "p":
                    p_text = "".join(child.itertext()).strip()[:80]
                    if anchor and anchor in p_text:
                        child.addnext(tbl_elem)
                        print(f'  非段落元素回插到 "{anchor[:40]}" 之后')
                        inserted = True
                        break
        if not inserted:
            ref.addnext(tbl_elem)
            ref = tbl_elem
            print(f'  非段落元素插到章末（锚点 "{anchor[:40]}" 未匹配）')

    doc.save(output_file)
    print(f"\n已保存到 {output_file}")

    # Verification pass
    doc2 = Document(output_file)
    print("\n--- 验证 ---")
    for i in range(
        max(0, start_idx - 1),
        min(start_idx + len(blocks) + 5, len(doc2.paragraphs)),
    ):
        p = doc2.paragraphs[i]
        if "Heading" in p.style.name:
            print(f"  [{i}] ({p.style.name}) {p.text[:80]}")

    return output_file


def apply_path(docx_path=None, args=None) -> dict:
    """pipeline: 把一份 md 合并进 docx 的指定节区间。

    2026-07-30 重写。旧版做的是三件都不该做的事：① 忽略传进来的 docx_path/args，
    直接调 `main()` 去读 **sys.argv** —— 在 pipeline 里 sys.argv 是驱动程序的参数，
    等于拿别人的命令行当自己的输入；② 把 SystemExit 和一切异常都包成
    `{"status": ...}` 返回，调用方看到的是「正常返回」，于是失败被当成成功；
    ③ 没有 `changed` 计数，报告里永远是一行「无计数」。
    现在：参数从 args 取，缺一个就抛 —— 这个动作没有能猜的默认值。
    """
    need = ("md_merge_md_file", "md_merge_start_idx", "md_merge_end_idx")
    missing = [k for k in need if getattr(args, k, None) is None]
    if missing:
        raise ValueError(
            f"md_merge_impl.apply_path 缺参数 {missing} —— "
            f"合并哪份 md、并进哪一节，没有能猜的默认值")
    if docx_path is None:
        raise ValueError("md_merge_impl.apply_path 需要 docx_path")
    out = merge_md_into_docx(
        str(getattr(args, "md_merge_md_file")), str(docx_path),
        int(getattr(args, "md_merge_start_idx")), int(getattr(args, "md_merge_end_idx")),
        output_file=getattr(args, "md_merge_output", None),
        in_place=bool(getattr(args, "md_merge_in_place", True)),
        no_backup=bool(getattr(args, "no_backup", True)),
        track_changes=bool(getattr(args, "md_merge_track_changes", False)),
        tc_author=str(getattr(args, "md_merge_tc_author", "CC")),
    )
    return {"changed": 1, "output": str(out)}


def main() -> int:
    ap = argparse.ArgumentParser(
        prog="md-merge",
        description="用 MD 内容替换 DOCX 指定节 (表格 anchor 安全回插)。"
                    "支持位置索引 或 --start-anchor/--end-anchor 按标题定位；"
                    "--in-place 原地改 + 自动 .bak-时间戳 (Work §1.5 协议)。",
    )
    ap.add_argument("md_file", help="源 Markdown")
    ap.add_argument("docx_file", help="目标 DOCX")
    ap.add_argument("start_idx", nargs="?", type=int, help="起始段落 idx (节标题段)")
    ap.add_argument("end_idx", nargs="?", type=int, help="结束段落 idx (不含, 下一节标题)")
    ap.add_argument("output_file", nargs="?", help="输出路径 (默认 <docx>-merged.docx; --in-place 时忽略)")
    ap.add_argument("--start-anchor", help="按标题文本定位 start_idx (替代位置参数)")
    ap.add_argument("--end-anchor", help="按标题文本定位 end_idx (下一节标题)")
    ap.add_argument("--in-place", action="store_true", help="原地改 + 自动备份 .bak-时间戳")
    ap.add_argument("--no-backup", action="store_true", help="配合 --in-place 跳过备份")
    ap.add_argument("--track-changes", action="store_true",
                    help="以 Word 修订标记写入：旧内容标 w:del 原位保留、新内容标 w:ins "
                         "(目标 docx 本身在修订模式协作时用)")
    ap.add_argument("--tc-author", default="CC",
                    help="修订标记署名 (配合 --track-changes, 默认 CC)")
    args = ap.parse_args()

    start_idx, end_idx = args.start_idx, args.end_idx
    if args.start_anchor or args.end_anchor:
        doc = Document(args.docx_file)
        if args.start_anchor:
            start_idx = resolve_anchor(doc, args.start_anchor)
        if args.end_anchor:
            end_idx = resolve_anchor(doc, args.end_anchor, from_idx=(start_idx or 0) + 1)
        print(f"锚点解析: start_idx={start_idx} end_idx={end_idx}")

    if start_idx is None or end_idx is None:
        print("[md-merge] 需位置参数 start_idx/end_idx, 或 --start-anchor/--end-anchor", file=sys.stderr)
        return 2

    merge_md_into_docx(args.md_file, args.docx_file, start_idx, end_idx,
          args.output_file, in_place=args.in_place, no_backup=args.no_backup,
          track_changes=args.track_changes, tc_author=args.tc_author)
    return 0


if __name__ == "__main__":
    sys.exit(main())
