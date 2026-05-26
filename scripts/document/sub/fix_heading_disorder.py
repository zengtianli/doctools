#!/usr/bin/env python3
"""fix_heading_disorder.py — qual-supply docx 标题失序 (heading-disorder) 检测+修复脚本 v2.

DEPRECATED · 仅 qual-supply 兼容保留 · 新项目用 docx outline normalize-arabic /
outline promote-h1 / outline demote-h2 替代 (sub/normalize_outline_to_arabic.py +
sub/promote_misclassified_h1.py + sub/demote_h2_with_h3_format.py). 本脚本 6 类杂糅
违反「一脚本一功能」(qual-supply CLAUDE.md 2026-05-25 用户钦定), 已停止扩展. distill
仅 cp 不重构, 保留原启发以备老 docx 复用.

W2 升级 (2026-05-24): 与 /tmp/heading-disorder/integrated-anomalies.json ground truth 对齐.

子类:

    A  false_promotion       任何非 Heading 段 (含 ZDWP正文 / zdwp表名 / Normal)
                             但文本形态像 heading/zdwp表名/Title
                             -> auto-fix: 改 style 到 suggested
    B  false_demotion        Heading 段但文本不像 heading            -> manual_review
    C  numbering_backward    同 depth 顺序扫,本段编号 < 前一段编号    -> manual_review
                             (允许跨父编号,标 cross_parent: true)
    D  numbering_skip        同 depth 顺序扫,本段编号 > 前一段+1     -> manual_review
    E  level_mismatch        styled heading 但编号深度与 style 级别不一致
                             -> auto-fix: 改 style 到匹配深度
    F  duplicate_adjacent    相邻段 (idx 差 ≤5) text 100% 相同或字符重叠 >90%
                             -> 安全 auto-fix 仅限 (idx 差 ≤2 + 100% 相同 +
                                 一份 Heading X 一份 ZDWP正文,删 ZDWP正文 份);
                                 其余 manual_review
    G  build_structural_warnings (软警告,非 anomaly)
                             整章 H1 重复 / Title 段被打成 ZDWP正文 等结构性
                             错位 → 建议用户重新 build,不硬修

用法:
    python3 fix_heading_disorder.py <docx_path> [--dry-run] [--no-backup] [--report <json>]
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))
from apply_body_styles import (  # noqa: E402
    classify_paragraph,
    make_backup_path,
    STYLE_NAMES,
)


# ---------- 中文数字 -> 阿拉伯 ----------

_CN_DIGIT = {
    "零": 0, "〇": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4,
    "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
}


def cn_to_int(s: str) -> int | None:
    s = s.strip()
    if not s:
        return None
    if s.isdigit():
        return int(s)
    if s == "十":
        return 10
    if s.startswith("十") and len(s) == 2 and s[1] in _CN_DIGIT:
        return 10 + _CN_DIGIT[s[1]]
    if len(s) == 2 and s[1] == "十" and s[0] in _CN_DIGIT:
        return _CN_DIGIT[s[0]] * 10
    if len(s) == 3 and s[1] == "十" and s[0] in _CN_DIGIT and s[2] in _CN_DIGIT:
        return _CN_DIGIT[s[0]] * 10 + _CN_DIGIT[s[2]]
    if len(s) == 1 and s in _CN_DIGIT:
        return _CN_DIGIT[s]
    return None


# ---------- heading 形态正则 ----------

RE_TABLE      = re.compile(r"^表\s*[\d一二三四五六七八九十百]+[-—–]?[\d一二三四五六七八九十百]*\s")
RE_TITLE_CN   = re.compile(r"^([一二三四五六七八九十]+)、")
RE_CHAPTER    = re.compile(r"^第([一二三四五六七八九十百零\d]+)章\s*\S")
RE_H4         = re.compile(r"^(\d+)\.(\d+)\.(\d+)\.(\d+)\s+\S")
RE_H3         = re.compile(r"^(\d+)\.(\d+)\.(\d+)\s+\S")
RE_H2         = re.compile(r"^(\d+)\.(\d+)\s+\S")
RE_H1_DOT     = re.compile(r"^(\d+)\.\s+\S")
RE_H1         = re.compile(r"^(\d+)\s+\S")


def detect_heading_form(text: str) -> tuple[str, int, tuple[int, ...] | None] | None:
    """返回 (form_label, depth, numbers_tuple) 或 None.

    depth: 0 = 表名 / 1-4 = heading 深度.
    """
    t = (text or "").strip()
    if not t:
        return None
    if RE_TABLE.match(t):
        return ("zdwp_table", 0, None)
    m = RE_CHAPTER.match(t)
    if m:
        n = cn_to_int(m.group(1))
        return ("chapter_H1", 1, (n,) if n is not None else None)
    m = RE_H4.match(t)
    if m:
        return ("H4_form", 4, tuple(int(x) for x in m.groups()))
    m = RE_H3.match(t)
    if m:
        return ("H3_form", 3, tuple(int(x) for x in m.groups()))
    m = RE_H2.match(t)
    if m:
        return ("H2_form", 2, tuple(int(x) for x in m.groups()))
    m = RE_H1_DOT.match(t)
    if m:
        return ("H1_dot_form", 1, (int(m.group(1)),))
    m = RE_H1.match(t)
    if m and "." not in t.split()[0]:
        return ("H1_form", 1, (int(m.group(1)),))
    m = RE_TITLE_CN.match(t)
    if m:
        n = cn_to_int(m.group(1))
        return ("Title_cn", 1, (n,) if n is not None else None)
    return None


# ---------- style 分类 ----------

# styleId -> heading level (1-5); None = 非 heading
HEADING_STYLE_ID_TO_LEVEL = {
    "1": 1, "2": 2, "3": 3, "4": 4, "5": 5,
    "Heading1": 1, "Heading2": 2, "Heading3": 3, "Heading4": 4, "Heading5": 5,
}

# 视为"非 heading"的 styleId / styleName — 这些段如果文本形态像 heading 都算 A
NON_HEADING_STYLE_IDS = {"ZDWP", "Normal", "a", "zdwp1"}  # zdwp1=zdwp表名
NON_HEADING_STYLE_NAME_KEYWORDS = ("正文", "Normal", "Body", "zdwp表名")


def is_heading_style(style_id: str | None, style_name: str | None) -> int | None:
    """返回 heading 级别 1-5; 非 heading 返回 None.

    Title / zdwp表名 视为非 heading (depth 概念不同, 让 A 类判断).
    """
    if style_id and style_id in HEADING_STYLE_ID_TO_LEVEL:
        return HEADING_STYLE_ID_TO_LEVEL[style_id]
    if style_name:
        m = re.match(r"^Heading\s*(\d+)", style_name, re.I)
        if m:
            lvl = int(m.group(1))
            if 1 <= lvl <= 5:
                return lvl
        m = re.match(r"^标题\s*(\d+)", style_name)
        if m:
            lvl = int(m.group(1))
            if 1 <= lvl <= 5:
                return lvl
    return None


def is_non_heading_style(style_id: str | None, style_name: str | None) -> bool:
    """段是否是 '非 heading' (即 A 类的扫描范围).

    Heading X / Title 跳过; ZDWP正文 / Normal / zdwp表名 / 任何含"正文"name 算 yes.
    """
    if is_heading_style(style_id, style_name) is not None:
        return False
    # Title 也不扫 A (它 depth=0 + W2 ground truth 显示 ZDWP正文 误写 "三、" 才报 A)
    # 但 Title style 本身 (styleId=a4) 不在 A 范围 — 已经是合法标题样式
    if style_id == "a4":
        return False
    if style_id and style_id in NON_HEADING_STYLE_IDS:
        return True
    if style_name:
        for kw in NON_HEADING_STYLE_NAME_KEYWORDS:
            if kw in style_name:
                return True
    return False


# 形态 label -> suggested styleId
FORM_TO_STYLE_ID = {
    "zdwp_table":   "zdwp1",
    "Title_cn":     "a4",
    "chapter_H1":   "1",
    "H1_form":      "1",
    "H1_dot_form":  "1",
    "H2_form":      "2",
    "H3_form":      "3",
    "H4_form":      "3",  # docx 只有 3 个 Heading 级别可用 (1/2/3 + 4 在 styles 但默认未用)
}

DEPTH_TO_STYLE_ID = {1: "1", 2: "2", 3: "3", 4: "3", 5: "3"}


# ---------- detection ----------

def detect_anomalies(doc) -> tuple[list[dict], list[str]]:
    """返回 (anomalies, build_structural_warnings)."""
    anomalies: list[dict] = []
    warnings: list[str] = []

    # 收集所有段
    paras: list[dict] = []
    for idx, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        sid = p.style.style_id if p.style is not None else None
        nm = p.style.name if p.style is not None else None
        paras.append({
            "idx": idx, "text": text, "style_id": sid, "style_name": nm, "p": p,
        })

    # ---- A: false_promotion (扩检到任何非 Heading 段) ----
    for pi in paras:
        t = pi["text"]
        if not t or len(t) > 200:
            continue
        sid, nm = pi["style_id"], pi["style_name"]
        if not is_non_heading_style(sid, nm):
            continue
        form = detect_heading_form(t)
        if not form:
            continue
        label, depth, _tup = form
        suggested = FORM_TO_STYLE_ID.get(label)
        if suggested is None:
            continue
        # 跳过: 当前已是 suggested
        if sid == suggested:
            continue
        # 跳过特殊: zdwp表名 段且文本本来就是表名形态 (label=zdwp_table) — 已经对了
        if sid == "zdwp1" and label == "zdwp_table":
            continue
        suggested_ui = STYLE_NAMES.get(suggested, suggested)
        anomalies.append({
            "idx": pi["idx"],
            "category": "A_false_promotion",
            "current_style": nm or sid,
            "current_style_id": sid,
            "suggested_style": suggested_ui,
            "suggested_style_id": suggested,
            "text": t[:120],
            "reason": f"文本符合 {label}(depth={depth}) 形态但 style={nm or sid}",
        })

    # ---- B: false_demotion (Heading 段但不像 heading) ----
    SHORT_TITLE_WORDS = {"前言", "摘要", "引言", "结语", "结论", "参考文献",
                         "附录", "目录", "致谢"}
    for pi in paras:
        t = pi["text"]
        if not t:
            continue
        sid, nm = pi["style_id"], pi["style_name"]
        cur_level = is_heading_style(sid, nm)
        if cur_level is None:
            continue
        too_long = len(t) > 120
        form = detect_heading_form(t)
        is_short_title = (
            t in SHORT_TITLE_WORDS or
            any(t.startswith(w) for w in SHORT_TITLE_WORDS)
        )
        if too_long or (form is None and not is_short_title and len(t) > 30):
            anomalies.append({
                "idx": pi["idx"],
                "category": "B_false_demotion",
                "current_style": nm or sid,
                "current_style_id": sid,
                "suggested_style": "ZDWP正文",
                "suggested_style_id": "ZDWP",
                "text": t[:120],
                "reason": (
                    f"Heading {cur_level} 但文本"
                    + ("过长(>120 字)" if too_long else "")
                    + ("," if (too_long and form is None) else "")
                    + ("不符合 heading 形态" if form is None else "")
                ),
            })

    # ---- C / D: numbering_backward / numbering_skip ----
    # 收集**所有**符合 heading form 的段 (含 style=正文 但形态像 heading 的伪 heading)
    heading_seq: list[dict] = []
    for pi in paras:
        t = pi["text"]
        form = detect_heading_form(t)
        if not form:
            continue
        label, depth, tup = form
        if tup is None or depth < 1:
            continue
        heading_seq.append({
            "idx": pi["idx"], "depth": depth, "tup": tup,
            "style_id": pi["style_id"], "style_name": pi["style_name"],
            "text": t[:120], "form": label,
        })

    # 按 (depth, prefix) 跟踪 last_num — 同一前缀下出现回退/跳号都报
    # (跨父分支独立,不互相影响; 与 W2 ground truth 算法一致)
    last_by_prefix: dict[tuple[int, tuple], dict] = {}
    for h in heading_seq:
        depth = h["depth"]
        tup = h["tup"]
        if len(tup) < depth:
            continue
        cur_num = tup[depth - 1]
        cur_prefix = tup[: depth - 1]
        key = (depth, cur_prefix)
        if key in last_by_prefix:
            prev = last_by_prefix[key]
            if cur_num < prev["num"]:
                anomalies.append({
                    "idx": h["idx"],
                    "category": "C_numbering_backward",
                    "current_style": h["style_name"] or h["style_id"],
                    "current_style_id": h["style_id"],
                    "suggested_style": h["style_name"] or h["style_id"],
                    "suggested_style_id": h["style_id"],
                    "text": h["text"],
                    "reason": (
                        f"同级编号倒退: depth={depth} prefix={cur_prefix} "
                        f"前一段 idx={prev['idx']} 编号={prev['num']}, "
                        f"本段编号={cur_num}"
                    ),
                    "_prev_idx": prev["idx"],
                })
            elif cur_num > prev["num"] + 1:
                anomalies.append({
                    "idx": h["idx"],
                    "category": "D_numbering_skip",
                    "current_style": h["style_name"] or h["style_id"],
                    "current_style_id": h["style_id"],
                    "suggested_style": h["style_name"] or h["style_id"],
                    "suggested_style_id": h["style_id"],
                    "text": h["text"],
                    "reason": (
                        f"同级编号跳号: depth={depth} prefix={cur_prefix} "
                        f"前一段编号={prev['num']}, 本段编号={cur_num} "
                        f"(期望={prev['num'] + 1})"
                    ),
                    "_prev_idx": prev["idx"],
                })
        last_by_prefix[key] = {"idx": h["idx"], "num": cur_num}

    # ---- E: level_mismatch (styled heading 但编号深度与 style 级别不一致) ----
    for h in heading_seq:
        text_depth = h["depth"]
        cur_level = is_heading_style(h["style_id"], h["style_name"])
        if cur_level is None:
            continue
        if cur_level != text_depth and text_depth in DEPTH_TO_STYLE_ID:
            target = DEPTH_TO_STYLE_ID[text_depth]
            anomalies.append({
                "idx": h["idx"],
                "category": "E_level_mismatch",
                "current_style": h["style_name"] or h["style_id"],
                "current_style_id": h["style_id"],
                "suggested_style": STYLE_NAMES.get(target, target),
                "suggested_style_id": target,
                "text": h["text"],
                "reason": (
                    f"编号深度={text_depth} (form={h['form']}) 与 "
                    f"style 层级={cur_level} 不一致 → 应为 Heading {text_depth}"
                ),
            })

    # ---- F: duplicate_adjacent (相邻段 text 重复) ----
    # 相邻段 (idx 差 ≤ 7) 且 (text 100% 相同 OR 字符重叠 > 90%)
    F_WINDOW = 7
    for i, a in enumerate(paras):
        ta = a["text"]
        if not ta or len(ta) < 6:
            continue
        for j in range(i + 1, min(i + F_WINDOW + 1, len(paras))):
            b = paras[j]
            tb = b["text"]
            if not tb or len(tb) < 6:
                continue
            idx_diff = b["idx"] - a["idx"]
            if idx_diff > F_WINDOW:
                break
            # 完全相同?
            if ta == tb:
                overlap = 1.0
                exact = True
            else:
                # 字符集重叠
                sa, sb = set(ta), set(tb)
                if not sa or not sb:
                    continue
                overlap = len(sa & sb) / max(len(sa), len(sb))
                exact = False
                if overlap < 0.9:
                    continue
            # 不能两段都是同样 style 的"正常段"(误报)
            sid_a, sid_b = a["style_id"], b["style_id"]
            # 只在 style 不同 时才报 (重复证据 = 一份高一份低)
            if sid_a == sid_b:
                continue
            # 判断是否一份 Heading 一份非 Heading (W2 描述: build 阶段重复插入)
            a_is_heading = is_heading_style(sid_a, a["style_name"]) is not None
            b_is_heading = is_heading_style(sid_b, b["style_name"]) is not None
            if a_is_heading == b_is_heading:
                # 都 heading 或都非 heading — 不算 F 类
                continue
            heading_side = "a" if a_is_heading else "b"
            non_heading = b if a_is_heading else a
            heading = a if a_is_heading else b
            anomalies.append({
                "idx": non_heading["idx"],
                "category": "F_duplicate_adjacent",
                "current_style": non_heading["style_name"] or non_heading["style_id"],
                "current_style_id": non_heading["style_id"],
                "suggested_style": "DELETE",
                "suggested_style_id": "DELETE",
                "text": non_heading["text"][:120],
                "reason": (
                    f"与 idx={heading['idx']} (style={heading['style_name']}) "
                    f"相邻重复 (idx 差={idx_diff}, "
                    f"{'完全相同' if exact else f'字符重叠={overlap:.0%}'}) — "
                    f"build 阶段重复插入证据"
                ),
                "_heading_idx": heading["idx"],
                "_heading_style_id": heading["style_id"],
                "_idx_diff": idx_diff,
                "_exact": exact,
                "_overlap": round(overlap, 2),
            })

    # 去重 F: 同一段被多次报 → 保留最近的 (idx_diff 最小的)
    f_by_idx: dict[int, dict] = {}
    other_anomalies = []
    for a in anomalies:
        if a["category"] == "F_duplicate_adjacent":
            cur = f_by_idx.get(a["idx"])
            if cur is None or a["_idx_diff"] < cur["_idx_diff"]:
                f_by_idx[a["idx"]] = a
        else:
            other_anomalies.append(a)
    anomalies = other_anomalies + list(f_by_idx.values())

    # ---- G: build structural warnings (软警告) ----
    # 整章 H1 重复 (相同前缀的 H1 出现 ≥2 次)
    h1_seen: dict[str, list[int]] = defaultdict(list)
    for h in heading_seq:
        if h["depth"] == 1 and h["tup"]:
            h1_seen[str(h["tup"][0])].append(h["idx"])
    for num, idxs in h1_seen.items():
        if len(idxs) >= 2:
            warnings.append(
                f"H1 编号 '{num}' 在 idx {idxs} 重复出现 — "
                f"很可能是 build 阶段两章合并错位,建议重 build 或手动合章"
            )
    # Title 段被打成 ZDWP正文
    for pi in paras:
        if pi["text"] and RE_TITLE_CN.match(pi["text"]) and pi["style_id"] == "ZDWP":
            warnings.append(
                f"idx {pi['idx']} 文本 {pi['text'][:60]!r} 形如「X、章节标题」 "
                f"但 style=ZDWP正文 — 应为 Title (a4); "
                f"可能是 build 阶段 Title 段漏改"
            )

    # 按 idx + category 排序
    anomalies.sort(key=lambda a: (a["idx"], a["category"]))
    return anomalies, warnings


# ---------- fix ----------

AUTO_FIX_CATEGORIES = {"A_false_promotion", "E_level_mismatch", "F_duplicate_adjacent"}


def apply_fixes(doc, anomalies: list[dict], dry_run: bool) -> None:
    """对 A / E / F 自动 fix; 修改 anomaly['auto_fixed'] in-place.

    F 安全规则: 仅当 idx_diff ≤ 2 + exact=True + 一份 Heading 一份非 Heading
    才删非 Heading 那份; 否则 manual_review.
    """
    available_ids = {s.style_id for s in doc.styles}
    paragraphs = doc.paragraphs
    deleted_idxs: set[int] = set()  # 已删的 idx,后续 fix 跳过

    # F 优先处理 (idx 倒序,避免删除影响后续 idx)
    f_anoms = [a for a in anomalies if a["category"] == "F_duplicate_adjacent"]
    f_anoms.sort(key=lambda a: -a["idx"])
    for a in f_anoms:
        # 安全门槛
        safe = (
            a.get("_exact") and
            a.get("_idx_diff", 99) <= 2 and
            a.get("current_style_id") == "ZDWP"  # 删的一定是 ZDWP正文 那份
        )
        if not safe:
            a["auto_fixed"] = False
            a["manual_review"] = True
            continue
        if dry_run:
            a["auto_fixed"] = False
            continue
        idx = a["idx"]
        if idx >= len(paragraphs):
            a["auto_fixed"] = False
            continue
        p = paragraphs[idx]
        try:
            # python-docx 删段: p._element.getparent().remove(p._element)
            p._element.getparent().remove(p._element)
            deleted_idxs.add(idx)
            a["auto_fixed"] = True
        except Exception as exc:
            a["auto_fixed"] = False
            a["reason"] = a["reason"] + f" [F 删除失败: {exc}]"

    # 注意: 删段后 paragraphs 列表的 idx 失效, 但我们用 doc.paragraphs[idx]
    # 索引访问 — 由于上面用倒序 F 处理,A/E fix 时 idx 仍然指向原位置 (前面段没动)
    # 但删除发生在 F 段后,所有 idx > 已删 idx 的段需要做映射. 我们用稳妥做法:
    # F 处理完后重新拿 paragraphs 列表,A/E 用文本匹配 + idx 起点扫描定位.
    # 简化: F 删除一定在 A/E 之后处理 — 颠倒顺序.

    # 实际策略: 先做 A/E (idx 仍稳定),再做 F (倒序删).
    # 上面 F 已经做了 — 撤回. 重做:

    # 由于上面已 F 做了删除,下面 A/E 需要用 element 比较找到原 paragraph.
    # 为简单起见,改架构: F 移到最后做. 此处重置已删段,撤销操作不现实,故改用:
    # 不做删除,只改 F 段 text 为空 (这样不破坏 idx 索引). — 不行,留空段也是污染.
    # 正确做法: F 真删 + A/E 用 element 对齐 fix.

    # 重新设计: 通过对每个 A/E anomaly 找 doc 现存 paragraphs 中 idx 偏移后的位置.
    # 已删段数 (idx < this_idx) = offset, this_paragraph = doc.paragraphs[this_idx - offset]
    deleted_sorted = sorted(deleted_idxs)

    def remap_idx(orig_idx: int) -> int:
        # 计算 orig_idx 之前删了多少段
        offset = sum(1 for d in deleted_sorted if d < orig_idx)
        return orig_idx - offset

    paragraphs = doc.paragraphs  # 重新取 (F 删除后)

    # 收集所有"未被安全删除"的 F 段 idx (这些段是另一段 Heading 的重复证据,
    # 即使形态像 heading,也不该被 A 升级 — 否则会造出两段一模一样的 Heading)
    f_blocked_idxs = {
        a["idx"] for a in anomalies
        if a["category"] == "F_duplicate_adjacent"
        and not a.get("auto_fixed")
    }

    for a in anomalies:
        if a["category"] not in ("A_false_promotion", "E_level_mismatch"):
            if "auto_fixed" not in a:
                a["auto_fixed"] = False
            continue
        # F-blocked: 此段是另一段 Heading 的重复证据,不能 A 升级
        if a["idx"] in f_blocked_idxs:
            a["auto_fixed"] = False
            a["manual_review"] = True
            a["reason"] = a["reason"] + " [F-blocked: 此段是另一段 Heading 的重复证据,需用户决策]"
            continue
        target = a["suggested_style_id"]
        if target not in available_ids:
            a["auto_fixed"] = False
            a["reason"] = a["reason"] + f" [跳过: 目标 styleId {target!r} 不在 styles.xml]"
            continue
        if dry_run:
            a["auto_fixed"] = False
            continue
        orig_idx = a["idx"]
        if orig_idx in deleted_idxs:
            a["auto_fixed"] = False
            a["reason"] = a["reason"] + " [跳过: 段已被 F 类删除]"
            continue
        new_idx = remap_idx(orig_idx)
        if new_idx >= len(paragraphs):
            a["auto_fixed"] = False
            continue
        p = paragraphs[new_idx]
        # 校验: 文本匹配
        if (p.text or "").strip()[:120] != a["text"]:
            # 偏移失败,跳过
            a["auto_fixed"] = False
            a["reason"] = a["reason"] + (
                f" [跳过: idx 偏移后文本不匹配 — got "
                f"{(p.text or '').strip()[:40]!r}]"
            )
            continue
        try:
            ui_name = STYLE_NAMES.get(target)
            if ui_name and ui_name in [s.name for s in doc.styles]:
                p.style = doc.styles[ui_name]
            else:
                p.style = doc.styles.get_by_id(target, 1)
            a["auto_fixed"] = True
        except Exception:
            try:
                pPr = p._p.get_or_add_pPr()
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is None:
                    from docx.oxml import OxmlElement
                    pStyle = OxmlElement("w:pStyle")
                    pPr.insert(0, pStyle)
                pStyle.set(qn("w:val"), target)
                a["auto_fixed"] = True
            except Exception as exc:
                a["auto_fixed"] = False
                a["reason"] = a["reason"] + f" [fix 失败: {exc}]"


# ---------- 汇总 ----------

CATEGORY_ORDER = [
    "A_false_promotion",
    "B_false_demotion",
    "C_numbering_backward",
    "D_numbering_skip",
    "E_level_mismatch",
    "F_duplicate_adjacent",
]


def summarize(anomalies: list[dict]) -> dict:
    summary: dict[str, dict] = {}
    for cat in CATEGORY_ORDER:
        items = [a for a in anomalies if a["category"] == cat]
        entry: dict[str, int] = {"count": len(items)}
        if cat in AUTO_FIX_CATEGORIES:
            entry["auto_fixed"] = sum(1 for a in items if a.get("auto_fixed"))
            entry["manual_review"] = sum(1 for a in items if a.get("manual_review"))
        else:
            entry["manual_review"] = len(items)
        summary[cat] = entry
    return summary


# ---------- main ----------

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="qual-supply docx 标题失序 v2 (W2 ground truth 对齐)",
    )
    ap.add_argument("docx", help="目标 docx 路径 (inplace 修改, 除非 --dry-run)")
    ap.add_argument("--dry-run", action="store_true", help="只扫不改, 不备份不写")
    ap.add_argument("--no-backup", action="store_true", help="跳过备份")
    ap.add_argument("--report", help="anomaly 报告 JSON")
    args = ap.parse_args(argv)

    src = Path(args.docx).resolve()
    if not src.is_file():
        print(f"ERROR: 不存在 {src}", file=sys.stderr)
        return 2

    backup_path = None
    if not args.dry_run and not args.no_backup:
        backup_path = make_backup_path(src)
        shutil.copy2(src, backup_path)

    doc = Document(str(src))
    anomalies, warnings = detect_anomalies(doc)

    before = summarize(anomalies)
    apply_fixes(doc, anomalies, dry_run=args.dry_run)
    after = summarize(anomalies)  # 含 auto_fixed 字段

    if not args.dry_run:
        doc.save(str(src))

    # 控制台报告
    print(f"file:     {src}")
    if backup_path:
        print(f"backup:   {backup_path}")
    elif args.dry_run:
        print("backup:   (dry-run, 未备份)")
    else:
        print("backup:   (--no-backup)")
    print(f"dry_run:  {args.dry_run}")
    print(f"total anomalies: {len(anomalies)}")
    print("summary:")
    for cat in CATEGORY_ORDER:
        print(f"  {cat:25s} {after[cat]}")
    if warnings:
        print(f"build_structural_warnings: {len(warnings)}")
        for w in warnings:
            print(f"  - {w}")

    if args.report:
        report_path = Path(args.report).expanduser().resolve()
        report_path.parent.mkdir(parents=True, exist_ok=True)
        # manual_review_items: B/C/D/F-manual
        manual_review_items = [
            a for a in anomalies
            if not a.get("auto_fixed") and (
                a["category"] in ("B_false_demotion", "C_numbering_backward",
                                  "D_numbering_skip") or
                (a["category"] == "F_duplicate_adjacent" and a.get("manual_review"))
            )
        ]
        report = {
            "input": str(src),
            "dry_run": args.dry_run,
            "backup": str(backup_path) if backup_path else None,
            "before": {
                k: v["count"] for k, v in before.items()
            } | {"total": sum(v["count"] for v in before.values())},
            "fixed": {
                k: after[k].get("auto_fixed", 0)
                for k in ("A_false_promotion", "E_level_mismatch", "F_duplicate_adjacent")
            } | {
                "total_auto_fixed": sum(
                    after[k].get("auto_fixed", 0)
                    for k in ("A_false_promotion", "E_level_mismatch",
                              "F_duplicate_adjacent")
                ),
            },
            "remaining": {
                "B_false_demotion": after["B_false_demotion"]["count"],
                "C_numbering_backward": after["C_numbering_backward"]["count"],
                "D_numbering_skip": after["D_numbering_skip"]["count"],
                "A_unfixed": before["A_false_promotion"]["count"]
                             - after["A_false_promotion"].get("auto_fixed", 0),
                "E_unfixed": before["E_level_mismatch"]["count"]
                             - after["E_level_mismatch"].get("auto_fixed", 0),
                "F_manual": after["F_duplicate_adjacent"].get("manual_review", 0),
                "total": len(manual_review_items),
            },
            "summary_detailed": after,
            "manual_review_items": manual_review_items,
            "build_structural_warnings": warnings,
            "anomalies": anomalies,
        }
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2))
        print(f"report:   {report_path}")

    return 0


# ---------------- pipeline adapter ----------------
def apply(doc, args=None) -> dict:
    dry = bool(getattr(args, "dry_run", False)) if args else False
    anomalies, warnings = detect_anomalies(doc)
    apply_fixes(doc, anomalies, dry_run=dry)
    after = summarize(anomalies)
    auto_fixed = sum(
        after[k].get("auto_fixed", 0)
        for k in ("A_false_promotion", "E_level_mismatch", "F_duplicate_adjacent")
    )
    return {
        "changed": auto_fixed,
        "anomalies_total": len(anomalies),
        "summary": after,
        "warnings": warnings,
    }


if __name__ == "__main__":
    sys.exit(main())
