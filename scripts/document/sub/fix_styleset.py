"""fix_styleset.py — group module: style-set fix family (7 subcommands · W13/W15 2026-05-26)

shape_contract / 结构对账 / _verify_structure_invariants (GOAL I6 anchors)

Subcommands:
  fix style-rebrand        ← 批量段样式迁移 (Normal → 项目正文样式, 或按 --role 取 profile.roles 主样式)
  fix style-pool-cleanup   ← 删 docx 内"定义但段落未用且非系统 default"的样式
  fix style-pane-filter    ← 设 word/settings.xml 的 stylePaneFormatFilter, 白名单 profile.roles 样式
  fix role-fill            ← 检查 profile.roles 每角色; 缺则按模板克隆建对应样式
  fix style-rename         ← (W15) 改样式 .name 字段, 不动 styleId, 不动段引用
  fix clear-direct-format  ← (W15) 清段 inline 直接格式 (pPr/rPr 直接子元素), 只保留 pStyle/rStyle
  fix style-create         ← (W15) 按 base style 克隆新空 style 定义到 styles.xml

CLI 通用 args (per subcommand):
    <docx_path> [--dry-run] [--inplace] [--no-backup] [--force] [--report json]
    --profile <name>          (default zdwp; 走 doctools.lib.styles.load_profile)
    --from / --to / --role / 其它 per-subcommand

默认 --dry-run; --inplace 才真改 (自动留 .bak-N-DATE).
shape_contract 跑前后对账, 漂移 → exit 3 (除非 --force, 打 4 行黄底 WARNING).
allowed_deltas:
  style-rebrand:       全 0 (段数 / heading 数 / caption 集合不变)
  style-pool-cleanup:  段数 0; styles 定义减是允许的 (不进 shape_contract, 而是另算)
  style-pane-filter:   全 0 (settings.xml 改不影响 body 结构)
  role-fill:           段数 0; styles 定义增是允许的 (另算)
"""

from __future__ import annotations

import argparse
import datetime as _dt
import json
import os
import shutil
import subprocess
import sys
import zipfile
from collections import Counter
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from .shape_contract import (
    capture_structure,
    diff_structure,
    verify_no_structural_drift,
    format_snapshot,
)

# styles profile loader (same fallback pattern as sub/styles.py)
try:
    from doctools.lib.styles import load_profile, StylesProfile  # type: ignore
except ImportError:
    import importlib.util as _ilu
    _styles_path = (Path(__file__).resolve().parent.parent.parent.parent
                    / "lib" / "styles.py")
    _spec = _ilu.spec_from_file_location("_doctools_styles", _styles_path)
    _m = _ilu.module_from_spec(_spec)
    _spec.loader.exec_module(_m)  # type: ignore
    load_profile = _m.load_profile
    StylesProfile = _m.StylesProfile


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ═════════════════════════════════════════════════════════════════════════════
# common helpers (lsof / backup / save / report)
# ═════════════════════════════════════════════════════════════════════════════
def _lsof_check(path: Path) -> Optional[str]:
    try:
        r = subprocess.run(["lsof", str(path)], capture_output=True,
                           text=True, timeout=5)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def _pick_backup_path(src: Path) -> Path:
    today = _dt.date.today().isoformat()
    parent, stem, suffix = src.parent, src.stem, src.suffix
    n = 1
    while True:
        cand = parent / f"{stem}.bak-{n}-{today}{suffix}"
        if not cand.exists():
            return cand
        n += 1


def _will_write(args) -> bool:
    """True iff actually writing (inplace + not dry-run)."""
    return bool(getattr(args, "inplace", False)) and not getattr(args, "dry_run", False)


def _common_setup(args) -> Path:
    src: Path = Path(args.docx_path) if hasattr(args, "docx_path") else Path(args.docx)
    if not src.exists():
        print(f"[ERR] 文件不存在: {src}", file=sys.stderr)
        sys.exit(2)
    src = src.resolve()
    if _will_write(args):
        occ = _lsof_check(src)
        if occ:
            print(f"[ABORT] 文件被占用 (Word/WPS):\n{occ}", file=sys.stderr)
            sys.exit(3)
    return src


def _save_with_backup(src: Path, doc, args) -> Optional[Path]:
    if not getattr(args, "inplace", False) or getattr(args, "dry_run", False):
        return None
    bak = None
    if not getattr(args, "no_backup", False):
        bak = _pick_backup_path(src)
        shutil.copy2(src, bak)
    doc.save(str(src))
    return bak


def _emit_report(report: dict, args):
    rp = getattr(args, "report", None)
    if not rp:
        return
    p = Path(rp)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps(report, ensure_ascii=False, indent=2, default=str),
                 encoding="utf-8")


def _force_warning(reason: str):
    sys.stderr.write("\033[1;33m" + "!" * 78 + "\033[0m\n")
    sys.stderr.write(f"\033[1;33m!! --force 旁路 shape_contract: {reason}\033[0m\n")
    sys.stderr.write("\033[1;33m!! 文档可能被破坏; 你最好知道自己在干什么\033[0m\n")
    sys.stderr.write("\033[1;33m" + "!" * 78 + "\033[0m\n")
    sys.stderr.flush()


def _refuse_msg(violations: list[str]) -> str:
    lines = [
        "",
        "\033[1;31m" + "=" * 78 + "\033[0m",
        "\033[1;31m[REFUSED] shape_contract — 结构对账失败, 拒绝写盘\033[0m",
        "\033[1;31m" + "=" * 78 + "\033[0m",
    ]
    for v in violations:
        lines.append(f"  · {v}")
    lines += [
        "",
        "  跑下去会破坏 docx 结构 (段数 / heading 数 / caption / 图片对象漂移).",
        "  --force 旁路 (不推荐).",
        "",
    ]
    return "\n".join(lines) + "\n"


# ═════════════════════════════════════════════════════════════════════════════
# helpers · styles introspection
# ═════════════════════════════════════════════════════════════════════════════
def _all_para_style_ids(doc) -> set[str]:
    """所有段落实际引用的 styleId 集合 (含 cell 段)."""
    ids: set[str] = set()
    for p in doc.paragraphs:
        sid = getattr(p.style, "style_id", None) if p.style else None
        if sid:
            ids.add(sid)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    sid = getattr(p.style, "style_id", None) if p.style else None
                    if sid:
                        ids.add(sid)
    return ids


def _all_run_style_ids(doc) -> set[str]:
    """所有 run-level rStyle 引用 (char styles)."""
    ids: set[str] = set()
    body = doc.element.body
    for r in body.findall(".//" + qn("w:rStyle")):
        v = r.get(qn("w:val"))
        if v:
            ids.add(v)
    return ids


def _all_table_style_ids(doc) -> set[str]:
    ids: set[str] = set()
    body = doc.element.body
    for ts in body.findall(".//" + qn("w:tblStyle")):
        v = ts.get(qn("w:val"))
        if v:
            ids.add(v)
    return ids


def _style_id_referenced_in_xml_part(docx_path: Path, part_name: str,
                                     style_id: str) -> bool:
    """Check if styleId appears in given zip part (numbering.xml / settings.xml etc.)."""
    try:
        with zipfile.ZipFile(str(docx_path)) as z:
            if part_name not in z.namelist():
                return False
            xml = z.read(part_name)
    except (zipfile.BadZipFile, KeyError):
        return False
    # cheap substring match (styles 引用都是 w:val="<id>" 形式)
    needle = f'w:val="{style_id}"'.encode("utf-8")
    return needle in xml


def _styles_defined(doc) -> list[tuple[str, str, str]]:
    """List of (style_id, style_name, type) for all defined styles."""
    out = []
    for s in doc.styles:
        sid = getattr(s, "style_id", "") or ""
        name = getattr(s, "name", "") or ""
        try:
            stype = s.type
            stype_name = stype.name if hasattr(stype, "name") else str(stype)
        except Exception:
            stype_name = "?"
        out.append((sid, name, str(stype_name)))
    return out


def _profile_role_styles(profile: StylesProfile) -> dict[str, list[str]]:
    """Extract per-role candidate style names/ids from profile *_STYLES fields."""
    return {
        "body":          list(getattr(profile, "BODY_STYLES", []) or []),
        "h1":            list(getattr(profile, "H1_STYLES", []) or []),
        "h2":            list(getattr(profile, "H2_STYLES", []) or []),
        "h3":            list(getattr(profile, "H3_STYLES", []) or []),
        "h4":            list(getattr(profile, "H4_STYLES", []) or []),
        "title":         list(getattr(profile, "TITLE_STYLES", []) or []),
        "table_caption": list(getattr(profile, "TABLE_CAPTION_STYLES", []) or []),
        "fig_caption":   list(getattr(profile, "FIG_CAPTION_STYLES", []) or []),
        "table_cell":   list(getattr(profile, "TABLE_CELL_STYLES", []) or []),
    }


def _profile_target_for_role(profile: StylesProfile, role: str) -> Optional[str]:
    """主样式 styleId (target) per role; fall back to first *_STYLES item."""
    mapping = {
        "body":          getattr(profile, "BODY_TARGET_STYLE_ID", None),
        "h1":            getattr(profile, "H1_TARGET_STYLE_ID", None),
        "h2":            getattr(profile, "H2_TARGET_STYLE_ID", None),
        "h3":            getattr(profile, "H3_TARGET_STYLE_ID", None),
        "title":         getattr(profile, "TITLE_TARGET_STYLE_ID", None),
        "table_caption": getattr(profile, "TABLE_CAPTION_TARGET_STYLE_ID", None),
        "table_cell":    getattr(profile, "TABLE_CELL_STYLE_ID", None),
    }
    sid = mapping.get(role)
    if sid:
        return sid
    cands = _profile_role_styles(profile).get(role, [])
    return cands[0] if cands else None


# ═════════════════════════════════════════════════════════════════════════════
# CORE — shape_contract gate (跑前后对账)
# ═════════════════════════════════════════════════════════════════════════════
def _shape_gate(src: Path, work_fn, args,
                allowed_deltas: dict | None = None,
                label: str = "fix") -> tuple[dict, dict, list[str], bool]:
    """通用 fix gate:
        1. capture before
        2. work_fn(doc) -> 改 doc (in-memory), 返回 stats dict
        3. (if --inplace and not dry-run) save -> reopen -> capture after
           else: 复用 doc 算 after (不存盘)
        4. diff -> violations
        5. violations and not --force: refuse, return BEFORE without save
        6. violations and --force: warn, return AFTER (already saved if --inplace)

    Returns: (before_snap, after_snap_or_estimate, violations, refused)
    """
    before = capture_structure(src)

    # in-memory pass: open doc, run work_fn
    doc = Document(str(src))
    stats = work_fn(doc)

    if not _will_write(args):
        # estimate after = re-snapshot in-memory state via temp save
        # (轻量做法: 写到 tmp, capture, 不动 src)
        tmp = src.parent / f".{src.stem}._shape_tmp_{_dt.datetime.now().strftime('%H%M%S%f')}{src.suffix}"
        try:
            doc.save(str(tmp))
            after = capture_structure(tmp)
        finally:
            try:
                tmp.unlink()
            except FileNotFoundError:
                pass
        violations = diff_structure(before, after, allowed_deltas)
        return before, after, violations, False

    # --inplace: transactional — write to tmp, capture, only swap to src if shape OK
    # (backup is for user reference; rollback uses tmp staging not backup)
    bak = _pick_backup_path(src) if not getattr(args, "no_backup", False) else None
    if bak:
        shutil.copy2(src, bak)
    tmp_out = src.parent / f".{src.stem}._shape_stage_{_dt.datetime.now().strftime('%H%M%S%f')}{src.suffix}"
    try:
        doc.save(str(tmp_out))
        after = capture_structure(tmp_out)
        violations = diff_structure(before, after, allowed_deltas)

        if violations and not getattr(args, "force", False):
            # REFUSE — tmp_out discarded, src untouched
            try:
                tmp_out.unlink()
            except FileNotFoundError:
                pass
            return before, after, violations, True
        if violations and getattr(args, "force", False):
            _force_warning(f"{label} drift accepted")

        # commit: replace src atomically
        shutil.move(str(tmp_out), str(src))
        setattr(args, "_actual_backup", str(bak) if bak else None)
        return before, after, violations, False
    except Exception:
        if tmp_out.exists():
            try:
                tmp_out.unlink()
            except FileNotFoundError:
                pass
        raise


# ═════════════════════════════════════════════════════════════════════════════
# subcommand 1: style-rebrand
# ═════════════════════════════════════════════════════════════════════════════
def _set_para_pStyle(p, style_id: str):
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p._p.insert(0, pPr)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style_id)


def cmd_style_rebrand(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)

    # 决定 from / to
    from_match = args.from_style
    to_id = args.to_style
    if args.role and not to_id:
        to_id = _profile_target_for_role(profile, args.role)
        if not to_id:
            print(f"[ERR] profile {args.profile!r} 无 role={args.role!r} 主样式", file=sys.stderr)
            return 2
    if not from_match:
        from_match = "Normal"
    if not to_id:
        print("[ERR] --to / --role 必须给一个", file=sys.stderr)
        return 2

    stats_holder = {"matched": 0, "changed": 0, "skipped_same": 0}

    def _work(doc):
        # 验目标 styleId 在 docx 里存在
        available = {s.style_id for s in doc.styles}
        if to_id not in available:
            stats_holder["target_missing"] = to_id
        for p in doc.paragraphs:
            name = (p.style.name if p.style else None) or ""
            sid = getattr(p.style, "style_id", None) if p.style else None
            if from_match in (name, sid):
                stats_holder["matched"] += 1
                if sid == to_id:
                    stats_holder["skipped_same"] += 1
                    continue
                if to_id in available:
                    _set_para_pStyle(p, to_id)
                    stats_holder["changed"] += 1
        return stats_holder

    before, after, violations, refused = _shape_gate(
        src, _work, args,
        allowed_deltas={  # 段数 / 表 / heading / caption 全 0
            "paragraph_count": 0,
            "table_count": 0,
            "heading_counts": 0,
            "caption_figure_count": 0,
            "caption_table_count": 0,
        },
        label="style-rebrand",
    )

    if refused:
        sys.stderr.write(_refuse_msg(violations))

    report = {
        "subcommand": "fix style-rebrand",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "from": from_match,
        "to": to_id,
        "role": args.role,
        "dry_run": getattr(args, "dry_run", False),
        "inplace": getattr(args, "inplace", False),
        "matched": stats_holder.get("matched", 0),
        "changed": stats_holder.get("changed", 0),
        "skipped_same": stats_holder.get("skipped_same", 0),
        "target_missing_in_docx": stats_holder.get("target_missing"),
        "shape_violations": violations,
        "refused": refused,
        "backup": getattr(args, "_actual_backup", None),
    }
    _emit_report(report, args)
    print(f"[fix style-rebrand] from={from_match!r} to={to_id!r} "
          f"matched={report['matched']} changed={report['changed']} "
          f"skipped_same={report['skipped_same']} "
          f"violations={len(violations)} refused={refused}")
    return 3 if refused else 0


# ═════════════════════════════════════════════════════════════════════════════
# subcommand 2: style-pool-cleanup
# ═════════════════════════════════════════════════════════════════════════════
def _is_system_default_style(s) -> bool:
    """W:style 含 w:default="1" 或 styleId 是 Word built-in 常用."""
    try:
        el = s.element
    except Exception:
        return False
    if el is None:
        return False
    if el.get(qn("w:default")) == "1":
        return True
    sid = (getattr(s, "style_id", "") or "").lower()
    name = (getattr(s, "name", "") or "").lower()
    builtins = {
        "normal", "defaultparagraphfont", "tablenormal", "nolist",
        "title", "subtitle", "heading 1", "heading 2", "heading 3",
        "heading 4", "heading 5", "heading 6", "heading 7", "heading 8",
        "heading 9", "caption", "header", "footer", "footnotetext",
        "footnotereference", "hyperlink", "endnotetext",
    }
    return sid in builtins or name in builtins


def cmd_style_pool_cleanup(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))

    used_para = _all_para_style_ids(doc)
    used_run = _all_run_style_ids(doc)
    used_tbl = _all_table_style_ids(doc)
    used = used_para | used_run | used_tbl

    # also: any style referenced in numbering.xml / settings.xml
    aux_referenced: set[str] = set()
    for part in ("word/numbering.xml", "word/settings.xml", "word/footnotes.xml",
                 "word/endnotes.xml", "word/header1.xml", "word/footer1.xml"):
        for sid, _, _ in _styles_defined(doc):
            if sid in aux_referenced or sid in used:
                continue
            if _style_id_referenced_in_xml_part(src, part, sid):
                aux_referenced.add(sid)

    # also: styles referenced via basedOn / next / link from another style
    style_link_refs: set[str] = set()
    for s in doc.styles:
        el = getattr(s, "element", None)
        if el is None:
            continue
        for tag in ("w:basedOn", "w:next", "w:link"):
            for n in el.findall(qn(tag)):
                v = n.get(qn("w:val"))
                if v:
                    style_link_refs.add(v)

    deletable = []
    for s in doc.styles:
        sid = getattr(s, "style_id", "") or ""
        name = getattr(s, "name", "") or ""
        if not sid:
            continue
        if sid in used or sid in aux_referenced or sid in style_link_refs:
            continue
        if _is_system_default_style(s):
            continue
        deletable.append((sid, name))

    # dry-run reports only
    if not args.inplace:
        report = {
            "subcommand": "fix style-pool-cleanup",
            "docx": str(src),
            "profile": getattr(profile, "_name", args.profile),
            "dry_run": True,
            "total_defined": len(_styles_defined(doc)),
            "used_para": len(used_para),
            "used_run": len(used_run),
            "used_tbl": len(used_tbl),
            "aux_referenced": sorted(aux_referenced),
            "style_link_refs": len(style_link_refs),
            "deletable_count": len(deletable),
            "deletable_sample": [{"id": s, "name": n} for s, n in deletable[:20]],
        }
        _emit_report(report, args)
        print(f"[fix style-pool-cleanup] DRY-RUN total={report['total_defined']} "
              f"used_para={len(used_para)} deletable={len(deletable)}")
        return 0

    # --inplace: capture before, delete style elements, save, capture after
    def _work(doc2):
        # re-resolve in this doc2 instance for deletion
        used2 = _all_para_style_ids(doc2) | _all_run_style_ids(doc2) | _all_table_style_ids(doc2)
        link_refs2: set[str] = set()
        for s2 in doc2.styles:
            el = getattr(s2, "element", None)
            if el is None:
                continue
            for tag in ("w:basedOn", "w:next", "w:link"):
                for n in el.findall(qn(tag)):
                    v = n.get(qn("w:val"))
                    if v:
                        link_refs2.add(v)
        deleted_ids = []
        styles_el = doc2.styles.element
        for s2 in list(doc2.styles):
            sid = getattr(s2, "style_id", "") or ""
            if not sid:
                continue
            if sid in used2 or sid in link_refs2 or sid in aux_referenced:
                continue
            if _is_system_default_style(s2):
                continue
            el = getattr(s2, "element", None)
            if el is not None and el.getparent() is not None:
                el.getparent().remove(el)
                deleted_ids.append(sid)
        return {"deleted_count": len(deleted_ids), "deleted_ids": deleted_ids}

    before, after, violations, refused = _shape_gate(
        src, _work, args,
        allowed_deltas={
            "paragraph_count": 0,
            "table_count": 0,
            "heading_counts": 0,
            "caption_figure_count": 0,
            "caption_table_count": 0,
            "drawings_count": 0,
        },
        label="style-pool-cleanup",
    )

    if refused:
        sys.stderr.write(_refuse_msg(violations))

    report = {
        "subcommand": "fix style-pool-cleanup",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": False,
        "inplace": True,
        "deletable_planned": len(deletable),
        "shape_violations": violations,
        "refused": refused,
        "backup": getattr(args, "_actual_backup", None),
    }
    _emit_report(report, args)
    print(f"[fix style-pool-cleanup] deletable={len(deletable)} "
          f"violations={len(violations)} refused={refused}")
    return 3 if refused else 0


# ═════════════════════════════════════════════════════════════════════════════
# subcommand 3: style-pane-filter
# ═════════════════════════════════════════════════════════════════════════════
def cmd_style_pane_filter(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)

    role_styles = _profile_role_styles(profile)
    # 白名单 = profile 所有 *_STYLES 集合 + 目标 styleId
    whitelist: set[str] = set()
    for v in role_styles.values():
        whitelist.update(v)
    for f in ("BODY_TARGET_STYLE_ID", "H1_TARGET_STYLE_ID", "H2_TARGET_STYLE_ID",
              "H3_TARGET_STYLE_ID", "TITLE_TARGET_STYLE_ID",
              "TABLE_CAPTION_TARGET_STYLE_ID", "TABLE_CELL_STYLE_ID"):
        v = getattr(profile, f, None)
        if v:
            whitelist.add(v)

    if args.dry_run or not args.inplace:
        report = {
            "subcommand": "fix style-pane-filter",
            "docx": str(src),
            "profile": getattr(profile, "_name", args.profile),
            "dry_run": True,
            "whitelist_count": len(whitelist),
            "whitelist_sample": sorted(whitelist)[:20],
            "note": "dry-run: would set stylePaneFormatFilter visible mask + customStyles whitelist",
        }
        _emit_report(report, args)
        print(f"[fix style-pane-filter] DRY-RUN whitelist={len(whitelist)}")
        return 0

    def _work(doc2):
        # stylePaneFormatFilter 是 settings.xml 一项: <w:stylePaneFormatFilter w:val="..."
        # /w:allStyles="0" w:customStyles="1" w:latentStyles="0" w:stylesInUse="1" .../>
        settings_el = doc2.settings.element
        # remove existing
        for n in settings_el.findall(qn("w:stylePaneFormatFilter")):
            settings_el.remove(n)
        n = OxmlElement("w:stylePaneFormatFilter")
        # bit mask "5824" 常见 = visible heading + numbered (Word 默认); 这里给 "1F08"
        # 让用户只看到 in-use + custom (latent 隐藏)
        n.set(qn("w:val"), "1F08")
        n.set(qn("w:allStyles"), "0")
        n.set(qn("w:customStyles"), "1")
        n.set(qn("w:latentStyles"), "0")
        n.set(qn("w:stylesInUse"), "1")
        n.set(qn("w:headingStyles"), "1")
        n.set(qn("w:numberingStyles"), "0")
        n.set(qn("w:tableStyles"), "0")
        n.set(qn("w:directFormattingOnRuns"), "0")
        n.set(qn("w:directFormattingOnParagraphs"), "0")
        n.set(qn("w:directFormattingOnNumbering"), "0")
        n.set(qn("w:directFormattingOnTables"), "0")
        n.set(qn("w:clearFormatting"), "1")
        n.set(qn("w:top3HeadingStyles"), "1")
        n.set(qn("w:visibleStyles"), "0")
        n.set(qn("w:alternateStyleNames"), "0")
        settings_el.append(n)
        return {"set": True}

    before, after, violations, refused = _shape_gate(
        src, _work, args,
        allowed_deltas={  # settings 改不影响 body — 全 0
            "paragraph_count": 0,
            "table_count": 0,
            "section_count": 0,
            "heading_counts": 0,
            "caption_figure_count": 0,
            "caption_table_count": 0,
            "drawings_count": 0,
        },
        label="style-pane-filter",
    )
    if refused:
        sys.stderr.write(_refuse_msg(violations))

    report = {
        "subcommand": "fix style-pane-filter",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": False,
        "inplace": True,
        "whitelist_count": len(whitelist),
        "shape_violations": violations,
        "refused": refused,
        "backup": getattr(args, "_actual_backup", None),
    }
    _emit_report(report, args)
    print(f"[fix style-pane-filter] whitelist={len(whitelist)} "
          f"violations={len(violations)} refused={refused}")
    return 3 if refused else 0


# ═════════════════════════════════════════════════════════════════════════════
# subcommand 4: role-fill
# ═════════════════════════════════════════════════════════════════════════════
def _clone_style_as(doc, src_id: str, new_id: str, new_name: str) -> bool:
    """Clone existing style element as new styleId + name. Return True on success."""
    styles_el = doc.styles.element
    src_el = None
    for s in styles_el.findall(qn("w:style")):
        if s.get(qn("w:styleId")) == src_id:
            src_el = s
            break
    if src_el is None:
        return False
    new_el = etree.fromstring(etree.tostring(src_el))
    new_el.set(qn("w:styleId"), new_id)
    # change <w:name w:val="..."/>
    nm = new_el.find(qn("w:name"))
    if nm is None:
        nm = OxmlElement("w:name")
        new_el.insert(0, nm)
    nm.set(qn("w:val"), new_name)
    styles_el.append(new_el)
    return True


def cmd_role_fill(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))

    roles_needed = _profile_role_styles(profile)
    defined_ids = {sid for sid, _, _ in _styles_defined(doc)}
    defined_names = {n for _, n, _ in _styles_defined(doc)}

    # 哪些角色没命中 docx 任何 style
    missing: dict[str, list[str]] = {}
    for role, cands in roles_needed.items():
        if not cands:
            continue
        if any((c in defined_ids or c in defined_names) for c in cands):
            continue
        missing[role] = cands

    # 模板源 (用 Normal 克隆作基线; body role 例外 — 克隆 Body Text)
    template_id_for_role = {
        "body":          "Normal",
        "h1":            "Heading 1",
        "h2":            "Heading 2",
        "h3":            "Heading 3",
        "h4":            "Heading 4",
        "title":         "Title",
        "table_caption": "Caption",
        "fig_caption":   "Caption",
        "table_cell":    "Normal",
    }

    if args.dry_run or not args.inplace:
        report = {
            "subcommand": "fix role-fill",
            "docx": str(src),
            "profile": getattr(profile, "_name", args.profile),
            "dry_run": True,
            "roles_in_profile": list(roles_needed.keys()),
            "missing_roles": list(missing.keys()),
            "missing_detail": {r: cands for r, cands in missing.items()},
            "clone_plan": [
                {"role": r, "new_id": cands[0], "new_name": cands[0],
                 "from_template": template_id_for_role.get(r, "Normal")}
                for r, cands in missing.items()
            ],
        }
        _emit_report(report, args)
        print(f"[fix role-fill] DRY-RUN missing_roles={list(missing.keys())} "
              f"plan_count={len(missing)}")
        return 0

    def _work(doc2):
        created = []
        skipped = []
        for role, cands in missing.items():
            new_id = cands[0]
            new_name = cands[0]
            src_tmpl = template_id_for_role.get(role, "Normal")
            # find src template by id or by name
            real_src = None
            for s in doc2.styles:
                if getattr(s, "style_id", "") == src_tmpl or getattr(s, "name", "") == src_tmpl:
                    real_src = getattr(s, "style_id", None)
                    break
            if not real_src:
                skipped.append({"role": role, "reason": f"no template {src_tmpl!r}"})
                continue
            ok = _clone_style_as(doc2, real_src, new_id, new_name)
            if ok:
                created.append({"role": role, "new_id": new_id, "new_name": new_name,
                                "from": real_src})
            else:
                skipped.append({"role": role, "reason": "clone failed"})
        return {"created": created, "skipped": skipped}

    before, after, violations, refused = _shape_gate(
        src, _work, args,
        allowed_deltas={
            "paragraph_count": 0,
            "table_count": 0,
            "heading_counts": 0,
            "caption_figure_count": 0,
            "caption_table_count": 0,
            "drawings_count": 0,
        },
        label="role-fill",
    )
    if refused:
        sys.stderr.write(_refuse_msg(violations))

    report = {
        "subcommand": "fix role-fill",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": False,
        "inplace": True,
        "missing_roles_before": list(missing.keys()),
        "shape_violations": violations,
        "refused": refused,
        "backup": getattr(args, "_actual_backup", None),
    }
    _emit_report(report, args)
    print(f"[fix role-fill] missing={list(missing.keys())} "
          f"violations={len(violations)} refused={refused}")
    return 3 if refused else 0


# ═════════════════════════════════════════════════════════════════════════════
# subcommand 5: style-rename (W15)
#   改样式 .name 字段, 不动 styleId, 段引用 (pStyle.val=styleId) 不需变 → 0 段修改
# ═════════════════════════════════════════════════════════════════════════════
def cmd_style_rename(args) -> int:
    src = _common_setup(args)
    from_name = args.from_style
    to_name = args.to_style
    if not from_name or not to_name:
        print("[ERR] --from / --to 都必须给 (style .name 值)", file=sys.stderr)
        return 2

    stats_holder = {"matched_styles": [], "renamed": 0}

    def _work(doc):
        styles_el = doc.styles.element
        for st in styles_el.findall(qn("w:style")):
            nm = st.find(qn("w:name"))
            if nm is None:
                continue
            cur = nm.get(qn("w:val"))
            if cur == from_name:
                sid = st.get(qn("w:styleId")) or ""
                stats_holder["matched_styles"].append(
                    {"styleId": sid, "from": cur, "to": to_name}
                )
                nm.set(qn("w:val"), to_name)
                stats_holder["renamed"] += 1
        return stats_holder

    before, after, violations, refused = _shape_gate(
        src, _work, args,
        allowed_deltas={  # 改 styles.xml name 不影响 body 结构
            "paragraph_count": 0,
            "table_count": 0,
            "section_count": 0,
            "heading_counts": 0,
            "caption_figure_count": 0,
            "caption_table_count": 0,
            "drawings_count": 0,
        },
        label="style-rename",
    )

    if refused:
        sys.stderr.write(_refuse_msg(violations))

    report = {
        "subcommand": "fix style-rename",
        "docx": str(src),
        "from": from_name,
        "to": to_name,
        "dry_run": getattr(args, "dry_run", False),
        "inplace": getattr(args, "inplace", False),
        "renamed": stats_holder["renamed"],
        "matched_styles": stats_holder["matched_styles"],
        "shape_violations": violations,
        "refused": refused,
        "backup": getattr(args, "_actual_backup", None),
    }
    _emit_report(report, args)
    print(f"[fix style-rename] from={from_name!r} to={to_name!r} "
          f"renamed={stats_holder['renamed']} "
          f"matched_ids={[m['styleId'] for m in stats_holder['matched_styles']]} "
          f"violations={len(violations)} refused={refused}")
    return 3 if refused else 0


# ═════════════════════════════════════════════════════════════════════════════
# subcommand 6: clear-direct-format (W15)
#   清段 inline 直接格式 — pPr 只保留 pStyle; rPr 只保留 rStyle
# ═════════════════════════════════════════════════════════════════════════════
def cmd_clear_direct_format(args) -> int:
    src = _common_setup(args)
    style_filter = getattr(args, "style", None)  # None = 清所有段

    stats_holder = {
        "paragraphs_scanned": 0,
        "paragraphs_matched": 0,
        "pPr_children_removed": 0,
        "rPr_children_removed": 0,
        "runs_scanned": 0,
    }

    def _work(doc):
        # 收集 styleId 与 style.name 双向映射: filter 可以是 style.name
        name_to_id: dict[str, str] = {}
        for s in doc.styles:
            sid = getattr(s, "style_id", "") or ""
            nm = getattr(s, "name", "") or ""
            if sid and nm:
                name_to_id[nm] = sid
        target_id: Optional[str] = None
        if style_filter:
            if style_filter in name_to_id:
                target_id = name_to_id[style_filter]
            else:
                # also accept styleId直接给
                if style_filter in {v for v in name_to_id.values()}:
                    target_id = style_filter
                else:
                    target_id = style_filter  # 让后续判断兜底

        def _para_styleid(p_el) -> Optional[str]:
            pPr = p_el.find(qn("w:pPr"))
            if pPr is None:
                return None
            ps = pPr.find(qn("w:pStyle"))
            if ps is None:
                return None
            return ps.get(qn("w:val"))

        def _process_paragraph(p_el):
            stats_holder["paragraphs_scanned"] += 1
            if target_id is not None:
                sid = _para_styleid(p_el)
                if sid != target_id:
                    return
            stats_holder["paragraphs_matched"] += 1
            # pPr: keep pStyle (style ref) + sectPr (section structural, not formatting)
            # sectPr 是节边界标记, 删它会让 section_count 漂移; 严格说不是"直接格式"
            pPr = p_el.find(qn("w:pPr"))
            if pPr is not None:
                _PPR_KEEP = {"pStyle", "sectPr"}
                for child in list(pPr):
                    tag = etree.QName(child.tag).localname
                    if tag not in _PPR_KEEP:
                        pPr.remove(child)
                        stats_holder["pPr_children_removed"] += 1
            # runs: rPr keep only rStyle
            for r in p_el.findall(qn("w:r")):
                stats_holder["runs_scanned"] += 1
                rPr = r.find(qn("w:rPr"))
                if rPr is None:
                    continue
                for child in list(rPr):
                    tag = etree.QName(child.tag).localname
                    if tag != "rStyle":
                        rPr.remove(child)
                        stats_holder["rPr_children_removed"] += 1

        body = doc.element.body
        # 顶层 paragraphs
        for p_el in body.findall(".//" + qn("w:p")):
            _process_paragraph(p_el)
        return stats_holder

    before, after, violations, refused = _shape_gate(
        src, _work, args,
        allowed_deltas={  # 清 inline 直接格式不动结构
            "paragraph_count": 0,
            "table_count": 0,
            "section_count": 0,
            "heading_counts": 0,
            "caption_figure_count": 0,
            "caption_table_count": 0,
            "drawings_count": 0,
        },
        label="clear-direct-format",
    )

    if refused:
        sys.stderr.write(_refuse_msg(violations))

    report = {
        "subcommand": "fix clear-direct-format",
        "docx": str(src),
        "style_filter": style_filter,
        "dry_run": getattr(args, "dry_run", False),
        "inplace": getattr(args, "inplace", False),
        **stats_holder,
        "shape_violations": violations,
        "refused": refused,
        "backup": getattr(args, "_actual_backup", None),
    }
    _emit_report(report, args)
    print(f"[fix clear-direct-format] style_filter={style_filter!r} "
          f"scanned={stats_holder['paragraphs_scanned']} "
          f"matched={stats_holder['paragraphs_matched']} "
          f"pPr_removed={stats_holder['pPr_children_removed']} "
          f"rPr_removed={stats_holder['rPr_children_removed']} "
          f"violations={len(violations)} refused={refused}")
    return 3 if refused else 0


# ═════════════════════════════════════════════════════════════════════════════
# subcommand 7: style-create (W15)
#   按 base style 克隆新空 style 定义到 styles.xml (改 styleId + name, 段不受影响)
# ═════════════════════════════════════════════════════════════════════════════
def cmd_style_create(args) -> int:
    src = _common_setup(args)
    base = args.base
    new_id = args.new_id
    new_name = args.new_name
    new_type = getattr(args, "type", "paragraph") or "paragraph"
    if not base or not new_id or not new_name:
        print("[ERR] --base / --new-id / --new-name 都必须给", file=sys.stderr)
        return 2

    stats_holder = {
        "base_found": False,
        "base_styleId": None,
        "created": False,
        "collision": False,
    }

    def _work(doc):
        styles_el = doc.styles.element
        # 找 base: 优先 styleId, 再 name
        base_el = None
        for st in styles_el.findall(qn("w:style")):
            if st.get(qn("w:styleId")) == base:
                base_el = st
                break
        if base_el is None:
            for st in styles_el.findall(qn("w:style")):
                nm = st.find(qn("w:name"))
                if nm is not None and nm.get(qn("w:val")) == base:
                    base_el = st
                    break
        if base_el is None:
            return stats_holder
        stats_holder["base_found"] = True
        stats_holder["base_styleId"] = base_el.get(qn("w:styleId"))

        # collision check: new_id already exists?
        for st in styles_el.findall(qn("w:style")):
            if st.get(qn("w:styleId")) == new_id:
                stats_holder["collision"] = True
                return stats_holder

        # deepcopy + 改 styleId + 改 name + 改 type
        new_el = etree.fromstring(etree.tostring(base_el))
        new_el.set(qn("w:styleId"), new_id)
        if new_type:
            new_el.set(qn("w:type"), new_type)
        nm = new_el.find(qn("w:name"))
        if nm is None:
            nm = OxmlElement("w:name")
            new_el.insert(0, nm)
        nm.set(qn("w:val"), new_name)
        # default="1" 不该跟着克隆 (避免和原 base 冲突 default 标记)
        if new_el.get(qn("w:default")) == "1":
            del new_el.attrib[qn("w:default")]
        styles_el.append(new_el)
        stats_holder["created"] = True
        return stats_holder

    before, after, violations, refused = _shape_gate(
        src, _work, args,
        allowed_deltas={  # 加一个空 style def 不动 body
            "paragraph_count": 0,
            "table_count": 0,
            "section_count": 0,
            "heading_counts": 0,
            "caption_figure_count": 0,
            "caption_table_count": 0,
            "drawings_count": 0,
        },
        label="style-create",
    )

    if refused:
        sys.stderr.write(_refuse_msg(violations))

    report = {
        "subcommand": "fix style-create",
        "docx": str(src),
        "base": base,
        "new_id": new_id,
        "new_name": new_name,
        "new_type": new_type,
        "dry_run": getattr(args, "dry_run", False),
        "inplace": getattr(args, "inplace", False),
        **stats_holder,
        "shape_violations": violations,
        "refused": refused,
        "backup": getattr(args, "_actual_backup", None),
    }
    _emit_report(report, args)
    print(f"[fix style-create] base={base!r} new_id={new_id!r} new_name={new_name!r} "
          f"base_found={stats_holder['base_found']} "
          f"collision={stats_holder['collision']} "
          f"created={stats_holder['created']} "
          f"violations={len(violations)} refused={refused}")
    if stats_holder["collision"]:
        return 2
    if not stats_holder["base_found"]:
        return 2
    return 3 if refused else 0


# ═════════════════════════════════════════════════════════════════════════════
# argparse register
# ═════════════════════════════════════════════════════════════════════════════
_SUBCMDS = {
    "style-rebrand":       (cmd_style_rebrand,
                            "批量段样式迁移 (Normal → 项目正文样式)"),
    "style-pool-cleanup":  (cmd_style_pool_cleanup,
                            "删 docx 内定义但未用的样式"),
    "style-pane-filter":   (cmd_style_pane_filter,
                            "设 stylePaneFormatFilter 白名单 profile.roles 样式"),
    "role-fill":           (cmd_role_fill,
                            "缺角色时按 profile 自动建样式 (从 Normal/Heading X 克隆)"),
    "style-rename":        (cmd_style_rename,
                            "改样式 .name 字段, 不动 styleId, 段引用 0 修改"),
    "clear-direct-format": (cmd_clear_direct_format,
                            "清段 inline 直接格式 (pPr/rPr 直接子元素), 保留 pStyle/rStyle"),
    "style-create":        (cmd_style_create,
                            "按 base style 克隆新空 style 定义到 styles.xml"),
}


def _add_common_args(p: argparse.ArgumentParser):
    p.add_argument("docx_path", type=Path, help="target docx path")
    p.add_argument("--dry-run", action="store_true",
                   help="只列计划不写盘 (default when --inplace not given)")
    p.add_argument("--inplace", action="store_true",
                   help="真改 (自动留 .bak-N-DATE); 不给 --inplace 默认 dry-run")
    p.add_argument("--no-backup", action="store_true", help="跳过备份 (慎用)")
    p.add_argument("--force", action="store_true",
                   help="旁路 shape_contract (打 4 行 WARNING; 不推荐)")
    p.add_argument("--report", type=Path, default=None, help="JSON report path")
    p.add_argument("--profile", type=str, default=None,
                   help="styles_registry profile (zdwp / eco-flow / ...)")


def register(subparsers) -> None:
    """Register `fix <subcmd>` group on doctools CLI subparsers."""
    from ._dispatch import get_or_add_group, get_or_add_subparsers

    fix_p = get_or_add_group(subparsers, "fix",
                              "style-set fix family (rebrand / pool-cleanup / pane-filter / role-fill)")
    fix_sub = get_or_add_subparsers(fix_p, dest="fix_target")
    existing = getattr(fix_sub, "choices", {}) or {}

    for name, (fn, helptxt) in _SUBCMDS.items():
        if name in existing:
            continue
        sp = fix_sub.add_parser(name, help=helptxt)
        _add_common_args(sp)
        if name == "style-rebrand":
            sp.add_argument("--from", dest="from_style", default=None,
                            help="match by style name OR styleId (default 'Normal')")
            sp.add_argument("--to", dest="to_style", default=None,
                            help="target styleId")
            sp.add_argument("--role", default=None,
                            help="alt to --to: 'body'/'h1'/.../ take profile.roles.<role> target")
        elif name == "style-rename":
            sp.add_argument("--from", dest="from_style", required=True,
                            help="样式 .name 当前值 (e.g. '0 图名称')")
            sp.add_argument("--to", dest="to_style", required=True,
                            help="样式 .name 新值 (e.g. 'ZDWP图名')")
        elif name == "clear-direct-format":
            sp.add_argument("--style", dest="style", default=None,
                            help="只清匹配该 style .name (或 styleId) 的段; 不给 = 清所有段")
        elif name == "style-create":
            sp.add_argument("--base", required=True,
                            help="基样式 styleId 或 .name (e.g. 'Normal')")
            sp.add_argument("--new-id", dest="new_id", required=True,
                            help="新 styleId")
            sp.add_argument("--new-name", dest="new_name", required=True,
                            help="新 style .name")
            sp.add_argument("--type", dest="type", default="paragraph",
                            choices=["paragraph", "character", "table", "numbering"],
                            help="新 style 类型 (default paragraph)")
        sp.set_defaults(func=fn)


# ═════════════════════════════════════════════════════════════════════════════
# cmd_restore (W-restore 2026-05-26) — 9-step 综合 styleset 修复链
#
# eco-flow 项目专属配方; profile yaml = profiles/eco_flow_health.yaml
# (与 fix_styleset 其它命令用的 styles_registry profile 不是同一个 SSOT)
#
# 9 step (每步过 shape_contract via 各 cmd_xxx 内置 gate):
#   0  capture before (audit)
#   1  rename × 3   : 0 图名称→ZDWP图名 / 0 表格标题→ZDWP 表名 / 0表格内容→ZDWP表格内容
#   2  rebrand × 2  : Normal→ZDWP正文 / 01正文→ZDWP正文  (走 BODY_TARGET)
#   3  create × 5   : ZDWP附表 / zdwp题目0 / zdwp题目1 / zdwp作者 / zdwp封面日期
#   4  firstLine    : 给 ZDWP正文 样式定义加 firstLineChars=200 + firstLine=480
#   5  clear-direct : --style "ZDWP正文" 清 inline pPr/rPr
#   6  cover-assign : LLM 识别封面 4 idx → 设 pStyle
#   7  pool-cleanup : 留 KEEP set
#   8  final shape_contract verify (vs step-0 before)
# ═════════════════════════════════════════════════════════════════════════════

# eco-flow 配方常量
_ECO_FLOW_PROFILE_YAML = (
    Path(__file__).resolve().parent.parent / "profiles" / "eco_flow_health.yaml"
)

_RECIPE_RENAME = [
    # (style.name FROM, style.name TO)
    ("0 图名称", "ZDWP图名"),
    ("0 表格标题", "ZDWP 表名"),
    ("0表格内容", "ZDWP表格内容"),
]
_RECIPE_REBRAND_FROM = ["Normal", "01正文"]
_RECIPE_REBRAND_TO_ID = "ZDWP正文"

_RECIPE_CREATE = [
    # (base styleId/name, new_id, new_name, type)
    ("ZDWP正文", "ZDWP附表", "ZDWP附表", "paragraph"),
    ("Title",    "zdwp题目0", "zdwp题目0", "paragraph"),
    ("Title",    "zdwp题目1", "zdwp题目1", "paragraph"),
    ("ZDWP正文", "zdwp作者",  "zdwp作者",  "paragraph"),
    ("ZDWP正文", "zdwp封面日期", "zdwp封面日期", "paragraph"),
]

# 兜底 KEEP (除 profile.roles 派生的样式之外, 防误删 Word built-in)
_BUILTIN_KEEP = (
    *(f"Heading {i}" for i in range(1, 10)),
    "Normal", "Default Paragraph Font", "Title", "Subtitle",
    "Caption", "Header", "Footer", "FootnoteText", "FootnoteReference",
    "EndnoteText", "Hyperlink", "TOC 1", "TOC 2", "TOC 3", "TOC 4",
    "TOC 5", "No List", "Table Normal", "页眉", "页脚", "超链接",
)


def _load_yaml_profile(yaml_path: Path) -> dict:
    """Load eco_flow_health.yaml-style profile (separate from styles_registry)."""
    import yaml as _yaml
    p = Path(os.path.expanduser(str(yaml_path))).resolve() if False else Path(yaml_path).expanduser().resolve()
    with open(p, "r", encoding="utf-8") as f:
        return _yaml.safe_load(f) or {}


def _derive_keep_set(yaml_profile: dict) -> set[str]:
    """KEEP set = all role styles (派生自 profile.roles) + builtins.

    Args:
        yaml_profile: eco_flow_health.yaml loaded dict
    Returns:
        set of style names (matches by .name; styles.xml stores name + styleId)
    """
    keep: set[str] = set()
    for role_styles in (yaml_profile.get("roles") or {}).values():
        if role_styles:
            keep.update(role_styles)
    keep.update(yaml_profile.get("tolerated_styles") or [])
    keep.update(_BUILTIN_KEEP)
    return keep


def _add_first_line_indent_to_style(docx_path: Path, style_id: str,
                                    chars: int = 200, twips: int = 480) -> dict:
    """Open docx via zipfile + lxml, add w:ind firstLineChars + firstLine to
    the given styleId in word/styles.xml. Save back to docx_path.

    Returns: {"found": bool, "modified": bool}
    """
    stats = {"found": False, "modified": False}
    import zipfile as _zf
    import tempfile

    # read all parts → temp dir; modify styles.xml; rezip
    src = Path(docx_path)
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        with _zf.ZipFile(str(src), "r") as zin:
            zin.extractall(td_path)
        styles_xml = td_path / "word" / "styles.xml"
        if not styles_xml.exists():
            return stats
        tree = etree.parse(str(styles_xml))
        root = tree.getroot()
        ns = {"w": W_NS}
        # find style by styleId OR by name
        target_el = None
        for st in root.findall(qn("w:style")):
            sid = st.get(qn("w:styleId"))
            if sid == style_id:
                target_el = st
                break
            nm = st.find(qn("w:name"))
            if nm is not None and nm.get(qn("w:val")) == style_id:
                target_el = st
                break
        if target_el is None:
            return stats
        stats["found"] = True
        pPr = target_el.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            target_el.append(pPr)
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:firstLineChars"), str(chars))
        ind.set(qn("w:firstLine"), str(twips))
        stats["modified"] = True

        # write styles.xml back
        tree.write(str(styles_xml), xml_declaration=True,
                   encoding="UTF-8", standalone=True)

        # repackage zip
        tmp_out = src.with_suffix(src.suffix + ".firstLine.tmp")
        with _zf.ZipFile(str(tmp_out), "w", _zf.ZIP_DEFLATED) as zout:
            for f in td_path.rglob("*"):
                if f.is_file():
                    arc = f.relative_to(td_path).as_posix()
                    zout.write(str(f), arcname=arc)
        shutil.move(str(tmp_out), str(src))
    return stats


def _cover_assign_pstyle(docx_path: Path, role_to_styleid: dict) -> dict:
    """Open docx, for each non-None idx in role_to_styleid (using
    LLM-identified cover paragraph idxs), set pStyle to mapped styleId.

    role_to_styleid: e.g. {3: "zdwp题目0", 4: "zdwp题目1", 11: "zdwp作者", 20: "zdwp封面日期"}
    Returns: {"assigned": int, "skipped_missing_idx": int}
    """
    stats = {"assigned": 0, "skipped_missing_idx": 0, "skipped_missing_style": 0,
             "details": []}
    doc = Document(str(docx_path))
    available = {s.style_id for s in doc.styles}
    n_paras = len(doc.paragraphs)
    for idx, sid in role_to_styleid.items():
        if idx is None:
            stats["skipped_missing_idx"] += 1
            continue
        if idx < 0 or idx >= n_paras:
            stats["skipped_missing_idx"] += 1
            stats["details"].append({"idx": idx, "issue": "out of range", "n_paras": n_paras})
            continue
        if sid not in available:
            stats["skipped_missing_style"] += 1
            stats["details"].append({"idx": idx, "issue": f"styleId {sid} missing"})
            continue
        p = doc.paragraphs[idx]
        _set_para_pStyle(p, sid)
        stats["assigned"] += 1
        stats["details"].append({"idx": idx, "set_styleId": sid,
                                 "text_preview": (p.text or "")[:40]})
    doc.save(str(docx_path))
    return stats


def _pool_cleanup_with_keep(docx_path: Path, keep_names: set[str]) -> dict:
    """Delete styles whose .name is not in keep_names AND not actually used
    by any paragraph/run/table reference AND not a system default.

    Returns: {"deleted_count": N, "deleted": [...], "kept": N}
    """
    stats = {"deleted_count": 0, "deleted": [], "kept": 0, "skipped_in_use": 0}
    doc = Document(str(docx_path))

    used = _all_para_style_ids(doc) | _all_run_style_ids(doc) | _all_table_style_ids(doc)
    link_refs: set[str] = set()
    for s in doc.styles:
        el = getattr(s, "element", None)
        if el is None:
            continue
        for tag in ("w:basedOn", "w:next", "w:link"):
            for n in el.findall(qn(tag)):
                v = n.get(qn("w:val"))
                if v:
                    link_refs.add(v)

    for s in list(doc.styles):
        sid = getattr(s, "style_id", "") or ""
        name = getattr(s, "name", "") or ""
        if not sid:
            continue
        if name in keep_names or sid in keep_names:
            stats["kept"] += 1
            continue
        if sid in used or sid in link_refs:
            stats["skipped_in_use"] += 1
            stats["kept"] += 1
            continue
        if _is_system_default_style(s):
            stats["kept"] += 1
            continue
        el = getattr(s, "element", None)
        if el is not None and el.getparent() is not None:
            el.getparent().remove(el)
            stats["deleted_count"] += 1
            stats["deleted"].append({"id": sid, "name": name})

    doc.save(str(docx_path))
    return stats


def _resolve_styleid_by_name(docx_path: Path, name: str) -> Optional[str]:
    """Lookup style.name -> style_id in docx. Used by cover-assign (need styleId not .name)."""
    doc = Document(str(docx_path))
    for s in doc.styles:
        if getattr(s, "name", "") == name or getattr(s, "style_id", "") == name:
            return getattr(s, "style_id", None)
    return None


def _make_subcmd_args(docx_path: Path, **kw):
    """Build argparse.Namespace for invoking other cmd_xxx in-process."""
    ns = argparse.Namespace()
    ns.docx_path = docx_path
    ns.dry_run = False
    ns.inplace = True
    ns.no_backup = True  # we manage backup at restore-level
    ns.force = False
    ns.report = None
    ns.profile = kw.pop("profile", None)
    for k, v in kw.items():
        setattr(ns, k, v)
    return ns


def cmd_restore(args) -> int:
    """eco-flow styleset 9-step 综合修复.

    args attrs:
      docx_path: Path
      dry_run, inplace, no_backup, force, report
      output: Path | None   — 输出位置 (--inplace 时忽略)
      no_llm: bool          — 用启发式 fallback 替代 LLM
      yaml_profile: Path | None — eco_flow_health.yaml 路径
    """
    src = _common_setup(args)
    yaml_path = Path(getattr(args, "yaml_profile", None)
                     or _ECO_FLOW_PROFILE_YAML).expanduser().resolve()
    if not yaml_path.exists():
        print(f"[ERR] profile yaml not found: {yaml_path}", file=sys.stderr)
        return 2
    yaml_profile = _load_yaml_profile(yaml_path)
    keep_set = _derive_keep_set(yaml_profile)
    dry_run = getattr(args, "dry_run", False)
    inplace = getattr(args, "inplace", False)
    output: Optional[Path] = getattr(args, "output", None)
    no_llm: bool = getattr(args, "no_llm", False)

    # decide staging path
    ts = _dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    staging = src.parent / f".{src.stem}.restore-staging-{ts}{src.suffix}"

    step_log: list[dict] = []

    def _log(step_id: str, label: str, stats: dict, ok: bool = True, note: str = ""):
        step_log.append({"step": step_id, "label": label, "ok": ok,
                         "stats": stats, "note": note})

    # Step 0: capture before
    try:
        before = capture_structure(src)
    except Exception as e:
        print(f"[ERR] step0 capture failed: {e}", file=sys.stderr)
        return 1
    _log("0", "capture_before", {"paragraph_count": before.get("paragraph_count")})

    if dry_run:
        # dry-run: report plan only
        report = {
            "subcommand": "styleset restore",
            "mode": "dry-run",
            "docx": str(src),
            "yaml_profile": str(yaml_path),
            "no_llm": no_llm,
            "plan": {
                "step1_rename":  [{"from": f, "to": t} for f, t in _RECIPE_RENAME],
                "step2_rebrand": [{"from": f, "to": _RECIPE_REBRAND_TO_ID} for f in _RECIPE_REBRAND_FROM],
                "step3_create":  [{"base": b, "new_id": nid, "new_name": nname}
                                  for b, nid, nname, _ in _RECIPE_CREATE],
                "step4_firstLine": {"style_id": _RECIPE_REBRAND_TO_ID,
                                    "chars": 200, "firstLine_twips": 480},
                "step5_clear_direct": {"style_filter": _RECIPE_REBRAND_TO_ID},
                "step6_cover_assign": {"method": "heuristic" if no_llm else "llm"},
                "step7_pool_cleanup": {"keep_count": len(keep_set),
                                       "keep_sample": sorted(keep_set)[:15]},
                "step8_verify": "shape_contract diff vs before"
            },
        }
        _emit_report(report, args)
        print(f"[styleset restore] DRY-RUN docx={src.name} "
              f"plan_steps=8 keep_count={len(keep_set)} no_llm={no_llm}")
        return 0

    # real run: pick backup, copy src → staging
    bak: Optional[Path] = None
    if not getattr(args, "no_backup", False):
        bak = _pick_backup_path(src)
        shutil.copy2(src, bak)
    shutil.copy2(src, staging)

    try:
        # Step 1: rename × 3
        for from_name, to_name in _RECIPE_RENAME:
            ns = _make_subcmd_args(staging, from_style=from_name, to_style=to_name)
            rc = cmd_style_rename(ns)
            if rc != 0:
                raise RuntimeError(f"step1 rename {from_name!r} → {to_name!r} rc={rc}")
        _log("1", "rename×3", {"renames": [{"from": f, "to": t} for f, t in _RECIPE_RENAME]})

        # Step 2: rebrand × 2 (Normal → ZDWP正文, 01正文 → ZDWP正文)
        for from_match in _RECIPE_REBRAND_FROM:
            ns = _make_subcmd_args(staging, from_style=from_match,
                                   to_style=_RECIPE_REBRAND_TO_ID, role=None)
            rc = cmd_style_rebrand(ns)
            if rc != 0:
                raise RuntimeError(f"step2 rebrand {from_match!r} → {_RECIPE_REBRAND_TO_ID!r} rc={rc}")
        _log("2", "rebrand×2", {"from_list": _RECIPE_REBRAND_FROM, "to": _RECIPE_REBRAND_TO_ID})

        # Step 3: create × 5 (skip if styleId already exists — cmd_style_create returns 2 on collision)
        created, collided = [], []
        for base, new_id, new_name, ntype in _RECIPE_CREATE:
            ns = _make_subcmd_args(staging, base=base, new_id=new_id,
                                   new_name=new_name, type=ntype)
            rc = cmd_style_create(ns)
            if rc == 0:
                created.append(new_id)
            elif rc == 2:
                collided.append(new_id)  # OK: already exists, skip
            else:
                raise RuntimeError(f"step3 create {new_id!r} unexpected rc={rc}")
        _log("3", "create×5", {"created": created, "skipped_existing": collided})

        # Step 4: firstLineChars=200 + firstLine=480 on ZDWP正文 style
        s4 = _add_first_line_indent_to_style(staging, _RECIPE_REBRAND_TO_ID, 200, 480)
        if not s4.get("found"):
            raise RuntimeError(f"step4 firstLine: style {_RECIPE_REBRAND_TO_ID!r} not found")
        _log("4", "firstLine_indent", s4)

        # Step 5: clear-direct-format --style ZDWP正文
        ns = _make_subcmd_args(staging, style=_RECIPE_REBRAND_TO_ID)
        rc = cmd_clear_direct_format(ns)
        if rc != 0:
            raise RuntimeError(f"step5 clear-direct rc={rc}")
        _log("5", "clear_direct_format", {"style_filter": _RECIPE_REBRAND_TO_ID})

        # Step 6: cover-assign via LLM/heuristic
        from .cover_identifier import (
            identify_cover_roles, identify_cover_roles_heuristic
        )
        try:
            if no_llm:
                cover = identify_cover_roles_heuristic(staging)
                cover_method = "heuristic"
            else:
                try:
                    cover = identify_cover_roles(staging)
                    cover_method = "llm"
                except (ImportError, RuntimeError, ValueError) as e:
                    print(f"[WARN] LLM unavailable in step6, fallback: {e}", file=sys.stderr)
                    cover = identify_cover_roles_heuristic(staging)
                    cover_method = "heuristic-fallback"
        except Exception as e:
            raise RuntimeError(f"step6 cover-identify failed: {e}")

        role_to_styleid = {
            cover.get("primary_title_idx"): "zdwp题目0",
            cover.get("subtitle_idx"):      "zdwp题目1",
            cover.get("author_idx"):        "zdwp作者",
            cover.get("date_idx"):          "zdwp封面日期",
        }
        s6 = _cover_assign_pstyle(staging, role_to_styleid)
        _log("6", "cover_assign", {"method": cover_method, "cover_idx": cover, **s6})

        # Step 7: pool-cleanup with KEEP set
        s7 = _pool_cleanup_with_keep(staging, keep_set)
        _log("7", "pool_cleanup",
             {"deleted_count": s7["deleted_count"], "kept": s7["kept"],
              "deleted_sample": s7["deleted"][:10]})

        # Step 8: final shape_contract verify (full doc, vs before)
        after = capture_structure(staging)
        # default allowed_deltas: paragraph/table/heading/caption全0;
        # styles 池减是允许的 (不在 shape_contract 指标里 — 它只看 body 结构)
        violations = diff_structure(before, after, allowed_deltas={
            "paragraph_count": 0,
            "table_count": 0,
            "section_count": 0,
            "heading_counts": 0,
            "caption_figure_count": 0,
            "caption_table_count": 0,
            "drawings_count": 0,
        })
        if violations and not getattr(args, "force", False):
            _log("8", "shape_contract_verify", {"violations": violations}, ok=False)
            raise RuntimeError(f"step8 shape_contract failed: {len(violations)} violation(s)")
        if violations and getattr(args, "force", False):
            _force_warning("styleset restore drift accepted")
        _log("8", "shape_contract_verify",
             {"violations": violations, "before_para": before.get("paragraph_count"),
              "after_para": after.get("paragraph_count")})

        # commit: where to?
        if inplace:
            shutil.move(str(staging), str(src))
            final_path = src
        elif output:
            output = Path(output).expanduser()
            output.parent.mkdir(parents=True, exist_ok=True)
            shutil.move(str(staging), str(output))
            final_path = output
        else:
            # neither --inplace nor --output: default = .restored.docx beside src
            final_path = src.with_name(f"{src.stem}.restored{src.suffix}")
            shutil.move(str(staging), str(final_path))

        report = {
            "subcommand": "styleset restore",
            "mode": "applied",
            "docx_src": str(src),
            "docx_out": str(final_path),
            "backup": str(bak) if bak else None,
            "yaml_profile": str(yaml_path),
            "no_llm": no_llm,
            "keep_count": len(keep_set),
            "shape_violations": violations,
            "step_log": step_log,
        }
        _emit_report(report, args)
        print(f"[styleset restore] OK src={src.name} out={final_path.name} "
              f"steps=8/8 violations={len(violations)} bak={bak.name if bak else None}")
        return 0

    except Exception as e:
        # rollback: keep src untouched, drop staging
        if staging.exists():
            try:
                staging.unlink()
            except FileNotFoundError:
                pass
        print(f"[styleset restore] FAILED: {type(e).__name__}: {e}", file=sys.stderr)
        report = {
            "subcommand": "styleset restore",
            "mode": "rolled_back",
            "docx_src": str(src),
            "backup": str(bak) if bak else None,
            "error": f"{type(e).__name__}: {e}",
            "step_log": step_log,
        }
        _emit_report(report, args)
        return 1


# ─── ProcessPool batch wrapper (top-level for pickling) ────────────────
def _restore_worker(task: dict) -> dict:
    """Worker entry for ProcessPoolExecutor. task = {docx, output_dir, dry_run, ...}"""
    import time as _t
    t0 = _t.perf_counter()
    ns = argparse.Namespace()
    ns.docx_path = Path(task["docx"])
    ns.dry_run = task.get("dry_run", False)
    ns.inplace = task.get("inplace", False)
    ns.no_backup = task.get("no_backup", False)
    ns.force = task.get("force", False)
    ns.report = task.get("report")
    ns.no_llm = task.get("no_llm", False)
    ns.yaml_profile = task.get("yaml_profile")
    # decide output for this docx
    out_dir = task.get("output_dir")
    output_explicit = task.get("output")
    if out_dir and not ns.inplace and not ns.dry_run:
        out_dir_p = Path(out_dir).expanduser()
        out_dir_p.mkdir(parents=True, exist_ok=True)
        ns.output = out_dir_p / Path(task["docx"]).name
    elif output_explicit and not ns.inplace and not ns.dry_run:
        ns.output = Path(output_explicit).expanduser()
    else:
        ns.output = None
    try:
        rc = cmd_restore(ns)
        ok = (rc == 0)
        err = None
    except Exception as exc:
        rc = 1
        ok = False
        err = f"{type(exc).__name__}: {exc}"
    return {"docx": task["docx"], "ok": ok, "rc": rc, "error": err,
            "duration": _t.perf_counter() - t0}


def cmd_restore_batch(docx_paths: list[Path], output_dir: Optional[Path] = None,
                      output_single: Optional[Path] = None,
                      dry_run: bool = False, inplace: bool = False,
                      no_backup: bool = False, no_llm: bool = False,
                      force: bool = False, max_workers: Optional[int] = None,
                      yaml_profile: Optional[Path] = None) -> int:
    """Drive batch restore (N=1 serial, N≥2 ProcessPool)."""
    import time as _t
    n = len(docx_paths)
    if n == 0:
        print("[styleset restore] no docx given", file=sys.stderr)
        return 2

    if n == 1:
        # serial path: invoke cmd_restore directly
        ns = argparse.Namespace()
        ns.docx_path = docx_paths[0]
        ns.dry_run = dry_run
        ns.inplace = inplace
        ns.no_backup = no_backup
        ns.force = force
        ns.report = None
        ns.no_llm = no_llm
        ns.yaml_profile = yaml_profile
        if output_dir and not inplace and not dry_run:
            output_dir.mkdir(parents=True, exist_ok=True)
            ns.output = output_dir / docx_paths[0].name
        elif output_single and not inplace and not dry_run:
            ns.output = output_single
        else:
            ns.output = None
        return cmd_restore(ns)

    # parallel path
    from concurrent.futures import ProcessPoolExecutor, as_completed
    mw = max_workers if max_workers and max_workers > 0 else min(n, os.cpu_count() or 4)
    print(f"[styleset restore] BATCH n={n} workers={mw} no_llm={no_llm} dry_run={dry_run}")
    tasks = [
        {
            "docx": str(p),
            "output_dir": str(output_dir) if output_dir else None,
            "output": str(output_single) if (output_single and n == 1) else None,
            "dry_run": dry_run,
            "inplace": inplace,
            "no_backup": no_backup,
            "no_llm": no_llm,
            "force": force,
            "yaml_profile": str(yaml_profile) if yaml_profile else None,
        }
        for p in docx_paths
    ]
    t0 = _t.perf_counter()
    results: list[dict] = []
    with ProcessPoolExecutor(max_workers=mw) as ex:
        futs = {ex.submit(_restore_worker, t): t for t in tasks}
        for f in as_completed(futs):
            try:
                r = f.result()
            except Exception as exc:
                t = futs[f]
                r = {"docx": t["docx"], "ok": False, "rc": 1,
                     "error": f"{type(exc).__name__}: {exc}", "duration": 0}
            results.append(r)
            tag = "OK" if r["ok"] else "FAIL"
            print(f"  [{tag}] {Path(r['docx']).name}  {r['duration']:.2f}s "
                  + (f"err={r['error']}" if r.get("error") else ""))
    wall = _t.perf_counter() - t0
    ok = sum(1 for r in results if r["ok"])
    ser_sum = sum(r.get("duration", 0) for r in results)
    print(f"[styleset restore] BATCH done {ok}/{n} ok · wall={wall:.2f}s "
          f"serial_sum={ser_sum:.2f}s speedup≈{ser_sum/wall:.2f}x")
    return 0 if ok == n else 1


# ─── Argparse registration for `audit-styleset restore` / `styleset restore` ──
def _add_restore_args(p: argparse.ArgumentParser):
    p.add_argument("docx_paths", nargs="+", type=Path,
                   help="target docx path(s). N≥2 → ProcessPool auto-parallel")
    p.add_argument("--dry-run", action="store_true", help="不写盘, 仅报告计划")
    p.add_argument("--inplace", action="store_true",
                   help="原地修改 (自动备份); 与 -o 互斥")
    p.add_argument("--no-backup", action="store_true", help="跳过备份 (慎用)")
    p.add_argument("--force", action="store_true",
                   help="旁路 final shape_contract")
    p.add_argument("--no-llm", action="store_true",
                   help="用启发式 fallback 替代 LLM 封面识别 (CC agent 内推荐)")
    p.add_argument("--yaml-profile", dest="yaml_profile", type=Path, default=None,
                   help=f"eco_flow_health yaml (default: {_ECO_FLOW_PROFILE_YAML.name})")
    p.add_argument("--max-workers", dest="max_workers", type=int, default=None,
                   help="ProcessPool 并发数; N≥2 时生效, 默认 = min(N, cpu_count)")
    p.add_argument("-o", "--output", type=Path, default=None,
                   help="单文件输出路径 (N=1) 或多文件输出目录 (N≥2)")
    p.add_argument("--report", type=Path, default=None,
                   help="JSON report path (单文件模式)")


def _restore_dispatch(args) -> int:
    docx_paths = [Path(p).expanduser().resolve() for p in args.docx_paths]
    missing = [p for p in docx_paths if not p.exists()]
    if missing:
        for m in missing:
            print(f"[ERR] not found: {m}", file=sys.stderr)
        return 2

    output: Optional[Path] = getattr(args, "output", None)
    output_dir: Optional[Path] = None
    output_single: Optional[Path] = None
    if output and len(docx_paths) >= 2:
        output_dir = output  # treat as directory
    elif output:
        output_single = output

    return cmd_restore_batch(
        docx_paths,
        output_dir=output_dir,
        output_single=output_single,
        dry_run=getattr(args, "dry_run", False),
        inplace=getattr(args, "inplace", False),
        no_backup=getattr(args, "no_backup", False),
        no_llm=getattr(args, "no_llm", False),
        force=getattr(args, "force", False),
        max_workers=getattr(args, "max_workers", None),
        yaml_profile=getattr(args, "yaml_profile", None),
    )


def register_restore(subparsers) -> None:
    """Register `restore` as a new target under `audit-styleset` (alias `styleset`).

    This is a separate registration entry from `register()` (which sets up
    the `fix <subcmd>` group). Called from sub/__init__.py register_all
    AFTER audit_styleset.register has set up the audit-styleset group.
    """
    from ._dispatch import get_or_add_group, get_or_add_subparsers

    # 'audit-styleset' is the canonical name; 'styleset' is its alias.
    # get_or_add_group looks up by name in choices map — alias should resolve.
    grp = get_or_add_group(subparsers, "audit-styleset",
                           "style-set health audits + restore (9-step)")
    grp_sub = get_or_add_subparsers(grp, dest="styleset_target")

    existing = getattr(grp_sub, "choices", {}) or {}
    if "restore" in existing:
        return
    sp = grp_sub.add_parser(
        "restore",
        help="eco-flow styleset 9-step 综合修复 (rename×3 → rebrand → create×5 → firstLine → clear → cover → pool-cleanup)",
    )
    _add_restore_args(sp)
    sp.set_defaults(func=_restore_dispatch)


# ─── extend module-level register() to also wire register_restore() ────
_orig_register = register


def register(subparsers) -> None:  # type: ignore[no-redef]
    _orig_register(subparsers)
    register_restore(subparsers)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="doctools-fix-styleset",
        description="docx-health-v2 · style-set fix family with shape_contract gate",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)
    register(sub)
    args = parser.parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    sys.exit(main())
