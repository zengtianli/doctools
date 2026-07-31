#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""biddiff.py — 标书对照家族三合一（2026-07-31 家族折叠；三件同为
bid-diff-and-revise 蒸馏批次）

子命令 ↔ 原脚本（函数体逐字搬移；模块级 main/apply_path 改名 _<sub> 后缀；
normalize/best_match 两版实现不同（seq 版四档中文标点表 vs ref 版七对+阈值常量），
ref 版改名 *_ref 各自保留，不合并语义）：

    seq  ← seqdiff.py（逐段 sequence diff: src vs dst → MD 报告）
    ref  ← compare_vs_ref.py（改动草稿 MD 改为段 vs 参考 docx 雷同检查）
    gen  ← gen_rules.py（改动草稿 MD → rules JSON for track-changes）

各子命令 CLI 与原独立脚本逐字一致：python3 sub/biddiff.py <sub> …。
退役原件在 ~/.Trash/consolidation-20260731/biddiff/（含 MANIFEST.md）。
"""
from __future__ import annotations

import argparse  # noqa: E402
import json  # noqa: E402
import re  # noqa: E402
import sys  # noqa: E402
from difflib import SequenceMatcher  # noqa: E402
from pathlib import Path  # noqa: E402

from docx import Document  # noqa: E402


# ══════════ seq ← seqdiff.py ══════════

EXACT = "原样照搬"
HIGH = "小改"   # >= 0.85
MID = "改写"    # 0.6-0.85
NEW = "新增"


def normalize(text: str) -> str:
    t = text.strip()
    t = re.sub(r"\s+", "", t)
    pairs = [
        ("\uff0c", ","), ("\u3002", "."), ("\uff1b", ";"), ("\uff1a", ":"),
        ("\uff08", "("), ("\uff09", ")"), ("\uff01", "!"), ("\uff1f", "?"),
        ("\u201c", '"'), ("\u201d", '"'), ("\u2018", "'"), ("\u2019", "'"),
        ("\u3001", ","),
    ]
    for a, b in pairs:
        t = t.replace(a, b)
    return t


def heading_level(p) -> int:
    name = (p.style.name or "").strip()
    m = re.match(r"^Heading\s+(\d+)$", name, re.IGNORECASE)
    if m:
        return int(m.group(1))
    m = re.match(r"^标题\s*(\d+)$", name)
    if m:
        return int(m.group(1))
    return 0


def iter_block_items(parent):
    from docx.document import Document as _Doc
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _Doc):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_paragraphs(doc_path: Path):
    doc = Document(str(doc_path))
    seq = 0
    heading_stack = {}
    out = []

    def emit(text, is_table):
        nonlocal seq
        text = text.strip()
        if not text:
            return
        chain = " > ".join(heading_stack[k] for k in sorted(heading_stack)) or "(前言)"
        out.append((seq, chain, is_table, text, normalize(text)))
        seq += 1

    def walk_block(block, in_table=False):
        if hasattr(block, "rows"):
            for row in block.rows:
                for cell in row.cells:
                    for sub in iter_block_items(cell):
                        walk_block(sub, in_table=True)
        else:
            lvl = heading_level(block)
            text = block.text.strip()
            if lvl > 0 and text and not in_table:
                heading_stack[lvl] = text
                for k in list(heading_stack):
                    if k > lvl:
                        del heading_stack[k]
                emit(text, False)
            else:
                emit(text, in_table)

    for blk in iter_block_items(doc):
        walk_block(blk)
    return out


def best_match(target_norm, src_index, src_norms_list):
    if not target_norm:
        return (NEW, None, None)
    if target_norm in src_index:
        return (EXACT, 1.0, src_index[target_norm])
    target_len = len(target_norm)
    if target_len < 4:
        return (NEW, None, None)
    best_r = 0.0
    best_seq = None
    sm = SequenceMatcher(autojunk=False)
    sm.set_seq2(target_norm)
    for src_norm, src_seq in src_norms_list:
        if abs(len(src_norm) - target_len) > target_len * 0.6:
            continue
        sm.set_seq1(src_norm)
        r = sm.ratio()
        if r > best_r:
            best_r = r
            best_seq = src_seq
            if r >= 0.99:
                break
    if best_r >= 0.85:
        return (HIGH, best_r, best_seq)
    if best_r >= 0.6:
        return (MID, best_r, best_seq)
    return (NEW, best_r if best_r > 0 else None, best_seq)


def main_seq():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--src", required=True, help="源 docx（旧版）")
    ap.add_argument("--dst", required=True, help="目标 docx（新版，查重对象）")
    ap.add_argument("--out", required=True, help="输出 MD 路径")
    ap.add_argument("--noise-len", type=int, default=8, help="短于 N 字符的段视为表格字段噪声（默认 8）")
    args = ap.parse_args()

    src_paras = extract_paragraphs(Path(args.src))
    dst_paras = extract_paragraphs(Path(args.dst))
    print(f"源 {len(src_paras)} 段 / 目标 {len(dst_paras)} 段", file=sys.stderr)

    src_index, src_chain, src_text = {}, {}, {}
    src_norms_list = []
    for s, ch, _, raw, norm in src_paras:
        if norm not in src_index:
            src_index[norm] = s
        src_chain[s] = ch
        src_text[s] = raw
        src_norms_list.append((norm, s))

    rows = []
    counts = {EXACT: 0, HIGH: 0, MID: 0, NEW: 0}
    cross_chapter = []
    for seq, chain, is_table, raw, norm in dst_paras:
        status, ratio, src_seq = best_match(norm, src_index, src_norms_list)
        counts[status] += 1
        rows.append((seq, chain, is_table, raw, status, ratio, src_seq))
        if status in (EXACT, HIGH) and src_seq is not None:
            if src_chain.get(src_seq) != chain:
                cross_chapter.append((seq, chain, raw, status, src_seq, src_chain.get(src_seq, "?")))

    L = []
    L.append(f"# 逐段对照 — {Path(args.src).stem} vs {Path(args.dst).stem}\n")
    L.append("> **判定**：归一化精确=原样照搬；ratio≥0.85=小改；0.6≤ratio<0.85=改写；其余=新增\n")
    L.append("")
    L.append("## 总体统计\n")
    L.append("| 类别 | 段数 | 占比 |")
    L.append("|------|------|------|")
    total = len(dst_paras)
    for k in [EXACT, HIGH, MID, NEW]:
        L.append(f"| {k} | {counts[k]} | {counts[k]*100/total:.1f}% |")
    L.append(f"| **总段数** | **{total}** | 100% |")
    L.append(f"\n**雷同风险段** = 原样照搬 + 小改 = {counts[EXACT] + counts[HIGH]} 段（占 {(counts[EXACT]+counts[HIGH])*100/total:.1f}%）\n")

    if cross_chapter:
        L.append("\n## ⚠️ 跨章节迁移\n")
        L.append("| 新版段 | 新章节 | → 源章节 | 状态 | 原文摘要 |")
        L.append("|--------|--------|---------|------|----------|")
        for seq, chain, raw, status, src_seq, src_ch in cross_chapter[:50]:
            snippet = raw[:40].replace("|", "\\|")
            L.append(f"| {seq} | {chain[:30]} | {src_ch[:30]} | {status} | {snippet}… |")
        if len(cross_chapter) > 50:
            L.append(f"\n（另有 {len(cross_chapter)-50} 条略）")
        L.append("")

    def is_noise(raw):
        n = normalize(raw)
        return len(n) < args.noise_len or bool(re.match(r"^[\d\W]+$", n))

    L.append("\n## 雷同风险段落（按新版章节分组，已过滤短字段）\n")
    grouped, noise_count = {}, 0
    for seq, chain, is_table, raw, status, ratio, src_seq in rows:
        if status in (EXACT, HIGH):
            if is_noise(raw):
                noise_count += 1
                continue
            grouped.setdefault(chain, []).append((seq, raw, status, ratio, src_seq))
    L.append(f"> 噪声过滤 {noise_count} 段；实质雷同 {sum(len(v) for v in grouped.values())} 段\n")

    if not grouped:
        L.append("✅ 无雷同段落\n")
    else:
        for chain in sorted(grouped):
            items = grouped[chain]
            L.append(f"### 「{chain}」 — {len(items)} 段")
            for seq, raw, status, ratio, src_seq in items:
                src_ch = src_chain.get(src_seq, "?")
                same = "✓同章节" if src_ch == chain else f"✗→「{src_ch}」"
                rs = f"{ratio:.2f}" if ratio is not None else "-"
                L.append(f"- **#{seq}** [{status} ratio={rs} {same}]")
                L.append(f"  > {raw}")
                if status == HIGH and src_seq is not None:
                    L.append(f"  > **源原文(#{src_seq})**: {src_text.get(src_seq,'')}")
            L.append("")

    L.append("\n## 附录：完整明细\n")
    L.append("| # | 章节链 | T? | 状态 | ratio | 源# | 原文 |")
    L.append("|---|--------|----|------|-------|-----|------|")
    for seq, chain, is_table, raw, status, ratio, src_seq in rows:
        snippet = raw[:50].replace("|", "\\|").replace("\n", " ")
        rs = f"{ratio:.2f}" if ratio is not None else ""
        srcs = str(src_seq) if src_seq is not None else ""
        tbl = "T" if is_table else ""
        L.append(f"| {seq} | {chain[:25]} | {tbl} | {status} | {rs} | {srcs} | {snippet} |")

    Path(args.out).write_text("\n".join(L), encoding="utf-8")
    print(f"OK -> {args.out}", file=sys.stderr)
    print(f"照搬 {counts[EXACT]} | 小改 {counts[HIGH]} | 改写 {counts[MID]} | 新增 {counts[NEW]}", file=sys.stderr)


# ---------------- pipeline adapter ----------------
def apply_path_seq(docx_path, args=None) -> dict:
    """pipeline-compatible adapter (跨文件 analyzer).

    docx_path = dst 新版 (查重对象). args 透传:
      - src (必需): 源 docx 旧版
      - noise_len: 噪声过滤短段阈值 (默认 8)
      - out / out_dir: 输出 MD 路径
    """
    from pathlib import Path as _P
    src_path = getattr(args, "src", None) if args else None
    if not src_path:
        return {"skipped": "no --src; seqdiff needs old-version docx"}
    src_path = _P(src_path)
    dst_path = _P(docx_path)
    noise_len = int(getattr(args, "noise_len", 8) or 8)
    out_path = getattr(args, "out", None)
    out_dir = getattr(args, "out_dir", None)
    if out_path:
        out = _P(out_path)
    elif out_dir:
        out = _P(out_dir) / f"seqdiff-{src_path.stem}-vs-{dst_path.stem}.md"
    else:
        out = dst_path.parent / "reports" / f"seqdiff-{src_path.stem}-vs-{dst_path.stem}.md"
    out.parent.mkdir(parents=True, exist_ok=True)

    src_paras = extract_paragraphs(src_path)
    dst_paras = extract_paragraphs(dst_path)
    src_index, src_chain, src_text = {}, {}, {}
    src_norms_list = []
    for s, ch, _, raw, norm in src_paras:
        if norm not in src_index:
            src_index[norm] = s
        src_chain[s] = ch
        src_text[s] = raw
        src_norms_list.append((norm, s))

    rows = []
    counts = {EXACT: 0, HIGH: 0, MID: 0, NEW: 0}
    for seq, chain, is_table, raw, norm in dst_paras:
        status, ratio, src_seq = best_match(norm, src_index, src_norms_list)
        counts[status] += 1
        rows.append((seq, chain, is_table, raw, status, ratio, src_seq))

    L = [f"# 逐段对照 — {src_path.stem} vs {dst_path.stem}\n"]
    total = len(dst_paras)
    L.append("## 总体统计\n")
    L.append("| 类别 | 段数 | 占比 |")
    L.append("|------|------|------|")
    for k in [EXACT, HIGH, MID, NEW]:
        pct = counts[k] * 100 / total if total else 0
        L.append(f"| {k} | {counts[k]} | {pct:.1f}% |")
    L.append(f"| **总段数** | **{total}** | 100% |")
    out.write_text("\n".join(L), encoding="utf-8")

    return {
        "src": str(src_path),
        "dst": str(dst_path),
        "src_paras": len(src_paras),
        "dst_paras": total,
        "exact": counts[EXACT],
        "high": counts[HIGH],
        "mid": counts[MID],
        "new": counts[NEW],
        "out": str(out),
    }


# ══════════ ref ← compare_vs_ref.py ══════════

THRESHOLD_HIGH = 0.85
THRESHOLD_MID = 0.6


def normalize_ref(t):
    t = re.sub(r"\s+", "", t)
    pairs = [("，", ","), ("。", "."), ("；", ";"), ("（", "("), ("）", ")"),
             ("\u201c", '"'), ("\u201d", '"')]
    for a, b in pairs:
        t = t.replace(a, b)
    return t


def extract_ref_paragraphs(path):
    doc = Document(str(path))
    out = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t and len(t) >= 8:
            out.append((f"P{i}", t, normalize_ref(t)))
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                t = cell.text.strip()
                if t and len(t) >= 8:
                    out.append((f"T{ti}.{ri}.{ci}", t, normalize_ref(t)))
    return out


def extract_revisions(drafts_dir: Path, glob_pattern: str):
    out = []
    for f in sorted(drafts_dir.glob(glob_pattern)):
        if "vs" in f.name or "清理" in f.name:
            continue
        text = f.read_text(encoding="utf-8")
        for m in re.finditer(r"\*\*改为[^*]*\*\*：?\s*\n((?:>\s.*\n)+)", text):
            block = m.group(1)
            content_lines = [re.sub(r"^>\s?", "", ln).strip() for ln in block.split("\n") if ln.strip()]
            content = " ".join(content_lines).strip()
            if len(content) >= 8:
                out.append((f.name, m.start(), content, normalize_ref(content)))
    return out


def best_match_ref(target_norm, ref_index, ref_norms_list):
    if target_norm in ref_index:
        return ("EXACT", 1.0, ref_index[target_norm])
    tl = len(target_norm)
    best_r, best_id = 0.0, None
    sm = SequenceMatcher(autojunk=False)
    sm.set_seq2(target_norm)
    for rn, ri in ref_norms_list:
        if abs(len(rn) - tl) > tl * 0.6:
            continue
        sm.set_seq1(rn)
        r = sm.ratio()
        if r > best_r:
            best_r, best_id = r, ri
            if r >= 0.99:
                break
    if best_r >= THRESHOLD_HIGH:
        return ("HIGH", best_r, best_id)
    if best_r >= THRESHOLD_MID:
        return ("MID", best_r, best_id)
    return ("OK", best_r, best_id)


def main_ref():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--drafts-dir", required=True, help="改动草稿 MD 所在目录")
    ap.add_argument("--ref", required=True, help="参考 docx（主标或基准）")
    ap.add_argument("--out", required=True, help="输出 MD 路径")
    ap.add_argument("--glob", default="*改动草稿.md", help="MD 文件 glob 模式（默认 *改动草稿.md）")
    args = ap.parse_args()

    ref_paras = extract_ref_paragraphs(Path(args.ref))
    print(f"参考 {len(ref_paras)} 段")
    ref_index = {n: i for i, _, n in ref_paras}
    ref_text = {i: t for i, t, _ in ref_paras}
    ref_norms_list = [(n, i) for i, _, n in ref_paras]

    revs = extract_revisions(Path(args.drafts_dir), args.glob)
    print(f"改为段 {len(revs)} 条")

    risks = []
    for fname, _, content, norm in revs:
        status, ratio, ref_id = best_match_ref(norm, ref_index, ref_norms_list)
        if status in ("EXACT", "HIGH", "MID"):
            risks.append((fname, content, status, ratio, ref_id, ref_text.get(ref_id, "")))

    L = [f"# 改动后内容 vs 参考 雷同检查\n",
         f"> 参考：`{Path(args.ref).name}`（{len(ref_paras)} 段）",
         f"> 改为段：{len(revs)} 条",
         f"> 阈值：HIGH ratio≥{THRESHOLD_HIGH} → 必须重写；MID {THRESHOLD_MID}≤ratio<{THRESHOLD_HIGH} → 建议调整\n"]
    if not risks:
        L.append("✅ 无雷同风险\n")
    else:
        n_h = sum(1 for r in risks if r[2] in ("EXACT", "HIGH"))
        n_m = sum(1 for r in risks if r[2] == "MID")
        L.append(f"## 总览：{len(risks)} 条风险（高 {n_h} / 中 {n_m}）\n")
        by_file = {}
        for r in risks:
            by_file.setdefault(r[0], []).append(r)
        for fname in sorted(by_file):
            L.append(f"\n## {fname}\n")
            for _, content, status, ratio, ref_id, ref_t in by_file[fname]:
                L.append(f"### [{status} ratio={ratio:.2f}] vs 参考 {ref_id}")
                L.append(f"**改为**：\n> {content[:300]}")
                L.append(f"**参考对应段**：\n> {ref_t[:300]}\n")

    Path(args.out).write_text("\n".join(L), encoding="utf-8")
    print(f"OK -> {args.out}")
    print(f"风险 {len(risks)} 条")


# ---------------- pipeline adapter ----------------
def apply_path_ref(docx_path, args=None) -> dict:
    """pipeline-compatible adapter (跨文件 analyzer).

    docx_path 充当 *主目标* (此 step 它就是 ``--ref`` 默认值);
    可被 args.ref 覆盖。args 透传:
      - drafts_dir (必需): 改动草稿 MD 目录
      - ref: 参考 docx (覆盖 docx_path)
      - glob: MD glob 模式 (默认 *改动草稿.md)
      - out / out_dir: 输出路径 (out 优先; 否则 out_dir/vs-<ref-stem>-雷同检查.md)
    """
    from pathlib import Path as _P
    drafts_dir = getattr(args, "drafts_dir", None) if args else None
    if not drafts_dir:
        return {"skipped": "no --drafts-dir; compare_vs_ref needs MD drafts dir"}
    ref_path = _P(getattr(args, "ref", None) or docx_path)
    glob_pattern = getattr(args, "glob", None) or "*改动草稿.md"
    out_path = getattr(args, "out", None)
    out_dir = getattr(args, "out_dir", None)
    if out_path:
        out = _P(out_path)
    elif out_dir:
        out = _P(out_dir) / f"vs-{ref_path.stem}-雷同检查.md"
    else:
        out = _P(docx_path).parent / "reports" / f"vs-{ref_path.stem}-雷同检查.md"
    out.parent.mkdir(parents=True, exist_ok=True)

    ref_paras = extract_ref_paragraphs(ref_path)
    ref_index = {n: i for i, _, n in ref_paras}
    ref_text = {i: t for i, t, _ in ref_paras}
    ref_norms_list = [(n, i) for i, _, n in ref_paras]
    revs = extract_revisions(_P(drafts_dir), glob_pattern)
    risks = []
    for fname, _, content, norm in revs:
        status, ratio, ref_id = best_match_ref(norm, ref_index, ref_norms_list)
        if status in ("EXACT", "HIGH", "MID"):
            risks.append((fname, content, status, ratio, ref_id, ref_text.get(ref_id, "")))
    L = [f"# 改动后内容 vs 参考 雷同检查\n",
         f"> 参考：`{ref_path.name}`({len(ref_paras)} 段)",
         f"> 改为段：{len(revs)} 条\n"]
    if not risks:
        L.append("✅ 无雷同风险\n")
    else:
        L.append(f"## 总览：{len(risks)} 条风险\n")
        by_file = {}
        for r in risks:
            by_file.setdefault(r[0], []).append(r)
        for fname in sorted(by_file):
            L.append(f"\n## {fname}\n")
            for _, content, status, ratio, ref_id, ref_t in by_file[fname]:
                L.append(f"### [{status} ratio={ratio:.2f}] vs 参考 {ref_id}")
                L.append(f"**改为**：\n> {content[:300]}")
                L.append(f"**参考对应段**：\n> {ref_t[:300]}\n")
    out.write_text("\n".join(L), encoding="utf-8")
    return {
        "ref": str(ref_path),
        "ref_paras": len(ref_paras),
        "revisions": len(revs),
        "risks": len(risks),
        "high": sum(1 for r in risks if r[2] in ("EXACT", "HIGH")),
        "mid": sum(1 for r in risks if r[2] == "MID"),
        "out": str(out),
    }


# ══════════ gen ← gen_rules.py ══════════

PAIR_BLOCK_RE = re.compile(
    r"\*\*原文(?:\s*#?([\w\d]+))?[^*]*\*\*[^\n]*\n((?:>\s?.*\n)+)", re.MULTILINE
)
REVISE_BLOCK_RE = re.compile(
    r"\*\*改为(?:\s*#?([\w\d]+))?[^*]*\*\*[^\n]*\n((?:>\s?.*\n)+)", re.MULTILINE
)
SECTION_RE = re.compile(r"^## 改动\s*\d+", re.MULTILINE)


def extract_quote(block):
    lines = []
    for ln in block.split("\n"):
        if not ln.startswith(">"):
            continue
        content = re.sub(r"^>\s?", "", ln).rstrip()
        if content.startswith("**") and "原文" in content[:20]:
            continue
        lines.append(content)
    return "\n".join(lines).strip()


def parse_md(md_path):
    text = md_path.read_text(encoding="utf-8")
    rules = []
    starts = [m.start() for m in SECTION_RE.finditer(text)]
    if not starts:
        return rules
    sections = []
    for i, start in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(text)
        sections.append(text[start:end])

    for sec_idx, sec in enumerate(sections, 1):
        first_line = sec.split("\n", 1)[0]
        title = first_line.replace("## 改动", "").strip()
        originals = [(m.group(1), extract_quote(m.group(2))) for m in PAIR_BLOCK_RE.finditer(sec)]
        revises = [(m.group(1), extract_quote(m.group(2))) for m in REVISE_BLOCK_RE.finditer(sec)]
        if not originals or not revises:
            continue
        orig_ids = {oid: txt for oid, txt in originals if oid}
        rev_ids = {rid: txt for rid, txt in revises if rid}
        common = orig_ids.keys() & rev_ids.keys()
        if common:
            for cid in sorted(common, key=lambda x: int(re.sub(r"\D", "", x) or 0)):
                rules.append({
                    "find": orig_ids[cid],
                    "replace": rev_ids[cid],
                    "comment": f"[{md_path.stem} 改动{sec_idx} #{cid}] {title[:40]}"
                })
            continue
        if len(originals) == len(revises):
            for (oid, otxt), (rid, rtxt) in zip(originals, revises):
                rules.append({
                    "find": otxt,
                    "replace": rtxt,
                    "comment": f"[{md_path.stem} 改动{sec_idx}] {title[:40]}"
                })
        else:
            print(f"  ⚠️ {md_path.name} 改动{sec_idx}: 原文{len(originals)}/改为{len(revises)} 不匹配，跳过")
    return rules


def try_quote_swap(s):
    """直引号 → 弯引号候选。"""
    out = []
    toggle = False
    for ch in s:
        if ch == '"':
            out.append("\u201d" if toggle else "\u201c")
            toggle = not toggle
        elif ch == "'":
            out.append("\u2019" if toggle else "\u2018")
        else:
            out.append(ch)
    return "".join(out)


def is_title(text):
    t = text.strip()
    if len(t) > 80:
        return False
    if re.match(r"^[（(][一二三四五六七八九十]+[）)]", t):
        return True
    if re.match(r"^第[一二三四五六七八九十]+[章节部分]", t):
        return True
    if any(kw in t for kw in ["子方案", "总体技术", "工作部署", "技术路线", "技术思路", "技术方案"]) and len(t) < 60:
        return True
    return False


def main_gen():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--drafts-dir", required=True, help="改动草稿 MD 目录")
    ap.add_argument("--docx", required=True, help="目标 docx（用于 find 存在性校验 + 引号自适配）")
    ap.add_argument("--out", required=True, help="输出 rules JSON 路径")
    ap.add_argument("--glob", default="*-改动草稿.md", help="MD 文件 glob（默认 *-改动草稿.md）")
    ap.add_argument("--no-title-skip", action="store_true", help="不跳过 title-only rule（默认跳过）")
    ap.add_argument("--no-paren-strip", action="store_true", help="不剥除说明性括号（默认剥除）")
    args = ap.parse_args()

    md_files = sorted(Path(args.drafts_dir).glob(args.glob))
    md_files = [f for f in md_files if "vs" not in f.name and "清理-清单" not in f.name]
    print(f"待处理 MD: {len(md_files)} 份")
    all_rules = []
    for f in md_files:
        rules = parse_md(f)
        print(f"  {f.name}: {len(rules)} 条")
        all_rules.extend(rules)

    # 加载 docx 全文用于 find 验证
    doc = Document(args.docx)
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                docx_text += "\n" + cell.text

    # Pass 1: 去空 + 去 find==replace + 引号 swap 守卫
    valid = []
    quote_fixed = 0
    for r in all_rules:
        if not r["find"].strip():
            continue
        if r["find"] == r["replace"]:
            continue
        if r["find"] not in docx_text:
            alt = try_quote_swap(r["find"])
            if alt != r["find"] and alt in docx_text:
                r["find"] = alt
                r["replace"] = try_quote_swap(r["replace"])
                quote_fixed += 1
        for m in re.finditer(r"〔\s*〕", r["find"] + r["replace"]):
            print(f"  ⚠️ 空〔〕：{r['comment']}")
        valid.append(r)
    if quote_fixed:
        print(f"  ✓ 引号 swap 修复 {quote_fixed} 条")

    # Pass 2: title-only 跳过守卫
    if not args.no_title_skip:
        skipped = 0
        kept = []
        for r in valid:
            if is_title(r["find"]) and is_title(r["replace"]):
                skipped += 1
                continue
            kept.append(r)
        if skipped:
            print(f"  ✓ 跳过 title-only rule: {skipped} 条")
        valid = kept

    # Pass 3: 说明性括号剥除守卫
    if not args.no_paren_strip:
        PAREN = re.compile(r"[（(](?:面向|针对|含|本节|破除|重写|说明)[^）)]{0,40}[）)]")
        stripped = 0
        for r in valid:
            new_r = PAREN.sub("", r["replace"])
            if new_r != r["replace"]:
                r["replace"] = new_r.strip()
                stripped += 1
        if stripped:
            print(f"  ✓ 剥括号说明: {stripped} 条")

    Path(args.out).parent.mkdir(parents=True, exist_ok=True)
    with open(args.out, "w", encoding="utf-8") as fp:
        json.dump(valid, fp, ensure_ascii=False, indent=2)
    print(f"\nOK -> {args.out}")
    print(f"最终 {len(valid)} 条规则（原 {len(all_rules)}，过滤 {len(all_rules)-len(valid)}）")


# ---------------- pipeline adapter ----------------
def apply_path_gen(docx_path, args=None) -> dict:
    """pipeline-compatible adapter (跨文件 analyzer).

    docx_path = 目标 docx (用于 find 存在性校验 + 引号自适配)。
    args 透传:
      - drafts_dir (必需): 改动草稿 MD 目录
      - glob: MD glob 模式 (默认 *-改动草稿.md)
      - no_title_skip / no_paren_strip: 守卫开关
      - out / out_dir: 输出 JSON 路径
    """
    from pathlib import Path as _P
    drafts_dir = getattr(args, "drafts_dir", None) if args else None
    if not drafts_dir:
        return {"skipped": "no --drafts-dir; gen_rules needs MD drafts dir"}
    glob_pattern = getattr(args, "glob", None) or "*-改动草稿.md"
    out_path = getattr(args, "out", None)
    out_dir = getattr(args, "out_dir", None)
    if out_path:
        out = _P(out_path)
    elif out_dir:
        out = _P(out_dir) / f"_revise_rules-{_P(docx_path).stem}.json"
    else:
        out = _P(docx_path).parent / "reports" / f"_revise_rules-{_P(docx_path).stem}.json"
    out.parent.mkdir(parents=True, exist_ok=True)

    no_title_skip = bool(getattr(args, "no_title_skip", False))
    no_paren_strip = bool(getattr(args, "no_paren_strip", False))

    md_files = sorted(_P(drafts_dir).glob(glob_pattern))
    md_files = [f for f in md_files if "vs" not in f.name and "清理-清单" not in f.name]
    all_rules = []
    for f in md_files:
        all_rules.extend(parse_md(f))

    doc = Document(str(docx_path))
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                docx_text += "\n" + cell.text

    valid = []
    quote_fixed = 0
    for r in all_rules:
        if not r["find"].strip():
            continue
        if r["find"] == r["replace"]:
            continue
        if r["find"] not in docx_text:
            alt = try_quote_swap(r["find"])
            if alt != r["find"] and alt in docx_text:
                r["find"] = alt
                r["replace"] = try_quote_swap(r["replace"])
                quote_fixed += 1
        valid.append(r)
    skipped = 0
    if not no_title_skip:
        kept = []
        for r in valid:
            if is_title(r["find"]) and is_title(r["replace"]):
                skipped += 1
                continue
            kept.append(r)
        valid = kept
    stripped = 0
    if not no_paren_strip:
        PAREN = re.compile(r"[(（](?:面向|针对|含|本节|破除|重写|说明)[^)）]{0,40}[)）]")
        for r in valid:
            new_r = PAREN.sub("", r["replace"])
            if new_r != r["replace"]:
                r["replace"] = new_r.strip()
                stripped += 1

    with open(out, "w", encoding="utf-8") as fp:
        json.dump(valid, fp, ensure_ascii=False, indent=2)

    return {
        "md_files": len(md_files),
        "raw_rules": len(all_rules),
        "valid_rules": len(valid),
        "quote_fixed": quote_fixed,
        "title_skipped": skipped,
        "paren_stripped": stripped,
        "out": str(out),
    }


# ──────────────────────────── 家族入口（子命令分发）────────────────────────────

SUBCOMMANDS = {
    "seq": main_seq,
    "ref": main_ref,
    "gen": main_gen,
}


def main(argv: list[str] | None = None) -> int:
    args = list(sys.argv[1:] if argv is None else argv)
    if not args or args[0] in ("-h", "--help"):
        print("usage: biddiff.py {" + ",".join(SUBCOMMANDS) + "} <args…>\n"
              "每个子命令的参数与原独立脚本逐字一致：biddiff.py <sub> --help 查看。")
        return 0 if args else 2
    sub, rest = args[0], args[1:]
    fn = SUBCOMMANDS.get(sub)
    if fn is None:
        print(f"[biddiff] unknown subcommand: {sub!r}; choices={list(SUBCOMMANDS)}",
              file=sys.stderr)
        return 2
    saved = sys.argv[:]
    sys.argv = [sys.argv[0]] + rest
    try:
        rc = fn()
        return int(rc) if isinstance(rc, int) else 0
    finally:
        sys.argv = saved


if __name__ == "__main__":
    sys.exit(main())
