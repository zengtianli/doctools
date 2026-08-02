#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""split.py — 拆分/换身家族二合一（2026-07-31 家族折叠）

子命令 ↔ 原脚本（函数体逐字搬移；模块级 main 改名 _<sub> 后缀）：

    by-h1         ← split_by_h1.py（plan_slices / write_slice / sanitize_filename
                    公有名保留：chapters_sync 与 pipeline_lib 靠它）
    body-replace  ← body_replace.py（_is_h1_elem/_paragraph_text 与 by-h1 段真重复，
                    保 by-h1 的富实现（带 styleId→styleName 映射），调用点改传映射）

各子命令 CLI 与原独立脚本逐字一致：python3 sub/split.py <sub> …。
退役原件在 ~/.Trash/consolidation-20260731/split/（含 MANIFEST.md）。
"""
from __future__ import annotations

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[3] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py

# sub/ 自身进 sys.path —— docx_cli 的 _dispatch 用 spec_from_file_location 加载,
# 不带脚本目录, 裸 import _cli_common 会 ImportError (append 不是 insert(0))
_sys.path.append(str(_Path(__file__).resolve().parent))
import _cli_common as _cc  # noqa: E402  家族 main() 样板 SSOT

import argparse  # noqa: E402
import re  # noqa: E402
import shutil  # noqa: E402
import sys  # noqa: E402
from copy import deepcopy  # noqa: E402
from pathlib import Path  # noqa: E402
from typing import Optional  # noqa: E402

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx 未安装 (pip install python-docx)", file=sys.stderr)
    sys.exit(2)

# 媒体去冗余(默认开): 每片只留它 body 真引用的 media + 同步裁 rels。
# 不做的话每片都扛整本 media(bug: 整本 27MB → 每章都 27MB)。复用 strip orphan-media 的
# deep 扫描(专为 split/table-extract "body 已裁 rels 未裁" 场景设计)。三态 import 兼容:
# 包内(pipeline `from . import`) / 脚本(sys.path[0]=sub/) / docx_cli runpy(__main__)。
_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))
try:
    from . import strip as _som  # type: ignore  # scan_orphans/rewrite_skip 公有名保留
except ImportError:
    try:
        import strip as _som  # type: ignore
    except ImportError:
        _som = None  # 缺 lxml 等极端情况: 退化为不去冗余, 不让 split 失败


# ══════════ by-h1 ← split_by_h1.py ══════════

H1_STYLES = {
    "Heading 1",
    "标题 1",
    "heading 1",
    "1",
    "10",
    "1.1.1.1 N级标题",
    # 院方系列报告（浙江省水利水电勘测设计院可用水量/生态流量报告）样式名
    "1一级标题",
}


_ILLEGAL_FILENAME_RE = re.compile(r'[/\\:*?"<>|\r\n\t]')
_MULTI_WS_RE = re.compile(r"\s+")


def sanitize_filename(name: str, max_len: int = 100) -> str:
    """Replace illegal filename chars with _, compress whitespace, strip, truncate."""
    if not name:
        return "untitled"
    n = _ILLEGAL_FILENAME_RE.sub("_", name)
    n = _MULTI_WS_RE.sub(" ", n).strip()
    if not n:
        return "untitled"
    if len(n) > max_len:
        n = n[:max_len].rstrip()
    return n


def get_style_name(p) -> str:
    try:
        return (p.style.name or "") if p.style is not None else ""
    except Exception:
        return ""


def is_h1_paragraph(p) -> bool:
    """True if paragraph carries an H1-class style (per H1_STYLES set)."""
    return get_style_name(p) in H1_STYLES


def _body_children(doc) -> list:
    """Return ordered list of direct child elements of <w:body> (excludes sectPr)."""
    body = doc.element.body
    # body 的最后一个子通常是 sectPr (文档级节属性), 不动它
    return list(body)


def _iter_paragraphs_in_element(elem):
    """Find all <w:p> descendants of a body-level element (for tables, etc.)."""
    return elem.iter(qn("w:p"))


def _paragraph_text(p_elem) -> str:
    """Extract concatenated text from <w:t> nodes under a <w:p>."""
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def _build_style_id_to_name(doc) -> dict:
    """Map styleId → styleName for the doc (院方系列 styleId 如 '1f8',
    与 styleName '1一级标题' 不同 → 必须建映射才能命中 H1_STYLES)."""
    out = {}
    try:
        styles_xml = doc.styles.element
        for s in styles_xml.findall(qn("w:style")):
            sid = s.get(qn("w:styleId"))
            name_el = s.find(qn("w:name"))
            nm = name_el.get(qn("w:val")) if name_el is not None else None
            if sid and nm:
                out[sid] = nm
    except Exception:
        pass
    return out


def _is_h1_elem(elem, style_id_to_name: Optional[dict] = None) -> bool:
    """Check if a body-level <w:p> element carries an H1 style."""
    if elem.tag != qn("w:p"):
        return False
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return False
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return False
    style_val = pStyle.get(qn("w:val")) or ""
    if style_val in H1_STYLES:
        return True
    # docx style id may differ from style name; also try common normalizations
    if style_val.lower() in {s.lower() for s in H1_STYLES}:
        return True
    # Heuristic: some templates use "1" or "10" as Heading 1 styleId
    if style_val in {"1", "10", "Heading1"}:
        return True
    # 院方系列: styleId (如 '1f8') ≠ styleName (如 '1一级标题') → 查 map
    if style_id_to_name is not None:
        nm = style_id_to_name.get(style_val)
        if nm and (nm in H1_STYLES or nm.lower() in {s.lower() for s in H1_STYLES}):
            return True
    return False


def plan_slices(docx_path: Path, include_frontmatter: bool, doc=None):
    """Inspect docx → return (slices, sect_idx, h1_count).

    slices: list[{idx, title, start, end, is_frontmatter}]
    h1_count: number of H1 elements detected in body (independent of slice emission)

    `doc`: if provided, skip the source parse (pipeline reuse).
    """
    if doc is None:
        doc = Document(str(docx_path))
    body = doc.element.body
    children = list(body)
    style_id_to_name = _build_style_id_to_name(doc)
    # Locate sectPr (final node) — we exclude it from slicing range
    sect_idx = len(children)
    for i in range(len(children) - 1, -1, -1):
        if children[i].tag == qn("w:sectPr"):
            sect_idx = i
            break

    # Find H1 positions among children [0, sect_idx)
    h1_positions: list[tuple[int, str]] = []
    for i in range(sect_idx):
        elem = children[i]
        if _is_h1_elem(elem, style_id_to_name):
            title = _paragraph_text(elem).strip() or "untitled"
            h1_positions.append((i, title))

    slices: list[dict] = []
    idx_counter = 0
    if h1_positions:
        first_h1 = h1_positions[0][0]
        if include_frontmatter and first_h1 > 0:
            slices.append({
                "idx": idx_counter,
                "title": "frontmatter",
                "start": 0,
                "end": first_h1,
                "is_frontmatter": True,
            })
            idx_counter += 1
        for k, (pos, title) in enumerate(h1_positions):
            end = h1_positions[k + 1][0] if k + 1 < len(h1_positions) else sect_idx
            slices.append({
                "idx": idx_counter,
                "title": title,
                "start": pos,
                "end": end,
                "is_frontmatter": False,
            })
            idx_counter += 1
    else:
        # No H1 found; treat full body as one slice if frontmatter requested
        if include_frontmatter and sect_idx > 0:
            slices.append({
                "idx": 0,
                "title": "frontmatter",
                "start": 0,
                "end": sect_idx,
                "is_frontmatter": True,
            })
    return slices, sect_idx, len(h1_positions)


def write_slice(src_docx: Path, dst_docx: Path, start: int, end: int,
                prune_media: bool = True) -> tuple[int, int]:
    """Copy src→dst, then prune body keeping only children [start, end) + sectPr.

    prune_media (默认 True): 切完后丢掉本片 body 不再引用的 media + 同步裁 rels,
    避免每片都扛整本媒体(否则整本 27MB → 每片 27MB)。

    Returns (paragraph_count_in_slice, file_bytes).
    """
    dst_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(str(src_docx), str(dst_docx))
    doc = Document(str(dst_docx))
    body = doc.element.body
    children = list(body)
    # Identify sectPr (last)
    sect_elem = None
    for i in range(len(children) - 1, -1, -1):
        if children[i].tag == qn("w:sectPr"):
            sect_elem = children[i]
            break
    # Determine elements to KEEP — by their original index in children list
    keep_set = set(range(start, end))
    # Remove all body children NOT in keep_set (and not the trailing sectPr)
    for i, elem in enumerate(children):
        if elem is sect_elem:
            continue
        if i not in keep_set:
            body.remove(elem)
    doc.save(str(dst_docx))
    # 去冗余 media: 本片 body 已裁但 rels/media 仍是整本 → deep 扫只留真引用的, 同步裁 rels。
    # best-effort: 任何异常都不让 split 失败(去冗余是优化, 不是正确性前提)。
    if prune_media and _som is not None:
        try:
            scan = _som.scan_orphans(dst_docx, deep=True)
            if scan.get("orphan_count"):
                _som.rewrite_skip(dst_docx, dst_docx, set(scan["orphans"]))
        except Exception:
            pass
    # Count paragraphs in saved slice
    doc2 = Document(str(dst_docx))
    para_count = len(doc2.paragraphs)
    file_bytes = dst_docx.stat().st_size
    return para_count, file_bytes


def run_split(
    src_docx: Path,
    out_dir: Path,
    include_frontmatter: bool = False,
    allow_no_h1: bool = False,
    dry_run: bool = False,
    name_pattern: str = "{idx:02d}-{title}.docx",
    doc=None,
    prune_media: bool = True,
) -> dict:
    """Execute split-by-h1; reuses provided `doc` for planning if given.

    Returns a report dict (used by pipeline built-in step). Raises on
    fail-fast conditions (0 H1 + no allow_no_h1) for caller to handle.
    """
    src = Path(src_docx).expanduser().resolve()
    if not src.exists():
        return {"error": f"input docx not found: {src}", "exit_code": 2}
    out_dir = Path(out_dir).expanduser().resolve()

    slices, sect_idx, h1_count = plan_slices(src, include_frontmatter, doc=doc)

    if h1_count == 0 and not allow_no_h1:
        return {
            "error": "0 Heading-1 detected (docx unhealthy); run /docx health first or pass --allow-no-h1",
            "exit_code": 3,
            "h1_count": 0,
        }

    if not slices:
        return {"h1_count": h1_count, "slices_emitted": 0, "note": "no slices to emit"}

    if dry_run:
        plan = []
        for s in slices:
            safe = sanitize_filename(s["title"])
            fname = name_pattern.format(idx=s["idx"], title=safe)
            plan.append({"idx": s["idx"], "fname": fname, "title": s["title"]})
        return {"h1_count": h1_count, "slices_planned": plan, "dry_run": True}

    out_dir.mkdir(parents=True, exist_ok=True)
    emitted = []
    failed = []
    for s in slices:
        safe = sanitize_filename(s["title"])
        fname = name_pattern.format(idx=s["idx"], title=safe)
        dst = out_dir / fname
        try:
            paras, nbytes = write_slice(src, dst, s["start"], s["end"], prune_media=prune_media)
            emitted.append({
                "idx": s["idx"], "fname": fname, "paragraphs": paras, "bytes": nbytes,
            })
        except Exception as e:
            failed.append({"idx": s["idx"], "fname": fname,
                           "error": f"{type(e).__name__}: {e}"})
    return {
        "h1_count": h1_count,
        "slices_emitted": len(emitted),
        "slices_failed": len(failed),
        "out_dir": str(out_dir),
        "emitted": emitted,
        "failed": failed,
    }


def main_by_h1() -> int:
    ap = argparse.ArgumentParser(
        description="Split a DOCX by Heading 1 into N independent DOCX files.",
    )
    ap.add_argument("--docx", required=True, help="input docx path")
    ap.add_argument("--out-dir", required=True, help="output directory (mkdir -p)")
    ap.add_argument(
        "--name-pattern",
        default="{idx:02d}-{title}.docx",
        help="output filename pattern, default '{idx:02d}-{title}.docx'",
    )
    ap.add_argument(
        "--include-frontmatter", action="store_true",
        help="emit content before first H1 as 00-frontmatter.docx",
    )
    ap.add_argument(
        "--dry-run", action="store_true",
        help="print plan only, don't write files",
    )
    ap.add_argument(
        "--allow-no-h1", action="store_true",
        help="suppress unhealthy-docx fail-fast when 0 H1 detected (rarely needed; "
             "default behavior is to FAIL and instruct user to run /docx health first)",
    )
    ap.add_argument(
        "--keep-all-media", action="store_true",
        help="禁用默认的媒体去冗余: 每片保留整本 media(默认 OFF=去冗余, 每片只留自己引用的图)",
    )
    args = ap.parse_args()

    src = Path(args.docx).expanduser().resolve()
    if not src.exists():
        print(f"ERROR: input docx not found: {src}", file=sys.stderr)
        return 2
    out_dir = Path(args.out_dir).expanduser().resolve()

    slices, sect_idx, h1_count = plan_slices(src, args.include_frontmatter)

    # Health gate: 0 H1 detected = docx unhealthy signal (default fail-fast).
    # Iron rule [[docx-split-fail-run-health-first]]: don't patch around bad data,
    # tell the user to run /docx health first (scaffold already exists).
    if h1_count == 0 and not args.allow_no_h1:
        print(
            "\n".join([
                "",
                "ERROR: 0 Heading-1 detected in body — docx likely UNHEALTHY.",
                f"  file: {src}",
                "",
                "  Recognized H1 styles: " + ", ".join(sorted(H1_STYLES)),
                "",
                "  Likely causes:",
                "    - chapter titles styled as Normal/body paragraphs (not Heading 1)",
                "    - heading-level-skew (real H1 demoted to H2/H3/...)",
                "    - caption-outline-pollution (figure/table captions stole H1 slot)",
                "    - custom Chinese style name not in H1_STYLES whitelist",
                "",
                "  Fix path (scaffold already in place):",
                f"    /docx health diagnose '{src}'    # see which病种 hits",
                f"    /docx health full     '{src}'    # diagnose + auto-fix safe + re-diagnose",
                "    # then re-run this split script",
                "",
                "  Escape hatch (rare): pass --allow-no-h1 if you really want the current behavior",
                "  (emits only frontmatter.docx, requires --include-frontmatter).",
                "",
            ]),
            file=sys.stderr,
        )
        return 3

    if not slices:
        # h1_count > 0 path can't reach here; this is the --allow-no-h1 + no frontmatter case
        print(f"[split-by-h1] no H1 (and no frontmatter requested) — nothing to do")
        return 0

    if args.dry_run:
        print(f"# DRY RUN — would write {len(slices)} files to {out_dir}/")
        for s in slices:
            safe = sanitize_filename(s["title"])
            fname = args.name_pattern.format(idx=s["idx"], title=safe)
            print(f"  [{s['idx']:>2}] children[{s['start']}:{s['end']}]  → {fname}  (title={s['title']!r})")
        return 0

    out_dir.mkdir(parents=True, exist_ok=True)
    n_ok = 0
    for s in slices:
        safe = sanitize_filename(s["title"])
        fname = args.name_pattern.format(idx=s["idx"], title=safe)
        dst = out_dir / fname
        try:
            paras, nbytes = write_slice(src, dst, s["start"], s["end"],
                                        prune_media=not args.keep_all_media)
            print(f"  [{s['idx']:>2}] {fname}  · {paras} paragraphs · {nbytes:,} bytes")
            n_ok += 1
        except Exception as e:
            print(f"  [{s['idx']:>2}] FAILED {fname}: {type(e).__name__}: {e}",
                  file=sys.stderr)
    print(f"OK: {n_ok} files written to {out_dir}")
    return 0 if n_ok == len(slices) else 1


# ══════════ body-replace ← body_replace.py ══════════

def _find_first_h1_index(children: list, sect_idx: int,
                         style_map: Optional[dict] = None) -> int:
    """Return index of first H1 element in children[0:sect_idx), or -1 if none."""
    for i in range(sect_idx):
        if _is_h1_elem(children[i], style_map):
            return i
    return -1


def _find_sect_idx(body) -> int:
    """Find index of trailing <w:sectPr> in body. Returns len(children) if none."""
    children = list(body)
    for i in range(len(children) - 1, -1, -1):
        if children[i].tag == qn("w:sectPr"):
            return i
    return len(children)


def _collect_used_styles(doc) -> set:
    """Collect all pStyle val referenced by paragraphs in body. Used for fallback check."""
    styles = set()
    for p in doc.element.body.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        v = pStyle.get(qn("w:val"))
        if v:
            styles.add(v)
    return styles


def _available_style_ids(doc) -> set:
    """Return set of styleId defined in doc's styles.xml (canonical resolvable keys).

    Note: w:val on w:pStyle MUST be a styleId, not a name. Names are only useful
    as a lookup target — we surface the name→id map separately via
    _available_name_to_id.
    """
    ids = set()
    try:
        for s in doc.styles:
            try:
                if s.style_id:
                    ids.add(s.style_id)
            except Exception:
                pass
    except Exception:
        pass
    return ids


def _available_name_to_id(doc) -> dict:
    """Map style.name -> style.style_id (for content val that uses names not ids)."""
    name_to_id: dict = {}
    try:
        for s in doc.styles:
            try:
                if s.name and s.style_id:
                    name_to_id[s.name] = s.style_id
            except Exception:
                pass
    except Exception:
        pass
    return name_to_id


def _has_image_refs(elem) -> int:
    """Count <a:blip> or <w:drawing> descendants — proxy for inline images."""
    n = 0
    for _ in elem.iter(qn("w:drawing")):
        n += 1
    return n


# Common heading style aliases (content styleId → shell preferred styleId/name).
# Tried in order; first one present in `available` wins.
_HEADING_ALIASES = {
    "Heading1":  ["Heading 1", "1", "标题 1"],
    "Heading2":  ["Heading 2", "2", "标题 2"],
    "Heading3":  ["Heading 3", "3", "标题 3"],
    "Heading4":  ["Heading 4", "4", "标题 4"],
    "Heading5":  ["Heading 5", "5", "标题 5"],
    "Heading6":  ["Heading 6", "6", "标题 6"],
    "Heading 1": ["Heading 1", "1", "标题 1"],
    "Heading 2": ["Heading 2", "2", "标题 2"],
    "Heading 3": ["Heading 3", "3", "标题 3"],
    "Heading 4": ["Heading 4", "4", "标题 4"],
    "Heading 5": ["Heading 5", "5", "标题 5"],
    "Heading 6": ["Heading 6", "6", "标题 6"],
    "标题 1":    ["Heading 1", "1", "标题 1"],
    "标题 2":    ["Heading 2", "2", "标题 2"],
    "标题 3":    ["Heading 3", "3", "标题 3"],
    "标题 4":    ["Heading 4", "4", "标题 4"],
    # Common pandoc-produced body styles → shell's body
    "FirstParagraph": ["Normal", "正文"],
    "BodyText":       ["Normal", "正文"],
    "Compact":        ["Normal", "正文"],
}


def _resolve_alias(v: str, available_ids: set, name_to_id: dict):
    """Return target styleId via alias map; tries direct id, then via name_to_id.

    Returns None if no alias hit resolvable styleId.
    """
    aliases = _HEADING_ALIASES.get(v, [])
    for a in aliases:
        # direct styleId hit
        if a in available_ids:
            return a
        # alias is a name → resolve to id
        if a in name_to_id:
            return name_to_id[a]
    return None


def _remap_styles_in_elem(elem, available_ids: set, name_to_id: dict,
                           fallback: str = "Normal") -> list:
    """Walk <w:p> descendants, rewrite pStyle val to a valid styleId.

    Order: (1) val is already a valid styleId → keep
           (2) val is a known name in shell → swap to its styleId
           (3) val matches _HEADING_ALIASES → use alias's resolved styleId
           (4) fall back to `fallback` styleId + record warning
    """
    warnings: list = []
    for p in elem.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        v = pStyle.get(qn("w:val"))
        if not v:
            continue
        # (1) already a valid styleId in shell
        if v in available_ids:
            continue
        # (2) val is a style.name in shell → swap to id
        if v in name_to_id:
            pStyle.set(qn("w:val"), name_to_id[v])
            continue
        # (3) alias map
        alias_target = _resolve_alias(v, available_ids, name_to_id)
        if alias_target is not None:
            pStyle.set(qn("w:val"), alias_target)
            continue
        # (4) fallback
        txt = _paragraph_text(p)[:30]
        warnings.append((v, txt))
        pStyle.set(qn("w:val"), fallback)
    return warnings


def body_replace(
    shell_path: Path,
    content_path: Path,
    out_path: Path,
    keep_shell_h1: bool = True,
) -> dict:
    """Core operation. Returns dict of stats for reporting."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(str(shell_path), str(out_path))

    out_doc = Document(str(out_path))
    out_body = out_doc.element.body
    out_children = list(out_body)
    out_sect_idx = _find_sect_idx(out_body)

    # Step 1: prune out body
    if keep_shell_h1:
        first_h1 = _find_first_h1_index(out_children, out_sect_idx,
                                        _build_style_id_to_name(out_doc))
        if first_h1 < 0:
            print(
                "WARN: --keep-shell-h1 set but shell has 0 H1 — will keep nothing from shell body",
                file=sys.stderr,
            )
            keep_until = 0
        else:
            keep_until = first_h1 + 1  # keep [0, first_h1] inclusive
    else:
        keep_until = 0

    sect_elem = None
    if out_sect_idx < len(out_children):
        sect_elem = out_children[out_sect_idx]

    # Remove everything from index keep_until up to (but not including) sectPr
    for i, elem in enumerate(out_children):
        if elem is sect_elem:
            continue
        if i >= keep_until and i < out_sect_idx:
            out_body.remove(elem)

    # Step 2: collect from content
    content_doc = Document(str(content_path))
    content_body = content_doc.element.body
    content_children = list(content_body)
    content_sect_idx = _find_sect_idx(content_body)

    if keep_shell_h1:
        c_first_h1 = _find_first_h1_index(content_children, content_sect_idx,
                                          _build_style_id_to_name(content_doc))
        if c_first_h1 < 0:
            content_start = 0  # no H1 in content → take everything
        else:
            content_start = c_first_h1 + 1  # skip content's first H1
    else:
        content_start = 0

    # Step 3: available styles in out (= shell's styles)
    available_ids = _available_style_ids(out_doc)
    name_to_id = _available_name_to_id(out_doc)

    # Step 4: deepcopy content children into out body, before sectPr
    appended = 0
    image_drops = 0
    style_warnings: list = []

    for i in range(content_start, content_sect_idx):
        src_elem = content_children[i]
        # Skip sectPr if it appears mid-body (rare; nested section breaks)
        if src_elem.tag == qn("w:sectPr"):
            continue
        new_elem = deepcopy(src_elem)
        # style fallback
        warns = _remap_styles_in_elem(new_elem, available_ids, name_to_id, fallback="Normal")
        style_warnings.extend(warns)
        # count images (will be broken — rId points to content's media)
        n_img = _has_image_refs(new_elem)
        image_drops += n_img
        # insert before sectPr
        if sect_elem is not None:
            sect_elem.addprevious(new_elem)
        else:
            out_body.append(new_elem)
        appended += 1

    out_doc.save(str(out_path))

    # Build stats
    out_doc2 = Document(str(out_path))
    final_paras = len(out_doc2.paragraphs)
    final_h1 = [p for p in out_doc2.paragraphs if p.style.name in H1_STYLES]
    first_h1_text = final_h1[0].text if final_h1 else ""
    used_styles = sorted({p.style.name for p in out_doc2.paragraphs if p.style is not None})

    # Print warnings to stderr
    if style_warnings:
        unique_styles = {}
        for v, txt in style_warnings:
            unique_styles.setdefault(v, txt)
        print(
            f"WARN: {len(style_warnings)} paragraphs referenced styles missing in shell — fell back to Normal",
            file=sys.stderr,
        )
        for v, txt in unique_styles.items():
            print(f"  · style {v!r} not in shell (sample: {txt!r})", file=sys.stderr)

    if image_drops:
        print(
            f"WARN: {image_drops} <w:drawing> elements copied without remapping image parts — "
            f"images will likely render broken. (known limitation, follow-up)",
            file=sys.stderr,
        )

    return {
        "shell": str(shell_path),
        "content": str(content_path),
        "out": str(out_path),
        "keep_shell_h1": keep_shell_h1,
        "appended_elements": appended,
        "image_drops": image_drops,
        "style_warnings": len(style_warnings),
        "final_paragraphs": final_paras,
        "first_h1_text": first_h1_text,
        "used_styles": used_styles,
    }


def main_body_replace() -> int:
    ap = argparse.ArgumentParser(
        description="Replace docx body keeping shell's styles + (optional) first H1.",
    )
    ap.add_argument("--shell", required=True, help="shell docx (styles/cover/H1 source)")
    ap.add_argument("--content", required=True, help="content docx (body source)")
    ap.add_argument("--out", required=True, help="output docx path")

    # Default: keep-shell-h1 ON (mutually exclusive)
    grp = ap.add_mutually_exclusive_group()
    grp.add_argument(
        "--keep-shell-h1", dest="keep_shell_h1", action="store_true",
        default=True,
        help="(default) preserve shell's first H1 paragraph; drop content's first H1",
    )
    grp.add_argument(
        "--no-keep-shell-h1", dest="keep_shell_h1", action="store_false",
        help="drop all shell body; take content from its first element",
    )
    ap.add_argument(
        "--dry-run", action="store_true",
        help="print plan only, don't write output",
    )
    args = ap.parse_args()

    shell = Path(args.shell).expanduser().resolve()
    content = Path(args.content).expanduser().resolve()
    out = Path(args.out).expanduser().resolve()

    if not shell.exists():
        print(f"ERROR: shell docx not found: {shell}", file=sys.stderr)
        return 2
    if not content.exists():
        print(f"ERROR: content docx not found: {content}", file=sys.stderr)
        return 2

    if args.dry_run:
        # Inspect both, print plan
        s_doc = Document(str(shell))
        s_children = list(s_doc.element.body)
        s_sect = _find_sect_idx(s_doc.element.body)
        s_h1 = _find_first_h1_index(s_children, s_sect, _build_style_id_to_name(s_doc))
        c_doc = Document(str(content))
        c_children = list(c_doc.element.body)
        c_sect = _find_sect_idx(c_doc.element.body)
        c_h1 = _find_first_h1_index(c_children, c_sect, _build_style_id_to_name(c_doc))
        print(f"# DRY RUN — body-replace plan")
        print(f"  shell:   {shell}")
        print(f"    body children: {s_sect}, first H1 idx: {s_h1}")
        if s_h1 >= 0:
            print(f"    shell H1 text: {_paragraph_text(s_children[s_h1])!r}")
        print(f"  content: {content}")
        print(f"    body children: {c_sect}, first H1 idx: {c_h1}")
        if c_h1 >= 0:
            print(f"    content H1 text: {_paragraph_text(c_children[c_h1])!r}")
        print(f"  out:     {out}")
        print(f"  keep_shell_h1: {args.keep_shell_h1}")
        if args.keep_shell_h1:
            keep_until = (s_h1 + 1) if s_h1 >= 0 else 0
            c_start = (c_h1 + 1) if c_h1 >= 0 else 0
        else:
            keep_until = 0
            c_start = 0
        print(f"  → keep shell children [0:{keep_until}) + content children [{c_start}:{c_sect})")
        return 0

    stats = body_replace(shell, content, out, keep_shell_h1=args.keep_shell_h1)
    print(f"OK: wrote {out}")
    print(f"  appended {stats['appended_elements']} elements from content")
    print(f"  final paragraphs: {stats['final_paragraphs']}")
    print(f"  first H1: {stats['first_h1_text']!r}")
    print(f"  styles used: {stats['used_styles']}")
    if stats["image_drops"]:
        print(f"  WARN: {stats['image_drops']} images may render broken (see stderr)")
    if stats["style_warnings"]:
        print(f"  WARN: {stats['style_warnings']} style fallbacks (see stderr)")
    return 0


# ──────────────────────────── 家族入口（子命令分发）────────────────────────────

SUBCOMMANDS = {
    "by-h1": main_by_h1,
    "body-replace": main_body_replace,
}


def main(argv: list[str] | None = None) -> int:
    return _cc.family_main(SUBCOMMANDS, argv, file=__file__)


if __name__ == "__main__":
    sys.exit(main())
