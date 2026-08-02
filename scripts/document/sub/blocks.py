#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""blocks.py — 标题块结构家族三合一（2026-07-31 家族折叠）

子命令 ↔ 原脚本（函数体逐字搬移；模块级 apply/main 改名 _<sub> 后缀）：

    fix-heading-disorder  ← fix_heading_disorder.py（legacy deprecated 动词，原名保留）
    reorder               ← reorder_heading_blocks.py（原 from fix_heading_disorder
                            import 的 detect_heading_form/is_heading_style 变同文件内引用）
    relocate              ← relocate_orphan_blocks.py

⚠ fix-heading-disorder 与 reorder 依赖 apply_body_styles（qual-supply 仓的脚本，
本仓没有）——旧独立脚本 import 期即 ModuleNotFoundError。折叠后改为这两个子命令的
入口首行引爆（同样的异常、同样的 rc），不连坐同文件里本来能跑的 relocate。

各子命令 CLI 与原独立脚本逐字一致：python3 sub/blocks.py <sub> <docx> …。
退役原件在 ~/.Trash/consolidation-20260731/blocks/（含 MANIFEST.md）。
"""
from __future__ import annotations

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[3] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py
from cn_number import cn_to_int  # noqa: E402,F401  中文数字 SSOT
import caption_re  # noqa: E402  题注判据 SSOT
_sys.path.append(str(_Path(__file__).resolve().parent))
import _cli_common as _cc  # noqa: E402  备份路径 SSOT（原来伸手找外仓的 make_backup_path）

import argparse  # noqa: E402
import json  # noqa: E402
import shutil  # noqa: E402
import re  # noqa: E402
import sys  # noqa: E402
from collections import Counter, defaultdict  # noqa: E402
from pathlib import Path  # noqa: E402
from typing import Optional  # noqa: E402

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))


def _ui_name(doc, style_id: str) -> str:
    """styleId → 该文档里的 UI 样式名；查不到就原样返回 styleId。

    这是 2026-08-02 之前那个 `apply_body_styles.STYLE_NAMES` 静态字典的真实语义 ——
    而映射本来就在文档自己身上，不需要外仓给一张表（原实现的 fallback
    `doc.styles.get_by_id(target, 1)` 干的就是同一件事）。

    ⚠ 不能直接信 `get_by_id` 的返回：**它查不到时静默回落到默认段落样式**，
    于是 `_ui_name(doc, "zdwp1")` 在一份没有 zdwp 样式的文档里返回 `"Normal"`，
    报告就会写成「建议改成 Normal」（2026-08-02 实测，本函数第一版就是这么错的）。
    所以必须回查 styleId 是否真的就是要的那个。
    """
    try:
        st = doc.styles.get_by_id(style_id, 1)
        if st is None or st.style_id != style_id:      # 回落了 → 文档里没有这个样式
            return style_id
        return getattr(st, "name", None) or style_id
    except Exception:                       # noqa: BLE001 — 查不到样式不是错误
        return style_id


# ══════════ fix-heading-disorder ← fix_heading_disorder.py ══════════

# ---------- 中文数字 -> 阿拉伯 ----------
# 本文件原有一份「4 条硬编码形态分支、不支持百千」的局部实现，2026-08-01 下沉到
# lib/cn_number.py（见文件顶部 import）。行为变化：「第一百零五章」这类过去返 None、
# 被 numbering 连贯性检查整段跳过，现在返 105 参与序列比对。

# ---------- heading 形态正则 ----------

# 2026-08-01 判据下沉 lib/caption_re.TABLE_NAME_HEURISTIC（与 styles.RE_TABLE_NAME
# 原是两份手抄，字符类差一个 en dash、中文数字集差一个「百」）。行为变化：章号
# 现在允许小数点 —— `表 3.1-2 成本对照表` 旧实现 detect_heading_form 返 None
# （落到后面的 heading 候选分支按正文处理），现在返 ("zdwp_table", 0, None)。
RE_TABLE      = caption_re.pattern(caption_re.TABLE_NAME_HEURISTIC)
RE_TITLE_CN   = re.compile(r"^([一二三四五六七八九十]+)、")
RE_CHAPTER    = re.compile(r"^第([一二三四五六七八九十百零\d]+)章\s*\S")
RE_H4         = re.compile(r"^(\d+)\.(\d+)\.(\d+)\.(\d+)\s+\S")
RE_H3         = re.compile(r"^(\d+)\.(\d+)\.(\d+)\s+\S")
RE_H2         = re.compile(r"^(\d+)\.(\d+)\s+\S")
RE_H1_DOT     = re.compile(r"^(\d+)\.\s+\S")
RE_H1         = re.compile(r"^(\d+)\s+\S")


def detect_heading_form(text: str) -> tuple[str, int, tuple[int, ...] | None] | None:
    """返回 (form_label, depth, numbers_tuple) 或 None.

    depth: 0 = 表名 / 1-4 = heading 深度.
    """
    t = (text or "").strip()
    if not t:
        return None
    if RE_TABLE.match(t):
        return ("zdwp_table", 0, None)
    m = RE_CHAPTER.match(t)
    if m:
        n = cn_to_int(m.group(1))
        return ("chapter_H1", 1, (n,) if n is not None else None)
    m = RE_H4.match(t)
    if m:
        return ("H4_form", 4, tuple(int(x) for x in m.groups()))
    m = RE_H3.match(t)
    if m:
        return ("H3_form", 3, tuple(int(x) for x in m.groups()))
    m = RE_H2.match(t)
    if m:
        return ("H2_form", 2, tuple(int(x) for x in m.groups()))
    m = RE_H1_DOT.match(t)
    if m:
        return ("H1_dot_form", 1, (int(m.group(1)),))
    m = RE_H1.match(t)
    if m and "." not in t.split()[0]:
        return ("H1_form", 1, (int(m.group(1)),))
    m = RE_TITLE_CN.match(t)
    if m:
        n = cn_to_int(m.group(1))
        return ("Title_cn", 1, (n,) if n is not None else None)
    return None


# ---------- style 分类 ----------

# styleId -> heading level (1-5); None = 非 heading
HEADING_STYLE_ID_TO_LEVEL = {
    "1": 1, "2": 2, "3": 3, "4": 4, "5": 5,
    "Heading1": 1, "Heading2": 2, "Heading3": 3, "Heading4": 4, "Heading5": 5,
}

# 视为"非 heading"的 styleId / styleName — 这些段如果文本形态像 heading 都算 A
NON_HEADING_STYLE_IDS = {"ZDWP", "Normal", "a", "zdwp1"}  # zdwp1=zdwp表名
NON_HEADING_STYLE_NAME_KEYWORDS = ("正文", "Normal", "Body", "zdwp表名")


def is_heading_style(style_id: str | None, style_name: str | None) -> int | None:
    """返回 heading 级别 1-5; 非 heading 返回 None.

    Title / zdwp表名 视为非 heading (depth 概念不同, 让 A 类判断).
    """
    if style_id and style_id in HEADING_STYLE_ID_TO_LEVEL:
        return HEADING_STYLE_ID_TO_LEVEL[style_id]
    if style_name:
        m = re.match(r"^Heading\s*(\d+)", style_name, re.I)
        if m:
            lvl = int(m.group(1))
            if 1 <= lvl <= 5:
                return lvl
        m = re.match(r"^标题\s*(\d+)", style_name)
        if m:
            lvl = int(m.group(1))
            if 1 <= lvl <= 5:
                return lvl
    return None


def is_non_heading_style(style_id: str | None, style_name: str | None) -> bool:
    """段是否是 '非 heading' (即 A 类的扫描范围).

    Heading X / Title 跳过; ZDWP正文 / Normal / zdwp表名 / 任何含"正文"name 算 yes.
    """
    if is_heading_style(style_id, style_name) is not None:
        return False
    # Title 也不扫 A (它 depth=0 + W2 ground truth 显示 ZDWP正文 误写 "三、" 才报 A)
    # 但 Title style 本身 (styleId=a4) 不在 A 范围 — 已经是合法标题样式
    if style_id == "a4":
        return False
    if style_id and style_id in NON_HEADING_STYLE_IDS:
        return True
    if style_name:
        for kw in NON_HEADING_STYLE_NAME_KEYWORDS:
            if kw in style_name:
                return True
    return False


# 形态 label -> suggested styleId
FORM_TO_STYLE_ID = {
    "zdwp_table":   "zdwp1",
    "Title_cn":     "a4",
    "chapter_H1":   "1",
    "H1_form":      "1",
    "H1_dot_form":  "1",
    "H2_form":      "2",
    "H3_form":      "3",
    "H4_form":      "3",  # docx 只有 3 个 Heading 级别可用 (1/2/3 + 4 在 styles 但默认未用)
}

DEPTH_TO_STYLE_ID = {1: "1", 2: "2", 3: "3", 4: "3", 5: "3"}


# ---------- detection ----------

def detect_anomalies(doc) -> tuple[list[dict], list[str]]:
    """返回 (anomalies, build_structural_warnings)."""
    anomalies: list[dict] = []
    warnings: list[str] = []

    # 收集所有段
    paras: list[dict] = []
    for idx, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        sid = p.style.style_id if p.style is not None else None
        nm = p.style.name if p.style is not None else None
        paras.append({
            "idx": idx, "text": text, "style_id": sid, "style_name": nm, "p": p,
        })

    # ---- A: false_promotion (扩检到任何非 Heading 段) ----
    for pi in paras:
        t = pi["text"]
        if not t or len(t) > 200:
            continue
        sid, nm = pi["style_id"], pi["style_name"]
        if not is_non_heading_style(sid, nm):
            continue
        form = detect_heading_form(t)
        if not form:
            continue
        label, depth, _tup = form
        suggested = FORM_TO_STYLE_ID.get(label)
        if suggested is None:
            continue
        # 跳过: 当前已是 suggested
        if sid == suggested:
            continue
        # 跳过特殊: zdwp表名 段且文本本来就是表名形态 (label=zdwp_table) — 已经对了
        if sid == "zdwp1" and label == "zdwp_table":
            continue
        suggested_ui = _ui_name(doc, suggested)
        anomalies.append({
            "idx": pi["idx"],
            "category": "A_false_promotion",
            "current_style": nm or sid,
            "current_style_id": sid,
            "suggested_style": suggested_ui,
            "suggested_style_id": suggested,
            "text": t[:120],
            "reason": f"文本符合 {label}(depth={depth}) 形态但 style={nm or sid}",
        })

    # ---- B: false_demotion (Heading 段但不像 heading) ----
    SHORT_TITLE_WORDS = {"前言", "摘要", "引言", "结语", "结论", "参考文献",
                         "附录", "目录", "致谢"}
    for pi in paras:
        t = pi["text"]
        if not t:
            continue
        sid, nm = pi["style_id"], pi["style_name"]
        cur_level = is_heading_style(sid, nm)
        if cur_level is None:
            continue
        too_long = len(t) > 120
        form = detect_heading_form(t)
        is_short_title = (
            t in SHORT_TITLE_WORDS or
            any(t.startswith(w) for w in SHORT_TITLE_WORDS)
        )
        if too_long or (form is None and not is_short_title and len(t) > 30):
            anomalies.append({
                "idx": pi["idx"],
                "category": "B_false_demotion",
                "current_style": nm or sid,
                "current_style_id": sid,
                "suggested_style": "ZDWP正文",
                "suggested_style_id": "ZDWP",
                "text": t[:120],
                "reason": (
                    f"Heading {cur_level} 但文本"
                    + ("过长(>120 字)" if too_long else "")
                    + ("," if (too_long and form is None) else "")
                    + ("不符合 heading 形态" if form is None else "")
                ),
            })

    # ---- C / D: numbering_backward / numbering_skip ----
    # 收集**所有**符合 heading form 的段 (含 style=正文 但形态像 heading 的伪 heading)
    heading_seq: list[dict] = []
    for pi in paras:
        t = pi["text"]
        form = detect_heading_form(t)
        if not form:
            continue
        label, depth, tup = form
        if tup is None or depth < 1:
            continue
        heading_seq.append({
            "idx": pi["idx"], "depth": depth, "tup": tup,
            "style_id": pi["style_id"], "style_name": pi["style_name"],
            "text": t[:120], "form": label,
        })

    # 按 (depth, prefix) 跟踪 last_num — 同一前缀下出现回退/跳号都报
    # (跨父分支独立,不互相影响; 与 W2 ground truth 算法一致)
    last_by_prefix: dict[tuple[int, tuple], dict] = {}
    for h in heading_seq:
        depth = h["depth"]
        tup = h["tup"]
        if len(tup) < depth:
            continue
        cur_num = tup[depth - 1]
        cur_prefix = tup[: depth - 1]
        key = (depth, cur_prefix)
        if key in last_by_prefix:
            prev = last_by_prefix[key]
            if cur_num < prev["num"]:
                anomalies.append({
                    "idx": h["idx"],
                    "category": "C_numbering_backward",
                    "current_style": h["style_name"] or h["style_id"],
                    "current_style_id": h["style_id"],
                    "suggested_style": h["style_name"] or h["style_id"],
                    "suggested_style_id": h["style_id"],
                    "text": h["text"],
                    "reason": (
                        f"同级编号倒退: depth={depth} prefix={cur_prefix} "
                        f"前一段 idx={prev['idx']} 编号={prev['num']}, "
                        f"本段编号={cur_num}"
                    ),
                    "_prev_idx": prev["idx"],
                })
            elif cur_num > prev["num"] + 1:
                anomalies.append({
                    "idx": h["idx"],
                    "category": "D_numbering_skip",
                    "current_style": h["style_name"] or h["style_id"],
                    "current_style_id": h["style_id"],
                    "suggested_style": h["style_name"] or h["style_id"],
                    "suggested_style_id": h["style_id"],
                    "text": h["text"],
                    "reason": (
                        f"同级编号跳号: depth={depth} prefix={cur_prefix} "
                        f"前一段编号={prev['num']}, 本段编号={cur_num} "
                        f"(期望={prev['num'] + 1})"
                    ),
                    "_prev_idx": prev["idx"],
                })
        last_by_prefix[key] = {"idx": h["idx"], "num": cur_num}

    # ---- E: level_mismatch (styled heading 但编号深度与 style 级别不一致) ----
    for h in heading_seq:
        text_depth = h["depth"]
        cur_level = is_heading_style(h["style_id"], h["style_name"])
        if cur_level is None:
            continue
        if cur_level != text_depth and text_depth in DEPTH_TO_STYLE_ID:
            target = DEPTH_TO_STYLE_ID[text_depth]
            anomalies.append({
                "idx": h["idx"],
                "category": "E_level_mismatch",
                "current_style": h["style_name"] or h["style_id"],
                "current_style_id": h["style_id"],
                "suggested_style": _ui_name(doc, target),
                "suggested_style_id": target,
                "text": h["text"],
                "reason": (
                    f"编号深度={text_depth} (form={h['form']}) 与 "
                    f"style 层级={cur_level} 不一致 → 应为 Heading {text_depth}"
                ),
            })

    # ---- F: duplicate_adjacent (相邻段 text 重复) ----
    # 相邻段 (idx 差 ≤ 7) 且 (text 100% 相同 OR 字符重叠 > 90%)
    F_WINDOW = 7
    for i, a in enumerate(paras):
        ta = a["text"]
        if not ta or len(ta) < 6:
            continue
        for j in range(i + 1, min(i + F_WINDOW + 1, len(paras))):
            b = paras[j]
            tb = b["text"]
            if not tb or len(tb) < 6:
                continue
            idx_diff = b["idx"] - a["idx"]
            if idx_diff > F_WINDOW:
                break
            # 完全相同?
            if ta == tb:
                overlap = 1.0
                exact = True
            else:
                # 字符集重叠
                sa, sb = set(ta), set(tb)
                if not sa or not sb:
                    continue
                overlap = len(sa & sb) / max(len(sa), len(sb))
                exact = False
                if overlap < 0.9:
                    continue
            # 不能两段都是同样 style 的"正常段"(误报)
            sid_a, sid_b = a["style_id"], b["style_id"]
            # 只在 style 不同 时才报 (重复证据 = 一份高一份低)
            if sid_a == sid_b:
                continue
            # 判断是否一份 Heading 一份非 Heading (W2 描述: build 阶段重复插入)
            a_is_heading = is_heading_style(sid_a, a["style_name"]) is not None
            b_is_heading = is_heading_style(sid_b, b["style_name"]) is not None
            if a_is_heading == b_is_heading:
                # 都 heading 或都非 heading — 不算 F 类
                continue
            heading_side = "a" if a_is_heading else "b"
            non_heading = b if a_is_heading else a
            heading = a if a_is_heading else b
            anomalies.append({
                "idx": non_heading["idx"],
                "category": "F_duplicate_adjacent",
                "current_style": non_heading["style_name"] or non_heading["style_id"],
                "current_style_id": non_heading["style_id"],
                "suggested_style": "DELETE",
                "suggested_style_id": "DELETE",
                "text": non_heading["text"][:120],
                "reason": (
                    f"与 idx={heading['idx']} (style={heading['style_name']}) "
                    f"相邻重复 (idx 差={idx_diff}, "
                    f"{'完全相同' if exact else f'字符重叠={overlap:.0%}'}) — "
                    f"build 阶段重复插入证据"
                ),
                "_heading_idx": heading["idx"],
                "_heading_style_id": heading["style_id"],
                "_idx_diff": idx_diff,
                "_exact": exact,
                "_overlap": round(overlap, 2),
            })

    # 去重 F: 同一段被多次报 → 保留最近的 (idx_diff 最小的)
    f_by_idx: dict[int, dict] = {}
    other_anomalies = []
    for a in anomalies:
        if a["category"] == "F_duplicate_adjacent":
            cur = f_by_idx.get(a["idx"])
            if cur is None or a["_idx_diff"] < cur["_idx_diff"]:
                f_by_idx[a["idx"]] = a
        else:
            other_anomalies.append(a)
    anomalies = other_anomalies + list(f_by_idx.values())

    # ---- G: build structural warnings (软警告) ----
    # 整章 H1 重复 (相同前缀的 H1 出现 ≥2 次)
    h1_seen: dict[str, list[int]] = defaultdict(list)
    for h in heading_seq:
        if h["depth"] == 1 and h["tup"]:
            h1_seen[str(h["tup"][0])].append(h["idx"])
    for num, idxs in h1_seen.items():
        if len(idxs) >= 2:
            warnings.append(
                f"H1 编号 '{num}' 在 idx {idxs} 重复出现 — "
                f"很可能是 build 阶段两章合并错位,建议重 build 或手动合章"
            )
    # Title 段被打成 ZDWP正文
    for pi in paras:
        if pi["text"] and RE_TITLE_CN.match(pi["text"]) and pi["style_id"] == "ZDWP":
            warnings.append(
                f"idx {pi['idx']} 文本 {pi['text'][:60]!r} 形如「X、章节标题」 "
                f"但 style=ZDWP正文 — 应为 Title (a4); "
                f"可能是 build 阶段 Title 段漏改"
            )

    # 按 idx + category 排序
    anomalies.sort(key=lambda a: (a["idx"], a["category"]))
    return anomalies, warnings


# ---------- fix ----------

AUTO_FIX_CATEGORIES = {"A_false_promotion", "E_level_mismatch", "F_duplicate_adjacent"}


def apply_fixes(doc, anomalies: list[dict], dry_run: bool) -> None:
    """对 A / E / F 自动 fix; 修改 anomaly['auto_fixed'] in-place.

    F 安全规则: 仅当 idx_diff ≤ 2 + exact=True + 一份 Heading 一份非 Heading
    才删非 Heading 那份; 否则 manual_review.
    """
    available_ids = {s.style_id for s in doc.styles}
    paragraphs = doc.paragraphs
    deleted_idxs: set[int] = set()  # 已删的 idx,后续 fix 跳过

    # F 优先处理 (idx 倒序,避免删除影响后续 idx)
    f_anoms = [a for a in anomalies if a["category"] == "F_duplicate_adjacent"]
    f_anoms.sort(key=lambda a: -a["idx"])
    for a in f_anoms:
        # 安全门槛
        safe = (
            a.get("_exact") and
            a.get("_idx_diff", 99) <= 2 and
            a.get("current_style_id") == "ZDWP"  # 删的一定是 ZDWP正文 那份
        )
        if not safe:
            a["auto_fixed"] = False
            a["manual_review"] = True
            continue
        if dry_run:
            a["auto_fixed"] = False
            continue
        idx = a["idx"]
        if idx >= len(paragraphs):
            a["auto_fixed"] = False
            continue
        p = paragraphs[idx]
        try:
            # python-docx 删段: p._element.getparent().remove(p._element)
            p._element.getparent().remove(p._element)
            deleted_idxs.add(idx)
            a["auto_fixed"] = True
        except Exception as exc:
            a["auto_fixed"] = False
            a["reason"] = a["reason"] + f" [F 删除失败: {exc}]"

    # 注意: 删段后 paragraphs 列表的 idx 失效, 但我们用 doc.paragraphs[idx]
    # 索引访问 — 由于上面用倒序 F 处理,A/E fix 时 idx 仍然指向原位置 (前面段没动)
    # 但删除发生在 F 段后,所有 idx > 已删 idx 的段需要做映射. 我们用稳妥做法:
    # F 处理完后重新拿 paragraphs 列表,A/E 用文本匹配 + idx 起点扫描定位.
    # 简化: F 删除一定在 A/E 之后处理 — 颠倒顺序.

    # 实际策略: 先做 A/E (idx 仍稳定),再做 F (倒序删).
    # 上面 F 已经做了 — 撤回. 重做:

    # 由于上面已 F 做了删除,下面 A/E 需要用 element 比较找到原 paragraph.
    # 为简单起见,改架构: F 移到最后做. 此处重置已删段,撤销操作不现实,故改用:
    # 不做删除,只改 F 段 text 为空 (这样不破坏 idx 索引). — 不行,留空段也是污染.
    # 正确做法: F 真删 + A/E 用 element 对齐 fix.

    # 重新设计: 通过对每个 A/E anomaly 找 doc 现存 paragraphs 中 idx 偏移后的位置.
    # 已删段数 (idx < this_idx) = offset, this_paragraph = doc.paragraphs[this_idx - offset]
    deleted_sorted = sorted(deleted_idxs)

    def remap_idx(orig_idx: int) -> int:
        # 计算 orig_idx 之前删了多少段
        offset = sum(1 for d in deleted_sorted if d < orig_idx)
        return orig_idx - offset

    paragraphs = doc.paragraphs  # 重新取 (F 删除后)

    # 收集所有"未被安全删除"的 F 段 idx (这些段是另一段 Heading 的重复证据,
    # 即使形态像 heading,也不该被 A 升级 — 否则会造出两段一模一样的 Heading)
    f_blocked_idxs = {
        a["idx"] for a in anomalies
        if a["category"] == "F_duplicate_adjacent"
        and not a.get("auto_fixed")
    }

    for a in anomalies:
        if a["category"] not in ("A_false_promotion", "E_level_mismatch"):
            if "auto_fixed" not in a:
                a["auto_fixed"] = False
            continue
        # F-blocked: 此段是另一段 Heading 的重复证据,不能 A 升级
        if a["idx"] in f_blocked_idxs:
            a["auto_fixed"] = False
            a["manual_review"] = True
            a["reason"] = a["reason"] + " [F-blocked: 此段是另一段 Heading 的重复证据,需用户决策]"
            continue
        target = a["suggested_style_id"]
        if target not in available_ids:
            a["auto_fixed"] = False
            a["reason"] = a["reason"] + f" [跳过: 目标 styleId {target!r} 不在 styles.xml]"
            continue
        if dry_run:
            a["auto_fixed"] = False
            continue
        orig_idx = a["idx"]
        if orig_idx in deleted_idxs:
            a["auto_fixed"] = False
            a["reason"] = a["reason"] + " [跳过: 段已被 F 类删除]"
            continue
        new_idx = remap_idx(orig_idx)
        if new_idx >= len(paragraphs):
            a["auto_fixed"] = False
            continue
        p = paragraphs[new_idx]
        # 校验: 文本匹配
        if (p.text or "").strip()[:120] != a["text"]:
            # 偏移失败,跳过
            a["auto_fixed"] = False
            a["reason"] = a["reason"] + (
                f" [跳过: idx 偏移后文本不匹配 — got "
                f"{(p.text or '').strip()[:40]!r}]"
            )
            continue
        try:
            # 原实现先按 UI 名查、查不到再按 id 查；而 UI 名本来就是从 id 派生的，
            # 那一层是循环，直接按 id 取即可（取不到落到下面的 pStyle 兜底分支）。
            p.style = doc.styles.get_by_id(target, 1)
            a["auto_fixed"] = True
        except Exception:
            try:
                pPr = p._p.get_or_add_pPr()
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is None:
                    from docx.oxml import OxmlElement
                    pStyle = OxmlElement("w:pStyle")
                    pPr.insert(0, pStyle)
                pStyle.set(qn("w:val"), target)
                a["auto_fixed"] = True
            except Exception as exc:
                a["auto_fixed"] = False
                a["reason"] = a["reason"] + f" [fix 失败: {exc}]"


# ---------- 汇总 ----------

CATEGORY_ORDER = [
    "A_false_promotion",
    "B_false_demotion",
    "C_numbering_backward",
    "D_numbering_skip",
    "E_level_mismatch",
    "F_duplicate_adjacent",
]


def summarize(anomalies: list[dict]) -> dict:
    summary: dict[str, dict] = {}
    for cat in CATEGORY_ORDER:
        items = [a for a in anomalies if a["category"] == cat]
        entry: dict[str, int] = {"count": len(items)}
        if cat in AUTO_FIX_CATEGORIES:
            entry["auto_fixed"] = sum(1 for a in items if a.get("auto_fixed"))
            entry["manual_review"] = sum(1 for a in items if a.get("manual_review"))
        else:
            entry["manual_review"] = len(items)
        summary[cat] = entry
    return summary


# ---------- main ----------

def main_fix_heading_disorder(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="qual-supply docx 标题失序 v2 (W2 ground truth 对齐)",
    )
    ap.add_argument("docx", help="目标 docx 路径 (inplace 修改, 除非 --dry-run)")
    ap.add_argument("--dry-run", action="store_true", help="只扫不改, 不备份不写")
    ap.add_argument("--no-backup", action="store_true", help="跳过备份")
    ap.add_argument("--report", help="anomaly 报告 JSON")
    args = ap.parse_args(argv)

    src = Path(args.docx).resolve()
    if not src.is_file():
        print(f"ERROR: 不存在 {src}", file=sys.stderr)
        return 2

    backup_path = None
    if not args.dry_run and not args.no_backup:
        backup_path = _cc.find_next_backup(src)
        shutil.copy2(src, backup_path)

    doc = Document(str(src))
    anomalies, warnings = detect_anomalies(doc)

    before = summarize(anomalies)
    apply_fixes(doc, anomalies, dry_run=args.dry_run)
    after = summarize(anomalies)  # 含 auto_fixed 字段

    if not args.dry_run:
        doc.save(str(src))

    # 控制台报告
    print(f"file:     {src}")
    if backup_path:
        print(f"backup:   {backup_path}")
    elif args.dry_run:
        print("backup:   (dry-run, 未备份)")
    else:
        print("backup:   (--no-backup)")
    print(f"dry_run:  {args.dry_run}")
    print(f"total anomalies: {len(anomalies)}")
    print("summary:")
    for cat in CATEGORY_ORDER:
        print(f"  {cat:25s} {after[cat]}")
    if warnings:
        print(f"build_structural_warnings: {len(warnings)}")
        for w in warnings:
            print(f"  - {w}")

    if args.report:
        report_path = Path(args.report).expanduser().resolve()
        report_path.parent.mkdir(parents=True, exist_ok=True)
        # manual_review_items: B/C/D/F-manual
        manual_review_items = [
            a for a in anomalies
            if not a.get("auto_fixed") and (
                a["category"] in ("B_false_demotion", "C_numbering_backward",
                                  "D_numbering_skip") or
                (a["category"] == "F_duplicate_adjacent" and a.get("manual_review"))
            )
        ]
        report = {
            "input": str(src),
            "dry_run": args.dry_run,
            "backup": str(backup_path) if backup_path else None,
            "before": {
                k: v["count"] for k, v in before.items()
            } | {"total": sum(v["count"] for v in before.values())},
            "fixed": {
                k: after[k].get("auto_fixed", 0)
                for k in ("A_false_promotion", "E_level_mismatch", "F_duplicate_adjacent")
            } | {
                "total_auto_fixed": sum(
                    after[k].get("auto_fixed", 0)
                    for k in ("A_false_promotion", "E_level_mismatch",
                              "F_duplicate_adjacent")
                ),
            },
            "remaining": {
                "B_false_demotion": after["B_false_demotion"]["count"],
                "C_numbering_backward": after["C_numbering_backward"]["count"],
                "D_numbering_skip": after["D_numbering_skip"]["count"],
                "A_unfixed": before["A_false_promotion"]["count"]
                             - after["A_false_promotion"].get("auto_fixed", 0),
                "E_unfixed": before["E_level_mismatch"]["count"]
                             - after["E_level_mismatch"].get("auto_fixed", 0),
                "F_manual": after["F_duplicate_adjacent"].get("manual_review", 0),
                "total": len(manual_review_items),
            },
            "summary_detailed": after,
            "manual_review_items": manual_review_items,
            "build_structural_warnings": warnings,
            "anomalies": anomalies,
        }
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2))
        print(f"report:   {report_path}")

    return 0


# ---------------- pipeline adapter ----------------
def apply_fix_heading_disorder(doc, args=None) -> dict:
    dry = bool(getattr(args, "dry_run", False)) if args else False
    anomalies, warnings = detect_anomalies(doc)
    apply_fixes(doc, anomalies, dry_run=dry)
    after = summarize(anomalies)
    auto_fixed = sum(
        after[k].get("auto_fixed", 0)
        for k in ("A_false_promotion", "E_level_mismatch", "F_duplicate_adjacent")
    )
    return {
        "changed": auto_fixed,
        "anomalies_total": len(anomalies),
        "summary": after,
        "warnings": warnings,
    }


# ══════════ reorder ← reorder_heading_blocks.py ══════════

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_P = f"{{{W_NS}}}p"


# ---------- 段分类 ----------

def classify_paragraph(p, idx: int) -> dict:
    """返回段信息 dict: {idx, text, style_id, style_name, h_level, h_number}.

    h_level: 1-5 = heading 层级 (styled OR form-only); None = 非 heading.
    h_number: tuple 编号 (如 (3,1,2)) 或 None.

    判定优先:
      - styled Heading X → level = X, number 用 text 形态解析(如能)
      - 否则 detect_heading_form: 若 form depth ∈ {1..5} 且像 heading → level = depth
        但要小心: 普通正文段以数字开头(如"1. xxx 是 ..." 列表)别误判 → 仅当文本 ≤ 50 字算"像 heading"
    """
    text = (p.text or "").strip()
    sid = p.style.style_id if p.style is not None else None
    nm = p.style.name if p.style is not None else None

    styled_level = is_heading_style(sid, nm)
    form = detect_heading_form(text) if text else None

    h_level = None
    h_number = None

    if styled_level is not None:
        h_level = styled_level
        if form and form[2] is not None:
            h_number = form[2]
    elif form is not None and form[1] in (1, 2, 3, 4, 5) and form[2] is not None:
        # form-only heading: 文本短且形态明确 → 认作 heading
        # 但纯 "1 xxx" 单数字 H1_form 可能是正文列表项 → 仅在很短(≤30 字)时认
        label, depth, _ = form
        if label in ("H2_form", "H3_form", "H4_form", "chapter_H1"):
            if len(text) <= 80:
                h_level = depth
                h_number = form[2]
        # H1_form / H1_dot_form / Title_cn 不当 heading 处理(避免列表项误判)

    return {
        "idx": idx,
        "text": text,
        "style_id": sid,
        "style_name": nm,
        "h_level": h_level,
        "h_number": h_number,
        "p": p,
        "styled": styled_level is not None,
    }


# ---------- 块切分 ----------

def slice_blocks(paras: list[dict]) -> list[dict]:
    """切段块. 每个块 = heading + 下属直到下一个同级或更高级 heading.

    返回 list[block], block 含:
      level, heading_idx, heading_text, number, end_idx (exclusive),
      paragraph_indices, styled, parent_block_idx (即包含本块的更高级块的 heading_idx;
      若为 H1 块或无更高级, 则 -1).
    """
    headings = [pi for pi in paras if pi["h_level"] is not None]
    blocks: list[dict] = []
    # 用栈追踪当前每层 heading 的 heading_idx, 以确定 parent
    # stack[level] = heading_idx of currently open block at that level
    open_stack: dict[int, int] = {}
    for i, h in enumerate(headings):
        # 当前 heading 出现 → 关闭所有 level >= h.level 的 open block
        for lvl in list(open_stack.keys()):
            if lvl >= h["h_level"]:
                del open_stack[lvl]
        # parent = max level < h.level still open
        parent_levels = [lvl for lvl in open_stack if lvl < h["h_level"]]
        parent_idx = open_stack[max(parent_levels)] if parent_levels else -1

        # 下一个同级或更高级别 heading = level <= h.level
        end_idx = len(paras)
        for nh in headings[i + 1:]:
            if nh["h_level"] <= h["h_level"]:
                end_idx = nh["idx"]
                break
        blocks.append({
            "level": h["h_level"],
            "heading_idx": h["idx"],
            "heading_text": h["text"],
            "number": h["h_number"],
            "end_idx": end_idx,
            "paragraph_indices": list(range(h["idx"], end_idx)),
            "styled": h["styled"],
            "style_id": h["style_id"],
            "parent_block_idx": parent_idx,
        })
        open_stack[h["h_level"]] = h["idx"]

    # Rehome pass — 让每个 H{N} 块尽量找到 number 匹配的 L{N-1} 父块.
    # 用于修两种常见反模式:
    #   (a) H3 出现在其 L2 父之前 (3.1.2 物理在 3.1 之前) — parent fall-through 到 H1
    #   (b) H3 跨章漂浮 (在错误的 L2 之下) — number prefix 与 parent number 不匹配
    by_idx = {b["heading_idx"]: b for b in blocks}
    for b in blocks:
        if b["number"] is None or len(b["number"]) <= 1:
            continue
        want_prefix = b["number"][:-1]
        want_level = len(want_prefix)
        # 候选 = number 完全匹配 want_prefix 且 level 严格 = want_level 的块
        candidates = [c for c in blocks
                      if c["number"] == want_prefix
                      and c["level"] == want_level]
        if not candidates:
            continue
        cur_parent = by_idx.get(b["parent_block_idx"])
        if cur_parent and cur_parent["heading_idx"] in [c["heading_idx"] for c in candidates]:
            continue
        # 同 H1 范围内的 candidates 优先
        # 若 cur_parent 没有 → 用所有 candidates
        cur_h1 = _find_h1_ancestor(b, by_idx)
        same_h1_cands = [c for c in candidates
                         if _find_h1_ancestor(c, by_idx) == cur_h1]
        pool = same_h1_cands if same_h1_cands else candidates
        # 在 pool 内选 heading_idx <= b.heading_idx 且最大的 (本块前最近); 否则第一个
        before = [c for c in pool if c["heading_idx"] < b["heading_idx"]]
        if before:
            target = max(before, key=lambda x: x["heading_idx"])
        else:
            target = min(pool, key=lambda x: x["heading_idx"])
        if not _same_h1(b, target, by_idx):
            continue
        b["parent_block_idx"] = target["heading_idx"]
    return blocks


_RE_NUM_PREFIX = None
def _strip_number_prefix(text: str) -> str:
    """剥离 heading 文本的编号前缀, 如 '3.1.2 办公室工作机制' → '办公室工作机制'."""
    global _RE_NUM_PREFIX
    if _RE_NUM_PREFIX is None:
        import re
        _RE_NUM_PREFIX = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s+")
    m = _RE_NUM_PREFIX.match(text or "")
    if m:
        return (text or "")[m.end():].strip()
    return (text or "").strip()


def _texts_similar(texts: list[str]) -> bool:
    """判定一组文本是否"实质相同". 严格 — 完全相同或一方是另一方前缀(>= 4 字符)."""
    if not texts:
        return True
    base = texts[0]
    for t in texts[1:]:
        if t == base:
            continue
        # 一方是另一方前缀且长度 >=4
        if len(t) >= 4 and len(base) >= 4:
            if t.startswith(base) or base.startswith(t):
                continue
        return False
    return True


def _find_h1_ancestor(b: dict, by_idx: dict) -> Optional[int]:
    """沿 parent_block_idx 链向上找 H1 块的 heading_idx; 找不到返回 None."""
    seen_path = set()
    cur = b
    while cur and cur["heading_idx"] not in seen_path:
        seen_path.add(cur["heading_idx"])
        if cur["level"] == 1:
            return cur["heading_idx"]
        pidx = cur["parent_block_idx"]
        if pidx < 0 or pidx not in by_idx:
            return None
        cur = by_idx[pidx]
    return None


def _same_h1(b1: dict, b2: dict, by_idx: dict) -> bool:
    h1a = _find_h1_ancestor(b1, by_idx)
    h1b = _find_h1_ancestor(b2, by_idx)
    if h1a is None or h1b is None:
        # 无 H1 上下文 → 视为同一(顶级)
        return h1a == h1b
    return h1a == h1b


def parent_prefix(number: Optional[tuple]) -> Optional[tuple]:
    """块的父块编号前缀. (3,1,2) -> (3,1); (3,1) -> (3,); (3,) -> ()."""
    if number is None:
        return None
    return number[:-1]


# ---------- 错位 / 重复检测 ----------

def detect_issues(blocks: list[dict]) -> tuple[list[dict], list[dict]]:
    """返回 (misordered_groups, duplicate_groups).

    misordered_group: 同父下兄弟块按文档顺序 vs number 排序顺序不一致.
    duplicate_group: 同父下 number 完全相同的多个块.
    """
    # 按 (level, parent_block_idx) 分组兄弟块 — parent_block_idx 是文档位置上
    # 的真实父块, 不是 number prefix. 这样跨章/重排过的章节里的同号 heading 不会被
    # 误归为兄弟.
    # **不处理 H1 (level=1)** — 顶级章节由多份合并文档而来, 同号 H1 (如多个 "2 xxx") 是
    # 不同章节的合法重复, 不能当 duplicate 删. H1 顺序按文档顺序保留.
    groups: dict[tuple, list[dict]] = {}
    for b in blocks:
        if b["number"] is None:
            continue
        if b["level"] == 1:
            continue
        key = (b["level"], b["parent_block_idx"])
        groups.setdefault(key, []).append(b)

    # 找每个 block 的父块, 取出父块 number prefix 用于"合法兄弟"过滤
    block_by_idx = {b["heading_idx"]: b for b in blocks}

    misordered = []
    duplicates = []
    for key, sibs in groups.items():
        level, parent_idx = key
        # 合法兄弟过滤: 只考虑 number 的 prefix 与父块 number 一致的块.
        # 比如父块=(4,) "4 水价管理机制" 下, 只算 (4,x) 的 H2 块为兄弟; (3,4)/(3,5) 跨章
        # 漂浮块不当兄弟 (它们结构上不该在这里, 但 number-sort 会越权把它们排进来).
        if parent_idx >= 0 and parent_idx in block_by_idx:
            parent_num = block_by_idx[parent_idx]["number"]
        else:
            parent_num = ()
        if parent_num is None:
            parent_num = ()

        valid_sibs = []
        invalid_sibs = []
        for b in sibs:
            if b["number"][:len(parent_num)] == parent_num:
                valid_sibs.append(b)
            else:
                invalid_sibs.append(b)

        if invalid_sibs:
            # 记录但不处理(留 issue 给后续人工/其他脚本)
            pass

        if not valid_sibs:
            continue

        # 文档顺序 (按 heading_idx)
        sibs_doc = sorted(valid_sibs, key=lambda x: x["heading_idx"])

        # 检重复 — 仅当 number 相同 AND heading 文本除编号外一致 → 真重复.
        # number 相同但文本不同 = number-collision (不同标题用同一编号), 不删.
        # 在同 number 内部按 stripped-text 聚类, 每簇若 >= 2 则记一组 duplicates.
        seen: dict[tuple, list[dict]] = {}
        for b in sibs_doc:
            seen.setdefault(b["number"], []).append(b)
        for num, group in seen.items():
            if len(group) <= 1:
                continue
            # 按 stripped text 聚类
            clusters: dict[str, list[dict]] = {}
            for b in group:
                k = _strip_number_prefix(b["heading_text"])
                # 简化: 用前 20 字作 key
                key_text = k[:20]
                clusters.setdefault(key_text, []).append(b)
            for ktext, cluster in clusters.items():
                if len(cluster) >= 2:
                    duplicates.append({
                        "key": key,
                        "number": num,
                        "blocks": cluster,
                    })

        # 检错位 — 从 sibs 中剔除"真 duplicate 的非 keeper" (那些块会被删, 不参与排序),
        # 但 number-collision 的两块都保留参与排序.
        to_delete_idx: set[int] = set()
        for num, group in seen.items():
            if len(group) <= 1:
                continue
            clusters: dict[str, list[dict]] = {}
            for b in group:
                key_text = _strip_number_prefix(b["heading_text"])[:20]
                clusters.setdefault(key_text, []).append(b)
            for ktext, cluster in clusters.items():
                if len(cluster) >= 2:
                    keeper = pick_keeper(cluster)
                    for b in cluster:
                        if b is not keeper:
                            to_delete_idx.add(b["heading_idx"])
        survivors = [b for b in valid_sibs if b["heading_idx"] not in to_delete_idx]
        # 排序键: 主键 number, 副键 heading_idx 保稳定
        kept_doc = sorted(survivors, key=lambda x: x["heading_idx"])
        kept_sorted = sorted(survivors, key=lambda x: (x["number"], x["heading_idx"]))
        if [b["heading_idx"] for b in kept_doc] != [b["heading_idx"] for b in kept_sorted]:
            misordered.append({
                "key": key,
                "before": [b["number"] for b in kept_doc],
                "after": [b["number"] for b in kept_sorted],
                "kept_doc": kept_doc,
                "kept_sorted": kept_sorted,
            })

    return misordered, duplicates


def pick_keeper(group: list[dict]) -> dict:
    """从同 number 的重复块中挑保留者.

    优先: ① styled (Heading X 样式) ② 块内段数多 ③ 后出现 (heading_idx 大)
    """
    def score(b):
        styled_score = 1 if b["styled"] else 0
        size_score = len(b["paragraph_indices"])
        idx_score = b["heading_idx"]  # 后出现优先
        return (styled_score, size_score, idx_score)

    return max(group, key=score)


# ---------- 重排 / 删除 (lxml) ----------

def apply_changes(doc, blocks: list[dict], misordered: list[dict],
                  duplicates: list[dict]) -> tuple[int, int]:
    """对 docx body 应用重排+删除.

    返回 (moves_applied, deletions_applied).

    策略:
      1. 先标记所有要删除的段 idx (重复块的非 keeper)
      2. 对每个 misordered group, 按 number 排序后, 找出兄弟块的 XML 片段集合,
         在 body 中 detach 再按新顺序 insert
      3. 兄弟块范围 = [heading_idx, end_idx); 跨章节边界检查: 所有兄弟在同 H1 内

    实现: 全用 lxml `<w:body>` 上操作 `<w:p>`.
    """
    body = doc.element.body
    # 收集 body 下所有顶层 <w:p> (注意 body 还含 <w:sectPr> 等其他元素)
    # paragraphs idx 对齐 doc.paragraphs (跳过 sectPr/tbl)
    # doc.paragraphs 只含 body 下 <w:p>, idx 与 body 中 <w:p> 顺序一致.
    all_p_elements = list(body.iter(W_P))
    # 但 doc.paragraphs 可能含 table 内的 p — 用 doc.paragraphs[i]._element 更安全
    para_elems = [p._element for p in doc.paragraphs]

    # ---- step 1: 删除重复 ----
    delete_idx_set: set[int] = set()
    delete_records = []
    for dup in duplicates:
        keeper = pick_keeper(dup["blocks"])
        for b in dup["blocks"]:
            if b is keeper:
                continue
            for idx in b["paragraph_indices"]:
                delete_idx_set.add(idx)
            delete_records.append({
                "number": list(dup["number"]),
                "kept_idx": keeper["heading_idx"],
                "deleted_idx": b["heading_idx"],
                "reason": ("kept styled" if keeper["styled"] and not b["styled"]
                           else "kept larger/later block"),
            })

    # ---- step 2: 重排兄弟块 ----
    # 关键: 兄弟块的段范围在 docx 中可能 交错(因为错位本身就是源 idx 乱). 我们要按 "新顺序" 把每个 keeper 块的 <w:p> 连续起来, 放在原最早兄弟的位置.
    # 实现:
    #   a) 收集本 group 所有 keeper 的段 idx 集合 union (排除 delete_idx_set 里的)
    #   b) 这些段在 body 中 detach
    #   c) 按 keeper 的 number 排序顺序, 依次把每个 keeper 块的段 XML 列表合并
    #   d) 在原最早段位置之前 insert 进 body
    moves_applied = 0

    # 同一个 idx 不能被两个 group 重排 → 但 group 按 (level, parent_prefix) 分,
    # 不同 group 的段不相交. 安全.
    # 先处理: 收集要重排的段 idx 集合, 这些段不能在 delete 阶段被 remove (会丢失).
    # 顺序: 先处理删除(纯 remove), 再处理重排(detach + insert).
    # 但删除涉及的段如果在某 group keeper 块内 — 不会, keeper 不会被标记为删.
    # 非 keeper 块的段 idx 在 delete_idx_set 中, 重排时这些块也不在 keeper 列表 → ok.

    # 但还有一种情况: 非 keeper 块的段, 如果它在某 misordered group 中作为 "before" 出现 → 因 pick_keeper 已被剔除, 不在 kept_sorted 中. ok.

    # 先执行 删除
    deletions_applied = 0
    for idx in sorted(delete_idx_set, reverse=True):
        elem = para_elems[idx]
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            deletions_applied += 1

    # 重新刷新 para_elems? — 用旧引用, 因为 lxml elem 即使 detach 仍可用,
    # 而我们只对剩余 keeper 的段 idx 操作, 它们的 elem 引用仍有效.

    # 跨章检查: 仅在同一 H1 (chapter) 范围内允许重排.
    # 找出每个 H1 块的 idx 范围
    h1_ranges = []
    for b in blocks:
        if b["level"] == 1:
            h1_ranges.append((b["heading_idx"], b["end_idx"]))
    if not h1_ranges:
        # 无 H1 — 整 doc 一个范围
        h1_ranges = [(0, len(para_elems))]

    def same_chapter(idx_list: list[int]) -> bool:
        if not idx_list:
            return True
        for lo, hi in h1_ranges:
            if all(lo <= i < hi for i in idx_list):
                return True
        return False

    for mis in misordered:
        kept_sorted = mis["kept_sorted"]
        kept_doc = mis["kept_doc"]
        if [b["number"] for b in kept_doc] == [b["number"] for b in kept_sorted]:
            continue  # 已经对了

        # 跨章检查
        heading_idxs = [b["heading_idx"] for b in kept_sorted]
        if not same_chapter(heading_idxs):
            print(f"[skip] cross-chapter group {mis['key']} — refused")
            continue

        # 收集每个 keeper 块的 elem list (按 keeper 的 paragraph_indices, 跳过 delete)
        keeper_elems_per_block = []
        for b in kept_sorted:
            elems = [para_elems[i] for i in b["paragraph_indices"]
                     if i not in delete_idx_set]
            keeper_elems_per_block.append(elems)

        all_elems_flat = [e for lst in keeper_elems_per_block for e in lst]
        if not all_elems_flat:
            continue
        # anchor = 文档原顺序中最早的 keeper 段 (kept_doc 第一个块的首段)
        first_doc_block = min(kept_doc, key=lambda x: x["heading_idx"])
        anchor_elem = para_elems[first_doc_block["heading_idx"]]
        anchor_parent = anchor_elem.getparent()
        if anchor_parent is None:
            continue
        # 用 anchor 前面的非 detach 元素作为 "前驱锚"
        # 找 anchor_elem 的 previous sibling 不在 all_elems_flat 集合里的
        detach_set = set(id(e) for e in all_elems_flat)
        pred_elem = anchor_elem.getprevious()
        while pred_elem is not None and id(pred_elem) in detach_set:
            pred_elem = pred_elem.getprevious()

        # detach
        for e in all_elems_flat:
            parent = e.getparent()
            if parent is not None:
                parent.remove(e)

        # 重新计算 insert_pos: pred_elem 之后第一个位置
        if pred_elem is None:
            insert_pos = 0
        else:
            insert_pos = list(anchor_parent).index(pred_elem) + 1

        for block_elems in keeper_elems_per_block:
            for e in block_elems:
                anchor_parent.insert(insert_pos, e)
                insert_pos += 1

        moves_applied += 1

    return moves_applied, deletions_applied


# ---------- main ----------

def main_reorder(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("docx", help="输入 docx 路径")
    ap.add_argument("--dry-run", action="store_true", help="只生成 plan, 不动 docx")
    ap.add_argument("--no-backup", action="store_true", help="不自动备份")
    ap.add_argument("--report", help="JSON plan 输出路径")
    args = ap.parse_args(argv)

    src = Path(args.docx)
    if not src.exists():
        print(f"[error] 找不到 {src}", file=sys.stderr)
        return 2

    # 备份
    backup_path = None
    if not args.dry_run and not args.no_backup:
        cand = _cc.find_next_backup(src)
        shutil.copy2(src, cand)
        backup_path = cand
        print(f"[backup] {cand.name}")

    doc = Document(str(src))

    # 段分类
    paras = [classify_paragraph(p, i) for i, p in enumerate(doc.paragraphs)]
    blocks = slice_blocks(paras)

    misordered, duplicates = detect_issues(blocks)

    plan = {
        "input": str(src),
        "backup": str(backup_path) if backup_path else None,
        "blocks_total": len(blocks),
        "misordered_pairs": [
            {
                "level": m["key"][0],
                "parent_idx": m["key"][1],
                "before": [list(n) for n in m["before"]],
                "after": [list(n) for n in m["after"]],
            }
            for m in misordered
        ],
        "duplicates": [
            {
                "level": d["key"][0],
                "parent_idx": d["key"][1],
                "number": list(d["number"]),
                "kept_idx": pick_keeper(d["blocks"])["heading_idx"],
                "deleted_idxs": [b["heading_idx"] for b in d["blocks"]
                                 if b is not pick_keeper(d["blocks"])],
            }
            for d in duplicates
        ],
        "dry_run": args.dry_run,
    }

    print(f"[plan] blocks={len(blocks)}  misordered_groups={len(misordered)}  "
          f"duplicate_groups={len(duplicates)}")
    for m in misordered[:20]:
        b = [".".join(str(x) for x in n) for n in m["before"]]
        a = [".".join(str(x) for x in n) for n in m["after"]]
        print(f"  misorder L{m['key'][0]} parent_idx={m['key'][1]}: {b} → {a}")
    for d in duplicates[:20]:
        num = ".".join(str(x) for x in d["number"])
        kept = pick_keeper(d["blocks"])["heading_idx"]
        delkk = [b["heading_idx"] for b in d["blocks"] if b is not pick_keeper(d["blocks"])]
        print(f"  dup L{d['key'][0]} parent_idx={d['key'][1]} num={num}: keep idx={kept} del={delkk}")

    if not args.dry_run:
        moves, dels = apply_changes(doc, blocks, misordered, duplicates)
        plan["moves_applied"] = moves
        plan["deletions_applied"] = dels
        doc.save(str(src))
        print(f"[apply] moves={moves} deletions={dels} → saved")
    else:
        plan["moves_applied"] = 0
        plan["deletions_applied"] = 0
        print("[dry-run] no changes written")

    if args.report:
        Path(args.report).parent.mkdir(parents=True, exist_ok=True)
        Path(args.report).write_text(json.dumps(plan, indent=2, ensure_ascii=False))
        print(f"[report] {args.report}")

    return 0


# ---------------- pipeline adapter ----------------
def apply_reorder(doc, args=None) -> dict:
    dry = bool(getattr(args, "dry_run", False)) if args else False
    paras = [classify_paragraph(p, i) for i, p in enumerate(doc.paragraphs)]
    blocks = slice_blocks(paras)
    misordered, duplicates = detect_issues(blocks)
    moves = dels = 0
    if not dry:
        moves, dels = apply_changes(doc, blocks, misordered, duplicates)
    return {
        "changed": moves + dels,
        "blocks_total": len(blocks),
        "misordered_groups": len(misordered),
        "duplicate_groups": len(duplicates),
        "moves_applied": moves,
        "deletions_applied": dels,
    }


# ══════════ relocate ← relocate_orphan_blocks.py ══════════


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_P = f"{{{W_NS}}}p"
W_TBL = f"{{{W_NS}}}tbl"
W_SECTPR = f"{{{W_NS}}}sectPr"
def lsof_check(docx_path: Path) -> Optional[str]:
    """检测 docx 是否被 Word 等进程占用. 返回占用进程描述, 否则 None.

    → `_cli_common.lsof_check`（2026-08-02 四种语义收敛为一种，取的正是本文件
    原来这一版：`lsof --` + 行数>1）。"""
    return _cc.lsof_check(docx_path)


def get_body_block_children(body):
    """返回 body 直系子元素中的 block-level 元素列表 (按文档顺序).

    block-level = <w:p> 或 <w:tbl>; <w:sectPr> 与其他元素跳过.
    这是与 docx.paragraphs 的关键差异: paragraphs 不含 <w:tbl>.
    我们用全集合 blocks (p+tbl) 做 detach/insert, idx 体系也基于 blocks.
    """
    blocks = []
    for child in body:
        tag = child.tag
        if tag == W_P or tag == W_TBL:
            blocks.append(child)
    return blocks


def block_text(elem) -> str:
    """提取 block (p or tbl) 文本, strip 空白. tbl 取所有 <w:t> 文本拼接."""
    texts = []
    for t in elem.iter(f"{{{W_NS}}}t"):
        if t.text:
            texts.append(t.text)
    return "".join(texts).strip()


def find_block_idx_by_text(blocks: list, expected_text: str,
                          hint_idx: Optional[int] = None,
                          search_window: int = 50) -> Optional[int]:
    """在 blocks 中找首个文本完全匹配 expected_text 的 block idx.

    优先策略:
      1. 若 hint_idx 给了, 先检查 hint_idx (零漂移情形)
      2. 否则在 hint_idx ± search_window 范围内找
      3. 仍不中 → 全文档扫
      4. 找到多个 → 取最接近 hint_idx 的; 无 hint 取首个
    """
    expected = (expected_text or "").strip()
    if not expected:
        return None

    # 完整候选: 文本完全相等
    matches_exact = [i for i, b in enumerate(blocks) if block_text(b) == expected]
    if not matches_exact:
        # fallback: 前缀匹配 (>= 10 chars 或全 text 长度)
        min_len = min(len(expected), 30)
        prefix = expected[:min_len]
        matches_exact = [i for i, b in enumerate(blocks)
                        if block_text(b).startswith(prefix)
                        and len(block_text(b)) >= min_len]

    if not matches_exact:
        return None

    if hint_idx is None:
        return matches_exact[0]

    # 取最接近 hint 的
    return min(matches_exact, key=lambda i: abs(i - hint_idx))


def find_block_idx_by_context(blocks: list, context_text: str,
                              hint_idx: Optional[int] = None) -> Optional[int]:
    """target_context_text 是 "插入位置前一段的 text 前 50 字" — 按前缀匹配查."""
    ctx = (context_text or "").strip()
    if not ctx:
        return None

    # 完全匹配优先
    exact = [i for i, b in enumerate(blocks) if block_text(b) == ctx]
    if exact:
        if hint_idx is None:
            return exact[0]
        return min(exact, key=lambda i: abs(i - hint_idx))

    # 前缀匹配 (block text 以 ctx 开头, 或 ctx 以 block text 开头)
    # 用 ctx 前 30-50 字片段
    snippet = ctx[: min(50, len(ctx))]
    cand = []
    for i, b in enumerate(blocks):
        bt = block_text(b)
        if not bt:
            continue
        # 双向前缀容错
        n = min(len(snippet), len(bt), 30)
        if n < 8:
            continue
        if snippet[:n] == bt[:n]:
            cand.append(i)
    if not cand:
        return None
    if hint_idx is None:
        return cand[0]
    return min(cand, key=lambda i: abs(i - hint_idx))


# ---------- move 执行 ----------

def execute_move(body, move: dict, move_no: int, dry_run: bool = False) -> dict:
    """执行单个 move. 返回 result dict.

    流程:
      1. blocks = body 当前 block 子元素 (fresh)
      2. text 锚定 source_heading_idx → src_h_idx
      3. text 锚定 target_insert_after_idx → tgt_idx
      4. source range = [src_h_idx, src_h_idx + (orig_end - orig_head)]
         (用 plan 里的 end - head 差作长度)
      5. detach source 段 (collect refs), 在 tgt_idx 之后顺序 insert
      6. 防自杀: target 不能落在 source range 内
    """
    blocks = get_body_block_children(body)
    n = len(blocks)

    src_head_hint = move.get("source_heading_idx")
    src_end_hint = move.get("source_block_end_idx")
    tgt_hint = move.get("target_insert_after_idx")
    src_text = move.get("source_heading_text", "")
    tgt_text = move.get("target_context_text", "")

    # 二次锚定
    src_h_idx = find_block_idx_by_text(blocks, src_text, hint_idx=src_head_hint)
    if src_h_idx is None:
        return {
            "move_no": move_no, "status": "skip",
            "reason": f"source_heading_text not found: {src_text[:40]!r}",
        }

    tgt_idx = find_block_idx_by_context(blocks, tgt_text, hint_idx=tgt_hint)
    if tgt_idx is None:
        return {
            "move_no": move_no, "status": "skip",
            "reason": f"target_context_text not found: {tgt_text[:40]!r}",
        }

    # 计算 source range 长度 — 用 plan 里 end-head 差 (块长度对漂移不敏感)
    if src_head_hint is None or src_end_hint is None:
        return {
            "move_no": move_no, "status": "skip",
            "reason": "plan missing source_heading_idx / source_block_end_idx",
        }
    block_len = src_end_hint - src_head_hint + 1
    if block_len < 1:
        return {
            "move_no": move_no, "status": "skip",
            "reason": f"invalid block_len={block_len}",
        }

    src_end_idx = src_h_idx + block_len - 1
    if src_end_idx >= n:
        # 块长度超出当前 blocks 尾部 (可能 plan 与现状对不上) → 截到 n-1
        src_end_idx = n - 1

    # 自杀检测
    if src_h_idx <= tgt_idx <= src_end_idx:
        return {
            "move_no": move_no, "status": "skip",
            "reason": f"target idx {tgt_idx} falls inside source range [{src_h_idx},{src_end_idx}]",
        }

    if dry_run:
        return {
            "move_no": move_no, "status": "dry-ok",
            "resolved_source_head_idx": src_h_idx,
            "resolved_source_end_idx": src_end_idx,
            "resolved_target_idx": tgt_idx,
            "block_len": block_len,
        }

    # 收集 source 段的 elem refs
    source_elems = blocks[src_h_idx: src_end_idx + 1]

    # target elem ref (在 detach 前抓引用)
    target_elem = blocks[tgt_idx]

    # detach source
    parent = body
    for e in source_elems:
        if e.getparent() is parent:
            parent.remove(e)

    # 在 target_elem 之后顺序 insert
    # parent 是 body, 用 addnext 链或 index-based insert
    # 我们用 index-based: 先取 target 在 body 中的当前 index, 然后 insert(index+1, ...)
    # 注意 detach 后 target 索引会变, 重新 query
    body_children = list(body)
    try:
        tgt_pos = body_children.index(target_elem)
    except ValueError:
        return {
            "move_no": move_no, "status": "error",
            "reason": "target_elem lost from body after source detach",
        }

    insert_at = tgt_pos + 1
    for offset, e in enumerate(source_elems):
        body.insert(insert_at + offset, e)

    return {
        "move_no": move_no, "status": "ok",
        "resolved_source_head_idx": src_h_idx,
        "resolved_source_end_idx": src_end_idx,
        "resolved_target_idx": tgt_idx,
        "block_len": block_len,
        "moved_block_count": len(source_elems),
    }


# ---------- main ----------

def main_relocate(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="按 plan JSON 把 docx 孤儿段块挪到正确位置 (机械操作, 不算编号)"
    )
    ap.add_argument("docx", help="输入 docx 路径")
    ap.add_argument("--plan", required=True, help="plan JSON 路径 (含 moves[])")
    ap.add_argument("--dry-run", action="store_true", help="只 verify plan, 不动 docx")
    ap.add_argument("--no-backup", action="store_true", help="不自动备份")
    ap.add_argument("--self-test", action="store_true",
                    help="生成内置 fake plan 做 self-test (开发用)")
    args = ap.parse_args(argv)

    src = Path(args.docx)
    if not src.exists():
        print(f"[error] docx 不存在: {src}", file=sys.stderr)
        return 2

    plan_path = Path(args.plan)
    if not plan_path.exists():
        print(f"[error] plan 不存在: {plan_path}", file=sys.stderr)
        return 2

    # lsof 检查
    occupied = lsof_check(src)
    if occupied:
        print(f"[error] docx 被进程占用, 关闭 Word/WPS 后重试:\n{occupied}",
              file=sys.stderr)
        return 3

    # 读 plan + schema 校验 (doctools v1)
    try:
        plan = json.loads(plan_path.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"[error] plan JSON 解析失败: {e}", file=sys.stderr)
        return 2

    try:
        from lib.schemas import validate as _validate_schema
        _err = _validate_schema(plan, "plan")
        if _err:
            print(f"[error] plan schema 校验失败 (v1): {_err}", file=sys.stderr)
            return 2
    except Exception:
        # schemas lib 不可用时降级到原 ad-hoc 检查
        pass

    moves = plan.get("moves", [])
    if not isinstance(moves, list):
        print(f"[error] plan.moves 不是 list", file=sys.stderr)
        return 2

    print(f"[plan] {plan_path.name} — moves={len(moves)}  source_docx={plan.get('source_docx','?')}")

    # 备份
    backup_path = None
    if not args.dry_run and not args.no_backup:
        cand = _cc.find_next_backup(src)
        shutil.copy2(src, cand)
        backup_path = cand
        print(f"[backup] {cand.name}")

    # 加载 docx
    doc = Document(str(src))
    body = doc.element.body

    results = []
    for i, move in enumerate(moves, start=1):
        res = execute_move(body, move, i, dry_run=args.dry_run)
        results.append(res)
        # 简报
        status = res["status"]
        if status in ("ok", "dry-ok"):
            print(f"  [{status}] move#{i}  src_h={res.get('resolved_source_head_idx')} "
                  f"end={res.get('resolved_source_end_idx')} "
                  f"→ after idx={res.get('resolved_target_idx')}  "
                  f"len={res.get('block_len')}  "
                  f"src={move.get('source_heading_text','')[:30]!r}")
        else:
            print(f"  [{status}] move#{i}  {res.get('reason','?')}")

    # 写盘
    if not args.dry_run:
        ok_cnt = sum(1 for r in results if r["status"] == "ok")
        if ok_cnt > 0:
            doc.save(str(src))
            print(f"[save] {src.name} ← {ok_cnt} moves applied")
            # 重读自证 OOXML 合法
            try:
                _verify = Document(str(src))
                _para_n = len(_verify.paragraphs)
                print(f"[verify] re-read ok, paragraphs={_para_n}")
            except Exception as e:
                print(f"[error] re-read failed: {e}", file=sys.stderr)
                if backup_path:
                    print(f"  → 可用备份恢复: {backup_path.name}")
                return 4
        else:
            print(f"[save] 无 ok move, 不写盘")

    # 抽样 print 末尾段顺序 (前后各 5 段)
    if not args.dry_run:
        post_doc = Document(str(src))
        post_blocks = get_body_block_children(post_doc.element.body)
        print(f"\n[post-sample] total blocks={len(post_blocks)}, first 5 + last 5:")
        for i, b in enumerate(post_blocks[:5]):
            t = block_text(b)[:60]
            print(f"  [{i}] {t!r}")
        if len(post_blocks) > 10:
            print(f"  ...")
            for i, b in enumerate(post_blocks[-5:], start=len(post_blocks) - 5):
                t = block_text(b)[:60]
                print(f"  [{i}] {t!r}")

    # 汇总
    ok = sum(1 for r in results if r["status"] == "ok")
    dry = sum(1 for r in results if r["status"] == "dry-ok")
    skip = sum(1 for r in results if r["status"] == "skip")
    err = sum(1 for r in results if r["status"] == "error")
    print(f"\n[summary] ok={ok}  dry-ok={dry}  skip={skip}  error={err}  total={len(results)}")

    return 0 if err == 0 else 5


# ---------------- pipeline adapter ----------------
def apply_relocate(doc, args=None) -> dict:
    """pipeline: 仅当 args.relocate_plan 提供时执行;否则 noop"""
    plan_path = getattr(args, "relocate_plan", None) if args else None
    if not plan_path:
        return {"changed": 0, "skipped": "no relocate_plan in args"}
    dry = bool(getattr(args, "dry_run", False)) if args else False
    try:
        plan = json.loads(Path(plan_path).read_text(encoding="utf-8"))
    except Exception as exc:
        return {"error": f"plan read failed: {exc}"}
    moves = plan.get("moves", [])
    body = doc.element.body
    results = []
    for i, move in enumerate(moves, start=1):
        res = execute_move(body, move, i, dry_run=dry)
        results.append(res)
    ok = sum(1 for r in results if r["status"] == "ok")
    return {
        "changed": ok,
        "ok": ok,
        "skip": sum(1 for r in results if r["status"] == "skip"),
        "error": sum(1 for r in results if r["status"] == "error"),
    }


# ──────────────────────────── 家族入口（子命令分发）────────────────────────────

SUBCOMMANDS = {
    "fix-heading-disorder": main_fix_heading_disorder,
    "reorder": main_reorder,
    "relocate": main_relocate,
}


def main(argv: list[str] | None = None) -> int:
    return _cc.family_main(SUBCOMMANDS, argv, file=__file__)


if __name__ == "__main__":
    sys.exit(main())
