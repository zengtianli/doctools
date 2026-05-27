#!/usr/bin/env python3
# distilled from qual-supply/scripts/ (2026-05-25 W2)
# merged 5 scripts:
#   apply_body_styles.py         → style body
#   apply_table_styles.py        → style table
#   apply_caption_styles.py      → style caption
#   number_captions_by_style.py  → caption number-by-style
#   renumber_h4_figures.py       → renumber h4-figures
#
# 项目硬编码样式集合 → 抽到 doctools/config/styles_registry.yaml SSOT,通过 --profile <name> 切换。
"""doctools style/caption/renumber group · profile-driven 样式批改.

子命令:
    style body         启发式套 H1/H2/H3/Title/表名/正文 (基于段文本形态匹配 + profile styleId 目标)
    style table        所有 cell 段套 table-cell style + 表全网格 tblBorders + cell 居中
    style caption      给匹配 ^(表|图)\\d+\\.\\d+-\\d+ 的段套对 ZDWP 表名/ZDWP 图名 (或 profile 等价)
    caption number-by-style  pStyle-aware 补 "表 X-Y" / "图 X-Y" 编号 (跨 H1 重置)
    renumber h4-figures      重派 H4 编号 + 图/表 caption H1.H2-N 格式

通用 CLI:
    <docx_path> [--dry-run] [--no-backup] [--report <json>] [--profile <name>]

profile 默认 'zdwp' (兼容 qual-supply 老用法);eco-flow / generic 等通过 --profile 切换。
"""

from __future__ import annotations

import argparse
import datetime as _dt
import json
import re
import shutil
import subprocess
import sys
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Optional

# python-docx / lxml — 顶部 import 让所有子命令共享
from docx import Document
from docx.document import Document as DocType
from docx.oxml.ns import qn
from lxml import etree

# doctools.lib.styles loader
# 兼容两种 import 路径: (a) doctools package; (b) 同 repo 相对导入
try:
    from doctools.lib.styles import load_profile, StylesProfile  # type: ignore
except ImportError:
    # fallback: 直接从文件加载 (脚本被独立调用时)
    import importlib.util as _ilu
    _styles_path = Path(__file__).resolve().parent.parent.parent.parent / "lib" / "styles.py"
    _spec = _ilu.spec_from_file_location("_doctools_styles", _styles_path)
    _m = _ilu.module_from_spec(_spec)
    _spec.loader.exec_module(_m)  # type: ignore
    load_profile = _m.load_profile
    StylesProfile = _m.StylesProfile


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


# =============================================================================
# 公共辅助
# =============================================================================

def _qn_local(tag: str) -> str:
    return f"{W}{tag}"


def lsof_check(path: Path) -> Optional[str]:
    """检测 docx 是否被 Word/WPS 占用. 返回 lsof 输出 (occupied) 或 None (free)."""
    try:
        r = subprocess.run(
            ["lsof", str(path)], capture_output=True, text=True, timeout=5
        )
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def pick_backup_path(src: Path) -> Path:
    today = _dt.date.today().isoformat()
    parent, stem, suffix = src.parent, src.stem, src.suffix
    n = 1
    while True:
        cand = parent / f"{stem}.bak-{n}-{today}{suffix}"
        if not cand.exists():
            return cand
        n += 1


def _common_setup(args, dry_skip_lsof: bool = False) -> Path:
    """通用 setup: 验文件存在 + lsof 自检 (除 dry-run/--no-backup 时仍检)."""
    src: Path = args.docx if hasattr(args, "docx") else args.docx_path
    if not Path(src).exists():
        print(f"[ERR] 文件不存在: {src}", file=sys.stderr)
        sys.exit(2)
    src = Path(src).resolve()
    if not args.dry_run and not dry_skip_lsof:
        occ = lsof_check(src)
        if occ:
            print(f"[ABORT] 文件被占用 (Word/WPS):\n{occ}", file=sys.stderr)
            sys.exit(3)
    return src


def _save_with_backup(src: Path, doc, args, wrote_needed: bool = True) -> Optional[Path]:
    """统一 save 流程: dry-run/no-need 不写; --no-backup 跳备份."""
    if args.dry_run:
        return None
    if not wrote_needed:
        return None
    bak = None
    if not args.no_backup:
        bak = pick_backup_path(src)
        shutil.copy2(src, bak)
    doc.save(str(src))
    return bak


# =============================================================================
# 子命令 1: style body — apply_body_styles
# =============================================================================
# 启发式分类 (按优先级, top 先匹) — 文本形态识别 (项目无关 regex)

RE_TABLE_NAME = re.compile(r"^表\s*[\d一二三四五六七八九十]+[-—]?[\d一二三四五六七八九十]*\s+")
RE_TOP_CN     = re.compile(r"^[一二三四五六七八九十]+、")
RE_CHAPTER    = re.compile(r"^第[一二三四五六七八九十\d]+章\s+")
RE_NUM_TITLE  = re.compile(r"^\d+\s+\S")
RE_H3         = re.compile(r"^\d+\.\d+\.\d+\s+")
RE_H2         = re.compile(r"^\d+\.\d+\s+")

# ──────────────────────────────────────────────────────────────────────────────
# 白/黑名单 (CRITICAL FIX 2026-05-26):
# style body 仅按文本形态判,缺现有 style 守门 → 把 Title/Heading 1-4/自定义"N级标题"
# 误套成 caption。下面两个 PROTECTED / BODY-LIKE 集合在 _apply_body_impl 入口拦截:
#   - PROTECTED: 命中即 skip,绝不修改 (heading/title/toc/caption/中文 N级标题/
#                ZDWP H*/数字打头"N 级标题"等)
#   - BODY_LIKE: 命中或样式为空/Normal 才允许 apply body 启发式
# 反模式: 用 fnmatch/in 简单匹配,中文样式 ("1一级标题"/"1.1.1.1 N级标题") 必漏。
# 这里用 regex 同时覆盖 styleId 与 style name 两条线 (docx 里 styleId 常是 "10"/"1"
# 等数字, name 是 "Heading 1"/"标题 1"/"1一级标题")。
# ──────────────────────────────────────────────────────────────────────────────

# 保护类(命中绝不改)。两条线: style name 与 styleId 都过。
_PROTECTED_NAME_PATTERNS = [
    re.compile(r"^Heading\s*[1-9]$", re.I),
    re.compile(r"^Title$", re.I),
    re.compile(r"^Subtitle$", re.I),
    re.compile(r"^TOC.*", re.I),
    re.compile(r"^toc.*", re.I),
    re.compile(r"^目录"),
    re.compile(r"标题"),                              # 含"标题"二字 (含 Word 中文"标题 1" / "1一级标题" / "1.1.1.1 N级标题")
    re.compile(r"题目"),
    re.compile(r"^(一|二|三|四|五|六|七|八|九|十)级标题"),
    re.compile(r"^\d+\s*(一|二|三|四|五|六|七|八|九|十)级"),  # "1一级" "2二级" 数字打头
    re.compile(r"^ZDWP\s*H[1-9]", re.I),
    re.compile(r"^ZDWP\s*标题", re.I),
    # 已有 caption 样式 — 别重套
    re.compile(r"图\s*名"),
    re.compile(r"表\s*名"),
    re.compile(r"图名称"),
    re.compile(r"表格标题"),
    re.compile(r"^caption", re.I),
    re.compile(r"Image\s*Caption", re.I),
    re.compile(r"图\s*注"),
    re.compile(r"表\s*注"),
    re.compile(r"表\s*题"),
]

# 正向白名单 — 可被 body 启发式套样式的段必须 style 命中以下之一(或为空/None)。
_BODY_LIKE_NAME_PATTERNS = [
    re.compile(r"^Normal$", re.I),
    re.compile(r"^Normal\s+Indent$", re.I),
    re.compile(r"^Body\s*Text\s*[123]?$", re.I),
    re.compile(r"^正文"),                              # "正文" / "01正文" / "02正文"
    re.compile(r"^\d+\s*正文"),
    re.compile(r"ZDWP正文"),
    re.compile(r"^ZDWP$"),                            # styleId
    re.compile(r"^Default\s*Paragraph", re.I),
]


def _style_name_id_of(p) -> tuple[Optional[str], Optional[str]]:
    """Return (style_name, style_id) of paragraph; (None, None) if no style."""
    try:
        s = p.style
    except Exception:
        return (None, None)
    if s is None:
        return (None, None)
    name = getattr(s, "name", None) or None
    sid = getattr(s, "style_id", None) or None
    return (name, sid)


# ──────────────────────────────────────────────────────────────────────────────
# CAPTION-RISK GUARD (CRITICAL FIX 2026-05-26 W10 灾难补丁):
# profile.{BODY,TITLE,H1,H2,H3}_TARGET_STYLE_ID 在 docx 里若漂移指向 caption 样式
# (style.name 含 "图名" / "表名" / "Caption" 等, 或 style XML 含 <w:numPr> / SEQ),
# style body 会把 Normal 段全套成 caption 自动编号样式 → 整本文档报废
# ("图6.1-1 → 磐安县小型水库..." 这类前缀)。
# 防御: 跑主循环之前先 inspect profile 的所有 NON-CAPTION 目标 styleId, 解析其 docx
# 内实际 style 节点 (name + XML), 命中下列条件之一即 refuse:
#   - style.name 含 caption-style 关键字 (中英都覆盖)
#   - style XML <w:pPr><w:numPr> 存在 (paragraph 进入自动编号 list)
#   - style XML 任何 <w:instrText> 含 SEQ (字段计数器)
#   - style XML 任何 <w:fldChar> 存在 (field begin/separate/end)
# --force 旁路 (打 4 行 WARNING)。
# ──────────────────────────────────────────────────────────────────────────────

_CAPTION_TARGET_NAME_PATTERNS: list[re.Pattern] = [
    re.compile(r"图\s*名"),
    re.compile(r"表\s*名"),
    re.compile(r"图名称"),
    re.compile(r"表名称"),
    re.compile(r"表格标题"),
    re.compile(r"图\s*注"),
    re.compile(r"表\s*注"),
    re.compile(r"题\s*注"),
    re.compile(r"题\s*名"),
    re.compile(r"^caption$", re.I),
    re.compile(r"image\s*caption", re.I),
    re.compile(r"table\s*caption", re.I),
    re.compile(r"figure\s*caption", re.I),
    re.compile(r"^figure$", re.I),
    re.compile(r"^table$", re.I),
]

# profile field 名 → 该 target 是否允许 caption-like 名字
# (TABLE_CAPTION_TARGET 本身就是 caption, 名字 OK; 但仍不许 w:numPr / SEQ)
_NON_CAPTION_TARGET_FIELDS = (
    "BODY_TARGET_STYLE_ID",
    "TITLE_TARGET_STYLE_ID",
    "H1_TARGET_STYLE_ID",
    "H2_TARGET_STYLE_ID",
    "H3_TARGET_STYLE_ID",
)
_ALL_TARGET_FIELDS = _NON_CAPTION_TARGET_FIELDS + ("TABLE_CAPTION_TARGET_STYLE_ID",)

# 哪些 target 允许带 <w:numPr> (Word 标准 heading 自带列表编号是正常的)。
# Body / Title / Caption 带 numPr = 灾难 (正文/标题/题注被自动编号)。
_NUMPR_ALLOWED_TARGET_FIELDS = frozenset({
    "H1_TARGET_STYLE_ID",
    "H2_TARGET_STYLE_ID",
    "H3_TARGET_STYLE_ID",
})


def _style_element_by_id(doc, style_id: str):
    """Return lxml <w:style> element matching styleId, or None."""
    try:
        s = doc.styles.get_by_id(style_id, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    except Exception:
        for st in doc.styles:
            if getattr(st, "style_id", None) == style_id:
                s = st
                break
        else:
            return None, None
    name = getattr(s, "name", None) or None
    el = getattr(s, "element", None)
    return name, el


def _inspect_style_xml_for_auto_numbering(el, allow_numpr: bool = False) -> list[str]:
    """Return list of reason strings for auto-numbering risk in this style XML.

    allow_numpr=True for H1/H2/H3 targets (heading numPr is by design).
    """
    reasons: list[str] = []
    if el is None:
        return reasons
    # <w:pPr><w:numPr> → list / numbered auto-prefix (灾难 for body/title/caption)
    if not allow_numpr:
        numpr_nodes = el.findall(".//" + qn("w:numPr"))
        if numpr_nodes:
            nums = []
            for np in numpr_nodes:
                nid = np.find(qn("w:numId"))
                if nid is not None:
                    nums.append(nid.get(qn("w:val")) or "?")
            reasons.append(f"<w:numPr> present (numId={nums or ['?']}) — 自动列表编号")
    # <w:instrText> containing SEQ field
    for it in el.findall(".//" + qn("w:instrText")):
        txt = (it.text or "").strip()
        if "SEQ" in txt.upper():
            reasons.append(f"<w:instrText> contains SEQ field: {txt!r}")
            break
    # any <w:fldChar>
    if el.findall(".//" + qn("w:fldChar")):
        reasons.append("<w:fldChar> present in style — field auto-prefix")
    return reasons


def _inspect_caption_risk(doc, profile: StylesProfile) -> dict:
    """Inspect every target styleId in profile; return dict of refusal info.

    Returns:
        {
            "refused": bool,
            "risks": [
                {"field": "BODY_TARGET_STYLE_ID", "style_id": "ZDWP",
                 "style_name": "ZDWP 图名", "reasons": ["name matches /图\\s*名/", ...]},
                ...
            ],
        }
    """
    risks: list[dict] = []
    for field in _ALL_TARGET_FIELDS:
        sid = getattr(profile, field, None)
        if not sid:
            continue
        style_name, el = _style_element_by_id(doc, sid)
        local_reasons: list[str] = []
        # Layer A: name pattern (only for non-caption targets — caption targets
        # legitimately have caption-like names)
        if field in _NON_CAPTION_TARGET_FIELDS and style_name:
            for pat in _CAPTION_TARGET_NAME_PATTERNS:
                if pat.search(style_name):
                    local_reasons.append(
                        f"style.name={style_name!r} matches caption pattern /{pat.pattern}/"
                    )
                    break
        # Layer B: auto-numbering / SEQ field in style XML
        allow_numpr = field in _NUMPR_ALLOWED_TARGET_FIELDS
        local_reasons.extend(
            _inspect_style_xml_for_auto_numbering(el, allow_numpr=allow_numpr)
        )
        if local_reasons:
            risks.append({
                "field": field,
                "style_id": sid,
                "style_name": style_name,
                "reasons": local_reasons,
            })
    return {"refused": bool(risks), "risks": risks}


def _format_caption_risk_message(profile_name: str, info: dict) -> str:
    lines = [
        "",
        "\033[1;31m" + "=" * 78 + "\033[0m",
        f"\033[1;31m[REFUSED] style body — profile {profile_name!r} 目标样式存在 caption / 自动编号风险\033[0m",
        "\033[1;31m" + "=" * 78 + "\033[0m",
        "",
    ]
    for r in info["risks"]:
        lines.append(
            f"  · profile.{r['field']} = {r['style_id']!r} → "
            f"docx 内 style.name = {r['style_name']!r}"
        )
        for reason in r["reasons"]:
            lines.append(f"      - {reason}")
        lines.append("")
    lines += [
        "  跑下去会把 Normal 段套成 caption / 自动编号样式,",
        "  正文将被 Word 自动插入 \"图X-Y\" / \"表X-Y\" / SEQ 编号前缀,文档报废。",
        "",
        "  常见根因: yaml profile 与本 docx 实际 styleId 漂移 (此 styleId 在本 docx",
        "  被命名为 caption 样式或含 w:numPr / SEQ 字段)。",
        "",
        "  修复路径:",
        "    A. 用 `style rename` 把当前 styleId 改名为正文样式 (推荐)",
        "    B. 修 profile.*_TARGET_STYLE_ID 指向真正的正文 styleId",
        "    C. --force 跳过此检查 (不推荐;若 profile/docx 真匹配错则文档会被毁)",
        "",
    ]
    return "\n".join(lines)


def _is_protected_paragraph(style_name: Optional[str], style_id: Optional[str], profile: StylesProfile) -> bool:
    """True 表示这段是 heading/title/toc/caption,style body 不许碰."""
    # 1. profile 自身的语义判定 (基于 yaml *_STYLES 集合)
    for predicate in (profile.is_h1, profile.is_h2, profile.is_h3, profile.is_h4,
                      profile.is_title, profile.is_table_caption, profile.is_fig_caption):
        if predicate(style_name) or predicate(style_id):
            return True
    # 2. 正则补充 — 防 profile 未覆盖的中文 / 数字打头自定义样式
    for s in (style_name, style_id):
        if not s:
            continue
        for pat in _PROTECTED_NAME_PATTERNS:
            if pat.search(s):
                return True
    return False


def _is_body_like_paragraph(style_name: Optional[str], style_id: Optional[str], profile: StylesProfile) -> bool:
    """True 表示这段是 Normal / 正文系列 / 无名样式,允许 body 启发式套样式."""
    # 无样式 / 空 → body-like
    if not style_name and not style_id:
        return True
    # profile 自身 body 集合命中
    if profile.is_body(style_name) or profile.is_body(style_id):
        return True
    # 正则白名单
    for s in (style_name, style_id):
        if not s:
            continue
        for pat in _BODY_LIKE_NAME_PATTERNS:
            if pat.search(s):
                return True
    return False


def _classify_body_paragraph(text: str, profile: StylesProfile) -> Optional[str]:
    """返回目标 styleId; None = 跳过 (空段). styleId 从 profile 取.

    NOTE: 仅按 text 形态判,不看 current style — 守门由 _apply_body_impl 入口的
    _is_protected_paragraph / _is_body_like_paragraph 完成。
    """
    s = (text or "").strip()
    if s == "":
        return None
    if RE_TABLE_NAME.match(s):
        return profile.TABLE_CAPTION_TARGET_STYLE_ID
    if RE_TOP_CN.match(s):
        return profile.TITLE_TARGET_STYLE_ID
    if RE_CHAPTER.match(s):
        return profile.H1_TARGET_STYLE_ID
    if RE_H3.match(s):
        return profile.H3_TARGET_STYLE_ID
    if RE_H2.match(s):
        return profile.H2_TARGET_STYLE_ID
    if RE_NUM_TITLE.match(s) and "." not in s.split()[0]:
        return profile.H1_TARGET_STYLE_ID
    return profile.BODY_TARGET_STYLE_ID


def _fix_zdwp_next_field(doc, body_style_id: str) -> str:
    """ZDWP 样式 <w:next w:val="ZDWP"/> 自指修复. profile.ZDWP_NEXT_FIELD=True 才调用."""
    from docx.oxml import OxmlElement
    styles_elem = doc.styles.element
    target = None
    for s in styles_elem.findall(qn("w:style")):
        if s.get(qn("w:styleId")) == body_style_id:
            target = s
            break
    if target is None:
        return "no-body-style"
    nxt = target.find(qn("w:next"))
    if nxt is None:
        nxt = OxmlElement("w:next")
        nxt.set(qn("w:val"), body_style_id)
        # FutureWarning fix: element truth-testing 已被 lxml deprecate,显式 is not None
        _bo = target.find(qn("w:basedOn"))
        anchor = _bo if _bo is not None else target.find(qn("w:name"))
        if anchor is not None:
            anchor.addnext(nxt)
        else:
            target.insert(0, nxt)
        return "added"
    if nxt.get(qn("w:val")) != body_style_id:
        nxt.set(qn("w:val"), body_style_id)
        return "updated"
    return "ok"


def _apply_body_impl(doc, profile: StylesProfile, dry_run: bool) -> dict:
    stats = {
        "total": 0,
        "changed": Counter(),
        "skipped_empty": 0,
        "skipped_already": Counter(),
        "skipped_no_target_in_profile": 0,
        "skipped_protected": Counter(),       # heading/title/toc/caption etc.
        "skipped_unknown_style": Counter(),   # 非 protected 也非 body-like — 不动
        "errors": [],
    }
    available_ids = {s.style_id for s in doc.styles}
    for idx, p in enumerate(doc.paragraphs):
        stats["total"] += 1
        text = p.text or ""
        # 0. 空段直接 skip (与原逻辑一致)
        if text.strip() == "":
            stats["skipped_empty"] += 1
            continue
        style_name, style_id = _style_name_id_of(p)
        # 1. PROTECTED — heading / title / toc / caption 等绝不改
        if _is_protected_paragraph(style_name, style_id, profile):
            key = f"{style_name or '?'}|{style_id or '?'}"
            stats["skipped_protected"][key] += 1
            continue
        # 2. 非 body-like (例如 ZDWP 表名/0 表格内容 等已有专用样式) — 不主动 reset
        if not _is_body_like_paragraph(style_name, style_id, profile):
            key = f"{style_name or '?'}|{style_id or '?'}"
            stats["skipped_unknown_style"][key] += 1
            continue
        # 3. 走启发式分类 → target styleId
        target = _classify_body_paragraph(text, profile)
        if target is None or target == "":
            stats["skipped_no_target_in_profile"] += 1
            continue
        if target not in available_ids:
            stats["errors"].append((idx, text[:60], f"styleId {target!r} not in styles.xml"))
            continue
        cur_id = style_id
        if cur_id == target:
            stats["skipped_already"][target] += 1
            continue
        if not dry_run:
            try:
                p.style = doc.styles.get_by_id(target, 1)
            except Exception:
                pPr = p._p.get_or_add_pPr()
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is None:
                    from docx.oxml import OxmlElement
                    pStyle = OxmlElement("w:pStyle")
                    pPr.insert(0, pStyle)
                pStyle.set(qn("w:val"), target)
        stats["changed"][target] += 1
    return stats


def cmd_body(args) -> int:
    # ── TEMPLATE INJECTION MODE (2026-05-26) ──────────────────────────────────
    # `style body --template std` 只注入预设样式骨架, 不跑启发式 body-apply。
    # 用例: 全裸 docx (无 ZDWP 系列样式) 先打底, 再跑 `styleset restore` 不抛
    # "ZDWP正文 not found". 与 body 启发式正交 — 用户想要套段样式再单跑
    # `style body X.docx --profile zdwp` 即可。
    tmpl = getattr(args, "template", None)
    if tmpl:
        return _cmd_template_inject(args)

    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))

    # ── CAPTION-RISK GUARD (CRITICAL FIX 2026-05-26 W10) ──────────────────────
    # 跑主循环前先查 profile.*_TARGET_STYLE_ID 在本 docx 里是不是漂移指向 caption /
    # 自动编号样式; 命中 → exit 2 (除非 --force).
    risk_info = _inspect_caption_risk(doc, profile)
    profile_name = getattr(profile, "_name", args.profile)
    if risk_info["refused"]:
        msg = _format_caption_risk_message(profile_name or "<?>", risk_info)
        if not getattr(args, "force", False):
            sys.stderr.write(msg)
            sys.stderr.flush()
            # also emit machine-readable report if --report given
            if getattr(args, "report", None):
                _emit_report({
                    "subcommand": "style body",
                    "docx": str(src),
                    "profile": profile_name,
                    "refused": True,
                    "caption_risk": risk_info["risks"],
                }, args)
            return 2
        # --force bypass
        sys.stderr.write(msg)
        sys.stderr.write("\033[1;33m" + "!" * 78 + "\033[0m\n")
        sys.stderr.write("\033[1;33m!! --force 已绕过 caption-risk guard\033[0m\n")
        sys.stderr.write("\033[1;33m!! 文档可能被毁; 你最好知道自己在干什么\033[0m\n")
        sys.stderr.write("\033[1;33m" + "!" * 78 + "\033[0m\n\n")
        sys.stderr.flush()

    stats = _apply_body_impl(doc, profile, args.dry_run)

    next_status = "(dry-run skip)"
    if not args.dry_run and profile.ZDWP_NEXT_FIELD and profile.BODY_TARGET_STYLE_ID:
        next_status = _fix_zdwp_next_field(doc, profile.BODY_TARGET_STYLE_ID)

    bak = _save_with_backup(src, doc, args, wrote_needed=True)

    report = {
        "subcommand": "style body",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": args.dry_run,
        "backup": str(bak) if bak else None,
        "total": stats["total"],
        "skipped_empty": stats["skipped_empty"],
        "changed_by_style": dict(stats["changed"]),
        "skipped_already": dict(stats["skipped_already"]),
        "skipped_no_target_in_profile": stats["skipped_no_target_in_profile"],
        "skipped_protected": dict(stats["skipped_protected"]),
        "skipped_unknown_style": dict(stats["skipped_unknown_style"]),
        "zdwp_next_field": next_status,
        "errors": [{"idx": i, "text": t, "reason": r} for (i, t, r) in stats["errors"]],
    }
    _emit_report(report, args)
    print(f"[style body] file={src.name} total={stats['total']} "
          f"changed={sum(stats['changed'].values())} "
          f"skipped_already={sum(stats['skipped_already'].values())} "
          f"skipped_protected={sum(stats['skipped_protected'].values())} "
          f"skipped_unknown={sum(stats['skipped_unknown_style'].values())} "
          f"zdwp_next={next_status}")
    return 0


# =============================================================================
# style body --template std — XML 骨架注入 (2026-05-26)
# =============================================================================
# 全裸 docx (fresh md2word 等无 ZDWP 系列样式) 用此打底 9 段样式 + 再跑
# `styleset restore` 即可走 9-step 不抛 "style ZDWP正文 not found".
#
# 模板来源: profiles/templates/std_styleset.xml = 磐安 v3 抽出的 9 个 <w:style>
# 节点 (ZDWP正文/ZDWP图名/ZDWP 表名/ZDWP附表/zdwp题目0/zdwp题目1/zdwp作者/
# zdwp封面日期/ZDWP表格内容). 注入时:
#   - 跳过已存在 (按 styleId + name 双键查)
#   - basedOn/link/next 引用了目标 docx 里没有的 styleId → 重定向到 docx 的
#     Normal styleId; 若仍找不到对应字符样式 (link) 则删 link 元素
#
# 与 styleset restore 解耦: 这里只补 styles.xml 定义, 不动 paragraph.

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "profiles" / "templates"

# 9 std styles (按磐安 v3 抽出, name 即 SSOT)
STD_STYLE_NAMES = (
    "ZDWP表格内容", "ZDWP图名", "ZDWP 表名", "ZDWP正文", "ZDWP附表",
    "zdwp题目0", "zdwp题目1", "zdwp作者", "zdwp封面日期",
)


def _read_template_xml(name: str) -> etree._Element:
    """读 profiles/templates/<name>_styleset.xml, 返回 <w:styles> root."""
    p = TEMPLATE_DIR / f"{name}_styleset.xml"
    if not p.exists():
        raise FileNotFoundError(f"template not found: {p}")
    return etree.fromstring(p.read_bytes())


def _docx_styles_index(styles_root) -> tuple[dict, dict, str]:
    """返回 (id2el, name2el, normal_style_id).

    name2el 用于 'name 已存在' 检测; id2el 用于 'basedOn/link 引用 sid 存在性'
    检测; normal_style_id = name=='Normal' 的 styleId (fallback 'Normal').
    """
    id2el = {}
    name2el = {}
    normal_sid = None
    for st in styles_root.findall(f"{W}style"):
        sid = st.get(f"{W}styleId")
        if sid:
            id2el[sid] = st
        nm_el = st.find(f"{W}name")
        if nm_el is not None:
            nm_v = nm_el.get(f"{W}val")
            if nm_v:
                name2el[nm_v] = st
                if nm_v in ("Normal", "正文") and st.get(f"{W}type") == "paragraph":
                    normal_sid = sid or normal_sid
    return id2el, name2el, (normal_sid or "Normal")


def _rewrite_refs(style_el, docx_id2el: dict, docx_normal_sid: str) -> list[tuple[str, str, str]]:
    """重写 basedOn / link / next: 若引用的 styleId 不在 docx_id2el, 改写或删除.

    返回 [(tag, old_val, new_val_or_'<deleted>'), ...] 用于汇报.
    规则:
        basedOn / next 段落引用 → 重写为 docx_normal_sid
        link 字符样式引用 → 直接删除该元素 (字符样式不强制)
    """
    actions = []
    for tag in ("basedOn", "next", "link"):
        el = style_el.find(f"{W}{tag}")
        if el is None:
            continue
        old = el.get(f"{W}val")
        if not old or old in docx_id2el:
            continue
        if tag == "link":
            style_el.remove(el)
            actions.append((tag, old, "<deleted>"))
        else:
            el.set(f"{W}val", docx_normal_sid)
            actions.append((tag, old, docx_normal_sid))
    return actions


def _inject_template_styles(docx_path: Path, template_name: str, dry_run: bool) -> dict:
    """注入 template 的 styles 到 docx/word/styles.xml. python-docx 不便操作整个
    styles part XML 树 → 直接用 zipfile 重写.
    """
    import zipfile
    template_root = _read_template_xml(template_name)

    # 读 docx 现有 styles.xml
    with zipfile.ZipFile(str(docx_path), "r") as zin:
        styles_bytes = zin.read("word/styles.xml")
        names = zin.namelist()
    docx_styles_root = etree.fromstring(styles_bytes)
    id2el, name2el, normal_sid = _docx_styles_index(docx_styles_root)

    added = []
    skipped = []
    ref_rewrites = []
    for tpl_st in template_root.findall(f"{W}style"):
        sid = tpl_st.get(f"{W}styleId")
        nm_el = tpl_st.find(f"{W}name")
        nm = nm_el.get(f"{W}val") if nm_el is not None else None
        # 双键查重
        if (sid and sid in id2el) or (nm and nm in name2el):
            skipped.append({"styleId": sid, "name": nm, "reason": "exists"})
            continue
        # 深拷贝 (避免污染 template_root)
        new_st = etree.fromstring(etree.tostring(tpl_st))
        # 重写跨 docx 失效引用
        acts = _rewrite_refs(new_st, id2el, normal_sid)
        if acts:
            ref_rewrites.append({"styleId": sid, "name": nm, "rewrites": acts})
        docx_styles_root.append(new_st)
        # 更新索引以便后续 template 内自引用 (本批 9 个之间未来扩展时)
        if sid:
            id2el[sid] = new_st
        if nm:
            name2el[nm] = new_st
        added.append({"styleId": sid, "name": nm})

    result = {
        "template": template_name,
        "added": added,
        "skipped": skipped,
        "ref_rewrites": ref_rewrites,
        "docx_normal_sid": normal_sid,
        "dry_run": dry_run,
    }
    if dry_run or not added:
        return result

    # 写回 docx — zip rebuild (替换 word/styles.xml)
    new_styles_bytes = etree.tostring(
        docx_styles_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    import os
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(str(docx_path), "r") as zin, \
         zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/styles.xml":
                zout.writestr(item, new_styles_bytes)
            else:
                zout.writestr(item, zin.read(item.filename))
    os.replace(str(tmp_path), str(docx_path))
    return result


def _cmd_template_inject(args) -> int:
    src = _common_setup(args)
    tmpl = args.template
    # 备份 (复用 _save_with_backup 的精神, 但我们直接 zip 操作, 单独 copy)
    bak = None
    if not args.dry_run and not args.no_backup:
        bak = pick_backup_path(src)
        shutil.copy2(src, bak)

    try:
        result = _inject_template_styles(src, tmpl, args.dry_run)
    except FileNotFoundError as e:
        print(f"[ERR] {e}", file=sys.stderr)
        return 4

    result["docx"] = str(src)
    result["backup"] = str(bak) if bak else None
    result["subcommand"] = "style body --template"
    _emit_report(result, args)
    n_add = len(result["added"])
    n_skip = len(result["skipped"])
    n_rw = len(result["ref_rewrites"])
    suffix = " (dry-run)" if args.dry_run else ""
    print(
        f"[style body --template {tmpl}] file={src.name} "
        f"added={n_add} skipped={n_skip} ref_rewrites={n_rw}{suffix}"
    )
    if n_add:
        names = ", ".join(s["name"] or s["styleId"] for s in result["added"])
        print(f"  + {names}")
    return 0


# =============================================================================
# 子命令 2: style table — apply_table_styles
# =============================================================================
TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
    "tblStyleRowBandSize", "tblStyleColBandSize",
    "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar",
    "tblLook", "tblCaption", "tblDescription", "tblPrChange",
]
TBLPR_ORDER_IDX = {n: i for i, n in enumerate(TBLPR_ORDER)}

TCPR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge",
    "tcBorders", "shd", "noWrap", "tcMar", "textDirection",
    "tcFitText", "vAlign", "hideMark", "headers",
    "cellIns", "cellDel", "cellMerge", "tcPrChange",
]
TCPR_ORDER_IDX = {n: i for i, n in enumerate(TCPR_ORDER)}

BORDER_SIDES = ("top", "left", "bottom", "right", "insideH", "insideV")

_PPR_AFTER_JC = {
    "textDirection", "textAlignment", "textboxTightWrap",
    "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr", "pPrChange",
}


def _insert_in_order(parent, child, order_idx=None):
    if order_idx is None:
        order_idx = TBLPR_ORDER_IDX
    tag_local = etree.QName(child).localname
    for existing in parent.findall(_qn_local(tag_local)):
        parent.remove(existing)
    target_idx = order_idx.get(tag_local, 999)
    inserted = False
    for i, sib in enumerate(list(parent)):
        sib_local = etree.QName(sib).localname
        sib_idx = order_idx.get(sib_local, 999)
        if sib_idx > target_idx:
            parent.insert(i, child)
            inserted = True
            break
    if not inserted:
        parent.append(child)


def _set_pPr_jc(pPr, val="center") -> bool:
    jc = pPr.find(_qn_local("jc"))
    if jc is not None:
        if jc.get(_qn_local("val")) == val:
            return False
        jc.set(_qn_local("val"), val)
        return True
    jc = etree.Element(_qn_local("jc"))
    jc.set(_qn_local("val"), val)
    for i, sib in enumerate(list(pPr)):
        if etree.QName(sib).localname in _PPR_AFTER_JC:
            pPr.insert(i, jc)
            return True
    pPr.append(jc)
    return True


def _apply_cell_style(doc, style_id: str) -> tuple[int, int]:
    changed = skipped = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p_el = p._p
                    pPr = p_el.find(_qn_local("pPr"))
                    if pPr is None:
                        pPr = etree.SubElement(p_el, _qn_local("pPr"))
                        p_el.insert(0, pPr)
                    pStyle = pPr.find(_qn_local("pStyle"))
                    if pStyle is not None and pStyle.get(_qn_local("val")) == style_id:
                        skipped += 1
                        continue
                    if pStyle is None:
                        pStyle = etree.SubElement(pPr, _qn_local("pStyle"))
                        pPr.insert(0, pStyle)
                    pStyle.set(_qn_local("val"), style_id)
                    changed += 1
    return changed, skipped


def _set_table_borders(tbl) -> int:
    tblPr = tbl.find(_qn_local("tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(tbl, _qn_local("tblPr"))
        tbl.insert(0, tblPr)
    tblBorders = etree.Element(_qn_local("tblBorders"))
    for side in BORDER_SIDES:
        b = etree.SubElement(tblBorders, _qn_local(side))
        b.set(_qn_local("val"), "single")
        b.set(_qn_local("sz"), "4")
        b.set(_qn_local("space"), "0")
        b.set(_qn_local("color"), "auto")
    _insert_in_order(tblPr, tblBorders)
    return 6


def _clear_cell_borders(tbl) -> int:
    removed = 0
    for tcBorders in tbl.findall(
        f".//{_qn_local('tc')}/{_qn_local('tcPr')}/{_qn_local('tcBorders')}"
    ):
        tcBorders.getparent().remove(tcBorders)
        removed += 1
    return removed


def _center_cell_content(doc) -> tuple[int, int]:
    paragraphs_centered = cells_valign_set = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p_el = p._p
                    pPr = p_el.find(_qn_local("pPr"))
                    if pPr is None:
                        pPr = etree.Element(_qn_local("pPr"))
                        p_el.insert(0, pPr)
                    if _set_pPr_jc(pPr, "center"):
                        paragraphs_centered += 1
                tc = cell._tc
                tcPr = tc.find(_qn_local("tcPr"))
                if tcPr is None:
                    tcPr = etree.Element(_qn_local("tcPr"))
                    tc.insert(0, tcPr)
                vAlign = tcPr.find(_qn_local("vAlign"))
                if vAlign is not None:
                    if vAlign.get(_qn_local("val")) != "center":
                        vAlign.set(_qn_local("val"), "center")
                        cells_valign_set += 1
                else:
                    vAlign = etree.Element(_qn_local("vAlign"))
                    vAlign.set(_qn_local("val"), "center")
                    _insert_in_order(tcPr, vAlign, TCPR_ORDER_IDX)
                    cells_valign_set += 1
    return paragraphs_centered, cells_valign_set


def cmd_table(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))

    cell_style_id = args.style_id or profile.TABLE_CELL_STYLE_ID or "0"

    if args.dry_run:
        # 仅统计 (dry-run 不写)
        n_tables = len(doc.tables)
        report = {
            "subcommand": "style table",
            "docx": str(src),
            "profile": getattr(profile, "_name", args.profile),
            "dry_run": True,
            "n_tables": n_tables,
            "cell_style_id": cell_style_id,
            "note": "dry-run: would set cell pStyle / tblBorders / jc-center / vAlign",
        }
        _emit_report(report, args)
        print(f"[style table] DRY-RUN tables={n_tables}")
        return 0

    changed, skipped = _apply_cell_style(doc, cell_style_id)
    border_set = tc_removed = 0
    for t in doc.tables:
        border_set += _set_table_borders(t._tbl)
        tc_removed += _clear_cell_borders(t._tbl)
    paragraphs_centered, cells_valign_set = _center_cell_content(doc)

    bak = _save_with_backup(src, doc, args, wrote_needed=True)
    report = {
        "subcommand": "style table",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": False,
        "backup": str(bak) if bak else None,
        "cell_style_id": cell_style_id,
        "n_tables": len(doc.tables),
        "cell_pStyle_changed": changed,
        "cell_pStyle_skipped": skipped,
        "border_sides_set": border_set,
        "tcBorders_removed": tc_removed,
        "paragraphs_centered": paragraphs_centered,
        "cells_valign_set": cells_valign_set,
    }
    _emit_report(report, args)
    print(f"[style table] tables={report['n_tables']} cells_changed={changed} "
          f"borders={border_set} centered={paragraphs_centered}/{cells_valign_set}")
    return 0


# =============================================================================
# 子命令 3: style caption — apply_caption_styles
# =============================================================================
_CAPTION_PATTERN = re.compile(r"^(表|图)\s*\d+\.\d+-\d+")


def _scan_and_apply_caption(doc, profile: StylesProfile, do_apply: bool) -> dict:
    available = {s.name for s in doc.styles}
    table_style = profile.pick_style(available, "TABLE_STYLE_PRIORITY")
    figure_style = profile.pick_style(available, "FIGURE_STYLE_PRIORITY")

    tables_styled = figures_styled = no_change_skip = 0
    manual_review: list[dict] = []
    details: list[dict] = []

    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        m = _CAPTION_PATTERN.match(text)
        if not m:
            continue
        kind = m.group(1)
        target = table_style if kind == "表" else figure_style
        current = p.style.name if p.style else None

        if target is None:
            priority = profile.TABLE_STYLE_PRIORITY if kind == "表" else profile.FIGURE_STYLE_PRIORITY
            manual_review.append({
                "idx": idx, "text": text[:80], "kind": kind,
                "reason": f"无可用样式 (tried: {priority})",
            })
            continue
        if current == target:
            no_change_skip += 1
            details.append({
                "idx": idx, "text": text[:60], "kind": kind,
                "style_before": current, "style_after": target, "action": "skip",
            })
            continue
        details.append({
            "idx": idx, "text": text[:60], "kind": kind,
            "style_before": current, "style_after": target, "action": "apply",
        })
        if do_apply:
            p.style = doc.styles[target]
        if kind == "表":
            tables_styled += 1
        else:
            figures_styled += 1

    return {
        "table_style_used": table_style,
        "figure_style_used": figure_style,
        "tables_styled": tables_styled,
        "figures_styled": figures_styled,
        "no_change_skip": no_change_skip,
        "manual_review": manual_review,
        "details": details,
    }


def cmd_caption(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))
    result = _scan_and_apply_caption(doc, profile, do_apply=not args.dry_run)

    wrote_needed = (result["tables_styled"] + result["figures_styled"]) > 0
    bak = _save_with_backup(src, doc, args, wrote_needed=wrote_needed)

    report = {
        "subcommand": "style caption",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": args.dry_run,
        "backup": str(bak) if bak else None,
        **result,
    }
    _emit_report(report, args)
    print(f"[style caption] tables={result['tables_styled']} figures={result['figures_styled']} "
          f"skip={result['no_change_skip']} manual={len(result['manual_review'])} "
          f"table_style={result['table_style_used']!r} figure_style={result['figure_style_used']!r}")
    return 0


# =============================================================================
# 子命令 4: caption number-by-style — number_captions_by_style
# =============================================================================
CN_NUM = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
    "十一": 11, "十二": 12, "十三": 13, "十四": 14, "十五": 15,
}

# 关键词集合(项目无关,中文报告通用)
TABLE_KEYWORDS = (
    "统计表", "配置表", "余缺水量表", "比选表", "拟定表",
    "计算结果表", "结果表", "情况", "对照表", "汇总表",
)
FIG_KEYWORDS = (
    "示意图", "分布图", "布局图", "平面图", "对比图",
    "布置图", "线路图", "范围图", "管线", "输水",
    "结构图", "流程图", "断面图",
)

RE_HAS_TABLE_NUM = re.compile(r"^\s*表\s*\d+\s*[-.–—]\s*\d+")
RE_HAS_FIG_NUM = re.compile(r"^\s*图\s*\d+\s*[-.–—]\s*\d+")
RE_FU_TU = re.compile(r"^\s*附图\s*\d+")
RE_ENUM_PREFIX = re.compile(r"^\s*[（(][一二三四五六七八九十0-9]+[）)]")
RE_SENTENCE_END = re.compile(r"[。；，.;,]$")
RE_CN_CHAPTER = re.compile(r"^([一二三四五六七八九十]+)、")
RE_AR_CHAPTER = re.compile(r"^(\d+)[\s　、.]")
RE_LEADING_NUM = re.compile(r"^(\d+)")


def _get_p_text(p) -> str:
    return "".join(t.text or "" for t in p.iter(qn("w:t")))


def _get_p_style(p) -> Optional[str]:
    pStyle = p.find(".//" + qn("w:pStyle"))
    return pStyle.get(qn("w:val")) if pStyle is not None else None


def _p_has_drawing(p) -> bool:
    return (
        next(p.iter(qn("w:drawing")), None) is not None
        or next(p.iter(qn("w:pict")), None) is not None
    )


def _is_empty_p(p) -> bool:
    return _get_p_text(p).strip() == "" and not _p_has_drawing(p)


def _parse_chapter_from_text(text: str) -> Optional[int]:
    t = text.strip()
    m = RE_LEADING_NUM.match(t)
    if m:
        return int(m.group(1))
    m = RE_CN_CHAPTER.match(t)
    if m:
        return CN_NUM.get(m.group(1))
    return None


def _is_h1_by_text_fallback(text: str) -> bool:
    t = text.strip()
    if len(t) > 30 or len(t) < 2:
        return False
    return bool(RE_CN_CHAPTER.match(t) or RE_AR_CHAPTER.match(t))


def _has_nearby_table(elements, i):
    for offset in (1, 2):
        j = i + offset
        if j >= len(elements):
            break
        el = elements[j]
        tag = el.tag.split("}")[-1]
        if tag == "tbl":
            return True
        if tag == "p" and not _is_empty_p(el):
            return False
    return False


def _has_nearby_drawing(elements, i):
    for offset in (-1, -2, -3, 1, 2, 3):
        j = i + offset
        if not (0 <= j < len(elements)):
            continue
        el = elements[j]
        if el.tag.split("}")[-1] != "p":
            continue
        if _p_has_drawing(el):
            step = 1 if offset > 0 else -1
            blocked = False
            for k in range(i + step, j, step):
                kel = elements[k]
                if kel.tag.split("}")[-1] == "p" and not _is_empty_p(kel) and not _p_has_drawing(kel):
                    blocked = True
                    break
            if not blocked:
                return True
    return False


def _prepend_run_text(p, prefix: str) -> bool:
    for r in p.iter(qn("w:r")):
        t_elems = list(r.iter(qn("w:t")))
        if not t_elems:
            continue
        first_t = t_elems[0]
        first_t.text = prefix + (first_t.text or "")
        first_t.set(qn("xml:space"), "preserve")
        return True
    return False


def _classify_caption_byst(el, text, style, profile: StylesProfile, elements, i) -> Optional[str]:
    if RE_FU_TU.match(text):
        return None
    if RE_HAS_TABLE_NUM.match(text) or RE_HAS_FIG_NUM.match(text):
        return None
    if RE_ENUM_PREFIX.match(text) and not profile.is_table_caption(style) and not profile.is_fig_caption(style):
        return None
    if len(text) >= 60 or RE_SENTENCE_END.search(text):
        return None
    if profile.is_table_caption(style):
        return "table"
    if profile.is_fig_caption(style):
        return "figure"
    has_tbl = _has_nearby_table(elements, i)
    has_draw = _has_nearby_drawing(elements, i)
    has_table_kw = any(kw in text for kw in TABLE_KEYWORDS)
    has_fig_kw = any(kw in text for kw in FIG_KEYWORDS) and "表" not in text
    if has_tbl and (has_table_kw or profile.is_table_caption(style)):
        return "table"
    if has_draw and (has_fig_kw or profile.is_fig_caption(style)):
        return "figure"
    if has_table_kw and has_tbl:
        return "table"
    if has_fig_kw and has_draw:
        return "figure"
    if has_table_kw and not has_fig_kw:
        return "table"
    if has_fig_kw and not has_table_kw:
        return "figure"
    return None


def _process_number_by_style(doc, profile: StylesProfile, dry_run: bool) -> dict:
    body = doc.element.body
    elements = list(body.iterchildren())

    chapter = 0
    tbl_y = fig_y = 0
    numbered: list[dict] = []
    manual_review: list[dict] = []
    chapters_detected: list[int] = []

    for i, el in enumerate(elements):
        if el.tag.split("}")[-1] != "p":
            continue
        text = _get_p_text(el)
        text_strip = text.strip()
        if not text_strip:
            continue
        style = _get_p_style(el)

        # H1 章节识别 (profile pStyle 优先, 文本兜底)
        is_h1 = profile.is_h1(style)
        if not is_h1 and _is_h1_by_text_fallback(text_strip):
            is_h1 = True
        if is_h1:
            ch = _parse_chapter_from_text(text_strip) or (chapter + 1)
            if ch != chapter:
                chapter = ch
                tbl_y = fig_y = 0
                if ch not in chapters_detected:
                    chapters_detected.append(ch)
            continue

        cap_type = _classify_caption_byst(el, text_strip, style, profile, elements, i)
        if cap_type is None:
            continue
        if chapter == 0:
            manual_review.append({"idx": i, "reason": "no-chapter-context", "text_snippet": text_strip[:60]})
            continue

        if cap_type == "table":
            tbl_y += 1
            prefix = profile.format_caption("table", H1=chapter, N=tbl_y)
            number = f"表 {chapter}-{tbl_y}"
        else:
            fig_y += 1
            prefix = profile.format_caption("figure", H1=chapter, N=fig_y)
            number = f"图 {chapter}-{fig_y}"

        if not dry_run:
            ok = _prepend_run_text(el, prefix)
            if not ok:
                manual_review.append({"idx": i, "reason": "no-run-with-text-to-prepend", "text_snippet": text_strip[:60]})
                if cap_type == "table":
                    tbl_y -= 1
                else:
                    fig_y -= 1
                continue

        numbered.append({"idx": i, "type": cap_type, "number": number, "style": style, "text_after": prefix + text_strip})

    summary = {
        "tables_numbered": sum(1 for x in numbered if x["type"] == "table"),
        "figures_numbered": sum(1 for x in numbered if x["type"] == "figure"),
        "manual_review_count": len(manual_review),
        "chapters_detected": chapters_detected,
    }
    return {"summary": summary, "numbered": numbered, "manual_review": manual_review}


def cmd_caption_number_by_style(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))
    result = _process_number_by_style(doc, profile, args.dry_run)
    wrote_needed = (result["summary"]["tables_numbered"] + result["summary"]["figures_numbered"]) > 0
    bak = _save_with_backup(src, doc, args, wrote_needed=wrote_needed)

    report = {
        "subcommand": "caption number-by-style",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": args.dry_run,
        "backup": str(bak) if bak else None,
        **result,
    }
    _emit_report(report, args)
    s = result["summary"]
    print(f"[caption number-by-style] tables={s['tables_numbered']} figures={s['figures_numbered']} "
          f"manual={s['manual_review_count']} chapters={s['chapters_detected']}")
    return 0


# =============================================================================
# 子命令 5: renumber h4-figures — renumber_h4_figures
# =============================================================================
H4_PREFIX_RE = re.compile(r"^\s*\d+(?:[.．]\d+){0,3}[.．]?\s*")
FIG_PREFIX_RE = re.compile(r"^\s*图\s*[\d一二三四五六七八九十百零]+(?:[.．\-–—][\d一二三四五六七八九十百零]+)*\s*")
TBL_PREFIX_RE = re.compile(r"^\s*表\s*[\d一二三四五六七八九十百零]+(?:[.．\-–—][\d一二三四五六七八九十百零]+)*\s*")


def _get_style_name(p) -> str:
    try:
        return (p.style.name or "") if p.style is not None else ""
    except Exception:
        return ""


def _get_text(p) -> str:
    return "".join(r.text or "" for r in p.runs)


def _rewrite_with_prefix(p, new_prefix: str, prefix_re: re.Pattern) -> tuple[str, str]:
    runs = p.runs
    full = "".join(r.text or "" for r in runs)
    stripped = prefix_re.sub("", full, count=1)
    removed_len = len(full) - len(stripped)
    new_text = new_prefix + stripped

    if not runs:
        p.add_run(new_prefix)
        return full, new_text

    cum = 0
    split_idx = None
    split_off = 0
    for i, r in enumerate(runs):
        t = r.text or ""
        if cum + len(t) >= removed_len:
            split_idx = i
            split_off = removed_len - cum
            break
        cum += len(t)

    if split_idx is None:
        for r in runs:
            r.text = ""
        runs[0].text = new_prefix
        return full, new_text

    for i in range(split_idx):
        runs[i].text = ""
    split_text = runs[split_idx].text or ""
    runs[split_idx].text = new_prefix + split_text[split_off:]
    return full, new_text


def _build_h4fig_plan(doc, profile: StylesProfile):
    h1 = h2 = h3 = h4 = 0
    fig_n_in_h2: dict[tuple[int, int], int] = {}
    tbl_n_in_h2: dict[tuple[int, int], int] = {}
    plan_h4, plan_fig, plan_tbl = [], [], []

    for idx, p in enumerate(doc.paragraphs):
        sn = _get_style_name(p)
        if profile.is_h1(sn):
            h1 += 1; h2 = 0; h3 = 0; h4 = 0
        elif profile.is_h2(sn):
            h2 += 1; h3 = 0; h4 = 0
        elif profile.is_h3(sn):
            h3 += 1; h4 = 0
        elif profile.is_h4(sn):
            h4 += 1
            new_prefix = profile.format_heading(4, h1, h2, h3, h4)
            plan_h4.append({"idx": idx, "new_prefix": new_prefix, "style": sn})
        elif profile.is_fig_caption(sn):
            text = _get_text(p)
            if text.lstrip().startswith("图"):
                key = (h1, h2)
                fig_n_in_h2[key] = fig_n_in_h2.get(key, 0) + 1
                n = fig_n_in_h2[key]
                new_prefix = profile.format_caption("figure_h2", H1=h1, H2=h2, N=n)
                plan_fig.append({"idx": idx, "new_prefix": new_prefix, "style": sn})
        elif profile.is_table_caption(sn):
            text = _get_text(p)
            if text.lstrip().startswith("表"):
                key = (h1, h2)
                tbl_n_in_h2[key] = tbl_n_in_h2.get(key, 0) + 1
                n = tbl_n_in_h2[key]
                new_prefix = profile.format_caption("table_h2", H1=h1, H2=h2, N=n)
                plan_tbl.append({"idx": idx, "new_prefix": new_prefix, "style": sn})
    return plan_h4, plan_fig, plan_tbl


def cmd_renumber_h4_figures(args) -> int:
    profile = load_profile(args.profile)
    src = _common_setup(args)
    doc = Document(str(src))
    plan_h4, plan_fig, plan_tbl = _build_h4fig_plan(doc, profile)

    report = {
        "subcommand": "renumber h4-figures",
        "docx": str(src),
        "profile": getattr(profile, "_name", args.profile),
        "dry_run": args.dry_run,
        "h4_count": len(plan_h4),
        "fig_count": len(plan_fig),
        "tbl_count": len(plan_tbl),
        "h4_changes": [],
        "fig_changes": [],
        "tbl_changes": [],
        "backup": None,
    }

    if args.dry_run:
        _emit_report(report, args)
        print(f"[renumber h4-figures] DRY-RUN H4={len(plan_h4)} 图={len(plan_fig)} 表={len(plan_tbl)}")
        return 0

    paragraphs = doc.paragraphs
    for item in plan_h4:
        p = paragraphs[item["idx"]]
        old, new = _rewrite_with_prefix(p, item["new_prefix"], H4_PREFIX_RE)
        report["h4_changes"].append({"idx": item["idx"], "old": old[:60], "new": new[:60]})
    for item in plan_fig:
        p = paragraphs[item["idx"]]
        old, new = _rewrite_with_prefix(p, item["new_prefix"], FIG_PREFIX_RE)
        report["fig_changes"].append({"idx": item["idx"], "old": old[:60], "new": new[:60]})
    for item in plan_tbl:
        p = paragraphs[item["idx"]]
        old, new = _rewrite_with_prefix(p, item["new_prefix"], TBL_PREFIX_RE)
        report["tbl_changes"].append({"idx": item["idx"], "old": old[:60], "new": new[:60]})

    bak = _save_with_backup(src, doc, args, wrote_needed=bool(plan_h4 or plan_fig or plan_tbl))
    report["backup"] = str(bak) if bak else None
    _emit_report(report, args)
    print(f"[renumber h4-figures] H4={report['h4_count']} 图={report['fig_count']} 表={report['tbl_count']}")
    return 0


# =============================================================================
# 报告 emitter (统一 --report path 写盘)
# =============================================================================
def _emit_report(report: dict, args):
    """如果 --report 提供, 写 JSON 到该路径."""
    rp = getattr(args, "report", None)
    if rp:
        rp = Path(rp)
        rp.parent.mkdir(parents=True, exist_ok=True)
        rp.write_text(json.dumps(report, ensure_ascii=False, indent=2, default=str), encoding="utf-8")


# =============================================================================
# CLI register & main
# =============================================================================
def _add_common_args(p: argparse.ArgumentParser):
    p.add_argument("docx", type=Path, help="目标 docx 路径")
    p.add_argument("--dry-run", action="store_true", help="只规划不写入")
    p.add_argument("--no-backup", action="store_true", help="跳过备份(慎用)")
    p.add_argument("--report", type=Path, default=None, help="写 JSON 报告到此路径")
    p.add_argument("--profile", type=str, default=None,
                   help="styles_registry profile (zdwp / eco-flow / generic / ...)")


def register(subparsers):
    """register 5 subcommands into a doctools CLI subparsers.

    嵌套命名:
        style body / style table / style caption  (3 个 style 下子命令)
        caption number-by-style    (shared `caption` group with caption.py/captions.py)
        renumber h4-figures         (shared `renumber` group with renumber.py)
    """
    from ._dispatch import get_or_add_group, get_or_add_subparsers

    # style group (new; no other module owns it)
    style_p = get_or_add_group(subparsers, "style", "样式批改(body/table/caption)")
    style_sub = get_or_add_subparsers(style_p, dest="style_cmd")
    existing_s = getattr(style_sub, "choices", {}) or {}

    if "body" not in existing_s:
        body_p = style_sub.add_parser("body", help="启发式套 H1/H2/H3/Title/正文 样式")
        _add_common_args(body_p)
        body_p.add_argument(
            "--force",
            action="store_true",
            help="绕过 caption-risk guard (不推荐: profile/docx 漂移时会把正文套成 caption 自动编号)",
        )
        body_p.add_argument(
            "--template",
            choices=["std"],
            default=None,
            help=(
                "注入 ZDWP 样式骨架到 docx (跳过已存在). 'std' = 磐安 v3 抽出的"
                " 9 段样式 (ZDWP正文/图名/表名/附表/题目0/1/作者/封面日期/表格内容). "
                "用例: 全裸 docx (无 ZDWP 系列) 先 --template std 打底, 再跑 "
                "`styleset restore` 不抛 'ZDWP正文 not found'."
            ),
        )
        body_p.set_defaults(func=cmd_body)

    if "table" not in existing_s:
        table_p = style_sub.add_parser("table", help="cell 段样式 + 表全网格 + cell 居中")
        _add_common_args(table_p)
        table_p.add_argument("--style-id", default=None,
                             help="cell 段目标 styleId (默认走 profile.TABLE_CELL_STYLE_ID)")
        table_p.set_defaults(func=cmd_table)

    if "caption" not in existing_s:
        caption_p = style_sub.add_parser("caption", help="给 ^(表|图)X.Y-Z 段套对样式")
        _add_common_args(caption_p)
        caption_p.set_defaults(func=cmd_caption)

    # caption 子命令 (pStyle-aware 编号补) — shared parent with caption.py/captions.py
    cap_p = get_or_add_group(subparsers, "caption", "caption ops (number / pair / number-by-style)")
    cap_sub = get_or_add_subparsers(cap_p, dest="caption_target")
    existing_c = getattr(cap_sub, "choices", {}) or {}
    if "number-by-style" not in existing_c:
        nbs_p = cap_sub.add_parser("number-by-style",
                                    help="pStyle-aware 补 '表 X-Y' / '图 X-Y' 编号")
        _add_common_args(nbs_p)
        nbs_p.set_defaults(func=cmd_caption_number_by_style)

    # renumber 子命令 — shared parent with renumber.py
    rn_p = get_or_add_group(subparsers, "renumber", "renumber headings + caption numbers")
    rn_sub = get_or_add_subparsers(rn_p, dest="renumber_target")
    existing_r = getattr(rn_sub, "choices", {}) or {}
    if "h4-figures" not in existing_r:
        h4_p = rn_sub.add_parser("h4-figures",
                                  help="重派 H4 + 图/表 caption (H1.H2-N)")
        _add_common_args(h4_p)
        h4_p.set_defaults(func=cmd_renumber_h4_figures)


def main(argv: list[str] | None = None) -> int:
    """Module entrypoint — 5 subcommands."""
    parser = argparse.ArgumentParser(
        prog="doctools-styles",
        description="docx style / caption / renumber group · profile-driven 样式批改 (distilled from qual-supply)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)
    register(sub)
    args = parser.parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    sys.exit(main())
