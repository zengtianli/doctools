#!/usr/bin/env python3
# distilled from qual-supply/scripts/freeze_heading_numbers.py (2026-05-25 W1)
# -*- coding: utf-8 -*-
"""
freeze_heading_numbers.py
==========================

单功能描述
----------
把 docx 内 Heading 1-4 段的"自动编号"按 Word 渲染算法**写死成段文本前缀**(`1 概述`/`1.1 项目背景`),
并(可选)从 word/styles.xml 移除 heading 1-4 样式自带的 `<w:numPr>`, **断开自动编号链**。
合稿时 Word 不再重算编号、不再因 numId 合并而污染。

触发场景
--------
- 多份 docx 合稿前, H 段自动编号来自共享 numId, 合稿后 Word 重算导致编号乱跳
- 主报告"准成品"阶段, 需把编号"冻结"为字面文本以便人审 + 终稿稳定
- TOC / 导航大纲仍依赖 Heading 1-4 outlineLvl, 仅去 numPr 不影响大纲

CLI
---
    python3 scripts/freeze_heading_numbers.py <docx_path> \\
        [--levels 1,2,3,4] [--unlink-style] [--no-unlink-style] \\
        [--dry-run] [--no-backup] [--report <json>]

默认行为
--------
- `--levels 1,2,3,4` (全级别)
- `--unlink-style` 启用 (= 同时改 styles.xml 移除 heading 1-4 的 numPr)
- 自动 `--backup` 写 `.bak-N-YYYY-MM-DD.docx` (除非 `--no-backup`)
- 写前 `lsof` 自检 Word/WPS 占用

算法
----
1. 解析 word/numbering.xml 拿 numId=4 对应 abstractNum 各 ilvl 的 lvlText 模板 (`%1`/`%1.%2`/...)
   实际本 docx 各 ilvl 模板 = decimal + 起始 1 + 默认 restart
2. 按段物理顺序遍历, 维护 h1/h2/h3/h4 计数器 (高级别自增时低级别清 0)
3. 按 lvlText 模板把 `%1/%2/%3/%4` 替换成当前计数, 得到段编号 (如 `1.2.3`)
4. 找段第一个非空 run, 在其首个 `<w:t>` 前 prepend `"<编号> "` (空格半角)
5. 已含 `^\\d+(\\.\\d+)*\\s` 前缀的段 skip (避免重复 prepend)
6. (可选) 移除 word/styles.xml 内 heading 1-4 的 numPr

不许做
------
- 改 H 段以外的段
- 改 numbering.xml (只读)
- 改 heading 1-4 以外的样式 (heading 5-9 / Title / 标题 / TOC / Normal ...)
- 段级赋值 `paragraph.text = ...` (会丢 bold/字号)
- 用 sed 改 XML
- commit / push
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
NS = {"w": W_NS}

TARGET_STYLE_NAMES = {"heading 1", "heading 2", "heading 3", "heading 4"}
PREFIX_PATTERN = re.compile(r"^\d+(\.\d+)*\s")
STYLE_TO_LEVEL = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
}


# ----------------------------- helpers ----------------------------- #

def find_next_backup(docx_path: Path) -> Path:
    today = date.today().isoformat()
    stem = docx_path.stem
    parent = docx_path.parent
    n = 1
    while True:
        cand = parent / f"{stem}.bak-{n}-{today}.docx"
        if not cand.exists():
            return cand
        n += 1


def lsof_check(docx_path: Path) -> str | None:
    try:
        out = subprocess.run(
            ["lsof", str(docx_path)],
            capture_output=True, text=True, timeout=5,
        )
        if out.returncode == 0 and out.stdout.strip():
            return out.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


# ----------------------------- numbering.xml ----------------------------- #

def read_lvltext_templates(docx_path: Path, num_id: str = "4") -> dict[int, dict]:
    """读 numbering.xml, 返回 {ilvl: {"lvlText": "%1.%2", "start": 1, "numFmt": "decimal"}}"""
    with zipfile.ZipFile(docx_path, "r") as z:
        try:
            xml_bytes = z.read("word/numbering.xml")
        except KeyError:
            return {}
    root = etree.fromstring(xml_bytes)
    # find num with numId
    abs_id = None
    for n in root.findall(f"{W}num"):
        if n.get(W + "numId") == num_id:
            ani = n.find(f"{W}abstractNumId")
            if ani is not None:
                abs_id = ani.get(W + "val")
            break
    if abs_id is None:
        return {}
    out: dict[int, dict] = {}
    for an in root.findall(f"{W}abstractNum"):
        if an.get(W + "abstractNumId") != abs_id:
            continue
        for lvl in an.findall(f"{W}lvl"):
            ilvl_s = lvl.get(W + "ilvl")
            try:
                ilvl = int(ilvl_s)
            except (TypeError, ValueError):
                continue
            lt = lvl.find(f"{W}lvlText")
            st = lvl.find(f"{W}start")
            nf = lvl.find(f"{W}numFmt")
            out[ilvl] = {
                "lvlText": lt.get(W + "val") if lt is not None else None,
                "start": int(st.get(W + "val")) if st is not None else 1,
                "numFmt": nf.get(W + "val") if nf is not None else "decimal",
            }
        break
    return out


def render_number(template: str | None, counters: list[int], level: int) -> str:
    """把 %1/%2/%3/%4 替换为 counters[0..level-1]; 默认 fallback 用 . 连接"""
    if template:
        s = template
        for i in range(1, 5):
            s = s.replace(f"%{i}", str(counters[i - 1]) if i <= level else "")
        return s
    # fallback
    return ".".join(str(c) for c in counters[:level])


# ----------------------------- styles.xml patch ----------------------------- #

def remove_heading_numpr_in_styles(
    docx_path: Path, levels: set[int], dry_run: bool
) -> tuple[list[dict], bytes | None]:
    """从 word/styles.xml 移除指定 heading 级别样式的 <w:numPr>"""
    target_names = {f"heading {lv}" for lv in levels}
    with zipfile.ZipFile(docx_path, "r") as z:
        styles_bytes = z.read("word/styles.xml")
    root = etree.fromstring(styles_bytes)
    removed: list[dict] = []
    for s in root.findall(f".//{W}style"):
        n = s.find(f"{W}name")
        nm = n.get(W + "val") if n is not None else None
        if nm is None or nm not in target_names:
            continue
        pPr = s.find(f"{W}pPr")
        if pPr is None:
            continue
        numPr = pPr.find(f"{W}numPr")
        if numPr is None:
            continue
        sid = s.get(W + "styleId")
        # capture detail
        numId_el = numPr.find(f"{W}numId")
        ilvl_el = numPr.find(f"{W}ilvl")
        removed.append({
            "name": nm,
            "styleId": sid,
            "numId": numId_el.get(W + "val") if numId_el is not None else None,
            "ilvl": ilvl_el.get(W + "val") if ilvl_el is not None else None,
        })
        if not dry_run:
            pPr.remove(numPr)
    if dry_run or not removed:
        return removed, None
    new_bytes = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    return removed, new_bytes


def write_styles_xml(docx_path: Path, new_styles_bytes: bytes) -> None:
    """把新 styles.xml 写回 docx (其他文件原样保留)"""
    tmp = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for it in zin.infolist():
            if it.filename == "word/styles.xml":
                zout.writestr(it, new_styles_bytes)
            else:
                zout.writestr(it, zin.read(it.filename))
    shutil.move(str(tmp), str(docx_path))


# ----------------------------- paragraph prefix ----------------------------- #

def prepend_to_first_run(paragraph, prefix: str) -> bool:
    """在段第一个非空 run 的首个 <w:t> 前 prepend prefix; 返回是否成功"""
    p_el = paragraph._p
    # 找第一个含 w:t 的 run
    for r in p_el.findall(qn("w:r")):
        # 跳过纯 fldChar / numbering instr 等 run
        t_list = r.findall(qn("w:t"))
        if not t_list:
            continue
        # 找首个 w:t (含或不含文本都行, 要 prepend 到最前面)
        first_t = t_list[0]
        # 保留 xml:space="preserve" (prefix 末尾有空格)
        existing = first_t.text or ""
        first_t.text = prefix + existing
        first_t.set(
            "{http://www.w3.org/XML/1998/namespace}space", "preserve"
        )
        return True
    return False


def freeze_headings(
    doc: Document, lvltext: dict[int, dict], levels: set[int], apply: bool
) -> dict:
    """扫所有段, 按物理顺序给 Heading 1-4 编号并 prepend 到段文本"""
    counters = [0, 0, 0, 0]  # h1, h2, h3, h4
    headings_seen = {1: 0, 2: 0, 3: 0, 4: 0}
    frozen_count = 0
    skip_already_prefixed = 0
    skip_empty_run = 0
    samples = {1: [], 2: [], 3: [], 4: []}
    manual_review: list[dict] = []

    for idx, p in enumerate(doc.paragraphs):
        style_name = p.style.name
        level = STYLE_TO_LEVEL.get(style_name)
        if level is None:
            continue
        if level not in levels:
            continue

        # update counters
        counters[level - 1] += 1
        for j in range(level, 4):
            counters[j] = 0
        headings_seen[level] += 1

        text = p.text.strip()
        # skip if already prefixed
        if PREFIX_PATTERN.match(text):
            skip_already_prefixed += 1
            manual_review.append({
                "idx": idx,
                "level": level,
                "reason": "already_prefixed",
                "text": text[:60],
            })
            continue

        # render number from lvlText (ilvl = level-1)
        ilvl = level - 1
        tpl_info = lvltext.get(ilvl, {})
        template = tpl_info.get("lvlText")
        number = render_number(template, counters, level)
        prefix = f"{number} "

        if apply:
            ok = prepend_to_first_run(p, prefix)
            if not ok:
                skip_empty_run += 1
                manual_review.append({
                    "idx": idx,
                    "level": level,
                    "reason": "no_run_with_t",
                    "text": text[:60],
                })
                continue
        frozen_count += 1

        if len(samples[level]) < 3:
            samples[level].append({
                "idx": idx,
                "level": level,
                "number": number,
                "before": text[:50],
                "after": (prefix + text)[:60],
            })

    return {
        "headings_seen": headings_seen,
        "frozen_count": frozen_count,
        "skip_already_prefixed": skip_already_prefixed,
        "skip_empty_run": skip_empty_run,
        "samples": samples,
        "manual_review": manual_review,
    }


# ----------------------------- main ----------------------------- #

def parse_levels(s: str) -> set[int]:
    out: set[int] = set()
    for tok in s.split(","):
        tok = tok.strip()
        if not tok:
            continue
        try:
            v = int(tok)
        except ValueError:
            raise argparse.ArgumentTypeError(f"invalid level: {tok!r}")
        if v not in (1, 2, 3, 4):
            raise argparse.ArgumentTypeError(f"level out of range 1-4: {v}")
        out.add(v)
    if not out:
        raise argparse.ArgumentTypeError("no levels specified")
    return out


def main() -> int:
    p = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    p.add_argument("docx", type=Path)
    p.add_argument("--levels", type=parse_levels, default=parse_levels("1,2,3,4"))
    p.add_argument("--unlink-style", dest="unlink_style", action="store_true", default=True)
    p.add_argument("--no-unlink-style", dest="unlink_style", action="store_false")
    p.add_argument("--num-id", default="4", help="numbering.xml numId (default 4)")
    p.add_argument("--dry-run", action="store_true")
    p.add_argument("--no-backup", action="store_true")
    p.add_argument("--report", type=Path, default=None)
    args = p.parse_args()

    if not args.docx.exists():
        print(f"[ERR] 文件不存在: {args.docx}", file=sys.stderr)
        return 2

    if not args.dry_run:
        occ = lsof_check(args.docx)
        if occ:
            print(f"[ERR] 文件被占用 (Word/WPS 在开?):\n{occ}", file=sys.stderr)
            return 3

    # Step 1: read lvlText templates
    lvltext = read_lvltext_templates(args.docx, num_id=args.num_id)
    print(f"[INFO] numbering.xml numId={args.num_id} lvlText 模板:")
    for ilvl in (0, 1, 2, 3):
        info = lvltext.get(ilvl)
        if info is None:
            print(f"  ilvl={ilvl}: <MISSING> (fallback %1.%2.%3.%4)")
        else:
            print(
                f"  ilvl={ilvl}: lvlText={info['lvlText']!r} "
                f"start={info['start']} numFmt={info['numFmt']}"
            )

    # Step 2-3: freeze H paragraph prefixes
    doc = Document(str(args.docx))
    freeze_result = freeze_headings(
        doc, lvltext, args.levels, apply=not args.dry_run
    )

    print(f"[INFO] H 段统计 (seen):")
    for lv in (1, 2, 3, 4):
        print(f"  H{lv}: {freeze_result['headings_seen'][lv]}")
    print(f"[INFO] frozen_count = {freeze_result['frozen_count']}")
    print(f"[INFO] skip_already_prefixed = {freeze_result['skip_already_prefixed']}")
    print(f"[INFO] skip_empty_run = {freeze_result['skip_empty_run']}")

    # samples
    for lv in (1, 2, 3, 4):
        for s in freeze_result["samples"][lv]:
            print(f"  H{lv} idx={s['idx']:4d} | {s['before']!r:50s} -> {s['after']!r}")

    # Step 4: backup + write doc
    backup_path = None
    wrote_doc = False
    if args.dry_run:
        print("[DRY-RUN] 不写文件")
    elif freeze_result["frozen_count"] == 0 and not args.unlink_style:
        print("[INFO] 无需写文件 (无 prefix 改动且 --no-unlink-style)")
    else:
        if not args.no_backup:
            backup_path = find_next_backup(args.docx)
            shutil.copy2(args.docx, backup_path)
            print(f"[INFO] 备份 -> {backup_path.name}")
        if freeze_result["frozen_count"] > 0:
            doc.save(str(args.docx))
            wrote_doc = True
            print(f"[OK] H 段 prefix 写入 {freeze_result['frozen_count']} 处")

    # Step 5: unlink heading 1-4 numPr in styles.xml
    style_removed: list[dict] = []
    if args.unlink_style:
        removed, new_bytes = remove_heading_numpr_in_styles(
            args.docx, args.levels, dry_run=args.dry_run
        )
        style_removed = removed
        if args.dry_run:
            print(f"[DRY-RUN] 将移除 styles.xml numPr: {len(removed)} 处")
            for d in removed:
                print(f"  {d['name']} (styleId={d['styleId']}) numId={d['numId']} ilvl={d['ilvl']}")
        else:
            if new_bytes:
                write_styles_xml(args.docx, new_bytes)
                print(f"[OK] styles.xml 移除 {len(removed)} 个 numPr (heading 1-4)")
                for d in removed:
                    print(f"  - {d['name']} (styleId={d['styleId']}) numId={d['numId']} ilvl={d['ilvl']}")
            else:
                print("[INFO] styles.xml 无需修改 (heading 1-4 无 numPr)")

    report = {
        "docx": str(args.docx.resolve()),
        "dry_run": args.dry_run,
        "levels": sorted(args.levels),
        "unlink_style": args.unlink_style,
        "lvltext_templates": {
            str(k): v for k, v in lvltext.items()
        },
        "headings_seen": freeze_result["headings_seen"],
        "headings_frozen": freeze_result["frozen_count"],
        "levels_count": freeze_result["headings_seen"],
        "skip_already_prefixed": freeze_result["skip_already_prefixed"],
        "skip_empty_run": freeze_result["skip_empty_run"],
        "samples": freeze_result["samples"],
        "manual_review": freeze_result["manual_review"],
        "style_numPr_removed": style_removed,
        "backup": str(backup_path) if backup_path else None,
        "wrote_doc": wrote_doc,
    }

    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(
            json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(f"[INFO] report -> {args.report}")

    return 0


# ---------------- pipeline adapter ----------------
def apply_path(docx_path, args=None) -> dict:
    """pipeline: freeze heading 1-4 编号 + unlink style numPr"""
    levels = getattr(args, "freeze_levels", None) if args else None
    if levels is None:
        levels = {1, 2, 3, 4}
    elif isinstance(levels, str):
        levels = parse_levels(levels)
    unlink_style = getattr(args, "unlink_style", True) if args else True
    num_id = getattr(args, "num_id", "4") if args else "4"
    dry = bool(getattr(args, "dry_run", False)) if args else False
    p = Path(docx_path)
    lvltext = read_lvltext_templates(p, num_id=num_id)
    doc = Document(str(p))
    freeze_result = freeze_headings(doc, lvltext, levels, apply=not dry)
    if not dry and freeze_result["frozen_count"] > 0:
        doc.save(str(p))
    style_removed: list[dict] = []
    if unlink_style:
        removed, new_bytes = remove_heading_numpr_in_styles(p, levels, dry_run=dry)
        style_removed = removed
        if not dry and new_bytes:
            write_styles_xml(p, new_bytes)
    return {
        "changed": freeze_result["frozen_count"] + len(style_removed),
        "headings_frozen": freeze_result["frozen_count"],
        "headings_seen": freeze_result["headings_seen"],
        "style_numPr_removed": len(style_removed),
        "skip_already_prefixed": freeze_result["skip_already_prefixed"],
    }


if __name__ == "__main__":
    sys.exit(main())
