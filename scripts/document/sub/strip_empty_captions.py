#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
strip_empty_captions.py
=======================

单功能描述
----------
从 docx 移除**样式为 caption (图名称 / 表格标题 / 图名 / 表名 / Caption 等)
且段落文本完全为空**的段。模板里常残留 "图前 placeholder 空段" 这种结构性
噪音, 影响 audit / renumber 等下游脚本判定。

算法
----
1. ``python-docx`` 打开 docx
2. 遍历 ``doc.paragraphs``, 识别 caption 样式 (style.name 大小写不敏感匹配:
   "图名称" / "图名" / "表格标题" / "表标题" / "表名" / 含 "caption" )
3. 段文本 ``.text.strip() == ""`` **且** 无 ``<w:drawing>`` / ``<w:pict>`` /
   ``<w:object>`` (inline 图/对象) **且** 无 ``<w:r>`` 含非空 ``<w:t>`` →
   strict 空段, 标记删除
4. 经 lxml 真删 (``parent.remove(p_elem)``), 不是设隐藏
5. ``doc.save``

谨慎边界
--------
- **不删 inline drawing**: 段含 ``<w:drawing>`` (浮动图也含) 一律保留
- **不删非 strict 空**: 段含 inline image / 含非空 run / 含表格引用都保留
- **不动正文段**: 仅处理 caption 样式段
- 同一 caption 段紧贴下一个非空 caption 段时 (典型 "图占位空段 + 图名" 配对),
  本脚本会删 placeholder 段; 这是预期 (用户已确认 strict 空段为模板残留)

CLI
---
    python3 sub/strip_empty_captions.py <docx_path> \\
        [-o OUT | --inplace] [--dry-run] [--no-backup] [--report <json>]

默认行为
--------
- 默认 ``--inplace`` (留 ``.bak-N-YYYY-MM-DD.docx``)
- ``-o OUT`` 写到新路径, 不动原文件 + 不留 bak
- ``--dry-run`` 列将删的段 (样式 + 段索引 + 上下文 1 行), 不写

不做
----
- 不删非空 caption 段
- 不删含 inline drawing 的 caption 段 (图本身不动)
- 不改 caption 编号 (caption renumber 的活)
- 不动正文 / heading 段
"""
from __future__ import annotations

import argparse
import datetime
import json
import shutil
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("[ERR] 缺 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(2)

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{NS}}}"

# 触发删除的样式名关键词 (大小写不敏感, substring 匹配)
CAPTION_STYLE_KEYWORDS = (
    "图名称", "图名", "表格标题", "表标题", "表名", "caption",
)


def lsof_check(p: Path) -> str | None:
    try:
        r = subprocess.run(
            ["lsof", str(p)], capture_output=True, text=True, timeout=5
        )
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def find_next_backup(docx_path: Path) -> Path:
    today = datetime.date.today().isoformat()
    n = 1
    while True:
        cand = docx_path.with_name(
            f"{docx_path.stem}.bak-{n}-{today}{docx_path.suffix}"
        )
        if not cand.exists():
            return cand
        n += 1


def is_caption_style(style_name: str | None) -> bool:
    if not style_name:
        return False
    low = style_name.lower()
    for kw in CAPTION_STYLE_KEYWORDS:
        if kw.lower() in low:
            return True
    return False


def is_strict_empty_paragraph(p_elem) -> bool:
    """段是 strict 空: 无 <w:t> 含文字 且 无 <w:drawing>/<w:pict>/<w:object>."""
    # 任何 inline 图/对象都视为非空 (内容保留)
    for tag in (f"{W}drawing", f"{W}pict", f"{W}object"):
        if p_elem.find(f".//{tag}") is not None:
            return False
    # 任何 <w:t> 含非空文本视为非空
    for t in p_elem.findall(f".//{W}t"):
        if t.text and t.text.strip():
            return False
    # 任何 <w:tab> / <w:br> / <w:noBreakHyphen> 也不算文本 (这些段就是结构噪音);
    # 但 <w:sym> / <w:fldChar> 含字段 (如 SEQ 编号) 不算文本但可能是有意义的占位 —
    # 保守起见, 段含 <w:fldChar> (字段) 视为非空, 保留
    if p_elem.find(f".//{W}fldChar") is not None:
        return False
    if p_elem.find(f".//{W}sym") is not None:
        return False
    return True


def scan_empty_captions(doc) -> list[dict]:
    """扫: 返回每个待删段的 {idx, style, prev_text, next_text}."""
    hits: list[dict] = []
    paragraphs = list(doc.paragraphs)
    for i, p in enumerate(paragraphs):
        try:
            style_name = p.style.name if p.style is not None else None
        except (AttributeError, KeyError):
            style_name = None
        if not is_caption_style(style_name):
            continue
        if not is_strict_empty_paragraph(p._element):
            continue
        prev_text = paragraphs[i - 1].text.strip()[:40] if i > 0 else ""
        next_text = paragraphs[i + 1].text.strip()[:40] if i + 1 < len(paragraphs) else ""
        hits.append({
            "idx": i,
            "style": style_name,
            "prev_text": prev_text,
            "next_text": next_text,
        })
    return hits


def delete_marked(doc, hits: list[dict]) -> int:
    """按段索引删除 (从后往前以保索引有效)."""
    paragraphs = list(doc.paragraphs)
    idx_set = {h["idx"] for h in hits}
    deleted = 0
    for i in sorted(idx_set, reverse=True):
        if i >= len(paragraphs):
            continue
        p_elem = paragraphs[i]._element
        parent = p_elem.getparent()
        if parent is None:
            continue
        parent.remove(p_elem)
        deleted += 1
    return deleted


def main() -> int:
    parser = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("docx", type=Path)
    mx = parser.add_mutually_exclusive_group()
    mx.add_argument("-o", "--output", type=Path, default=None,
                    help="写到新路径(不动原文件,不留 bak)")
    mx.add_argument("--inplace", action="store_true", default=True,
                    help="原地改写(默认),自动留 .bak-N-YYYY-MM-DD")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--no-backup", action="store_true")
    parser.add_argument("--report", type=Path, default=None)
    args = parser.parse_args()

    if not args.docx.exists():
        print(f"[ERR] 文件不存在: {args.docx}", file=sys.stderr)
        return 2

    inplace = args.output is None
    if not args.dry_run and inplace:
        occ = lsof_check(args.docx)
        if occ:
            print(f"[ERR] 文件被占用 (Word/WPS 在开?), 立即停止:\n{occ}", file=sys.stderr)
            return 3

    print(f"[INFO] 扫描 {args.docx.name}")
    doc = Document(str(args.docx))
    hits = scan_empty_captions(doc)
    print(f"  [scan] empty_caption_paragraphs={len(hits)}")

    report = {
        "docx": str(args.docx.resolve()),
        "dry_run": args.dry_run,
        "inplace": inplace,
        "output": str(args.output.resolve()) if args.output else None,
        "backup": None,
        "wrote": False,
        "hits": hits,
        "deleted": 0,
    }

    if not hits:
        print("[INFO] 无空 caption 段, 不写")
        if args.report:
            args.report.parent.mkdir(parents=True, exist_ok=True)
            args.report.write_text(
                json.dumps(report, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            print(f"[INFO] report -> {args.report}")
        return 0

    if args.dry_run:
        print("[DRY-RUN] 将删除以下空 caption 段:")
        for h in hits[:50]:
            print(f"  - idx={h['idx']:>4}  style={h['style']!r:<20}  "
                  f"prev='{h['prev_text']}'  next='{h['next_text']}'")
        if len(hits) > 50:
            print(f"  ... 共 {len(hits)} 段 (仅显示前 50)")
        if args.report:
            args.report.parent.mkdir(parents=True, exist_ok=True)
            args.report.write_text(
                json.dumps(report, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            print(f"[INFO] report -> {args.report}")
        return 0

    if inplace and not args.no_backup:
        bak = find_next_backup(args.docx)
        shutil.copy2(args.docx, bak)
        report["backup"] = str(bak)
        print(f"[INFO] 备份 -> {bak.name}")

    deleted = delete_marked(doc, hits)
    report["deleted"] = deleted

    out_path = args.output if args.output else args.docx
    doc.save(str(out_path))
    report["wrote"] = True
    print(f"[OK] 删除 {deleted} 个空 caption 段, 写回 {out_path.name}")

    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(
            json.dumps(report, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        print(f"[INFO] report -> {args.report}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
