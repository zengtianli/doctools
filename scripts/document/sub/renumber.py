#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""renumber.py — 标题重编号家族二合一（2026-07-31 家族折叠）

子命令 ↔ 原脚本（函数体逐字搬移；模块级 apply/main 改名 _<sub> 后缀；
seq 原 from renumber_headings import 的 4 个符号变同文件内引用；
make_backup 机制收敛 _cli_common（同构：<stem>.bak-N-YYYY-MM-DD.docx 自增+copy2））：

    headings  ← renumber_headings.py（plan_renumber / rewrite_paragraph_with_prefix 等
                                      公有名保留：health.py safe-fix 与 seq 段靠它）
    seq       ← renumber_headings_seq.py（doc_dispatch.py do_renum 的标题号引擎）

各子命令 CLI 与原独立脚本逐字一致：python3 sub/renumber.py <sub> <docx> …。
退役原件在 ~/.Trash/consolidation-20260731/renumber/（含 MANIFEST.md）。
"""
from __future__ import annotations

# ── surgical 收口：python-docx 存盘只重写点名的部件（炸开面 60→1）─────────────
import sys as _sys
from pathlib import Path as _Path
_sys.path.append(str(_Path(__file__).resolve().parents[3] / "lib"))
import docx_safe_save  # noqa: E402,F401  详见 lib/docx_safe_save.py

# sub/ 自身进 sys.path —— docx_cli 的 _dispatch 用 spec_from_file_location 加载,
# 不带脚本目录, 裸 import _cli_common 会 ImportError (append 不是 insert(0))
_sys.path.append(str(_Path(__file__).resolve().parent))
import _cli_common as _cc  # noqa: E402  备份机制 SSOT

import argparse  # noqa: E402
import json  # noqa: E402
import re  # noqa: E402
import shutil  # noqa: E402,F401
import sys  # noqa: E402
from datetime import date  # noqa: E402,F401
from pathlib import Path  # noqa: E402
from typing import Optional  # noqa: E402

try:
    from docx import Document
    from docx.document import Document as _DocType  # noqa: F401
except ImportError:
    print("ERROR: python-docx 未安装,请 pip install python-docx", file=sys.stderr)
    sys.exit(2)


# ══════════ headings ← renumber_headings.py ══════════

# --------- 编号剥离正则 ----------
# heading 旧编号: 阿拉伯数字点串 / 第X章 / 一、 / 1、 等
HEADING_PREFIX_RE = re.compile(
    r"^\s*(?:\d+(?:[.．。]\d+)+[\.．。]?"      # 1.2 / 1.2.3 (多级)
    r"|\d+[\.．。、,]"                          # 1. / 1、
    r"|\d+(?=\s)"                              # 单数字后跟空格
    r"|第\s*[一二三四五六七八九十百零\d]+\s*章"
    r"|[一二三四五六七八九十百零]+\s*[、,])\s*"
)

# 表名旧前缀: "表 X-Y" / "表 X.Y" / "表 X" / 中文数字 / 全角空格 / em-dash 等
TABLE_PREFIX_RE = re.compile(
    r"^\s*表\s*[\d一二三四五六七八九十百零]+\s*"
    r"[-–—―\.．]?\s*[\d一二三四五六七八九十百零]*\s*"
)


def get_style_name(p) -> str:
    try:
        return (p.style.name or "") if p.style is not None else ""
    except Exception:
        return ""


def get_paragraph_text(p) -> str:
    """段文本: 拼所有 run 的 text (与 paragraph.text 等价但不依赖 attribute)."""
    return "".join(r.text or "" for r in p.runs)


def strip_old_prefix(text: str, kind: str) -> str:
    """剥旧编号. kind = 'heading' | 'table'."""
    if kind == "heading":
        return HEADING_PREFIX_RE.sub("", text, count=1)
    if kind == "table":
        return TABLE_PREFIX_RE.sub("", text, count=1)
    return text


def rewrite_paragraph_with_prefix(p, new_prefix: str) -> None:
    """在段 run[0] 改 <w:t> text, 替换旧编号为 new_prefix.

    new_prefix 末尾应已含分隔空格 (例如 "3.2.1 " 或 "表 3-1 ").
    """
    runs = p.runs
    if not runs:
        # 空段(无 run) — 创建一个 run 写入 prefix
        p.add_run(new_prefix)
        return

    # 收集所有 run text, 找第一个非空 run
    first_idx = None
    for i, r in enumerate(runs):
        if (r.text or "") != "":
            first_idx = i
            break
    if first_idx is None:
        # 全空 run, 写到 run[0]
        first_idx = 0

    # 当前 full text
    full_text = "".join(r.text or "" for r in runs)

    # 判断 kind: heading / table by prefix in new_prefix
    if new_prefix.startswith("表"):
        stripped = strip_old_prefix(full_text, "table")
    else:
        stripped = strip_old_prefix(full_text, "heading")

    # 计算: full_text 前面被剥掉了多少字符
    removed_len = len(full_text) - len(stripped)

    # 现在要把 "new_prefix + 剩余正文" 装回 runs, 尽量只改前面的 runs.
    # 策略:
    #   - 计算每个 run 在 full_text 的 [start, end)
    #   - 找到 removed_len 落在哪个 run (split run)
    #   - run[0..split-1].text = ""  (这些被旧编号覆盖)
    #   - run[split].text = new_prefix + run[split].text[removed_len - cumulative:]
    #     -- 即在 split run 的合适偏移处, 砍掉前缀, 拼新编号
    #   - run[split+1..] 保持不变
    new_text_for_first_nonempty = None

    cum = 0
    split_run_idx = None
    split_offset = 0  # 在 split run 内, 旧编号末尾的偏移
    for i, r in enumerate(runs):
        t = r.text or ""
        if cum + len(t) >= removed_len:
            split_run_idx = i
            split_offset = removed_len - cum
            break
        cum += len(t)

    if split_run_idx is None:
        # removed_len 超过全文 (理论上不会), 全清掉重写
        for r in runs:
            r.text = ""
        runs[0].text = new_prefix
        return

    # 清空 split 之前所有 run
    for i in range(split_run_idx):
        runs[i].text = ""

    # split run: 砍掉前 split_offset 字符, 前面拼 new_prefix
    split_run_text = runs[split_run_idx].text or ""
    runs[split_run_idx].text = new_prefix + split_run_text[split_offset:]

    _ = new_text_for_first_nonempty  # placeholder, unused


def detect_h1_base(doc) -> int:
    """auto 基准: 第一个 H1 段现有的前导数字(如标书从 10 起);无 → 1。"""
    for p in doc.paragraphs:
        if get_style_name(p) == "Heading 1":
            m = re.match(r"^\s*(\d+)", get_paragraph_text(p))
            return int(m.group(1)) if m else 1
    return 1


def plan_renumber(doc, h1_base: int = 1) -> tuple[list[dict], list[dict], dict]:
    """扫一遍, 返回 (before, after, stats). 不改 doc. h1_base=首个 H1 的目标编号(默认 1)."""
    h1 = h1_base - 1
    h2 = 0
    h3 = 0
    tbl = 0  # H1 章内表序

    before: list[dict] = []
    after: list[dict] = []

    h1_max = 0
    h2_max = 0
    h3_max = 0
    tbl_total = 0

    for idx, p in enumerate(doc.paragraphs):
        sname = get_style_name(p)
        text = get_paragraph_text(p)

        new_prefix: Optional[str] = None
        kind: Optional[str] = None

        if sname == "Heading 1":
            h1 += 1
            h2 = 0
            h3 = 0
            tbl = 0
            new_prefix = f"{h1} "
            kind = "h1"
            h1_max = max(h1_max, h1)
        elif sname == "Heading 2":
            h2 += 1
            h3 = 0
            new_prefix = f"{h1}.{h2} "
            kind = "h2"
            h2_max = max(h2_max, h2)
        elif sname == "Heading 3":
            h3 += 1
            new_prefix = f"{h1}.{h2}.{h3} "
            kind = "h3"
            h3_max = max(h3_max, h3)
        elif sname == "zdwp表名":
            tbl += 1
            new_prefix = f"表 {h1}-{tbl} "
            kind = "table"
            tbl_total += 1
        else:
            # Title / Normal / 其他: 不动
            continue

        # 计算新 text 用于 report
        if kind == "table":
            stripped = strip_old_prefix(text, "table")
        else:
            stripped = strip_old_prefix(text, "heading")
        new_text = new_prefix + stripped

        before.append({"idx": idx, "old_text": text, "style": sname})
        after.append({"idx": idx, "new_text": new_text, "style": sname, "new_prefix": new_prefix})

    stats = {
        "h1_count": h1_max,
        "h2_count": h2_max,
        "h3_count": h3_max,
        "table_count": tbl_total,
    }
    return before, after, stats


def apply_renumber(doc, after_plan: list[dict]) -> None:
    """按 plan 改 doc 的 paragraphs."""
    paragraphs = doc.paragraphs
    for item in after_plan:
        idx = item["idx"]
        new_prefix = item["new_prefix"]
        p = paragraphs[idx]
        rewrite_paragraph_with_prefix(p, new_prefix)


def verify_strict_sequence(doc, h1_base: int = 1) -> tuple[bool, list[str]]:
    """重读后, 验证 H1/H2/H3 段 text 开头编号严格递增连贯.

    返回 (ok, errors).
    """
    errors: list[str] = []
    # 跟踪 (cur_h1) 下的 h2 序; (cur_h1, cur_h2) 下的 h3 序
    h2_in_h1: dict[int, int] = {}
    h3_in_h2: dict[tuple[int, int], int] = {}
    tbl_in_h1: dict[int, int] = {}
    cur_h1 = 0
    cur_h2 = 0
    h1_seen = h1_base - 1

    for idx, p in enumerate(doc.paragraphs):
        sname = get_style_name(p)
        text = get_paragraph_text(p)
        if sname == "Heading 1":
            m = re.match(r"^\s*(\d+)\s", text)
            if not m:
                errors.append(f"#{idx} H1 缺新编号: {text[:30]!r}")
                continue
            n = int(m.group(1))
            expected = h1_seen + 1
            if n != expected:
                errors.append(f"#{idx} H1 编号 {n} != 期望 {expected}")
            h1_seen = n
            cur_h1 = n
            cur_h2 = 0
        elif sname == "Heading 2":
            m = re.match(r"^\s*(\d+)\.(\d+)\s", text)
            if not m:
                errors.append(f"#{idx} H2 缺新编号: {text[:30]!r}")
                continue
            a, b = int(m.group(1)), int(m.group(2))
            if a != cur_h1:
                errors.append(f"#{idx} H2 章号 {a} != 当前 H1 {cur_h1}")
            prev = h2_in_h1.get(cur_h1, 0)
            if b != prev + 1:
                errors.append(f"#{idx} H2 {a}.{b} 序号不连贯 (上一个 {a}.{prev})")
            h2_in_h1[cur_h1] = b
            cur_h2 = b
        elif sname == "Heading 3":
            m = re.match(r"^\s*(\d+)\.(\d+)\.(\d+)\s", text)
            if not m:
                errors.append(f"#{idx} H3 缺新编号: {text[:30]!r}")
                continue
            a, b, c = int(m.group(1)), int(m.group(2)), int(m.group(3))
            if (a, b) != (cur_h1, cur_h2):
                errors.append(f"#{idx} H3 父号 {a}.{b} != 当前 {cur_h1}.{cur_h2}")
            key = (cur_h1, cur_h2)
            prev = h3_in_h2.get(key, 0)
            if c != prev + 1:
                errors.append(f"#{idx} H3 {a}.{b}.{c} 序号不连贯 (上一个 {a}.{b}.{prev})")
            h3_in_h2[key] = c
        elif sname == "zdwp表名":
            m = re.match(r"^\s*表\s*(\d+)-(\d+)\s", text)
            if not m:
                errors.append(f"#{idx} 表号缺新编号: {text[:30]!r}")
                continue
            a, b = int(m.group(1)), int(m.group(2))
            if a != cur_h1:
                errors.append(f"#{idx} 表号章号 {a} != 当前 H1 {cur_h1}")
            prev = tbl_in_h1.get(cur_h1, 0)
            if b != prev + 1:
                errors.append(f"#{idx} 表号 {a}-{b} 序号不连贯 (上一个 {a}-{prev})")
            tbl_in_h1[cur_h1] = b

    return len(errors) == 0, errors


def make_backup(docx_path: Path) -> Path:
    """生成 .bak-N-YYYY-MM-DD.docx, N 自增不覆盖."""
    return _cc.make_backup(docx_path)


def main_headings(argv: Optional[list[str]] = None) -> int:
    ap = argparse.ArgumentParser(
        description="按当前段物理顺序统一重编 H1/H2/H3 + 表号"
    )
    ap.add_argument("docx_path", help="目标 docx 路径")
    ap.add_argument("--dry-run", action="store_true", help="只输出 plan, 不动 docx")
    ap.add_argument("--no-backup", action="store_true", help="不生成 .bak 备份")
    ap.add_argument("--report", default=None, help="写 JSON 报告到该路径")
    ap.add_argument("--h1-base", default="1",
                    help="首个 H1 的目标编号: 整数 或 auto(沿用文档第一个 H1 现有编号,适配从第10章起的标书)")
    args = ap.parse_args(argv)

    docx_path = Path(args.docx_path)
    if not docx_path.exists():
        print(f"ERROR: {docx_path} 不存在", file=sys.stderr)
        return 2

    doc = Document(str(docx_path))
    h1_base = detect_h1_base(doc) if args.h1_base == "auto" else int(args.h1_base)
    if h1_base != 1:
        print(f"[base] H1 从 {h1_base} 起编")
    before, after, stats = plan_renumber(doc, h1_base)

    print(f"[plan] H1={stats['h1_count']} H2(max-in-section)={stats['h2_count']} "
          f"H3(max-in-section)={stats['h3_count']} 表={stats['table_count']}")
    print(f"[plan] 共 {len(after)} 段需重编")

    if args.report:
        report = {
            "before": before,
            "after": [{k: v for k, v in a.items() if k != "new_prefix"} for a in after],
            **stats,
        }
        Path(args.report).write_text(
            json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(f"[report] 已写入 {args.report}")

    if args.dry_run:
        print("[dry-run] 不动 docx, 退出")
        return 0

    # backup
    if not args.no_backup:
        bak = make_backup(docx_path)
        print(f"[backup] {bak.name}")

    apply_renumber(doc, after)
    doc.save(str(docx_path))
    print(f"[saved] {docx_path}")

    # OOXML 合法性 verify: 重读
    try:
        doc2 = Document(str(docx_path))
    except Exception as e:
        print(f"ERROR: 重读失败 (OOXML 可能损坏): {e}", file=sys.stderr)
        return 3

    ok, errors = verify_strict_sequence(doc2, h1_base)
    if not ok:
        print("[verify] 编号序列检测失败:", file=sys.stderr)
        for e in errors[:20]:
            print(f"  - {e}", file=sys.stderr)
        return 4
    print("[verify] 编号严格递增连贯 ✓")
    return 0


# ---------------- pipeline adapter ----------------
def apply_renumber_headings(doc, args=None) -> dict:
    """按层级重排标题编号（内存版，不开文件不存盘）。

    h1_base 走 args 而不是写死 1：main() 一直是 `plan_renumber(doc, h1_base)`
    （--h1-base 专为「标书从第 10 章起编」而加），抽 apply() 时这个值漏搬了，
    于是所有走 pipeline / spec 的调用一律被悄悄按 1 起编 —— 「配了没生效」是
    这条链上最难查的一类 bug，宁可补回来。缺省 1 与漏搬前的行为逐字相同。
    """
    dry = bool(getattr(args, "dry_run", False)) if args else False
    raw_base = getattr(args, "h1_base", 1) if args else 1
    if raw_base is None:
        raw_base = 1
    h1_base = detect_h1_base(doc) if str(raw_base) == "auto" else int(raw_base)
    before, after, stats = plan_renumber(doc, h1_base)
    if not dry:
        apply_renumber(doc, after)
    return {
        "changed": len(after),
        "h1_base": h1_base,
        **stats,
    }


# ══════════ seq ← renumber_headings_seq.py ══════════

HEAD_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
NUM_RE = re.compile(r"^\s*(\d+(?:[.．]\d+)*)[\s　.．、]")


def plan(doc):
    """返回 (plan_items, warnings)。plan_items=[(para_idx, old_num, new_num)]。"""
    items, warns = [], []
    path: list[int] = []  # 修正后的编号链,如 [12, 3] = 12.3
    base_locked = False
    for idx, p in enumerate(doc.paragraphs):
        if get_style_name(p) not in HEAD_STYLES:
            continue
        m = NUM_RE.match(get_paragraph_text(p))
        if not m:
            continue  # 无编号标题(前言/附录)不发明编号
        old = m.group(1).replace("．", ".")
        depth = old.count(".") + 1
        if depth == 1:
            if not base_locked:
                path = [int(old)]  # 首个一级号 = base,原样保留(标书从10起仍是10)
                base_locked = True
            else:
                path = [_next_top(items, int(old))]
        else:
            if depth > len(path) + 1:
                warns.append(f"#{idx} 深度突跳: {old} (当前链 {'.'.join(map(str, path))})")
                while depth > len(path) + 1:
                    path.append(1)
            path = path[: depth - 1] + [_next_sibling(items, path[: depth - 1])]
        new = ".".join(map(str, path))
        items.append((idx, old, new, depth))
    return items, warns


def _next_top(items, _old):
    tops = [int(n.split(".")[0]) for _, _, n, d in items if d == 1]
    return tops[-1] + 1 if tops else 1


def _next_sibling(items, parent: list[int]):
    pre = ".".join(map(str, parent)) + "."
    sibs = [int(n[len(pre):]) for _, _, n, d in items
            if d == len(parent) + 1 and n.startswith(pre) and "." not in n[len(pre):]]
    return sibs[-1] + 1 if sibs else 1


def verify(doc):
    """复检: 再按 plan 模型扫一遍,所有 old==new 即连贯。"""
    items, _ = plan(doc)
    return [(i, o, n) for i, o, n, _ in items if o != n]


def apply_seq(doc, args=None) -> dict:
    """在**已打开**的 doc 上按现有编号深度重排标题号（不动层级）。

    改写走 renumber_headings 的 run 级 rewrite_paragraph_with_prefix，只碰
    document.xml 的段首 run，所以是 doc-based step。

    dry_run 只出计划不改树 —— 计划和真改用的是同一个 plan()，不存在「预览的和
    实际改的不是一回事」。存盘后的连贯性复检留给调用方（verify() 要重开文件，
    那是 apply() 契约里明确不做的事）。
    """
    dry_run = bool(getattr(args, "dry_run", False)) if args else False
    items, warns = plan(doc)
    changes = [(i, o, n) for i, o, n, _ in items if o != n]
    if not dry_run:
        paras = doc.paragraphs
        for i, _o, n in changes:
            rewrite_paragraph_with_prefix(paras[i], f"{n} ")
    return {
        "changed": len(changes),
        "headings": len(items),
        "dry_run": dry_run,
        "changes": [{"idx": i, "old": o, "new": n} for i, o, n in changes],
        "warnings": warns,
    }


def main_seq(argv=None) -> int:
    ap = argparse.ArgumentParser(description="标题编号按现有深度重排序(修断号/重号/错号,不动层级)")
    ap.add_argument("docx_path")
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--no-backup", action="store_true")
    ap.add_argument("--report", default=None)
    a = ap.parse_args(argv)
    path = Path(a.docx_path)
    if not path.exists():
        print(f"ERROR: {path} 不存在", file=sys.stderr)
        return 2

    doc = Document(str(path))
    rep = apply_seq(doc, a)
    changes = rep["changes"]
    print(f"[plan] 带编号标题 {rep['headings']} 段, 需修 {len(changes)} 段")
    for c in changes:
        print(f"  #{c['idx']} {c['old']} → {c['new']}")
    for w in rep["warnings"]:
        print(f"  ⚠ {w}")
    if a.report:
        Path(a.report).write_text(json.dumps(
            {"changes": changes, "warnings": rep["warnings"]},
            ensure_ascii=False, indent=2), encoding="utf-8")
    if a.dry_run:
        print("[dry-run] 未写文件")
        return 0
    if not changes:
        print("✓ 标题编号已连贯,无需改写")
        return 0
    # 备份放在内存改写之后不影响内容：磁盘上的文件直到下面 doc.save 才动
    if not a.no_backup:
        print(f"[backup] {make_backup(path).name}")
    doc.save(str(path))
    left = verify(Document(str(path)))
    if left:
        print(f"✗ 复检仍有 {len(left)} 处不连贯: {left[:5]}", file=sys.stderr)
        return 4
    print(f"✓ 已修 {len(changes)} 处,复检连贯")
    return 0


# ──────────────────────────── 家族入口（子命令分发）────────────────────────────

SUBCOMMANDS = {
    "headings": main_headings,
    "seq": main_seq,
}


def main(argv: list[str] | None = None) -> int:
    return _cc.family_main(SUBCOMMANDS, argv, file=__file__)


if __name__ == "__main__":
    sys.exit(main())
