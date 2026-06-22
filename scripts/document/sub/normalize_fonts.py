#!/usr/bin/env python3
"""normalize_fonts — 统一 docx 中西文字体 + 把标题颜色压成黑色。

Why（踩坑根因）：
  pandoc 把 md 转 docx 时套 Word 内置 Heading 1/2/3 样式，这套样式自带主题强调蓝
  （Heading 1 无显式色→继承主题蓝、Heading 3/4/5 = 0F4761 青蓝）。若生成脚本只
  改了字号/字体没改颜色，标题就是蓝/青的——投稿要求黑色。
  同时 pandoc 常在 run 上留直接 rFonts（西文 Calibri 等），盖过样式级字体 →
  中西文字体花。

本命令两手抓（样式级 + run 级 belt-and-suspenders）：
  · 标题段（样式名含 heading/title/subtitle/toc）：中文→黑体、西文→Times New
    Roman、字体颜色→黑（000000）。
  · 其余段（正文/图注/表格内文字）：中文→宋体、西文→Times New Roman；颜色不动
    （正文一般本就是 auto/黑，避免误伤刻意着色）。

遍历 body + 所有表格（含嵌套表）的段落与 run。原地修改，默认先备份 .bak-时间戳。

Usage:
  python3 normalize_fonts.py 论文.docx                       # 默认：标题黑体黑色/正文宋体/西文TNR
  python3 normalize_fonts.py 论文.docx --body-cjk 宋体 --heading-cjk 黑体 --latin "Times New Roman"
  python3 normalize_fonts.py 论文.docx --heading-color 000000
  python3 normalize_fonts.py 论文.docx --dry-run
  python3 normalize_fonts.py 论文.docx --no-backup
"""
from __future__ import annotations

import argparse
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# 标题类样式判定关键字（小写匹配 style.name）
_HEADING_KEYS = ("heading", "标题", "title", "subtitle", "副标题", "toc")


def _is_heading_style_name(name: str) -> bool:
    n = (name or "").lower()
    return any(k in n or k in (name or "") for k in _HEADING_KEYS)


def _set_rfonts(rpr_parent, cjk: str, latin: str) -> None:
    """在 rPr/style-rPr 上设 rFonts：eastAsia=cjk, ascii/hAnsi/cs=latin。"""
    rpr = rpr_parent.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), cjk)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:cs"), latin)


def _set_color(rpr_parent, hexval: str) -> None:
    """在 rPr/style-rPr 上设 w:color。"""
    rpr = rpr_parent.get_or_add_rPr()
    c = rpr.find(qn("w:color"))
    if c is None:
        c = rpr.makeelement(qn("w:color"), {})
        rpr.append(c)
    c.set(qn("w:val"), hexval)
    # 去掉主题色引用（themeColor 会盖过 val）
    for a in (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade")):
        if a in c.attrib:
            del c.attrib[a]


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for t in cell.tables:  # 嵌套表
                yield from _iter_table_paragraphs(t)


def _iter_all_paragraphs(doc):
    yield from doc.paragraphs
    for t in doc.tables:
        yield from _iter_table_paragraphs(t)


def process(docx_path: Path, body_cjk: str, heading_cjk: str, latin: str,
            heading_color: str, dry_run: bool, backup: bool) -> dict:
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document(str(docx_path))
    styles = doc.styles

    # ── 1) 样式级 ──
    style_hits = 0
    for st in styles:
        if st.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        heading = _is_heading_style_name(st.name)
        _set_rfonts(st.element, heading_cjk if heading else body_cjk, latin)
        if heading:
            _set_color(st.element, heading_color)
        style_hits += 1

    # ── 2) run 级（belt-and-suspenders）──
    run_hits = 0
    heading_runs = 0
    for p in _iter_all_paragraphs(doc):
        heading = _is_heading_style_name(p.style.name if p.style else "")
        cjk = heading_cjk if heading else body_cjk
        for r in p.runs:
            _set_rfonts(r._element, cjk, latin)
            if heading:
                _set_color(r._element, heading_color)
                heading_runs += 1
            run_hits += 1

    result = {
        "file": str(docx_path),
        "style_hits": style_hits,
        "run_hits": run_hits,
        "heading_runs": heading_runs,
        "policy": f"标题={heading_cjk}/#{heading_color} · 正文={body_cjk} · 西文={latin}",
        "dry_run": dry_run,
    }
    if dry_run:
        result["written"] = False
        return result

    if backup:
        bak = docx_path.with_name(
            docx_path.name + ".bak-" + datetime.now().strftime("%Y%m%d-%H%M%S"))
        shutil.copy2(docx_path, bak)
        result["backup"] = str(bak)

    doc.save(str(docx_path))
    result["written"] = True
    return result


def _fmt(r: dict) -> str:
    head = f"[fonts] {Path(r['file']).name}"
    tail = ""
    if r.get("written"):
        tail = f" → 样式 {r['style_hits']} · run {r['run_hits']}(标题run {r['heading_runs']})"
        if r.get("backup"):
            tail += f" · bak={Path(r['backup']).name}"
    elif r["dry_run"]:
        tail = (f" (dry-run) 样式 {r['style_hits']} · run {r['run_hits']}"
                f"(标题run {r['heading_runs']})")
    return f"{head}: {r['policy']}{tail}"


def main() -> int:
    ap = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("docx_pos", nargs="?", help="(positional) docx 路径，等价 --docx")
    ap.add_argument("--docx", dest="docx_kw", help="目标 docx（原地修改）")
    ap.add_argument("--body-cjk", default="宋体", help="正文中文字体，默认 宋体")
    ap.add_argument("--heading-cjk", default="黑体", help="标题中文字体，默认 黑体")
    ap.add_argument("--latin", default="Times New Roman", help="中西文中的西文字体，默认 Times New Roman")
    ap.add_argument("--heading-color", default="000000", help="标题字体颜色 hex，默认 000000（黑）")
    ap.add_argument("--no-backup", action="store_true", help="不创建 .bak-时间戳 备份")
    ap.add_argument("--dry-run", action="store_true", help="只报告不写盘")
    args = ap.parse_args()

    docx = args.docx_kw or args.docx_pos
    if not docx:
        print("[fonts] missing docx (positional or --docx)", file=sys.stderr)
        return 2
    docx_path = Path(docx)
    if not docx_path.exists():
        print(f"[fonts] not found: {docx_path}", file=sys.stderr)
        return 2

    r = process(docx_path, args.body_cjk, args.heading_cjk, args.latin,
                args.heading_color, dry_run=args.dry_run, backup=not args.no_backup)
    print(_fmt(r))
    return 0


# ---------------- pipeline adapter ----------------
def apply_path(docx_path, args=None) -> dict:
    body = getattr(args, "body_cjk", "宋体") if args else "宋体"
    heading = getattr(args, "heading_cjk", "黑体") if args else "黑体"
    latin = getattr(args, "latin", "Times New Roman") if args else "Times New Roman"
    color = getattr(args, "heading_color", "000000") if args else "000000"
    dry = bool(getattr(args, "dry_run", False)) if args else False
    backup = not bool(getattr(args, "no_backup", False)) if args else True
    return process(Path(docx_path), body, heading, latin, color, dry, backup)


if __name__ == "__main__":
    sys.exit(main())
