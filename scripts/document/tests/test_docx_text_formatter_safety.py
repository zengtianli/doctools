"""tests for 规范化的**安全红线** —— 2026-07-26 审计实测出的 4 条,每条都真出过错。

1. 制表符原位:run 内 `[t, tab, t]` 改写后 tab 不许被挤到末尾(目录条/签署行会错位)
2. 页眉页脚默认保留:clean 不再删 headerReference/footerReference,要删得显式 --strip-headers
   (旧默认删掉了一切,连 typeset 套完院模板的 14/13 个引用也一并没了)
3. 单位词边界:小时候 ✗→ h候 · 毫米波 ✗→ mm波 · Item2 ✗→ Item²
4. 引号计数是真改动数,不是匹配数(本来就是中文引号的不该报"替换了 N 个")
"""

from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

_DOC = Path(__file__).resolve().parent.parent
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

spec = importlib.util.spec_from_file_location("dtf_safety_uut", str(_DOC / "docx_text_formatter.py"))
assert spec and spec.loader
dtf = importlib.util.module_from_spec(spec)
sys.modules["dtf_safety_uut"] = dtf
spec.loader.exec_module(dtf)

sys.path.insert(0, str(_DOC.parent.parent / "lib"))
from text_fixes import fix_quotes, fix_units  # noqa: E402


def _stats():
    return {"quotes": 0, "punctuation": 0, "units": 0, "scopes": {}}


def _kids(r):
    return [c.tag.split("}")[1] for c in r]


# ── 1. 制表符原位 ────────────────────────────────────────────────

def test_tab_stays_between_text():
    """run 子节点顺序 [t, tab, t] 必须原样保留 —— 合并文本会把 tab 挤到末尾。"""
    d = Document()
    p = d.add_paragraph()
    p._p.append(parse_xml(
        f'<w:r {nsdecls("w")}><w:t>第一项,</w:t><w:tab/><w:t>第二项(A)</w:t></w:r>'))
    r = p._p.findall(qn("w:r"))[0]
    assert _kids(r) == ["t", "tab", "t"]

    dtf.process_paragraph_element(p._p, _stats())

    r = p._p.findall(qn("w:r"))[0]
    assert _kids(r) == ["t", "tab", "t"], "tab 被挪位 = 目录条/签署行对齐错乱"
    assert [t.text for t in r.findall(qn("w:t"))] == ["第一项，", "第二项（A）"]


def test_br_also_preserved():
    d = Document()
    p = d.add_paragraph()
    p._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t>上行,</w:t><w:br/><w:t>下行;</w:t></w:r>'))
    dtf.process_paragraph_element(p._p, _stats())
    assert _kids(p._p.findall(qn("w:r"))[0]) == ["t", "br", "t"]


# ── 2. 页眉页脚默认保留 ──────────────────────────────────────────

@pytest.fixture
def docx_with_header(tmp_path: Path) -> Path:
    d = Document()
    sec = d.sections[0]
    sec.header.is_linked_to_previous = False
    sec.header.paragraphs[0].text = "页眉:标题(H)"
    sec.footer.is_linked_to_previous = False
    sec.footer.paragraphs[0].text = "页脚:第1页"
    d.add_paragraph("正文,内容(A)")
    src = tmp_path / "withheader.docx"
    d.save(src)
    return src


def _header_refs(docx: Path) -> tuple[int, int]:
    d = Document(docx)
    h = sum(len(s._sectPr.findall(qn("w:headerReference"))) for s in d.sections)
    f = sum(len(s._sectPr.findall(qn("w:footerReference"))) for s in d.sections)
    return h, f


def _run(src: Path) -> Path:
    assert dtf.process_docx(str(src)) is True
    return src.with_name(f"{src.stem}_fixed{src.suffix}")


def test_headers_kept_by_default(docx_with_header: Path):
    """默认跑一遍规范化,页眉页脚引用必须一个不少。"""
    before = _header_refs(docx_with_header)
    assert before == (1, 1)
    assert dtf.STRIP_HEADERS is False
    assert _header_refs(_run(docx_with_header)) == before


def test_header_text_normalized(docx_with_header: Path):
    """保留的同时,页眉页脚里的半角标点也要被规范化。"""
    out = _run(docx_with_header)
    d = Document(out)
    assert "（H）" in d.sections[0].header.paragraphs[0].text


def test_strip_headers_opt_in(docx_with_header: Path):
    """显式 --strip-headers 时才删。"""
    dtf.STRIP_HEADERS = True
    try:
        assert _header_refs(_run(docx_with_header)) == (0, 0)
    finally:
        dtf.STRIP_HEADERS = False


# ── 3. 单位词边界 ────────────────────────────────────────────────

@pytest.mark.parametrize("text", [
    "小时候在毫米波实验室",
    "秒钟表和分钟表",
    "Item2 Form3",
    "毫升装置",
])
def test_units_not_touched_without_number(text: str):
    """没有数字打头的中文词 / 英文单词内的 m2 一律不动。"""
    assert fix_units(text) == (text, 0)


@pytest.mark.parametrize("text,expect", [
    ("面积5平方米", "面积5m²"),
    ("体积3 立方米", "体积3 m³"),
    ("长2公里", "长2km"),
    ("5m2", "5m²"),
    ("温度25摄氏度", "温度25℃"),
])
def test_units_still_convert_after_number(text: str, expect: str):
    assert fix_units(text)[0] == expect


# ── 4. 引号计数 = 真改动数 ───────────────────────────────────────

def test_quote_count_is_real_changes():
    assert fix_quotes("已是中文引号“甲”")[1] == 0
    assert fix_quotes('英文"甲"')[1] == 2


# ── 5. 域勾选 × 引号计数器（唯一会静默产出错误文档的路径）─────────

def _para_with_ins():
    """同一段落三个 run：正文 → 修订插入 → 正文，引号跨 run 配对。"""
    d = Document()
    p = d.add_paragraph()
    p._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t>他说"甲</w:t></w:r>'))
    ins = parse_xml(
        f'<w:ins {nsdecls("w")} w:id="1" w:author="t" w:date="2026-01-01T00:00:00Z">'
        f'<w:r><w:t>插入"</w:t></w:r></w:ins>')
    p._p.append(ins)
    p._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t>乙"结束</w:t></w:r>'))
    return d, p


def _run_texts(p):
    return [r.findall(qn("w:t"))[0].text for r in p._p.iter(qn("w:r"))
            if r.findall(qn("w:t"))]


def test_skipped_scope_does_not_flip_quote_direction():
    """勾掉「审阅修订」后，正文里后续引号的左右方向必须与全勾时一致。

    被跳过的 run 仍要推进段落级奇偶计数器 —— 否则第 3 个 run 的 “ 会翻成 ”。
    """
    d_all, p_all = _para_with_ins()
    dtf.process_paragraph_element(p_all._p, _stats(), cfg=dtf.FormatConfig(quote_font=False))
    full = _run_texts(p_all)

    d_body, p_body = _para_with_ins()
    cfg = dtf.FormatConfig(quote_font=False,
                           scopes=frozenset({"body", "table", "comments", "notes", "headers"}))
    dtf.process_paragraph_element(p_body._p, _stats(), cfg=cfg)
    partial = _run_texts(p_body)

    assert partial[1] == '插入"', "勾掉的域被改了"
    assert partial[0] == full[0] and partial[2] == full[2], (
        f"正文引号方向被域过滤带偏: 全勾={full} 只正文={partial}")
    # 第 3 个引号是本段第 3 个 → 奇数位 → 左引号；关键是「跳过 revision 不改变它」
    assert partial[2].startswith("乙“"), f"第 3 个 run 引号方向错: {partial}"


def test_scope_filter_actually_skips():
    d, p = _para_with_ins()
    cfg = dtf.FormatConfig(scopes=frozenset({"revision"}), quote_font=False)
    dtf.process_paragraph_element(p._p, _stats(), cfg=cfg)
    texts = _run_texts(p)
    assert texts[0] == '他说"甲' and texts[2] == '乙"结束', "只勾修订时正文不该被改"
    assert texts[1] == '插入”', f"修订段该被改: {texts}"


def test_quote_font_can_be_turned_off():
    """关掉「引号设为宋体」后不再拆 run、不再写字体。"""
    d = Document()
    p = d.add_paragraph()
    p.add_run('标题“重点”部分')
    dtf.process_paragraph_element(p._p, _stats(), cfg=dtf.FormatConfig(quote_font=False))
    assert len(p._p.findall(qn("w:r"))) == 1, "关掉后不该拆 run"

    d2 = Document()
    p2 = d2.add_paragraph()
    p2.add_run('标题“重点”部分')
    dtf.process_paragraph_element(p2._p, _stats(), cfg=dtf.FormatConfig(quote_font=True))
    assert len(p2._p.findall(qn("w:r"))) > 1, "开着时应拆 run 设宋体"


# ── 6. 半角→全角必须绕开「机器读的」片段 ─────────────────────────

sys.path.insert(0, str(_DOC.parent.parent / "lib"))
from text_fixes import fix_punctuation  # noqa: E402


@pytest.mark.parametrize("text", [
    "https://example.com/a?b=1;c=2",
    "见 www.gov.cn/x?a=1;b=2",
    "联系 a.b@c.com",
    "`x = {1: 2}`",
    "1,234",
    "1,234,567.89",
])
def test_punct_skips_machine_readable(text: str):
    """URL / 邮箱 / 行内代码 / 千分位整段原样保留 —— 转全角=内容改坏。"""
    assert fix_punctuation(text)[0] == text


def test_punct_skips_fenced_code_block():
    md = "说明,如下:\n```python\nd = {'a': 1, 'b': 2}\n```\n结束,完毕"
    out, _ = fix_punctuation(md)
    assert "{'a': 1, 'b': 2}" in out, "围栏代码块被全角化"
    assert "说明，如下：" in out and "结束，完毕" in out, "围栏外的正文该照常改"


def test_punct_still_converts_prose():
    assert fix_punctuation("正常中文,标点:全角化(应该改)")[0] == "正常中文，标点：全角化（应该改）"
