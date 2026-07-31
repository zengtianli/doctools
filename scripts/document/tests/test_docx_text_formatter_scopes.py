"""tests for docx_fmt.py text（原 docx_text_formatter.py）—— 规范化的**覆盖面**回归门。

2026-07-26 起因：规范化对「审阅」无效。python-docx 的 `Document.paragraphs` 只认
`w:body/./w:p`、`Paragraph.runs` 只认 `./w:r`，于是 `w:ins`/`w:del`（修订）、
`w:hyperlink`、文本框、嵌套表里的 run 被静默跳过，批注/脚注更是整个 part 没碰。
本测试手搓一份「什么坑都有」的 docx，逐域断言字符真被改到，且修订结构没被写坏。

红线（改坏了必须红）：
  · 修订删除态里的文本载体必须仍是 `w:delText`（写成 `w:t` = 删掉的字变回正文）
  · `w:moveFrom`（审阅移动的原位置）同上
  · 幂等：跑两遍文本一致、run 数不增（拆引号 run 不能每跑一遍翻一倍）
"""

from __future__ import annotations

import importlib.util
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

_DOC = Path(__file__).resolve().parent.parent
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

spec = importlib.util.spec_from_file_location("dtf_uut", str(_DOC / "docx_fmt.py"))
assert spec and spec.loader
dtf = importlib.util.module_from_spec(spec)
sys.modules["dtf_uut"] = dtf
spec.loader.exec_module(dtf)


def _rev(tag: str, rid: str, inner: str):
    """建一个修订节点（w:ins / w:del / w:moveFrom / w:moveTo）并塞一个 run 进去。"""
    node = OxmlElement(tag)
    node.set(qn("w:id"), rid)
    node.set(qn("w:author"), "tester")
    node.set(qn("w:date"), "2026-01-01T00:00:00Z")
    node.append(parse_xml(f'<w:r {nsdecls("w")}>{inner}</w:r>'))
    return node


def _t(text: str) -> str:
    return f'<w:t xml:space="preserve">{text}</w:t>'


def _del_t(text: str) -> str:
    return f'<w:delText xml:space="preserve">{text}</w:delText>'


@pytest.fixture
def messy_docx(tmp_path: Path) -> Path:
    """一份含 修订/移动/批注/超链接/嵌套表/文本框 的 docx，每域各埋一处半角标点。"""
    d = Document()
    d.add_paragraph().add_run("正文:面积5平方米(A);")

    p = d.add_paragraph()
    p._p.append(_rev("w:ins", "801", _t('插入:体积3立方米(B);"引"')))
    p._p.append(_rev("w:del", "802", _del_t('删除:长度2公里(C);"引"')))
    p._p.append(_rev("w:moveFrom", "803", _del_t("移出:原处(D);")))
    p._p.append(_rev("w:moveTo", "804", _t("移入:新处(E);")))

    d.add_paragraph()._p.append(parse_xml(
        f'<w:hyperlink {nsdecls("w", "r")} r:id="rIdX"><w:r>{_t("超链接:见附件(F);")}</w:r></w:hyperlink>'))

    cell = d.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run("外层:数据(G);")
    cell.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0].add_run("嵌套表:数据(H);")

    d.add_paragraph()._p.append(parse_xml(
        f'<w:r {nsdecls("w")}><w:pict><v:shape xmlns:v="urn:schemas-microsoft-com:vml"'
        ' style="width:100pt;height:50pt"><v:textbox><w:txbxContent><w:p><w:r>'
        f'{_t("文本框:说明(I);")}</w:r></w:p></w:txbxContent></v:textbox></v:shape></w:pict></w:r>'))

    d.add_comment(d.paragraphs[0].runs, text='批注:注释(J);"引"', author="tester")

    src = tmp_path / "messy.docx"
    d.save(src)
    return src


def _part_text(docx: Path, name: str) -> str:
    from lxml import etree

    zf = zipfile.ZipFile(docx)
    if name not in zf.namelist():
        return ""
    root = etree.fromstring(zf.read(name))
    return "".join(e.text or "" for e in root.iter(W + "t", W + "delText"))


def _iter_tags(docx: Path, name: str, tag: str):
    from lxml import etree

    return list(etree.fromstring(zipfile.ZipFile(docx).read(name)).iter(W + tag))


def _run_engine(src: Path) -> Path:
    assert dtf.process_docx(str(src)) is True
    out = src.with_name(f"{src.stem}_fixed{src.suffix}")
    assert out.exists()
    return out


# ── 覆盖面：每个域都得改到 ──────────────────────────────────────

@pytest.mark.parametrize("marker", ["A", "B", "C", "D", "E", "F", "G", "H", "I"])
def test_every_body_scope_normalized(messy_docx: Path, marker: str):
    """正文 / 修订(插入·删除·移出·移入) / 超链接 / 嵌套表 / 文本框 全都得半角→全角。"""
    body = _part_text(_run_engine(messy_docx), "word/document.xml")
    assert f"（{marker}）；" in body, f"域 {marker} 没被规范化：{body}"
    assert f"({marker});" not in body


def test_comment_normalized(messy_docx: Path):
    """批注（comments.xml）是独立 part，python-docx 的段落遍历根本碰不到。"""
    txt = _part_text(_run_engine(messy_docx), "word/comments.xml")
    assert "批注：注释（J）；" in txt
    assert "“" in txt and "”" in txt


def test_units_in_revision(messy_docx: Path):
    body = _part_text(_run_engine(messy_docx), "word/document.xml")
    assert "3m³" in body and "2km" in body      # w:ins / w:del 里的中文单位


# ── 红线：修订结构不能被写坏 ────────────────────────────────────

@pytest.mark.parametrize("wrapper", ["del", "moveFrom"])
def test_deleted_text_stays_delText(messy_docx: Path, wrapper: str):
    """删除态内不允许出现 w:t —— 那等于把「删掉的字」变回正文。"""
    out = _run_engine(messy_docx)
    nodes = _iter_tags(out, "word/document.xml", wrapper)
    assert nodes, f"w:{wrapper} 丢了"
    for node in nodes:
        assert not list(node.iter(W + "t")), f"w:{wrapper} 里混入了 w:t"
        assert list(node.iter(W + "delText"))


def test_revision_marks_survive(messy_docx: Path):
    """修订节点数量不变（不能把 w:ins/w:del 拆散或吞掉）。"""
    before = {t: len(_iter_tags(messy_docx, "word/document.xml", t))
              for t in ("ins", "del", "moveFrom", "moveTo")}
    after = {t: len(_iter_tags(_run_engine(messy_docx), "word/document.xml", t))
             for t in ("ins", "del", "moveFrom", "moveTo")}
    assert before == after == {"ins": 1, "del": 1, "moveFrom": 1, "moveTo": 1}


def test_idempotent(messy_docx: Path, tmp_path: Path):
    """跑第二遍：文本一致、run 数不增（引号拆 run 不能反复膨胀）。"""
    first = _run_engine(messy_docx)
    again = tmp_path / "second.docx"
    again.write_bytes(first.read_bytes())
    second = _run_engine(again)
    assert _part_text(first, "word/document.xml") == _part_text(second, "word/document.xml")
    assert len(_iter_tags(first, "word/document.xml", "r")) == \
           len(_iter_tags(second, "word/document.xml", "r"))
