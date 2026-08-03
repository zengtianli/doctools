"""renum.py 的空集出口回归门 —— 「枚举为空必须判非 0」不许被悄悄改回 return 0。

## 为什么单开一个文件（2026-08-03 立）

`renum.py` 2026-08-03 加了 `EMPTY_RC = 3` 与 `_empty_set_exit()`，把「一个对象都没
枚举到却打『✓ 每节连续 1..N』并 rc=0」这个假绿堵掉了。**但当天的核验镜头指出：
这条新判据在冒烟轴里零覆盖** —— `tests/smoke/_verb_specs.py` 的 fixture 被同一轮
加强成「有 6 张图」了，于是 `renumber-fig` 走的是非空分支，谁把 `_empty_set_exit`
改回 `return 0`，92 条 smoke 一条都不会红。

冒烟轴按设计只跑「一条动词一份标准 fixture」，覆盖不了「同一条动词在空集上的行为」，
所以这件事不该塞进 `_verb_specs`（那会污染它「一动词一行」的形状），而是单开本文件。

## 判据

对每一个 `_empty_set_exit()` 调用点，构造一份**该类对象确实为空**的输入，断言：

  1. 退出码 == `EMPTY_RC`（3），不是 0，也不是 2
     （3 与 2 必须分开：2 在本仓语义是「判定出了问题」，doc_dispatch.do_renum 收到 2
      会 break、typeset_pipeline 收到非 0 会回滚 —— 把空集混进 2 会让「这份文档没有图」
      被误读成「图号有问题」，实测两处都坏）
  2. stdout/stderr 里**不出现任何通过语**（✓ / 连续 / 通过 / OK）
"""
from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[3]
RENUM = ROOT / "scripts" / "document" / "renum.py"
sys.path.insert(0, str(ROOT / "lib"))

EMPTY_RC = 3
# 通过语的特征：出现任何一个都说明它在空集上装作做出了肯定判定
PASS_MARKS = re.compile(r"[✓✔]|每节连续|全部连续|检查通过|\bOK\b")


def _blank_docx(tmp_path: Path, name: str = "empty.docx") -> Path:
    """造一份**一个图/表题注都没有**的 docx。"""
    import docx_safe_save  # noqa: F401  存盘收口（本仓硬约束）
    from docx import Document

    d = Document()
    d.add_heading("第一章 概述", level=1)
    d.add_paragraph("这是一段没有任何图表题注的正文。")
    d.add_heading("第二章 方法", level=1)
    d.add_paragraph("同样没有题注。")
    out = tmp_path / name
    d.save(out)
    return out


def _run(*argv: str) -> subprocess.CompletedProcess:
    return subprocess.run([sys.executable, str(RENUM), *argv],
                          capture_output=True, text=True, timeout=180)


def _assert_empty_set(r: subprocess.CompletedProcess, what: str) -> None:
    blob = r.stdout + r.stderr
    assert r.returncode == EMPTY_RC, (
        f"{what}：空集上 rc={r.returncode}，期望 {EMPTY_RC}。\n"
        f"rc=0 = 假绿（空集报通过），rc=2 = 与「判定出问题」撞码，"
        f"会让 doc_dispatch break / typeset_pipeline 回滚。\n{blob[-600:]}")
    hit = PASS_MARKS.search(blob)
    assert not hit, f"{what}：空集上仍打了通过语 {hit.group(0)!r}\n{blob[-600:]}"


# ── figures：docx 侧图号重排 ────────────────────────────────────────────────

def test_figures_default_on_docx_without_captions(tmp_path):
    """`renum.py figures <无题注的 docx>` —— 本轮修掉的那条假绿的原案。"""
    _assert_empty_set(_run("figures", str(_blank_docx(tmp_path))), "figures 默认模式")


def test_figures_dry_run_on_docx_without_captions(tmp_path):
    _assert_empty_set(_run("figures", str(_blank_docx(tmp_path)), "--dry-run"),
                      "figures --dry-run")


def test_figures_wrong_kind_is_also_empty(tmp_path):
    """`--kind 表` 在只有图的文档上同样是空集 —— 不许因为「文档里有别的题注」就放行。"""
    _assert_empty_set(_run("figures", str(_blank_docx(tmp_path)), "--kind", "表"),
                      "figures --kind 表")


# ── tabfig：md 侧 表/图 题注号对齐 ──────────────────────────────────────────

def test_tabfig_on_dir_without_tokens(tmp_path):
    (tmp_path / "ch1-intro.md").write_text("# 第一章\n\n没有任何编号 token。\n",
                                           encoding="utf-8")
    _assert_empty_set(_run("tabfig", str(tmp_path)), "tabfig（目录里无编号 token）")


def test_tabfig_on_empty_dir(tmp_path):
    """扫描根存在但一个 ch<N>-*.md 都没有 —— 本仓铁律：拒绝在空集上报绿。"""
    _assert_empty_set(_run("tabfig", str(tmp_path)), "tabfig（目录里没有章节文件）")


# ── chapter：md 侧章号位移 ─────────────────────────────────────────────────

def test_chapter_with_empty_sequence(tmp_path):
    cfg = tmp_path / "chapters.yaml"
    cfg.write_text("number_base: 1\nsequence: []\n", encoding="utf-8")
    _assert_empty_set(_run("chapter", str(cfg)), "chapter（sequence 为空）")


# ── 反向：非空集必须照常工作，别把这道门做成「一律判 3」──────────────────

def test_nonempty_docx_does_not_hit_empty_exit(tmp_path):
    """没有这条，把 `_empty_set_exit` 挪到函数开头无条件调用也能让上面全绿。"""
    import docx_safe_save  # noqa: F401
    from docx import Document

    d = Document()
    d.add_heading("第一章 概述", level=1)
    d.add_paragraph("正文。")
    d.add_paragraph("图1-1 第一张图的题注")
    d.add_paragraph("图1-2 第二张图的题注")
    src = tmp_path / "has-fig.docx"
    d.save(src)

    # ⚠ `figures` 的**默认 kind 是英文 `Figure`**，中文题注要显式 `--cn-section --kind 图`。
    # 写这条测试时按「默认模式 + 中文题注」跑，拿到的是 rc=3「未发现任何 Figure 题注」——
    # 判据没错，是对照组构造错了。留这行注释免得下一个人把它当成空集判据的 bug。
    r = _run("figures", str(src), "--cn-section", "--kind", "图", "--dry-run")
    assert r.returncode != EMPTY_RC, (
        f"有 2 条图题注的文档被判成空集（rc={EMPTY_RC}）—— 空集判据把非空集也吃了\n"
        f"{(r.stdout + r.stderr)[-600:]}")


if __name__ == "__main__":
    raise SystemExit(pytest.main([__file__, "-q"]))
