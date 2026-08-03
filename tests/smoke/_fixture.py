#!/usr/bin/env python3
"""_fixture.py — smoke 用的 fixture 造件器（纯 stdlib + python-docx，不读仓外真实文件）。

为什么造而不是拿现成的：`templates/template.docx` 是指向 `~/Work` 真实交付件的
symlink（15MB 真业务数据）。smoke 一旦指向它，① 测试依赖一个不在版本控制里的文件
② 写模式动词有机会写进用户在跑的项目。所以 fixture 全部现造，且只落 `tmp_path`。

造出来的东西要足够"真"，否则大半动词是**空跑**——rc=0 但一件事没干，
而空跑的 rc=0 和真干活的 rc=0 长得一模一样。所以 fixture 里塞齐：

    多级标题(H1~H4) · 正文 · 3 张表 · 6 张图 · 交叉引用 · 两个分节 · 页眉页脚
    书签(_Toc_*) · 域(REF/TOC/PAGE) · 重复图片(供去重)
    直接格式脏段(供 clear-direct-format / scan-ppr) · 等线 run(供 fonts normalize)

题注这一层 2026-08-03 加强过一轮，**三件事同时成立才不空跑**（缺一件都会让一
整族动词扫到 0 个对象、rc 照样 0）：

① **真的题注样式**，不是 `Normal`。图题套内置 ``Caption``、表题套自造的 ``表题``。
   加强前 6 个题注段全是 `Normal` → `renumber h4-figures` 打 `图=0 表=0`、
   `caption` 族与 `audit captions` 的样式侧判据全部落空。
   ⚠ **图/表用两个不同样式名是必须的，不是审美**：`styles_registry.yaml` 的
   `Caption` 同时列在 FIG_CAPTION_STYLES 与 TABLE_CAPTION_STYLES 里，而
   `styles._build_h4fig_plan` 是 `elif` 链、先判图后判表 —— 表题若也套 `Caption`
   会被图分支先认领、再被 `startswith("图")` 否掉，`plan_tbl` 永远是 0。
   `表题` 在 TABLE_CAPTION_STYLES / `caption_re.CAPTION_STYLES_EXACT` /
   `is_table_caption_style` 三处都在，只不在 FIG 那边，正好躲开这条链。

② **五种短横全覆盖**（ASCII `-` / U+2011 / U+2013 / U+2014 / 全角 `－`）。
   `caption_re` 声明认五种，只用 ASCII 的 fixture 抓不住「某处判据退回只认 ASCII」
   这类回归 —— 2026-08-02 `sub/strip.py` 就真漏网过一次。
   分布见 `build_docx` 里每条题注后面的行内注释。

③ **两段式与三段式都要有**：`图1-1` 那种两段式撑不起 `SECTIONED_CAPTION`
   （`style caption` / `strip outlinelvl` 的写盘范围，有意只认 `图3.1-1` 三段式）。
   所以第 3 章带一对三段式题注，且**故意写了 `w:outlineLvl`** 给 `strip outlinelvl`
   当活干（真实病灶：合稿后题注段被写进 Word 导航大纲）。

图号重排两条线是**零共享**的（见 `caption_re.en_caption_pattern` 那段），所以
中文题注再多也喂不到 `renumber-fig` 的默认模式。附录 A 因此放了两条**乱序**的
英文题注（`Figure 3` 在 `Figure 1` 前面）+ 正文引用 —— 乱序是刻意的，顺序对的
话 remap 是恒等映射，写盘动词照样等于没干活。

⚠ **fixture 内容与若干动词的 expect_rc / mutates 是绑定的**（`audit-styleset *`
靠样式集判 severity、`para scan-ppr` 靠脏段数、`health diagnose` 靠病种命中、
一批写盘动词靠上面这几类对象存不存在）。改这个文件 = 可能改掉 `_verb_specs`
里那几条，改完必须重跑 smoke 并**重新实测**那两列，不能照抄邻居。
"""
from __future__ import annotations

import json
import re
import struct
import sys
import zipfile
import zlib
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
sys.path.append(str(REPO / "lib"))
import docx_safe_save  # noqa: E402,F401  （存盘收口，见 CLAUDE.md）

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.style import WD_STYLE_TYPE  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402

#: 图题样式名。落在 `styles_registry` 的 FIG_CAPTION_STYLES、
#: `caption_re.is_fig_caption_style`、`is_caption_family_style` 三处判据里。
#:
#: ⚠ **不用内置的 `Caption`**，虽然它也在 FIG_CAPTION_STYLES 里：
#: ① `caption_re.is_fig_caption_style("Caption")` 是 **False**（那条正则要的是
#:    `图名/图注/image caption/^figure$`），于是 `shape_contract` 数出来的
#:    `caption_figure_count` 恒为 0 —— 一整个快照字段空跑。
#: ② `styles` 的 FIGURE_STYLE_PRIORITY 本身就是 `["ZDWP 图名", "Caption"]`，
#:    题注要是已经叫 `Caption`，`style caption` 每条都落 `no_change_skip`，
#:    apply 分支一次都走不到。
#: `Image Caption` 三处判据全中，且不在 TABLE_CAPTION_STYLES 里（见 docstring ①
#: 那条 elif 链），是同时满足这几条的写法。
FIG_CAPTION_STYLE = "Image Caption"
#: 表题样式名。**有意与图题不同**，理由见模块 docstring ①。
#: `表题` 同时在 TABLE_CAPTION_STYLES / `CAPTION_STYLES_EXACT` /
#: `is_table_caption_style` 里，且不在 FIG_CAPTION_STYLES 里。
TBL_CAPTION_STYLE = "表题"


# ─────────────────────────── 零件 ───────────────────────────

def make_png(path: Path, w: int = 160, h: int = 120, rgb=(66, 133, 244)) -> Path:
    """纯 stdlib 造 PNG（不引 PIL —— smoke 不该因为缺一个绘图库就跑不起来）。"""
    def chunk(typ: bytes, data: bytes) -> bytes:
        c = typ + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))
    png = (b"\x89PNG\r\n\x1a\n"
           + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
           + chunk(b"IDAT", zlib.compress(raw, 9))
           + chunk(b"IEND", b""))
    path.write_bytes(png)
    return path


def _add_bookmark(paragraph, name: str, bid: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_field(paragraph, instr: str = "PAGE") -> None:
    """塞一个完整的 w:fldChar begin/instrText/separate/结果/end 域。"""
    r1 = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r1.append(fc1)

    r2 = OxmlElement("w:r")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {instr} "
    r2.append(it)

    r3 = OxmlElement("w:r")
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "separate")
    r3.append(fc3)

    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r4.append(t)

    r5 = OxmlElement("w:r")
    fc5 = OxmlElement("w:fldChar")
    fc5.set(qn("w:fldCharType"), "end")
    r5.append(fc5)

    for r in (r1, r2, r3, r4, r5):
        paragraph._p.append(r)


def _set_outline_lvl(paragraph, lvl: int) -> None:
    """给段落写 ``w:pPr/w:outlineLvl``（段级直接格式）。

    这是**故意造的病灶**：合稿后题注段被写进 Word 导航大纲，`strip outlinelvl`
    正是来摘它的。不造它的话那条动词扫到 0 个对象、rc 照样 0。
    """
    pPr = paragraph._p.get_or_add_pPr()
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(lvl))
    pPr.append(ol)


def _ensure_caption_styles(doc) -> None:
    """确保两个题注段落样式都存在（内置模板都没有，得自己建）。

    内置只有 ``Caption`` 一个，而它同时属图题族与表题族 —— 正是要避开的那种。
    """
    for name in (FIG_CAPTION_STYLE, TBL_CAPTION_STYLE):
        try:
            doc.styles[name]
        except KeyError:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style = doc.styles["Normal"]


def _caption(doc, text: str, kind: str, *, outline_lvl: int | None = None):
    """加一个**真题注段**：套题注样式 + 居中（+ 可选 outlineLvl 病灶）。

    `kind` ∈ {"图", "表"}，决定套哪个样式 —— 两者不能共用一个名字，见模块 docstring ①。
    """
    p = doc.add_paragraph(text)
    p.style = doc.styles[FIG_CAPTION_STYLE if kind == "图"
                         else TBL_CAPTION_STYLE]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if outline_lvl is not None:
        _set_outline_lvl(p, outline_lvl)
    return p


# ─────────────────────────── 主 fixture docx ───────────────────────────

def build_docx(dst: Path, img_dir: Path | None = None) -> Path:
    """造主 fixture docx。内容固定 —— 固定是特性，rc 才可断言。"""
    img_dir = img_dir or dst.parent
    img_dir.mkdir(parents=True, exist_ok=True)
    img1 = make_png(img_dir / "fig1.png", 160, 120, (66, 133, 244))
    img2 = make_png(img_dir / "fig2.png", 160, 120, (219, 68, 55))
    img3 = img_dir / "fig3.png"
    img3.write_bytes(img1.read_bytes())      # 逐字节重复 → 供 seqdiff image 去重
    img4 = make_png(img_dir / "fig4.png", 160, 120, (15, 157, 88))
    img5 = make_png(img_dir / "fig5.png", 160, 120, (244, 180, 0))
    img6 = make_png(img_dir / "fig6.png", 160, 120, (123, 31, 162))

    doc = Document()
    _ensure_caption_styles(doc)

    p = doc.add_paragraph("某某水利工程可行性研究报告")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h1 = doc.add_heading("第1章 概述", level=1)
    _add_bookmark(h1, "_Toc_chapter1", 1)
    doc.add_paragraph(
        "本报告依据《水利水电工程可行性研究报告编制规程》(SL/T 618-2021)编制。"
        "工程位于某省某市，总投资约 12000 万元，工期 24 个月。参见\"设计标准\"一节。")
    doc.add_heading("1.1 工程概况", level=2)
    doc.add_paragraph(
        "工程规模为中型，设计洪水标准 50 年一遇，校核 500 年一遇。渠道全长 12.5km，"
        "设计流量 8.5m3/s。温度控制在 25 摄氏度以内，间距 100 毫米。")
    doc.add_heading("1.1.1 地理位置", level=3)
    doc.add_paragraph("工程区地处东经 118°21′，北纬 32°05′，属亚热带季风气候区。")

    doc.add_picture(str(img1), width=Cm(8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _caption(doc, "图1-1 工程地理位置示意图", "图")        # 短横①: ASCII HYPHEN-MINUS

    doc.add_paragraph("如图1-1所示，工程区交通便利。")
    fp_ref = doc.add_paragraph("详见第 ")
    _add_field(fp_ref, r"REF _Toc_chapter2 \h")
    fp_ref.add_run(" 章。")
    fp_toc = doc.add_paragraph()
    _add_field(fp_toc, r"TOC \o \"1-3\" \h \z \u")

    _caption(doc, "表1-1 工程特性表", "表")               # 短横①: ASCII HYPHEN-MINUS
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Table Grid"
    for i, row in enumerate([["序号", "项目", "数值"],
                             ["1", "设计流量", "8.5 m3/s"],
                             ["2", "渠道长度", "12.5 km"],
                             ["3", "总投资", "12000 万元"]]):
        for j, v in enumerate(row):
            tbl.cell(i, j).text = v
    doc.add_paragraph("表1-1 列出了主要工程特性指标。")

    doc.add_section(WD_SECTION.NEW_PAGE)

    h2 = doc.add_heading("第2章 水文分析", level=1)
    _add_bookmark(h2, "_Toc_chapter2", 2)
    doc.add_paragraph("采用 1980-2020 年共 41 年实测资料进行频率分析。数据来源于某某水文站。")
    doc.add_heading("2.1 设计标准", level=2)
    doc.add_paragraph("按照规范要求，本工程防洪标准取 50 年一遇。")
    doc.add_heading("2.1.1 径流计算", level=3)
    doc.add_paragraph("多年平均径流量 3.2 亿 m3，径流系数 0.42。")
    doc.add_heading("2.1.1.1 参数率定", level=4)
    doc.add_paragraph("率定期 NSE=0.85，验证期 NSE=0.81。")

    doc.add_picture(str(img2), width=Cm(8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _caption(doc, "图2‑1 年径流过程线", "图")          # 短横②: U+2011 非断行连字符

    _caption(doc, "表2–1 设计洪水成果表", "表")        # 短横③: U+2013 EN DASH
    tbl2 = doc.add_table(rows=3, cols=4)
    tbl2.style = "Table Grid"
    for i, row in enumerate([["频率", "洪峰(m3/s)", "洪量(万m3)", "备注"],
                             ["2%", "520", "3200", "设计"],
                             ["0.2%", "780", "4800", "校核"]]):
        for j, v in enumerate(row):
            tbl2.cell(i, j).text = v

    doc.add_picture(str(img3), width=Cm(8))
    _caption(doc, "图2—2 工程地理位置示意图（重复）", "图")  # 短横④: U+2014 EM DASH

    doc.add_heading("第3章 结论与建议", level=1)
    doc.add_paragraph("本工程技术可行、经济合理，建议尽快开展初步设计。"
                      "联系人：张三，电话 13800138000。")
    doc.add_paragraph("“引号测试”与'单引号'，以及 100mm、25℃ 等单位。")

    # 直接格式脏段：供 fix clear-direct-format / para scan-ppr / fonts normalize
    dirty = doc.add_paragraph()
    r = dirty.add_run("这一段带直接格式：加粗、变色、等线。")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r.font.name = "等线"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")

    # ── 段级直接格式脏段：供 para scan-ppr / para fix-ppr ────────────────────
    # ⚠ 上面那个 `dirty` 是 **run 级** 脏（bold/颜色/字体），`scan-ppr` 看不见它 ——
    # 它比的是正文流段的 `w:pPr` 里 jc / ind 签名与本样式**模态**的偏离。
    # 2026-08-03 之前这条动词的 4 条 suspects 全是「题注段没套题注样式、只手动
    # 居中」——也就是说它抓的是 fixture 自己的缺陷；题注样式一补，suspects 归零、
    # rc 从 3 掉到 0。所以这里显式补两段**真正的正文流脏格式**，让它继续有对象。
    ppr_center = doc.add_paragraph("本段是正文却被手动居中，与本样式的模态对齐方式不一致。")
    ppr_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ppr_ind = doc.add_paragraph("本段被手动加了首行缩进与左缩进，缩进签名偏离模态。")
    ppr_ind.paragraph_format.first_line_indent = Pt(24)
    ppr_ind.paragraph_format.left_indent = Pt(12)

    # ── 三段式题注（图3.1-1 / 表3.1－1）+ outlineLvl 病灶 ───────────────────
    # 两段式的 `图1-1` 撑不起 SECTIONED_CAPTION（`style caption` 与
    # `strip outlinelvl` 的写盘范围，有意只认三段）。没有这一节，那两条动词
    # 一个对象都扫不到，rc 却照样 0。
    doc.add_heading("3.1 推荐方案工程布置", level=2)
    doc.add_picture(str(img4), width=Cm(8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _caption(doc, "图3.1-1 灌区渠系布置图", "图", outline_lvl=6)   # 三段式 + 短横①
    doc.add_paragraph("渠系布置见图3.1-1，投资估算见表3.1－1。")

    _caption(doc, "表3.1－1 投资估算表", "表", outline_lvl=6)      # 三段式 + 短横⑤ U+FF0D
    tbl3 = doc.add_table(rows=3, cols=2)
    tbl3.style = "Table Grid"
    for i, row in enumerate([["项目", "投资(万元)"],
                             ["建筑工程", "7200"],
                             ["设备购置", "2100"]]):
        for j, v in enumerate(row):
            tbl3.cell(i, j).text = v

    # ── 附录 A：英文题注线（renumber-fig 默认模式唯一认得的形态）─────────────
    # 中英两条编号线在 `caption_re` 里是零共享的，所以中文题注再多也喂不到
    # `renumber-fig` 的默认模式。**号是乱的（3 在 1 前面）也是刻意的**：顺序对了
    # remap 就是恒等映射，写盘动词跑完等于没干活，mutates 断言什么也证明不了。
    doc.add_heading("附录A 英文摘要（English Summary）", level=1)
    doc.add_paragraph(
        "This appendix summarises the canal layout and the annual runoff process.")
    doc.add_picture(str(img5), width=Cm(8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _caption(doc, "Figure 3 Layout of the irrigation canal system", "图")
    doc.add_picture(str(img6), width=Cm(8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _caption(doc, "Figure 1 Annual runoff hydrograph", "图")
    doc.add_paragraph(
        "The canal layout is shown in Figure 3, and Figure 1 gives the annual "
        "runoff process at the gauging station.")

    # ── 附录 B：**没编号**的题注（`caption number` 的活）────────────────────
    # 上面那些题注全都已经有号了，`caption.py` 见到号就 `continue`，所以它们喂不到
    # 这条动词。它要的是「一个它认得的章上下文 + 一段没号的题注」，两个条件缺一
    # 就落 manual_review / 静默跳过、rc 照样 0。
    # ⚠ 标题写成 `四、…` 而不是 `第4章 …` 不是笔误：`caption.is_h1_chapter` 只认
    # `RE_CN_CHAPTER = ^([一二三四五六七八九十]+)、`，`第N章` 这一形态它压根不认
    # （`parse_chapter` 里的阿拉伯分支 is_h1_chapter 也没调）。整篇都用 `第N章`
    # 的话 `chapter` 恒为 0，所有候选一律记成 `no-chapter-context`。
    doc.add_heading("四、补充图件与指标", level=1)
    doc.add_picture(str(img4), width=Cm(8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _caption(doc, "灌区水系分布图", "图")            # 无编号 → 等着被补 `图 4-1　`

    _caption(doc, "工程效益指标一览", "表")          # 无编号 → 等着被补 `表 4-1　`
    tbl4 = doc.add_table(rows=2, cols=2)
    tbl4.style = "Table Grid"
    for i, row in enumerate([["指标", "数值"], ["灌溉面积(万亩)", "12.4"]]):
        for j, v in enumerate(row):
            tbl4.cell(i, j).text = v

    for sec in doc.sections:
        sec.header.is_linked_to_previous = False
        sec.footer.is_linked_to_previous = False
        hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
        hp.text = "某某水利工程可行性研究报告"
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_field(fp, "PAGE")

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    return dst


def build_docx_v2(dst: Path, img_dir: Path | None = None) -> Path:
    """主 fixture 的"下一版"：改几句话 + 少一张图，供 compare / seqdiff / track 用。

    完全相同的两份文件会让 diff 类动词跑出「无差异」——那也是 rc=0，
    但它证明不了 diff 真的在比。所以这里必须真的不一样。
    """
    build_docx(dst, img_dir=img_dir)
    doc = Document(str(dst))
    for p in doc.paragraphs:
        if p.text.startswith("按照规范要求"):
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = "按照最新规范要求，本工程防洪标准调整为 100 年一遇。"
        elif p.text.startswith("率定期"):
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = "率定期 NSE=0.88，验证期 NSE=0.83（重新率定）。"
    doc.add_paragraph("本段是 v2 新增的补充说明，用于 sequence diff 验证新增段被识别。")
    doc.save(str(dst))
    return dst


def build_template_docx(dst: Path) -> Path:
    """一份"模板" docx：只有样式与页眉页脚，没有正文。

    `template` / `format apply` / `chrome` 这类动词默认会去读
    `templates/template.docx`（→ ~/Work 的 symlink）。smoke 必须显式传这一份把
    那条依赖断掉，否则测试既碰真实业务文件、又在别的机器上必然失败。
    """
    doc = Document()
    doc.add_heading("模板标题样式", level=1)
    doc.add_heading("模板二级样式", level=2)
    doc.add_paragraph("模板正文样式示例。")
    sec = doc.sections[0]
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    sec.header.paragraphs[0].text = "模板页眉"
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_field(fp, "PAGE")
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    return dst


# ─────────────────────────── 配套的非 docx 造件 ───────────────────────────

FIXTURE_MD = """\
# 某某水利工程可行性研究报告

## 第1章 概述

本报告依据《水利水电工程可行性研究报告编制规程》编制。工程位于某省某市。

![图1-1 工程地理位置示意图](fig1.png)

| 序号 | 项目 | 数值 |
|---|---|---|
| 1 | 设计流量 | 8.5 m3/s |
| 2 | 渠道长度 | 12.5 km |

## 第2章 水文分析

采用 1980-2020 年共 41 年实测资料进行频率分析。

### 2.1 设计标准

按照规范要求，本工程防洪标准取 50 年一遇。
"""

FIXTURE_MD2 = """\
# 附录

## 第3章 结论与建议

本工程技术可行、经济合理，建议尽快开展初步设计。
"""

# 这一份同时喂 `compare-ref ref` 与 `revise-rules gen`，**两者认的形状不一样**，
# 下面每一处写法都是被其中一条正则逼出来的（`sub/biddiff.py`）：
#
#   · 章节头必须是 `## 改动 <数字>` —— `SECTION_RE = ^## 改动\\s*\\d+`。
#     写成 `## 意见 1` → `parse_md` 里 `starts` 为空 → **直接 return []**，
#     于是 `revise-rules gen` 打「0 条」、写出一个 `[]`、**rc 照样 0**。
#     2026-08-03 之前这份草稿就是这么写的，那条动词从立表起 100% 空跑。
#   · `**原文**` 与 `**改为**` **必须成对**：`if not originals or not revises: continue`。
#     只有「改为」没有「原文」 → 整节被静默跳过（同样 rc=0）。
#   · 引用块必须**紧贴**在 `**原文**…` 那一行的下一行，中间**不能有空行** ——
#     `PAIR_BLOCK_RE` 是 `\\*\\*…\\*\\*[^\\n]*\\n((?:>\\s?.*\\n)+)`，`[^\\n]*\\n` 只吃掉
#     本行剩余部分。（`compare-ref` 那条 `extract_revisions` 的 `\\s*\\n` 反而容得下
#     空行，所以「不留空行」是同时满足两者的唯一写法。）
#   · `>` 后面要有一个空格：`extract_revisions` 用的是 `>\\s`（不是 `>\\s?`）。
#   · `find` 取**文档里逐字存在**的整句，让 gen 的 Pass 1 存在性校验走真分支。
DRAFT_MD = """\
# 第1章改动草稿

## 改动 1 防洪标准

**原文**：
> 按照规范要求，本工程防洪标准取 50 年一遇。

**改为**：
> 按照最新规范要求，本工程防洪标准调整为 100 年一遇，并补充校核工况说明。

## 改动 2 径流成果

**原文**：
> 多年平均径流量 3.2 亿 m3，径流系数 0.42。

**改为**：
> 多年平均径流量 3.2 亿 m3，径流系数 0.42，另需补充蒸发折算系数。
"""

# `scan-sensitive` 的输入是**目录**，且目录里必须有 .md
BID_MD = """\
# 投标承诺

我方为该产品的唯一指定供应商，绝对保证工期最短，产品为国家级免检产品。
"""


def build_side_files(root: Path) -> dict:
    """造齐非 docx 的配套件，返回路径字典。"""
    root.mkdir(parents=True, exist_ok=True)

    md_dir = root / "md"
    md_dir.mkdir(exist_ok=True)
    (md_dir / "01-第1章.md").write_text(FIXTURE_MD, encoding="utf-8")
    (md_dir / "02-第3章.md").write_text(FIXTURE_MD2, encoding="utf-8")
    make_png(md_dir / "fig1.png")

    drafts = root / "drafts"
    drafts.mkdir(exist_ok=True)
    # ⚠ 文件名里那个短横不是随手加的：`compare-ref ref` 的默认 glob 是
    # `*改动草稿.md`，`revise-rules gen` 的默认 glob 是 `*-改动草稿.md`
    # （biddiff.py:399 vs :604 —— 两条默认值本身就不一致）。
    # `第1章-改动草稿.md` 是同时满足两者的唯一写法；去掉短横，revise-rules
    # 就变成「待处理 MD: 0 份」的空跑，而它照样 rc=0，测试看不出来。
    (drafts / "第1章-改动草稿.md").write_text(DRAFT_MD, encoding="utf-8")

    bid = root / "bid"
    bid.mkdir(exist_ok=True)
    (bid / "bid.md").write_text(BID_MD, encoding="utf-8")

    out_dir = root / "out"
    out_dir.mkdir(exist_ok=True)

    # `caption pair` 的 --decision 是 required，且过 schema 校验（v1）
    decision = root / "decision.json"
    decision.write_text(json.dumps({
        "version": "1",
        "source_docx": "fixture.docx",
        "ops": [{"op": "renumber-all-tables"}],
        "operations": [{"op": "renumber-all-tables"}],
    }, ensure_ascii=False, indent=1), encoding="utf-8")

    # `blocks relocate` 的 --plan 是 required；source_heading_text 必须在文档里找得到，
    # 否则该 move 被记成 skip（仍 rc=0，所以这里得用真实存在的标题文本）
    plan = root / "relocate-plan.json"
    plan.write_text(json.dumps({
        "version": "1",
        "source_docx": "fixture.docx",
        "moves": [{
            "source_idx": 19,
            "target_idx": 5,
            "source_heading_idx": 19,
            "source_block_end_idx": 19,
            "target_insert_after_idx": 5,
            "source_heading_text": "2.1.1 径流计算",
            "target_context_text": "1.1.1 地理位置",
            "reason": "smoke: 把 2.1.1 挪到 1.1.1 后面",
        }],
    }, ensure_ascii=False, indent=1), encoding="utf-8")

    return {
        "md_dir": md_dir,
        "md": md_dir / "01-第1章.md",
        "md2": md_dir / "02-第3章.md",
        "drafts": drafts,
        "bid_dir": bid,
        "out_dir": out_dir,
        "decision": decision,
        "relocate_plan": plan,
    }


# ─────────────────────────── fixture 自身的富度门 ───────────────────────────

#: 每一项 = (名字, 实测值取法, 下界)。**下界不是「当前值」而是「低于此即失去断言力」**，
#: 所以留了余量：改 fixture 加东西不会碰它，砍掉一整类对象才会。
def selfcheck(docx: Path, drafts_dir: Path) -> list[str]:
    """核对 fixture 里各类对象**真的存在**，返回失败项清单（空 = 通过）。

    为什么要这一层：fixture 的价值全在「有东西可跑」，而**空跑的 rc=0 和真干活的
    rc=0 长得一模一样** —— 谁不小心把题注样式改回 `Normal`、把短横统一成 ASCII、
    或者把附录删了，smoke 会**继续全绿**，只是不再证明任何事。这道门把「fixture
    有多富」从注释变成机器判据。

    fail-closed：任何一类计数落到 0 / 低于下界都判失败，不打招呼、不「跳过」。
    """
    from caption_re import (parse, SECTIONED_CAPTION, en_caption_pattern)

    doc = Document(str(docx))
    paras = doc.paragraphs
    styles = [(p.style.name if p.style else "", p.text.strip()) for p in paras]

    fig_caps = [t for s, t in styles if s == FIG_CAPTION_STYLE and t]
    tbl_caps = [t for s, t in styles if s == TBL_CAPTION_STYLE and t]
    cn_fig_caps = [t for t in fig_caps if t.startswith("图")]
    sectioned = [t for _s, t in styles if t and parse(t, SECTIONED_CAPTION)]

    en_re = en_caption_pattern("Figure")
    en_nums = [int(m.group(1)) for t in fig_caps
               if (m := en_re.match(t)) is not None]

    with zipfile.ZipFile(docx) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]

    outlined = 0
    for p in paras:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:outlineLvl")) is not None:
            if p.text.strip().startswith(("图", "表")):
                outlined += 1

    all_cap_text = "".join(fig_caps + tbl_caps)
    dashes = {name: ch for name, ch in
              (("ASCII", "-"), ("U+2011", "‑"), ("U+2013", "–"),
               ("U+2014", "—"), ("U+FF0D", "－"))
              if ch in all_cap_text}

    draft_files = sorted(drafts_dir.glob("*-改动草稿.md"))
    draft_text = "".join(f.read_text(encoding="utf-8") for f in draft_files)
    n_sections = len(re.findall(r"^## 改动\s*\d+", draft_text, re.M))
    n_pairs = min(draft_text.count("**原文"), draft_text.count("**改为"))

    checks = [
        ("图题段（真题注样式）",            len(cn_fig_caps),        4),
        ("表题段（真题注样式）",            len(tbl_caps),           4),
        ("三段式题注（SECTIONED_CAPTION）", len(sectioned),          2),
        ("带 outlineLvl 的题注（strip 的活）", outlined,             2),
        ("题注里的短横种类",                len(dashes),             3),
        ("内嵌图对象（w:drawing）",         len(doc.inline_shapes),  4),
        ("媒体部件（word/media/*）",        len(media),              3),
        ("英文图题（renumber-fig 默认线）",  len(en_nums),            2),
        ("改动草稿的 `## 改动 N` 节",       n_sections,              2),
        ("草稿里 原文/改为 成对块",          n_pairs,                 2),
    ]
    bad = [f"{name}: 实测 {got} < 下界 {lo}"
           for name, got, lo in checks if got < lo]

    # 英文图号必须是**乱的** —— 顺序对了 remap 是恒等映射，写盘动词跑完等于没干活。
    if en_nums and en_nums == sorted(en_nums):
        bad.append(f"英文图号 {en_nums} 是升序 —— renumber-fig 的 remap 会退化成恒等映射")

    if not bad:
        print("fixture 富度自检：")
        for name, got, lo in checks:
            print(f"  ✓ {name:32} = {got}  (≥{lo})")
        print(f"  ✓ 短横覆盖: {sorted(dashes)}")
        print(f"  ✓ 英文图号顺序: {en_nums}（乱序，remap 非恒等）")
    return bad


if __name__ == "__main__":  # 手工造一份看看
    target = Path(sys.argv[1] if len(sys.argv) > 1 else "/tmp/doctools-smoke/probe")
    target.mkdir(parents=True, exist_ok=True)
    main_docx = build_docx(target / "fixture.docx")
    build_docx_v2(target / "fixture-v2.docx")
    build_template_docx(target / "template.docx")
    side = build_side_files(target)
    failures = selfcheck(main_docx, side["drafts"])
    if failures:
        print("\n⛔ fixture 富度自检不通过（空跑的 rc=0 与真干活的 rc=0 长得一样，"
              "所以这里不许放行）：", file=sys.stderr)
        for f in failures:
            print(f"   ✗ {f}", file=sys.stderr)
        sys.exit(1)
    print("OK ->", target)
